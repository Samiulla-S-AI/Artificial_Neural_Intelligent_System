import tkinter as tk
from tkinter import ttk, scrolledtext, Label, Frame, Button,Entry, Text, END, NORMAL, DISABLED, WORD,Canvas, Tk, filedialog
import customtkinter as ctk
import pyttsx3
import datetime
import speech_recognition as sr
from PIL import Image, ImageGrab, ImageDraw, ImageFont
import webbrowser
import numpy as np
import requests
import io
import base64
from openai import OpenAI
from urllib.parse import urlparse
import os
import pyperclip
import pandas as pd
import time
import math
import threading
import logging
from urllib3.util.retry import Retry
from requests.adapters import HTTPAdapter
import pygame
import screen_brightness_control as sbc
import keyboard
import pyautogui
from PIL import Image, ImageGrab
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import smtplib
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import psutil
from tkinter import messagebox
from deep_translator import GoogleTranslator
import urllib.parse
import sys
import psutil
import random
import platform
import subprocess
import importlib
from pathlib import Path
import sqlite3
import google.generativeai as genai
from urllib.parse import quote
import shutil
import re
import pygetwindow as gw
import win32clipboard
from screen_analyzer import ConversationLogManager, save_coordinates
import json
log_manager = ConversationLogManager()
DATABASE_FILE ="conversation_log.db"
def get_app_directory():
    try:
        home = str(Path.home())
        app_dir = os.path.join(home, "VoiceAssistant")
        return app_dir
    except Exception as e:
        print(f"Error getting app directory: {str(e)}")
        return os.path.join(os.path.dirname(os.path.abspath(__file__)), "VoiceAssistant")
def ensure_directory_exists(directory):
    try:
        os.makedirs(directory, exist_ok=True)
        return True
    except Exception as e:
        print(f"Error creating directory {directory}: {str(e)}")
        return False
def get_log_directory():
    try:
        app_dir = get_app_directory()
        log_dir = os.path.join(app_dir, "logs")
        return log_dir
    except Exception as e:
        print(f"Error getting log directory: {str(e)}")
        return os.path.join(os.path.dirname(os.path.abspath(__file__)), "logs")
def setup_logging():
    try:
        log_dir = get_log_directory()
        if not ensure_directory_exists(log_dir):
            print("Warning: Failed to create log directory. Using basic console logging.")
            logging.basicConfig(
                level=logging.INFO,
                format='%(asctime)s - %(levelname)s - %(message)s',
                handlers=[logging.StreamHandler()]
            )
            return False
        log_file = os.path.join(log_dir, "voice_assistant.log")
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler(log_file),
                logging.StreamHandler()
            ]
        )
        logging.info("Logging initialized successfully")
        return True
    except Exception as e:
        print(f"Warning: Error in logging setup: {str(e)}. Using basic console logging.")
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[logging.StreamHandler()]
        )
        return False
def load_setting(setting_name):
    """Load setting from database"""
    try:
        conn = sqlite3.connect("settings.db")
        cursor = conn.cursor()
        cursor.execute('''CREATE TABLE IF NOT EXISTS settings
                        (name TEXT PRIMARY KEY, value TEXT)''')
        cursor.execute('SELECT value FROM settings WHERE name = ?', (setting_name,))
        result = cursor.fetchone()
        conn.close()
        return result[0] if result else None
    except Exception as e:
        logging.error(f"Error loading setting: {str(e)}")
        return None
try:
    gemini_model = genai.GenerativeModel(model_name=load_setting('gemini_model')) or genai.GenerativeModel("....")
except Exception as e:
    logging.error(f'Error Intializing Gemini MOdel:{str(e)}')
    gemini_model=None
def create_required_directories():
    try:
        app_dir = get_app_directory()
        required_dirs = {
            'app': app_dir,
            'logs': os.path.join(app_dir, "logs"),
            'data': os.path.join(app_dir, "data"),
            'config': os.path.join(app_dir, "configure"),
            'temp': os.path.join(app_dir, "temp")
        }
        success = True
        for dir_name, dir_path in required_dirs.items():
            if not ensure_directory_exists(dir_path):
                print(f"Warning: Failed to create {dir_name} directory at {dir_path}")
                success = False
        return success
    except Exception as e:
        print(f"Error creating required directories: {str(e)}")
        return False
ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("blue")
engine = pyttsx3.init('sapi5')
voices = engine.getProperty('voices')
engine.setProperty('voice', voices[0].id)
def play_sound():
    try:
        pygame.mixer.init()
        pygame.mixer.music.load(r'C:\Users\HP\Desktop\Ai\start_sounds\notification sound 2_15-03-25_15-27-59-839.mp3')
        pygame.mixer.music.play()
        while pygame.mixer.music.get_busy():
            pygame.time.Clock().tick(10)
    except Exception as e:
        print(f"Failed to play sound: {e}")
def play_sound2():
    try:
        pygame.mixer.init()
        pygame.mixer.music.load(r'C:\Users\HP\Desktop\Ai\start_sounds\mixkit-sci-fi-click-900.mp3')
        pygame.mixer.music.play()
        while pygame.mixer.music.get_busy():
            pygame.time.Clock().tick(10)
    except Exception as e:
        print(f"Failed to play sound: {e}")
def play_sound3():
    try:
        pygame.mixer.init()
        pygame.mixer.music.load(r'C:\Users\HP\Desktop\Ai\start_sounds\mixkit-fast-sci-fi-bleep-903.mp3')
        pygame.mixer.music.play()
        while pygame.mixer.music.get_busy():
            pygame.time.Clock().tick(10)
    except Exception as e:
        print(f"Failed to play sound: {e}")
def play_sound4():
    try:
        pygame.mixer.init()
        pygame.mixer.music.load(r'C:\Users\HP\Desktop\Ai\start_sounds\mixkit-fast-sci-fi-bleep-903.mp3')
        pygame.mixer.music.play()
        while pygame.mixer.music.get_busy():
            pygame.time.Clock().tick(10)
    except Exception as e:
        print(f"Failed to play sound: {e}")
class WaveVisualizer(tk.Canvas):
    def _convert_tk_properties(self, **kwargs):
        """Convert customtkinter properties to tkinter equivalents"""
        property_map = {
            'fg_color': 'bg',
            'border_color': 'highlightbackground',
            'border_width': 'highlightthickness'
        }
        converted_kwargs = {}
        for key, value in kwargs.items():
            if key in property_map:
                converted_kwargs[property_map[key]] = value
            else:
                converted_kwargs[key] = value
        return converted_kwargs
    def __init__(self, parent, **kwargs):
        kwargs = self._convert_tk_properties(**kwargs)
        super().__init__(parent, **kwargs)
        self.configure(bg='#000000')  # Black background
        self.is_active = True
        self.animation_running = False
        self.after_id = None
        self._lock = threading.Lock()
        self.particles = []
        self.colors = ['#FF0000', '#FF7F00', '#FFFF00', '#00FF00', '#0000FF', '#4B0082', '#8B00FF']  # Rainbow colors
        self.init_particles()
    def init_particles(self):
        """Initialize particles for the visualization"""
        self.particles = []
        num_particles = 50
        for i in range(num_particles):
            angle = (2 * math.pi * i) / num_particles
            self.particles.append({
                'angle': angle,
                'radius': 20,
                'speed': 0.1 + random.random() * 0.2,
                'color': random.choice(self.colors),
                'size': random.randint(2, 5)
            })
    def speaking_animation(self):
        if not self.is_active or self.animation_running:
            return
        def _animate():
            if not self.is_active or not self.animation_running:
                return False
            try:
                with self._lock:
                    self.delete("wave")
                    width = self.winfo_width()
                    height = self.winfo_height()
                    center_x = width / 2
                    center_y = height / 2
                    current_time = time.time()
                    for particle in self.particles:
                        particle['radius'] += math.sin(current_time * 2) * 0.5
                        particle['angle'] += particle['speed']
                        x = center_x + particle['radius'] * math.cos(particle['angle']) * 5
                        y = center_y + particle['radius'] * math.sin(particle['angle']) * 5
                        size = particle['size'] * (1 + math.sin(current_time * 4) * 0.3)
                        glow_size = size * 2
                        self.create_oval(
                            x - glow_size, y - glow_size,
                            x + glow_size, y + glow_size,
                            fill='', outline=particle['color'],
                            width=1, tags="wave",
                            stipple='gray50'
                        )
                        self.create_oval(
                            x - size, y - size,
                            x + size, y + size,
                            fill=particle['color'],
                            outline='',
                            tags="wave"
                        )
                        if random.random() < 0.3:
                            next_particle = random.choice(self.particles)
                            next_x = center_x + next_particle['radius'] * math.cos(next_particle['angle']) * 5
                            next_y = center_y + next_particle['radius'] * math.sin(next_particle['angle']) * 5
                            self.create_line(
                                x, y, next_x, next_y,
                                fill=particle['color'],
                                width=1,
                                tags="wave",
                                stipple='gray50'
                            )
                    return True
            except tk.TclError:
                self.is_active = False
                return False
            except Exception as e:
                logging.error(f"Animation error: {str(e)}")
                return False
        def animate_loop():
            if _animate():
                self.after_id = self.after(30, animate_loop)
            else:
                self.animation_running = False
                self.after_id = None
        with self._lock:
            if not self.animation_running and self.is_active:
                self.animation_running = True
                self.after_id = self.after(30, animate_loop)
    def start_listening_animation(self):
        if not self.is_active or self.animation_running:
            return
        def _animate():
            if not self.is_active or not self.animation_running:
                return False
            try:
                with self._lock:
                    self.delete("wave")
                    width = self.winfo_width()
                    height = self.winfo_height()
                    center_x = width / 2
                    center_y = height / 2
                    current_time = time.time()
                    max_radius = min(width, height) / 3
                    num_circles = 8
                    for i in range(num_circles):
                        radius = (max_radius * (i + 1) / num_circles) + (math.sin(current_time * 2) * 10)
                        color = self.colors[i % len(self.colors)]
                        self.create_oval(
                            center_x - radius, center_y - radius,
                            center_x + radius, center_y + radius,
                            outline=color,
                            width=2,
                            tags="wave"
                        )
                        for j in range(8):
                            angle = (2 * math.pi * j / 8) + current_time * 2
                            dot_x = center_x + radius * math.cos(angle)
                            dot_y = center_y + radius * math.sin(angle)
                            dot_size = 3 + math.sin(current_time * 4 + j) * 2
                            self.create_oval(
                                dot_x - dot_size, dot_y - dot_size,
                                dot_x + dot_size, dot_y + dot_size,
                                fill=color,
                                outline='',
                                tags="wave"
                            )
                    return True
            except tk.TclError:
                self.is_active = False
                return False
            except Exception as e:
                logging.error(f"Animation error: {str(e)}")
                return False
        def animate_loop():
            if _animate():
                self.after_id = self.after(30, animate_loop)
            else:
                self.animation_running = False
                self.after_id = None
        with self._lock:
            if not self.animation_running and self.is_active:
                self.animation_running = True
                self.after_id = self.after(30, animate_loop)
    def stop_listening_animation(self):
        with self._lock:
            self.animation_running = False
            if self.after_id:
                self.after_cancel(self.after_id)
                self.after_id = None
            if self.is_active:
                try:
                    self.delete("wave")
                except tk.TclError:
                    self.is_active = False
    def cleanup(self):
        """Clean up the visualizer before window destruction"""
        with self._lock:
            self.is_active = False
            self.animation_running = False
            if self.after_id:
                self.after_cancel(self.after_id)
                self.after_id = None
            try:
                self.delete("all")
            except tk.TclError:
                pass
class LiveAssistantWindow(ctk.CTkToplevel):
    def __init__(self, parent_assistant):
        super().__init__(parent_assistant.root)
        self.parent_assistant = parent_assistant
        self.geometry("600x600")
        self.title("Live Assistant")
        self.attributes("-topmost", True)
        self.protocol("WM_DELETE_WINDOW", self.cleanup)
        self.content_history = []
        self.current_history_index = -1
        self.is_generating = False
        self.after_ids = []
        default_frame_fg_color = self._apply_appearance_mode(ctk.ThemeManager.theme["CTkFrame"]["fg_color"])
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(2, weight=1)
        self.grid_rowconfigure(4, weight=3)
        self.type_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.type_frame.grid(row=0, column=0, padx=10, pady=(10, 0), sticky="ew")
        self.type_label = ctk.CTkLabel(self.type_frame, text="Content type:")
        self.type_label.pack(side="left", padx=(0, 10))
        self.content_type = tk.StringVar(value="General")
        self.type_options = ["General", "Article", "Email", "Code/Program", "Letter", "Creative Writing", "Technical Documentation"]
        self.type_menu = ctk.CTkOptionMenu(
            self.type_frame, 
            values=self.type_options,
            variable=self.content_type,
            width=150
        )
        self.type_menu.pack(side="left")
        self.search_frame = ctk.CTkFrame(self.type_frame, fg_color="transparent")
        self.search_frame.pack(side="right")
        self.search_entry = ctk.CTkEntry(self.search_frame, placeholder_text="Search history...")
        self.search_entry.pack(side="left", padx=5)
        self.search_button = ctk.CTkButton(
            self.search_frame, 
            text="🔍", 
            width=30,
            command=self.search_history
        )
        self.search_button.pack(side="left")
        self.prompt_label = ctk.CTkLabel(self, text="Enter your prompt:", font=("Arial", 12, "bold"))
        self.prompt_label.grid(row=1, column=0, padx=10, pady=(10, 0), sticky="w")
        self.prompt_frame = ctk.CTkFrame(self, border_width=2, border_color="#4a8af4", fg_color="transparent")
        self.prompt_frame.grid(row=2, column=0, padx=10, pady=5, sticky="nsew")
        self.prompt_frame.grid_columnconfigure(0, weight=1)
        self.prompt_frame.grid_rowconfigure(0, weight=1)
        self.prompt_entry = ctk.CTkTextbox(self.prompt_frame, height=80, wrap="word")
        self.prompt_entry.grid(row=0, column=0, padx=2, pady=2, sticky="nsew")
        self.prompt_entry.bind("<Return>", lambda event: self.generate_response() if event.widget != self.prompt_entry else None)
        self.output_label = ctk.CTkLabel(self, text="Response:", font=("Arial", 12, "bold"))
        self.output_label.grid(row=3, column=0, padx=10, pady=(10, 0), sticky="w")
        self.output_frame = ctk.CTkFrame(self, border_width=2, border_color="#50C878", fg_color=default_frame_fg_color)
        self.output_frame.grid(row=4, column=0, padx=10, pady=5, sticky="nsew")
        self.output_frame.grid_columnconfigure(0, weight=1)
        self.output_frame.grid_rowconfigure(0, weight=1)
        self.output_text = ctk.CTkTextbox(
            self.output_frame,
            height=300,
            state="disabled",
            wrap="word",
            fg_color=default_frame_fg_color,
            text_color=self._apply_appearance_mode(ctk.ThemeManager.theme["CTkTextbox"]["text_color"])
        )
        self.output_text.grid(row=0, column=0, padx=5, pady=5, sticky="nsew")
        self.button_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.button_frame.grid(row=5, column=0, padx=10, pady=5, sticky="ew")
        self.button_frame.grid_columnconfigure((0, 1, 2, 3, 4), weight=1)
        self.submit_button = ctk.CTkButton(
            self.button_frame, 
            text="Generate (Enter)", 
            command=self.generate_response,
            fg_color="#2B7539",  # Green color for primary action
            hover_color="#1E5C2D"
        )
        self.submit_button.grid(row=0, column=0, padx=5, pady=5)
        self.insert_button = ctk.CTkButton(
            self.button_frame, 
            text="Insert (Ctrl+I)", 
            state="disabled", 
            command=self.insert_text_at_cursor
        )
        self.insert_button.grid(row=0, column=1, padx=5, pady=5)
        self.modify_button = ctk.CTkButton(
            self.button_frame, 
            text="Modify (Ctrl+M)", 
            state="disabled", 
            command=self.modify_current_response
        )
        self.modify_button.grid(row=0, column=2, padx=5, pady=5)
        self.retry_button = ctk.CTkButton(
            self.button_frame, 
            text="Retry (Ctrl+R)", 
            state="disabled", 
            command=self.retry_generation
        )
        self.retry_button.grid(row=0, column=3, padx=5, pady=5)
        self.export_button = ctk.CTkButton(
            self.button_frame, 
            text="Export (Ctrl+E)", 
            state="disabled", 
            command=self.export_content
        )
        self.export_button.grid(row=0, column=4, padx=5, pady=5)
        self.history_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.history_frame.grid(row=6, column=0, padx=10, pady=5, sticky="ew")
        self.history_frame.grid_columnconfigure((0, 1, 2, 3), weight=1)
        self.prev_button = ctk.CTkButton(
            self.history_frame, 
            text="◀ Previous (Ctrl+←)", 
            state="disabled", 
            command=lambda: self.navigate_history(-1)
        )
        self.prev_button.grid(row=0, column=0, padx=5, pady=5)
        self.history_label = ctk.CTkLabel(self.history_frame, text="New Content")
        self.history_label.grid(row=0, column=1, padx=5, pady=5)
        self.next_button = ctk.CTkButton(
            self.history_frame, 
            text="Next ▶ (Ctrl+→)", 
            state="disabled", 
            command=lambda: self.navigate_history(1)
        )
        self.next_button.grid(row=0, column=2, padx=5, pady=5)
        self.delete_button = ctk.CTkButton(
            self.history_frame,
            text="Delete (Del)",
            state="disabled",
            command=self.delete_history_item,
            fg_color="#e74c3c",  # Red color
            hover_color="#c0392b"  # Darker red on hover
        )
        self.delete_button.grid(row=0, column=3, padx=5, pady=5)
        self.progress_bar = ctk.CTkProgressBar(self, mode="indeterminate")
        self.progress_bar.grid(row=7, column=0, padx=10, pady=(0, 5), sticky="ew")
        self.progress_bar.grid_remove()
        self.status_label = ctk.CTkLabel(self, text="Ready", text_color="gray")
        self.status_label.grid(row=8, column=0, padx=10, pady=(0, 10), sticky="w")
        self.generated_response = ""
        self.bind("<Return>", lambda event: self.generate_response() if event.widget != self.prompt_entry else None)
        self.bind("<Control-i>", lambda event: self.insert_text_at_cursor() if self.insert_button.cget("state") == "normal" else None)
        self.bind("<Control-m>", lambda event: self.modify_current_response() if self.modify_button.cget("state") == "normal" else None)
        self.bind("<Control-r>", lambda event: self.retry_generation() if self.retry_button.cget("state") == "normal" else None)
        self.bind("<Control-e>", lambda event: self.export_content() if self.export_button.cget("state") == "normal" else None)
        self.bind("<Control-Left>", lambda event: self.navigate_history(-1) if self.prev_button.cget("state") == "normal" else None)
        self.bind("<Control-Right>", lambda event: self.navigate_history(1) if self.next_button.cget("state") == "normal" else None)
        self.bind("<Control-n>", lambda event: self.new_content())
        self.bind("<Control-f>", lambda event: self.search_entry.focus_set())
        self.bind("<Escape>", lambda event: self.cancel_generation() if self.is_generating else None)
        self.bind("<Delete>", lambda event: self.delete_history_item() if self.delete_button.cget("state") == "normal" else None)
        self.load_history()
        self.update_history_buttons()
    def load_history(self):
        """Load history from database"""
        try:
            conn = sqlite3.connect("assistant_history.db")
            cursor = conn.cursor()
            cursor.execute('''
            CREATE TABLE IF NOT EXISTS live_assistant_history (
                id INTEGER PRIMARY KEY,
                prompt TEXT,
                response TEXT,
                content_type TEXT,
                timestamp TEXT,
                tags TEXT DEFAULT ''
            )
            ''')
            cursor.execute("PRAGMA table_info(live_assistant_history)")
            columns = [column[1] for column in cursor.fetchall()]
            if 'tags' not in columns:
                logging.info("Adding tags column to live_assistant_history table")
                cursor.execute("ALTER TABLE live_assistant_history ADD COLUMN tags TEXT DEFAULT ''")
                conn.commit()
            cursor.execute("SELECT prompt, response, content_type, timestamp, tags FROM live_assistant_history ORDER BY id DESC")
            rows = cursor.fetchall()
            self.content_history = []
            for row in rows:
                self.content_history.append({
                    "prompt": row[0],
                    "response": row[1],
                    "type": row[2],
                    "timestamp": row[3],
                    "tags": row[4] if len(row) > 4 else ""
                })
            conn.close()
            self.current_history_index = -1
            self.update_history_buttons()
        except Exception as e:
            logging.error(f"Error loading Live Assistant history: {e}")
            self.content_history = []
            self.current_history_index = -1
    def save_history(self):
        """Save history to database"""
        try:
            conn = sqlite3.connect("assistant_history.db")
            cursor = conn.cursor()
            cursor.execute('''
            CREATE TABLE IF NOT EXISTS live_assistant_history (
                id INTEGER PRIMARY KEY,
                prompt TEXT,
                response TEXT,
                content_type TEXT,
                timestamp TEXT,
                tags TEXT DEFAULT ''
            )
            ''')
            cursor.execute("DELETE FROM live_assistant_history")
            for item in self.content_history:
                cursor.execute(
                    "INSERT INTO live_assistant_history (prompt, response, content_type, timestamp, tags) VALUES (?, ?, ?, ?, ?)",
                    (item["prompt"], item["response"], item["type"], item["timestamp"], item.get("tags", ""))
                )
            conn.commit()
            conn.close()
        except Exception as e:
            logging.error(f"Error saving Live Assistant history: {e}")
            self.status_label.configure(text=f"Error saving history: {e}", text_color="red")
    def update_history_buttons(self):
        """Update history navigation buttons based on current position"""
        if len(self.content_history) == 0:
            self.prev_button.configure(state="disabled")
            self.next_button.configure(state="disabled")
            self.delete_button.configure(state="disabled")
            self.history_label.configure(text="No history")
            return
        if self.current_history_index == -1:
            self.prev_button.configure(state="normal")
            self.next_button.configure(state="disabled")
            self.delete_button.configure(state="disabled")
            self.history_label.configure(text="New Content")
        elif self.current_history_index == 0:
            self.prev_button.configure(state="disabled")
            self.next_button.configure(state="normal" if len(self.content_history) > 1 else "disabled")
            self.delete_button.configure(state="normal")
            self.history_label.configure(text=f"Item 1 of {len(self.content_history)}")
        elif self.current_history_index == len(self.content_history) - 1:
            self.prev_button.configure(state="normal")
            self.next_button.configure(state="disabled")
            self.delete_button.configure(state="normal")
            self.history_label.configure(text=f"Item {len(self.content_history)} of {len(self.content_history)}")
        else:
            self.prev_button.configure(state="normal")
            self.next_button.configure(state="normal")
            self.delete_button.configure(state="normal")
            self.history_label.configure(text=f"Item {self.current_history_index + 1} of {len(self.content_history)}")
    def navigate_history(self, direction):
        """Navigate through content history
        Args:
            direction: -1 for previous, 1 for next
        """
        new_index = self.current_history_index + direction
        if 0 <= new_index < len(self.content_history):
            self.current_history_index = new_index
            history_item = self.content_history[new_index]
            content_type = history_item["type"]
            self.content_type.set(content_type)
            self.prompt_entry.delete("1.0", "end")
            self.prompt_entry.insert("1.0", history_item["prompt"])
            self.output_text.configure(state="normal")
            self.output_text.delete("1.0", "end")
            if content_type.lower() in ["code", "program"]:
                font = ("Courier New", 12)
                text_bg_color = "#f0f8ff"  # Light blue for code
                text_color = "black"
            else:
                font = ("Arial", 12)
                text_bg_color = self._apply_appearance_mode(ctk.ThemeManager.theme["CTkTextbox"]["fg_color"])
                text_color = self._apply_appearance_mode(ctk.ThemeManager.theme["CTkTextbox"]["text_color"])
            self.output_text.configure(font=font, fg_color=text_bg_color, text_color=text_color)
            self.output_text.insert("1.0", history_item["response"])
            self.output_text.configure(state="disabled")
            self.generated_response = history_item["response"]
            self.insert_button.configure(state="normal")
            self.modify_button.configure(state="normal")
            self.export_button.configure(state="normal")
            self.status_label.configure(
                text=f"Viewing saved content from {history_item['timestamp']}", 
                text_color="blue"
            )
        self.update_history_buttons()
    def delete_history_item(self):
        """Delete the current history item"""
        if self.current_history_index < 0 or self.current_history_index >= len(self.content_history):
            return
        confirm = messagebox.askyesno(
            "Confirm Delete", 
            f"Are you sure you want to delete this history item?\n\nType: {self.content_history[self.current_history_index]['type']}\nTimestamp: {self.content_history[self.current_history_index]['timestamp']}",
            icon="warning"
        )
        if not confirm:
            return
        deleted_item = self.content_history.pop(self.current_history_index)
        logging.info(f"Deleted history item from {deleted_item['timestamp']}")
        self.save_history()
        if len(self.content_history) == 0:
            self.new_content()
        elif self.current_history_index >= len(self.content_history):
            self.current_history_index = len(self.content_history) - 1
            self.navigate_history(0)
        else:
            self.navigate_history(0)
        self.status_label.configure(text="History item deleted", text_color="green")
    def new_content(self):
        """Switch to creating new content"""
        self.current_history_index = -1
        self.prompt_entry.delete("1.0", "end")
        self.output_text.configure(state="normal")
        self.output_text.delete("1.0", "end")
        self.output_text.configure(font=("Arial", 12))
        self.output_text.configure(fg_color="white")
        self.output_text.configure(state="disabled")
        self.insert_button.configure(state="disabled")
        self.modify_button.configure(state="disabled")
        self.retry_button.configure(state="disabled")
        self.export_button.configure(state="disabled")
        self.status_label.configure(text="Ready for new content", text_color="gray")
        self.update_history_buttons()
    def generate_response(self):
        """Generate response based on prompt and content type"""
        prompt = self.prompt_entry.get("1.0", "end-1c").strip()
        content_type = self.content_type.get()
        if not prompt:
            self.status_label.configure(text="Please enter a prompt.", text_color="orange")
            return
        self.status_label.configure(text=f"Creating {content_type.lower()}...", text_color="cyan")
        self.submit_button.configure(state="disabled")
        self.insert_button.configure(state="disabled")
        self.modify_button.configure(state="disabled")
        self.retry_button.configure(state="disabled")
        self.progress_bar.grid()
        self.progress_bar.start()
        self.is_generating = True
        self.update()
        try:
            screenshot = ImageGrab.grab()
            logging.info("Screen captured successfully for Live Assistant.")
        except Exception as e:
            logging.error(f"Live Assistant: Error capturing screen: {e}")
            self.status_label.configure(text=f"Error capturing screen: {e}", text_color="red")
            self.submit_button.configure(state="normal")
            self.progress_bar.stop()
            self.progress_bar.grid_remove()
            self.is_generating = False
            return
        enhanced_prompt = f"[{content_type}] {prompt}"
        threading.Thread(
            target=self._generate_in_thread, 
            args=(enhanced_prompt, prompt, content_type, screenshot),
            daemon=True
        ).start()
    def _generate_in_thread(self, enhanced_prompt, prompt, content_type, screenshot):
        """Generate the response in a background thread"""
        try:
            if screenshot is not None and (not hasattr(screenshot, 'size') or screenshot.size[0] <= 0 or screenshot.size[1] <= 0):
                logging.warning("Invalid screenshot detected before API call, setting to None")
                screenshot = None
            response_text = self.parent_assistant.get_response_with_screen(enhanced_prompt, screenshot)
            response_text = self.clean_markdown_formatting(response_text)
            self.after(0, lambda: self._update_with_response(response_text, prompt, content_type))
        except Exception as e:
            error_message = str(e)
            if "Blob" in error_message and "NoneType" in error_message:
                error_message = "Error processing screenshot. Please try again without requesting a screenshot."
            elif "API key" in error_message.lower():
                error_message = "API key error. Please check your Gemini API key configuration."
            elif "rate limit" in error_message.lower():
                error_message = "Rate limit exceeded. Please wait a moment and try again."
            logging.error(f"Live Assistant: Error getting AI response: {e}")
            self.after(0, lambda: self._handle_generation_error(error_message))
    def _update_with_response(self, response_text, prompt, content_type):
        """Update the UI with the generated response"""
        self.generated_response = response_text
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        history_item = {
            "prompt": prompt,
            "response": response_text,
            "type": content_type,
            "timestamp": timestamp,
            "tags": self._extract_tags(prompt, response_text)
        }
        if self.current_history_index == -1:
            self.content_history.insert(0, history_item)
            self.current_history_index = 0
        else:
            self.content_history[self.current_history_index] = history_item
        self.save_history()
        if content_type.lower() in ["code", "program"]:
            font = ("Courier New", 12)
            text_bg_color = "#f0f8ff" # Light blue for code
            text_color = "black"
        else:
            font = ("Arial", 12)
            text_bg_color = self._apply_appearance_mode(ctk.ThemeManager.theme["CTkTextbox"]["fg_color"])
            text_color = self._apply_appearance_mode(ctk.ThemeManager.theme["CTkTextbox"]["text_color"])
        self.output_text.configure(state="normal")
        self.output_text.delete("1.0", "end")
        self.output_text.configure(font=font, fg_color=text_bg_color, text_color=text_color)
        self.output_text.insert("1.0", self.generated_response)
        self.output_text.configure(state="disabled")
        self.status_label.configure(text=f"{content_type} created successfully", text_color="green")
        self.insert_button.configure(state="normal")
        self.modify_button.configure(state="normal")
        self.retry_button.configure(state="normal")
        self.export_button.configure(state="normal")
        self.update_history_buttons()
        self.progress_bar.stop()
        self.progress_bar.grid_remove()
        self.submit_button.configure(state="normal")
        self.is_generating = False
    def _handle_generation_error(self, error_message):
        """Handle errors during response generation"""
        self.status_label.configure(text=f"AI Error: {error_message}", text_color="red")
        self.output_text.configure(state="normal")
        self.output_text.delete("1.0", "end")
        self.output_text.insert("1.0", f"Error: {error_message}")
        self.output_text.configure(state="disabled")
        self.retry_button.configure(state="normal")
        self.submit_button.configure(state="normal")
        self.progress_bar.stop()
        self.progress_bar.grid_remove()
        self.is_generating = False
    def _extract_tags(self, prompt, response):
        """Extract potential tags from content for better searchability"""
        words = set((prompt + " " + response).lower().split())
        common_words = {"the", "a", "an", "and", "or", "but", "in", "on", "at", "to", "for", "with"}
        potential_tags = [word for word in words if len(word) > 3 and word not in common_words]
        return ",".join(sorted(potential_tags)[:10])
    def clean_markdown_formatting(self, text):
        """Remove markdown formatting from text"""
        text = re.sub(r'```(?:\w+)?\n(.*?)```', r'\1', text, flags=re.DOTALL)
        text = re.sub(r'#{1,6}\s+(.*?)(?:\n|$)', r'\1\n', text)
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        text = re.sub(r'__(.*?)__', r'\1', text)
        text = re.sub(r'_(.*?)_', r'\1', text)
        text = re.sub(r'^\s*[-*+]\s+', '', text, flags=re.MULTILINE)
        text = re.sub(r'^\s*\d+\.\s+', '', text, flags=re.MULTILINE)
        text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
        text = re.sub(r'^\s*[-*_]{3,}\s*$', '', text, flags=re.MULTILINE)
        return text.strip()
    def modify_current_response(self):
        """Open a dialog to modify the current response"""
        if not self.generated_response:
            return
        edit_window = ctk.CTkToplevel(self)
        edit_window.title("Modify Content")
        edit_window.geometry("700x500")
        edit_window.transient(self)
        edit_window.grab_set()
        edit_window.grid_columnconfigure(0, weight=1)
        edit_window.grid_rowconfigure(1, weight=1)
        label = ctk.CTkLabel(edit_window, text="Edit your content:")
        label.grid(row=0, column=0, padx=10, pady=(10, 0), sticky="w")
        edit_box = ctk.CTkTextbox(edit_window, wrap="word")
        edit_box.grid(row=1, column=0, padx=10, pady=10, sticky="nsew")
        edit_box.insert("1.0", self.generated_response)
        button_frame = ctk.CTkFrame(edit_window, fg_color="transparent")
        button_frame.grid(row=2, column=0, padx=10, pady=(0, 10), sticky="ew")
        button_frame.grid_columnconfigure((0, 1), weight=1)
        cancel_button = ctk.CTkButton(
            button_frame, 
            text="Cancel", 
            command=edit_window.destroy,
            fg_color="gray"
        )
        cancel_button.grid(row=0, column=0, padx=5, pady=5)
        def save_edits():
            edited_text = edit_box.get("1.0", "end-1c")
            if edited_text:
                self.generated_response = edited_text
                self.output_text.configure(state="normal")
                self.output_text.delete("1.0", "end")
                self.output_text.insert("1.0", edited_text)
                self.output_text.configure(state="disabled")
                if 0 <= self.current_history_index < len(self.content_history):
                    self.content_history[self.current_history_index]["response"] = edited_text
                    self.content_history[self.current_history_index]["timestamp"] = \
                        datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S") + " (edited)"
                    self.save_history()
                self.status_label.configure(text="Content updated", text_color="green")
            edit_window.destroy()
        save_button = ctk.CTkButton(
            button_frame, 
            text="Save Changes", 
            command=save_edits,
            fg_color="#2B7539"
        )
        save_button.grid(row=0, column=1, padx=5, pady=5)
    def insert_text_at_cursor(self):
        """Prepare to insert generated text at cursor position in another application"""
        if not self.generated_response:
            return
        self.status_label.configure(text="Position cursor to insert in 3 seconds...", text_color="yellow")
        self.attributes("-topmost", False)
        self.update()
        after_id = self.after(3000, self._perform_insertion)
        self.after_ids.append(after_id)
    def _perform_insertion(self):
        """Insert the generated text at the current cursor position.
        Attempts smart insertion for Emails (Subject/Body) and code execution.
        """
        try:
            logging.info("Attempting insertion...")
            current_content_type = self.content_type.get()
            text_to_insert = self.generated_response
            logging.info(f"Content type: {current_content_type}, Text length: {len(text_to_insert)}")
            self.iconify()
            logging.info("Window iconified.")
            self.update()
            time.sleep(1.2)
            detected_type = current_content_type
            if "subject:" in text_to_insert.lower()[:50]:
                logging.info("Email pattern detected in content")
                detected_type = "Email"
            elif any(keyword in text_to_insert.lower()[:100] for keyword in ["sub ", "vba", "function ", "dim ", "excel"]):
                logging.info("Code pattern detected in content")
                detected_type = "Code"
            logging.info(f"Using content type: {detected_type} (user selected: {current_content_type})")
            if detected_type == "Email" or current_content_type == "Email":
                logging.info("Email insertion logic triggered.")
                if "subject:" in text_to_insert.lower():
                    subject_match = re.search(r'(?i)subject:\s*(.*?)(?:\n|$)', text_to_insert)
                    if subject_match:
                        subject = subject_match.group(1).strip()
                        body_start = subject_match.end()
                        body = text_to_insert[body_start:].strip()
                        logging.info(f"Extracted Subject: '{subject[:50]}...'")
                        logging.info(f"Extracted Body length: {len(body)} chars")
                        logging.info("Pasting subject...")
                        pyperclip.copy(subject)
                        time.sleep(0.5)
                        pyautogui.hotkey('ctrl', 'v')
                        time.sleep(0.7)
                        logging.info("Pressing Enter to move to body field...")
                        pyautogui.press('enter')
                        time.sleep(0.7)
                        if body:
                            logging.info("Pasting body...")
                            pyperclip.copy(body)
                            time.sleep(0.5)
                            pyautogui.hotkey('ctrl', 'v')
                        self.status_label.configure(text="Email content inserted (Subject & Body)", text_color="green")
                        logging.info("Email insertion complete with Subject/Body split")
                    else:
                        logging.info("Could not extract subject properly, pasting full content...")
                        pyperclip.copy(text_to_insert)
                        time.sleep(0.5)
                        pyautogui.hotkey('ctrl', 'v')
                        self.status_label.configure(text="Email content inserted (as is)", text_color="green")
                else:
                    logging.info("No subject found in email, pasting full content...")
                    pyperclip.copy(text_to_insert)
                    time.sleep(0.5)
                    pyautogui.hotkey('ctrl', 'v')
                    self.status_label.configure(text="Email content inserted (as is)", text_color="green")
            elif detected_type.lower() in ["code", "program"] or current_content_type.lower() in ["code", "program"]:
                logging.info("Code/Program content type detected")
                pyperclip.copy(text_to_insert)
                time.sleep(0.5)
                pyautogui.hotkey('ctrl', 'v')
                logging.info("Code pasted successfully")
                time.sleep(1.2)
                logging.info("Pressing F5 to execute code...")
                pyautogui.press('f5')
                time.sleep(1.0)
                logging.info("Pressing Alt+F11 to return to worksheet...")
                pyautogui.hotkey('alt', 'f11')
                time.sleep(0.5)
                self.status_label.configure(text="Code pasted, executed, and returned to worksheet", text_color="green")
                logging.info("Code insertion and execution complete")
            else:
                logging.info("General content pasting...")
                pyperclip.copy(text_to_insert)
                time.sleep(0.5)
                pyautogui.hotkey('ctrl', 'v')
                self.status_label.configure(text="Content inserted", text_color="green")
                logging.info("Content insertion complete")
            logging.info("Restoring window...")
            self.after(1000, self.deiconify)
            self.attributes("-topmost", True)
        except Exception as e:
            logging.error(f"Error inserting text: {e}", exc_info=True)
            try:
                logging.info("Error occurred, attempting to restore window...")
                self.deiconify()
                self.attributes("-topmost", True)
            except tk.TclError:
                logging.warning("Window likely already destroyed during error handling.")
                pass
            self.status_label.configure(text=f"Error inserting text: {e}", text_color="red")
    def retry_generation(self):
        """Retry the current generation with the same prompt"""
        self.generate_response()
    def search_history(self):
        """Search through history based on keywords"""
        search_term = self.search_entry.get().lower().strip()
        if not search_term:
            self.status_label.configure(text="Please enter a search term", text_color="orange")
            return
        self.status_label.configure(text=f"Searching for '{search_term}'...", text_color="blue")
        results = []
        for i, item in enumerate(self.content_history):
            content_to_search = (
                item["prompt"].lower() + " " + 
                item["response"].lower() + " " + 
                item["type"].lower() + " " + 
                item.get("tags", "").lower()
            )
            if search_term in content_to_search:
                results.append((i, item))
        if not results:
            self.status_label.configure(text=f"No results found for '{search_term}'", text_color="orange")
            return
        self._show_search_results(search_term, results)
    def _show_search_results(self, search_term, results):
        """Display search results in a dialog"""
        results_window = ctk.CTkToplevel(self)
        results_window.title(f"Search Results: {search_term}")
        results_window.geometry("600x400")
        results_window.transient(self)
        results_window.grid_columnconfigure(0, weight=1)
        results_window.grid_rowconfigure(1, weight=1)
        label = ctk.CTkLabel(results_window, text=f"Found {len(results)} results for '{search_term}':")
        label.grid(row=0, column=0, padx=10, pady=(10, 0), sticky="w")
        scroll_frame = ctk.CTkScrollableFrame(results_window)
        scroll_frame.grid(row=1, column=0, padx=10, pady=10, sticky="nsew")
        scroll_frame.grid_columnconfigure(0, weight=1)
        for i, (index, item) in enumerate(results):
            result_frame = ctk.CTkFrame(scroll_frame)
            result_frame.grid(row=i, column=0, padx=5, pady=5, sticky="ew")
            result_frame.grid_columnconfigure(1, weight=1)
            header = ctk.CTkLabel(
                result_frame, 
                text=f"{item['type']} - {item['timestamp']}", 
                font=("Arial", 12, "bold")
            )
            header.grid(row=0, column=0, columnspan=2, padx=5, pady=(5,0), sticky="w")
            prompt_preview = item["prompt"][:50] + "..." if len(item["prompt"]) > 50 else item["prompt"]
            prompt = ctk.CTkLabel(result_frame, text=f"Prompt: {prompt_preview}")
            prompt.grid(row=1, column=0, columnspan=2, padx=5, pady=2, sticky="w")
            def make_open_func(idx):
                return lambda: self._load_search_result(idx, results_window)
            open_btn = ctk.CTkButton(
                result_frame, 
                text="Open", 
                width=80, 
                command=make_open_func(index)
            )
            open_btn.grid(row=2, column=1, padx=5, pady=5, sticky="e")
        self.status_label.configure(text=f"Found {len(results)} results for '{search_term}'", text_color="green")
    def _load_search_result(self, index, results_window):
        """Load a search result by its index"""
        self.current_history_index = index
        history_item = self.content_history[index]
        self.content_type.set(history_item["type"])
        self.prompt_entry.delete("1.0", "end")
        self.prompt_entry.insert("1.0", history_item["prompt"])
        self.output_text.configure(state="normal")
        self.output_text.delete("1.0", "end")
        self.output_text.insert("1.0", history_item["response"])
        self.output_text.configure(state="disabled")
        self.generated_response = history_item["response"]
        self.insert_button.configure(state="normal")
        self.modify_button.configure(state="normal")
        self.export_button.configure(state="normal")
        self.status_label.configure(
            text=f"Loaded content from {history_item['timestamp']}", 
            text_color="blue"
        )
        self.update_history_buttons()
        results_window.destroy()
    def export_content(self):
        """Export the current content to a file"""
        if not self.generated_response:
            return
        content_type = self.content_type.get()
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        extension = ".txt"
        if content_type.lower() == "code/program":
            extension = ".py"
        elif content_type.lower() == "email":
            extension = ".eml"
        elif content_type.lower() == "article" or content_type.lower() == "technical documentation":
            extension = ".md"
        default_filename = f"{content_type.lower().replace('/', '_')}_{timestamp}{extension}"
        file_path = filedialog.asksaveasfilename(
            defaultextension=extension,
            filetypes=[("Text files", "*.txt"), ("All files", "*.*")],
            initialfile=default_filename
        )
        if not file_path:
            return
        try:
            with open(file_path, "w", encoding="utf-8") as f:
                if content_type.lower() in ["article", "technical documentation"]:
                    prompt = self.prompt_entry.get("1.0", "end-1c").strip()
                    f.write(f"# {prompt.splitlines()[0] if prompt else content_type}\n\n")
                    f.write(self.generated_response)
                else:
                    f.write(self.generated_response)
            self.status_label.configure(text=f"Content exported to {os.path.basename(file_path)}", text_color="green")
        except Exception as e:
            logging.error(f"Error exporting content: {e}")
            self.status_label.configure(text=f"Error exporting: {e}", text_color="red")
    def cancel_generation(self):
        """Cancel the current generation process"""
        if not self.is_generating:
            return
        self.status_label.configure(text="Generation cancelled", text_color="orange")
        self.submit_button.configure(state="normal")
        self.progress_bar.stop()
        self.progress_bar.grid_remove()
        self.is_generating = False
    def cleanup(self):
        """Clean up resources when closing the window"""
        try:
            self.save_history()
            for after_id in self.after_ids:
                self.after_cancel(after_id)
            if hasattr(self.parent_assistant, 'live_assistant_win'):
                delattr(self.parent_assistant, 'live_assistant_win')
            self.destroy()
        except Exception as e:
            logging.error(f"Error cleaning up LiveAssistantWindow: {e}")
            self.destroy()
class VoiceAssistant:
    def __init__(self, root):
        self.root = root
        self.programming_window = None
        self.root.title("A.N.I.S AI Assistant")
        self.root.geometry("650x700")
        self.root.resizable(False, False)
        self.stop_listening_func = None
        self.default_prompts = {
            'study_notes': """You are a helpful study assistant focusing on {topic}.,You should make friendly conversation and simple conscise and short answers  within limit of 4 sentenxce
            Please provide a clear and informative response to this question: {question},should be short and should be undersatdnble by beginner level students
            Include key points, examples, and explanations where relevant.swer should be short and concise in 1 minutes or less."""
        }
        self.programming_window = None
        self.study_mode = False
        self.current_topic = None
        self.study_progress = {}
        self.api_key = self.load_setting('gemini_api_key')
        self.quiz_window = None
        self.quiz_questions = []
        self.quiz_answers = []
        self.quiz_correct_answers = []
        self.quiz_score = 0
        self.quiz_total_questions = 0
        self.recognizer = sr.Recognizer()
        self.recognizer.energy_threshold = 10000000
        self.recognizer.pause_threshold = 0.8
        self.recognizer.non_speaking_duration = 0.8
        self.recognizer.phrase_threshold = 0.3
        self.microphone = sr.Microphone()
        self.stop_listening_func = None 
        self.microphone = sr.Microphone()
        self.main_container = ctk.CTkFrame(
            self.root,
            fg_color='#0a0a1f',
            corner_radius=0
        )
        self.main_container.pack(fill=tk.BOTH, expand=True, padx=30, pady=20)
        self.banner_frame = ctk.CTkFrame(
            self.main_container,
            fg_color='#0a0a1f',
            height=70,
            corner_radius=0
        )
        self.banner_frame.pack(fill=tk.X, pady=(0, 10))
        self.gradient_frame = ctk.CTkFrame(
            self.main_container,
            fg_color='#0a0a1f',
            border_color='#2a2a4a',
            border_width=2,
            corner_radius=10
        )
        self.gradient_frame.pack(fill=tk.X, pady=20)
        self.wave_frame = ctk.CTkFrame(
            self.gradient_frame,
            fg_color='#0a0a1f',
            corner_radius=8
        )
        self.wave_frame.pack(fill=tk.X, padx=15, pady=15)
        self.themes = {
            "Dark Blue": {
                 "bg": "#0D0221",
                "fg": "#00FF9F",
                "accent": "#FF2975",
                "secondary": "#2DE2E6",
                "highlight": "#261447",
                "text_bg": "#0F0326",
                "success": "#00FF9F",
                "error": "#FF2975",
                "text": "#00FF9F",
                "text_secondary": "#2DE2E6",
                "input_bg": "#0F0326",
                "border": "#261447"
            },
            "Cyberpunk": {
                  "bg": "#0a0a1f",
                "fg": "#e2e2e2",
                "accent": "#4361ee",
                "secondary": "#8f9aff",
                "highlight": "#2a2a4a",
                "text_bg": "#12122a",
                "success": "#4CAF50",
                "error": "#ff4d4d",
                "text": "#e2e2e2",
                "text_secondary": "#8f9aff",
                "input_bg": "#12122a",
                "border": "#2a2a4a"
            },
            "Forest": {
                "bg": "#1A2F1D",
                "fg": "#C7F9CC",
                "accent": "#57BA5E",
                "secondary": "#80ED99",
                "highlight": "#2D4A31",
                "text_bg": "#1F3823",
                "success": "#57BA5E",
                "error": "#FF6B6B",
                "text": "#C7F9CC",
                "text_secondary": "#80ED99",
                "input_bg": "#1F3823",
                "border": "#2D4A31"
            },
            "Ocean": {
                "bg": "#0A192F",
                "fg": "#E6F1FF",
                "accent": "#64FFDA",
                "secondary": "#8892B0",
                "highlight": "#172A45",
                "text_bg": "#0C1B2B",
                "success": "#64FFDA",
                "error": "#FF647F",
                "text": "#E6F1FF",
                "text_secondary": "#8892B0",
                "input_bg": "#0C1B2B",
                "border": "#172A45"
            },
            "Tron Legacy": {
                "bg": "#0C141F",
                "fg": "#00F2FF",
                "accent": "#15f4ee",
                "secondary": "#2a2a4a",
                "highlight": "#2a2a4a",
                "text_bg": "#12122a",
                "success": "#4CAF50",
                "error": "#ff4d4d",
                "text": "#00F2FF",
                "text_secondary": "#15f4ee",
                "input_bg": "#12122a",
                "border": "#2a2a4a"
            }
        }
        self.current_theme = "Dark Blue"
        self.wave_vis = WaveVisualizer(
            self.wave_frame, 
            width=820,
            height=100,
            bg='#000000',
            highlightthickness=1,
            highlightbackground='#2a2a4a'
        )
        self.wave_vis.pack(pady=5)
        self.create_gui()
        self.apply_theme(self.current_theme)
        self.root.protocol("WM_DELETE_WINDOW", self.exit_assistant)
        try:
            import google.generativeai as genai
            api_key = self.load_setting('gemini_api_key')
            model_name=self.load_setting('gemini_model')
            genai.configure(api_key=api_key)
            self.gemini_model = genai.GenerativeModel(model_name)
            self.chat = self.gemini_model.start_chat(history=[])
        except Exception as e:
            self.handle_errors(e, "Error initializing Gemini API")
        self.file_assistant = FileHandlingAssistant(self.root, self)
        self.conversation_history = []
        self.max_history_length = 5
        self.last_topic = None
    def _create_general_settings(self, parent):
        """Create general settings tab content"""
        general_frame = ctk.CTkFrame(parent)
        general_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        title_label = ctk.CTkLabel(
            general_frame,
            text="General Settings",
            font=ctk.CTkFont(size=18, weight="bold")
        )
        title_label.pack(anchor=tk.W, pady=(0, 15))
        api_frame = ctk.CTkFrame(general_frame)
        api_frame.pack(fill=tk.X, pady=10)
        api_label = ctk.CTkLabel(
            api_frame,
            text="API Key",
            font=ctk.CTkFont(size=16, weight="bold")
        )
        api_label.pack(anchor=tk.W, padx=15, pady=(10, 5))
        api_desc = ctk.CTkLabel(
            api_frame,
            text="Your Google Gemini API key for natural language processing.",
            wraplength=700
        )
        api_desc.pack(anchor=tk.W, padx=15, pady=(0, 5))
        api_entry = ctk.CTkEntry(api_frame, width=400, show="•")
        api_entry.pack(anchor=tk.W, padx=15, pady=5)
        api_entry.insert(0, self.api_key)
        self.api_entry = api_entry
        startup_frame = ctk.CTkFrame(general_frame)
        startup_frame.pack(fill=tk.X, pady=10)
        startup_label = ctk.CTkLabel(
            startup_frame,
            text="Startup Options",
            font=ctk.CTkFont(size=16, weight="bold")
        )
        startup_label.pack(anchor=tk.W, padx=15, pady=(10, 5))
        self.startup_var = tk.BooleanVar(value=False)
        startup_check = ctk.CTkCheckBox(
            startup_frame,
            text="Start listening on launch",
            variable=self.startup_var,
            onvalue=True,
            offvalue=False
        )
        startup_check.pack(anchor=tk.W, padx=15, pady=5)
        history_frame = ctk.CTkFrame(general_frame)
        history_frame.pack(fill=tk.X, pady=10)
        history_label = ctk.CTkLabel(
            history_frame,
            text="Conversation History",
            font=ctk.CTkFont(size=16, weight="bold")
        )
        history_label.pack(anchor=tk.W, padx=15, pady=(10, 5))
        history_desc = ctk.CTkLabel(
            history_frame,
            text="Control how many conversation turns to remember for context.",
            wraplength=700
        )
        history_desc.pack(anchor=tk.W, padx=15, pady=(0, 5))
        history_slider = ctk.CTkSlider(
            history_frame,
            from_=5,
            to=50,
            number_of_steps=9,
            width=300
        )
        history_slider.pack(anchor=tk.W, padx=15, pady=5)
        history_slider.set(self.max_history_length)
        self.history_slider = history_slider
        history_value = ctk.CTkLabel(
            history_frame,
            text=f"Maximum history length: {self.max_history_length}"
        )
        history_value.pack(anchor=tk.W, padx=15, pady=(0, 10))
        self.history_value_label = history_value
        def update_history_label(value):
            value = int(value)
            self.history_value_label.configure(text=f"Maximum history length: {value}")
        history_slider.configure(command=update_history_label)
    def browse_and_process_file(self, command):
        """Browse for a file and process it based on the command"""
        file_path = filedialog.askopenfilename()
        if file_path:
            self.speak(f"Selected file: {os.path.basename(file_path)}")
            self.file_assistant.file_path_var.set(file_path)
            self.file_assistant.show_window()
    def show_recent_files(self):
        """Show recent files window"""
        self.file_assistant.show_recent_files_window()
        self.speak("Showing your recently processed files")
    def show_recent_files_window(self, initial_tab_name="All Files", filter_op_type=None, filter_date=None, filter_weekday=None):
        """Show a window with recently processed files by type, with optional filtering."""
        if hasattr(self, 'recent_window') and self.recent_window is not None and self.recent_window.winfo_exists():
            try:
                 self.recent_window.destroy()
            except tk.TclError:
                 pass
        self.recent_window = ctk.CTkToplevel(self.window or self.parent)
        self.recent_window.title("Recent Files")
        self.recent_window.geometry("700x550")
        if self.window and self.window.winfo_exists():
            self.recent_window.transient(self.window)
        else:
             self.recent_window.transient(self.parent)
        theme = self.assistant.themes[self.assistant.current_theme]
        self.recent_window.configure(fg_color=theme["bg"])
        self.recent_window.grab_set()
        filter_info_text = "Filters: None"
        filters_applied = []
        if filter_op_type:
            filters_applied.append(f"Type={filter_op_type}")
        if filter_date:
            filters_applied.append(f"Date={filter_date}")
        if filter_weekday:
            weekday_names = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
            weekday_int_map = {'1': 0, '2': 1, '3': 2, '4': 3, '5': 4, '6': 5, '0': 6}
            try:
                 filters_applied.append(f"Day={weekday_names[weekday_int_map[filter_weekday]]}")
            except KeyError:
                 filters_applied.append(f"Day=Invalid({filter_weekday})")
        if filters_applied:
            filter_info_text = "Filters: " + ", ".join(filters_applied)
        filter_label = ctk.CTkLabel(
             self.recent_window,
             text=filter_info_text,
             font=("Segoe UI", 10),
             text_color=theme["secondary"]
        )
        filter_label.pack(pady=(5, 0), padx=20, anchor='w')
        tabview = ctk.CTkTabview(
            self.recent_window,
            fg_color=theme["highlight"],
            segmented_button_fg_color=theme["bg"],
            segmented_button_selected_color=theme["accent"],
            segmented_button_unselected_color=theme["bg"]
        )
        tabview.pack(fill=tk.BOTH, expand=True, padx=20, pady=(5, 20))
        tab_names = ["All Files", "Summaries", "Converted", "Extracted", "Exported", "Opened", "Other"]
        tabs = {}
        for name in tab_names:
             tabs[name] = tabview.add(name)
        try:
            db_path = os.path.join(self.assistant_folder, "files_database.db")
            if not os.path.exists(db_path):
                conn = sqlite3.connect(db_path)
                cursor = conn.cursor()
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS processed_files (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        file_path TEXT,
                        file_name TEXT,
                        operation_type TEXT,
                        timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
                    )
                ''')
                conn.commit()
                conn.close()
                all_files_data = []
            else:
                conn = sqlite3.connect(db_path)
                conn.execute("PRAGMA journal_mode=WAL;")
                cursor = conn.cursor()
                query = "SELECT id, file_path, file_name, operation_type, timestamp FROM processed_files"
                conditions = []
                params = []
                if filter_op_type:
                     if filter_op_type.startswith('exported'):
                          conditions.append("operation_type LIKE ?")
                          params.append('exported_%')
                     elif filter_op_type == 'other':
                          known_types = ('summary', 'converted', 'extract', 'opened')
                          conditions.append(f"operation_type NOT LIKE 'exported_%' AND operation_type NOT IN ({','.join('?'*len(known_types))})")
                          params.extend(known_types)
                     else:
                         conditions.append("operation_type = ?")
                         params.append(filter_op_type)
                if filter_date:
                    conditions.append("DATE(timestamp) = DATE(?)")
                    params.append(filter_date)
                elif filter_weekday:
                    conditions.append("STRFTIME('%w', timestamp) = ?")
                    params.append(filter_weekday)
                if conditions:
                    query += " WHERE " + " AND ".join(conditions)
                query += " ORDER BY timestamp DESC LIMIT 100"
                cursor.execute(query, params)
                all_files_data = cursor.fetchall()
                conn.close()
            tab_data = {
                "All Files": [],
                "Summaries": [],
                "Converted": [],
                "Extracted": [],
                "Exported": [],
                "Opened": [],
                "Other": []
            }
            for file_row in all_files_data:
                file_id, file_path, file_name, op_type, timestamp = file_row
                op_type_lower = op_type.lower()
                tab_data["All Files"].append(file_row)
                if op_type_lower == 'summary':
                    tab_data["Summaries"].append(file_row)
                elif op_type_lower == 'converted':
                    tab_data["Converted"].append(file_row)
                elif op_type_lower == 'extract':
                    tab_data["Extracted"].append(file_row)
                elif op_type_lower.startswith('exported'):
                     tab_data["Exported"].append(file_row)
                elif op_type_lower == 'opened':
                     tab_data["Opened"].append(file_row)
                else:
                     tab_data["Other"].append(file_row)
            for tab_name in tab_names:
                 parent_tab = tabs.get(tab_name)
                 if parent_tab:
                     self.create_files_list(parent_tab, tab_data[tab_name], show_type=(tab_name == "All Files"))
            try:
                 if initial_tab_name in tab_names:
                     tabview.set(initial_tab_name)
                 else:
                     tabview.set("All Files")
            except Exception as tab_e:
                 logging.error(f"Error setting tab: {tab_e}")
                 tabview.set("All Files")
        except sqlite3.Error as db_e:
             logging.error(f"Database error in show_recent_files: {db_e}")
             error_frame = ctk.CTkFrame(tabs["All Files"], fg_color=theme["bg"])
             error_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
             ctk.CTkLabel(
                 error_frame,
                 text=f"Error loading files from database: {db_e}",
                 font=("Segoe UI", 14),
                 text_color=theme["error"]
             ).pack(pady=20)
        except Exception as e:
            logging.error(f"Error showing recent files window: {str(e)}")
            error_frame = ctk.CTkFrame(tabs["All Files"], fg_color=theme["bg"])
            error_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
            ctk.CTkLabel(
                error_frame,
                text=f"An unexpected error occurred: {str(e)}",
                font=("Segoe UI", 14),
                text_color=theme["error"]
            ).pack(pady=20)
    def show_coordinates_manager(self):
        """Show the coordinates manager window"""
        if hasattr(self, 'show_command_history'):
            self.show_command_history(filter_type="coordinates")
        else:
            messagebox.showinfo("Coordinates Manager", "The coordinates manager feature is not fully implemented yet.")
    def show_live_assistant_window(self):
        """Opens the Live Assistant window."""
        if hasattr(self, 'live_assistant_win') and self.live_assistant_win.winfo_exists():
            self.live_assistant_win.lift()
            self.live_assistant_win.focus()
        else:
            self.live_assistant_win = LiveAssistantWindow(self)
            logging.info("Live Assistant window opened.")
    def get_response_with_screen(self, user_prompt, pil_image):
        """Sends user prompt and screen image to Gemini for focused, contextual response."""
        try:
            model = genai.GenerativeModel('gemini-2.0-flash')
            prompt_parts = []
            content_type = "General"
            content_types = ["Article", "Email", "Code", "Program", "Letter", "General"]
            type_match = re.match(r'\[(.*?)\]\s*(.*)', user_prompt)
            if type_match:
                detected_type = type_match.group(1).strip()
                if any(t.lower() == detected_type.lower() for t in content_types):
                    content_type = detected_type
                    user_prompt = type_match.group(2).strip()
            if content_type.lower() == "article":
                base_prompt = (
                    "You are a professional content writer tasked with creating a well-structured article.\n"
                    "Your article should:\n"
                    "1. Have a clear introduction, body paragraphs, and conclusion\n"
                    "2. Be informative, engaging, and factually accurate\n"
                    "3. Use appropriate headings where necessary\n"
                    "4. Be directly ready to copy and paste (no markdown or formatting commands)\n"
                    "5. Be complete, thorough, and directly address the user's request\n"
                    "\nAnalyze the screenshot if relevant to the article topic."
                )
            elif content_type.lower() in ["code", "program"]:
                base_prompt = (
                    "You are an expert programmer tasked with writing high-quality, working code.\n"
                    "Your code should:\n"
                    "1. Be complete and ready to run\n"
                    "2. Include necessary imports and setup\n"
                    "3. Follow best practices and proper style conventions\n"
                    "4. Be well-commented where appropriate\n"
                    "5. Solve the user's problem efficiently\n"
                    "\nAnalyze the screenshot for any relevant context, such as existing code, error messages, or problem specifications."
                )
            elif content_type.lower() == "email":
                base_prompt = (
                    "You are an expert email writer tasked with creating a professional email.\n"
                    "Your email should:\n"
                    "1. Have a clear subject line (if requested)\n"
                    "2. Include appropriate greeting and closing\n"
                    "3. Be professionally written with appropriate tone\n"
                    "4. Be concise yet complete\n"
                    "5. Be ready to copy and paste (no formatting commands)\n"
                    "\nAnalyze the screenshot for any relevant context or information to include in the email."
                )
            elif content_type.lower() == "letter":
                base_prompt = (
                    "You are an expert writer tasked with creating a formal letter.\n"
                    "Your letter should:\n"
                    "1. Include proper sender/recipient details (if provided)\n"
                    "2. Have appropriate opening and closing salutations\n"
                    "3. Be formally structured with clear paragraphs\n"
                    "4. Maintain a consistent and appropriate tone\n"
                    "5. Be ready to copy and paste (no formatting commands)\n"
                    "\nAnalyze the screenshot for any relevant context or information to include in the letter."
                )
            else:
                base_prompt = (
                    "You are a versatile AI assistant with screen analysis capabilities.\n"
                    "Analyze the screenshot and respond directly to the user's request.\n"
                    "Your response should be:\n"
                    "1. Helpful and insightful\n"
                    "2. Correct and precise\n"
                    "3. Directly answering what was asked\n"
                    "4. As concise as possible while still being complete\n"
                    "5. Ready to paste as-is (don't use code blocks or markdown formatting)\n"
                    "\nIf the user asks for code or a specific solution, provide a complete, working implementation."
                )
            prompt_parts.append(base_prompt)
            prompt_parts.append(f"\nUser Request: {user_prompt}")
            prompt_parts.append(f"\nContent Type: {content_type}")
            prompt_parts.append("\nScreenshot to analyze:")
            prompt_parts.append(pil_image)
            logging.info(f"Sending {content_type} request to Gemini with prompt: {user_prompt}")
            response = model.generate_content(prompt_parts)
            if response and response.text:
                response_text = response.text.strip()
                logging.info(f"Received {content_type} response from Gemini for live assistant.")
                log_manager.log_conversation(f"[Live Assistant:{content_type}] {user_prompt}", 
                                            response_text, True)
                return response_text
            else:
                logging.warning("Gemini returned an empty response for live assistant.")
                try:
                    logging.warning(f"Gemini response details: {response.prompt_feedback}")
                except Exception:
                    pass
                return f"Error: Failed to generate {content_type.lower()}. Please try a different prompt or check API access."
        except Exception as e:
            logging.error(f"Error in get_response_with_screen: {e}")
            return f"Error: {str(e)}"
    def show_command_history(self, filter_type="all"):
        """Display a window with the command history."""
        try:
            history_window = ctk.CTkToplevel(self.root)
            history_window.title("Command History")
            history_window.geometry("900x600")
            history_window.transient(self.root)
            history_window.grab_set()
            history_window.grid_columnconfigure(0, weight=1)
            history_window.grid_rowconfigure(3, weight=1)
            filter_frame = ctk.CTkFrame(history_window)
            filter_frame.grid(row=0, column=0, padx=20, pady=(20, 10), sticky="ew")
            filter_label = ctk.CTkLabel(filter_frame, text="Filter by:")
            filter_label.grid(row=0, column=0, padx=10, pady=10)
            self.filter_var = tk.StringVar(value=filter_type)
            all_radio = ctk.CTkRadioButton(filter_frame, text="All", variable=self.filter_var, 
                                        value="all", command=lambda: self.refresh_history_display())
            all_radio.grid(row=0, column=1, padx=10, pady=10)
            commands_radio = ctk.CTkRadioButton(filter_frame, text="Commands", variable=self.filter_var, 
                                            value="commands", command=lambda: self.refresh_history_display())
            commands_radio.grid(row=0, column=2, padx=10, pady=10)
            coords_radio = ctk.CTkRadioButton(filter_frame, text="Coordinates", variable=self.filter_var, 
                                        value="coordinates", command=lambda: self.refresh_history_display())
            coords_radio.grid(row=0, column=3, padx=10, pady=10)
            search_frame = ctk.CTkFrame(history_window)
            search_frame.grid(row=1, column=0, padx=20, pady=(0, 10), sticky="ew")
            search_label = ctk.CTkLabel(search_frame, text="Search:")
            search_label.grid(row=0, column=0, padx=5, pady=10, sticky="w")
            self.search_entry = ctk.CTkEntry(search_frame, placeholder_text="Search...")
            self.search_entry.grid(row=0, column=1, padx=5, pady=10, sticky="ew")
            search_button = ctk.CTkButton(search_frame, text="Search", 
                                    command=lambda: self._search_coordinates_display())
            search_button.grid(row=0, column=2, padx=5, pady=10, sticky="e")
            button_frame = ctk.CTkFrame(history_window)
            button_frame.grid(row=2, column=0, padx=20, pady=(0, 10), sticky="ew")
            refresh_button = ctk.CTkButton(button_frame, text="Refresh", 
                                        command=lambda: self._reload_and_refresh_display())
            refresh_button.grid(row=0, column=0, padx=10, pady=10)
            clear_button = ctk.CTkButton(button_frame, text="Clear History", 
                                    command=lambda: self._clear_history_display())
            clear_button.grid(row=0, column=1, padx=10, pady=10)
            self.history_content_frame = ctk.CTkScrollableFrame(history_window)
            self.history_content_frame.grid(row=3, column=0, padx=20, pady=10, sticky="nsew")
            self.history_content_frame.grid_columnconfigure(0, weight=1)
            self.history_window = history_window
            self.coordinates_items = log_manager.get_saved_coordinates()
            self.history_items = log_manager.get_all_conversations(limit=100)
            self.filtered_coordinates = None
            self.refresh_history_display()
        except Exception as e:
            self.display_message("Error", f"Failed to display history: {e}")
            print(f"Error displaying history: {e}")
    def _search_coordinates_display(self):
        """Search and filter coordinates based on user input."""
        search_text = self.search_entry.get().lower().strip()
        if not search_text:
            self.filtered_coordinates = None
        else:
            self.filtered_coordinates = [
                item for item in self.coordinates_items
                if (search_text in str(item.get('element_description', '')).lower() or
                    search_text in str(item.get('app_name', '')).lower())
            ]
        self.refresh_history_display()
    def _reload_and_refresh_display(self):
        """Reload data from the database and refresh the display."""
        self.coordinates_items = log_manager.get_saved_coordinates()
        self.history_items = log_manager.get_all_conversations(limit=100)
        self.filtered_coordinates = None
        if hasattr(self, 'search_entry'):
            self.search_entry.delete(0, tk.END)
        self.refresh_history_display()
        self.display_message("System", "History data refreshed from database.")
    def refresh_history_display(self):
        """Refresh the history display based on selected filter."""
        for widget in self.history_content_frame.winfo_children():
            widget.destroy()
        filter_type = self.filter_var.get()
        if filter_type == "coordinates":
            self._display_coordinates()
        else:
            self._display_history(filter_type)
    def _display_coordinates(self):
        """Display the coordinates in the history window."""
        display_items = self.filtered_coordinates if self.filtered_coordinates is not None else self.coordinates_items
        if not display_items:
            if self.filtered_coordinates is not None:
                label = ctk.CTkLabel(self.history_content_frame, text="No coordinates match your search criteria.")
                label.grid(row=0, column=0, pady=10, sticky="w")
            else:
                label = ctk.CTkLabel(self.history_content_frame, text="No saved coordinates available.")
                label.grid(row=0, column=0, pady=10, sticky="w")
            return
        for i, item in enumerate(display_items):
            self._create_coordinate_frame(i+2, item)
    def _display_history(self, filter_type):
        """Display command or conversation history."""
        items = []
        if filter_type == "all" or filter_type == "commands":
            commands = [item for item in self.history_items if item.get('command_type')]
            items.extend(commands)
        if not items:
            label = ctk.CTkLabel(self.history_content_frame, text="No history items to display.")
            label.grid(row=0, column=0, pady=10, sticky="w")
            return
        for i, item in enumerate(items):
            self._create_history_frame(i+2, item)
    def _create_coordinate_frame(self, row_position, item):
        """Create a frame for a single coordinate entry with edit and delete buttons."""
        item_id = item.get('id', 0)
        timestamp = item.get('timestamp', 'Unknown')
        app_name = item.get('app_name', 'Unknown App')
        element = item.get('element_description', 'Unknown Element')
        x, y = item.get('x', 0), item.get('y', 0)
        success_count = item.get('success_count', 0)
        fail_count = item.get('fail_count', 0)
        item_frame = ctk.CTkFrame(self.history_content_frame, fg_color=("#f0f0f0", "#333333"))
        item_frame.grid(row=row_position, column=0, pady=(0, 10), sticky="ew")
        item_frame.grid_columnconfigure(0, weight=1)
        header_frame = ctk.CTkFrame(item_frame, fg_color="transparent")
        header_frame.grid(row=0, column=0, columnspan=2, padx=10, pady=(5, 0), sticky="ew")
        header_frame.grid_columnconfigure(0, weight=1)
        header_label = ctk.CTkLabel(header_frame, 
                                text=f"[{timestamp}] COORDINATES #{item_id} - {app_name}",
                                font=ctk.CTkFont(weight="bold"))
        header_label.grid(row=0, column=0, sticky="w")
        element_label = ctk.CTkLabel(item_frame, text=f"Element: {element}")
        element_label.grid(row=1, column=0, padx=10, sticky="w")
        coords_label = ctk.CTkLabel(item_frame, text=f"Position: ({x}, {y})")
        coords_label.grid(row=2, column=0, padx=10, sticky="w")
        stats_label = ctk.CTkLabel(item_frame, text=f"Success/Fail: {success_count}/{fail_count}")
        stats_label.grid(row=3, column=0, padx=10, sticky="w")
        button_frame = ctk.CTkFrame(item_frame, fg_color="transparent")
        button_frame.grid(row=4, column=0, padx=10, pady=(0, 5), sticky="w")
        edit_button = ctk.CTkButton(button_frame, text="Edit", width=80,
                                command=lambda id=item_id, elm=element, x_val=x, y_val=y, app=app_name: 
                                    self._edit_coordinates_dialog(id, elm, x_val, y_val, app))
        edit_button.grid(row=0, column=0, padx=(0, 5))
        delete_button = ctk.CTkButton(button_frame, text="Delete", width=80,
                                    fg_color="darkred", hover_color="red",
                                    command=lambda id=item_id: self._delete_coordinates_dialog(id))
        delete_button.grid(row=0, column=1)
    def _create_history_frame(self, row_position, item):
        """Create a frame for displaying a history item (command or conversation)."""
        item_id = item.get('id', 0)
        timestamp = item.get('timestamp', 'Unknown')
        command_type = item.get('command_type', '')
        target = item.get('target', '')
        success = item.get('success', 1) == 1
        details = item.get('details', '')
        item_frame = ctk.CTkFrame(self.history_content_frame, fg_color=("#f0f0f0", "#333333"))
        item_frame.grid(row=row_position, column=0, pady=(0, 10), sticky="ew")
        item_frame.grid_columnconfigure(0, weight=1)
        header_frame = ctk.CTkFrame(item_frame, fg_color="transparent")
        header_frame.grid(row=0, column=0, padx=10, pady=(5, 0), sticky="ew")
        header_frame.grid_columnconfigure(0, weight=1)
        header_label = ctk.CTkLabel(header_frame, 
                                text=f"[{timestamp}] COMMAND #{item_id} - {command_type.upper()}",
                                font=ctk.CTkFont(weight="bold"),
                                text_color="green" if success else "red")
        header_label.grid(row=0, column=0, sticky="w")
        if target:
            target_label = ctk.CTkLabel(item_frame, text=f"Target: {target}")
            target_label.grid(row=1, column=0, padx=10, sticky="w")
        if details:
            details_label = ctk.CTkLabel(item_frame, text=f"Details: {details}", wraplength=800)
            details_label.grid(row=2, column=0, padx=10, sticky="w")
        status_text = "Succeeded" if success else "Failed"
        status_label = ctk.CTkLabel(item_frame, 
                                text=f"Status: {status_text}",
                                text_color="green" if success else "red")
        status_label.grid(row=3, column=0, padx=10, pady=(0, 5), sticky="w")
    def _edit_coordinates_dialog(self, coord_id, element, x, y, app_name):
        """Open a dialog to edit a coordinate record."""
        edit_window = ctk.CTkToplevel(self.history_window)
        edit_window.title(f"Edit Coordinates #{coord_id}")
        edit_window.geometry("400x300")
        edit_window.transient(self.history_window)
        edit_window.grab_set()
        edit_window.grid_columnconfigure(0, weight=1)
        edit_window.grid_columnconfigure(1, weight=2)
        ctk.CTkLabel(edit_window, text="Element:").grid(row=0, column=0, padx=10, pady=10, sticky="w")
        element_entry = ctk.CTkEntry(edit_window, width=250)
        element_entry.grid(row=0, column=1, padx=10, pady=10, sticky="ew")
        element_entry.insert(0, element)
        ctk.CTkLabel(edit_window, text="X Coordinate:").grid(row=1, column=0, padx=10, pady=10, sticky="w")
        x_entry = ctk.CTkEntry(edit_window, width=100)
        x_entry.grid(row=1, column=1, padx=10, pady=10, sticky="w")
        x_entry.insert(0, str(x))
        ctk.CTkLabel(edit_window, text="Y Coordinate:").grid(row=2, column=0, padx=10, pady=10, sticky="w")
        y_entry = ctk.CTkEntry(edit_window, width=100)
        y_entry.grid(row=2, column=1, padx=10, pady=10, sticky="w")
        y_entry.insert(0, str(y))
        ctk.CTkLabel(edit_window, text="App Name:").grid(row=3, column=0, padx=10, pady=10, sticky="w")
        app_entry = ctk.CTkEntry(edit_window, width=250)
        app_entry.grid(row=3, column=1, padx=10, pady=10, sticky="ew")
        app_entry.insert(0, app_name)
        status_label = ctk.CTkLabel(edit_window, text="")
        status_label.grid(row=4, column=0, columnspan=2, padx=10, pady=10)
        button_frame = ctk.CTkFrame(edit_window, fg_color="transparent")
        button_frame.grid(row=5, column=0, columnspan=2, padx=10, pady=10)
        def save_changes():
            try:
                new_element = element_entry.get()
                new_x = int(x_entry.get())
                new_y = int(y_entry.get())
                new_app = app_entry.get()
                result = log_manager.update_coordinates(coord_id, new_element, new_x, new_y, new_app)
                if result:
                    status_label.configure(text="Changes saved successfully!", text_color="green")
                    self.coordinates_items = log_manager.get_saved_coordinates()
                    self.refresh_history_display()
                    edit_window.after(1500, edit_window.destroy)
                else:
                    status_label.configure(text="Failed to save changes.", text_color="red")
            except ValueError:
                status_label.configure(text="X and Y must be integers.", text_color="red")
            except Exception as e:
                status_label.configure(text=f"Error: {e}", text_color="red")
        save_button = ctk.CTkButton(button_frame, text="Save Changes", command=save_changes)
        save_button.grid(row=0, column=0, padx=(0, 10))
        cancel_button = ctk.CTkButton(button_frame, text="Cancel", command=edit_window.destroy)
        cancel_button.grid(row=0, column=1)
    def _delete_coordinates_dialog(self, coord_id):
        """Delete a coordinate record after confirmation."""
        confirm = messagebox.askyesno("Confirm Delete", f"Are you sure you want to delete coordinate #{coord_id}?")
        if confirm:
            result = log_manager.delete_coordinates(coord_id)
            if result:
                self.coordinates_items = log_manager.get_saved_coordinates()
                self.refresh_history_display()
                self.display_message("System", f"Coordinate #{coord_id} deleted successfully.")
            else:
                self.display_message("Error", f"Failed to delete coordinate #{coord_id}.")
    def _clear_history_display(self):
        """Clears the history display."""
        confirm = messagebox.askyesno("Confirm", "Are you sure you want to clear the history display?")
        if confirm:
            filter_type = self.filter_var.get()
            if filter_type == "coordinates":
                self.filtered_coordinates = []
            self.refresh_history_display()
    def stop_typing(self):
        """Stop typing mode"""
        self.typing_active = False
        def update_typing_status_stop():
             try:
                  if hasattr(self, 'status_label') and self.status_label.winfo_exists():
                       self.status_label.configure(text="Listening...")
             except tk.TclError:
                  pass
             except Exception as e:
                  logging.error(f"Error updating status label in stop_typing: {e}")
        if hasattr(self, 'root') and self.root.winfo_exists():
            self.root.after(0, update_typing_status_stop)
    def process_text_input_from_gui(self, event=None):
        """Process text entered in the GUI input field."""
        command_text = self.text_input_entry.get().strip()
        if command_text:
            self.insert_with_timestamp(f"You: {command_text}", "user_msg")
            try:
                if self.text_input_entry.winfo_exists():
                    self.text_input_entry.delete(0, tk.END)
            except tk.TclError:
                pass
            thread = threading.Thread(target=self.process_command, args=(command_text.lower(),), daemon=True)
            thread.start()
    def create_gui(self):
        top_control_frame = ctk.CTkFrame(
            self.main_container,
            fg_color='#0a0a1f',
            height=60,
            corner_radius=0
        )
        top_control_frame.pack(fill=tk.X, pady=(0, 10))
        top_control_frame.pack_propagate(False)
        stats_frame = ctk.CTkFrame(
            top_control_frame,
            fg_color='#0a0a1f',
            corner_radius=0
        )
        stats_frame.pack(side=tk.LEFT, padx=20)
        self.cpu_label = ctk.CTkLabel(
            stats_frame,
            text="CPU: 0%",
            font=("Segoe UI", 10),
            text_color='#8f9aff'
        )
        self.cpu_label.pack(side=tk.LEFT, padx=(0, 15))
        self.memory_label = ctk.CTkLabel(
            stats_frame,
            text="RAM: 0%",
            font=("Segoe UI", 10),
            text_color='#8f9aff'
        )
        self.memory_label.pack(side=tk.LEFT, padx=(0, 15))
        self.time_label = ctk.CTkLabel(
            stats_frame,
            text="",
            font=("Segoe UI", 10),
            text_color='#8f9aff'
        )
        self.time_label.pack(side=tk.LEFT)
        self.update_time()
        control_buttons_frame = ctk.CTkFrame(
            top_control_frame,
            fg_color='#0a0a1f',
        )
        control_buttons_frame.pack(side=tk.RIGHT, padx=20)
        self.settings_button = ctk.CTkButton(
            control_buttons_frame,
            text="⚙",
            font=('Segoe UI', 16),
            fg_color='#0a0a1f',
            text_color='#8f9aff',
            border_width=0,
            corner_radius=8,
            hover_color='#2a2a4a',
            command=self.show_settings
        )
        self.settings_button.pack(side=tk.RIGHT, padx=(10))
        play_container = ctk.CTkFrame(
            control_buttons_frame,
            fg_color='#0a0a1f',
            border_color='#00ff00',
        )
        play_container.pack(side=tk.RIGHT, padx=10)
        self.start_button = ctk.CTkButton(
            play_container,
            text="▶",
            command=self.toggle_listening,
            font=('Segoe UI', 16, 'bold'),
            fg_color='#00ff00',
            text_color='#000000',
            hover_color='#00cc00',
            width=40,
            height=40,
            corner_radius=20,
            border_width=0
        )
        self.start_button.pack(padx=2, pady=2)
        self.mic_frame = ctk.CTkFrame(
            self.main_container,
            fg_color='#0a0a1f',
            height=30,
            corner_radius=0
        )
        self.mic_frame.pack(fill=tk.X, pady=(0, 10))
        self.mic_dots = []
        for i in range(10):
            dot = ctk.CTkLabel(
                self.mic_frame,
                text="●",
                font=("Segoe UI", 8),
                text_color='#1f1f3f'  # Start with dark color
            )
            dot.pack(side=tk.LEFT, padx=2)
            self.mic_dots.append(dot)
        self.start_button._clicked = False
        self.start_button._listening = False
        self.listening_thread = None
        self.update_system_stats()
        self.header_label = ctk.CTkLabel(
            self.banner_frame, 
            text="A.N.I.S",
            font=("Segoe UI", 50, "bold"),
            text_color='#4361ee'
        )
        self.header_label.pack(pady=(5, 0))
        self.subtitle_label = ctk.CTkLabel(
            self.banner_frame,
            text="Artificial Neural Intelligence System",
            font=("Segoe UI Light", 14),
            text_color='#8f9aff'
        )
        self.subtitle_label.pack()
        self.status_frame = ctk.CTkFrame(self.main_container, fg_color='#0a0a1f', corner_radius=0)
        self.status_frame.pack(fill=tk.X, pady=10)
        self.status_label = ctk.CTkLabel(
            self.status_frame,
            text="● Not Listening",
            font=("Segoe UI", 12),
            text_color='#ff4d4d'  # Red when not listening
        )
        self.status_label.pack(side=tk.LEFT, padx=20)
        input_frame = ctk.CTkFrame(
            self.main_container,
            fg_color=self.themes[self.current_theme]["bg"],
            border_color=self.themes[self.current_theme]["highlight"],
            border_width=1,
            corner_radius=10
        )
        input_frame.pack(fill=tk.X, padx=0, pady=(10, 10))
        self.input_frame = input_frame
        self.text_input_entry = ctk.CTkEntry(
            input_frame,
            placeholder_text="Type your command here...",
            font=("Segoe UI", 12),
            fg_color=self.themes[self.current_theme]["input_bg"],
            text_color=self.themes[self.current_theme]["text"],
            border_width=0,
            height=35
        )
        self.text_input_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(15, 10), pady=10)
        self.text_input_entry.bind("<Return>", self.process_text_input_from_gui)
        self.send_button = ctk.CTkButton(
            input_frame,
            text="Send",
            width=70,
            height=35,
            command=self.process_text_input_from_gui,
            fg_color=self.themes[self.current_theme]["accent"],
            text_color='#ffffff', # Assuming white text on accent color works generally
            hover_color=self.themes[self.current_theme].get("highlight", self.themes[self.current_theme]["accent"])
        )
        self.send_button.pack(side=tk.RIGHT, padx=(0, 15), pady=10)
        text_container = ctk.CTkFrame(
            self.main_container,
            fg_color='#0a0a1f',
            border_color='#2a2a4a',
            border_width=1,
            corner_radius=10
        )
        text_container.pack(fill=tk.BOTH, expand=True, pady=20)
        text_frame = ctk.CTkFrame(
            text_container,
            fg_color='#0a0a1f',
            corner_radius=8
        )
        text_frame.pack(fill=tk.BOTH, expand=True, padx=2, pady=2)
        style = ttk.Style()
        style.configure(
            "Custom.Vertical.TScrollbar",
            background='#4361ee',
            troughcolor='#1a1a35',
            width=10,
            arrowsize=0
        )
        self.output_text = ctk.CTkTextbox(
            text_frame,
            height=350,
            width=750,
            fg_color='#12122a',
            text_color='#e2e2e2',
            font=("JetBrains Mono", 10),
            corner_radius=8,
            border_width=0,
            wrap="word"
        )
        self.output_text.pack(fill='both', expand=True, padx=3, pady=3)
        self.output_text._textbox.tag_configure(
            "user_msg", 
            foreground="#BF00FF",
            font=("JetBrains Mono", 9, "bold")
        )
        self.output_text._textbox.tag_configure(
            "assistant_msg", 
            foreground="#4CAF50",
            font=("JetBrains Mono", 10)
        )
        self.output_text._textbox.tag_configure(
            "time_stamp", 
            foreground="#666666",
            font=("JetBrains Mono", 8)
        )
        def insert_with_timestamp(self, text, tag):
            timestamp = datetime.datetime.now().strftime("%H:%M:%S")
            def update_gui():
                try:
                    if hasattr(self, 'output_text') and self.output_text.winfo_exists():
                        self.output_text.configure(state=NORMAL)
                        self.output_text._textbox.insert(END, f"[{timestamp}] ", "time_stamp")
                        self.output_text._textbox.insert(END, f"{text}\n", tag)
                        self.output_text.configure(state=DISABLED)
                        self.output_text.see(END)
                except tk.TclError as e:
                    logging.warning(f"GUI Error in insert_with_timestamp: {e}")
                except Exception as e:
                    logging.error(f"Unexpected error in insert_with_timestamp GUI update: {e}")
            if hasattr(self, 'root') and self.root.winfo_exists():
                self.root.after(0, update_gui)
        self.insert_with_timestamp = insert_with_timestamp.__get__(self)
    def update_time(self):
        """Update the time display"""
        try:
            if hasattr(self, 'time_label') and self.time_label.winfo_exists():
                current_time = datetime.datetime.now().strftime("%I:%M:%S %p")
                self.time_label.configure(text=current_time)
                self._after_id = self.root.after(1000, self.update_time)
        except Exception as e:
            logging.error(f"Error updating time: {str(e)}")
            return
    def update_status(self, is_listening=False):
        """Update the status label safely"""
        try:
            if hasattr(self, 'status_label') and self.status_label.winfo_exists():
                if is_listening:
                    self.status_label.configure(
                        text="● Listening",
                        text_color='#4CAF50'  # Green when listening
                    )
                else:
                    self.status_label.configure(
                        text="● Not Listening",
                        text_color='#ff4d4d'  # Red when not listening
                    )
        except tk.TclError:
            pass
    def type_and_enter(self, command):
        """Type the text at cursor position with minimal delay"""
        try:
            if command.startswith('type '):
                text = command[5:].strip()
                if text:
                    self.typing_active = True
                    self.status_label.configure(text="Typing Mode Active")
                    pyautogui.typewrite(text, interval=0.01)
                    self.speak('Done')
                else:
                    self.speak('Please provide text to type')
        except Exception as e:
            logging.error(f"Error in typing: {str(e)}")
            self.speak("Error typing")
    def show_chats_window(self):
        """Show AI chat window for study notes and interactions using customtkinter"""
        self.speak("Opening AI Tutor Bot")
        self.chat_window = ctk.CTkToplevel(self.root)
        self.chat_window.title(f"Study Notes - {self.current_topic}")
        self.chat_window.geometry("900x700")
        self.chat_window.grid_rowconfigure(0, weight=1)
        self.chat_window.grid_columnconfigure(0, weight=1)
        main_container = ctk.CTkFrame(
            self.chat_window,
            fg_color=self.themes[self.current_theme]["bg"],
            corner_radius=15
        )
        main_container.grid(row=0, column=0, padx=20, pady=20, sticky="nsew")
        main_container.grid_rowconfigure(1, weight=1)
        main_container.grid_columnconfigure(0, weight=1)
        header_frame = ctk.CTkFrame(
            main_container,
            fg_color=self.themes[self.current_theme]["highlight"],
            corner_radius=10
        )
        header_frame.grid(row=0, column=0, padx=20, pady=(20, 10), sticky="ew")
        title_label = ctk.CTkLabel(
            header_frame,
            text=f"Study Session: {self.current_topic}",
            font=ctk.CTkFont(size=24, weight="bold"),
            text_color=self.themes[self.current_theme]["accent"]
        )
        title_label.pack(side="left", padx=10, pady=10)
        status_frame = ctk.CTkFrame(
            header_frame,
            fg_color=self.themes[self.current_theme]["bg"],
            corner_radius=5,
            width=100,
            height=30
        )
        status_frame.pack(side="right", padx=10, pady=10)
        self.chat_status = ctk.CTkLabel(
            status_frame,
            text="● Ready",
            font=ctk.CTkFont(size=14),
            text_color=self.themes[self.current_theme]["success"]
        )
        self.chat_status.pack(fill="both", expand=True)
        chat_frame = ctk.CTkFrame(
            main_container,
            fg_color=self.themes[self.current_theme]["text_bg"],
            corner_radius=10
        )
        chat_frame.grid(row=1, column=0, padx=20, pady=10, sticky="nsew")
        chat_frame.grid_rowconfigure(0, weight=1)
        chat_frame.grid_columnconfigure(0, weight=1)
        self.chat_display = ctk.CTkTextbox(
            chat_frame,
            wrap="word",
            font=ctk.CTkFont(family="Segoe UI", size=12),
            fg_color=self.themes[self.current_theme]["text_bg"],
            text_color=self.themes[self.current_theme]["fg"],
            corner_radius=10,
            border_width=0
        )
        self.chat_display.grid(row=0, column=0, padx=10, pady=10, sticky="nsew")
        scrollbar = ctk.CTkScrollbar(
            chat_frame,
            command=self.chat_display.yview,
            fg_color="transparent",
            button_color=self.themes[self.current_theme]["accent"],
            button_hover_color=self.themes[self.current_theme]["secondary"]
        )
        scrollbar.grid(row=0, column=1, sticky="ns", padx=(0, 10), pady=10)
        self.chat_display.configure(yscrollcommand=scrollbar.set)
        input_frame = ctk.CTkFrame(
            main_container,
            fg_color=self.themes[self.current_theme]["highlight"],
            corner_radius=10
        )
        input_frame.grid(row=2, column=0, padx=20, pady=(10, 20), sticky="ew")
        input_frame.grid_columnconfigure(0, weight=1)
        self.chat_input = ctk.CTkEntry(
            input_frame,
            placeholder_text="Type your question here...",
            font=ctk.CTkFont(size=14),
            fg_color=self.themes[self.current_theme]["text_bg"],
            text_color=self.themes[self.current_theme]["fg"],
            corner_radius=10,
            height=40,
            border_width=0
        )
        self.chat_input.grid(row=0, column=0, padx=(10, 10), pady=10, sticky="ew")
        send_button = ctk.CTkButton(
            input_frame,
            text="Send",
            command=self.send_chat_message,
            font=ctk.CTkFont(size=14, weight="bold"),
            fg_color=self.themes[self.current_theme]["accent"],
            hover_color=self.themes[self.current_theme]["secondary"],
            corner_radius=10,
            width=100,
            height=36
        )
        send_button.grid(row=0, column=1, padx=(0, 10), pady=10)
        self.chat_input.bind("<Return>", lambda e: self.send_chat_message())
        self.chat_window.protocol("WM_DELETE_WINDOW", self.on_chat_window_close)
        def on_enter(e):
            send_button.configure(fg_color=self.themes[self.current_theme]["secondary"])
        def on_leave(e):
            send_button.configure(fg_color=self.themes[self.current_theme]["accent"])
        send_button.bind("<Enter>", on_enter)
        send_button.bind("<Leave>", on_leave)
        self.chat_display.configure(state="normal")
        welcome_message = f"Welcome to your AI Study Session on {self.current_topic}!\n\nI'm here to help you learn and understand this topic better. Feel free to ask any questions, request explanations, or discuss concepts you're struggling with.\n\nHow can I assist you today?"
        self.chat_display.configure(state="disabled")
        self.append_to_chat("AI", welcome_message)
        self.chat_input.focus()
    def on_chat_window_close(self):
        """Handle chat window closing"""
        if hasattr(self, 'chat_window') and self.chat_window.winfo_exists():
            self.chat_window.destroy()
    def append_to_chat(self, sender, message):
        """Append a message to the chat display with enhanced formatting"""
        self.chat_display.configure(state="normal")
        self.chat_display.insert("end", "─" * 50 + "\n")
        timestamp = datetime.datetime.now().strftime("%H:%M")
        self.chat_display.insert("end", f"[{timestamp}] ")
        current_pos = self.chat_display.index("end-1c")
        length = len(f"[{timestamp}] ")
        start_pos = self.chat_display.index(f"end-{length}c")
        self.chat_display._textbox.tag_add("timestamp", start_pos, current_pos)
        self.chat_display._textbox.tag_configure("timestamp", foreground=self.themes[self.current_theme]["secondary"])
        sender_color = self.themes[self.current_theme]["success"] if sender == "AI" else self.themes[self.current_theme]["secondary"]
        self.chat_display.insert("end", f"{sender}\n")
        sender_length = len(sender) + 1
        sender_start = self.chat_display.index(f"end-{sender_length}c")
        sender_end = self.chat_display.index("end-1c")
        self.chat_display._textbox.tag_add(f"sender_{sender}", sender_start, sender_end)
        self.chat_display._textbox.tag_configure(f"sender_{sender}", foreground=sender_color)
        formatted_message = self.format_study_message(message)
        self.chat_display.insert("end", f"{formatted_message}\n\n")
        self.chat_display.configure(state="disabled")
        self.chat_display.see("end")
        self.chat_status.configure(text="● Ready")
    def send_chat_message(self):
        """Handle sending chat messages"""
        message = self.chat_input.get().strip()
        if not message:
            return
        self.chat_input.delete(0, tk.END)
        self.chat_status.configure(text="● Processing...", text_color=self.themes[self.current_theme]["accent"])
        self.append_to_chat("You", message)
        try:
            study_prompt = self.load_prompt('study_notes')
            if not study_prompt:
                study_prompt = self.default_prompts['study_notes']
            formatted_prompt = study_prompt.format(
                topic=self.current_topic,
                question=message
            )
            threading.Thread(target=self._animate_status).start()
            response = self.chat.send_message(formatted_prompt)
            self.append_to_chat("AI", response.text)
            self.save_study_notes(message, response.text)
            self.chat_status.configure(text="● Ready", text_color=self.themes[self.current_theme]["success"])
        except Exception as e:
            error_msg = f"Error generating response: {str(e)}"
            logging.error(error_msg)
            self.append_to_chat("AI", "I apologize, but I encountered an error. Please try again.")
            self.chat_status.configure(text="● Error", text_color=self.themes[self.current_theme]["error"])
    def _animate_status(self):
        """Animate the status indicator during processing"""
        statuses = ["● Processing", "● Processing.", "● Processing..", "● Processing..."]
        for _ in range(10):
            for status in statuses:
                if not hasattr(self, 'chat_status') or not self.chat_window.winfo_exists():
                    return
                self.chat_status.configure(text=status)
                time.sleep(0.2)
    def save_study_notes(self, question, answer):
        """Save study notes to file"""
        try:
            notes_dir = os.path.join('data', 'study_notes')
            os.makedirs(notes_dir, exist_ok=True)
            filename = os.path.join(notes_dir, f"{self.current_topic.replace(' ', '_')}_notes.txt")
            timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            with open(filename, 'a', encoding='utf-8') as f:
                f.write(f"\n[{timestamp}]\n")
                f.write(f"Q: {question}\n")
                f.write(f"A: {answer}\n")
                f.write("-" * 50 + "\n")
        except Exception as e:
            logging.error(f"Error saving study notes: {str(e)}")
    def format_study_message(self, message):
        """Format the study message with proper indentation and structure"""
        lines = message.split('\n')
        formatted_lines = []
        for line in lines:
            line = line.strip()
            if line.startswith('•'):
                formatted_lines.append(f"  {line}")
            elif line.startswith('-'):
                formatted_lines.append(f"  {line}")
            elif ':' in line and not any(line.startswith(prefix) for prefix in ['http', 'https']):
                key, value = line.split(':', 1)
                formatted_lines.append(f"  {key.strip()}:{value}")
            else:
                formatted_lines.append(line)
        return '\n'.join(formatted_lines)
    def start_certification_exam(self):
        """Start the certification exam with 20 questions using customtkinter"""
        try:
            questions = self.generate_certification_questions()
            correct_answers = 0
            total_questions = len(questions)
            exam_window = ctk.CTkToplevel(self.root)
            exam_window.title(f"A.N.I.S Certification Exam - {self.current_topic}")
            exam_window.geometry("1000x800")
            exam_window.grid_rowconfigure(0, weight=1)
            exam_window.grid_columnconfigure(0, weight=1)
            main_container = ctk.CTkFrame(
                exam_window,
                fg_color=self.themes[self.current_theme]["bg"],
                corner_radius=15
            )
            main_container.grid(row=0, column=0, padx=20, pady=20, sticky="nsew")
            main_container.grid_rowconfigure(1, weight=1)
            main_container.grid_columnconfigure(0, weight=1)
            header_frame = ctk.CTkFrame(
                main_container,
                fg_color="transparent",
                corner_radius=10
            )
            header_frame.grid(row=0, column=0, padx=20, pady=(20, 10), sticky="ew")
            title_label = ctk.CTkLabel(
                header_frame,
                text=f"Certification Exam: {self.current_topic}",
                font=ctk.CTkFont(size=28, weight="bold"),
                text_color=self.themes[self.current_theme]["accent"]
            )
            title_label.pack(pady=(0, 5))
            subtitle_label = ctk.CTkLabel(
                header_frame,
                text="Complete all 20 questions • Pass mark: 60%",
                font=ctk.CTkFont(size=16),
                text_color=self.themes[self.current_theme]["secondary"]
            )
            subtitle_label.pack()
            scrollable_frame = ctk.CTkScrollableFrame(
                main_container,
                fg_color="transparent",
                corner_radius=10
            )
            scrollable_frame.grid(row=1, column=0, padx=20, pady=10, sticky="nsew")
            answers = []
            for i, (question, options) in enumerate(questions):
                q_frame = ctk.CTkFrame(
                    scrollable_frame,
                    fg_color=self.themes[self.current_theme]["text_bg"],
                    corner_radius=10
                )
                q_frame.pack(fill="x", padx=10, pady=8)
                q_num = ctk.CTkLabel(
                    q_frame,
                    text=f"Question {i+1}",
                    font=ctk.CTkFont(size=16, weight="bold"),
                    text_color=self.themes[self.current_theme]["accent"]
                )
                q_num.pack(anchor='w', padx=20, pady=(15, 5))
                q_text = ctk.CTkLabel(
                    q_frame,
                    text=question,
                    wraplength=800,
                    justify='left',
                    font=ctk.CTkFont(size=14),
                    text_color=self.themes[self.current_theme]["fg"]
                )
                q_text.pack(anchor='w', padx=20, pady=(0, 15))
                answer_var = tk.StringVar()
                answers.append(answer_var)
                options_frame = ctk.CTkFrame(
                    q_frame,
                    fg_color="transparent",
                    corner_radius=5
                )
                options_frame.pack(fill='x', padx=20, pady=(0, 15))
                for idx, option in enumerate(options[:-1]):
                    option_frame = ctk.CTkFrame(
                        options_frame,
                        fg_color=self.themes[self.current_theme]["highlight"],
                        corner_radius=5
                    )
                    option_frame.pack(fill='x', pady=2)
                    radio = ctk.CTkRadioButton(
                        option_frame,
                        text=f"{chr(65+idx)}. {option}",
                        variable=answer_var,
                        value=option,
                        font=ctk.CTkFont(size=12),
                        fg_color=self.themes[self.current_theme]["accent"],
                        hover_color=self.themes[self.current_theme]["secondary"],
                        text_color=self.themes[self.current_theme]["fg"]
                    )
                    radio.pack(anchor='w', padx=10, pady=5)
                    def on_enter(e, frame=option_frame):
                        frame.configure(fg_color=self.themes[self.current_theme]["secondary"])
                    def on_leave(e, frame=option_frame):
                        frame.configure(fg_color=self.themes[self.current_theme]["highlight"])
                    option_frame.bind("<Enter>", on_enter)
                    option_frame.bind("<Leave>", on_leave)
            def submit_exam():
                nonlocal correct_answers
                unanswered = sum(1 for var in answers if not var.get())
                if unanswered > 0:
                    if not messagebox.askyesno(
                        "Confirm Submission",
                        f"You have {unanswered} unanswered questions.\nAre you sure you want to submit?",
                        icon='warning'
                    ):
                        return
                for i, (_, options) in enumerate(questions):
                    if answers[i].get() == options[-1]:
                        correct_answers += 1
                score = (correct_answers / total_questions) * 100
                results_window = ctk.CTkToplevel(exam_window)
                results_window.title("Exam Results")
                results_window.geometry("400x300")
                results_window.transient(exam_window)
                results_window.grab_set()
                results_container = ctk.CTkFrame(
                    results_window,
                    fg_color=self.themes[self.current_theme]["bg"],
                    corner_radius=15
                )
                results_container.pack(fill="both", expand=True, padx=20, pady=20)
                results_title = ctk.CTkLabel(
                    results_container,
                    text="Exam Results",
                    font=ctk.CTkFont(size=24, weight="bold"),
                    text_color=self.themes[self.current_theme]["accent"]
                )
                results_title.pack(pady=(20, 10))
                score_text = f"Score: {score:.1f}%"
                score_color = self.themes[self.current_theme]["success"] if score >= 60 else self.themes[self.current_theme]["error"]
                score_label = ctk.CTkLabel(
                    results_container,
                    text=score_text,
                    font=ctk.CTkFont(size=32, weight="bold"),
                    text_color=score_color
                )
                score_label.pack(pady=10)
                pass_text = "Congratulations! You passed!" if score >= 60 else "Sorry, you did not pass. Try again!"
                pass_label = ctk.CTkLabel(
                    results_container,
                    text=pass_text,
                    font=ctk.CTkFont(size=16),
                    text_color=self.themes[self.current_theme]["fg"]
                )
                pass_label.pack(pady=10)
                close_button = ctk.CTkButton(
                    results_container,
                    text="Close",
                    command=lambda: [results_window.destroy(), exam_window.destroy()],
                    font=ctk.CTkFont(size=14, weight="bold"),
                    fg_color=self.themes[self.current_theme]["accent"],
                    hover_color=self.themes[self.current_theme]["secondary"],
                    corner_radius=10,
                    width=120
                )
                close_button.pack(pady=20)
                if score >= 60:
                    cert_button = ctk.CTkButton(
                        results_container,
                        text="Generate Certificate",
                        command=self.generate_certificate,
                        font=ctk.CTkFont(size=14, weight="bold"),
                        fg_color=self.themes[self.current_theme]["success"],
                        hover_color=self.themes[self.current_theme]["secondary"],
                        corner_radius=10,
                        width=160
                    )
                    cert_button.pack(pady=10)
                exam_window.destroy()
                return score
            submit_frame = ctk.CTkFrame(
                scrollable_frame,
                fg_color="transparent",
                corner_radius=10
            )
            submit_frame.pack(fill='x', pady=20)
            submit_btn = ctk.CTkButton(
                submit_frame,
                text="Submit Exam",
                command=submit_exam,
                font=ctk.CTkFont(size=16, weight="bold"),
                fg_color=self.themes[self.current_theme]["accent"],
                hover_color=self.themes[self.current_theme]["secondary"],
                corner_radius=10,
                width=200,
                height=40
            )
            submit_btn.pack(pady=10)
            def on_enter(e):
                submit_btn.configure(fg_color=self.themes[self.current_theme]["secondary"])
            def on_leave(e):
                submit_btn.configure(fg_color=self.themes[self.current_theme]["accent"])
            submit_btn.bind("<Enter>", on_enter)
            submit_btn.bind("<Leave>", on_leave)
            exam_window.protocol("WM_DELETE_WINDOW", lambda: [exam_window.destroy()])
            exam_window.wait_window()
            return (correct_answers / total_questions) * 100
        except Exception as e:
            self.handle_errors(e, "Error in certification exam")
            return 0
    def generate_certification_questions(self):
        """Generate 20 questions for certification based on the current topic using Google's Generative AI"""
        try:
            api_key = self.load_setting('gemini_api_key')
            model_name = self.load_setting("gemini_model")
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel(model_name)
            prompt = f"""Generate a certification exam about {self.current_topic}.
            Return exactly 20 multiple choice questions in this Python list format:
            [
                ("What is X?", ["Option A", "Option B", "Option C", "Option D", "Option A"]),
                ("What is Y?", ["Choice 1", "Choice 2", "Choice 3", "Choice 4", "Choice 2"])
            ]
            Where the last item in each options list is the correct answer and matches one of the options exactly.
            Make questions progressively harder and cover different aspects of {self.current_topic}."""
            response = model.generate_content(prompt)
            response_text = response.text.strip()
            response_text = response_text.replace('```python', '').replace('```', '')
            try:
                questions = eval(response_text)
                if isinstance(questions, list) and len(questions) >= 20:
                    return questions[:20]
                raise ValueError("Invalid question format")
            except Exception as e:
                print(f"Error parsing AI response: {str(e)}")
                print(f"Response was: {response_text[:200]}...")
                return self.generate_fallback_questions()
        except Exception as e:
            self.handle_errors(e, "Error in question generation, using fallback questions")
            return self.generate_fallback_questions()
    def generate_fallback_questions(self):
        """Generate fallback questions if API fails"""
        question_templates = [
            {
                "template": "What is the best practice for {topic} in {subject}?",
                "options": [
                    "Following established conventions",
                    "Using latest techniques",
                    "Implementing custom solutions",
                    "Ignoring standards",
                    "Following established conventions"
                ]
            },
            {
                "template": "Which approach is most effective for {topic} implementation?",
                "options": [
                    "Systematic planning and execution",
                    "Trial and error",
                    "Copy-paste from examples",
                    "Random experimentation",
                    "Systematic planning and execution"
                ]
            },
            {
                "template": "What is the key consideration when working with {topic}?",
                "options": [
                    "Understanding core principles",
                    "Quick implementation",
                    "Minimal documentation",
                    "Avoiding testing",
                    "Understanding core principles"
                ]
            }
        ]
        topics = [
            "fundamentals", "best practices", "architecture",
            "implementation", "debugging", "optimization",
            "security", "maintenance", "testing",
            "documentation", "deployment", "scalability",
            "performance", "reliability", "compatibility",
            "integration", "monitoring", "troubleshooting",
            "updates", "version control"
        ]
        questions = []
        for topic in topics:
            template = random.choice(question_templates)
            question = template["template"].format(topic=topic, subject=self.current_topic)
            questions.append((question, template["options"]))
        return questions
    def generate_certificate(self):
        """Generate a professional certificate with A.N.I.S branding"""
        try:
            certificates_dir = os.path.join("C:\\Users\\HP\\Desktop\\AI", "certificates")
            if not os.path.exists(certificates_dir):
                os.makedirs(certificates_dir)
            name_dialog = tk.Toplevel(self.root)
            name_dialog.title("Professional Certificate Information")
            name_dialog.geometry("400x200")
            name_dialog.configure(bg=self.themes[self.current_theme]["bg"])
            name_dialog.transient(self.root)
            name_dialog.grab_set()
            tk.Label(
                name_dialog,
                text="Please enter your full name for the certificate:",
                font=("Arial", 12),
                bg=self.themes[self.current_theme]["bg"],
                fg=self.themes[self.current_theme]["fg"]
            ).pack(pady=20)
            name_entry = tk.Entry(
                name_dialog,
                font=("Arial", 12),
                width=30,
                bg=self.themes[self.current_theme]["text_bg"],
                fg=self.themes[self.current_theme]["fg"]
            )
            name_entry.pack(pady=10)
            name_entry.focus()
            def create_certificate():
                user_name = name_entry.get().strip()
                if not user_name:
                    messagebox.showerror("Error", "Please enter your name")
                    return
                name_dialog.destroy()
                width = 1500
                height = 1050
                certificate = Image.new('RGB', (width, height), 'white')
                try:
                    bg_image = Image.open("C:\\Users\\HP\\Desktop\\Ai\\bgimage.png")
                    bg_image = bg_image.resize((width, height))
                    bg_image = bg_image.convert('RGBA')
                    bg_image.putalpha(64)
                    certificate = Image.new('RGB', (width, height), 'white')
                    certificate.paste(bg_image, (0, 0), bg_image)
                except Exception as e:
                    print(f"Background image error: {str(e)}")
                draw = ImageDraw.Draw(certificate)
                try:
                    title_font = ImageFont.truetype("arial.ttf", 78)
                    name_font = ImageFont.truetype("arial.ttf", 48)
                    body_font = ImageFont.truetype("arial.ttf", 36)
                    small_font = ImageFont.truetype("arial.ttf", 24)
                except Exception:
                    title_font = ImageFont.load_default()
                    name_font = ImageFont.load_default()
                    body_font = ImageFont.load_default()
                    small_font = ImageFont.load_default()
                for i in range(20):
                    color = (
                        int(25 + (i * 2)),
                        int(50 + (i * 2)),
                        int(120 + (i * 2))
                    )
                    draw.rectangle(
                        [i, i, width-i-1, height-i-1],
                        outline=color,
                        width=1
                    )
                inner_border = 50
                draw.rectangle(
                    [inner_border, inner_border, width-inner_border, height-inner_border],
                    outline='navy',
                    width=3
                )
                draw.rectangle(
                    [inner_border+10, inner_border+10, width-inner_border-10, height-inner_border-10],
                    outline='gold',
                    width=2
                )
                corner_size = 100
                for x, y in [(inner_border, inner_border), 
                            (width-inner_border-corner_size, inner_border),
                            (inner_border, height-inner_border-corner_size),
                            (width-inner_border-corner_size, height-inner_border-corner_size)]:
                    draw.rectangle([x, y, x+corner_size, y+corner_size], outline='gold', width=2)
                    draw.rectangle([x+10, y+10, x+corner_size-10, y+corner_size-10], outline='navy', width=2)
                    draw.line([x, y, x+corner_size, y+corner_size], fill='gold', width=1)
                    draw.line([x+corner_size, y, x, y+corner_size], fill='gold', width=1)
                logo_text = "A.N.I.S"
                logo_bbox = draw.textbbox((0, 0), logo_text, font=title_font)
                logo_width = logo_bbox[2] - logo_bbox[0]
                logo_x = (width - logo_width) / 2
                logo_y = 80
                shadow_offset = 3
                draw.text((logo_x + shadow_offset, logo_y + shadow_offset), logo_text, 
                         fill='gray', font=title_font)
                draw.text((logo_x, logo_y), logo_text, fill='navy', font=title_font)
                line_length = logo_width + 100
                line_x = (width - line_length) / 2
                draw.line([(line_x, logo_y - 10), (line_x + line_length, logo_y - 10)], 
                         fill='gold', width=2)
                draw.line([(line_x + 20, logo_y - 20), (line_x + line_length - 20, logo_y - 20)], 
                         fill='navy', width=2)
                draw.line([(line_x, logo_y + 80), (line_x + line_length, logo_y + 80)], 
                         fill='gold', width=2)
                draw.line([(line_x + 20, logo_y + 90), (line_x + line_length - 20, logo_y + 90)], 
                         fill='navy', width=2)
                platform_text = "Learning Platform"
                platform_bbox = draw.textbbox((0, 0), platform_text, font=body_font)
                platform_width = platform_bbox[2] - platform_bbox[0]
                platform_x = (width - platform_width) / 2
                platform_y = 180
                for i in range(3):
                    color = (
                        int(25 + (i * 40)),
                        int(50 + (i * 40)),
                        int(120 + (i * 40))
                    )
                    draw.text((platform_x - i, platform_y - i), platform_text, 
                            fill=color, font=body_font)
                title = "Certificate of Excellence"
                title_bbox = draw.textbbox((0, 0), title, font=title_font)
                title_width = title_bbox[2] - title_bbox[0]
                draw.text(((width - title_width) / 2, 280), title, fill='navy', font=title_font)
                name_text = "This is to certify that"
                name_bbox = draw.textbbox((0, 0), name_text, font=name_font)
                name_width = name_bbox[2] - name_bbox[0]
                draw.text(((width - name_width) / 2, 400), name_text, fill='black', font=name_font)
                name_bbox = draw.textbbox((0, 0), user_name, font=name_font)
                name_width = name_bbox[2] - name_bbox[0]
                name_x = (width - name_width) / 2
                name_y = 480
                draw.text((name_x, name_y), user_name, fill='navy', font=name_font)
                line_y = name_y + 60
                for i in range(3):
                    draw.line([(name_x - 50 + i, line_y + i), 
                              (name_x + name_width + 50 - i, line_y + i)], 
                             fill='gold', width=1)
                completion_text = "has successfully completed the course in"
                comp_bbox = draw.textbbox((0, 0), completion_text, font=body_font)
                comp_width = comp_bbox[2] - comp_bbox[0]
                draw.text(((width - comp_width) / 2, 580), completion_text, fill='black', font=body_font)
                topic_bbox = draw.textbbox((0, 0), self.current_topic, font=name_font)
                topic_width = topic_bbox[2] - topic_bbox[0]
                topic_x = (width - topic_width) / 2
                topic_y = 650
                line_length = topic_width + 200
                line_x = (width - line_length) / 2
                for i in range(3):
                    draw.line([(line_x + i, topic_y - 10 + i), 
                              (line_x + line_length - i, topic_y - 10 + i)], 
                             fill='gold', width=1)
                    draw.line([(line_x + i, topic_y + 70 + i), 
                              (line_x + line_length - i, topic_y + 70 + i)], 
                             fill='gold', width=1)
                draw.text((topic_x + 2, topic_y + 2), self.current_topic, fill='#1a1a1a', font=name_font)  # Shadow
                draw.text((topic_x, topic_y), self.current_topic, fill='navy', font=name_font)
                dot_radius = 3
                for i in range(2):
                    draw.ellipse([line_x - dot_radius + i*line_length, topic_y - 12, 
                                line_x + dot_radius + i*line_length, topic_y - 8], 
                               fill='gold')
                    draw.ellipse([line_x - dot_radius + i*line_length, topic_y + 68, 
                                line_x + dot_radius + i*line_length, topic_y + 72], 
                               fill='gold')
                current_time = datetime.datetime.now()
                date_text = f"Awarded on {current_time.strftime('%B %d, %Y')}"
                date_bbox = draw.textbbox((0, 0), date_text, font=body_font)
                date_width = date_bbox[2] - date_bbox[0]
                date_x = (width - date_width) / 2
                date_y = 800
                draw.text((date_x + 2, date_y + 2), date_text, fill='gray', font=body_font)
                draw.text((date_x, date_y), date_text, fill='black', font=body_font)
                cert_id = f"Certificate ID: {current_time.strftime('%Y%m%d%H%M%S')}"
                verify_text = "A.N.I.S Learning Platform - Verified Certificate"
                for i in range(3):
                    color = (100 + i * 20, 100 + i * 20, 100 + i * 20)
                    draw.text((inner_border + 10 - i, height - inner_border - 40 - i), 
                            cert_id, fill=color, font=small_font)
                    draw.text((width - inner_border - 350 - i, height - inner_border - 40 - i), 
                            verify_text, fill=color, font=small_font)
                certificate_filename = f"certificate_{self.current_topic.replace(' ', '_')}_{current_time.strftime('%Y%m%d_%H%M%S')}.png"
                certificate_path = os.path.join(certificates_dir, certificate_filename)
                certificate.save(certificate_path)
                messagebox.showinfo("Certificate Generated", 
                                  f"Professional Certificate has been generated and saved to:\n{certificate_path}")
                os.startfile(certificate_path)
            def on_enter(event=None):
                create_certificate()
            generate_button = tk.Button(
                name_dialog,
                text="Generate Professional Certificate",
                command=create_certificate,
                bg=self.themes[self.current_theme]["accent"],
                fg=self.themes[self.current_theme]["fg"],
                font=("Arial", 11)
            )
            generate_button.pack(pady=20)
            name_entry.bind("<Return>", on_enter)
        except Exception as e:
            print(f"Error generating certificate: {str(e)}")
            messagebox.showerror("Error", "Failed to generate certificate. Please try again.")
    def get_time_date(self):
        """Get current time and date information"""
        try:
            now = datetime.datetime.now()
            date_str = now.strftime("%B %d, %Y")
            time_str = now.strftime("%I:%M %p")
            day_str = now.strftime("%A")
            self.speak(f"Today is {day_str}, {date_str}. The current time is {time_str}")
        except Exception as e:
            self.handle_errors(e, "time and date")
            self.speak("Sorry, I couldn't get the time and date information")   
    def read_highlighted_text(self):
        """Read highlighted text with visual tracking and controls"""
        try:
            self.speak("Sure sir")
            pyautogui.hotkey('ctrl', 'c')
            time.sleep(0.2)
            highlighted_text = pyperclip.paste().strip()
            if not highlighted_text:
                self.speak("No text highlighted. Please select some text first.")
                return
            self.reading_window = ctk.CTkToplevel(self.root)
            self.reading_window.title("Text Reader")
            self.reading_window.geometry("400x550")
            self.reading_window.attributes('-topmost', True)
            self.reading_window.configure(fg_color=self.themes[self.current_theme]["bg"])
            self.reading_active = True
            self.reading_paused = False
            self.current_segment_index = 0
            self.reading_thread = None
            self._process_text_with_gemini(highlighted_text)
            main_frame = ctk.CTkFrame(
                self.reading_window,
                fg_color=self.themes[self.current_theme]["bg"],
                corner_radius=10
            )
            main_frame.pack(fill="both", expand=True, padx=15, pady=15)
            title_label = ctk.CTkLabel(
                main_frame,
                text="Text Reader",
                font=ctk.CTkFont(size=18, weight="bold"),
                text_color=self.themes[self.current_theme]["accent"]
            )
            title_label.pack(pady=(10, 5))
            self.reading_status = ctk.CTkLabel(
                main_frame,
                text="● Reading",
                font=ctk.CTkFont(size=14),
                text_color=self.themes[self.current_theme]["success"]
            )
            self.reading_status.pack(pady=(0, 10))
            progress_frame = ctk.CTkFrame(
                main_frame,
                fg_color=self.themes[self.current_theme]["highlight"],
                corner_radius=8
            )
            progress_frame.pack(fill="x", padx=10, pady=10)
            self.progress_count = ctk.CTkLabel(
                progress_frame,
                text=f"Segment: 1/{len(self.text_segments)}",
                font=ctk.CTkFont(size=12),
                text_color=self.themes[self.current_theme]["secondary"]
            )
            self.progress_count.pack(pady=(5, 0))
            self.progress_bar = ctk.CTkProgressBar(
                progress_frame,
                width=250,
                height=10,
                progress_color=self.themes[self.current_theme]["accent"],
                fg_color=self.themes[self.current_theme]["text_bg"]
            )
            self.progress_bar.pack(pady=(5, 5), padx=10)
            self.progress_bar.set(1/len(self.text_segments) if self.text_segments else 0)
            self.preview_box = ctk.CTkTextbox(
                main_frame,
                width=250,
                height=150,
                font=ctk.CTkFont(family="Segoe UI", size=12),
                fg_color=self.themes[self.current_theme]["text_bg"],
                text_color=self.themes[self.current_theme]["fg"],
                corner_radius=8,
                border_width=0
            )
            self.preview_box.pack(fill="both", expand=True, padx=10, pady=10)
            self.preview_box.insert("1.0", highlighted_text)
            self.preview_box.configure(state="disabled")
            controls_frame = ctk.CTkFrame(
                main_frame,
                fg_color="transparent"
            )
            controls_frame.pack(fill="x", padx=10, pady=(10, 5))
            def stop_reading():
                self.reading_active = False
                engine.stop()
                self.reading_window.destroy()
            def prev_segment():
                if self.current_segment_index > 0:
                    engine.stop()
                    self.current_segment_index -= 1
                    self._highlight_current_segment()
                    if not self.reading_paused and self.reading_active:
                        threading.Thread(target=self._read_current_segment, daemon=True).start()
            def next_segment():
                if self.current_segment_index < len(self.text_segments) - 1:
                    engine.stop()
                    self.current_segment_index += 1
                    self._highlight_current_segment()
                    if not self.reading_paused and self.reading_active:
                        threading.Thread(target=self._read_current_segment, daemon=True).start()
            button_font = ctk.CTkFont(size=13, weight="bold")
            nav_frame = ctk.CTkFrame(
                main_frame,
                fg_color="transparent"
            )
            nav_frame.pack(fill="x", padx=10, pady=(0, 5))
            prev_button = ctk.CTkButton(
                nav_frame,
                text="◀ Prev",
                command=prev_segment,
                font=button_font,
                fg_color=self.themes[self.current_theme]["highlight"],
                hover_color=self.themes[self.current_theme]["secondary"],
                corner_radius=8,
                width=80,
                height=30
            )
            prev_button.pack(side="left", padx=(0, 5), fill="x", expand=True)
            next_button = ctk.CTkButton(
                nav_frame,
                text="Next ▶",
                command=next_segment,
                font=button_font,
                fg_color=self.themes[self.current_theme]["highlight"],
                hover_color=self.themes[self.current_theme]["secondary"],
                corner_radius=8,
                width=80,
                height=30
            )
            next_button.pack(side="right", padx=(5, 0), fill="x", expand=True)
            speed_frame = ctk.CTkFrame(
                main_frame,
                fg_color=self.themes[self.current_theme]["highlight"],
                corner_radius=8
            )
            speed_frame.pack(fill="x", padx=10, pady=(10, 15))
            speed_label = ctk.CTkLabel(
                speed_frame,
                text="Reading Speed",
                font=ctk.CTkFont(size=12),
                text_color=self.themes[self.current_theme]["fg"]
            )
            speed_label.pack(pady=(5, 0))
            self.speed_value = ctk.DoubleVar(value=1.0)
            def update_speed(value):
                self.speed_value.set(float(value))
                speed_display.configure(text=f"{float(value):.1f}x")
                if hasattr(self, 'reading_active') and self.reading_active:
                    engine.setProperty('rate', int(190 * self.speed_value.get()))
            speed_slider = ctk.CTkSlider(
                speed_frame,
                from_=0.5,
                to=2.0,
                number_of_steps=15,
                variable=self.speed_value,
                command=update_speed,
                width=230,
                height=15,
                button_color=self.themes[self.current_theme]["accent"],
                button_hover_color=self.themes[self.current_theme]["secondary"],
                progress_color=self.themes[self.current_theme]["accent"]
            )
            speed_slider.pack(pady=10)
            speed_display = ctk.CTkLabel(
                speed_frame,
                text="1.0x",
                font=ctk.CTkFont(size=12, weight="bold"),
                text_color=self.themes[self.current_theme]["secondary"]
            )
            speed_display.pack(pady=(0, 5))
            self.reading_window.bind("<Left>", lambda e: prev_segment())
            self.reading_window.bind("<Right>", lambda e: next_segment())
            self.reading_window.protocol("WM_DELETE_WINDOW", self.close_reading_window)
            self._highlight_current_segment()
            self.root.after(100, lambda: self._start_reading_automatically())
        except Exception as e:
            self.handle_errors(e, "Error starting text reader")
    def close_reading_window(self):
        """Properly close the reading window and clean up resources"""
        try:
            self.reading_active = False
            self.reading_paused = False
            if hasattr(self, 'engine') and self.engine:
                try:
                    engine.stop()
                except:
                    pass
            if hasattr(self, 'reading_thread') and self.reading_thread is not None and self.reading_thread.is_alive():
                self.reading_thread.join(timeout=0.1)
            self.reading_thread = None
            self.current_segment_index = 0
            if hasattr(self, 'reading_window') and self.reading_window is not None:
                try:
                    if self.reading_window.winfo_exists():
                        self.reading_window.destroy()
                    self.reading_window = None
                except Exception as e:
                    logging.error(f"Error destroying reading window: {str(e)}")
        except Exception as e:
            self.handle_errors(e, "closing reading window")
            try:
                if hasattr(self, 'reading_window') and self.reading_window is not None and self.reading_window.winfo_exists():
                    self.reading_window.destroy()
                    self.reading_window = None
            except:
                pass
    def _start_reading_automatically(self):
        """Start reading text automatically after window is loaded"""
        try:
            if hasattr(self, 'reading_active') and self.reading_active and not self.reading_paused:
                if hasattr(self, 'reading_thread') and self.reading_thread is not None and self.reading_thread.is_alive():
                    self.reading_thread.join(timeout=0.1)
                self.reading_thread = threading.Thread(target=self._read_text_thread, daemon=True)
                self.reading_thread.start()
        except Exception as e:
            logging.error(f"Error auto-starting reading: {str(e)}")
            self.root.after(1000, self.close_reading_window)
    def _process_text_with_gemini(self, text):
        """Process text into intelligent reading segments using Gemini API"""
        try:
            if len(text) < 100:
                self.process_text_for_reading(text)
                return
            prompt = f"""
            Process the following text into natural reading segments.
            Guidelines:
            1. Break at natural sentence boundaries
            2. Maintain coherent thought units
            3. Consider pauses at punctuation
            4. Keep segments at a reasonable length (not too short or too long)
            5. Preserve meaning and context
            Return only the segmented text with each segment on a new line. No explanation or additional text.
            Text: {text}
            """
            try:
                response = self.gemini_model.generate_content(prompt)
                segmented_text = response.text.strip()
                segments = [s.strip() for s in segmented_text.split('\n') if s.strip()]
                if segments and len(segments) > 2:
                    self.text_segments = segments
                else:
                    self.process_text_for_reading(text)
            except Exception as e:
                logging.error(f"Error using Gemini API: {str(e)}")
                self.process_text_for_reading(text)
        except Exception as e:
            logging.error(f"Error in Gemini text processing: {str(e)}")
            self.process_text_for_reading(text)
    def process_text_for_reading(self, text):
        """Process text into natural reading segments (fallback method)"""
        segments = []
        sentences = re.split(r'(?<=[.!?])\s+', text)
        for sentence in sentences:
            if len(sentence) > 80:
                parts = re.split(r'(?<=[:;,])\s+', sentence)
                segments.extend(parts)
            else:
                segments.append(sentence)
        self.text_segments = [seg.strip() for seg in segments if seg.strip()]
    def _highlight_current_segment(self):
        """Highlight the current text segment"""
        if not hasattr(self, 'preview_box') or not self.preview_box.winfo_exists():
            return
        try:
            if not hasattr(self, 'text_segments') or not self.text_segments or self.current_segment_index >= len(self.text_segments):
                return
            current_segment = self.text_segments[self.current_segment_index]
            if hasattr(self, 'progress_count') and self.progress_count.winfo_exists():
                self.progress_count.configure(text=f"Segment: {self.current_segment_index + 1}/{len(self.text_segments)}")
                self.progress_bar.set((self.current_segment_index + 1) / len(self.text_segments))
            self.preview_box.configure(state="normal")
            self.preview_box._textbox.tag_remove("highlight", "1.0", "end")
            start_pos = "1.0"
            while True:
                segment_to_find = re.escape(current_segment)
                start_pos = self.preview_box.search(current_segment, start_pos, stopindex="end", regexp=False)
                if not start_pos:
                    break
                end_pos = f"{start_pos}+{len(current_segment)}c"
                self.preview_box._textbox.tag_add("highlight", start_pos, end_pos)
                self.preview_box._textbox.tag_config("highlight", background=self.themes[self.current_theme]["accent"], foreground="#ffffff")
                self.preview_box.see(start_pos)
                break
            self.preview_box.configure(state="disabled")
        except Exception as e:
            logging.error(f"Error highlighting segment: {str(e)}")
    def _read_current_segment(self):
        """Read only the current segment"""
        try:
            if not self.reading_active or self.reading_paused:
                return
            if self.current_segment_index >= len(self.text_segments):
                return
            current_segment = self.text_segments[self.current_segment_index]
            if hasattr(self, 'reading_status') and self.reading_status.winfo_exists():
                self.reading_status.configure(text="● Reading", text_color=self.themes[self.current_theme]["success"])
            engine.setProperty('rate', int(190 * self.speed_value.get()))
            engine.say(current_segment)
            engine.runAndWait()
        except Exception as e:
            logging.error(f"Error reading current segment: {str(e)}")
    def _read_text_thread(self):
        """Thread for reading text with visual tracking"""
        try:
            if self.current_segment_index >= len(self.text_segments):
                self.current_segment_index = 0
            if hasattr(self, 'reading_status') and self.reading_status.winfo_exists():
                self.reading_status.configure(text="● Reading", text_color=self.themes[self.current_theme]["success"])
            while self.reading_active and self.current_segment_index < len(self.text_segments):
                if self.reading_paused:
                    time.sleep(0.1)
                    continue
                self._highlight_current_segment()
                current_segment = self.text_segments[self.current_segment_index]
                engine.setProperty('rate', int(190 * self.speed_value.get()))
                engine.say(current_segment)
                engine.runAndWait()
                if not self.reading_active or self.reading_paused:
                    break
                self.current_segment_index += 1
                pause_time = 0.2 / self.speed_value.get()
                start_time = time.time()
                while time.time() - start_time < pause_time:
                    if not self.reading_active or self.reading_paused:
                        break
                    time.sleep(0.05)
            if self.reading_active and not self.reading_paused and self.current_segment_index >= len(self.text_segments):
                if hasattr(self, 'reading_status') and self.reading_status.winfo_exists():
                    self.reading_status.configure(text="● Complete", text_color=self.themes[self.current_theme]["success"])
                self.root.after(3000, self.close_reading_window)
                self.current_segment_index = 0
                self._highlight_current_segment()
        except Exception as e:
            logging.error(f"Error in reading thread: {str(e)}")
            if hasattr(self, 'reading_status') and self.reading_status.winfo_exists():
                self.reading_status.configure(text="● Error", text_color=self.themes[self.current_theme]["error"])
            self.root.after(2000, self.close_reading_window)
    def create_specialized_chatbot(self, specialization):
        """Create and open a specialized chatbot window"""
        try:
            SpecializedChatWindow(self.root, specialization)
            self.log_interaction(f"Create specialized chatbot for {specialization}", "Opened specialized chat window")
        except Exception as e:
            self.handle_errors(e, "Error creating specialized chatbot")
            self.speak("Sorry, I couldn't create the specialized chatbot. Please try again later.")
    def load_setting(self, setting_name):
        """Load setting from database"""
        try:
            conn = sqlite3.connect("settings.db")
            cursor = conn.cursor()
            cursor.execute('''CREATE TABLE IF NOT EXISTS settings
                            (name TEXT PRIMARY KEY, value TEXT)''')
            cursor.execute('SELECT value FROM settings WHERE name = ?', (setting_name,))
            result = cursor.fetchone()
            conn.close()
            return result[0] if result else None
        except Exception as e:
            logging.error(f"Error loading setting: {str(e)}")
            return None
    def save_setting(self, setting_name, value):
        """Save setting to database"""
        try:
            conn = sqlite3.connect("settings.db")
            cursor = conn.cursor()
            cursor.execute('''CREATE TABLE IF NOT EXISTS settings
                            (name TEXT PRIMARY KEY, value TEXT)''')
            cursor.execute('''INSERT OR REPLACE INTO settings (name, value)
                            VALUES (?, ?)''', (setting_name, value))
            conn.commit()
            conn.close()
            return True
        except Exception as e:
            logging.error(f"Error saving setting: {str(e)}")
            return False
    def save_gemini_settings(self):
        """Save Gemini API settings to database and update current configuration"""
        try:
            api_key = self.api_key_entry.get().strip()
            model_name = self.model_entry.get().strip()
            if not api_key:
                messagebox.showerror("Error", "API key cannot be empty")
                return
            if not model_name:
                messagebox.showerror("Error", "Model name cannot be empty")
                return
            self.save_setting('gemini_api_key', api_key)
            self.save_setting('gemini_model', model_name)
            try:
                genai.configure(api_key=api_key)
                self.gemini_model = genai.GenerativeModel(model_name)
                self.chat = self.gemini_model.start_chat(history=[])
                messagebox.showinfo("Success", "Gemini settings saved and applied successfully!")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to configure Gemini with provided settings: {str(e)}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save settings: {str(e)}")
    def save_email_settings(self):
        """Save email settings to database"""
        try:
            email = self.email_entry.get().strip()
            password = self.password_entry.get()
            smtp_server = self.smtp_entry.get().strip()
            smtp_port = self.port_entry.get().strip()
            if not all([email, password, smtp_server, smtp_port]):
                messagebox.showerror("Error", "All email settings fields are required")
                return
            try:
                smtp_port = int(smtp_port)
            except ValueError:
                messagebox.showerror("Error", "SMTP port must be a number")
                return
            self.save_setting('email_address', email)
            self.save_setting('email_password', password)
            self.save_setting('smtp_server', smtp_server)
            self.save_setting('smtp_port', str(smtp_port))
            messagebox.showinfo("Success", "Email settings saved successfully!")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save email settings: {str(e)}")
    def handle_email_command(self, highlighted_text=None):
        """Handle sending email commands with voice"""
        try:
            recipient_email = None
            if not highlighted_text:
                self.speak("No text is highlighted. Please highlight an email address and try again.")
                self.speak("Would you like to proceed without a highlighted email?")
                response = self.take_command()
                if not any(word in response.lower() for word in ["yes", "yeah", "continue", "proceed", "okay"]):
                    self.speak("Email command canceled. Please highlight an email and try again.")
                    return False
            if highlighted_text:
                try:
                    api_key = self.load_setting('gemini_api_key')
                    genai.configure(api_key=api_key)
                    prompt = f"""
                    Extract just the email address from this text: "{highlighted_text}"
                    If there's no valid email, respond with "INVALID".
                    Respond with ONLY the email address or "INVALID", nothing else.
                    """
                    model = genai.GenerativeModel(self.load_setting("gemini_model"))
                    response = model.generate_content(prompt)
                    extracted_email = response.text.strip()
                    if extracted_email != "INVALID" and '@' in extracted_email and '.' in extracted_email:
                        recipient_email = extracted_email
                        self.speak(f"I found the email {recipient_email}")
                        confirmation = "yes"
                except Exception as e:
                    logging.error(f"Error extracting email with AI: {str(e)}")
                    highlighted_text = None
            self.speak("What should be the subject of the email?")
            subject = self.take_command()
            if subject.lower() == 'none':
                subject = "No subject"
            self.speak("What message would you like to send?")
            max_attempts = 3
            attempts = 0
            message = None
            while attempts < max_attempts:
                message_input = self.take_command()
                if message_input.lower() == 'none':
                    self.speak(f"I didn't hear you. Attempt {attempts+1} of {max_attempts}.")
                    attempts += 1
                    continue
                if len(message_input) > 10:
                    message = message_input
                    break
                else:
                    self.speak("That message seems too short. Please provide a more detailed message.")
                    attempts += 1
            if not message:
                self.speak("I couldn't understand your message after multiple attempts. Please try again later.")
                return False
            self.speak(f"I'll send an email to {recipient_email} with subject '{subject}")
            confirmation ="yes"
            if any(word in confirmation.lower() for word in ["yes", "yeah", "send", "correct", "proceed"]):
                try:
                    email_sender = EmailSender()
                    email = self.load_setting('email_address')
                    password = self.load_setting('email_password')
                    smtp_server = self.load_setting('smtp_server')
                    smtp_port = self.load_setting('smtp_port')
                    if all([email, password, smtp_server, smtp_port]):
                        try:
                            smtp_port = int(smtp_port)
                            email_sender.email_config = {
                                'smtp_server': smtp_server,
                                'smtp_port': smtp_port,
                                'email': email,
                                'password': password
                            }
                        except:
                            pass
                    self.speak("Sending email now...")
                    success, message_result = email_sender.send_email(recipient_email, subject, message)
                    if success:
                        self.speak("Email sent successfully!")
                        return True
                    else:
                        self.speak(f"Failed to send email: {message_result}")
                        self.speak("open the email window to try manually?")
                        manual_response = "yes"
                        return False
                except Exception as e:
                    logging.error(f"Error sending email directly: {str(e)}")
                    self.speak(f"Error sending email: {str(e)}")
                    self.speak("Would you like to open the email window instead?")
                    window_response = self.take_command()
                    if any(word in window_response.lower() for word in ["yes", "yeah", "open", "window"]):
                        self.root.after(0, lambda: self.show_email_window(recipient_email, subject, message))
                    return False
            else:
                self.speak("Email sending canceled.")
                return False
        except Exception as e:
            logging.error(f"Error handling email command: {str(e)}")
            self.speak("I encountered an error trying to process your email request.")
            return False
    def show_email_window(self, to_email=None, subject=None, message=None):
        """Show email window with optional pre-filled information"""
        try:
            email_window = EmailWindow(self.root)
            if to_email:
                email_window.to_entry.delete(0, END)
                email_window.to_entry.insert(0, to_email)
            if subject:
                email_window.subject_entry.delete(0, END)
                email_window.subject_entry.insert(0, subject)
            if message:
                email_window.message_text.delete("1.0", END)
                email_window.message_text.insert("1.0", message)
            email = self.load_setting('email_address')
            password = self.load_setting('email_password')
            smtp_server = self.load_setting('smtp_server')
            smtp_port = self.load_setting('smtp_port')
            if all([email, password, smtp_server, smtp_port]):
                try:
                    smtp_port = int(smtp_port)
                    email_window.email_sender.email_config = {
                        'smtp_server': smtp_server,
                        'smtp_port': smtp_port,
                        'email': email,
                        'password': password
                    }
                except:
                    pass
        except Exception as e:
            logging.error(f"Error showing email window: {str(e)}")
            self.speak("I encountered an error trying to open the email window.")
    def get_highlighted_text(self):
        """Get currently highlighted text in the system"""
        try:
            clipboard_text = ""
            self.root.clipboard_clear()
            self.root.clipboard_append('')
            self.root.update()
            pyautogui.hotkey('ctrl', 'c')
            time.sleep(0.1)
            try:
                clipboard_text = self.root.clipboard_get()
            except:
                clipboard_text = ""
            return clipboard_text if clipboard_text else None
        except Exception as e:
            logging.error(f"Error getting highlighted text: {str(e)}")
            return None
    def process_command(self, command):
        """Process voice commands"""
        if not isinstance(command, str):
            logging.warning(f"process_command received non-string input: {type(command)}. Skipping.")
            return False
        original_command_for_log = command 
        try:
            commands = re.split(r'\s+\b(and|then)\b\s+|\s*,\s*', command, flags=re.IGNORECASE)
            commands = [cmd.strip().lower() for cmd in commands if cmd and cmd.strip()]
            if len(commands) > 1:
                logging.info(f"Processing command sequence: {commands}")
                for sub_command in commands:
                    if sub_command:
                        logging.info(f"Executing sequence part: '{sub_command}'")
                        self.process_command(sub_command)
                return True
            elif not commands:
                    logging.warning(f"Command resulted in empty list after splitting: '{original_command_for_log}'")
                    return False
            else:
                    command = commands[0]
            ambiguous_words = ["query", "command", "request", "search", "google", "youtube", "and", "then"]
            if command in ambiguous_words:
                    logging.info(f"Ignoring ambiguous/split keyword command: '{command}'")
                    return False
            play_sound2()
            previous_context = self._analyze_conversation_history()
            logging.info(f"Previous context: {previous_context}")
            ambiguous_words = ["query", "command", "request", "search", "google", "youtube"]
            if command.lower().strip() in ambiguous_words:
                logging.info(f"Ignoring ambiguous single-word command: '{command}'")
                return False
            conversation_starters = [
                "hello", "hi", "hey", "greetings", "what's up", "how are you", 
                "good morning", "good afternoon", "good evening", "nice to meet", 
                "thank you", "thanks", "appreciate", "great job", "well done"
            ]
            question_indicators = [
                "what is", "who is", "when is", "where is", "why is", "how is",
                "can you tell me about", "explain", "describe", "define", "what are"
            ]
            is_casual_conversation = any(command.lower().startswith(starter) for starter in conversation_starters)
            is_knowledge_question = any(command.lower().startswith(indicator) for indicator in question_indicators)
            if is_casual_conversation or is_knowledge_question:
                reference_words = ["that", "it", "this", "those", "them", "these"]
                contains_reference = any(word in command.lower().split() for word in reference_words)
                if contains_reference and self.last_topic:
                    for ref in reference_words:
                        command = re.sub(fr'\b{ref}\b', self.last_topic, command, flags=re.IGNORECASE)
                system_prompt = """
                You are A.N.I.S, a friendly and helpful assistant.
                Respond to the user's message in a natural, conversational way.
                Your response should be:
                - Direct and concise (1-2 sentences)
                - Friendly without being overly casual
                - Helpful and accurate
                - Free of special formatting characters
                If the message is a greeting, respond naturally but briefly.
                If the message is a question, provide a factual, direct answer.
                Don't use phrases like "I can" or "I will" - just answer directly.
                don't use the phrase "I will search for" just use this "search google" or "search youtube" for that topic
                """
                user_message = command
                if contains_reference and self.conversation_history:
                    recent_context = "\nRecent conversation context:\n"
                    for i, interaction in enumerate(self.conversation_history[-3:]):
                        recent_context += f"User: {interaction['user']}\n"
                        recent_context += f"Assistant: {interaction['assistant']}\n"
                    system_prompt += f"\n{recent_context}"
                response = gemini_model.generate_content([system_prompt, user_message])
                response_text = self.clean_response(response.text.strip())
                self.log_interaction(command, response_text)
                if command.lower().startswith("what is "):
                    potential_topic = command[8:].strip("?").strip()
                    if potential_topic and len(potential_topic.split()) <= 5:
                        self.last_topic = potential_topic
                        logging.info(f"Setting topic to: {potential_topic}")
                self.speak(response_text)
                return False
            complex_pattern = re.compile(r'search\s+(?:this|that)\s+in\s+(google|youtube)\s+instead\s+of\s+(.*?)\s+in\s+(.*?)($|\s)', re.IGNORECASE)
            match = complex_pattern.search(command.lower())
            if match and hasattr(self, 'last_topic') and self.last_topic:
                platform = match.group(1).lower()
                old_lang = match.group(2).lower()
                new_lang = match.group(3).lower()
                modified_topic = self.last_topic
                if old_lang in modified_topic.lower():
                    modified_topic = modified_topic.lower().replace(old_lang, new_lang)
                else:
                    modified_topic = modified_topic.replace(f"in {old_lang}", f"in {new_lang}")
                    if f"in {old_lang}" not in modified_topic and f"in {new_lang}" not in modified_topic:
                        modified_topic = f"{modified_topic} in {new_lang}"
                self.last_topic = modified_topic
                if platform == "google":
                    self.speak(f"Searching Google for {modified_topic}")
                    google_url = f'https://www.google.com/search?q={urllib.parse.quote(modified_topic)}'
                    webbrowser.open(google_url)
                    return False
                else:
                    self.speak(f"Searching YouTube for {modified_topic}")
                    youtube_url = f"https://www.youtube.com/results?search_query={urllib.parse.quote(modified_topic)}"
                    webbrowser.open(youtube_url)
                    return False
            command_intent = self.analyze_command_with_gemini(command)
            if command_intent and command_intent != command:
                logging.info(f"Command reinterpreted: '{command}' → '{command_intent}'")
                command = command_intent.lower().strip()
            else:
                command = command.lower().strip()
            if any(phrase in command.lower() for phrase in ["activate training mode", "start training mode", "training mode on"]):
                self.activate_training_mode()
                return True
            if any(phrase in command.lower() for phrase in ["deactivate training mode", "stop training mode", "training mode off"]):
                self.deactivate_training_mode()
                return True
            if "it's a" in command.lower() or "this is a" in command.lower() or "this is the" in command.lower():
                self.process_training_command(command)
                return True
            if any(phrase in command.lower() for phrase in ["live assistant", "open live assistant", "launch live assistant"]):
                self.speak("Opening Live Assistant.")
                self.show_live_assistant_window()
                return True
            click_keywords = ["click", "press", "tap", "select"]
            for keyword in click_keywords:
                if command.lower().startswith(f"{keyword} "):
                    element = command[len(keyword)+1:].strip()
                    if element:
                        self.process_click_request(element)
                        return True
            screen_analysis_phrases = [
                "analyze screen", "analyze the screen", "what do you see", 
                "look at screen", "describe screen", "what's on screen"
            ]
            if any(phrase in command.lower() for phrase in screen_analysis_phrases):
                self.analyze_screen()
                return True
            if command in ["google", "youtube", "search google", "search youtube"] and hasattr(self, 'last_topic') and self.last_topic:
                if command.startswith("google") or command == "search google":
                    self.speak(f"Searching Google for {self.last_topic}")
                    google_url = f'https://www.google.com/search?q={urllib.parse.quote(self.last_topic)}'
                    webbrowser.open(google_url)
                    return False
                elif command.startswith("youtube") or command == "search youtube":
                    self.speak(f"Searching YouTube for {self.last_topic}")
                    youtube_url = f"https://www.youtube.com/results?search_query={urllib.parse.quote(self.last_topic)}"
                    webbrowser.open(youtube_url)
                    return False
            if command.startswith("search for "):
                search_query = command.replace("search for ", "").strip()
                self.speak(f"Searching Google for {search_query}")
                google_url = f'https://www.google.com/search?q={urllib.parse.quote(search_query)}'
                webbrowser.open(google_url)
                return False
            if command.startswith("search youtube for "):
                search_query = command.replace("search youtube for ", "").strip()
                self.speak(f"Searching YouTube for {search_query}")
                youtube_url = f"https://www.youtube.com/results?search_query={urllib.parse.quote(search_query)}"
                webbrowser.open(youtube_url)
                return False
            if "search results for" in command.lower():
                parts = command.lower().split("search results for")
                if len(parts) > 1 and parts[1].strip():
                    search_query = parts[1].strip().strip("'").strip('"').strip(".")
                    self.speak(f"Searching Google for {search_query}")
                    google_url = f'https://www.google.com/search?q={urllib.parse.quote(search_query)}'
                    webbrowser.open(google_url)
                    return False
            reference_words = ["that", "it", "this", "those", "them", "these"]
            contains_reference = any(word in command.split() for word in reference_words)
            if contains_reference and hasattr(self, 'last_topic') and self.last_topic:
                logging.info(f"Command contains reference to: '{self.last_topic}'")
                learning_phrases = ["i want to learn", "learn", "teach me", "show me how to"]
                level_terms = [""]
                if any(phrase in command for phrase in learning_phrases):
                    has_level = any(level in command for level in level_terms)
                    base_topic = self.last_topic
                    type_terms = ["types", "categories", "kinds", "examples"]
                    for type_term in type_terms:
                        if type_term in base_topic:
                            parts = base_topic.split(type_term)
                            if len(parts) > 1 and parts[1].strip().startswith("of "):
                                base_topic = parts[1].strip()[3:].strip()
                                break
                    if has_level:
                        specified_level = ""
                        for level in level_terms:
                            if level in command:
                                specified_level = level
                                break
                        search_query = f"{specified_level} level {base_topic}"
                    else:
                        search_query = base_topic
                    self.last_topic = base_topic
                    self.study_mode = True
                    self.current_topic = base_topic
                    level_qualifier = specified_level if has_level else "beginner"
                    if hasattr(self, 'update_command_flow'):
                        self.update_command_flow(
                            task_type="learning", 
                            subject=base_topic,
                            command_type="study",
                            qualifiers={"level": level_qualifier, "mode": "study"}
                        )
                    self.speak(f"I'll help you learn {search_query}")
                    google_url = f'https://www.google.com/search?q={urllib.parse.quote(search_query)}'
                    webbrowser.open(google_url)
                    return False
                direct_question_starters = ["what", "how", "why", "when", "where", "who", "can you", "could you", "will"]
                is_direct_question = any(command.startswith(starter) for starter in direct_question_starters)
                if is_direct_question:
                    context = self.get_current_command_flow() if hasattr(self, 'get_current_command_flow') else ""
                    specific_query = command
                    for ref in reference_words:
                        specific_query = specific_query.replace(f" {ref} ", f" {self.last_topic} ")
                    if any(term in specific_query for term in ["types", "kinds", "categories", "different", "various"]):
                        answer = self.get_specific_topic_information(self.last_topic, specific_query)
                    else:
                        answer = self.get_specific_topic_information(self.last_topic, specific_query)
                    self.speak(answer)
                    return False
                if "youtube" in command and any(ref in command for ref in reference_words):
                    self.speak(f"Searching YouTube for {self.last_topic}")
                    youtube_url = f"https://www.youtube.com/results?search_query={urllib.parse.quote(self.last_topic)}"
                    webbrowser.open(youtube_url)
                    return False
                level_terms = [""]
                has_level_term = any(level in command for level in level_terms)
                show_pattern = any(pattern in command for pattern in ["show me", "show", "display", "present"])
                topic_related_terms = ["types", "concepts", "examples", "applications", "uses", 
                                     "similar", "category", "kind", "variations", "alternatives"]
                has_topic_term = any(term in command for term in topic_related_terms)
                if (has_level_term or has_topic_term) and self.last_topic:
                    logging.info(f"Detected potential fragment reference to: '{self.last_topic}'")
                    level = ""
                    topic_term = ""
                    for term in level_terms:
                        if term in command:
                            level = term
                            break
                    for term in topic_related_terms:
                        if term in command:
                            topic_term = term
                            break
                    base_topic = self.last_topic
                    type_terms = ["types", "categories", "kinds", "examples"]
                    for type_term in type_terms:
                        if type_term in base_topic:
                            parts = base_topic.split(type_term)
                            if len(parts) > 1 and parts[1].strip().startswith("of "):
                                base_topic = parts[1].strip()[3:].strip()
                                break
                    previous_context = self._analyze_conversation_history()
                    if not level and previous_context["preferred_level"]:
                        level = previous_context["preferred_level"]
                        logging.info(f"Using level from context: {level}")
                    if not topic_term and previous_context["preferred_content_type"]:
                        topic_term = previous_context["preferred_content_type"]
                        logging.info(f"Using content type from context: {topic_term}")
                    if level and topic_term:
                        search_query = f"{level} {topic_term} of {base_topic}"
                    elif level and not topic_term:
                        search_query = f"{level} level {base_topic}"
                    elif topic_term and not level:
                        search_query = f"{topic_term} of {base_topic}"
                    else:
                        search_query = base_topic
                    video_indicators = ["youtube", "video", "videos", "watch"]
                    explicit_video_request = any(indicator in command.lower() for indicator in video_indicators)
                    use_youtube = explicit_video_request or previous_context["search_platform"] == "youtube"
                    if use_youtube:
                        self.speak(f"Searching YouTube for {search_query}")
                        youtube_url = f"https://www.youtube.com/results?search_query={urllib.parse.quote(search_query)}"
                        webbrowser.open(youtube_url)
                        self.log_interaction(command, f"Searching YouTube for {search_query}")
                    else:
                        self.speak(f"Searching Google for {search_query}")
                        google_url = f'https://www.google.com/search?q={urllib.parse.quote(search_query)}'
                        webbrowser.open(google_url)
                        self.log_interaction(command, f"Searching Google for {search_query}")
                    if topic_term and "type" in topic_term:
                        self.last_topic = f"{topic_term} of {base_topic}"
                    else:
                        self.last_topic = base_topic
                    if hasattr(self, 'update_command_flow'):
                        self.update_command_flow(
                            task_type="search" if not previous_context["has_learning_context"] else "learning",
                            subject=base_topic,
                            command_type="query",
                            qualifiers={
                                "level": level if level else "general",
                                "type": topic_term if topic_term else "general",
                                "platform": "youtube" if use_youtube else "google"
                            }
                        )
                    return False
                if has_level_term and show_pattern and "of" in command:
                    level = ""
                    for term in level_terms:
                        if term in command:
                            level = term
                            break
                    base_topic = self.last_topic
                    type_terms = ["types", "categories", "kinds", "examples"]
                    for type_term in type_terms:
                        if type_term in base_topic:
                            parts = base_topic.split(type_term)
                            if len(parts) > 1 and parts[1].strip().startswith("of "):
                                base_topic = parts[1].strip()[3:].strip()
                                break
                    search_query = f"{level} level {base_topic}"
                    if "video" in command or "youtube" in command:
                        self.speak(f"Searching YouTube for {search_query}")
                        youtube_url = f"https://www.youtube.com/results?search_query={urllib.parse.quote(search_query)}"
                        webbrowser.open(youtube_url)
                    else:
                        self.speak(f"Searching Google for {search_query}")
                        google_url = f'https://www.google.com/search?q={urllib.parse.quote(search_query)}'
                        webbrowser.open(google_url)
                    self.last_topic = base_topic
                    return False
                search_phrases = ["search for", "search", "look up", "find", "show me"]
                if any(phrase in command for phrase in search_phrases):
                    if "youtube" in command or "video" in command:
                        video_type_terms = ["types", "categories", "kinds", "examples", "tutorials", "lectures"]
                        level_terms = ["advanced", "beginner", "intermediate", "basic", "expert", "professional"]
                        has_type_term = any(term in command for term in video_type_terms)
                        has_level_term = any(term in command for term in level_terms)
                        search_query = self.last_topic
                        if has_level_term:
                            for term in level_terms:
                                if term in command:
                                    base_topic = self.last_topic
                                    for type_term in video_type_terms:
                                        if type_term in base_topic:
                                            parts = base_topic.split(type_term)
                                            if len(parts) > 1:
                                                base_topic = parts[1].strip()
                                                if base_topic.startswith("of "):
                                                    base_topic = base_topic[3:].strip()
                                            break
                                    search_query = f"{term} level {base_topic}"
                                    break
                        elif has_type_term:
                            for term in video_type_terms:
                                if term in command:
                                    search_query = f"{term} of {self.last_topic}"
                                    break
                        self.speak(f"Searching YouTube for {search_query}")
                        youtube_url = f"https://www.youtube.com/results?search_query={urllib.parse.quote(search_query)}"
                        webbrowser.open(youtube_url)
                        return False
                    elif "google" in command or "web" in command or "internet" in command:
                        self.speak(f"Searching Google for {self.last_topic}")
                        google_url = f'https://www.google.com/search?q={urllib.parse.quote(self.last_topic)}'
                        webbrowser.open(google_url)
                        return False
                info_request_indicators = ["tell me about", "give me information", "explain", "more about", 
                                         "what is", "how does", "key points", "summary", "give some", 
                                         "provide", "details", "explain"]
                if any(indicator in command for indicator in info_request_indicators) and hasattr(self, 'get_information_for_topic'):
                    context = self.get_current_command_flow() if hasattr(self, 'get_current_command_flow') else ""
                    information = self.get_information_for_topic(self.last_topic, context)
                    self.speak(information)
                    return False
                type_indicators = ["what are the types", "what types", "what kinds", "what categories", 
                                  "different types", "different kinds", "categories of", "types of", 
                                  "examples of", "applications of", "uses of", "aspects of"]
                if any(indicator in command for indicator in type_indicators) and hasattr(self, 'get_information_for_topic'):
                    specific_query = f"What are the main types or categories of {self.last_topic}?"
                    types_info = self.get_specific_topic_information(self.last_topic, specific_query)
                    self.speak(types_info)
                    return False
                expanded_cmd = command
                for ref in reference_words:
                    if ref in expanded_cmd.split():
                        expanded_cmd = expanded_cmd.replace(ref, self.last_topic)
                logging.info(f"Expanded command: '{expanded_cmd}'")
                command = expanded_cmd
            if  'comments' in command or 'in python' in command or 'coding' in command:
                command = command.replace('code', '').strip()
                command = command.replace('coding','').strip()
                command = command.replace('comments', '').strip()
                command = command.replace('in python', '').strip()
                self.show_program_output_code(command)
                return False
            elif 'read' in command and ('text' in command or 'this' in command or 'highlighted' in command):
                self.read_highlighted_text()
                return False
            elif any(phrase in command for phrase in ['file assistant', 'file handling', 'handle file', 'file prompt']):
                self.speak("Opening file handling assistant")
                self.root.after(0, self.file_assistant.show_window)
                return False
            elif 'browse file' in command:
                self.speak("Opening file browser")
                self.root.after(0, lambda: self.browse_and_process_file(command))
                return False
            elif any(phrase in command for phrase in ['show yesterday files', 'show yesterday file', 'yesterday file', 'recent files']):
                self.speak("Showing recent files")
                self.root.after(0, self.show_recent_files)
                return False
            elif "auto install" in command or 'install' in command:
                app_name = command.replace("auto install", "").replace("install","").replace("how to","").strip()
                if app_name:
                    self.speak(f"Attempting to install {app_name}")
                    self.root.after(0, lambda: self.auto_install_app(app_name))
                else:
                    self.speak("Please specify the app name to install.")
                return False 
            elif "create a chatbot" in command :
                match = re.search(r"create a chatbot \[(.*?)\]", command)
                if not match:
                    match = re.search(r"create a chatbot (.*)", command)
                if match:
                    specialization = match.group(1).replace('for', '').replace('create a','').replace('create a chatbot','').strip()
                    self.speak(f"Creating a specialized chatbot for {specialization}.")
                    self.root.after(0, lambda s=specialization: self.create_specialized_chatbot(s))
                else:
                    self.speak("Please specify a topic, for example: create a chatbot specially for data science")
                return False
            elif "create a chat box" in command:
                match = re.search(r"create a chat box \[(.*?)\]", command)
                if not match:
                    match = re.search(r"create a chat box (.*)", command)
                if match:
                    specialization = match.group(1).strip()
                    self.speak(f"Creating a specialized chatbot for {specialization}.")
                    self.root.after(0, lambda s=specialization: self.create_specialized_chatbot(s))
                else:
                    self.speak("Please specify a topic, for example: create a chatbot specially for data science")
                return False
            elif any(cmd in command for cmd in ['analyse news','analyse text','analyse the news','analyse the highlighted','analyse highlights','analyse highlight','analyse the highlight','check news','check text','check link','analyse link']):
                highlighted_text = pyautogui.hotkey('ctrl', 'c')
                time.sleep(0.1)
                text_to_analyze = pyperclip.paste()
                if text_to_analyze:
                    self.speak("Analyzing the selected content")
                    self.analyze_news(text_to_analyze)
                else:
                    self.speak("Please highlight some text or a link to analyze")
                return False
            elif 'study mode on' in command:
                self.study_mode = True
                self.speak("Study mode activated. What would you like to learn?")
                topic = self.take_command().lower()
                if topic != 'none':
                    self.current_topic = topic
                    self.speak(f"Great! Let's start learning about {topic}")
                    search_query = f"{topic} course tutorial"
                    google_url = f'https://www.google.com/search?q={urllib.parse.quote(search_query)}'
                    search_url = f"https://www.youtube.com/results?search_query={urllib.parse.quote(search_query)}"
                    html = urllib.request.urlopen(search_url)
                    video_ids = re.findall(r"watch\?v=(\S{11})", html.read().decode())
                    if video_ids:
                            first_video_url = f"https://www.youtube.com/watch?v={video_ids[0]}"
                    else:
                            self.speak("No videos found for your search")
                    webbrowser.open(google_url)
                    webbrowser.open(first_video_url)
                    pyautogui.hotkey('enter')
                    self.speak(f"Playing video for {search_query} on YouTube")
                    self.root.after(0, self.show_chats_window)
                    if topic not in self.study_progress:
                        self.study_progress[topic] = {
                            'level': 1,
                            'completed_tasks': 0,
                            'quiz_scores': []
                        }
                return False
            elif 'play music' in command or 'play random music' in command:
                    try:
                        music_dir = r'C:\Users\HP\Music'
                        songs = os.listdir(music_dir)
                        if songs:
                            random_song = np.random.choice(songs)
                            os.startfile(os.path.join(music_dir, random_song))
                            self.speak(f'Playing {random_song}')
                        else:
                            self.speak("No music files found in the directory.")
                    except Exception as e:
                        self.speak(f"Failed to play random music: {e}")    
            elif 'study mode off' in command:
                if self.study_mode:
                    self.study_mode = False
                    self.current_topic = None
                    self.speak("Study mode deactivated. Great job with your studies!")
                else:
                    self.speak("Study mode is not currently active.")
                return False
            elif 'quiz' in command:
                self.show_quiz_window()
            elif 'next' in command:
                current_level = self.study_progress[self.current_topic]['level']
                self.study_progress[self.current_topic]['level'] = current_level + 1
                self.study_progress[self.current_topic]['completed_tasks'] += 1
                self.speak(f"Moving to level {current_level + 1} of {self.current_topic}")
                search_query = f"{self.current_topic} advanced tutorial level {current_level + 1}"
                google_url = f'https://www.google.com/search?q={urllib.parse.quote(search_query)}'
                youtube_url = f"https://www.youtube.com/results?search_query={urllib.parse.quote(search_query)}"
                webbrowser.open(google_url)
                webbrowser.open(youtube_url)
            elif 'complete task' in command and self.study_mode:
                if self.current_topic:
                    try:
                        self.speak("Would you like to take a quiz, get certification, or move to the next level?")
                        choice = self.take_command()
                        if choice is None:
                            self.speak("I didn't catch that. Please try again.")
                            return False
                        choice = choice.lower()
                        if 'quiz' in choice:
                            self.show_quiz_window()
                        elif 'certification' in choice or 'certificate' in choice:
                            self.speak("Starting certification exam for " + self.current_topic)
                            score = self.start_certification_exam()
                            if score >= 60:
                                self.speak("Congratulations! You passed the certification exam.")
                                self.generate_certificate()
                            else:
                                self.speak(f"You scored {score}%. You need 60% to pass. Please try again after more practice.")
                        elif 'next' in choice:
                            current_level = self.study_progress[self.current_topic]['level']
                            self.study_progress[self.current_topic]['level'] = current_level + 1
                            self.study_progress[self.current_topic]['completed_tasks'] += 1
                            self.speak(f"Moving to level {current_level + 1} of {self.current_topic}")
                            search_query = f"{self.current_topic} advanced tutorial level {current_level + 1}"
                            google_url = f'https://www.google.com/search?q={urllib.parse.quote(search_query)}'
                            youtube_url = f"https://www.youtube.com/results?search_query={urllib.parse.quote(search_query)}"
                            webbrowser.open(google_url)
                            webbrowser.open(youtube_url)
                        else:
                            self.speak("Please say either 'quiz', 'certification', or 'next' to proceed.")
                    except Exception as e:
                        self.handle_errors(e, "Error processing complete task command")
                else:
                    self.speak("No topic is currently selected in study mode.")
                return False
            elif any(phrase in command for phrase in ['what do you see', 'analyze screen', 'describe screen', 'what is on screen', 'what is on the screen']):
                self.speak("Ok Sir")
                description = self.analyze_screen()
                self.speak(description)
                return False
            elif 'reset conversation' in command:
                if hasattr(self, 'chat_window'):
                    self.chat_window.destroy()
                self.root.after(0, self.show_chats_window)
                self.speak("Conversation has been reset")
                return False
            elif 'analyse database' in command or 'analyze database' in command or 'database analysis' in command:
                self.speak("Opening database analysis window")
                self.root.after(0, self.analyze_database)
                return False
            elif 'activate settings' in command or 'open settings' in command:
                setting_type = None
                if 'activate settings' in command:
                    setting_part = command.split('activate settings', 1)[1].strip()
                    if setting_part:
                        setting_type = setting_part
                elif 'open settings' in command:
                    setting_part = command.split('open settings', 1)[1].strip()
                    if setting_part:
                        setting_type = setting_part
                if setting_type:
                    success, message = self.access_windows_settings(setting_type)
                    if success:
                        self.speak(message)
                        return False
                    else:
                        self.speak(f"{message}. Would you like to try another setting?")
                        response = self.take_command()
                        if any(word in response.lower() for word in ['yes', 'yeah', 'sure', 'okay']):
                            self.speak("What settings would you like to access?")
                        else:
                            self.speak("Okay, let me know if you need anything else.")
                            return False
                else:
                    self.speak("What settings would you like to access?")
                max_attempts = 3
                attempts = 0
                while attempts < max_attempts:
                    if not setting_type:
                        setting_type = self.take_command()
                    if setting_type.lower() != 'none':
                        success, message = self.access_windows_settings(setting_type)
                        if success:
                            self.speak(message)
                            return False
                        else:
                            self.speak(f"{message}. Please try again with a different settings name.")
                            setting_type = None
                            attempts += 1
                    else:
                        self.speak(f"I didn't hear you. Attempt {attempts+1} of {max_attempts}.")
                        setting_type = None
                        attempts += 1
                if attempts >= max_attempts:
                    self.speak("Maximum attempts reached. You can try again by saying 'activate settings'.")
                return False
            elif any(cmd in command for cmd in ['open google']):
                            webbrowser.open('https://www.google.com')
                            self.speak("Opening Google")
                            return False
            elif 'search for' in command:
                search_query = command.replace('search for', '').strip()
                if search_query:
                    level_terms = [""]
                    has_level = any(level in search_query for level in level_terms)
                    if has_level and "of" in search_query:
                        for level in level_terms:
                            if level in search_query:
                                parts = search_query.split("of", 1)
                                if len(parts) > 1:
                                    core_topic = parts[1].strip()
                                    level_part = parts[0].strip()
                                    self.last_topic = core_topic
                                    logging.info(f"Setting last_topic to core: '{core_topic}' from level search")
                    else:
                        self.last_topic = search_query
                        logging.info(f"Setting last_topic to: '{search_query}' from search command")
                    search_url = f'https://www.google.com/search?q={urllib.parse.quote(search_query)}'
                    webbrowser.open(search_url)
                    self.speak(f"Searching Google for {search_query}")
                    return False
            elif 'search youtube' in command or 'search youtube for' in command:
                search_query = command.replace('search youtube', '').replace('for', '').strip()
                if search_query:
                    youtube_url = f"https://www.youtube.com/results?search_query={urllib.parse.quote(search_query)}"
                    webbrowser.open(youtube_url)
                    self.speak(f"Searching YouTube for {search_query}")
                else:
                    self.speak("Please specify what to search for on YouTube")
                return False
            elif 'youtube' in command and any(term in command for term in ['about', 'related to', 'videos on', 'watch']):
                topic = None
                for keyword in ['about', 'related to', 'videos on', 'watch']:
                    if keyword in command:
                        parts = command.split(keyword, 1)
                        if len(parts) > 1:
                            topic = parts[1].strip()
                            break
                if topic:
                    youtube_url = f"https://www.youtube.com/results?search_query={urllib.parse.quote(topic)}"
                    webbrowser.open(youtube_url)
                    self.speak(f"Searching YouTube for videos about {topic}")
                else:
                    webbrowser.open('https://www.youtube.com')
                    self.speak("Opening YouTube")
                return False
            elif 'open youtube' in command:
                webbrowser.open('https://www.youtube.com')
                self.speak("Opening YouTube")
                return False
            elif 'open wikipedia' in command:
                webbrowser.open('https://www.wikipedia.org')
                self.speak("Opening Wikipedia")
                return False
            elif command.startswith('open '):
                app_name = command.replace('open ', '').strip()
                if app_name:
                    self.open_application(app_name)
                    return False
                else:
                    self.speak("Please specify which application to open")
                    return False
            elif any(cmd in command for cmd in ['programming window','programming assistant','programming']):
                self.speak("Opening the programming assistant")
                try:
                    import google.generativeai as genai
                    self.root.after(0, self.show_programming_window)
                except ImportError:
                    self.speak("Please install the required package google-generativeai first")
                    logging.error("google-generativeai package not installed")
                return False
            elif any(exit_cmd in command for exit_cmd in ['exit','you can leave','quit', 'goodbye', 'bye','see you later','see you again']):
                self.speak("Thank you sir,see you soon")
                self.exit_assistant()
                return True
            elif 'brightness' in command:
                self.control_brightness(command)
                return False
            elif 'volume' in command:
                self.control_volume(command)
                return False
            elif 'battery' in command:
                self.get_battery_info()
                return False
            elif 'storage' in command or 'disk space' in command:
                self.get_storage_info()
                return False
            elif any(word in command for word in ['the time', 'the date', 'the day']):
                self.get_time_date()
                return False
            elif command.startswith('type '):
                self.type_and_enter(command)
                return False
            elif command in ['stop typing', 'end typing', 'exit typing']:
                self.stop_typing()
                return False
            elif 'contact manager' in command or 'show contacts' in command:
                self.speak("Opening contact manager")
                self.root.after(0, self.show_contact_manager)
                return False
            elif 'tab' in command:
                tab_number = None
                if 'first' in command or '1st' in command or 'one' in command:
                    tab_number = 1
                elif 'second' in command or '2nd' in command or 'two' in command:
                    tab_number = 2
                elif 'third' in command or '3rd' in command or 'three' in command:
                    tab_number = 3
                elif 'fourth' in command or '4th' in command or 'four' in command:
                    tab_number = 4
                else:
                    numbers = [int(s) for s in command.split() if s.isdigit()]
                    if numbers:
                        tab_number = numbers[0]
                if tab_number is not None:
                    self.switch_tab(tab_number - 1)
                    self.speak(f"Switched to tab {tab_number}")
                    return False
            elif 'send a message to' in command:
                contact_name = command.replace('send a message to', '').strip()
                if contact_name:
                    self.whatsapp_interaction('message', contact_name)
                    return True
            elif 'video call' in command:
                contact_name = command.replace('video call', '').strip()
                if contact_name:
                    self.whatsapp_interaction('video', contact_name)
                    return True
            elif 'whatsapp call' in command:
                contact_name = command.replace('whatsapp call', '').strip()
                if contact_name:
                    self.whatsapp_interaction('call', contact_name)
                    return True
            elif "auto login" in command:
                website = command.replace("auto login", "").strip()
                if website:
                    if self.is_valid_website(website):
                        self.auto_login(website)
                    else:
                        self.speak(f"I cannot log into {website}. Please specify a valid website.")
                else:
                    self.speak("Please specify the website for auto login.")
                return False
            elif "password manager" in command:
                self.show_password_manager()
                return False
            elif command in ['delete text','delete the text']:
                pyautogui.hotkey('delete')
                return False
            is_scheduled_task = False
            if command.startswith("__scheduled__:"):
                is_scheduled_task = True
                command = command.replace("__scheduled__:", "", 1).strip()
                logging.info(f"Processing scheduled task command: '{command}'")
            if is_scheduled_task:
                return self._execute_direct_command(command)
            elif not is_scheduled_task and 'schedule' in command and ('task' in command or 'reminder' in command or '' in command):
                max_retries = 3
                retry_count = 0
                while retry_count < max_retries:
                    self.speak("How long until the tasks should start?")
                    duration = self.take_command()
                    if duration and duration.lower() != 'none':
                        seconds = self.parse_duration(duration)
                        if seconds > 0:
                            self.speak("What tasks would you like to schedule?.")
                            tasks = self.take_command()
                            if tasks and tasks.lower() != 'none':
                                task_list = [t.strip() for t in tasks.split(',') if t.strip()]
                                if task_list:
                                    self.show_task_scheduler(preset_time=seconds, preset_tasks=','.join(task_list))
                                    return False
                            self.show_task_scheduler(preset_time=seconds)
                            return False
                        else:
                            retry_count += 1
                            if retry_count < max_retries:
                                self.speak("I couldn't understand the time duration. Please try saying it like '5 minutes' or '2 hours' or '30 seconds'.")
                            else:
                                self.speak("I'm still having trouble understanding the time. Would you like to enter it manually in the scheduler window?")
                                response = self.take_command()
                                if any(word in response.lower() for word in ['yes', 'yeah', 'sure', 'okay']):
                                    self.show_task_scheduler()
                                else:
                                    self.speak("Task scheduling cancelled.")
                                return False
                    else:
                        retry_count += 1
                        if retry_count < max_retries:
                            self.speak("I didn't catch that. Please try again with the duration.")
                        else:
                            self.speak("I'm having trouble understanding. Would you like to open the scheduler window to enter the time manually?")
                            response = self.take_command()
                            if any(word in response.lower() for word in ['yes', 'yeah', 'sure', 'okay']):
                                self.show_task_scheduler()
                            else:
                                self.speak("Task scheduling cancelled.")
                            return False
                return False
            elif 'new text file' in command:
                pyautogui.hotkey('ctrl', 'n')
                self.speak("Done")
                return False
            elif 'new text window' in command:
                pyautogui.hotkey('ctrl', 'shift', 'n')
                self.speak("Done")
                return False
            elif 'open file' in command:
                pyautogui.hotkey('ctrl', 'o')
                self.speak("Done")
                return False
            elif 'data analysis' in command:
                self.speak("Here You can analyse the Data")
                self.root.after(0, self.show_data_analysis_window)
                return False
            elif 'open folder' in command:
                pyautogui.hotkey('ctrl', 'shift', 'o')
                self.speak("Done")
                return False
            elif 'translate' in command:
                target_lang = None
                for lang_name in self.language_dict.keys():
                    if f"to {lang_name.lower()}" in command:
                        target_lang = lang_name
                        break
                self.show_translation_comparison(target_lang)
                return False
            elif 'undo' in command:
                pyautogui.hotkey('ctrl', 'z')
                self.speak("Done")
                return False
            elif 'select all' in command:
                pyautogui.hotkey('ctrl', 'a')
                self.speak("Done")
                return False
            elif any(phrase in command for phrase in ['program for highlighted', 'create a program']):
                self.speak("Creating a program")
                self.show_program_output()
                return False 
            elif 'run window' in command:
                pyautogui.hotkey('win', 'r')
                self.speak("Done")
                return False
            elif 'run' in command:
                pyautogui.hotkey('win', 'r')
                self.speak("Done")
                return False
            elif 'saved' in command:
                pyautogui.hotkey('ctrl', 's')
                self.speak("Done")
                return False
            elif 'save as' in command:
                pyautogui.hotkey('ctrl', 'shift', 's')
                self.speak("Done")
                return False
            elif 'minimize' in command:
                pyautogui.hotkey('win', 'm')
                self.speak("Done")
                return False
            elif 'minimize all' in command:
                pyautogui.hotkey('win', 'm')
                self.speak("Done")
                return False
            elif 'maximize' in command:
                pyautogui.hotkey('win', 'shift', 'm')
                self.speak("Done")
                return False
            elif 'maximize all' in command:
                pyautogui.hotkey('win', 'shift', 'm')
                self.speak("Done")
                return False
            elif 'screenshot' in command:
                self.take_screenshot()
                return False
            elif 'copy' in command:
                pyautogui.hotkey('ctrl', 'c')
                self.speak("Done")
                return False
            elif 'copy all' in command:
                pyautogui.hotkey('ctrl', 'a')
                pyautogui.hotkey('ctrl', 'c')
                self.speak("Done")
                return False
            elif 'paste' in command:
                pyautogui.hotkey('ctrl', 'v')
                self.speak("Done")
                return False
            elif 'cut' in command:
                pyautogui.hotkey('ctrl', 'x')
                self.speak("Done")
                return False
            elif 'find' in command:
                pyautogui.hotkey('ctrl', 'f')
                self.speak("Done")
                return False
            elif 'close window' in command:
                pyautogui.hotkey('alt', 'f4')
                self.speak("Done")
                return False
            elif 'close' in command:
                pyautogui.hotkey('alt', 'f4')
                self.speak("Done")
                return False
            elif 'close all' in command:
                pyautogui.hotkey('alt', 'f4')
                self.speak("Done")
                return False
            elif 'close all windows' in command:
                pyautogui.hotkey('alt', 'f4')
                self.speak("Done")
                return False
            elif 'delete' in command:
                pyautogui.hotkey('delete')
                self.speak("Done")
                return False
            elif 'enter' in command:
                pyautogui.hotkey('enter')
                self.speak("Done")
                return False
            elif command in ['check properties','check property','check property of']:
                pyautogui.hotkey('alt', 'enter')
                self.speak("Done")
                return False
            elif any(phrase in command for phrase in ['maximize all', 'maximize the all window', 'maximize all windows', 'maximize window', 'maximize the window', 'restore all window', 'restore window', 'restore all']):
                pyautogui.hotkey('win', 'shift', 'm')
                self.speak("All windows maximized")
                return False
            elif any(phrase in command for phrase in ['save notes in file','save note', 'save this note', 'save the note', 'save notes']):
                try:
                    notes_dir = r"C:\Users\HP\\Desktop\Notes"
                    now = datetime.datetime.now()
                    week_num = now.isocalendar()[1]
                    weekday = now.strftime('%A')
                    week_dir = os.path.join(notes_dir, f"Week_{week_num}")
                    day_dir = os.path.join(week_dir, weekday)
                    os.makedirs(day_dir, exist_ok=True)
                    active_window = pyautogui.getActiveWindow()
                    if "Notepad" in str(active_window.title):
                        pyautogui.hotkey('ctrl', 'a')
                        time.sleep(0.2)
                        pyautogui.hotkey('ctrl', 'c')
                        time.sleep(0.2)
                        import win32clipboard
                        win32clipboard.OpenClipboard()
                        try:
                            text = win32clipboard.GetClipboardData()
                        except:
                            text = ""
                        finally:
                            win32clipboard.CloseClipboard()
                        if text.strip():
                            filename = self.generate_filename_from_content(text)
                            filename = self.sanitize_filename(filename)
                            filepath = os.path.join(day_dir, filename)
                            with open(filepath, 'w', encoding='utf-8') as f:
                                f.write(text)
                            self.speak(f"Notes saved as {filename} in {weekday} folder for Week {week_num}")
                        else:
                            self.speak("No content to save")
                    else:
                        self.speak("Please make sure Notepad is active")
                except Exception as e:
                    logging.error(f"Error saving notes: {str(e)}")
                    self.speak("Sorry, I couldn't save the notes")
                return False
            elif 'switch window' in command:
                pyautogui.hotkey('alt', 'tab')
                self.speak("Done")
                return False
            elif 'minimise all' in command or 'minimise the all window' in command or 'minimise the window' in command or 'minimise window' in command:
                pyautogui.hotkey('win', 'm')
                self.speak("Done")
                return False
            elif 'task manager' in command:
                pyautogui.hotkey('ctrl', 'shift', 'esc')
                self.speak("Done")
                return False
            elif 'lock screen' in command:
                pyautogui.hotkey('win', 'l')
                self.speak("Done")
                return False
            elif 'search youtube' in command or 'youtube search' in command or 'youtube' in command:
                search_query = command.replace('search youtube', '').replace('youtube search', '').replace('youtube', '').strip()
                if search_query:
                    try:
                        search_url = f"https://www.youtube.com/results?search_query={urllib.parse.quote(search_query)}"
                        html = urllib.request.urlopen(search_url)
                        video_ids = re.findall(r"watch\?v=(\S{11})", html.read().decode())
                        if video_ids:
                            first_video_url = f"https://www.youtube.com/watch?v={video_ids[0]}"
                            webbrowser.open(first_video_url)
                            self.speak(f"Playing video for {search_query} on YouTube")
                            pyautogui.hotkey('enter')
                        else:
                            self.speak("No videos found for your search")
                    except Exception as e:
                        logging.error(f"Error searching YouTube: {str(e)}")
                        self.speak("Sorry, I encountered an error while searching YouTube")
                else:
                    self.speak("Please specify what to search on YouTube")
                return False
            elif any(cmd in command for cmd in ['chatbot', 'gemini chat','gemini','chat box' ]):
                self.speak("Opening ANIS chat")
                self.root.after(0, self.show_chat_window)
                return False
            elif 'translate window' in command or 'translator' in command:
                self.speak("Opening translator")
                self.root.after(0, self.show_translator_window)
                return False
            elif 'make notes' in command or 'summary' in command:
                self.speak("Opening summarizer window")
                self.root.after(0, self.show_summarizer_window)
                return False
            elif any(phrase in command for phrase in ['send email', 'send mail', 'compose email', 'write email']):
                highlighted_text = self.get_highlighted_text()
                self.handle_email_command(highlighted_text)
                return False
            elif 'show email' in command or 'show mail' in command:
                self.root.after(0, self.show_email_window)
                return False
            elif 'open notes' in command:
                self.root.after(0, self.show_notes_window)
                return False
            elif 'check system health' in command or 'check system' in command or 'check system status' in command or 'check system information' in command or 'check system details' in command or 'check system status' in command or 'check system information' in command or 'check system details' in command:
                self.speak("Checking system health")
                self.check_system_health()
                return False
            elif any(phrase in command for phrase in ['search and open', 'find and open', 'search for and open']):
                search_term = command
                for phrase in ['search and open', 'find and open', 'search for and open']:
                    search_term = search_term.replace(phrase, '').strip()
                if search_term:
                    self.search_and_click_first_result(search_term)
                else:
                    self.speak("Please specify what to search for")
                return False
            elif any(phrase in command for phrase in ['save notes in file', 'save note', 'save this note', 'save the note', 'save notes']):
                try:
                    notes_dir = r"C:\Users\HP\Desktop\Ai\Notes"
                    now = datetime.datetime.now()
                    week_num = now.isocalendar()[1]
                    weekday = now.strftime('%A')
                    week_dir = os.path.join(notes_dir, f"Week_{week_num}")
                    day_dir = os.path.join(week_dir, weekday)
                    os.makedirs(day_dir, exist_ok=True)
                    timestamp = now.strftime('%Y%m%d_%H%M%S')
                    filename = f"note_{timestamp}.txt"
                    filepath = os.path.join(day_dir, filename)
                    active_window = pyautogui.getActiveWindow()
                    if "Notepad" in active_window.title:
                        pyautogui.hotkey('ctrl', 'a')
                        time.sleep(0.1)
                        pyautogui.hotkey('ctrl', 'c')
                        time.sleep(0.1)
                        text = self.window.clipboard_get()
                        with open(filepath, 'w', encoding='utf-8') as f:
                            f.write(text)
                        self.speak(f"Notes saved in {weekday} folder for Week {week_num}")
                    else:
                        self.speak("Please make sure Notepad is active")
                except Exception as e:
                    logging.error(f"Error saving notes: {str(e)}")
                    self.speak("Sorry, I couldn't save the notes")
                return False
            elif 'open' in command:
                if self.open_system_app(command):
                    return False
                search_term = command.replace('open', '').strip()
                if search_term:
                    self.search_and_click_first_result(search_term)
                return False
            elif 'take screenshot' in command:
                    self.speak("Taking screenshot")
                    self.take_screenshot()
                    return False
            elif any(phrase in command for phrase in ['save highlighted text','save the highlights','save the hightlight','save selection', 'save selected text', 'save highlight', 'save the highlighted words',"save highlights"]):
                try:
                    notes_dir = r"C:\Users\HP\Desktop\Ai\Important Notes"
                    if not os.path.exists(notes_dir):
                        os.makedirs(notes_dir)
                    now = datetime.datetime.now()
                    week_num = now.isocalendar()[1]
                    weekday = now.strftime('%A')
                    week_dir = os.path.join(notes_dir, f"Week_{week_num}")
                    day_dir = os.path.join(week_dir, weekday)
                    os.makedirs(day_dir, exist_ok=True)
                    pyautogui.hotkey('ctrl', 'c')
                    time.sleep(0.5)
                    try:
                        selected_text = self.root.clipboard_get()
                    except:
                        self.speak("No text was selected")
                        return False
                    if not selected_text or selected_text.strip() == "":
                        self.speak("Please select some text first")
                        return False
                    try:
                        import google.generativeai as genai
                        api_key = self.load_setting('gemini_api_key')
                        model_name = self.load_setting("gemini_model")
                        genai.configure(api_key=api_key)
                        model = genai.GenerativeModel(model_name)
                        prompt = f"Generate a short, descriptive filename (2-3 words) for this text: {selected_text[:200]}"
                        response = model.generate_content(prompt)
                        suggested_name = response.text.strip()
                        timestamp = now.strftime('%Y%m%d_%H%M%S')
                        filename = f"{suggested_name}_{timestamp}.txt"
                    except Exception as e:
                        logging.error(f"Error generating filename: {str(e)}")
                        timestamp = now.strftime('%Y%m%d_%H%M%S')
                        filename = f"selected_text_{timestamp}.txt"
                    filename = self.sanitize_filename(filename)
                    filepath = os.path.join(day_dir, filename)
                    with open(filepath, 'w', encoding='utf-8') as f:
                        f.write(f"Date: {now.strftime('%Y-%m-%d %H:%M:%S')}\n")
                        f.write("-" * 50 + "\n\n")
                        f.write(selected_text)
                    if os.path.exists(filepath):
                        self.speak(f"Selected text saved successfully in {weekday} folder")
                        logging.info(f"Saved selected text to: {filepath}")
                    else:
                        self.speak("Failed to save the text")
                except Exception as e:
                    logging.error(f"Error saving selected text: {str(e)}")
                    self.speak("Sorry, I couldn't save the selected text")
                return False
            elif 'google search' in command or 'search google' in command:
                    search_query = command.replace('google search', '').replace('search google', '').replace('google', '').strip()
                    if search_query:
                        self.speak(f"Searching Google for {search_query}")
                        url = f"https://www.google.com/search?q={urllib.parse.quote(search_query)}"
                        webbrowser.open(url)
                    else:
                        self.speak("Opening Google")
                        webbrowser.open("https://www.google.com")
                    return False
            try:
                response = self.get_gemini_action(command)
                if response is None:
                    return False
                elif response and response.strip():
                    self.speak(response)
                    self.log_interaction(command, response)
                    return False
                self.type_and_enter("Processing your request...")
                if custom_prompt:
                    prompt = f"""GuideLines: {custom_prompt}Guidelines:
                - answer should be in two lines for any query
                - Use simple language
                - Avoid special characters or formatting
                - Use proper punctuation
                - Keep responses brief and direct
                -your name is ANIS AI integrated with Operation System With Help Of Gemini and under Developing by 
                Samiulla from 26 August 2024 don't  tell this in beginning of your response you should tell only when someone ask you name
                - No emojis or symbols
                - answer in a friendly manner
                - answer like a human
                - answer in a way that is easy to understand
                - answer in a way that is not too formal
                - answer in a way that is not too technical
                - if need techinical answer in techincal terms
                - answer should be short like 2 lines and concise
                - avoid geographical deep terms
                - try to answer in 2 lines simple and short
                - some of the time you can answer in a funny way
                \n\nUser Query: {command}"""
                else:
                    prompt = f"""Please provide a clear and concise response to: {command}
                    Guidelines:
                - Use simple language
                - Avoid special characters or formatting
                - Use proper punctuation
                - Keep responses brief and direct
                -your name is ANIS AI integrated with Operation System With Help Of Gemini and under Developing by 
                Samiulla from 26 August 2024 don't  tell this in beginning of your response you should tell only when someone ask you name
                - No emojis or symbols
                - answer in a friendly manner
                - answer like a human
                - answer in a way that is easy to understand
                - answer in a way that is not too formal
                - answer in a way that is not too technical
                - if need techinical answer in techincal terms
                - answer should be short like 2 lines and concise
                - avoid geographical deep terms
                - try to answer in 2 lines simple and short
                - some of the time you can answer in a funny way
                """
                response = self.chat.send_message(prompt)
                response_text = response.text.strip()
                response_text = response_text.replace('*', '').replace('#', '').replace('`', '')
                self.stop_typing()
                self.speak(response_text)
                self.log_interaction(command, response_text)
                return False
            except Exception as e:
                self.handle_errors(e, "Error getting Gemini response")
                self.speak("I'm having trouble understanding that right now. Could you please try again?")
                return False
        except Exception as e:
            self.handle_errors(e, "Error processing command")
            self.speak("Sorry, I encountered an error processing that command")
            return False
    def toggle_listening(self):
        """Toggle background listening on and off."""
        try:
            if not self.start_button._listening:
                self.start_button._listening = True
                self.start_button.configure(text="◼")
                self.status_label.configure(text="● Listening", text_color='#4CAF50')
                if hasattr(self, 'wave_vis'): self.wave_vis.start_listening_animation()
                self.stop_listening_func = self.recognizer.listen_in_background(
                    self.microphone,
                    self._audio_callback,
                    phrase_time_limit=9
                )
                logging.info("Background listener started.")
            else:
                self.start_button._listening = False
                if self.stop_listening_func:
                    self.stop_listening_func(wait_for_stop=False)
                    self.stop_listening_func = None
                    logging.info("Background listener stopped.")
                else:
                    logging.warning("Stop listening called, but no stop function found.")
                self.start_button.configure(text="▶")
                self.status_label.configure(text="● Not Listening", text_color='#ff4d4d')
                if hasattr(self, 'wave_vis'): self.wave_vis.stop_listening_animation()
        except Exception as e:
            logging.error(f"Error toggling listening state: {str(e)}")
            self.speak("Sorry, I encountered an error with the voice recognition.")
            self.start_button._listening = False
            if self.stop_listening_func:
                self.stop_listening_func(wait_for_stop=False)
                self.stop_listening_func = None
            self.start_button.configure(text="▶")
            self.status_label.configure(text="● Not Listening", text_color='#ff4d4d')
            if hasattr(self, 'wave_vis'): self.wave_vis.stop_listening_animation()
    def _audio_callback(self, recognizer, audio):
        """
        This function is called by the background thread when audio data is captured.
        """
        if not self.start_button._listening:
            logging.info("Audio captured, but listening is stopped. Ignoring.")
            return
        logging.info("Audio data received by callback. Attempting recognition...")
        self.root.after(0, lambda: self.status_label.configure(text="Recognizing..."))
        query = None
        try:
            query = recognizer.recognize_google(audio, language='en-in')
            logging.info(f"Google Speech Recognition thinks you said: {query}")
            self.root.after(0, lambda q=query: self.insert_with_timestamp(f"You: {q}", "user_msg"))
            self.root.after(0, lambda: self.status_label.configure(text="Processing...", text_color='#FFA500')) # Orange color for processing
            self.process_command(query.lower())
            logging.info(f"Command processing finished for: {query}")
        except sr.UnknownValueError:
            logging.info("Google Speech Recognition could not understand audio")
        except sr.RequestError as e:
            logging.error(f"Could not request results from Google Speech Recognition service; {e}")
            self.root.after(0, lambda: self.speak("Speech service error."))
        except Exception as e:
            logging.error(f"Error during audio callback processing: {e}")
            self.root.after(0, lambda: self.speak("An error occurred while processing audio."))
        finally:
            def reset_status():
                if self.start_button._listening:
                    logging.info("Resetting status to Listening.")
                    self.status_label.configure(text="● Listening", text_color='#4CAF50')
                    if hasattr(self, 'wave_vis'): self.wave_vis.start_listening_animation()
                else:
                    logging.info("Listening stopped during callback, ensuring status is Not Listening.")
                    self.status_label.configure(text="● Not Listening", text_color='#ff4d4d')
                    if hasattr(self, 'wave_vis'): self.wave_vis.stop_listening_animation()
            self.root.after(0, reset_status)
    def start_listening(self):
        """Start the listening process"""
        try:
            self.start_button._listening = True
            self.start_button.configure(text="◼")
            self.start_color_transition(self.start_button, '#27ae60')
            self.start_button.master.configure(highlightbackground='#27ae60')
            self.toggle_listening()
            self.status_label.configure(text="● Listening", text_color='#4CAF50')
        except Exception as e:
            logging.error(f"Error in start_listening: {str(e)}")
            self.status_label.configure(text="● Not Listening", text_color='#ff4d4d')
    def listen_for_commands(self):
        """Continuously listen for commands until paused or exit command is given"""
        logging.warning("listen_for_commands called, but listening is now handled by listen_in_background.")
        pass
    def exit_assistant(self):
        """Safely exit the voice assistant"""
        try:
            self.start_button._listening = False
            if self.stop_listening_func:
                self.stop_listening_func(wait_for_stop=False)
                self.stop_listening_func = None
                logging.info("Background listener stopped on exit.")
            if hasattr(self, '_after_id'):
                self.root.after_cancel(self._after_id)
            if hasattr(self, 'engine'):
                self.engine.stop()
            if hasattr(self, 'wave_vis'):
                self.wave_vis.cleanup()
            if hasattr(self, 'root') and self.root.winfo_exists():
                self.root.quit()
                self.root.destroy()
        except Exception as e:
            logging.error(f"Error during exit: {str(e)}")
            sys.exit(0)
    def execute_powershell_command(self, command):
        """Execute a PowerShell command using subprocess"""
        try:
            result = subprocess.run(
                ["powershell", "-Command", command],
                capture_output=True,
                text=True
            )
            if result.returncode == 0:
                return True, result.stdout
            else:
                logging.error(f"PowerShell error: {result.stderr}")
                return False, result.stderr
        except Exception as e:
            logging.error(f"Error executing PowerShell command: {str(e)}")
            return False, str(e)
    def access_windows_settings(self, setting_type):
        """Access Windows settings via PowerShell based on setting type using AI"""
        try:
            api_key = self.load_setting('gemini_api_key')
            model_name = self.load_setting("gemini_model")
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel(model_name)
            prompt = f"""
            I need to open Windows settings for "{setting_type}" using PowerShell.
            Return a valid PowerShell command using the ms-settings: URI scheme that opens the specific Windows settings page.
            For example, to open display settings, the command would be: Start-Process ms-settings:display
            Respond with ONLY the PowerShell command, nothing else. Use the most specific ms-settings URI for this request.
            If you're not sure about the exact URI, use the most likely match based on common Windows settings pages.
            For complex settings, consider these common URIs:
            - Network: ms-settings:network
            - WiFi: ms-settings:network-wifi
            - Bluetooth: ms-settings:bluetooth
            - Display: ms-settings:display
            - Sound: ms-settings:sound
            - Notifications: ms-settings:notifications
            - Power: ms-settings:batterysaver
            - Storage: ms-settings:storagesense
            - Privacy: ms-settings:privacy
            - Windows Update: ms-settings:windowsupdate
            - Accounts: ms-settings:yourinfo
            """
            response = self.chat.send_message(prompt)
            powershell_command = response.text.strip()
            if not self.is_safe_powershell_command(powershell_command):
                logging.warning(f"Unsafe PowerShell command rejected: {powershell_command}")
                return False, f"I couldn't create a safe command for {setting_type} settings"
            logging.info(f"AI generated PowerShell command: {powershell_command}")
            success, output = self.execute_powershell_command(powershell_command)
            if success:
                return True, f"Opened {setting_type} settings"
            else:
                fallback_prompt = f"""
                The previous command failed to open Windows settings for "{setting_type}".
                Provide an alternative PowerShell command using a more generic ms-settings: URI.
                Remember to respond with ONLY the command, nothing else.
                """
                fallback_response = self.chat.send_message(fallback_prompt)
                fallback_command = fallback_response.text.strip()
                if self.is_safe_powershell_command(fallback_command) and fallback_command != powershell_command:
                    logging.info(f"Trying fallback command: {fallback_command}")
                    success, fallback_output = self.execute_powershell_command(fallback_command)
                    if success:
                        return True, f"Opened {setting_type} settings using alternative method"
                return False, f"Failed to open {setting_type} settings"
        except Exception as e:
            logging.error(f"Error using Gemini to generate settings command: {str(e)}")
            return False, f"I encountered an error trying to open {setting_type} settings"
    def is_safe_powershell_command(self, command):
        """Validate that the PowerShell command is safe to execute"""
        if not command.startswith("Start-Process ms-settings:"):
            return False
        dangerous_patterns = [
            ";",
            "&",
            "|",
            ">",
            "<",
            "$",
            "`",
            "\\",
            "//",
            "http",
            "-Command",
            "-EncodedCommand",
            "-ExecutionPolicy",
            "Invoke-Expression",
            "iex",
            "Invoke-WebRequest",
            "curl",
            "wget",
            "New-Object",
            "Remove",
            "Delete",
        ]
        for pattern in dangerous_patterns:
            if pattern in command:
                return False
        return True
    def process_click_request(self, element_description):
        """Processes a request to click on a specified UI element."""
        try:
            self.speak(f"Looking for {element_description}")
            logging.info(f"Looking for element: {element_description}")
            active_info = get_active_app_info()
            app_name = active_info["app_name"]
            coordinates = self.find_coordinates_in_db(element_description, app_name)
            if coordinates:
                x, y = coordinates
                logging.info(f"Found saved coordinates for '{element_description}': ({x}, {y})")
                self.speak(f"Found {element_description}")
                click_at_coordinates(x, y)
                self.record_coordinate_attempt(app_name, element_description, x, y, True)
                log_manager.log_command(
                    "click",
                    f"Click {element_description}",
                    element_description,
                    True,
                    f"x={x}, y={y}, app={app_name}"
                )
                return True
            else:
                self.speak(f"I don't know where {element_description} is. Please position your mouse over it and press Ctrl+Shift+X")
                return False
        except Exception as e:
            error_msg = f"Error processing click request: {str(e)}"
            logging.error(error_msg)
            self.speak(f"I couldn't click on {element_description} due to an error")
            log_manager.log_command(
                "click",
                f"Click {element_description}",
                element_description,
                False,
                f"Error: {str(e)}"
            )
            return False
    def process_type_request(self, text_to_type):
        """Handles the request to type text."""
        try:
            self.speak(f"Typing {text_to_type}")
            type_text(text_to_type)
            log_manager.log_command("type", f"Type: {text_to_type}", text_to_type, True)
            return True
        except Exception as e:
            error_msg = f"Error typing text: {str(e)}"
            logging.error(error_msg)
            self.speak("Sorry, I couldn't type that")
            log_manager.log_command("type", f"Type failed: {text_to_type}", text_to_type, False, f"Error: {str(e)}")
            return False
    def get_ai_response(self, user_query, pil_image=None):
        """Sends the user query and optional image to Gemini for a response."""
        try:
            model = genai.GenerativeModel('gemini-2.0-flash')
            prompt_parts = []
            base_prompt = (
                "You are a friendly and helpful AI assistant designed for voice interaction. "
                "Analyze the user's voice query and the attached screenshot (if provided). "
                "Respond naturally and conversationally, as if you were talking to the user. "
                "Keep your answers concise and to the point (usually 1-2 sentences), focusing on being helpful. "
                "If asked a question, answer directly. If shown the screen, comment briefly on what you see or offer a relevant suggestion."
            )
            prompt_parts.append(base_prompt)
            if user_query:
                prompt_parts.append(f'\nUser Query: "{user_query}"')
            else:
                prompt_parts.append("\nUser did not provide a specific voice query, focus on the screen.")
            if pil_image:
                prompt_parts.append("\nAnalyze the attached screenshot in conjunction with the query:")
                prompt_parts.append(pil_image)
                has_screenshot = True
            else:
                prompt_parts.append("\n(No screenshot provided for this query)")
                has_screenshot = False
            print("Sending request to Gemini...")
            response = model.generate_content(prompt_parts)
            if response and response.text:
                ai_text = response.text.strip()
                print("Received response from Gemini.")
                log_manager.log_conversation(user_query or "(Screen Analysis Only)", ai_text, has_screenshot)
                return ai_text
            else:
                print("Gemini returned an empty response or failed.")
                try:
                    print(f"Gemini response details: {response.prompt_feedback}")
                except Exception:
                    pass
                return "Error: Failed to get analysis from AI. Check API key and model access."
        except Exception as e:
            print(f"Error interacting with Gemini API: {e}")
            return f"Error during AI analysis: {e}"
    def analyze_screen(self):
        """Captures and analyzes the current screen using AI."""
        try:
            self.speak("Analyzing your screen")
            screenshot = capture_screen_to_pil()
            if not screenshot:
                self.speak("Failed to capture screen")
                return False
            app_info = get_active_app_info()
            app_name = app_info["app_name"]
            self.speak("Processing screen content")
            analysis_prompt = "Analyze what's on this screen and describe the main elements you can see. Focus on identifying UI elements, content, and possible actions."
            ai_response = self.get_ai_response(analysis_prompt, screenshot)
            if ai_response:
                self.speak(ai_response)
                log_manager.log_command(
                    "analyze_screen", 
                    "Screen analysis", 
                    app_name, 
                    True, 
                    "Screen successfully analyzed"
                )
                return True
            else:
                self.speak("The AI couldn't analyze the screen content")
                log_manager.log_command("analyze_screen", "Screen analysis failed", app_name, False, "AI returned empty response")
                return False
        except Exception as e:
            error_msg = f"Error analyzing screen: {str(e)}"
            logging.error(error_msg)
            self.speak("Sorry, I couldn't analyze your screen")
            log_manager.log_command("analyze_screen", "Screen analysis failed", None, False, f"Error: {str(e)}")
            return False
    def activate_training_mode(self):
        """Activate training mode for UI element identification"""
        self.training_mode_active = True
        logging.info("Training mode activated")
        self.speak("Training mode activated.Pattern to train is 'It's a [element name]'")
        log_manager.log_command("training_mode", "Activated training mode", "on", True)
        return True
    def deactivate_training_mode(self):
        """Deactivate training mode"""
        self.training_mode_active = False
        logging.info("Training mode deactivated")
        self.speak("Training mode deactivated")
        log_manager.log_command("training_mode", "Deactivated training mode", "off", True)
        return True
    def process_training_command(self, command):
        """Process a training command to identify UI elements"""
        element_name = None
        if "it's a" in command.lower():
            match = re.search(r"it's a\s+(.+)", command, re.IGNORECASE)
            if match:
                element_name = match.group(1).strip()
        elif "it's" in command.lower():
            match = re.search(r"it's \s+(.+)", command, re.IGNORECASE)
            if match:
                element_name = match.group(1).strip()
        elif "this is a" in command.lower():
            match = re.search(r"this is a\s+(.+)", command, re.IGNORECASE)
            if match:
                element_name = match.group(1).strip()
        elif "this is" in command.lower():
            match = re.search(r"this is \s+(.+)", command, re.IGNORECASE)
            if match:
                element_name = match.group(1).strip()
        elif "this is the" in command.lower():
            match = re.search(r"this is the\s+(.+)", command, re.IGNORECASE)
            if match:
                element_name = match.group(1).strip()
        if element_name:
            self.capture_ui_element(element_name)
        else:
            self.speak("Could not understand the element name. Please try again with 'It's a [element name]'.")
    def capture_ui_element(self, element_name):
        """Capture the current mouse position and store it with the element name"""
        try:
            x, y = pyautogui.position()
            app_info = get_active_app_info()
            app_name = app_info["app_name"]
            save_coordinates(element_name, x, y, True, app_name)
            self.speak(f"Captured {element_name} at position {x}, {y}")
            log_manager.log_command("capture_element", f"Captured element: {element_name}", 
                                  element_name, True, f"Position: ({x}, {y}), App: {app_name}")
            return True
        except Exception as e:
            error_msg = f"Error capturing UI element: {str(e)}"
            logging.error(error_msg)
            self.speak(f"Failed to capture {element_name}")
            log_manager.log_command("capture_element", f"Capture {element_name}", 
                                  element_name, False, f"Error: {str(e)}")
            return False
    def record_coordinate_attempt(self, app_name, element_description, x, y, success):
        """Record a coordinate click attempt, updating success/fail counts."""
        try:
            conn = sqlite3.connect("conversation_log.db")
            cursor = conn.cursor()
            cursor.execute("""
                SELECT id, success_count, fail_count FROM coordinates 
                WHERE app_name = ? AND element_description = ?
            """, (app_name.lower(), element_description.lower()))
            existing = cursor.fetchone()
            if existing:
                entry_id, success_count, fail_count = existing
                if success:
                    success_count += 1
                else:
                    fail_count += 1
                cursor.execute("""
                    UPDATE coordinates
                    SET success_count = ?, fail_count = ?, timestamp = CURRENT_TIMESTAMP
                    WHERE id = ?
                """, (success_count, fail_count, entry_id))
            else:
                cursor.execute("""
                    INSERT INTO coordinates (app_name, element_description, x, y, 
                                            success_count, fail_count)
                    VALUES (?, ?, ?, ?, ?, ?)
                """, (app_name.lower(), element_description.lower(), x, y, 
                    1 if success else 0, 0 if success else 1))
            conn.commit()
            conn.close()
            return True
        except sqlite3.Error as e:
            print(f"Database Error: Failed to record coordinate attempt - {e}")
            return False
    def find_coordinates_in_db(self, element_description, app_name):
        """Find coordinates for an element in the database with flexible matching and screen resolution adaptation."""
        try:
            conn = sqlite3.connect("conversation_log.db")
            cursor = conn.cursor()
            cursor.execute("""
                SELECT x, y FROM coordinates 
                WHERE LOWER(element_description) = ? AND LOWER(app_name) = ?
            """, (element_description.lower(), app_name.lower()))
            result = cursor.fetchone()
            if not result:
                cursor.execute("""
                    SELECT x, y FROM coordinates 
                    WHERE LOWER(element_description) LIKE ? AND LOWER(app_name) = ?
                    ORDER BY success_count DESC, timestamp DESC
                    LIMIT 1
                """, (f"%{element_description.lower()}%", app_name.lower()))
                result = cursor.fetchone()
            if not result:
                cursor.execute("""
                    SELECT x, y FROM coordinates 
                    WHERE LOWER(element_description) LIKE ?
                    ORDER BY success_count DESC, timestamp DESC
                    LIMIT 1
                """, (f"%{element_description.lower()}%",))
                result = cursor.fetchone()
            conn.close()
            if result:
                if not hasattr(self, '_screen_adapter'):
                    self._screen_adapter = ScreenCoordinateAdapter()
                original_x, original_y = result
                adapted_x, adapted_y = self._screen_adapter.adapt_coordinates(original_x, original_y)
                logging.info(f"Adapting coordinates for '{element_description}' from ({original_x}, {original_y}) to ({adapted_x}, {adapted_y})")
                return adapted_x, adapted_y
            return None
        except sqlite3.Error as e:
            logging.error(f"Database Error: Failed to find coordinates - {e}")
            return None
    def display_message(self, sender, message):
        self.output_text.configure(state=NORMAL)
        self.output_text._textbox.insert(END, f"{sender}: {message}\n\n")
        self.output_text.configure(state=DISABLED)
        self.output_text.see(END)
    def log_interaction(self, user_input, assistant_response):
        try:
            self.conversation_history.append({"user": user_input, "assistant": assistant_response})
            if len(self.conversation_history) > self.max_history_length:
                self.conversation_history.pop(0)
            timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            log_entry = (
                f"Time: {timestamp}\n"
                f"User: {user_input}\n"
                f"Assistant: {assistant_response}\n"
                f"{'-'*50}\n"
            )
            with open("interaction_log.txt", "a", encoding="utf-8") as log_file:
                log_file.write(log_entry)
        except Exception as e:
            logging.error(f"Logging error: {str(e)}")
    def handle_errors(self, error, context=""):
        """Generic error handler for the assistant"""
        error_message = f"Error in {context}: {str(error)}"
        logging.error(error_message)
        user_message = "I encountered an error. Please try again."
        self.speak(user_message)
        self.update_status(False)
        self.log_interaction(context, user_message)
    def process_text_command(self, text_input):
        """Process text commands without voice input"""
        try:
            if text_input:
                self.output_text.configure(state=NORMAL)
                self.output_text._textbox.insert(END, f"Text Input: {text_input}\n")
                self.output_text.configure(state=DISABLED)
                self.output_text.see(END)
                return self.process_command(text_input.lower())
            return False
        except Exception as e:
            self.handle_errors(e, "text command processing")
            return False
    def setup_keyboard_shortcuts(self):
        """Setup keyboard shortcuts for the assistant"""
        try:
            keyboard_thread = threading.Thread(
                target=self.handle_keyboard_shortcuts,
                daemon=True
            )
            keyboard_thread.start()
        except Exception as e:
            self.handle_errors(e, "keyboard shortcut setup")
    def speak(self, text, emotion=None):
        try:
            if hasattr(self, 'wave_vis') and self.wave_vis.is_active:
                self.wave_vis.speaking_animation()
            self.insert_with_timestamp(f"A.N.I.S: {text}", "assistant_msg")
            engine.say(text)
            engine.runAndWait()
            if hasattr(self, 'wave_vis') and self.wave_vis.is_active:
                self.wave_vis.stop_listening_animation()
        except Exception as e:
            logging.error(f"Error in speak method: {str(e)}")
    def take_command(self):
        r = sr.Recognizer()
        with sr.Microphone() as source:
            try:
                if hasattr(self, 'wave_vis') and self.wave_vis.is_active:
                    self.wave_vis.start_listening_animation()
                self.status_label.configure(text="Listening...")
                self.root.update()
                r.adjust_for_ambient_noise(source, duration=1)
                r.dynamic_energy_threshold = True
                r.pause_threshold = 0.8
                r.non_speaking_duration = 0.8
                r.phrase_threshold = 0.3
                audio = r.listen(source, timeout=9 ,phrase_time_limit=8)
                self.status_label.configure(text="Recognizing...")
                self.root.update()
                query = r.recognize_google(audio, language='en-in')
                self.insert_with_timestamp(f"You: {query}", "user_msg")
                return query.lower()
            except sr.WaitTimeoutError:
                return "none"
            except sr.UnknownValueError:
                return "none"
            except sr.RequestError:
                self.speak("Sorry, there was an error with the speech recognition service")
                return "none"
            except Exception as e:
                logging.error(f"Error in speech recognition: {str(e)}")
                return "none"
            finally:
                if hasattr(self, 'wave_vis') and self.wave_vis.is_active:
                    self.wave_vis.stop_listening_animation()
                self.status_label.configure(text="Not Listening")
    def update_system_stats(self):
        """Update system statistics"""
        try:
            cpu_percent = psutil.cpu_percent(interval=0.1)
            self.cpu_label.configure(text=f"CPU: {cpu_percent}%")
            memory = psutil.virtual_memory()
            self.memory_label.configure(text=f"RAM: {memory.percent}%")
            self.root.after(1000, self.update_system_stats)
        except Exception as e:
            logging.error(f"Error updating system stats: {str(e)}")
    def show_settings(self):
        """Show settings window"""
        settings_window = ctk.CTkToplevel(self.root)
        settings_window.title("Settings")
        settings_window.geometry("600x700")
        settings_window.configure(fg_color=self.themes[self.current_theme]["bg"])
        settings_window.transient(self.root)
        settings_window.grab_set()
        tabview = ctk.CTkTabview(
            settings_window,
            fg_color=self.themes[self.current_theme]["bg"],
            segmented_button_fg_color=self.themes[self.current_theme]["bg"],
            segmented_button_selected_color=self.themes[self.current_theme]["accent"],
            segmented_button_selected_hover_color=self.themes[self.current_theme]["secondary"],
            segmented_button_unselected_color=self.themes[self.current_theme]["bg"],
            segmented_button_unselected_hover_color=self.themes[self.current_theme]["highlight"]
        )
        tabview.pack(fill='both', expand=True, padx=10, pady=10)
        tab_appearance = tabview.add("Appearance")
        tab_voice = tabview.add("Voice")
        tab_prompts = tabview.add("Prompts")
        tab_commanding = tabview.add("Commanding")
        tab_api = tabview.add("API Settings")
        tab_screen = tabview.add("Screen Interaction")
        title_style = {
            "font": ("Arial", 16, "bold"),
            "text_color": self.themes[self.current_theme].get("text", self.themes[self.current_theme]["fg"])
        }
        desc_style = {
            "font": ("Arial", 12),
            "text_color": self.themes[self.current_theme].get("text_secondary", self.themes[self.current_theme]["secondary"])
        }
        ctk.CTkLabel(
            tab_appearance, 
            text="Appearance Settings",
            **title_style
        ).pack(pady=(20, 5), padx=20, anchor="w")
        ctk.CTkLabel(
            tab_appearance,
            text="Customize the look and feel of the voice assistant",
            **desc_style
        ).pack(pady=(0, 20), padx=20, anchor="w")
        theme_var = tk.StringVar(value=self.current_theme)
        for theme_name in self.themes.keys():
            button_container = ctk.CTkFrame(
                tab_appearance,
                fg_color=self.themes[self.current_theme]["bg"],
                corner_radius=8
            )
            button_container.pack(fill='x', pady=5, padx=10)
            theme_button = ctk.CTkRadioButton(
                button_container,
                text=theme_name,
                variable=theme_var,
                value=theme_name,
                command=lambda t=theme_name: self.change_theme(t),
                fg_color=self.themes[self.current_theme]["accent"],
                border_color=self.themes[self.current_theme]["fg"],
                text_color=self.themes[self.current_theme]["fg"],
                hover_color=self.themes[self.current_theme]["secondary"]
            )
            theme_button.pack(fill='x', padx=15, pady=8)
            if theme_name == self.current_theme:
                theme_button.select()
        ctk.CTkLabel(
            tab_voice, 
            text="Voice Settings",
            **title_style
        ).pack(pady=(20, 5), padx=20, anchor="w")
        ctk.CTkLabel(
            tab_voice,
            text="Configure the voice assistant's speech settings",
            **desc_style
        ).pack(pady=(0, 20), padx=20, anchor="w")
        self.voice_scale = ctk.CTkSlider(
            tab_voice,
            from_=0,
            to=100,
            orientation='horizontal',
            number_of_steps=100,
            command=self.change_voice_volume
        )
        self.voice_scale.set(engine.getProperty('volume') * 100)
        self.voice_scale.pack(fill='x', pady=10)
        volume_labels_frame = ctk.CTkFrame(
            tab_voice,
            fg_color=self.themes[self.current_theme]["bg"],
            corner_radius=0
        )
        volume_labels_frame.pack(fill='x')
        ctk.CTkLabel(
            volume_labels_frame,
            text="0%",
            font=("Segoe UI", 11),
            text_color=self.themes[self.current_theme]["fg"]
        ).pack(side='left')
        ctk.CTkLabel(
            volume_labels_frame,
            text="100%",
            font=("Segoe UI", 11),
            text_color=self.themes[self.current_theme]["fg"]
        ).pack(side='right')
        ctk.CTkLabel(
            tab_prompts, 
            text="Prompt Settings",
            **title_style
        ).pack(pady=(20, 5), padx=20, anchor="w")
        ctk.CTkLabel(
            tab_prompts,
            text="Customize the assistant's responses",
            **desc_style
        ).pack(pady=(0, 20), padx=20, anchor="w")
        prompts_main_frame = ctk.CTkScrollableFrame(
            tab_prompts,
            fg_color=self.themes[self.current_theme]["bg"],
            corner_radius=0
        )
        prompts_main_frame.pack(fill='both', expand=True)
        ctk.CTkLabel(
            prompts_main_frame,
            text="Custom Agent Prompt Guidelines",
            font=("Segoe UI", 16, "bold"),
            text_color=self.themes[self.current_theme]["accent"]
        ).pack(anchor='w', pady=(0, 20))
        prompt_sections = {
            'voice_commands': "Voice Commands Prompt",
            'chat_window': "Chat Window Prompt",
            'article_mode': "Article Mode Prompt",
            'programming': "Programming Assistant Prompt",
            'header_analysis': "Data Headers Analysis Prompt",
            'full_analysis': "Full Dataset Analysis Prompt",
            'commanding_prompt': "Commanding Prompt",
        }
        for key, title in prompt_sections.items():
            section_frame = ctk.CTkFrame(
                prompts_main_frame,
                fg_color=self.themes[self.current_theme]["bg"],
                corner_radius=0
            )
            section_frame.pack(fill='x', pady=10)
            ctk.CTkLabel(
                section_frame,
                text=title,
                font=("Segoe UI", 12),
                text_color=self.themes[self.current_theme]["fg"]
            ).pack(anchor='w')
            placeholder_text = {
                'voice_commands': "{command} - Voice command, {context} - Command context",
                'chat_window': "{message} - User message, {history} - Chat history",
                'article_mode': "{text} - Article content, {mode} - Analysis mode",
                'programming': "{code} - Code snippet, {language} - Programming language",
                'header_analysis': "{headers} - Dataset column names",
                'full_analysis': "{data_summary} - Dataset statistics",
                'commanding_prompt': "{commands} - Command, {contexts} - Command context",
            }
            ctk.CTkLabel(
                section_frame,
                text=placeholder_text[key],
                font=("Segoe UI", 9, "italic"),
                text_color=self.themes[self.current_theme]["secondary"]
            ).pack(anchor='w', pady=(0, 5))
            text_widget = ctk.CTkTextbox(
                section_frame,
                height=100,
                fg_color=self.themes[self.current_theme]["text_bg"],
                text_color=self.themes[self.current_theme]["fg"],
                font=("Segoe UI", 10),
                corner_radius=6
            )
            text_widget.pack(fill='x', pady=5)
            current_prompt = self.load_prompt(key)
            if current_prompt:
                text_widget.insert('1.0', current_prompt)
            else:
                default_prompts = {
                    'voice_commands': """Process this voice command and determine the appropriate action.
    Consider the context and user preferences.
    Command: {command}
    Context: {context}""",
                    'chat_window': """You are a helpful AI assistant. Respond to the user's message
    while maintaining context of the conversation.
    Message: {message}
    History: {history}""",
                    'article_mode': """Analyze this article and provide key insights.
    Include main points and relevant details.
    Text: {text}
    Mode: {mode}""",
                    'programming': """Review this code and provide:
    1. Code analysis
    2. Suggestions for improvement
    3. Best practices
    4. Potential issues
    Code: {code}
    Language: {language}""",
                    'header_analysis': """Analyze these dataset headers and provide:
    1. Brief explanation of each header
    2. Potential relationships between headers
    3. What insights could be derived
    Headers: {headers}""",
                    'full_analysis': """Analyze this dataset and provide:
    1. Key insights and patterns
    2. Data quality issues
    3. Notable correlations
    4. Machine learning potential
    Data Summary: {data_summary}""",
                    'correlation_analysis': """Analyze the correlation matrix and explain:
    1. Strong positive correlations
    2. Strong negative correlations
    3. Important feature relationships
    """
                }
                text_widget.insert('1.0', default_prompts.get(key, ""))
            button_frame = ctk.CTkFrame(
                section_frame,
                fg_color=self.themes[self.current_theme]["bg"],
                corner_radius=0
            )
            button_frame.pack(fill='x', pady=5)
            ctk.CTkButton(
                button_frame,
                text="Save Prompt",
                command=lambda k=key, t=text_widget: self.save_prompt(k, t.get('1.0', 'end-1c')),
                fg_color=self.themes[self.current_theme]["accent"],
                text_color="#ffffff",
                font=('Segoe UI', 10),
                corner_radius=6,
                hover_color=self.themes[self.current_theme]["highlight"]
            ).pack(side='right', padx=5)
            ctk.CTkButton(
                button_frame,
                text="Reset",
                command=lambda k=key, t=text_widget: self.reset_single_prompt(k, t),
                fg_color=self.themes[self.current_theme]["secondary"],
                text_color="#ffffff",
                font=('Segoe UI', 10),
                corner_radius=6,
                hover_color=self.themes[self.current_theme]["highlight"]
            ).pack(side='right', padx=5)
        ctk.CTkLabel(
            tab_commanding, 
            text="Command Recognition Settings",
            **title_style
        ).pack(pady=(20, 5), padx=20, anchor="w")
        ctk.CTkLabel(
            tab_commanding,
            text="Configure how A.N.I.S understands and processes your commands",
            **desc_style
        ).pack(pady=(0, 20), padx=20, anchor="w")
        command_frame = self.create_tron_style_frame(tab_commanding)
        command_frame.pack(fill="both", expand=True, padx=20, pady=10)
        ctk.CTkLabel(
            command_frame,
            text="Command Recognition Prompt",
            font=("Arial", 14, "bold")
        ).pack(pady=(10, 5), padx=20, anchor="w")
        commanding_prompt = ctk.CTkTextbox(
            command_frame,
            height=300,
            width=500,
            fg_color=self.themes[self.current_theme]["input_bg"],
            text_color=self.themes[self.current_theme]["text"],
            border_color=self.themes[self.current_theme]["border"],
            border_width=1
        )
        commanding_prompt.pack(fill="both", expand=True, padx=20, pady=10)
        current_prompt = self.load_prompt('commanding_prompt')
        if not current_prompt:
            current_prompt = """You are A.N.I.S, an AI assistant that helps interpret natural language commands. 
Your task is to convert conversational requests into specific command formats.
and my name is Samiulla
You should recognize requests like:
- you should give human like coversation and Don't uneccessarily give command like -> "search google" 
- Opening applications: "Can you show me Notepad please" → "open notepad"
- Searching the web: "Look up information about climate change" → "search for climate change"
- Opening YouTube with topics: "I want to watch videos about AI" → "youtube about AI"
if my command "any possible to rcb win the cup"  → "no way"
- System commands: "What's the time" → "time"
-if my command:[shortcut_name] → "[shortcut_name]"
-if my command:"file assistant" or "file handling" or "file prompt" or "file ai" → "file assistant"
-if my command only :"create a program" → "create a program"
-if my command:"save" → "saved"
-if my command:"save us" or "save as" → "save as"
-if my command:"save highlighted text" or "save the highlights" or "save the hightlight" or "save selection" or "save selected text" or "save highlight" or "save the highlighted words" → "save highlighted text"
-if my command:"save note" or "save note" or "save the note" or "save the notes" or "save notes" → "save note"
-if my command "click [button_name]"→ "click [button_name]"
-if my command "this is a [button_name]"→ "this is a [button_name]"
-if my command:"send a message to [person_name]" → ""send a message to [person_name]""'
-if my command:"[query]":"[answer]"
-if my command:"live assistant" or "open live assistant" or "launch live assistant" or "I need assistant" → "live assistant"
-if my command:"[relate to live assistant]" → "live assistant"
-if my command:"[shortcut_name]":"[shortcut_name]"
if my command:"[settings_name]" or ""settings [setings_name] or "activate the settings [settings_name]" or "activate settings [settings_name]" or "activate settings" → "activate settings [settings_name]"
-if my command:"select all"  → "select all"
-if my command:"email window"  → "show email"
-if my command:"send" or "enter"  → "enter"
-if my command:"what is the time now" → "the time"
-if my command:"contact manager"→""contact manager""
-if my command:"set brightness [level]" → "set brightness [level]"
-if my command:"set volume [level]" → "set volume [level]"
-if my command :"current storage of my laptop" ->'storage'
-if my command:"[app_name] auto login" → "[app_name] auto login"
-if my command:"type" or "type here" →"type"
-if my command:"create a chatbot for  [query]"  or "create a chat box [query]" → "create a chatbot [query]"
-if my command:"i would like to work in [app_name]" → "open [app_name]"
-if my command:"analyse news" → "analyse news"
-if my command:"analyse text" or "analyse highlighted text" → "analyse news"
-if my command:"analyse link" → "analyse news"
-if my command:"i want to analyse this text is correct or not"-"analyse news"
if my command only:"the time" or "the date" or "the day" -> "the time"
-if my command:"i would like to analyse data or data analysis":"data analysis"
-if my command:"you can leave":"you can leave"
-if my command:"types of [query]" like this type of simple commands:search google [query]
-if my command:"read text" or "read" or "read highlighted"  → "read text"
if my command:"say your name":" I am A N I S AI"
-if my command:"system health":"check system health"
-if my command only:"auto install [query]":"auto install [correct_app_name_query]" (process this command only when say this otherwise normal response ai is okay)
if my command:"what's special day today" :"search google special day today"
-if my command:"anish chatbot" or "chatbot":"chatbot"
-if my command:"special day today" or "what's today special" or  "today's special"or "today special" or "what special today" ::"search google special day today"
-if my command:"is that you can install the application automatically":you should understand this is not i mention you to install application i am just asking you wheather you can or not ,and you re asnwer is yes and you have this capability to install certain application
-if my command:"[query] in python" -> "[query] in python"
-if my command:"you are built from" like this type of normal command it should not use the search google and youtube,don't use the search google and youtube unneccesarily
-if my command:"send email" or "send mail" → "send email"
Specific patterns to handle properly:
- For application opening requests with phrases like "please show me the X", "can you show me X", extract just the application name:
  * "can you please show me the RBI application" → "open rbi"
  * "can you please show me the Power BI application" → "open power bi"
  * "please show me another Notepad for me" → "open notepad"
  *"i would like to study" → "study mode on"
  *"study mode on" or "study mode activate"  → "study mode on"
*"could you please read the text for me"  → "read text"
*"could you please read the text"  → "read text"
* "study mode deactivate"  → "study mode off"
- Remove filler words like "me", "the", "please", "can you", etc.
-- When processing requests to open applications:
  * Extract just the core application name
  * Identify special applications like PowerBI, Power BI, PowerPoint, etc.
* if my request is "type 	[query]"   → "type [query]"
Context-aware queries:
- For phrases for phrsases"is that you can install the application automatically":you should understand this is not i mention you to install application i am just asking you wheather you can or not ,and you re asnwer is yes and you have this capability to install certain application : "response of ai"
- For phrases like "what are the types of this" or "show me examples of that" or "give me advanced tutorials about it":
  * Recognize these as follow-up requests related to the last topic discussed
  * Check for context words like "this", "that", "it", "these", "those","is that"
  * Extract the request type (types, examples, tutorials, concepts)
  * Extract the level if specified (beginner, intermediate, advanced)
  * don't use the phrase I will search for just search google or youtube for that topic"
- * you can call me Samiulla Master
Always return simple, direct commands that can be processed by the system.
Commands should be converted to their exact format without extra words or explanations.
"""
        commanding_prompt.insert("1.0", current_prompt)
        button_frame = ctk.CTkFrame(command_frame, fg_color="transparent")
        button_frame.pack(fill="x", padx=20, pady=(5, 15))
        save_button = self.create_tron_button(
            button_frame, 
            text="Save",
            command=lambda: self.save_prompt('commanding_prompt', commanding_prompt.get("1.0", "end-1c"))
        )
        save_button.pack(side="left", padx=(0, 10))
        reset_button = self.create_tron_button(
            button_frame,
            text="Reset to Default",
            command=lambda: self.reset_single_prompt('commanding_prompt', commanding_prompt)
        )
        reset_button.pack(side="left")
        def on_close():
            settings_window.destroy()
        ctk.CTkButton(
            settings_window,
            text="Close",
            command=on_close,
            fg_color=self.themes[self.current_theme]["accent"],
            text_color="#ffffff",
            font=('Segoe UI', 12, 'bold'),
            width=100,
            height=40,
            corner_radius=8,
            hover_color=self.themes[self.current_theme]["secondary"]
        ).pack(pady=20)
        ctk.CTkLabel(
            tab_api, 
            text="API Settings",
            **title_style
        ).pack(pady=(20, 5), padx=20, anchor="w")
        ctk.CTkLabel(
            tab_api,
            text="Configure API keys and email settings",
            **desc_style
        ).pack(pady=(0, 20), padx=20, anchor="w")
        gemini_frame = ctk.CTkFrame(
            tab_api,
            fg_color=self.themes[self.current_theme]["bg"],
            corner_radius=8
        )
        gemini_frame.pack(fill='x', pady=10, padx=20)
        ctk.CTkLabel(
            gemini_frame,
            text="Gemini AI Settings",
            font=("Segoe UI", 14, "bold"),
            text_color=self.themes[self.current_theme]["accent"]
        ).pack(anchor='w', pady=(10, 15), padx=15)
        api_key_frame = ctk.CTkFrame(
            gemini_frame,
            fg_color=self.themes[self.current_theme]["bg"],
            corner_radius=0
        )
        api_key_frame.pack(fill='x', pady=5, padx=15)
        ctk.CTkLabel(
            api_key_frame,
            text="API Key:",
            font=("Segoe UI", 11),
            text_color=self.themes[self.current_theme]["fg"]
        ).pack(side='left', padx=(0, 10))
        current_api_key = str(self.load_setting('gemini_api_key')) or "Enter your Api_key"
        self.api_key_entry = ctk.CTkEntry(
            api_key_frame,
            width=300,
            font=("Segoe UI", 10),
            fg_color=self.themes[self.current_theme]["input_bg"],
            text_color=self.themes[self.current_theme]["text"],
            border_color=self.themes[self.current_theme]["border"],
            border_width=1
        )
        self.api_key_entry.insert(0, current_api_key)
        self.api_key_entry.pack(side='left', fill='x', expand=True)
        model_frame = ctk.CTkFrame(
            gemini_frame,
            fg_color=self.themes[self.current_theme]["bg"],
            corner_radius=0
        )
        model_frame.pack(fill='x', pady=5, padx=15)
        ctk.CTkLabel(
            model_frame,
            text="Model Name:",
            font=("Segoe UI", 11),
            text_color=self.themes[self.current_theme]["fg"]
        ).pack(side='left', padx=(0, 10))
        current_model = str(self.load_setting('gemini_model')) or "Enter Model Name"
        self.model_entry = ctk.CTkEntry(
            model_frame,
            width=300,
            font=("Segoe UI", 10),
            fg_color=self.themes[self.current_theme]["input_bg"],
            text_color=self.themes[self.current_theme]["text"],
            border_color=self.themes[self.current_theme]["border"],
            border_width=1
        )
        self.model_entry.insert(0, current_model)
        self.model_entry.pack(side='left', fill='x', expand=True)
        ctk.CTkButton(
            gemini_frame,
            text="Save Gemini Settings",
            command=self.save_gemini_settings,
            fg_color=self.themes[self.current_theme]["accent"],
            text_color="#ffffff",
            font=('Segoe UI', 10, 'bold'),
            corner_radius=8,
            hover_color=self.themes[self.current_theme]["secondary"]
        ).pack(pady=15)
        email_frame = ctk.CTkFrame(
            tab_api,
            fg_color=self.themes[self.current_theme]["bg"],
            corner_radius=8
        )
        email_frame.pack(fill='x', pady=10, padx=20)
        ctk.CTkLabel(
            email_frame,
            text="Email Settings",
            font=("Segoe UI", 14, "bold"),
            text_color=self.themes[self.current_theme]["accent"]
        ).pack(anchor='w', pady=(10, 15), padx=15)
        email_addr_frame = ctk.CTkFrame(
            email_frame,
            fg_color=self.themes[self.current_theme]["bg"],
            corner_radius=0
        )
        email_addr_frame.pack(fill='x', pady=5, padx=15)
        ctk.CTkLabel(
            email_addr_frame,
            text="Email Address:",
            font=("Segoe UI", 11),
            text_color=self.themes[self.current_theme]["fg"]
        ).pack(side='left', padx=(0, 10))
        current_email = self.load_setting('email_address') or "samiullas831@gmail.com"
        self.email_entry = ctk.CTkEntry(
            email_addr_frame,
            width=300,
            font=("Segoe UI", 10),
            fg_color=self.themes[self.current_theme]["input_bg"],
            text_color=self.themes[self.current_theme]["text"],
            border_color=self.themes[self.current_theme]["border"],
            border_width=1
        )
        self.email_entry.insert(0, current_email)
        self.email_entry.pack(side='left', fill='x', expand=True)
        password_frame = ctk.CTkFrame(
            email_frame,
            fg_color=self.themes[self.current_theme]["bg"],
            corner_radius=0
        )
        password_frame.pack(fill='x', pady=5, padx=15)
        ctk.CTkLabel(
            password_frame,
            text="App Password:",
            font=("Segoe UI", 11),
            text_color=self.themes[self.current_theme]["fg"]
        ).pack(side='left', padx=(0, 10))
        current_password = self.load_setting('email_password') or "hikk vnel ychc play"
        self.password_entry = ctk.CTkEntry(
            password_frame,
            width=300,
            font=("Segoe UI", 10),
            fg_color=self.themes[self.current_theme]["input_bg"],
            text_color=self.themes[self.current_theme]["text"],
            border_color=self.themes[self.current_theme]["border"],
            border_width=1,
            show="•"
        )
        self.password_entry.insert(0, current_password)
        self.password_entry.pack(side='left', fill='x', expand=True)
        smtp_frame = ctk.CTkFrame(
            email_frame,
            fg_color=self.themes[self.current_theme]["bg"],
            corner_radius=0
        )
        smtp_frame.pack(fill='x', pady=5, padx=15)
        ctk.CTkLabel(
            smtp_frame,
            text="SMTP Server:",
            font=("Segoe UI", 11),
            text_color=self.themes[self.current_theme]["fg"]
        ).pack(side='left', padx=(0, 10))
        current_smtp = self.load_setting('smtp_server') or "smtp.gmail.com"
        self.smtp_entry = ctk.CTkEntry(
            smtp_frame,
            width=300,
            font=("Segoe UI", 10),
            fg_color=self.themes[self.current_theme]["input_bg"],
            text_color=self.themes[self.current_theme]["text"],
            border_color=self.themes[self.current_theme]["border"],
            border_width=1
        )
        self.smtp_entry.insert(0, current_smtp)
        self.smtp_entry.pack(side='left', fill='x', expand=True)
        port_frame = ctk.CTkFrame(
            email_frame,
            fg_color=self.themes[self.current_theme]["bg"],
            corner_radius=0
        )
        port_frame.pack(fill='x', pady=5, padx=15)
        ctk.CTkLabel(
            port_frame,
            text="SMTP Port:",
            font=("Segoe UI", 11),
            text_color=self.themes[self.current_theme]["fg"]
        ).pack(side='left', padx=(0, 10))
        current_port = self.load_setting('smtp_port') or "587"
        self.port_entry = ctk.CTkEntry(
            port_frame,
            width=100,
            font=("Segoe UI", 10),
            fg_color=self.themes[self.current_theme]["input_bg"],
            text_color=self.themes[self.current_theme]["text"],
            border_color=self.themes[self.current_theme]["border"],
            border_width=1
        )
        self.port_entry.insert(0, current_port)
        self.port_entry.pack(side='left')
        ctk.CTkButton(
            email_frame,
            text="Save Email Settings",
            command=self.save_email_settings,
            fg_color=self.themes[self.current_theme]["accent"],
            text_color="#ffffff",
            font=('Segoe UI', 10, 'bold'),
            corner_radius=8,
            hover_color=self.themes[self.current_theme]["secondary"]
        ).pack(pady=15)
        screen_frame = ctk.CTkFrame(tab_screen)
        screen_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        title_label = ctk.CTkLabel(
            screen_frame,
            text="Screen Interaction Settings",
            font=ctk.CTkFont(size=18, weight="bold")
        )
        title_label.pack(anchor=tk.W, pady=(0, 15))
        resolution_frame = ctk.CTkFrame(screen_frame)
        resolution_frame.pack(fill=tk.X, pady=10)
        resolution_label = ctk.CTkLabel(
            resolution_frame,
            text="Screen Resolution Adaptation",
            font=ctk.CTkFont(size=16, weight="bold")
        )
        resolution_label.pack(anchor=tk.W, padx=15, pady=(10, 5))
        if not hasattr(self, '_screen_adapter'):
            self._screen_adapter = ScreenCoordinateAdapter()
        resolution_desc = ctk.CTkLabel(
            resolution_frame,
            text="Coordinates are automatically adapted to your current screen resolution. Update the reference resolution to set your current screen as the new standard.",
            wraplength=700
        )
        resolution_desc.pack(anchor=tk.W, padx=15, pady=(0, 5))
        current_resolution = f"Current Resolution: {self._screen_adapter.current_width}x{self._screen_adapter.current_height}"
        reference_resolution = f"Reference Resolution: {self._screen_adapter.reference_width}x{self._screen_adapter.reference_height}"
        current_res_label = ctk.CTkLabel(
            resolution_frame,
            text=current_resolution
        )
        current_res_label.pack(anchor=tk.W, padx=15, pady=(0, 5))
        ref_res_label = ctk.CTkLabel(
            resolution_frame,
            text=reference_resolution
        )
        ref_res_label.pack(anchor=tk.W, padx=15, pady=(0, 5))
        def update_reference_resolution():
            if self._screen_adapter.update_reference_resolution():
                ref_res_label.configure(text=f"Reference Resolution: {self._screen_adapter.current_width}x{self._screen_adapter.current_height}")
                self.speak("Screen resolution reference updated")
                messagebox.showinfo("Resolution Updated", "Reference resolution has been updated to match your current screen size.")
            else:
                messagebox.showerror("Error", "Failed to update reference resolution.")
        resolution_btn = ctk.CTkButton(
            resolution_frame,
            text="Update Reference Resolution",
            command=update_reference_resolution,
            width=200,
            fg_color=self.themes[self.current_theme]["accent"]
        )
        resolution_btn.pack(anchor=tk.W, padx=15, pady=5)
        coords_frame = ctk.CTkFrame(screen_frame)
        coords_frame.pack(fill=tk.X, pady=10)
        coords_label = ctk.CTkLabel(
            coords_frame,
            text="Saved Coordinates",
            font=ctk.CTkFont(size=16, weight="bold")
        )
        coords_label.pack(anchor=tk.W, padx=15, pady=(10, 5))
        coords_desc = ctk.CTkLabel(
            coords_frame,
            text="Manage UI element coordinates that have been saved during screen interaction.",
            wraplength=700
        )
        coords_desc.pack(anchor=tk.W, padx=15, pady=(0, 5))
        try:
            conn = sqlite3.connect(DATABASE_FILE)
            cursor = conn.cursor()
            cursor.execute("SELECT COUNT(*) FROM coordinates")
            coord_count = cursor.fetchone()[0]
            conn.close()
        except:
            coord_count = 0
        coords_status = ctk.CTkLabel(
            coords_frame,
            text=f"Currently saved coordinates: {coord_count}",
        )
        coords_status.pack(anchor=tk.W, padx=15, pady=(0, 5))
        coords_btn = ctk.CTkButton(
            coords_frame,
            text="Open Coordinates Manager",
            command=self.show_coordinates_manager,
            width=200,
            fg_color=self.themes[self.current_theme]["accent"]
        )
        coords_btn.pack(anchor=tk.W, padx=15, pady=5)
        try:
            conn = sqlite3.connect(DATABASE_FILE)
            cursor = conn.cursor()
            cursor.execute("SELECT COUNT(*) FROM conversations WHERE command_type IS NOT NULL")
            cmd_count = cursor.fetchone()[0]
            conn.close()
        except:
            cmd_count = 0
    def load_prompt(self, prompt_type):
        """Load prompt from database"""
        try:
            db_path =  "prompts.db"
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            cursor.execute('''CREATE TABLE IF NOT EXISTS prompts
                            (type TEXT PRIMARY KEY, content TEXT)''')
            cursor.execute('SELECT content FROM prompts WHERE type = ?', (prompt_type,))
            result = cursor.fetchone()
            conn.close()
            return result[0] if result else ""
        except Exception as e:
            logging.error(f"Error loading prompt: {str(e)}")
            return ""
    def save_prompt(self, prompt_type, content):
        """Save prompt to database"""
        try:
            db_path =  "prompts.db"
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            cursor.execute('''CREATE TABLE IF NOT EXISTS prompts
                            (type TEXT PRIMARY KEY, content TEXT)''')
            cursor.execute('''INSERT OR REPLACE INTO prompts (type, content)
                            VALUES (?, ?)''', (prompt_type, content))
            conn.commit()
            conn.close()
            messagebox.showinfo("Success", "Prompt saved successfully!")
        except Exception as e:
            logging.error(f"Error saving prompt: {str(e)}")
            messagebox.showerror("Error", f"Failed to save prompt: {str(e)}")
    def change_theme(self, theme_name):
        """Change the application theme"""
        if theme_name in self.themes:
            self.current_theme = theme_name
            self.apply_theme(theme_name)
    def open_system_app(self, command):
        """Open system applications with proper path handling"""
        office_paths = {
            '64bit': [
                "C:\\Program Files\\Microsoft Office\\root\\Office16",
                "C:\\Program Files\\Microsoft Office\\Office16",
                "C:\\Program Files (x86)\\Microsoft Office\\root\\Office16",
                "C:\\Program Files (x86)\\Microsoft Office\\Office16"
            ]
        }
        app_commands = {
            'word': {
                'name': 'Microsoft Word',
                'cmd': 'WINWORD.EXE',
                'paths': office_paths['64bit']
            },
            'excel': {
                'name': 'Microsoft Excel',
                'cmd': 'EXCEL.EXE',
                'paths': office_paths['64bit']
            },
            'powerpoint': {
                'name': 'Microsoft PowerPoint',
                'cmd': 'POWERPNT.EXE',
                'paths': office_paths['64bit']
            }
        }
        try:
            command = command.lower()
            if command in app_commands:
                app_info = app_commands[command]
                executable_found = False
                for base_path in app_info['paths']:
                    full_path = os.path.join(base_path, app_info['cmd'])
                    if os.path.exists(full_path):
                        subprocess.Popen(full_path)
                        executable_found = True
                        break
                if not executable_found:
                    error_msg = f"Could not find {app_info['name']}. Please verify it is installed correctly."
                    logging.error(error_msg)
                    messagebox.showerror("Application Error", error_msg)
            else:
                pass
        except Exception as e:
            error_msg = f"Error opening application: {str(e)}"
            logging.error(error_msg)
            messagebox.showerror("Error", error_msg)
    def search_and_click_first_result(self, search_term):
        """
        Search for an application/item using Windows search and click the first result
        """
        try:
            logging.info(f"Starting Windows search for '{search_term}'")
            pyautogui.hotkey('win')
            time.sleep(0.7)
            pyautogui.write(search_term)
            time.sleep(1.5)
            screen_width, screen_height = pyautogui.size()
            x_pos = screen_width // 2
            y_pos = int(screen_height * 0.2)
            logging.info(f"Clicking at position ({x_pos}, {y_pos})")
            pyautogui.click(x_pos, y_pos)
            logging.info(f"Successfully clicked on first search result for '{search_term}'")
            return True
        except Exception as e:
            logging.error(f"Error in search_and_click_first_result: {str(e)}")
            try:
                logging.info("Trying fallback: pressing Enter")
                pyautogui.press('enter')
                return True
            except Exception as enter_error:
                logging.error(f"Error pressing enter: {str(enter_error)}")
                return False
    def change_volume(self, value):
        """Change the assistant's voice volume"""
        try:
            engine.setProperty('volume', float(value) / 100)
            logging.info(f"Voice volume changed to: {value}%")
        except Exception as e:
            logging.error(f"Error changing volume: {str(e)}")
    def change_voice_volume(self, value):
        """Change the assistant's voice volume (alias for change_volume)"""
        self.change_volume(value)
    def create_tron_style_frame(self, parent, has_glow=True):
        """Create a TRON Legacy styled frame"""
        frame = ctk.CTkFrame(
            parent,
            fg_color='#0c1221',  # Dark blue background
            border_color='#00f6ff',  # Neon blue border
            border_width=1 if has_glow else 0,
            corner_radius=0
        )
        if has_glow:
            glow = ctk.CTkFrame(
                frame,
                fg_color='#00f6ff',  # Neon blue glow
                height=2,
                corner_radius=0
            )
            glow.pack(fill='x', side='top')
            bottom_glow = ctk.CTkFrame(
                frame,
                fg_color='#00f6ff',  # Neon blue glow
                height=2,
                corner_radius=0
            )
            bottom_glow.pack(fill='x', side='bottom')
        return frame
    def create_tron_button(self, parent, text, command, width=None):
        """Create a TRON Legacy styled button"""
        button_frame = ctk.CTkFrame(
            parent,
            fg_color='#0c1221',
            border_color='#00f6ff',
            border_width=1,
            corner_radius=0
        )
        button = ctk.CTkButton(
            button_frame,
            text=text,
            command=command,
            fg_color='#0c1221',
            text_color='#00f6ff',
            font=('Orbitron', 10, 'bold'),
            corner_radius=0,
            hover_color='#0f1a2b',
            width=width if width else 100,
            height=30,
            border_width=0
        )
        button.pack(padx=1, pady=1)
        def on_enter(e):
            button_frame.configure(border_color='#00ffff')
            button.configure(fg_color='#0f1a2b')
        def on_leave(e):
            button_frame.configure(border_color='#00f6ff')
            button.configure(fg_color='#0c1221')
        button.bind('<Enter>', on_enter)
        button.bind('<Leave>', on_leave)
        return button_frame
    def get_storage_info(self):
        try:
            total, used, free = shutil.disk_usage("/")
            used_percentage = (used / total) * 100
            message = f"{used_percentage:.1f}% of storage is used"
            self.speak(message)
        except Exception as e:
            self.handle_errors(e, "Error getting storage info")
            self.speak("Sorry, I couldn't get the storage information")
    def reset_single_prompt(self, key, text_widget):
        """Reset a single prompt to default value"""
        try:
            default_prompts = {
                'chat_window': """You are a helpful AI assistant. Respond to the user's message
while maintaining context of the conversation.
Message: {message}
History: {history}""",
                'article_mode': """Analyze this article and provide key insights.
Include main points and relevant details.
Text: {text}
Mode: {mode}""",
                'programming': """Review this code and provide:
1. Code analysis
2. Suggestions for improvement
3. Best practices
4. Potential issues
Code: {code}
Language: {language}""",
                'header_analysis': """Analyze these dataset headers and provide:
1. Brief explanation of each header
2. Potential relationships between headers
3. What insights could be derived
Headers: {headers}""",
                'full_analysis': """Analyze this dataset and provide:
1. Key insights and patterns
2. Data quality issues
3. Notable correlations
4. Machine learning potential
Data Summary: {data_summary}""",
                'correlation_analysis': """Analyze the correlation matrix and explain:
1. Strong positive correlations
2. Strong negative correlations
3. Important feature relationships
"""
            }
            text_widget.delete('1.0', 'end')
            text_widget.insert('1.0', default_prompts.get(key, ""))
            db_path =  "prompts.db"
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            cursor.execute('DELETE FROM prompts WHERE type = ?', (key,))
            conn.commit()
            conn.close()
            messagebox.showinfo("Success", f"Reset {key} prompt to default")
        except Exception as e:
            logging.error(f"Error resetting prompt: {str(e)}")
            messagebox.showerror("Error", f"Failed to reset prompt: {str(e)}")
    def _convert_tk_properties(self, **kwargs):
        """Convert tkinter properties to customtkinter equivalents"""
        property_map = {
            'bg': 'fg_color',
            'highlightbackground': 'border_color',
            'highlightthickness': 'border_width',
            'bd': 'border_width'
        }
        converted_kwargs = {}
        for key, value in kwargs.items():
            if key in property_map:
                converted_kwargs[property_map[key]] = value
            else:
                converted_kwargs[key] = value
        return converted_kwargs
    def cleanup(self):
        """Clean up the visualizer before window destruction"""
        with self._lock:
            self.is_active = False
            self.animation_running = False
            if self.after_id:
                self.after_cancel(self.after_id)
                self.after_id = None
            try:
                self.delete("all")
            except tk.TclError:
                pass
    def _analyze_conversation_history(self):
        """
        Analyze conversation history to extract useful context for current command processing.
        Returns a dictionary with relevant context information.
        """
        context = {
            "last_action": None,
            "search_platform": None,
            "last_topics": [],
            "has_learning_context": False,
            "preferred_level": None,
            "preferred_content_type": None
        }
        if not hasattr(self, 'conversation_history') or not self.conversation_history:
            return context
        last_entries = self.conversation_history[-7:] if len(self.conversation_history) >= 7 else self.conversation_history
        for entry in last_entries:
            assistant_resp = entry.get('assistant', '').lower()
            if "searching" in assistant_resp:
                context["last_action"] = "search"
                if "youtube" in assistant_resp:
                    context["search_platform"] = "youtube"
                elif "google" in assistant_resp:
                    context["search_platform"] = "google"
                if "for " in assistant_resp:
                    search_topic = assistant_resp.split("for ", 1)[1].strip()
                    if search_topic and search_topic not in context["last_topics"]:
                        context["last_topics"].append(search_topic)
            if "learn" in assistant_resp or "learning" in assistant_resp:
                context["has_learning_context"] = True
            user_req = entry.get('user', '').lower()
            level_terms = [""]
            for level in level_terms:
                if level in user_req:
                    context["preferred_level"] = level
                    break
            type_terms = ["types", "concepts", "examples", "applications", "tutorial", "course"]
            for term in type_terms:
                if term in user_req:
                    context["preferred_content_type"] = term
                    break
        return context
    def analyze_command_with_gemini(self, natural_command):
        """
        Use Gemini AI to interpret natural language commands and map them to specific commands
        the assistant understands
        """
        try:
            lower_cmd = natural_command.lower().strip()
            topic_like_phrases = lower_cmd.split()
            if 1 <= len(topic_like_phrases) <= 3:
                command_verbs = ["open", "search", "play", "show", "tell", "what", "how", "why", "when", "is", "are", "can", "do", "does", "will", "google", "youtube", "type", "set", "get", "increase", "decrease"]
                if not any(verb in lower_cmd for verb in command_verbs):
                    logging.info(f"Identified '{lower_cmd}' as a likely topic, returning directly.")
                    self.last_topic = lower_cmd
                    return lower_cmd
            reference_words = ["that", "it", "this", "those", "them", "these","he","she",'we']
            contains_reference = any(word in lower_cmd.split() for word in reference_words)
            is_continuation = False
            recent_context = ""
            if hasattr(self, 'conversation_history') and self.conversation_history:
                recent_context = "\nRecent conversation history:\n"
                for i, interaction in enumerate(self.conversation_history[-7:]):
                    recent_context += f"User: {interaction['user']}\n"
                    recent_context += f"Assistant: {interaction['assistant']}\n"
                follow_up_indicators = ["and", "also", "what about", "but", "so", "then"]
                if any(lower_cmd.startswith(word) for word in follow_up_indicators) or len(lower_cmd.split()) <= 3:
                    is_continuation = True
            if lower_cmd.startswith("advanced level of") and contains_reference and self.last_topic:
                search_query = f"advanced {self.last_topic}"
                if hasattr(self, 'update_command_flow'):
                    self.update_command_flow(
                        task_type="learning",
                        subject=self.last_topic,
                        command_type="study",
                        qualifiers={"level": "advanced"}
                    )
                return f"search for {search_query}"
            if (lower_cmd.startswith("what are the") or lower_cmd.startswith("what is the") or 
                lower_cmd.startswith("tell me the") or lower_cmd.startswith("list the")) and contains_reference:
                if self.last_topic:
                    question_parts = lower_cmd.split("the", 1)
                    if len(question_parts) > 1:
                        question_type = question_parts[1].split("of", 1)[0].strip()
                        return f"What are the {question_type} of {self.last_topic}?"
            direct_question_starters = ["what", "how", "why", "when", "where", "who", "can", "are", "is"]
            if any(lower_cmd.startswith(starter) for starter in direct_question_starters) and contains_reference:
                if self.last_topic:
                    for ref in reference_words:
                        lower_cmd = re.sub(rf'\b{ref}\b', self.last_topic, lower_cmd)
                    return lower_cmd
            if self.last_topic and contains_reference:
                search_platform_pattern = re.compile(r'(?:search|look for|look up|find)\s+(?:this|that|it)\s+(?:in|on)\s+(youtube|google)', re.IGNORECASE)
                search_match = search_platform_pattern.search(lower_cmd)
                if search_match:
                    platform = search_match.group(1).lower()
                    if platform == "youtube":
                        return f"search youtube for {self.last_topic}"
                    else:
                        return f"search for {self.last_topic}"
                watch_platform_pattern = re.compile(r'(?:would like to |want to |please |)\s*(?:watch|see|view)\s+(?:this|that|it)\s+(?:in|on)?\s+(youtube|google)?', re.IGNORECASE)
                watch_match = watch_platform_pattern.search(lower_cmd)
                if watch_match:
                    platform = watch_match.group(1).lower() if watch_match.group(1) else "youtube"
                    if platform == "youtube":
                        return f"search youtube for {self.last_topic}"
                    else:
                        return f"search for {self.last_topic}"
            if self.last_topic and contains_reference:
                learning_keywords = ["learn", "learning", "study", "teach me", "show me how"]
                level_keywords = ["advanced", "beginner", "intermediate", "basic", "expert", "professional", "concept"]
                is_learning_request = any(keyword in lower_cmd for keyword in learning_keywords)
                has_level = any(level in lower_cmd for level in level_keywords)
                if is_learning_request:
                    level_term = ""
                    for level in level_keywords:
                        if level in lower_cmd:
                            level_term = level
                            break
                    if level_term:
                        if level_term == "concept":
                            search_query = f"advanced {self.last_topic} concepts"
                        else:
                            search_query = f"{level_term} {self.last_topic}"
                    else:
                        search_query = f"how to learn {self.last_topic}"
                    if hasattr(self, 'update_command_flow'):
                        self.update_command_flow(
                            task_type="learning",
                            subject=self.last_topic,
                            command_type="study",
                            qualifiers={"level": level_term if level_term else "beginner"}
                        )
                    return f"search for {search_query}"
            if self.last_topic:
                if (any(word in lower_cmd for word in ["watch", "view", "see"]) and 
                    any(ref in lower_cmd for ref in reference_words)):
                    if "english" in lower_cmd or "dub" in lower_cmd or "dubbed" in lower_cmd:
                        if "dub" in lower_cmd or "dubbed" in lower_cmd:
                            combined_term = f"{self.last_topic} english dub"
                        else:
                            combined_term = f"{self.last_topic} in english"
                        self.last_topic = combined_term
                        return f"search youtube for {combined_term}"
                if len(lower_cmd.split()) <= 6 and self.conversation_history and len(self.conversation_history) >= 1:
                    last_assistant_response = self.conversation_history[-1].get('assistant', '').lower()
                    if ("searching" in last_assistant_response and 
                        any(kw in lower_cmd for kw in ["in english", "english", "dub", "dubbed", "subbed", "subtitle"])):
                        if "dub" in lower_cmd or "dubbed" in lower_cmd:
                            new_search = f"{self.last_topic} english dub"
                        elif "english" in lower_cmd:
                            new_search = f"{self.last_topic} in english"
                        else:
                            new_search = f"{self.last_topic} {lower_cmd}"
                        self.last_topic = new_search
                        if "youtube" in last_assistant_response:
                            return f"search youtube for {new_search}"
                        else:
                            return f"search for {new_search}"
            if lower_cmd in ["google", "youtube"] and self.last_topic:
                if lower_cmd == "google":
                    return f"search for {self.last_topic}"
                elif lower_cmd == "youtube":
                    return f"search youtube for {self.last_topic}"
            if lower_cmd.startswith("get_with_context"):
                if hasattr(self, 'last_topic') and self.last_topic:
                    parts = lower_cmd.split()
                    if len(parts) >= 2:
                        request_type = parts[1] if len(parts) > 1 else "information"
                        level = parts[2] if len(parts) > 2 else None
                        base_topic = self.last_topic
                        type_terms = ["types", "categories", "kinds", "examples"]
                        for type_term in type_terms:
                            if type_term in base_topic:
                                parts = base_topic.split(type_term)
                                if len(parts) > 1 and parts[1].strip().startswith("of "):
                                    base_topic = parts[1].strip()[3:].strip()
                                    break
                        previous_context = self._analyze_conversation_history() if hasattr(self, '_analyze_conversation_history') else {"has_learning_context": False}
                        if level and request_type:
                            search_query = f"{level} {request_type} of {base_topic}"
                        elif level:
                            search_query = f"{level} level {base_topic}"
                        elif request_type:
                            search_query = f"{request_type} of {base_topic}"
                        else:
                            search_query = base_topic
                        if request_type in type_terms:
                            self.last_topic = f"{request_type} of {base_topic}"
                        if hasattr(self, 'update_command_flow'):
                            self.update_command_flow(
                                task_type="search" if not previous_context.get("has_learning_context", False) else "learning",
                                subject=base_topic,
                                command_type="query",
                                qualifiers={
                                    "level": level if level else "general",
                                    "type": request_type if request_type else "general",
                                    "platform": "google"
                                }
                            )
                        return f"search for {search_query}"
                else:
                    return "I don't have context for that request. Could you be more specific?"
            if contains_reference and self.last_topic:
                logging.info(f"Detected reference to previous topic: '{self.last_topic}'")
                if any(term in lower_cmd for term in ["show me", "play", "video", "watch"]):
                    if any(word in lower_cmd for word in reference_words):
                        expanded_command = lower_cmd.replace("that", self.last_topic)
                        expanded_command = expanded_command.replace("it", self.last_topic)
                        expanded_command = expanded_command.replace("this", self.last_topic)
                        logging.info(f"Expanded command: '{expanded_command}'")
                        return f"search youtube for {self.last_topic}"
            if "google" in lower_cmd and not lower_cmd.startswith("search") and not lower_cmd.startswith("google search"):
                if not any(lower_cmd.startswith(term + " google") for term in ["search", "open", "launch", "start"]):
                    for search_term in lower_cmd.split("google"):
                        if search_term.strip():
                            self.last_topic = search_term.strip()
                            break
                    return lower_cmd
            question_starters = ['what', 'who', 'when', 'where', 'why', 'how', 'is', 'are', 'can', 'do', 'does', 'will']
            if any(lower_cmd.startswith(q) for q in question_starters) and not any(cmd in lower_cmd for cmd in ['open', 'search', 'play', 'show']):
                potential_topic = ' '.join(lower_cmd.split()[1:])
                if potential_topic:
                    self.last_topic = potential_topic
                return lower_cmd
            direct_cmd_patterns = [
                'search for', 'play music', 'exit', 'quit', 'goodbye',
                'battery', 'storage', 'time', 'date', 'study mode'
            ]
            if lower_cmd.startswith('search for') or lower_cmd.startswith('search'):
                search_parts = lower_cmd.split('search')
                if len(search_parts) > 1:
                    topic_part = search_parts[1].replace('for', '', 1).strip()
                    if topic_part:
                        self.last_topic = topic_part
                        logging.info(f"Updated last topic to: '{self.last_topic}'")
            for pattern in direct_cmd_patterns:
                if lower_cmd.startswith(pattern) or lower_cmd == pattern:
                    return lower_cmd
            app_patterns = ['open', 'launch', 'start', 'run', 'show me', 'can you open', 'please open', 
                          'could you open', 'would you open', 'can you show', 'please show', 'show the', 
                          'open up', 'start up', 'bring up', 'pull up']
            for pattern in app_patterns:
                if pattern in lower_cmd:
                    app_mapping = {
                        'notepad': 'open notepad',
                        'calculator': 'open calculator',
                        'paint': 'open paint',
                        'browser': 'open browser',
                        'chrome': 'open chrome',
                        'edge': 'open edge',
                        'firefox': 'open firefox',
                        'word': 'open word',
                        'excel': 'open excel',
                        'powerpoint': 'open powerpoint',
                        'power bi': 'open power bi',
                        'powerbi': 'open powerbi',
                        'explorer': 'open explorer',
                        'file explorer': 'open explorer',
                        'file manager': 'open explorer',
                        'settings': 'open settings',
                        'control panel': 'open control panel',
                        'task manager': 'open task manager',
                        'cmd': 'open cmd',
                        'command prompt': 'open cmd'
                    }
                    for app_name, command in app_mapping.items():
                        if app_name in lower_cmd:
                            return command
                    for pattern in app_patterns:
                        if pattern in lower_cmd:
                            app_request = lower_cmd.split(pattern, 1)[1].strip()
                            if app_request and not app_request.startswith(('a ', 'an ', 'the ')):
                                for filler in ['for me', 'for us', 'application', 'app', 'program', 'software', 'another']:
                                    app_request = app_request.replace(filler, '').strip()
                                if app_request:
                                    return f"open {app_request}"
            custom_prompt = self.load_prompt('commanding_prompt')
            if custom_prompt:
                system_prompt = custom_prompt
            else:
                system_prompt = """
                You are A.N.I.S, an AI assistant that helps interpret natural language commands. 
                Your task is to convert conversational requests into specific command formats.
                Here are the command formats you should map to:
                -don't include hear are the search for in the response
                -don't include hear are the search results for in the response
                -if user query is what is ,which,who,like this means he asking for about information not for to process command then you should return the command .For example suer tell what is python programming ,it means you should only explain about that
                - "search for X" - For web searches
                - "open youtube" - To open YouTube
                - "search youtube for X" - To search YouTube for X
                - "youtube about X" or "youtube related to X" - To search YouTube for X
                - "open google" - To open Google
                - "open wikipedia" - To open Wikipedia
                - "open X" - To open application X (e.g., "open notepad", "open calculator", "open power bi")
                - "play music" - To play music
                - "time" or "date" - For time and date info
                - "battery" - For battery info
                - "storage" - For storage info
                - "exit", "quit", or "goodbye" - To exit the assistant
                - "programming" - To open programming assistant
                - "reset conversation" - To reset conversation
                - "study mode on/off" - To toggle study mode
                - "type X" - To type text X
                - "volume up/down" - To adjust volume
                - "analyze screen" - To analyze what's on screen
                - "analyze database" - To analyze database
                Specific patterns to handle properly:
                - For application opening requests with phrases like "please show me the X", "can you show me X", extract just the application name:
                  * "can you please show me the RBI application" → "open rbi"
                  * "can you please show me the Power BI application" → "open power bi"
                  * "please show me another Notepad for me" → "open notepad"
                - Remove filler words like "me", "the", "please", "can you", etc.
                Examples of conversational to command mapping:
                - "Can you show me Notepad please" → "open notepad"
                - "Please open another instance of Calculator" → "open calculator"
                - "Show me the file explorer" → "open explorer" 
                - "Can you please show me the Power BI application" → "open power bi"
                - "I'd like to watch videos about AI" → "youtube about AI"
                - "Can you show me YouTube videos related to cooking recipes" → "youtube related to cooking recipes"
                - "I want to see what's on YouTube about space exploration" → "search youtube for space exploration"
                - "Could you look up information about climate change" → "search for climate change"
                - "Tell me what time it is" → "time"
                - "I'd like to exit now" → "exit"
                - "Help me with some programming" → "programming"
                If you recognize a request that matches these command patterns, convert it to the exact command format.
                For application opening requests, the format should be "open [application name]" without extra words.
                If you don't recognize a specific command intent, just return the original text unchanged.
                DO NOT include explanations, just return the command or the original text.
                if the user ask about information you can use this prompt to answer GuideLines: {custom_prompt}Guidelines:
                - answer should be in two lines for any query
                - Use simple language
                - Avoid special characters or formatting
                - Use proper punctuation
                - Keep responses brief and direct
                -your name is A N I S AI integrated with Operation System With Help Of Gemini and under Developing by 
                Samiulla from 26 August 2024 don't  tell this in beginning of your response you should tell only when someone ask you name
                - No emojis or symbols
                - answer in a friendly manner
                - answer like a human
                - answer in a way that is easy to understand
                - answer in a way that is not too formal
                - answer in a way that is not too technical
                - if need techinical answer in techincal terms
                - answer should be short like 2 lines and concise
                - avoid geographical deep terms
                - try to answer in 2 lines simple and short
                - some of the time you can answer in a funny way
                """
            user_message = f"Interpret this request: '{natural_command}'"
            response = gemini_model.generate_content([system_prompt, user_message])
            interpreted_command = response.text.strip()
            if interpreted_command.startswith('"') and interpreted_command.endswith('"'):
                interpreted_command = interpreted_command[1:-1]
            logging.info(f"Original: '{natural_command}' -> Interpreted: '{interpreted_command}'")
            if contains_reference and self.last_topic and interpreted_command == natural_command:
                expanded_cmd = natural_command.lower()
                for ref in reference_words:
                    if ref in expanded_cmd:
                        expanded_cmd = expanded_cmd.replace(ref, self.last_topic)
                logging.info(f"Trying expanded reference command: '{expanded_cmd}'")
                user_message = f"Interpret this request: '{expanded_cmd}'"
                try:
                    response = gemini_model.generate_content([system_prompt, user_message])
                    expanded_interpreted = response.text.strip()
                    if expanded_interpreted != interpreted_command:
                        logging.info(f"Using expanded interpretation: '{expanded_interpreted}'")
                        interpreted_command = expanded_interpreted
                except Exception as e:
                    logging.error(f"Error expanding reference: {str(e)}")
            if interpreted_command.startswith("search for "):
                potential_topic = interpreted_command.replace("search for ", "", 1).strip()
                if potential_topic:
                    self.last_topic = potential_topic
                    logging.info(f"Setting last_topic to: {potential_topic}")
            elif interpreted_command.startswith("search youtube for "):
                potential_topic = interpreted_command.replace("search youtube for ", "", 1).strip()
                if potential_topic:
                    self.last_topic = potential_topic
                    logging.info(f"Setting last_topic to: {potential_topic}")
            return interpreted_command if interpreted_command else natural_command
        except Exception as e:
            logging.error(f"Error in analyze_command_with_gemini: {str(e)}")
            return natural_command
    def get_gemini_action(self, command):
        """Use Gemini to process general commands that don't match specific patterns"""
        try:
            is_continuation = False
            recent_context = ""
            if self.conversation_history:
                recent_context = "\nRecent conversation history:\n"
                for i, interaction in enumerate(self.conversation_history[-7:]):
                    recent_context += f"User: {interaction['user']}\n"
                    recent_context += f"Assistant: {interaction['assistant']}\n"
                follow_up_indicators = ["and", "also", "what about", "but", "so", "then"]
                if any(command.lower().startswith(word) for word in follow_up_indicators) or len(command.split()) <= 3:
                    is_continuation = True
            reference_words = ["that", "it", "this", "those", "them", "these"]
            contains_reference = any(ref in command.lower().split() for ref in reference_words)
            if (command.lower().startswith("what are the") or command.lower().startswith("what is the")) and contains_reference and self.last_topic:
                type_indicators = ["types", "kinds", "categories", "examples", "applications", "benefits", "advantages", 
                                 "disadvantages", "limitations", "components", "parts", "features", "characteristics"]
                question_type = None
                for indicator in type_indicators:
                    if indicator in command.lower():
                        question_type = indicator
                        break
                if question_type:
                    specific_question = f"What are the {question_type} of {self.last_topic}? Provide a concise answer."
                    system_prompt = """
                    You are A.N.I.S, an AI assistant providing factual, educational information.
                    Answer the following question directly and concisely.
                    Focus on providing clear, accurate information without unnecessary explanation.
                    Your response should:
                    1. Be direct and factual
                    2. Use simple language
                    3. Be limited to 2-3 short sentences
                    4. Avoid unnecessary introductions or conclusions
                    5. Not use markdown formatting or special characters
                    If you're uncertain about specific details, provide generally accepted information.
                    """
                    response = gemini_model.generate_content([system_prompt, specific_question])
                    answer = response.text.strip()
                    clean_answer = self.clean_response(answer)
                    self.log_interaction(command, clean_answer)
                    self.last_topic = f"{question_type} of {self.last_topic}"
                    return clean_answer
            if "google" in command.lower() and not command.lower().startswith(("search", "google search", "open google")):
                if not command.lower().startswith("google"):
                    command_parts = command.lower().split("google")
                    command_without_google = " ".join([part.strip() for part in command_parts if part.strip()])
                    command = command_without_google
                    logging.info(f"Removed 'google' keyword from command: '{command}'")
            video_indicators = ["show me video", "show video", "show me a video", "video about", "videos about", "watch video", "video of", "show me the video"]
            if any(indicator in command.lower() for indicator in video_indicators):
                type_terms = ["types", "categories", "kinds", "examples", "tutorials", "concepts", "applications"]
                level_terms = [""]
                has_type = False
                selected_type = None
                for term in type_terms:
                    if term in command.lower():
                        has_type = True
                        selected_type = term
                        break
                has_level = False
                selected_level = None
                for term in level_terms:
                    if term in command.lower():
                        has_level = True
                        selected_level = term
                        break
                reference_words = ["that", "it", "this", "those", "them", "these"]
                if any(ref in command.lower().split() for ref in reference_words) and self.last_topic:
                    base_topic = self.last_topic
                    for type_term in type_terms:
                        if type_term in base_topic:
                            parts = base_topic.split(type_term)
                            if len(parts) > 1 and parts[1].strip().startswith("of "):
                                base_topic = parts[1].strip()[3:].strip()
                                break
                    if selected_level and selected_type:
                        search_query = f"{selected_level} {selected_type} of {base_topic}"
                    elif selected_level:
                        search_query = f"{selected_level} level {base_topic}"
                    elif selected_type:
                        search_query = f"{selected_type} of {base_topic}"
                    else:
                        search_query = base_topic
                    if selected_type:
                        self.last_topic = f"{selected_type} of {base_topic}"
                    if hasattr(self, 'update_command_flow'):
                        self.update_command_flow(
                            task_type="search",
                            subject=base_topic,
                            command_type="query",
                            qualifiers={
                                "level": selected_level if selected_level else "general",
                                "type": selected_type if selected_type else "general",
                                "platform": "youtube"
                            }
                        )
                    youtube_url = f"https://www.youtube.com/results?search_query={urllib.parse.quote(search_query)}"
                    webbrowser.open(youtube_url)
                    return f"Searching YouTube for {search_query}"
                else:
                    for indicator in video_indicators:
                        if indicator in command.lower():
                            topic = command.lower().split(indicator, 1)[1].strip()
                            if topic:
                                self.last_topic = topic
                                return f"search youtube for {topic}"
            info_request_indicators = ["tell me about", "give me information", "can you tell me", "explain", "what is", 
                                     "how does", "key points", "summary of", "important facts", "give some", "provide"]
            if any(indicator in command.lower() for indicator in info_request_indicators):
                reference_words = ["that", "it", "this", "those", "them", "these"]
                contains_reference = any(ref in command.lower().split() for ref in reference_words)
                if contains_reference and self.last_topic:
                    context = self.get_current_command_flow() if hasattr(self, 'get_current_command_flow') else ""
                    response = self.get_information_for_topic(self.last_topic, context)
                    self.speak(response)
                    return False
                else:
                    for indicator in info_request_indicators:
                        if indicator in command.lower():
                            parts = command.lower().split(indicator, 1)
                            if len(parts) > 1:
                                topic = parts[1].strip()
                                if topic.startswith("about "):
                                    topic = topic[6:].strip()
                                if topic:
                                    self.last_topic = topic
                                    context = self.get_current_command_flow() if hasattr(self, 'get_current_command_flow') else ""
                                    response = self.get_information_for_topic(topic, context)
                                    self.speak(response)
                                    return False
            system_prompt = f"""
            You are A N I S AI, an AI assistant. The user has given you a command that doesn't match any of your 
            pre-defined command patterns. Analyze this request and provide an appropriate response.
            Available capabilities:
            - Web search (ONLY when explicitly asked, e.g., "search for X", "look up X")
            - Opening applications
            - Checking system status
            - Playing music
            - Setting reminders
            - Showing time/date
            - Answering general knowledge questions (without searching unless asked)
            Your personality traits:
            - Helpful and friendly
            - Concise and to-the-point
            - Intelligent and knowledgeable
            - Proactive in suggesting better ways to accomplish tasks
            When you recognize an intent that could be handled by a specific feature:
            - For application opening requests: "I'll open [application] for you."
            - For web search requests: "[query]." (Only if explicitly asked to search)
            - For media requests: "Playing [media] now."
            - For system information: "Your [battery/storage/etc] is currently [status]."
            VERY IMPORTANT:
            1. Keep your response to a SINGLE SENTENCE. Never write multiple sentences.
            2. Do not use markdown formatting like **, *, ~, etc.
            3. Do not use bullet points or numbered lists.
            4. Keep responses under 15 words when possible.
            5. Be extremely direct and concise.
            6. Don't use phrases like "I can" or "I will".
            7. Do NOT trigger a search unless the user explicitly uses search keywords like "search", "look up", "find", "google", "search google", "search youtube".
            8. If the user asks a question (e.g., "what is X?"), just answer it directly without opening a browser.
            9. If the command is just a topic name (e.g., "quantum computing"), ask what the user wants to do with it (e.g., "What about quantum computing?").
            10. don't use the phrase "I will search for" just use this "search google" or "search youtube" for that topic.
            if user ask about information instead of you can use this for response and folow this:
                - answer should be in two lines for any query
                - Use simple language
                - Avoid special characters or formatting
                - Use proper punctuation
                - Keep responses brief and direct
                -your name is A N I S AI integrated with Operation System With Help Of Gemini and under Developing by 
                Samiulla from 26 August 2024 don't  tell this in beginning of your response you should tell only when someone ask you name
                - No emojis or symbols
                - answer in a friendly manner
                - answer like a human
                - answer in a way that is easy to understand
                - answer in a way that is not too formal
                - answer in a way that is not too technical
                - if need techinical answer in techincal terms
                - answer should be short like 2 lines and concise
                - avoid geographical deep terms
                - try to answer in 2 lines simple and short
                - some of the time you can answer in a funny way
            """
            if is_continuation and recent_context:
                user_message = f"Process this request with the following conversation context: {recent_context}\n\nCurrent request: '{command}'"
            else:
                user_message = f"Process this request: '{command}'"
            system_prompt = """
            You are A.N.I.S, an AI assistant. The user has given you a command that doesn't match any of your 
            pre-defined command patterns. Analyze this request and provide an appropriate response.
            Available capabilities:
            - Web search (ONLY when explicitly asked, e.g., "search for X", "look up X")
            - Opening applications
            - Checking system status
            - Playing music
            - Setting reminders
            - Showing time/date
            - Answering general knowledge questions (without searching unless asked)
            Your personality traits:
            - Helpful and friendly
            - Concise and to-the-point
            - Intelligent and knowledgeable
            - Proactive in suggesting better ways to accomplish tasks
            When you recognize an intent that could be handled by a specific feature:
            - For application opening requests: "I'll open [application] for you."
            - For web search requests: "[query]." (Only if explicitly asked to search)
            - For media requests: "Playing [media] now."
            - For system information: "Your [battery/storage/etc] is currently [status]."
            VERY IMPORTANT:
            1. Keep your response to a SINGLE SENTENCE. Never write multiple sentences.
            2. Do not use markdown formatting like **, *, ~, etc.
            3. Do not use bullet points or numbered lists.
            4. Keep responses under 15 words when possible.
            5. Be extremely direct and concise.
            6. Don't use phrases like "I can" or "I will".
            7. Do NOT trigger a search unless the user explicitly uses search keywords like "search", "look up", "find", "google", "search google", "search youtube".
            8. If the user asks a question (e.g., "what is X?"), just answer it directly without opening a browser.
            9. If the command is just a topic name (e.g., "quantum computing"), ask what the user wants to do with it (e.g., "What about quantum computing?").
            10. don't use the phrase "I will search for" just use this "search google" or "search youtube" for that topic.
            if user ask about information instead of you can use this for response and folow this:
                - answer should be in two lines for any query
                - Use simple language
                - Avoid special characters or formatting
                - Use proper punctuation
                - Keep responses brief and direct
                -your name is A N I S AI integrated with Operation System With Help Of Gemini and under Developing by 
                Samiulla from 26 August 2024 don't  tell this in beginning of your response you should tell only when someone ask you name
                - No emojis or symbols
                - answer in a friendly manner
                - answer like a human
                - answer in a way that is easy to understand
                - answer in a way that is not too formal
                - answer in a way that is not too technical
                - if need techinical answer in techincal terms
                - answer should be short like 2 lines and concise
                - avoid geographical deep terms
                - try to answer in 2 lines simple and short
                - some of the time you can answer in a funny way
            """
            response = gemini_model.generate_content([system_prompt, user_message])
            response_text = response.text.strip()
            response_text = self.clean_response(response_text)
            lower_cmd = command.lower()
            if any(app_term in lower_cmd for app_term in ['open', 'launch', 'start', 'show me', 'can you open', 'please open']):
                for app in ['notepad', 'calculator', 'paint', 'chrome', 'browser', 'explorer', 'settings', 'word', 'excel', 'power bi', 'powerbi']:
                    if app in lower_cmd:
                        self.open_application(app)
                        return None
            if any(term in command.lower() for term in ['open', 'launch', 'run', 'start', 'show me']):
                matches = re.findall(r'(?:open|launch|run|start|show|use)\s+(?:me|the|a|an)?\s*([a-zA-Z0-9\s]+?)(?:\s+app|\s+application|\s+program|\s+for me|\s*$)', command.lower())
                if matches:
                    app_name = matches[0].strip()
                    if app_name:
                        self.open_application(app_name)
                        return None
            info_request_indicators = ["tell me about", "give me information", "can you tell me", "explain", "what is", 
                                     "how does", "key points", "summary of", "important facts", "give some", "provide"]
            if any(indicator in command.lower() for indicator in info_request_indicators):
                reference_words = ["that", "it", "this", "those", "them", "these"]
                contains_reference = any(ref in command.lower().split() for ref in reference_words)
                if contains_reference and self.last_topic:
                    context = self.get_current_command_flow() if hasattr(self, 'get_current_command_flow') else ""
                    response = self.get_information_for_topic(self.last_topic, context)
                    google_url = f'https://www.google.com/search?q={urllib.parse.quote(self.last_topic)}'
                    webbrowser.open(google_url)
                    return response
                else:
                    for indicator in info_request_indicators:
                        if indicator in command.lower():
                            parts = command.lower().split(indicator, 1)
                            if len(parts) > 1:
                                topic = parts[1].strip()
                                if topic.startswith("about "):
                                    topic = topic[6:].strip()
                                if topic:
                                    self.last_topic = topic
                                    context = self.get_current_command_flow() if hasattr(self, 'get_current_command_flow') else ""
                                    google_url = f'https://www.google.com/search?q={urllib.parse.quote(topic)}'
                                    webbrowser.open(google_url)
                                    return self.get_information_for_topic(topic, context)
            return response_text
        except Exception as e:
            logging.error(f"Error in Gemini general response: {str(e)}")
            return "I didn't understand that. Could you try asking in a different way?"
    def clean_response(self, text):
        """Clean up response text by removing special characters and limiting to 2-3 sentence"""
        cleaned = text.replace('**', '').replace('*', '').replace('~', '').replace('`', '').replace('#', '')
        sentences = cleaned.split('.')
        if len(sentences) > 1:
            first_sentence = sentences[0].strip()
            if len(first_sentence) > 5:
                return first_sentence + '.'
        cleaned = ' '.join(cleaned.split())
        return cleaned
    def open_application(self, app_name):
        """
        Open applications with Windows search as primary method 
        and other methods as fallbacks
        """
        try:
            app_name_clean = app_name.lower().strip()
            app_name_clean = re.sub(r'\b(me|the|to|for me|for)\b', '', app_name_clean)
            is_another_instance = any(word in app_name_clean for word in ['another', 'new', 'second'])
            app_name_clean = app_name_clean.replace('another', '').replace('new', '').replace('second', '').strip()
            logging.info(f"Attempting to open app: '{app_name_clean}'")
            special_apps = {
                'power bi': 'Power BI Desktop',
                'powerbi': 'Power BI Desktop',
                'word': 'Word',
                'excel': 'Excel',
                'powerpoint': 'PowerPoint',
                'power point': 'PowerPoint',
                'outlook': 'Outlook',
                'access': 'Access',
                'publisher': 'Publisher',
                'onenote': 'OneNote',
                'teams': 'Microsoft Teams',
                'edge': 'Microsoft Edge',
                'chrome': 'Google Chrome',
                'firefox': 'Firefox',
                'visual studio': 'Visual Studio',
                'visual studio code': 'VS Code',
                'vs code': 'VS Code',
                'cmd': 'Command Prompt',
                'command prompt': 'Command Prompt',
                'control panel': 'Control Panel',
                'task manager': 'Task Manager'
            }
            search_term = app_name_clean
            if app_name_clean in special_apps:
                search_term = special_apps[app_name_clean]
            message = f"Opening {search_term}"
            if is_another_instance:
                message = f"Opening another instance of {search_term}"
            self.speak(message)
            try:
                logging.info(f"Trying Windows search for: {search_term}")
                if self.search_and_click_first_result(search_term):
                    return True
            except Exception as e:
                logging.error(f"Error using Windows search: {str(e)}")
            app_mapping = {
                'notepad': 'notepad.exe',
                'calculator': 'calc.exe',
                'paint': 'mspaint.exe',
                'chrome': 'chrome.exe',
                'edge': 'msedge.exe',
                'firefox': 'firefox.exe',
                'word': 'winword.exe',
                'excel': 'excel.exe',
                'powerpoint': 'powerpnt.exe',
                'powerbi': 'PBIDesktop.exe',
                'power bi': 'PBIDesktop.exe',
                'power point': 'powerpnt.exe',
                'explorer': 'explorer.exe',
                'settings': 'ms-settings:',
                'control panel': 'control.exe',
                'task manager': 'taskmgr.exe',
                'cmd': 'cmd.exe',
                'command prompt': 'cmd.exe',
                'browser': 'start https://www.google.com',
                'google': 'start https://www.google.com',
                'youtube': 'start https://www.youtube.com',
                'wikipedia': 'start https://www.wikipedia.org'
            }
            executable = None
            app_name_display = app_name_clean
            if app_name_clean in app_mapping:
                executable = app_mapping[app_name_clean]
                app_name_display = app_name_clean
            else:
                for known_app, exec_path in app_mapping.items():
                    if known_app in app_name_clean or app_name_clean in known_app:
                        executable = exec_path
                        app_name_display = known_app
                        break
            if executable:
                try:
                    logging.info(f"Trying direct execution: {executable}")
                    if executable.startswith('start '):
                        os.system(executable)
                    else:
                        subprocess.Popen(executable)
                    return True
                except Exception as e:
                    logging.error(f"Error opening application with direct path: {str(e)}")
            try:
                logging.info(f"Trying Run command: {app_name_clean}")
                subprocess.Popen(['cmd.exe', '/c', 'start', app_name_clean])
                return True
            except Exception as e:
                logging.error(f"Error opening application with Run command: {str(e)}")
                self.speak(f"I couldn't find {app_name_clean}")
                return False
        except Exception as e:
            logging.error(f"Error in open_application: {str(e)}")
            self.speak(f"I encountered an error trying to open {app_name}")
            return False
    def search_and_click_first_result(self, search_term):
        """Ultra-fast Windows search and click"""
        try:
            pyautogui.PAUSE = 0.1
            pyautogui.MINIMUM_DURATION = 0
            pyautogui.hotkey('win', 's')
            time.sleep(0.1)
            pyautogui.write(search_term, interval=0.01)
            time.sleep(0.2)
            pyautogui.press('enter')
            return True
        except Exception as e:
            logging.error(f"Search error: {str(e)}")
            self.speak("Search error")
            return False
    def cleanup(self):
        """Clean up resources before closing the application"""
        try:
            if hasattr(self, 'visualizer'):
                self.visualizer.cleanup()
            if hasattr(self, 'voice_thread') and self.voice_thread and self.voice_thread.is_alive():
                self.voice_thread.join(timeout=0.1)
            if hasattr(self, 'listening_thread') and self.listening_thread and self.listening_thread.is_alive():
                self.listening_thread.join(timeout=0.1)
        except Exception as e:
            logging.error(f"Error during cleanup: {str(e)}")
    def get_information_for_topic(self, topic, context=""):
        """Get relevant information for a topic from Gemini"""
        try:
            system_prompt = """
            You are ANIS, an AI assistant providing concise, accurate information. 
            The user has asked for information about a specific topic.
            VERY IMPORTANT RULES:
            1. Keep your response brief and to the point.
            2. Do not use markdown formatting (no **, *, ~, _, etc.)
            3. Do not use bullet points or numbered lists.
            4. Write in plain text with short, clear sentences.
            5. Focus only on the most essential facts.
            6. Limit your response to 3-4 short sentences maximum.
            7. Do not include introductions or conclusions.
            8. Avoid phrases like "here are key points" or "in summary".
            9. do not include hear are the search results for in the response
            10.you should focus that user asked for information about the topic by what is ,which,who,like this or for to process command
            If you're not sure about specific details, if user ask about information instead of  you can use this for response and folow this:
                - answer should be in two lines for any query
                - Use simple language
                - Avoid special characters or formatting
                - Use proper punctuation
                - Keep responses brief and direct
                -your name is ANIS AI integrated with Operation System With Help Of Gemini and under Developing by 
                Samiulla from 26 August 2024 don't  tell this in beginning of your response you should tell only when someone ask you name
                - No emojis or symbols
                - answer in a friendly manner
                - answer like a human
                - answer in a way that is easy to understand
                - answer in a way that is not too formal
                - answer in a way that is not too technical
                - if need techinical answer in techincal terms
                - answer should be short like 3 lines and concise
                - avoid geographical deep terms
                - try to answer in 3 lines simple and short
                - some of the time you can answer in a funny way
            """
            if context:
                user_prompt = f"Provide a very concise summary about {topic}. Context of user's learning journey: {context}"
            else:
                user_prompt = f"Provide a very concise summary about {topic}."
            response = gemini_model.generate_content([system_prompt, user_prompt])
            clean_response = self.clean_response(response.text.strip())
            self.last_topic = topic
            self.log_interaction(f"Get information about {topic}", "Provided information summary")
            if hasattr(self, 'update_command_flow'):
                self.update_command_flow(
                    task_type="information",
                    subject=topic,
                    command_type="key_points",
                    qualifiers={"mode": "educational"}
                )
            return clean_response
        except Exception as e:
            logging.error(f"Error getting information: {str(e)}")
            return f"I couldn't find information about {topic}. Would you like me to search the web for it?"
    def get_specific_topic_information(self, topic, specific_query):
        """Get information about a specific aspect of a topic"""
        try:
            system_prompt = """
            You are ANIS, an AI assistant providing concise, accurate information. 
            The user has asked a specific question about a topic.
            VERY IMPORTANT RULES:
            1. Keep your response brief and to the point.
            2. Do not use markdown formatting (no **, *, ~, _, etc.)
            3. Do not use bullet points or numbered lists.
            4. Write in plain text with short, clear sentences.
            5. Focus only on answering the specific question asked.
            6. Limit your response to 2-3 short sentences maximum.
            7. Do not include introductions or conclusions.
            8. Avoid phrases like "here are key points" or "in summary".
            9. 10.you should focus that user asked for information about the topic by what is ,which,who,like this or for to process command
            If you're not sure about specific details,if user ask about information instead of  you can use this for response and folow this:
                - answer should be in two lines for any query
                - Use simple language
                - Avoid special characters or formatting
                - Use proper punctuation
                - Keep responses brief and direct
                -your name is ANIS AI integrated with Operation System With Help Of Gemini and under Developing by 
                Samiulla from 26 August 2024 don't  tell this in beginning of your response you should tell only when someone ask you name
                - No emojis or symbols
                - answer in a friendly manner
                - answer like a human
                - answer in a way that is easy to understand
                - answer in a way that is not too formal
                - answer in a way that is not too technical
                - if need techinical answer in techincal terms
                - answer should be short like 2 lines and concise
                - avoid geographical deep terms
                - try to answer in 2 lines simple and short
                - some of the time you can answer in a funny way
            """
            user_prompt = f"Answer this specific question: {specific_query}"
            response = gemini_model.generate_content([system_prompt, user_prompt])
            clean_response = self.clean_response(response.text.strip())
            self.last_topic = topic
            self.log_interaction(specific_query, "Provided specific information")
            if hasattr(self, 'update_command_flow'):
                self.update_command_flow(
                    task_type="information",
                    subject=topic,
                    command_type="specific_question",
                    qualifiers={"query_type": "types" if "type" in specific_query.lower() else "general"}
                )
            return clean_response
        except Exception as e:
            logging.error(f"Error getting specific information: {str(e)}")
            return f"I couldn't find specific information about that aspect of {topic}."
    def check_system_health(self):
        try:
            cpu_percent = psutil.cpu_percent(interval=1)
            memory = psutil.virtual_memory()
            memory_percent = memory.percent
            disk = psutil.disk_usage('/')
            disk_percent = disk.percent
            health_report = (
                f"System Health Report:\n"
                f"CPU Usage: {cpu_percent}%\n"
                f"Memory Usage: {memory_percent}%\n"
                f"Disk Usage: {disk_percent}%"
            )
            self.speak(health_report)
            logging.info(health_report)
            if cpu_percent > 90 or memory_percent > 90 or disk_percent > 90:
                self.speak("Warning: System resources are running very high!")
        except Exception as e:
            self.speak("Unable to check system health")
            logging.error(f"System health check error: {str(e)}")
    def get_notepad_content(self):
        """Get content from active notepad window"""
        try:
            original_clipboard = self.window.clipboard_get()
            pyautogui.hotkey('ctrl', 'a')
            time.sleep(0.1)
            pyautogui.hotkey('ctrl', 'c')
            time.sleep(0.1)
            text = self.window.clipboard_get()
            self.window.clipboard_clear()
            self.window.clipboard_append(original_clipboard)
            return text.strip()
        except Exception as e:
            logging.error(f"Error getting notepad content: {str(e)}")
            return None
    def generate_filename_from_content(self, content):
        """Generate appropriate filename by analyzing content using Gemini"""
        try:
            import google.generativeai as genai
            api_key = self.load_setting('gemini_api_key')
            model_name = self.load_setting("gemini_model")
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel(model_name)
            prompt = f"""Analyze this text content and generate a short, descriptive filename (1 words, separated by underscores).
            The filename should reflect the main topic or purpose of the content.
            Only return the filename, no other text.
            Content:
            {content[:500]}...
            """
            response = model.generate_content(prompt)
            filename = response.text.strip()
            timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
            return f"{filename}.txt"
        except Exception as e:
            logging.error(f"Error generating filename: {str(e)}")
            return None
    def show_email_window(self):
        """Open the email window"""
        try:
            EmailWindow(self.root)
        except Exception as e:
            self.speak("Sorry, I encountered an error while opening the email window")
            logging.error(f"Email window error: {str(e)}")
    def sanitize_filename(self, filename):
        """Ensure filename is valid by removing/replacing invalid characters"""
        import re
        filename = filename.replace('\n', ' ').replace('\t', ' ')
        filename = re.sub(r'\s+', '_', filename.strip())
        invalid_chars = r'[<>:"/\\|?*\'\[\]{}()!@#$%^&+=;,.]'
        filename = re.sub(invalid_chars, '_', filename)
        filename = re.sub(r'_+', '_', filename)
        filename = filename.strip('_')
        filename = filename.lower()
        if not filename:
            timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
            return f"note_{timestamp}"
        max_length = 255
        if len(filename) > max_length:
            name, ext = os.path.splitext(filename)
            filename = name[:max_length-len(ext)] + ext
        return filename
    def show_translator_window(self):
        """Open the translator window"""
        try:
            translator = TranslatorWindow(self.root)
            self.root.wait_window(translator.window)
            self.start_button.configure(text="◼", bg='#27ae60')
            self.status_label.configure(text="Listening...")
        except Exception as e:
            self.speak("Sorry, I encountered an error with the translator")
            logging.error(f"Translator window error: {str(e)}")
            self.start_button.configure(text="", bg='#27ae60')
            self.status_label.configure(text="Listening...")
    def take_screenshot(self):
        try:
            screenshot_dir = "screenshots"
            if not os.path.exists(screenshot_dir):
                os.makedirs(screenshot_dir)
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            screenshot_path = os.path.join(screenshot_dir, f"screenshot_{timestamp}.png")
            screenshot = ImageGrab.grab()
            screenshot.save(screenshot_path)
            self.speak(f"Screenshot saved")
            return screenshot_path
        except Exception as e:
            self.speak("Failed to take screenshot")
            logging.error(f"Screenshot error: {str(e)}")
            return None
    def show_summarizer_window(self):
        """Open the summarizer window"""
        summarizer = SummarizeWindow(self.root)
        self.root.wait_window(summarizer.window)
    def handle_keyboard_shortcuts(self):
        try:
            keyboard.add_hotkey('ctrl+shift+a', self.start_listening)
            keyboard.add_hotkey('ctrl+shift+x', self.exit_assistant)
            keyboard.add_hotkey('ctrl+shift+s', self.take_screenshot)
            keyboard.wait()
        except Exception as e:
            logging.error(f"Keyboard shortcut error: {str(e)}")
    def log_interaction(self, user_input, assistant_response):
        try:
            timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            log_entry = (
                f"Time: {timestamp}\n"
                f"User: {user_input}\n"
                f"Assistant: {assistant_response}\n"
                f"{'-'*50}\n"
            )
            with open("interaction_log.txt", "a", encoding="utf-8") as log_file:
                log_file.write(log_entry)
        except Exception as e:
            logging.error(f"Logging error: {str(e)}")
    def handle_errors(self, error, context=""):
        """Generic error handler for the assistant"""
        error_message = f"Error in {context}: {str(error)}"
        logging.error(error_message)
        user_message = "I encountered an error. Please try again."
        self.speak(user_message)
        self.update_status(False)
        self.log_interaction(context, user_message)
    def process_text_command(self, text_input):
        """Process text commands without voice input"""
        try:
            if text_input:
                self.insert_with_timestamp(f"You: {text_input}", "user_msg")
                return self.process_command(text_input.lower())
            return False
        except Exception as e:
            self.handle_errors(e, "text command processing")
            return False
    def setup_keyboard_shortcuts(self):
        """Setup keyboard shortcuts for the assistant"""
        try:
            keyboard_thread = threading.Thread(
                target=self.handle_keyboard_shortcuts,
                daemon=True
            )
            keyboard_thread.start()
        except Exception as e:
            self.handle_errors(e, "keyboard shortcut setup")
    def speak(self, text, emotion=None):
        try:
            if hasattr(self, 'wave_vis') and self.wave_vis.is_active:
                self.wave_vis.speaking_animation()
            self.insert_with_timestamp(f"A.N.I.S: {text}", "assistant_msg")
            engine.say(text)
            engine.runAndWait()
            if hasattr(self, 'wave_vis') and self.wave_vis.is_active:
                self.wave_vis.stop_listening_animation()
        except Exception as e:
            logging.error(f"Error in speak method: {str(e)}")
    def take_command(self):
        r = sr.Recognizer()
        with sr.Microphone() as source:
            try:
                if hasattr(self, 'wave_vis') and self.wave_vis.is_active:
                    self.wave_vis.start_listening_animation()
                self.status_label.configure(text="Listening...")
                self.root.update()
                r.adjust_for_ambient_noise(source, duration=1)
                r.dynamic_energy_threshold = True
                r.pause_threshold = 0.8
                r.non_speaking_duration = 0.8
                r.phrase_threshold = 0.3
                audio = r.listen(source, timeout=10 ,phrase_time_limit=8)
                self.status_label.configure(text="Recognizing...")
                self.root.update()
                query = r.recognize_google(audio, language='en-in')
                self.insert_with_timestamp(f"You: {query}", "user_msg")
                return query.lower()
            except sr.WaitTimeoutError:
                return "none"
            except sr.UnknownValueError:
                return "none"
            except sr.RequestError:
                self.speak("Sorry, there was an error with the speech recognition service")
                return "none"
            except Exception as e:
                logging.error(f"Error in speech recognition: {str(e)}")
                return "none"
            finally:
                if hasattr(self, 'wave_vis') and self.wave_vis.is_active:
                    self.wave_vis.stop_listening_animation()
                self.status_label.configure(text="Not Listening")
    def get_website_url(self, website_name):
        """Get website URL using the URL database or Gemini API."""
        url = get_url(website_name)
        if url:
            return url
        try:
            prompt = f"What is the login URL for {website_name}? Only return the direct URL, no other text or explanation."
            response = self.gemini_model.generate_content(prompt)
            url = response.text.strip()
            add_url(website_name, url)
            return url
        except Exception as e:
            logging.error(f"Error getting website URL: {str(e)}")
            return None
    def auto_login(self, website):
        """Auto login to the specified website."""
        try:
            retry_attempts = 3
            for attempt in range(retry_attempts):
                try:
                    website_url = self.get_website_url(website)
                    if website_url:
                        break
                    else:
                        if attempt < retry_attempts - 1:
                            self.speak("Could not find login URL. Retrying...")
                            time.sleep(2)
                        else:
                            self.speak("Could not find login URL after multiple attempts.")
                            return
                except Exception as e:
                    logging.error(f"Error getting website URL: {str(e)}")
                    if attempt < retry_attempts - 1:
                        self.speak("Connection error. Retrying...")
                        time.sleep(2)
                    else:
                        self.speak("Could not connect after multiple attempts.")
                        return
            conn = sqlite3.connect('auto_login.db')
            cursor = conn.cursor()
            cursor.execute('SELECT email, password FROM credentials WHERE website = ?', (website.lower(),))
            result = cursor.fetchone()
            conn.close()
            if result:
                email, password = result
                driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()))
                driver.get(website_url)
                try:
                    wait = WebDriverWait(driver, 0.1)
                    email_selectors = [
                        (By.NAME, "email"),
                        (By.NAME, "username"),
                        (By.ID, "email"),
                        (By.ID, "username"),
                        (By.CSS_SELECTOR, "input[type='email']"),
                        (By.CSS_SELECTOR, "input[type='text']"),
                        (By.XPATH, "//input[@placeholder='Email' or @placeholder='Username' or contains(@placeholder, 'email') or contains(@placeholder, 'username')]"),
                    ]
                    password_selectors = [
                        (By.NAME, "password"),
                        (By.ID, "password"),
                        (By.CSS_SELECTOR, "input[type='password']"),
                        (By.XPATH, "//input[@placeholder='Password' or contains(@placeholder, 'password')]"),
                    ]
                    next_selectors = [
                        (By.XPATH, "//button[contains(text(), 'Next') or contains(text(), 'Continue')]"),
                        (By.XPATH, "//div[contains(text(), 'Next') or contains(text(), 'Continue')]"),
                        (By.XPATH, "//span[contains(text(), 'Next') or contains(text(), 'Continue')]"),
                        (By.CSS_SELECTOR, "button[type='submit']"),
                        (By.NAME, "next"),
                        (By.ID, "next"),
                        (By.ID, "identifierNext"),
                        (By.ID, "passwordNext")
                    ]
                    login_selectors = [
                        (By.CSS_SELECTOR, "button[type='submit']"),
                        (By.NAME, "login"),
                        (By.ID, "login"),
                        (By.XPATH, "//button[contains(text(), 'Login') or contains(text(), 'Sign in')]"),
                        (By.XPATH, "//input[@type='submit' or @value='Login' or @value='Sign in']")
                    ]
                    def find_and_click_element(selectors, timeout=1):
                        """Find and click an element using multiple selectors"""
                        for by, selector in selectors:
                            try:
                                element = wait.until(EC.element_to_be_clickable((by, selector)))
                                element.click()
                                return element
                            except:
                                continue
                        return None
                    def find_and_fill_element(selectors, text, timeout=1):
                        """Find and fill an element using multiple selectors"""
                        for by, selector in selectors:
                            try:
                                element = wait.until(EC.presence_of_element_located((by, selector)))
                                element.clear()
                                element.send_keys(text)
                                return element
                            except:
                                continue
                        return None
                    email_field = find_and_fill_element(email_selectors, email)
                    if not email_field:
                        self.speak("Could not find email field")
                        driver.quit()
                        return
                    next_button = find_and_click_element(next_selectors)
                    if next_button:
                        time.sleep(0.1)
                    password_field = find_and_fill_element(password_selectors, password)
                    if not password_field:
                        time.sleep(5)
                        password_field = find_and_fill_element(password_selectors, password)
                        if not password_field:
                            self.speak("Could not find password field. Please check your login process.")
                            driver.quit()
                            return
                    login_button = find_and_click_element(login_selectors)
                    if not login_button:
                        next_button = find_and_click_element(next_selectors)
                        if not next_button:
                            driver.quit()
                            return
                    time.sleep(7)
                    error_texts = ['invalid', 'incorrect', 'failed', 'error', 'wrong']
                    page_source = driver.page_source.lower();
                    if any(error_text in page_source for error_text in error_texts):
                        self.speak("Login in")
                        self.show_password_manager(website)
                        driver.quit()
                    else:
                        self.speak("Login successful!")
                except Exception as e:
                    logging.error(f"Error during login process: {str(e)}")
                    self.speak("")
                    driver.quit()
            else:
                self.speak(f"No credentials found for {website}. Opening password manager to add them.")
                self.show_password_manager(website)
        except Exception as e:
            logging.error(f"Error during auto login: {str(e)}")
            self.speak("An error occurred during the auto login process.")
    def show_password_manager(self, website=None):
        """Show password manager window."""
        try:
            password_window = tk.Toplevel(self.root)
            password_window.title("Password Manager")
            password_window.geometry("400x300")
            password_window.configure(bg=self.themes[self.current_theme]["bg"])
            tk.Label(
                password_window, 
                text="Website:",
                bg=self.themes[self.current_theme]["bg"],
                fg=self.themes[self.current_theme]["fg"]
            ).pack(pady=5)
            website_entry = tk.Entry(password_window, width=30)
            website_entry.pack(pady=5)
            if website:
                website_entry.insert(0, website)
            tk.Label(
                password_window,
                text="Email:",
                bg=self.themes[self.current_theme]["bg"],
                fg=self.themes[self.current_theme]["fg"]
            ).pack(pady=5)
            email_entry = tk.Entry(password_window, width=30)
            email_entry.pack(pady=5)
            tk.Label(
                password_window,
                text="Password:",
                bg=self.themes[self.current_theme]["bg"],
                fg=self.themes[self.current_theme]["fg"]
            ).pack(pady=5)
            password_entry = tk.Entry(password_window, width=30, show="*")
            password_entry.pack(pady=5)
            def save_credentials():
                try:
                    website = website_entry.get().strip().lower()
                    email = email_entry.get().strip()
                    password = password_entry.get().strip()
                    if website and email and password:
                        conn = sqlite3.connect('auto_login.db')
                        cursor = conn.cursor()
                        cursor.execute('''
                            INSERT OR REPLACE INTO credentials (website, email, password)
                            VALUES (?, ?, ?)
                        ''', (website, email, password))
                        conn.commit()
                        conn.close()
                        self.speak("Credentials saved successfully!")
                        password_window.destroy()
                    else:
                        self.speak("Please fill in all fields.")
                except Exception as e:
                    logging.error(f"Error saving credentials: {str(e)}")
                    self.speak("Error saving credentials.")
            save_btn = tk.Button(
                password_window,
                text="Save Credentials",
                command=save_credentials,
                bg=self.themes[self.current_theme]["accent"],
                fg="white"
            )
            save_btn.pack(pady=20)
        except Exception as e:
            logging.error(f"Error showing password manager: {str(e)}")
            self.speak("Error opening password manager.")
    def show_chat_window(self):
        """Show the Gemini chat window"""
        try:
            try:
                api_key = self.load_setting('gemini_api_key')
                model_name = self.load_setting("gemini_model")
                import google.generativeai as genai
                genai.configure(api_key=api_key)
                model = genai.GenerativeModel(model_name)
            except ImportError:
                self.speak("Please install the google-generativeai package first")
                messagebox.showerror("Missing Dependency", 
                    "Please install the required package:\npip install google-generativeai")
                return
            chat_window = GeminiChatWindow(self)
            chat_window.window.transient(self.root)
            chat_window.window.grab_set()
        except Exception as e:
            self.handle_errors(e, "Error opening chat window")
            self.speak("Sorry, I encountered an error while opening the chat window.")
    def parse_duration(self, duration_str):
        """Parse duration string into seconds"""
        try:
            duration_str = duration_str.lower()
            total_seconds = 0
            if duration_str.startswith('after'):
                duration_str = duration_str[5:].strip()
            parts = duration_str.split()
            number = None
            for i, part in enumerate(parts):
                if part.isdigit():
                    number = int(part)
                    if i + 1 < len(parts):
                        unit = parts[i + 1].lower()
                        if 'hour' in unit:
                            total_seconds += number * 3600
                        elif 'min' in unit:
                            total_seconds += number * 60
                        elif 'sec' in unit:
                            total_seconds += number
            return total_seconds
        except Exception as e:
            logging.error(f"Error parsing duration: {str(e)}")
            return 0
    def show_task_scheduler(self, preset_time=None, preset_tasks=None):
        """Show task scheduler window with improved task execution"""
        scheduler_window = tk.Toplevel(self.root)
        scheduler_window.title("Task Scheduler")
        scheduler_window.geometry("600x700")
        scheduler_window.configure(bg=self.themes[self.current_theme]["bg"])
        scheduler_window.transient(self.root)
        scheduler_window.grab_set()
        main_frame = Frame(
            scheduler_window,
            bg=self.themes[self.current_theme]["bg"],
            padx=20,
            pady=20
        )
        main_frame.pack(fill='both', expand=True)
        Label(
            main_frame,
            text="Schedule Tasks",
            font=("Segoe UI", 16, "bold"),
            bg=self.themes[self.current_theme]["bg"],
            fg=self.themes[self.current_theme]["accent"]
        ).pack(pady=(0, 20))
        time_frame = Frame(main_frame, bg=self.themes[self.current_theme]["bg"])
        time_frame.pack(fill='x', pady=10)
        Label(
            time_frame,
            text="Time until tasks (seconds):",
            font=("Segoe UI", 12),
            bg=self.themes[self.current_theme]["bg"],
            fg=self.themes[self.current_theme]["fg"]
        ).pack(side='left', padx=(0, 10))
        time_var = tk.StringVar(value=str(preset_time) if preset_time else "60")
        time_entry = ttk.Entry(
            time_frame,
            textvariable=time_var,
            font=("Segoe UI", 12),
            width=10
        )
        time_entry.pack(side='left')
        task_frame = Frame(main_frame, bg=self.themes[self.current_theme]["bg"])
        task_frame.pack(fill='both', expand=True, pady=20)
        Label(
            task_frame,
            text="Selected Tasks:",
            font=("Segoe UI", 12),
            bg=self.themes[self.current_theme]["bg"],
            fg=self.themes[self.current_theme]["fg"]
        ).pack(anchor='w')
        task_list_frame = Frame(task_frame, bg=self.themes[self.current_theme]["bg"])
        task_list_frame.pack(fill='both', expand=True, pady=10)
        task_listbox = tk.Listbox(
            task_list_frame,
            bg=self.themes[self.current_theme]["text_bg"],
            fg=self.themes[self.current_theme]["fg"],
            font=("Segoe UI", 11),
            selectmode='multiple',
            relief='flat',
            bd=1,
            highlightthickness=1,
            highlightbackground=self.themes[self.current_theme]["accent"]
        )
        task_listbox.pack(side='left', fill='both', expand=True)
        scrollbar = ttk.Scrollbar(task_list_frame, orient='vertical', command=task_listbox.yview)
        scrollbar.pack(side='right', fill='y')
        task_listbox.configure(yscrollcommand=scrollbar.set)
        predefined_tasks = [
            "Check system health",
            "Take screenshot",
            "Create backup",
            "Open notes",
            "Open contact manager",
            "Data analysis",
            "Save notes",
            "Open Word",
            "Open Excel",
            "Open PowerPoint"
        ]
        custom_frame = Frame(main_frame, bg=self.themes[self.current_theme]["bg"])
        custom_frame.pack(fill='x', pady=20)
        Label(
            custom_frame,
            text="Enter custom commands or select files:",
            font=("Segoe UI", 12),
            bg=self.themes[self.current_theme]["bg"],
            fg=self.themes[self.current_theme]["fg"]
        ).pack(anchor='w')
        custom_entry = ttk.Entry(
            custom_frame,
            font=("Segoe UI", 12),
            width=40
        )
        custom_entry.pack(fill='x', pady=5)
        if preset_tasks:
            seen_tasks = set()
            unique_tasks = []
            initial_tasks = [t.strip() for t in preset_tasks.split(',') if t.strip()]
            for task in initial_tasks:
                sub_tasks = []
                if task.lower().startswith('open '):
                    parts = task.lower().split('open ')
                    for part in parts:
                        if part.strip():
                            sub_tasks.append(f"open {part.strip()}")
                else:
                    sub_tasks.append(task)
                for sub_task in sub_tasks:
                    normalized_task = sub_task.lower().strip()
                    if normalized_task not in seen_tasks:
                        seen_tasks.add(normalized_task)
                        if sub_task.lower().startswith('open '):
                            words = sub_task.split(' ', 1)
                            if len(words) > 1:
                                sub_task = f"Open {words[1]}"
                        unique_tasks.append(sub_task)
            for task in unique_tasks:
                if task not in task_listbox.get(0, 'end'):
                    task_listbox.insert('end', task)
                    last_index = task_listbox.size() - 1
                    task_listbox.selection_set(last_index)
            custom_entry.delete(0, 'end')
            custom_entry.insert(0, ','.join(unique_tasks))
        for task in predefined_tasks:
            if task not in task_listbox.get(0, 'end'):
                task_listbox.insert('end', task)
        file_paths = []
        def browse_file():
            file_path = filedialog.askopenfilename(
                title="Select a file to schedule",
                filetypes=(
                    ("All Files", "*.*"),
                    ("Text Files", "*.txt"),
                    ("Word Documents", "*.doc;*.docx"),
                    ("Excel Files", "*.xls;*.xlsx"),
                    ("PDF Files", "*.pdf")
                )
            )
            if file_path:
                file_paths.append(file_path)
                filename = os.path.basename(file_path)
                task_listbox.insert('end', f"Open {filename}")
                task_listbox.selection_set('end')
        Button(
            custom_frame,
            text="Browse File",
            command=browse_file,
            bg=self.themes[self.current_theme]["accent"],
            fg="#ffffff",
            font=('Segoe UI', 10),
            relief='flat'
        ).pack(side='right', padx=5)
        def schedule_tasks():
            try:
                seconds = int(time_var.get())
                selected_indices = task_listbox.curselection()
                if seconds <= 0:
                    messagebox.showerror("Error", "Please enter a positive number of seconds")
                    return
                tasks = []
                seen_tasks = set()
                custom_task = custom_entry.get().strip()
                if custom_task:
                    custom_parts = custom_task.split(',')
                    for part in custom_parts:
                        task = part.strip()
                        if task:
                            normalized_task = task.lower()
                            if normalized_task not in seen_tasks:
                                seen_tasks.add(normalized_task)
                                tasks.append(task)
                if selected_indices:
                    for index in selected_indices:
                        task_text = task_listbox.get(index)
                        normalized_task = task_text.lower()
                        if task_text.startswith("Open "):
                            file_idx = index - (task_listbox.size() - len(file_paths))
                            if 0 <= file_idx < len(file_paths):
                                file_task = f"open_file:{file_paths[file_idx]}"
                                if file_task.lower() not in seen_tasks:
                                    seen_tasks.add(file_task.lower())
                                    tasks.append(file_task)
                        else:
                            if normalized_task not in seen_tasks:
                                seen_tasks.add(normalized_task)
                                tasks.append(task_text)
                if not tasks:
                    messagebox.showerror("Error", "Please select tasks or enter custom commands")
                    return
                combined_tasks = ','.join(tasks)
                logging.info(f"Scheduling combined tasks: {combined_tasks}")
                db_path = "C:\\Users\\HP\\Desktop\\Ai\\task.db"
                conn = sqlite3.connect(db_path)
                cursor = conn.cursor()
                cursor.execute('''SELECT id, execution_time FROM scheduled_tasks 
                                WHERE task = ? AND status = 'scheduled' AND completed = 0''', 
                                (combined_tasks,))
                existing_task = cursor.fetchone()
                if existing_task:
                    task_id, execution_time = existing_task
                    response = messagebox.askyesno(
                        "Duplicate Task", 
                        f"This exact task is already scheduled to run at {execution_time}.\nDo you want to schedule it again?"
                    )
                    if not response:
                        conn.close()
                        return
                conn.close()
                task_id = self.store_scheduled_task(seconds, combined_tasks)
                if task_id:
                    execution_time = datetime.datetime.now() + datetime.timedelta(seconds=seconds)
                    timer_key = f"task_{task_id}"
                    if hasattr(self, '_task_timers') and timer_key in self._task_timers:
                        old_timer = self._task_timers[timer_key]
                        if old_timer.is_alive():
                            logging.info(f"Cancelling existing timer for task ID {task_id}")
                            old_timer.cancel()
                    if not hasattr(self, '_task_timers'):
                        self._task_timers = {}
                    timer = threading.Timer(seconds, lambda t=combined_tasks, tid=task_id: self.execute_task(t, tid))
                    timer.daemon = True
                    self._task_timers[timer_key] = timer
                    timer.start()
                    time_str = execution_time.strftime("%I:%M:%S %p")
                    messagebox.showinfo("Success", f"Scheduled {len(tasks)} task(s) to run at {time_str}")
                    scheduler_window.destroy()
                else:
                    messagebox.showerror("Error", "Failed to store scheduled tasks")
            except ValueError:
                messagebox.showerror("Error", "Please enter a valid number of seconds")
            except Exception as e:
                logging.error(f"Error scheduling tasks: {str(e)}")
                messagebox.showerror("Error", f"Failed to schedule tasks: {str(e)}")
        Button(
            main_frame,
            text="Schedule Tasks",
            command=schedule_tasks,
            bg=self.themes[self.current_theme]["accent"],
            fg="#ffffff",
            font=('Segoe UI', 12, 'bold'),
            width=15,
            height=1,
            relief='flat',
            cursor='hand2'
        ).pack(pady=20)
    def execute_task(self, task_name, task_id):
        """Execute a scheduled task with support for file opening and multiple tasks"""
        try:
            db_path = "C:\\Users\\HP\\Desktop\\Ai\\task.db"
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            cursor.execute('''SELECT status, completed FROM scheduled_tasks 
                            WHERE id = ?''', (task_id,))
            task_info = cursor.fetchone()
            if not task_info:
                logging.error(f"Task ID {task_id} not found in database")
                return
            status, completed = task_info
            if completed == 1 or status in ['executing', 'completed']:
                logging.warning(f"Task {task_id} has already been {status}. Skipping execution.")
                return
            if not hasattr(self, '_task_execution_locks'):
                self._task_execution_locks = {}
            if task_id in self._task_execution_locks:
                logging.warning(f"Task {task_id} is already being executed. Preventing duplicate execution.")
                return
            self._task_execution_locks[task_id] = True
            try:
                cursor.execute('''UPDATE scheduled_tasks 
                                SET status = 'executing' 
                                WHERE id = ?''', (task_id,))
                conn.commit()
                logging.info(f"Executing scheduled task: '{task_name}' (ID: {task_id})")
                if ',' in task_name:
                    tasks = [t.strip() for t in task_name.split(',') if t.strip()]
                    logging.info(f"Found {len(tasks)} tasks to execute: {tasks}")
                    executed_tasks = set()
                    for individual_task in tasks:
                        if individual_task in executed_tasks:
                            logging.info(f"Skipping duplicate task: '{individual_task}'")
                            continue
                        logging.info(f"Executing sub-task: '{individual_task}'")
                        executed_tasks.add(individual_task)
                        compound_commands = self._parse_compound_command(individual_task)
                        if len(compound_commands) > 1:
                            logging.info(f"Detected compound command with {len(compound_commands)} parts: {compound_commands}")
                            executed_commands = set()
                            for cmd in compound_commands:
                                if cmd in executed_commands:
                                    logging.info(f"Skipping duplicate command part: '{cmd}'")
                                    continue
                                executed_commands.add(cmd)
                                self.speak(f"task: {cmd}")
                                try:
                                    self._execute_direct_command(cmd)
                                    time.sleep(1.5)
                                except Exception as e:
                                    logging.error(f"Error executing compound command part '{cmd}': {str(e)}")
                                    self.speak(f"Error executing task: {cmd}")
                                    if cmd.startswith('open '):
                                        try:
                                            app_name = cmd.replace('open ', '', 1).strip()
                                            self.speak(f"Trying to search for {app_name} instead")
                                            self.search_and_click_first_result(app_name)
                                        except Exception as search_error:
                                            logging.error(f"Fallback search failed for '{app_name}': {str(search_error)}")
                            continue
                        if individual_task.startswith('open_file:'):
                            self._execute_file_task(individual_task)
                        else:
                            self.speak(f"task: {individual_task}")
                            try:
                                self._execute_direct_command(individual_task)
                                time.sleep(2)
                            except Exception as e:
                                logging.error(f"Error executing sub-task '{individual_task}': {str(e)}")
                                self.speak(f"Error executing task: {individual_task}")
                                if individual_task.startswith('open '):
                                    try:
                                        app_name = individual_task.replace('open ', '', 1).strip()
                                        self.speak(f"Trying to search for {app_name} instead")
                                        self.search_and_click_first_result(app_name)
                                    except Exception as search_error:
                                        logging.error(f"Fallback search failed for '{app_name}': {str(search_error)}")
                else:
                    compound_commands = self._parse_compound_command(task_name)
                    if len(compound_commands) > 1:
                        logging.info(f"Detected compound command with {len(compound_commands)} parts: {compound_commands}")
                        executed_commands = set()
                        for cmd in compound_commands:
                            if cmd in executed_commands:
                                logging.info(f"Skipping duplicate command part: '{cmd}'")
                                continue
                            executed_commands.add(cmd)
                            self.speak(f"Executing scheduled task: {cmd}")
                            try:
                                self._execute_direct_command(cmd)
                                time.sleep(1.5)
                            except Exception as e:
                                logging.error(f"Error executing compound command part '{cmd}': {str(e)}")
                                self.speak(f"Error executing task: {cmd}")
                                if cmd.startswith('open '):
                                    try:
                                        app_name = cmd.replace('open ', '', 1).strip()
                                        self.speak(f"Trying to search for {app_name} instead")
                                        self.search_and_click_first_result(app_name)
                                    except Exception as search_error:
                                        logging.error(f"Fallback search failed for '{app_name}': {str(search_error)}")
                    elif task_name.startswith('open_file:'):
                        self._execute_file_task(task_name)
                    else:
                        self.speak(f"task: {task_name}")
                        try:
                            self._execute_direct_command(task_name)
                        except Exception as e:
                            logging.error(f"Error executing task '{task_name}': {str(e)}")
                            self.speak(f"Error executing task: {task_name}")
                            if task_name.startswith('open '):
                                try:
                                    app_name = task_name.replace('open ', '', 1).strip()
                                    self.speak(f"search for {app_name} instead")
                                    self.search_and_click_first_result(app_name)
                                except Exception as search_error:
                                    logging.error(f"Fallback search failed for '{app_name}': {str(search_error)}")
                cursor.execute('''UPDATE scheduled_tasks 
                                SET status = 'completed', 
                                completed = 1 
                                WHERE id = ?''', (task_id,))
                conn.commit()
                if hasattr(self, '_task_timers'):
                    timer_key = f"task_{task_id}"
                    if timer_key in self._task_timers:
                        del self._task_timers[timer_key]
            finally:
                if hasattr(self, '_task_execution_locks') and task_id in self._task_execution_locks:
                    del self._task_execution_locks[task_id]
        except Exception as e:
            logging.error(f"Error executing task '{task_name}' (ID: {task_id}): {str(e)}")
            try:
                conn = sqlite3.connect(db_path)
                cursor = conn.cursor()
                cursor.execute('''UPDATE scheduled_tasks 
                                SET status = 'failed' 
                                WHERE id = ?''', (task_id,))
                conn.commit()
                if hasattr(self, '_task_execution_locks') and task_id in self._task_execution_locks:
                    del self._task_execution_locks[task_id]
            except Exception as inner_e:
                logging.error(f"Error updating failed task status: {str(inner_e)}")
        finally:
            if 'conn' in locals():
                conn.close()
    def _execute_file_task(self, task_name):
        """Helper method to execute file-opening tasks"""
        file_path = task_name.replace('open_file:', '', 1)
        if os.path.exists(file_path):
            try:
                if file_path.lower().endswith('.txt'):
                    subprocess.Popen(['notepad.exe', file_path])
                else:
                    os.startfile(file_path)
                self.speak(f"Opening {os.path.basename(file_path)}")
            except Exception as e:
                logging.error(f"Error opening file {file_path}: {str(e)}")
                self.speak(f"Error opening {os.path.basename(file_path)}")
    def store_scheduled_task(self, seconds, task):
        """Store scheduled task in database with enhanced tracking"""
        try:
            db_path = "C:\\Users\\HP\\Desktop\\Ai\\task.db"
            execution_time = datetime.datetime.now() + datetime.timedelta(seconds=seconds)
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            self.initialize_database()
            cursor.execute('''SELECT id FROM scheduled_tasks 
                             WHERE task = ? AND status = 'scheduled' AND completed = 0''', 
                             (task,))
            existing_task = cursor.fetchone()
            if existing_task:
                logging.warning(f"Duplicate task detected: '{task}'. Using existing task ID: {existing_task[0]}")
                conn.close()
                return existing_task[0]
            cursor.execute('''INSERT INTO scheduled_tasks 
                             (task, seconds, scheduled_time, execution_time, status)
                             VALUES (?, ?, datetime('now'), ?, 'scheduled')''', 
                             (task, seconds, execution_time.strftime('%Y-%m-%d %H:%M:%S')))
            task_id = cursor.lastrowid
            conn.commit()
            conn.close()
            logging.info(f"Task '{task}' scheduled for execution at {execution_time}")
            return task_id
        except Exception as e:
            logging.error(f"Error storing scheduled task: {str(e)}")
            raise
    def initialize_database(self):
        """Initialize the task database with proper schema"""
        try:
            db_path = "C:\\Users\\HP\\Desktop\\Ai\\task.db"
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            cursor.execute('''CREATE TABLE IF NOT EXISTS scheduled_tasks
                             (id INTEGER PRIMARY KEY AUTOINCREMENT,
                              task TEXT NOT NULL,
                              seconds INTEGER NOT NULL,
                              scheduled_time TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                              execution_time TIMESTAMP,
                              completed INTEGER DEFAULT 0,
                              status TEXT DEFAULT 'scheduled')''')
            conn.commit()
            conn.close()
            logging.info(f"Database initialized at {db_path}")
        except Exception as e:
            logging.error(f"Error initializing database: {str(e)}")
            raise
    def _execute_direct_command(self, command):
        """Execute a command directly, bypassing the scheduling logic.
        This ensures all commands are accessible to scheduled tasks."""
        try:
            command_execution_key = f"cmd_{hash(command)}_{int(time.time())}"
            if hasattr(self, '_recent_commands'):
                current_time = time.time()
                self._recent_commands = {k: v for k, v in self._recent_commands.items() 
                                        if current_time - v < 5}
                for cmd_key, timestamp in self._recent_commands.items():
                    if cmd_key.startswith(f"cmd_{hash(command)}_") and current_time - timestamp < 3:
                        logging.warning(f"Preventing duplicate execution of command: '{command}' (executed {current_time - timestamp:.1f}s ago)")
                        return False
            else:
                self._recent_commands = {}
            self._recent_commands[command_execution_key] = time.time()
            logging.info(f"Direct: '{command}'")
            compound_commands = self._parse_compound_command(command)
            if len(compound_commands) > 1:
                logging.info(f"Detected compound command with {len(compound_commands)} parts: {compound_commands}")
                for cmd in compound_commands:
                    logging.info(f"Executing part of compound command: '{cmd}'")
                    self._execute_direct_command(cmd)
                    time.sleep(1)
                return False
            if 'search youtube' in command or 'youtube search' in command:
                search_query = command.replace('search youtube', '').replace('youtube search', '').replace('youtube', '').strip()
                webbrowser.open(f'https://www.youtube.com/results?search_query={urllib.parse.quote(search_query)}')
                self.speak(f"Searching YouTube for {search_query}")
                return False
            if 'search for' in command or 'google search' in command or 'search google' in command:
                search_query = command
                for phrase in ['search for', 'google search', 'search google']:
                    search_query = search_query.replace(phrase, '').strip()
                if search_query:
                    search_url = f'https://www.google.com/search?q={urllib.parse.quote(search_query)}'
                    webbrowser.open(search_url)
                    self.speak(f"Searching Google for {search_query}")
                    return False
            if command.startswith('open '):
                app_name = command.replace('open ', '', 1).strip().lower()
                success = False
                if app_name == 'notepad':
                    subprocess.Popen(['notepad.exe'])
                    success = True
                elif app_name == 'calculator':
                    subprocess.Popen(['calc.exe'])
                    success = True
                elif app_name == 'paint':
                    subprocess.Popen(['mspaint.exe'])
                    success = True
                elif app_name == 'cmd' or app_name == 'command prompt':
                    subprocess.Popen(['cmd.exe'])
                    success = True
                elif app_name == 'control panel':
                    subprocess.Popen(['control.exe'])
                    success = True
                elif app_name == 'google':
                    webbrowser.open('https://www.google.com')
                    success = True
                elif app_name == 'youtube':
                    webbrowser.open('https://www.youtube.com')
                    success = True
                elif app_name == 'wikipedia':
                    webbrowser.open('https://www.wikipedia.org')
                    success = True
                elif app_name == 'chrome' or app_name == 'google chrome':
                    chrome_paths = [
                        'chrome.exe',
                        'C:\\Program Files\\Google\\Chrome\\Application\\chrome.exe',
                        'C:\\Program Files (x86)\\Google\\Chrome\\Application\\chrome.exe',
                        os.path.expanduser('~') + '\\AppData\\Local\\Google\\Chrome\\Application\\chrome.exe'
                    ]
                    for path in chrome_paths:
                        try:
                            subprocess.Popen([path])
                            success = True
                            logging.info(f"Successfully opened Chrome using path: {path}")
                            break
                        except Exception as chrome_error:
                            logging.debug(f"Failed to open Chrome with path {path}: {str(chrome_error)}")
                            continue
                elif app_name == 'firefox' or app_name == 'mozilla firefox':
                    firefox_paths = [
                        'firefox.exe',
                        'C:\\Program Files\\Mozilla Firefox\\firefox.exe',
                        'C:\\Program Files (x86)\\Mozilla Firefox\\firefox.exe'
                    ]
                    for path in firefox_paths:
                        try:
                            subprocess.Popen([path])
                            success = True
                            break
                        except:
                            continue
                elif app_name == 'edge' or app_name == 'microsoft edge':
                    edge_paths = [
                        'msedge.exe',
                        'C:\\Program Files (x86)\\Microsoft\\Edge\\Application\\msedge.exe',
                        'C:\\Program Files\\Microsoft\\Edge\\Application\\msedge.exe'
                    ]
                    for path in edge_paths:
                        try:
                            subprocess.Popen([path])
                            success = True
                            break
                        except:
                            continue
                elif app_name in ['word', 'excel', 'powerpoint']:
                    success = self.open_system_app(app_name)
                if success:
                    return False
                logging.info(f"Direct execution failed for '{app_name}', trying search and open")
                try:
                    search_success = self.search_and_click_first_result(app_name)
                    return False
                except Exception as search_error:
                    logging.error(f"Error in fallback search for '{app_name}': {str(search_error)}")
                    self.speak(f"I had trouble opening {app_name}")
                    return False
            if 'open notepad' in command:
                subprocess.Popen(['notepad.exe'])
                self.speak("Opening Notepad")
                return False
            if 'open calculator' in command:
                subprocess.Popen(['calc.exe'])
                self.speak("Opening Calculator")
                return False
            if 'open paint' in command:
                subprocess.Popen(['mspaint.exe'])
                self.speak("Opening Paint")
                return False
            if 'open cmd' in command or 'open command prompt' in command:
                subprocess.Popen(['cmd.exe'])
                self.speak("Opening Command Prompt")
                return False
            if 'open control panel' in command:
                subprocess.Popen(['control.exe'])
                self.speak("Opening Control Panel")
                return False
            if 'open google' in command:
                webbrowser.open('https://www.google.com')
                self.speak("Opening Google")
                return False
            if 'open youtube' in command:
                webbrowser.open('https://www.youtube.com')
                self.speak("Opening YouTube")
                return False
            if 'open wikipedia' in command:
                webbrowser.open('https://www.wikipedia.org')
                self.speak("Opening Wikipedia")
                return False
            if 'search and open' in command or 'find and open' in command:
                search_term = command
                for phrase in ['search and open', 'find and open']:
                    search_term = search_term.replace(phrase, '').strip()
                if search_term:
                    self.search_and_click_first_result(search_term)
                    return False
                else:
                    self.speak("Please specify what to search for")
                    return False
            if command.startswith('type '):
                self.type_and_enter(command)
                return False
            if command in ['stop typing', 'end typing', 'exit typing']:
                self.stop_typing()
                return False
            if 'contact manager' in command or 'show contacts' in command:
                self.speak("Opening contact manager")
                self.root.after(0, self.show_contact_manager)
                return False
            if 'tab' in command:
                tab_number = None
                if 'first' in command or '1st' in command or 'one' in command:
                    tab_number = 1
                elif 'second' in command or '2nd' in command or 'two' in command:
                    tab_number = 2
                elif 'third' in command or '3rd' in command or 'three' in command:
                    tab_number = 3
                elif 'fourth' in command or '4th' in command or 'four' in command:
                    tab_number = 4
                else:
                    numbers = [int(s) for s in command.split() if s.isdigit()]
                    if numbers:
                        tab_number = numbers[0]
                if tab_number is not None:
                    self.switch_tab(tab_number - 1)
                    self.speak(f"Switched to tab {tab_number}")
                    return False
            if 'open notes' in command:
                self.root.after(0, self.show_notes_window)
                return False
            if 'check system health' in command or 'check system' in command or 'system health' in command:
                self.speak("Checking system health")
                self.check_system_health()
                return False
            if any(phrase in command for phrase in ['search and open', 'find and open', 'search for and open']):
                search_term = command
                for phrase in ['search and open', 'find and open', 'search for and open']:
                    search_term = search_term.replace(phrase, '').strip()
                if search_term:
                    self.search_and_click_first_result(search_term)
                else:
                    self.speak("Please specify what to search for")
                return False
            if 'data analysis' in command:
                self.speak("Here You can analyse the Data")
                self.root.after(0, self.show_load_dataset_window)
                return False
            if any(phrase in command for phrase in ['save notes in file', 'save note', 'save this note', 'save the note', 'save notes']):
                self.save_notes()
                return False
            if not command.startswith('search') and not any(word in command for word in ['open', 'check', 'save', 'type']):
                self.speak(f"Searching for {command}")
                self.search_and_click_first_result(command)
                return False
            if command.startswith('open '):
                app_name = command.replace('open ', '', 1).strip()
                self.speak(f"I don't know how to open {app_name} directly. Searching for it...")
                self.search_and_click_first_result(app_name)
                return False
            self.speak(f"I'm not sure how to handle the command: {command}")
            return False
        except Exception as e:
            logging.error(f"Error in direct command execution: {str(e)}")
            if command.startswith('open '):
                try:
                    app_name = command.replace('open ', '', 1).strip()
                    self.speak(f"I encountered an error. Trying to search for {app_name} instead.")
                    self.search_and_click_first_result(app_name)
                    return False
                except Exception as search_error:
                    logging.error(f"Error in recovery attempt: {str(search_error)}")
            self.speak("I encountered an error executing that command.")
            return False
    def start_color_transition(self, widget, target_color, steps=10):
        """Start smooth color transition animation for customtkinter widgets"""
        if not hasattr(self, 'color_transitions'):
            self.color_transitions = {}
        if widget in self.color_transitions:
            self.root.after_cancel(self.color_transitions[widget])
        try:
            current_color = widget.cget('fg_color')
            if isinstance(current_color, tuple):
                current_color = current_color[1]
        except:
            current_color = '#00ff00'
        def interpolate_color(start_color, end_color, step, total_steps):
            """Interpolate between two colors"""
            start_color = start_color.lstrip('#')
            end_color = end_color.lstrip('#')
            if len(start_color) == 3:
                start_color = ''.join([c*2 for c in start_color])
            if len(end_color) == 3:
                end_color = ''.join([c*2 for c in end_color])
            start_rgb = [int(start_color[i:i+2], 16) for i in (0, 2, 4)]
            end_rgb = [int(end_color[i:i+2], 16) for i in (0, 2, 4)]
            current_rgb = [
                int(start_rgb[i] + (end_rgb[i] - start_rgb[i]) * step / total_steps)
                for i in range(3)
            ]
            return f'#{current_rgb[0]:02x}{current_rgb[1]:02x}{current_rgb[2]:02x}'
        def animate_step(step=0):
            if step <= steps:
                color = interpolate_color(current_color, target_color, step, steps)
                try:
                    widget.configure(fg_color=color)
                except:
                    try:
                        widget.configure(bg=color)
                    except:
                        pass
                self.color_transitions[widget] = self.root.after(
                    20,
                    lambda: animate_step(step + 1)
                )
            else:
                self.color_transitions.pop(widget, None)
        animate_step()
    def _parse_compound_command(self, command):
        """Parse a compound command into individual commands"""
        logging.info(f"Parsing compound command: '{command}'")
        separators = ['open', 'search', 'check', 'find', 'save', 'type', 'show']
        if ' and ' in command:
            return [cmd.strip() for cmd in command.split(' and ')]
        if 'open notepad' in command and ('open google' in command or 'search' in command):
            parts = []
            if 'open notepad' in command:
                parts.append('open notepad')
                command = command.replace('open notepad', '', 1).strip()
            if 'open google' in command:
                parts.append('open google')
                command = command.replace('open google', '', 1).strip()
            if 'search youtube' in command:
                search_query = command.replace('search youtube', '', 1).strip()
                if search_query:
                    parts.append(f"search youtube {search_query}")
                else:
                    parts.append("search youtube")
            elif 'search' in command:
                search_query = command.replace('search', '', 1).strip()
                if search_query:
                    parts.append(f"search for {search_query}")
                else:
                    parts.append("search")
            elif command.strip():
                parts.append(command.strip())
            if parts:
                return parts
        result = []
        result.append(command)
        for separator in separators:
            start_pos = command.find(separator)
            if start_pos == -1:
                continue
            pos = command.find(f" {separator} ", start_pos + 1)
            while pos != -1:
                first_part = command[:pos].strip()
                second_part = command[pos+1:].strip()
                if first_part and second_part and len(first_part.split()) >= 2:
                    result = [first_part, second_part]
                    break
                pos = command.find(f" {separator} ", pos + 1)
        if len(result) == 1 and len(command.split()) > 4:
            words = command.split()
            commands = []
            current_command = []
            for word in words:
                if word in separators and current_command:
                    commands.append(' '.join(current_command))
                    current_command = [word]
                else:
                    current_command.append(word)
            if current_command:
                commands.append(' '.join(current_command))
            if len(commands) > 1:
                result = commands
        if len(result) == 1 and len(command.split()) > 6:
            words = command.split()
            for i in range(2, len(words) - 1):
                if words[i] in separators:
                    first_part = ' '.join(words[:i])
                    second_part = ' '.join(words[i:])
                    if first_part and second_part:
                        result = [first_part, second_part]
                        break
        logging.info(f"Parsed into: {result}")
        return result
    def search_and_click_first_result(self, search_term):
        """Search for a term and click the first result"""
        try:
            self.speak(f"Searching for {search_term}")
            pyautogui.hotkey('win', 's')
            time.sleep(1)
            pyautogui.typewrite(search_term, interval=0.05)
            time.sleep(1.5)
            pyautogui.press('enter')
            return True
        except Exception as e:
            logging.error(f"Error in search_and_click_first_result: {str(e)}")
            self.speak(f"I couldn't find {search_term}")
            return False
    def analyze_news(self, text_or_url):
        """Analyze news content for authenticity and insights"""
        try:
            news_window = tk.Toplevel(self.root)
            news_window.title("News Analysis")
            news_window.geometry("800x600")
            news_window.configure(bg='#0a0a1f')  # Dark blue background
            main_frame = Frame(
                news_window,
                bg='#0a0a1f',
                padx=20,
                pady=20
            )
            main_frame.pack(fill=tk.BOTH, expand=True)
            title_label = Label(
                main_frame,
                text="News Analysis Results",
                font=("Segoe UI", 24, "bold"),
                bg='#0a0a1f',
                fg='#00f2ff'  # Bright cyan for better visibility
            )
            title_label.pack(pady=(0, 20))
            result_text = scrolledtext.ScrolledText(
                main_frame,
                wrap=WORD,
                height=20,
                width=70,
                bg='#12122a',  # Slightly lighter than background
                fg='#e2e2e2',
                font=("JetBrains Mono", 11),
                padx=15,
                pady=10
            )
            result_text.pack(fill=tk.BOTH, expand=True, pady=10)
            result_text.tag_configure("title", foreground="#00f2ff", font=("JetBrains Mono", 12, "bold"))  # Bright cyan
            result_text.tag_configure("authentic", foreground="#00ff7f")  # Bright green
            result_text.tag_configure("suspicious", foreground="#ff4d4d")  # Bright red
            result_text.tag_configure("warning", foreground="#ffd700")  # Gold
            result_text.tag_configure("neutral", foreground="#4361ee")  # Blue
            result_text.tag_configure("insight", foreground="#8f9aff")  # Light purple
            result_text.insert(END, "Analyzing news content...\n", "neutral")
            news_window.update()
            prompt = f"""Analyze this news content for authenticity and provide insights. Keep responses very concise.
            Content: {text_or_url}
            Provide a clear conclusion in this format:
            1. VERDICT: [REAL/FAKE/MISLEADING] (one word only)
            2. CONFIDENCE: [0-100]%
            3. KEY RED FLAGS: (if any, max 3 bullet points)
            4. QUICK FACTS: (2-3 key verified facts)
            5. BOTTOM LINE: (1-2 sentences max)
            Keep all responses extremely brief and clear."""
            response = self.gemini_model.generate_content(prompt)
            analysis = response.text
            result_text.delete(1.0, END)
            verdict = ""
            confidence = ""
            bottom_line = ""
            sections = analysis.split('\n\n')
            for section in sections:
                if 'VERDICT' in section:
                    result_text.insert(END, "\n🔍 FINAL VERDICT\n", "title")
                    if 'REAL' in section:
                        result_text.insert(END, f"{section}\n", "authentic")
                        verdict = "This content is authentic"
                    elif 'FAKE' in section:
                        result_text.insert(END, f"{section}\n", "suspicious")
                        verdict = "This content is fake"
                    else:
                        result_text.insert(END, f"{section}\n", "warning")
                        verdict = "This content is misleading"
                    if 'CONFIDENCE' in section:
                        confidence_match = re.search(r'CONFIDENCE:\s*(\d+)%', section)
                        if confidence_match:
                            confidence = f" with {confidence_match.group(1)}% confidence"
                elif 'RED FLAGS' in section:
                    result_text.insert(END, "\n⚠️ RED FLAGS\n", "title")
                    result_text.insert(END, f"{section}\n", "suspicious")
                elif 'QUICK FACTS' in section:
                    result_text.insert(END, "\n✓ VERIFIED FACTS\n", "title")
                    result_text.insert(END, f"{section}\n", "authentic")
                elif 'BOTTOM LINE' in section:
                    result_text.insert(END, "\n💡 CONCLUSION\n", "title")
                    result_text.insert(END, f"{section}\n", "insight")
                    bottom_line = section.replace("BOTTOM LINE:", "").strip()
                else:
                    result_text.insert(END, f"{section}\n", "neutral")
            result_text.configure(state='disabled')
            close_button = Button(
                main_frame,
                text="Close",
                command=news_window.destroy,
                bg='#4361ee',
                fg='#ffffff',
                font=('Segoe UI', 12, 'bold'),
                relief='flat',
                cursor='hand2',
                activebackground='#3251de',  # Slightly darker when clicked
                activeforeground='#ffffff'
            )
            close_button.pack(pady=20)
            if verdict:
                verdict_text = f"{verdict}{confidence}. "
                self.speak(verdict_text)
                if bottom_line:
                    self.root.after(2000, lambda: self.speak(bottom_line))
        except Exception as e:
            logging.error(f"Error in news analysis: {str(e)}")
            self.speak("Sorry, I encountered an error analyzing the news content")
    def show_translation_comparison(self, target_language=None):
        """Show sentence-by-sentence translation comparison for highlighted text"""
        try:
            pyautogui.hotkey('ctrl', 'c')
            time.sleep(0.1)
            try:
                highlighted_text = self.root.clipboard_get().strip()
            except:
                self.speak("No text was highlighted")
                return
            if not highlighted_text:
                self.speak("Please highlight some text first")
                return
            translation_window = tk.Toplevel(self.root)
            translation_window.title("Translation Comparison - TRON Interface")
            translation_window.geometry("1800x800")
            translation_window.configure(bg='#000000')
            border_frame = tk.Frame(
                translation_window,
                bg='#000000',
                highlightbackground='#15f4ee',
                highlightthickness=3
            )
            border_frame.pack(fill='both', expand=True, padx=6, pady=6)
            main_frame = tk.Frame(border_frame, bg='#000000')
            main_frame.pack(fill='both', expand=True, padx=25, pady=25)
            title_frame = tk.Frame(main_frame, bg='#000000')
            title_frame.pack(fill='x', pady=(0, 25))
            title_label = tk.Label(
                title_frame,
                text="TRANSLATION MATRIX",
                font=("Share Tech Mono", 28, "bold"),
                bg='#000000',
                fg='#15f4ee'
            )
            title_label.pack(side='left')
            lang_frame = tk.Frame(main_frame, bg='#000000')
            lang_frame.pack(fill='x', pady=(0, 25))
            tk.Label(
                lang_frame,
                text="FROM:",
                font=("Share Tech Mono", 14),
                bg='#000000',
                fg='#FF410D'
            ).pack(side='left', padx=(0, 15))
            source_lang = ttk.Combobox(
                lang_frame,
                values=sorted(self.language_dict.keys()),
                width=25,
                font=("Share Tech Mono", 12),
                state='readonly'
            )
            source_lang.set("English")
            source_lang.pack(side='left', padx=(0, 25))
            tk.Label(
                lang_frame,
                text="TO:",
                font=("Share Tech Mono", 14),
                bg='#000000',
                fg='#FF410D'
            ).pack(side='left', padx=(0, 15))
            target_lang = ttk.Combobox(
                lang_frame,
                values=sorted(self.language_dict.keys()),
                width=25,
                font=("Share Tech Mono", 12),
                state='readonly'
            )
            target_lang.set("Spanish" if not target_language else target_language)
            target_lang.pack(side='left')
            scroll_frame = tk.Frame(main_frame, bg='#000000')
            scroll_frame.pack(fill='both', expand=True, pady=(0, 25))
            canvas = tk.Canvas(
                scroll_frame,
                bg='#000000',
                highlightthickness=0
            )
            scrollbar = ttk.Scrollbar(scroll_frame, orient="vertical", command=canvas.yview)
            content_frame = tk.Frame(canvas, bg='#000000')
            canvas.configure(yscrollcommand=scrollbar.set)
            canvas.bind('<Configure>', lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
            canvas_window = canvas.create_window((0, 0), window=content_frame, anchor='nw')
            def update_content_width(event):
                canvas.itemconfig(canvas_window, width=event.width)
            canvas.bind('<Configure>', update_content_width)
            scrollbar.pack(side="right", fill="y")
            canvas.pack(side="left", fill="both", expand=True)
            def translate_text():
                try:
                    translate_btn.configure(state='disabled', text="TRANSLATING...")
                    source_code = self.language_dict[source_lang.get()]
                    target_code = self.language_dict[target_lang.get()]
                    for widget in content_frame.winfo_children():
                        widget.destroy()
                    loading_frame = tk.Frame(content_frame, bg='#000000')
                    loading_frame.pack(pady=20)
                    loading_label = tk.Label(
                        loading_frame,
                        text="TRANSLATING TEXT...",
                        font=("Share Tech Mono", 14),
                        bg='#000000',
                        fg='#15f4ee'
                    )
                    loading_label.pack()
                    progress_bar = ttk.Progressbar(
                        loading_frame,
                        mode='indeterminate',
                        length=400
                    )
                    progress_bar.pack(pady=10)
                    progress_bar.start(10)
                    def process_translation():
                        try:
                            translator = GoogleTranslator(source=source_code, target=target_code)
                            full_translation = translator.translate(highlighted_text)
                            original_sentences = [s.strip() for s in re.split('[.!?]+', highlighted_text) if s.strip()]
                            translated_sentences = [s.strip() for s in re.split('[.!?]+', full_translation) if s.strip()]
                            min_sentences = min(len(original_sentences), len(translated_sentences))
                            original_sentences = original_sentences[:min_sentences]
                            translated_sentences = translated_sentences[:min_sentences]
                            def update_ui():
                                try:
                                    loading_frame.destroy()
                                    for i, (original, translated) in enumerate(zip(original_sentences, translated_sentences)):
                                        pair_frame = tk.Frame(content_frame, bg='#000000')
                                        pair_frame.pack(fill='x', pady=10, padx=15)
                                        original_frame = tk.Frame(pair_frame, bg='#000000')
                                        original_frame.pack(side='left', fill='both', expand=True, padx=5)
                                        tk.Label(
                                            original_frame,
                                            text=f"ORIGINAL ({source_lang.get()})",
                                            font=("Share Tech Mono", 12),
                                            bg='#000000',
                                            fg='#15f4ee'
                                        ).pack(anchor='w', pady=(0, 5))
                                        original_text = tk.Text(
                                            original_frame,
                                            height=4,
                                            width=50,
                                            bg='#0D1B2A',
                                            fg='#15f4ee',
                                            font=("Share Tech Mono", 11),
                                            wrap=tk.WORD,
                                            padx=10,
                                            pady=5
                                        )
                                        original_text.pack(fill='both')
                                        original_text.insert('1.0', original)
                                        original_text.configure(state='disabled')
                                        translated_frame = tk.Frame(pair_frame, bg='#000000')
                                        translated_frame.pack(side='right', fill='both', expand=True, padx=5)
                                        tk.Label(
                                            translated_frame,
                                            text=f"TRANSLATED ({target_lang.get()})",
                                            font=("Share Tech Mono", 12),
                                            bg='#000000',
                                            fg='#FF410D'
                                        ).pack(anchor='w', pady=(0, 5))
                                        translated_text = tk.Text(
                                            translated_frame,
                                            height=4,
                                            width=50,
                                            bg='#0D1B2A',
                                            fg='#FF410D',
                                            font=("Share Tech Mono", 11),
                                            wrap=tk.WORD,
                                            padx=10,
                                            pady=5
                                        )
                                        translated_text.pack(fill='both')
                                        translated_text.insert('1.0', translated)
                                        translated_text.configure(state='disabled')
                                        buttons_frame = tk.Frame(pair_frame, bg='#000000')
                                        buttons_frame.pack(fill='x', pady=5)
                                        def create_copy_button(text_widget, label):
                                            def copy_text():
                                                translation_window.clipboard_clear()
                                                translation_window.clipboard_append(text_widget.get("1.0", tk.END).strip())
                                                notification = tk.Toplevel(translation_window)
                                                notification.overrideredirect(True)
                                                notification.configure(bg='#15f4ee')
                                                tk.Label(
                                                    notification,
                                                    text=f"{label} COPIED",
                                                    font=("Share Tech Mono", 11),
                                                    bg='#15f4ee',
                                                    fg='#000000',
                                                    padx=20,
                                                    pady=10
                                                ).pack()
                                                x = translation_window.winfo_x() + translation_window.winfo_width() - 200
                                                y = translation_window.winfo_y() + 50
                                                notification.geometry(f"+{x}+{y}")
                                                notification.after(1500, notification.destroy)
                                            return tk.Button(
                                                buttons_frame,
                                                text=f"COPY {label}",
                                                command=copy_text,
                                                font=("Share Tech Mono", 10, "bold"),
                                                bg='#0D1B2A',
                                                fg='#15f4ee',
                                                activebackground='#15f4ee',
                                                activeforeground='#0D1B2A',
                                                relief='flat',
                                                padx=15,
                                                pady=5,
                                                cursor='hand2'
                                            )
                                        copy_original = create_copy_button(original_text, "ORIGINAL")
                                        copy_translated = create_copy_button(translated_text, "TRANSLATION")
                                        copy_original.pack(side='left', padx=5)
                                        copy_translated.pack(side='right', padx=5)
                                        if i < min_sentences - 1:
                                            separator = tk.Frame(content_frame, height=2, bg='#15f4ee')
                                            separator.pack(fill='x', pady=10)
                                    translate_btn.configure(state='normal', text="TRANSLATE")
                                except Exception as e:
                                    messagebox.showerror("Error", f"Error updating UI: {str(e)}", parent=translation_window)
                            translation_window.after(0, update_ui)
                        except Exception as e:
                            translation_window.after(0, lambda: messagebox.showerror(
                                "Translation Error", 
                                f"Translation failed: {str(e)}", 
                                parent=translation_window
                            ))
                            translation_window.after(0, lambda: translate_btn.configure(
                                state='normal', 
                                text="TRANSLATE"
                            ))
                    threading.Thread(target=process_translation, daemon=True).start()
                except Exception as e:
                    messagebox.showerror("Error", f"Failed to start translation: {str(e)}", parent=translation_window)
                    translate_btn.configure(state='normal', text="TRANSLATE")
            translate_btn = tk.Button(
                main_frame,
                text="TRANSLATE",
                command=translate_text,
                font=("Share Tech Mono", 14, "bold"),
                bg='#0D1B2A',
                fg='#15f4ee',
                activebackground='#15f4ee',
                activeforeground='#0D1B2A',
                relief='flat',
                padx=25,
                pady=10,
                cursor='hand2'
            )
            translate_btn.pack(pady=20)
            translation_window.update_idletasks()
            width = translation_window.winfo_width()
            height = translation_window.winfo_height()
            x = (translation_window.winfo_screenwidth() // 2) - (width // 2)
            y = (translation_window.winfo_screenheight() // 2) - (height // 2)
            translation_window.geometry(f'{width}x{height}+{x}+{y}')
            translation_window.transient(self.root)
            translation_window.grab_set()
            if target_language:
                translation_window.after(500, translate_text)
        except Exception as e:
            self.handle_errors(e, "Translation window error")
            self.speak("Sorry, I encountered an error with the translation window")
    language_dict = {
        "English": "en",
        "Spanish": "es",
        "French": "fr",
        "German": "de",
        "Italian": "it",
        "Portuguese": "pt",
        "Russian": "ru",
        "Japanese": "ja",
        "Korean": "ko",
        "Chinese": "zh",
        "Arabic": "ar",
        "Hindi": "hi",
        "Bengali": "bn",
        "Tamil": "ta",
        "Telugu": "te",
        "Kannada": "kn",
        "Malayalam": "ml",
        "Marathi": "mr",
        "Gujarati": "gu",
        "Urdu": "ur"
    }
    def switch_tab(self, index):
        """Switch to specified tab using Alt+Tab"""
        try:
            pyautogui.keyDown('alt')
            if index > 0:
                for _ in range(index):
                    pyautogui.press('tab')
                    time.sleep(0.1)
            pyautogui.keyUp('alt')
        except Exception as e:
            self.handle_errors(e, "Error switching tabs")
            self.speak("Sorry, I couldn't switch tabs") 
    def is_valid_website(self, website):
        """Check if the website is valid for auto-login."""
        known_websites = ["google", "facebook", "twitter", "linkedin", "github"]
        return any(known in website for known in known_websites)
    def show_program_output_code(self,comments):
        """Show programming output in Tron Legacy themed window with syntax highlighting"""
        try:
            if not comments:
                self.speak("Please highlight some text first or say comments following word code")
                return
            output_window = tk.Toplevel(self.root)
            output_window.title("Program Output - TRON Interface")
            output_window.geometry("800x600")
            output_window.configure(bg='#000000')
            border_frame = tk.Frame(
                output_window,
                bg='#000000',
                highlightbackground='#15f4ee',
                highlightthickness=2
            )
            border_frame.pack(fill='both', expand=True, padx=4, pady=4)
            main_frame = tk.Frame(border_frame, bg='#000000')
            main_frame.pack(fill='both', expand=True, padx=20, pady=20)
            title_frame = tk.Frame(main_frame, bg='#000000')
            title_frame.pack(fill='x', pady=(0, 20))
            title_label = tk.Label(
                title_frame,
                text="PROGRAM SOLUTION",
                font=("Share Tech Mono", 18, "bold"),
                bg='#000000',
                fg='#15f4ee'
            )
            title_label.pack(side='left')
            loading_frame = tk.Frame(main_frame, bg='#000000')
            loading_frame.pack(fill='x', pady=20)
            loading_label = tk.Label(
                loading_frame,
                text="INITIALIZING CODE GENERATION",
                font=("Share Tech Mono", 14),
                bg='#000000',
                fg='#FF410D'
            )
            loading_label.pack()
            style = ttk.Style()
            style.configure(
                'Tron.Horizontal.TProgressbar',
                troughcolor='#000B19',
                background='#15f4ee',
                thickness=4,
                width=15,
                arrowsize=15
            )
            progress_bar = ttk.Progressbar(
                loading_frame,
                mode='indeterminate',
                length=300,
                style='Tron.Horizontal.TProgressbar'
            )
            progress_bar.pack(pady=10)
            dots_label = tk.Label(
                loading_frame,
                text="",
                font=("Share Tech Mono", 14),
                bg='#000000',
                fg='#FF410D'
            )
            dots_label.pack()
            output_text = scrolledtext.ScrolledText(
                main_frame,
                wrap=tk.NONE,
                font=("JetBrains Mono", 12),
                bg='#000B19',
                fg='#15f4ee',
                insertbackground='#15f4ee',
                height=20,
                relief='flat'
            )
            output_text.tag_configure('keyword', foreground='#FF410D')
            output_text.tag_configure('string', foreground='#95FF00')
            output_text.tag_configure('comment', foreground='#7A7A7A')
            output_text.tag_configure('function', foreground='#15f4ee')
            output_text.tag_configure('number', foreground='#FFB700')
            output_text.tag_configure('operator', foreground='#FF410D')
            output_text.tag_configure('builtin', foreground='#FF410D')
            def on_window_close():
                try:
                    output_window.is_generating = False
                    if hasattr(dots_label, 'animation_id'):
                        output_window.after_cancel(dots_label.animation_id)
                    if progress_bar.winfo_exists():
                        progress_bar.stop()
                    output_window.destroy()
                except Exception as e:
                    print(f"Error during cleanup: {e}")
                    output_window.destroy()
            def animate_dots():
                try:
                    if not hasattr(output_window, 'is_generating') or not output_window.is_generating:
                        return
                    if not dots_label.winfo_exists():
                        return
                    current_text = dots_label.cget("text")
                    if current_text == "....":
                        dots_label.configure(text="")
                    else:
                        dots_label.configure(text=current_text + ".")
                    if output_window.is_generating and dots_label.winfo_exists():
                        dots_label.animation_id = output_window.after(500, animate_dots)
                except tk.TclError:
                    return
                except Exception as e:
                    print(f"Animation error: {str(e)}")
                    return
            def update_loading_status(status):
                loading_label.configure(text=status)
                output_window.update_idletasks()
            def generate_code():
                try:
                    output_window.is_generating = True
                    progress_bar.start(10)
                    animate_dots()
                    api_key = self.load_setting('gemini_api_key')
                    model_name = self.load_setting("gemini_model")
                    import google.generativeai as genai
                    genai.configure(api_key=api_key)
                    model = genai.GenerativeModel(model_name)
                    update_loading_status("ANALYZING REQUIREMENTS")
                    time.sleep(1)
                    update_loading_status("GENERATING CODE SOLUTION")
                    prompt = f"""Create a complete, working program for: {comments}
                    Guidelines:
                    - Provide only the code without any markdown formatting
                    - Include necessary imports
                    - Use clear variable names
                    - Add brief inline comments
                    - Make it production-ready
                    - Return raw code only, no explanations or decorations
                    - Explanation should be short and sweet, don't take more lines to explain
                    """
                    response = model.generate_content(prompt)
                    program = response.text.strip()
                    program = program.replace('```python', '').replace('```', '').strip()
                    update_loading_status("APPLYING SYNTAX HIGHLIGHTING")
                    time.sleep(0.5)
                    output_window.is_generating = False
                    progress_bar.stop()
                    loading_frame.destroy()
                    output_text.pack(fill='both', expand=True, pady=(0, 10))
                    output_text.pack(fill='both', expand=True, pady=(0, 10))
                    keywords = ['def', 'class', 'import', 'from', 'return', 'if', 'else', 'elif',
                              'try', 'except', 'finally', 'for', 'while', 'in', 'and', 'or',
                              'not', 'is', 'None', 'True', 'False', 'with', 'as', 'break',
                              'continue', 'pass', 'raise', 'yield', 'async', 'await']
                    builtins = ['print', 'len', 'range', 'str', 'int', 'float', 'list', 'dict',
                               'set', 'tuple', 'sum', 'min', 'max', 'sorted', 'enumerate',
                               'zip', 'map', 'filter', 'any', 'all', 'round', 'abs', 'open']
                    for line in program.split('\n'):
                        if not line.strip():
                            output_text.insert('end', '\n')
                            continue
                        if line.strip().startswith('#'):
                            output_text.insert('end', line + '\n', 'comment')
                            continue
                        current_pos = 0
                        line_length = len(line)
                        while current_pos < line_length:
                            if line[current_pos].isspace():
                                output_text.insert('end', line[current_pos])
                                current_pos += 1
                                continue
                            if line[current_pos] in '"\'':
                                quote = line[current_pos]
                                end_pos = line.find(quote, current_pos + 1)
                                if end_pos == -1:
                                    end_pos = line_length
                                output_text.insert('end', line[current_pos:end_pos + 1], 'string')
                                current_pos = end_pos + 1
                            elif line[current_pos:].startswith('#'):  # Comment
                                output_text.insert('end', line[current_pos:], 'comment')
                                current_pos = line_length
                            else:
                                end_pos = current_pos
                                while end_pos < line_length and (line[end_pos].isalnum() or line[end_pos] == '_'):
                                    end_pos += 1
                                token = line[current_pos:end_pos]
                                if token in keywords:
                                    output_text.insert('end', token, 'keyword')
                                elif token in builtins:
                                    output_text.insert('end', token, 'builtin')
                                elif token.isdigit():
                                    output_text.insert('end', token, 'number')
                                elif token and token[0].isupper():
                                    output_text.insert('end', token, 'class')
                                else:
                                    output_text.insert('end', token)
                                if end_pos == current_pos:
                                    output_text.insert('end', line[current_pos])
                                    end_pos += 1
                                current_pos = end_pos
                        output_text.insert('end', '\n')
                except Exception as e:
                    output_window.is_generating = False
                    progress_bar.stop()
                    loading_frame.destroy()
                    output_text.pack(fill='both', expand=True, pady=(0, 10))
                    output_text.insert('1.0', f"Error generating program: {str(e)}")
            buttons_frame = tk.Frame(main_frame, bg='#000000')
            buttons_frame.pack(fill='x', pady=(10, 0))
            def create_glowing_button(parent, text, command, color='#15f4ee'):
                frame = tk.Frame(
                    parent,
                    bg='#000000',
                    highlightbackground=color,
                    highlightthickness=1
                )
                btn = tk.Button(
                    frame,
                    text=text,
                    command=command,
                    bg='#000B19',
                    fg=color,
                    font=("Share Tech Mono", 11),
                    relief='flat',
                    cursor='hand2',
                    activebackground='#000B19',
                    activeforeground='#FFFFFF',
                    bd=0,
                    padx=20,
                    pady=5
                )
                btn.pack(padx=1, pady=1)
                def on_enter(e):
                    frame.configure(highlightbackground='#FFFFFF')
                    btn.configure(fg='#FFFFFF')
                def on_leave(e):
                    frame.configure(highlightbackground=color)
                    btn.configure(fg=color)
                btn.bind('<Enter>', on_enter)
                btn.bind('<Leave>', on_leave)
                return frame
            create_glowing_button(
                buttons_frame,
                "COPY CODE",
                lambda: self.root.clipboard_append(output_text.get("1.0", tk.END))
            ).pack(side='left', padx=5)
            create_glowing_button(
                buttons_frame,
                "CLOSE",
                on_window_close,
                color='#FF410D'
            ).pack(side='right', padx=5)
            output_window.protocol("WM_DELETE_WINDOW", on_window_close)
            threading.Thread(target=generate_code, daemon=True).start()
            output_window.transient(self.root)
            output_window.grab_set()
        except Exception as e:
            self.handle_errors(e, "Error showing program output")
            self.speak("Sorry, I encountered an error generating the program")
    def show_program_output(self):
        """Show programming output in Tron Legacy themed window with syntax highlighting"""
        try:
            pyautogui.hotkey('ctrl','a')
            pyautogui.hotkey('ctrl', 'c')
            time.sleep(0.1)
            try:
                highlighted_text = self.root.clipboard_get().strip()
            except:
                self.speak("No text was highlighted")
                return
            if not highlighted_text:
                self.speak("Please highlight some text first")
                return
            output_window = tk.Toplevel(self.root)
            output_window.title("Program Output - TRON Interface")
            output_window.geometry("800x600")
            output_window.configure(bg='#000000')
            border_frame = tk.Frame(
                output_window,
                bg='#000000',
                highlightbackground='#15f4ee',
                highlightthickness=2
            )
            border_frame.pack(fill='both', expand=True, padx=4, pady=4)
            main_frame = tk.Frame(border_frame, bg='#000000')
            main_frame.pack(fill='both', expand=True, padx=20, pady=20)
            title_frame = tk.Frame(main_frame, bg='#000000')
            title_frame.pack(fill='x', pady=(0, 20))
            title_label = tk.Label(
                title_frame,
                text="PROGRAM SOLUTION",
                font=("Share Tech Mono", 18, "bold"),
                bg='#000000',
                fg='#15f4ee'
            )
            title_label.pack(side='left')
            loading_frame = tk.Frame(main_frame, bg='#000000')
            loading_frame.pack(fill='x', pady=20)
            loading_label = tk.Label(
                loading_frame,
                text="INITIALIZING CODE GENERATION",
                font=("Share Tech Mono", 14),
                bg='#000000',
                fg='#FF410D'
            )
            loading_label.pack()
            style = ttk.Style()
            style.configure(
                'Tron.Horizontal.TProgressbar',
                troughcolor='#000B19',
                background='#15f4ee',
                thickness=4,
                width=15,
                arrowsize=15
            )
            progress_bar = ttk.Progressbar(
                loading_frame,
                mode='indeterminate',
                length=300,
                style='Tron.Horizontal.TProgressbar'
            )
            progress_bar.pack(pady=10)
            dots_label = tk.Label(
                loading_frame,
                text="",
                font=("Share Tech Mono", 14),
                bg='#000000',
                fg='#FF410D'
            )
            dots_label.pack()
            output_text = scrolledtext.ScrolledText(
                main_frame,
                wrap=tk.NONE,
                font=("JetBrains Mono", 12),
                bg='#000B19',
                fg='#15f4ee',
                insertbackground='#15f4ee',
                height=20,
                relief='flat'
            )
            output_text.tag_configure('keyword', foreground='#FF410D')
            output_text.tag_configure('string', foreground='#95FF00')
            output_text.tag_configure('comment', foreground='#7A7A7A')
            output_text.tag_configure('function', foreground='#15f4ee')
            output_text.tag_configure('number', foreground='#FFB700')
            output_text.tag_configure('operator', foreground='#FF410D')
            output_text.tag_configure('builtin', foreground='#FF410D')
            def on_window_close():
                try:
                    output_window.is_generating = False
                    if hasattr(dots_label, 'animation_id'):
                        output_window.after_cancel(dots_label.animation_id)
                    if progress_bar.winfo_exists():
                        progress_bar.stop()
                    output_window.destroy()
                except Exception as e:
                    print(f"Error during cleanup: {e}")
                    output_window.destroy()
            def animate_dots():
                try:
                    if not hasattr(output_window, 'is_generating') or not output_window.is_generating:
                        return
                    if not dots_label.winfo_exists():
                        return
                    current_text = dots_label.cget("text")
                    if current_text == "....":
                        dots_label.configure(text="")
                    else:
                        dots_label.configure(text=current_text + ".")
                    if output_window.is_generating and dots_label.winfo_exists():
                        dots_label.animation_id = output_window.after(500, animate_dots)
                except tk.TclError:
                    return
                except Exception as e:
                    print(f"Animation error: {str(e)}")
                    return
            def update_loading_status(status):
                loading_label.configure(text=status)
                output_window.update_idletasks()
            def generate_code():
                try:
                    output_window.is_generating = True
                    progress_bar.start(10)
                    animate_dots()
                    api_key = self.load_setting('gemini_api_key')
                    model_name = self.load_setting("gemini_model")
                    import google.generativeai as genai
                    genai.configure(api_key=api_key)
                    model = genai.GenerativeModel(model_name)
                    update_loading_status("ANALYZING REQUIREMENTS")
                    time.sleep(1)
                    update_loading_status("GENERATING CODE SOLUTION")
                    prompt = f"""Create a complete, working program for: {highlighted_text}
                    Guidelines:
                    - Provide only the code without any markdown formatting
                    - Include necessary imports
                    - Use clear variable names
                    - Add brief inline comments
                    - Make it production-ready
                    - Return raw code only, no explanations or decorations
                    - Explanation should be short and sweet, don't take more lines to explain
                    """
                    response = model.generate_content(prompt)
                    program = response.text.strip()
                    program = program.replace('```python', '').replace('```', '').strip()
                    update_loading_status("APPLYING SYNTAX HIGHLIGHTING")
                    time.sleep(0.5)
                    output_window.is_generating = False
                    progress_bar.stop()
                    loading_frame.destroy()
                    output_text.pack(fill='both', expand=True, pady=(0, 10))
                    output_text.pack(fill='both', expand=True, pady=(0, 10))
                    keywords = ['def', 'class', 'import', 'from', 'return', 'if', 'else', 'elif',
                              'try', 'except', 'finally', 'for', 'while', 'in', 'and', 'or',
                              'not', 'is', 'None', 'True', 'False', 'with', 'as', 'break',
                              'continue', 'pass', 'raise', 'yield', 'async', 'await']
                    builtins = ['print', 'len', 'range', 'str', 'int', 'float', 'list', 'dict',
                               'set', 'tuple', 'sum', 'min', 'max', 'sorted', 'enumerate',
                               'zip', 'map', 'filter', 'any', 'all', 'round', 'abs', 'open']
                    for line in program.split('\n'):
                        if not line.strip():
                            output_text.insert('end', '\n')
                            continue
                        if line.strip().startswith('#'):
                            output_text.insert('end', line + '\n', 'comment')
                            continue
                        current_pos = 0
                        line_length = len(line)
                        while current_pos < line_length:
                            if line[current_pos].isspace():
                                output_text.insert('end', line[current_pos])
                                current_pos += 1
                                continue
                            if line[current_pos] in '"\'':
                                quote = line[current_pos]
                                end_pos = line.find(quote, current_pos + 1)
                                if end_pos == -1:
                                    end_pos = line_length
                                output_text.insert('end', line[current_pos:end_pos + 1], 'string')
                                current_pos = end_pos + 1
                            elif line[current_pos:].startswith('#'):  # Comment
                                output_text.insert('end', line[current_pos:], 'comment')
                                current_pos = line_length
                            else:
                                end_pos = current_pos
                                while end_pos < line_length and (line[end_pos].isalnum() or line[end_pos] == '_'):
                                    end_pos += 1
                                token = line[current_pos:end_pos]
                                if token in keywords:
                                    output_text.insert('end', token, 'keyword')
                                elif token in builtins:
                                    output_text.insert('end', token, 'builtin')
                                elif token.isdigit():
                                    output_text.insert('end', token, 'number')
                                elif token and token[0].isupper():
                                    output_text.insert('end', token, 'class')
                                else:
                                    output_text.insert('end', token)
                                if end_pos == current_pos:
                                    output_text.insert('end', line[current_pos])
                                    end_pos += 1
                                current_pos = end_pos
                        output_text.insert('end', '\n')
                except Exception as e:
                    output_window.is_generating = False
                    progress_bar.stop()
                    loading_frame.destroy()
                    output_text.pack(fill='both', expand=True, pady=(0, 10))
                    output_text.insert('1.0', f"Error generating program: {str(e)}")
            buttons_frame = tk.Frame(main_frame, bg='#000000')
            buttons_frame.pack(fill='x', pady=(10, 0))
            def create_glowing_button(parent, text, command, color='#15f4ee'):
                frame = tk.Frame(
                    parent,
                    bg='#000000',
                    highlightbackground=color,
                    highlightthickness=1
                )
                btn = tk.Button(
                    frame,
                    text=text,
                    command=command,
                    bg='#000B19',
                    fg=color,
                    font=("Share Tech Mono", 11),
                    relief='flat',
                    cursor='hand2',
                    activebackground='#000B19',
                    activeforeground='#FFFFFF',
                    bd=0,
                    padx=20,
                    pady=5
                )
                btn.pack(padx=1, pady=1)
                def on_enter(e):
                    frame.configure(highlightbackground='#FFFFFF')
                    btn.configure(fg='#FFFFFF')
                def on_leave(e):
                    frame.configure(highlightbackground=color)
                    btn.configure(fg=color)
                btn.bind('<Enter>', on_enter)
                btn.bind('<Leave>', on_leave)
                return frame
            create_glowing_button(
                buttons_frame,
                "COPY CODE",
                lambda: self.root.clipboard_append(output_text.get("1.0", tk.END))
            ).pack(side='left', padx=5)
            create_glowing_button(
                buttons_frame,
                "CLOSE",
                on_window_close,
                color='#FF410D'
            ).pack(side='right', padx=5)
            output_window.protocol("WM_DELETE_WINDOW", on_window_close)
            threading.Thread(target=generate_code, daemon=True).start()
            output_window.transient(self.root)
            output_window.grab_set()
        except Exception as e:
            self.handle_errors(e, "Error showing program output")
            self.speak("Sorry, I encountered an error generating the program")
    def get_direct_download_link(self, app_name):
        """
        Use Gemini AI to get accurate direct download links for applications
        specifically for Windows 10/11 with improved error handling
        """
        try:
            app_name_lower = app_name.lower()
            fallback_link = self.get_download_link_fallback(app_name)
            if fallback_link:
                return fallback_link
            import google.generativeai as genai
            api_key = self.load_setting('gemini_api_key')
            model_name = self.load_setting("gemini_model")
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel(model_name)
            prompt = f"""
            Find a direct download link for the latest version of {app_name} for Windows 10/11.
            IMPORTANT:
            - ONLY return a direct URL that starts with http:// or https://
            - Always find a direct download link, never mention Microsoft Store
            - If you cannot find a direct download link, return exactly: "NO_DIRECT_LINK"
            - Do not include any explanations, just the URL or the specified response
            """
            response = model.generate_content(prompt)
            result = response.text.strip()
            logging.info(f"Obtained download link for {app_name}: {result}")
            if result.startswith(('http://', 'https://')):
                return result
            elif result == "NO_DIRECT_LINK":
                return self.get_download_link_fallback(app_name)
            else:
                logging.warning(f"AI returned unexpected format: {result}")
                return self.get_download_link_fallback(app_name)
        except Exception as e:
            self.handle_errors(e, f"Error getting download link for {app_name}")
            return self.get_download_link_fallback(app_name)
    def get_download_link_fallback(self, app_name):
        """
        Enhanced fallback method to get download links using a comprehensive approach
        """
        try:
            problematic_apps = ["whatsapp", "telegram", "microsoft teams", "onedrive", "facebook", 
                               "skype", "zoom", "slack", "discord"]
            app_name_lower = app_name.lower().strip()
            if any(prob_app in app_name_lower for prob_app in problematic_apps):
                return None
            app_database = {
                "chrome": "https://dl.google.com/chrome/install/latest/Chrome_installer.exe",
                "firefox": "https://download.mozilla.org/?product=firefox-latest-ssl&os=win64&lang=en-US",
                "brave": "https://laptop-updates.brave.com/latest/winx64",
                "opera": "https://net.geo.opera.com/opera/stable/windows",
                "edge": "https://c2rsetup.officeapps.live.com/c2r/downloadEdge.aspx?platform=Default&source=EdgeStablePage&Channel=Stable&language=en",
                "vlc": "https://get.videolan.org/vlc/3.0.20/win64/vlc-3.0.20-win64.exe",
                "7zip": "https://www.7-zip.org/a/7z2301-x64.exe",
                "winrar": "https://www.win-rar.com/fileadmin/winrar-versions/winrar/winrar-x64-621.exe",
                "ccleaner": "https://download.ccleaner.com/ccsetup618.exe",
                "malwarebytes": "https://data-cdn.mbamupdates.com/web/mb4-setup-consumer/MBSetup.exe",
                'vs code': "https://code.visualstudio.com/sha/download?build=stable&os=win32-x64-user",
                "visual studio": "https://aka.ms/vs/17/release/vs_community.exe",
                "visual studio code": "https://code.visualstudio.com/sha/download?build=stable&os=win32-x64-user",
                "notepad++": "https://github.com/notepad-plus-plus/notepad-plus-plus/releases/download/v8.6.4/npp.8.6.4.Installer.x64.exe",
                "git": "https://github.com/git-for-windows/git/releases/download/v2.44.0.windows.1/Git-2.44.0-64-bit.exe",
                "python": "https://www.python.org/ftp/python/3.12.2/python-3.12.2-amd64.exe",
                "nodejs": "https://nodejs.org/dist/v20.11.1/node-v20.11.1-x64.msi",
                "cursor": "https://download.cursor.sh/Cursor.exe",
                "obs studio": "https://github.com/obsproject/obs-studio/releases/download/30.0.2/OBS-Studio-30.0.2-Full-Installer-x64.exe",
                "audacity": "https://github.com/audacity/audacity/releases/download/Audacity-3.4.2/audacity-win-3.4.2-64bit.exe",
                "gimp": "https://download.gimp.org/gimp/v2.10/windows/gimp-2.10.36-setup-1.exe",
                "blender": "https://ftp.nluug.nl/pub/graphics/blender/release/Blender4.0/blender-4.0.2-windows-x64.msi",
                "discord": "https://discord.com/api/downloads/distributions/app/installers/latest?channel=stable&platform=win&arch=x86",
                "zoom": "https://zoom.us/client/latest/ZoomInstallerFull.exe",
                "skype": "https://go.skype.com/windows.desktop.download",
                "slack": "https://downloads.slack-edge.com/releases/windows/4.36.140/prod/x64/SlackSetup.exe",
                "office": "https://officecdn.microsoft.com/pr/C1297A47-86C4-4C1F-97FA-950631F94777/OfficeSuiteSetup.exe",
                "spotify": "https://download.scdn.co/SpotifySetup.exe",
                "steam": "https://cdn.akamai.steamstatic.com/client/installer/SteamSetup.exe",
                "itunes": "https://www.apple.com/itunes/download/win64",
                "vlc": "https://get.videolan.org/vlc/3.0.20/win64/vlc-3.0.20-win64.exe",
                "avast": "https://www.avast.com/en-us/download-thank-you.php?product=FAV-ONLINE&locale=en-us",
                "avg": "https://www.avg.com/en-us/download-thank-you.php?product=FAV-ONLINE&locale=en-us",
                "avira": "https://package.avira.com/download/connect-client-win/avira_en_av_5f7c1e7b1b4e4__ws.exe",
                "adobe reader": "https://ardownload2.adobe.com/pub/adobe/reader/win/AcrobatDC/2300320269/AcroRdrDC2300320269_en_US.exe",
                "foxit reader": "https://cdn01.foxitsoftware.com/product/reader/desktop/win/2023.3.0/FoxitPDFReader1312_L10N_Setup_Prom.exe",
                "minecraft": "https://launcher.mojang.com/download/MinecraftInstaller.exe",
                "epic games": "https://launcher-public-service-prod06.ol.epicgames.com/launcher/api/installer/download/EpicGamesLauncherInstaller.msi",
                "origin": "https://origin-a.akamaihd.net/Origin-Client-Download/origin/live/OriginSetup.exe",
                "battle.net": "https://www.battle.net/download/getInstallerForGame?os=win&gameProgram=BATTLENET_APP",
                "teamviewer": "https://download.teamviewer.com/download/TeamViewer_Setup.exe",
                "anydesk": "https://download.anydesk.com/AnyDesk.exe",
                "whatsapp": "https://web.whatsapp.com/desktop/windows/release/x64/WhatsAppSetup.exe",
                "telegram": "https://telegram.org/dl/desktop/win64"
            }
            app_name_lower = app_name.lower()
            if app_name_lower in app_database:
                return app_database[app_name_lower]
            for key, url in app_database.items():
                if key in app_name_lower or app_name_lower in key:
                    return url
            return None
        except Exception as e:
            self.handle_errors(e, f"Error in fallback download method for {app_name}")
            return None
    def get_download_page_url(self, app_name):
        """
        Get the download page URL for an application
        This is used when we don't have a direct download link
        """
        try:
            download_pages = {
                "whatsapp": "https://www.whatsapp.com/download",
                "telegram": "https://desktop.telegram.org/",
                "discord": "https://discord.com/download",
                "skype": "https://www.skype.com/en/get-skype/",
                "microsoft teams": "https://www.microsoft.com/en-us/microsoft-teams/download-app",
                "slack": "https://slack.com/downloads/windows",
                "zoom": "https://zoom.us/download",
                "onedrive": "https://www.microsoft.com/en-us/microsoft-365/onedrive/download",
                "office": "https://www.microsoft.com/en-us/microsoft-365/download-office",
                "excel": "https://www.microsoft.com/en-us/microsoft-365/excel",
                "word": "https://www.microsoft.com/en-us/microsoft-365/word",
                "powerpoint": "https://www.microsoft.com/en-us/microsoft-365/powerpoint",
                "outlook": "https://www.microsoft.com/en-us/microsoft-365/outlook/email-and-calendar-software-microsoft-outlook",
                "visual studio": "https://visualstudio.microsoft.com/downloads/",
                "visual studio code": "https://code.visualstudio.com/download",
                "git": "https://git-scm.com/download/win",
                "python": "https://www.python.org/downloads/windows/",
                "nodejs": "https://nodejs.org/en/download/",
                "docker": "https://www.docker.com/products/docker-desktop/",
                "chrome": "https://www.google.com/chrome/",
                "firefox": "https://www.mozilla.org/en-US/firefox/new/",
                "edge": "https://www.microsoft.com/en-us/edge/download",
                "brave": "https://brave.com/download/",
                "opera": "https://www.opera.com/download",
                "adobe reader": "https://get.adobe.com/reader/",
                "adobe acrobat": "https://www.adobe.com/acrobat/acrobat-pro.html",
                "photoshop": "https://www.adobe.com/products/photoshop.html",
                "illustrator": "https://www.adobe.com/products/illustrator.html",
                "premiere pro": "https://www.adobe.com/products/premiere.html",
                "after effects": "https://www.adobe.com/products/aftereffects.html",
                "steam": "https://store.steampowered.com/about/",
                "epic games": "https://www.epicgames.com/store/en-US/download",
                "origin": "https://www.origin.com/usa/en-us/store/download",
                "battle.net": "https://www.blizzard.com/en-us/apps/battle.net/desktop",
                "minecraft": "https://www.minecraft.net/en-us/download",
                "7zip": "https://www.7-zip.org/download.html",
                "winrar": "https://www.win-rar.com/download.html",
                "ccleaner": "https://www.ccleaner.com/ccleaner/download",
                "vlc": "https://www.videolan.org/vlc/download-windows.html",
                "obs studio": "https://obsproject.com/download",
                "obs": "https://obsproject.com/download",
                "teamviewer": "https://www.teamviewer.com/en-us/download/windows/",
                "anydesk": "https://anydesk.com/en/downloads/windows"
            }
            app_name_lower = app_name.lower().strip()
            if app_name_lower in download_pages:
                self.speak(f"I found the official download page for {app_name}. Opening it for you.")
                self.open_download_page(download_pages[app_name_lower], app_name)
                return True
            for key, url in download_pages.items():
                if key in app_name_lower or app_name_lower in key:
                    self.speak(f"I found the official download page for {app_name}. Opening it for you.")
                    self.open_download_page(url, app_name)
                    return True
            import google.generativeai as genai
            api_key = self.load_setting('gemini_api_key')
            model_name = self.load_setting("gemini_model")
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel(model_name)
            prompt = f"""
            Find the official download page URL for {app_name} for Windows 10/11.
            IMPORTANT:
            - Return only the URL of the official download page
            - Make sure it's a page where the user can directly download or access the download
            - Do not include any explanations, just the URL
            """
            response = model.generate_content(prompt)
            download_page_url = response.text.strip()
            logging.info(f"Obtained download page URL for {app_name}: {download_page_url}")
            if download_page_url.startswith(('http://', 'https://')):
                self.speak(f"I found a download page for {app_name}. Opening it for you.")
                self.open_download_page(download_page_url, app_name)
                return True
            else:
                return False
        except Exception as e:
            self.handle_errors(e, f"Error getting download page for {app_name}")
            return False
    def open_download_page(self, url, app_name):
        """
        Opens a download page in the default web browser and provides installation guidance
        """
        try:
            webbrowser.open(url)
            app_name_lower = app_name.lower()
            if "microsoft" in app_name_lower or "office" in app_name_lower:
                self.speak(f"Opening the official {app_name} download page. You may need to sign in with your Microsoft account to download.")
            elif "adobe" in app_name_lower:
                self.speak(f"Opening the official Adobe {app_name} page. You may need an Adobe account to download.")
            elif "whatsapp" in app_name_lower:
                self.speak(f"Opening the WhatsApp download page. Click the green download button to get the installer.")
            elif "telegram" in app_name_lower:
                self.speak(f"Opening the Telegram download page. Look for the Windows download option.")
            elif "visual studio" in app_name_lower:
                self.speak(f"Opening the Visual Studio download page. You'll need to select which components you want to install.")
            else:
                self.speak(f"Opening the download page for {app_name}.")
        except Exception as e:
            self.handle_errors(e, f"Error opening download page for {app_name}")
            self.speak(f"I encountered an error while trying to open the download page for {app_name}. Please try searching for it manually.")
    def download_app_direct(self, app_name):
        """
        Enhanced download application using direct link, handling various edge cases properly
        """
        try:
            self.speak(f"Looking for the latest Windows compatible version of {app_name}")
            processed_app_name = app_name.lower().strip()
            problematic_apps = ["whatsapp", "telegram", "microsoft teams", "onedrive", "facebook", 
                               "skype", "zoom", "slack", "discord"]
            if any(prob_app in processed_app_name for prob_app in problematic_apps):
                self.speak(f"{app_name} often has download issues. Let me open the official download page for you.")
                if self.get_download_page_url(app_name):
                    return
            if "visual studio" in processed_app_name and "code" not in processed_app_name:
                self.speak("Visual Studio requires a special installer. Opening the download page.")
                webbrowser.open("https://visualstudio.microsoft.com/downloads/")
                return
            download_link = self.get_direct_download_link(app_name)
            if not download_link:
                if self.get_download_page_url(app_name):
                    return
                self.auto_install_app(app_name)
                return
            progress_window = tk.Toplevel(self.root)
            progress_window.title(f"Downloading {app_name}")
            progress_window.geometry("450x200")
            progress_window.configure(bg='#0a0a1f')
            progress_window.resizable(False, False)
            frame = self.create_tron_style_frame(progress_window)
            frame.pack(padx=10, pady=10, fill="both", expand=True)
            header = tk.Label(frame, text=f"Downloading {app_name}...", bg='#0a0a1f', fg='#00CCFF', font=("Arial", 12, "bold"))
            header.pack(pady=10)
            progress_var = tk.DoubleVar()
            style = ttk.Style()
            style.theme_use('default')
            style.configure("Tron.Horizontal.TProgressbar", 
                            troughcolor='#0a0a1f', 
                            background='#00CCFF',
                            borderwidth=0,
                            thickness=10)
            progress_bar = ttk.Progressbar(frame, variable=progress_var, length=400, style="Tron.Horizontal.TProgressbar")
            progress_bar.pack(pady=10)
            status_label = tk.Label(frame, text="Initializing download...", bg='#0a0a1f', fg='#00CCFF')
            status_label.pack(pady=5)
            window_exists = True
            def on_window_close():
                nonlocal window_exists
                window_exists = False
                progress_window.destroy()
            progress_window.protocol("WM_DELETE_WINDOW", on_window_close)
            def update_progress(percentage, status):
                nonlocal window_exists
                if not window_exists:
                    return
                try:
                    if progress_bar.winfo_exists() and status_label.winfo_exists():
                        progress_var.set(percentage)
                        status_label.configure(text=status)
                        progress_window.update()
                except Exception as e:
                    window_exists = False
                    logging.debug(f"Progress update error (widget might be destroyed): {str(e)}")
            def download_thread():
                try:
                    downloads_dir = os.path.join(os.path.expanduser("~"), "Downloads")
                    ensure_directory_exists(downloads_dir)
                    file_name = os.path.basename(urlparse(download_link).path)
                    if not file_name or '?' in file_name or len(file_name) < 4:
                        file_name = f"{app_name.replace(' ', '_')}_installer.exe"
                    if not file_name.lower().endswith(('.exe', '.msi', '.zip')):
                        file_name = f"{os.path.splitext(file_name)[0]}.exe"
                    output_path = os.path.join(downloads_dir, file_name)
                    if os.path.exists(output_path):
                        base_name, extension = os.path.splitext(file_name)
                        from datetime import datetime
                        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
                        file_name = f"{base_name}_{timestamp}{extension}"
                        output_path = os.path.join(downloads_dir, file_name)
                    update_progress(5, "Starting download...")
                    session = requests.Session()
                    retries = Retry(total=5, backoff_factor=0.5, status_forcelist=[500, 502, 503, 504])
                    session.mount('http://', HTTPAdapter(max_retries=retries))
                    session.mount('https://', HTTPAdapter(max_retries=retries))
                    headers = {
                        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
                        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                        'Accept-Language': 'en-US,en;q=0.5',
                        'DNT': '1',
                        'Connection': 'keep-alive',
                        'Upgrade-Insecure-Requests': '1'
                    }
                    download_success = False
                    try:
                        if not window_exists:
                            return
                        update_progress(10, "Connecting to server...")
                        response = session.get(download_link, headers=headers, stream=True, timeout=30)
                        response.raise_for_status()
                        if 'content-length' in response.headers and int(response.headers.get('content-length', 0)) > 1000:
                            total_size = int(response.headers.get('content-length', 0))
                            temp_path = output_path + ".download"
                            update_progress(15, "Download started...")
                            downloaded = 0
                            with open(temp_path, 'wb') as f:
                                for chunk in response.iter_content(chunk_size=8192):
                                    if not window_exists:
                                        return
                                    if chunk:
                                        f.write(chunk)
                                        downloaded += len(chunk)
                                        percentage = min(95, int(downloaded * 95 / total_size))
                                        update_progress(percentage, f"Downloaded: {percentage}% ({round(downloaded/1048576, 1)} MB / {round(total_size/1048576, 1)} MB)")
                            if os.path.exists(temp_path) and os.path.getsize(temp_path) > 1000:
                                shutil.move(temp_path, output_path)
                                download_success = True
                                update_progress(100, "Download complete!")
                            else:
                                raise Exception("Downloaded file is too small, likely corrupted")
                        else:
                            raise Exception("Server didn't provide valid content length")
                    except Exception as req_error:
                        if not window_exists:
                            return
                        logging.warning(f"Requests download failed: {str(req_error)}. Trying method 2.")
                        update_progress(20, "Retrying with alternative method...")
                        try:
                            opener = urllib.request.build_opener()
                            opener.addheaders = [
                                ('User-Agent', 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'),
                                ('Accept', 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8')
                            ]
                            temp_path = output_path + ".download"
                            with opener.open(download_link, timeout=30) as response:
                                meta = response.info()
                                total_size = int(meta.get('Content-Length', 0))
                                if total_size < 1000 and 'Content-Length' in meta:
                                    raise Exception("File too small, likely not a valid installer")
                                with open(temp_path, 'wb') as out_file:
                                    if total_size == 0:
                                        if not window_exists:
                                            return
                                        update_progress(50, "Size unknown, downloading...")
                                        data = response.read()
                                        out_file.write(data)
                                        if len(data) < 1000:
                                            raise Exception("Downloaded data too small, likely not a valid installer")
                                    else:
                                        downloaded = 0
                                        block_size = 8192
                                        while True:
                                            if not window_exists:
                                                return
                                            buffer = response.read(block_size)
                                            if not buffer:
                                                break
                                            downloaded += len(buffer)
                                            out_file.write(buffer)
                                            percentage = min(95, int(downloaded * 95 / total_size))
                                            update_progress(percentage, f"Downloaded: {percentage}% ({round(downloaded/1048576, 1)} MB / {round(total_size/1048576, 1)} MB)")
                            if os.path.exists(temp_path) and os.path.getsize(temp_path) > 1000:
                                shutil.move(temp_path, output_path)
                                download_success = True
                                update_progress(100, "Download complete!")
                            else:
                                raise Exception("Downloaded file is too small, likely corrupted")
                        except Exception as urllib_error:
                            if not window_exists:
                                return
                            logging.warning(f"urllib download failed: {str(urllib_error)}. Trying method 3 (PowerShell).")
                            update_progress(30, "Trying PowerShell download method...")
                            try:
                                temp_path = output_path + ".download"
                                ps_cmd = f"""
                                $ProgressPreference = 'SilentlyContinue'
                                Invoke-WebRequest -Uri "{download_link}" -OutFile "{temp_path}" -UseBasicParsing
                                """
                                process = subprocess.Popen(["powershell", "-Command", ps_cmd], 
                                                         shell=True, 
                                                         stdout=subprocess.PIPE,
                                                         stderr=subprocess.PIPE)
                                for i in range(30, 95, 5):
                                    time.sleep(1)
                                    if not window_exists:
                                        process.kill()
                                        return
                                    if process.poll() is not None:
                                        break
                                    update_progress(i, f"Downloading with PowerShell... {i}%")
                                process.wait()
                                if os.path.exists(temp_path) and os.path.getsize(temp_path) > 1000:
                                    shutil.move(temp_path, output_path)
                                    download_success = True
                                    update_progress(100, "Download complete!")
                                else:
                                    raise Exception("PowerShell download failed or file is corrupted")
                            except Exception as ps_error:
                                logging.error(f"PowerShell download failed: {str(ps_error)}")
                                raise Exception(f"All download methods failed. Please try manual installation.")
                    if window_exists and download_success and os.path.exists(output_path) and os.path.getsize(output_path) > 1000:
                        update_progress(100, f"Download successful: {file_name}")
                        self.speak(f"{app_name} has been downloaded successfully to your Downloads folder")
                        if window_exists and messagebox.askyesno("Installation", f"Would you like to run the {app_name} installer now?"):
                            try:
                                if os.path.exists(output_path) and os.path.getsize(output_path) > 1000:
                                    if output_path.lower().endswith('.exe'):
                                        subprocess.Popen(f'"{output_path}"', shell=True)
                                    elif output_path.lower().endswith('.msi'):
                                        subprocess.Popen(f'msiexec /i "{output_path}"', shell=True)
                                    elif output_path.lower().endswith('.zip'):
                                        subprocess.Popen(f'explorer /select,"{output_path}"', shell=True)
                                    else:
                                        subprocess.Popen(f'explorer /select,"{output_path}"', shell=True)
                                else:
                                    raise Exception("File verification failed before execution")
                            except Exception as exec_error:
                                logging.error(f"Error executing installer: {str(exec_error)}")
                                subprocess.Popen(f'explorer "{downloads_dir}"', shell=True)
                                self.speak(f"I couldn't automatically run the installer. I've opened the Downloads folder for you to run it manually.")
                    else:
                        raise Exception("Final verification failed. File may be corrupted or incomplete.")
                except Exception as e:
                    if not window_exists:
                        return
                    update_progress(0, f"Error: {str(e)}")
                    self.handle_errors(e, f"Error downloading {app_name}")
                    self.speak(f"I encountered a problem downloading {app_name}. Trying alternative installation method.")
                    if self.get_download_page_url(app_name):
                        pass
                    else:
                        self.auto_install_app(app_name)
                finally:
                    if window_exists:
                        try:
                            close_btn = self.create_tron_button(frame, "Close", progress_window.destroy, width=15)
                            close_btn.pack(pady=10)
                        except Exception as btn_error:
                            logging.debug(f"Could not create close button: {str(btn_error)}")
            threading.Thread(target=download_thread, daemon=True).start()
        except Exception as e:
            self.handle_errors(e, f"Error setting up download for {app_name}")
            self.speak(f"I encountered an error while trying to download {app_name}. Trying alternative installation method.")
            if not self.get_download_page_url(app_name):
                self.auto_install_app(app_name)
    def auto_install_app(self, app_name):
        """Auto install application using winget in PowerShell, bypassing permissions. 
        Shows download progress and saves to downloads folder."""
        try:
            problematic_apps = ["whatsapp", "telegram", "microsoft teams", "onedrive", "facebook"]
            if any(prob_app in app_name.lower() for prob_app in problematic_apps):
                if self.get_download_page_url(app_name):
                    return
            if "visual studio" in app_name.lower() and "code" not in app_name.lower():
                self.speak("Visual Studio requires a special installer. Opening the download page.")
                webbrowser.open("https://visualstudio.microsoft.com/downloads/")
                return
            if "visual studio code" in app_name.lower() or "vs code" in app_name.lower():
                self.speak("Visual Studio Code requires a special installer. Opening the download page.")
                webbrowser.open( "https://code.visualstudio.com/sha/download?build=stable&os=win32-x64-user")
                return
            progress_window = tk.Toplevel(self.root)
            progress_window.title(f"Installing {app_name}")
            progress_window.geometry("400x150")
            progress_window.configure(bg='#0a0a1f')
            progress_window.resizable(False, False)
            progress_window.lift()
            progress_window.focus_force()
            Label(progress_window, text=f"Installing {app_name}...", font=("Segoe UI", 12), bg='#0a0a1f', fg='white').pack(pady=(10, 5))
            progress_label = Label(progress_window, text="Initializing...", font=("Segoe UI", 10), bg='#0a0a1f', fg='white')
            progress_label.pack(pady=(0, 5))
            progress_bar = ttk.Progressbar(progress_window, orient="horizontal", length=350, mode="determinate")
            progress_bar.pack(pady=(0, 10))
            downloads_folder = os.path.join(os.path.expanduser("~"), "Downloads")
            if not os.path.exists(downloads_folder):
                os.makedirs(downloads_folder)
            window_exists = True
            def on_window_close():
                nonlocal window_exists
                window_exists = False
                progress_window.destroy()
            progress_window.protocol("WM_DELETE_WINDOW", on_window_close)
            def update_progress(percentage, status):
                nonlocal window_exists
                if not window_exists:
                    return
                try:
                    if progress_bar.winfo_exists() and progress_label.winfo_exists():
                        progress_bar["value"] = percentage
                        progress_label.configure(text=status)
                        progress_window.update()
                except Exception as e:
                    window_exists = False
                    logging.debug(f"Progress update error (widget might be destroyed): {str(e)}")
            def install_thread():
                try:
                    update_progress(5, "Initializing installation...")
                    ai_suggestion = self.get_installation_suggestion(app_name)
                    if ai_suggestion and "direct_url" in ai_suggestion:
                        direct_url = ai_suggestion["direct_url"]
                        update_progress(10, f"Downloading {app_name} from direct URL...")
                        file_extension = ".exe"
                        if "file_extension" in ai_suggestion:
                            file_extension = ai_suggestion["file_extension"]
                        download_path = os.path.join(downloads_folder, f"{app_name}{file_extension}")
                        with urllib.request.urlopen(direct_url) as response, open(download_path, 'wb') as out_file:
                            file_size = int(response.info().get('Content-Length', 0))
                            downloaded = 0
                            block_size = 8192
                            while True:
                                if not window_exists:
                                    return
                                buffer = response.read(block_size)
                                if not buffer:
                                    break
                                downloaded += len(buffer)
                                out_file.write(buffer)
                                if file_size > 0:
                                    percent = min(int(downloaded / file_size * 100), 100)
                                    update_progress(10 + int(percent * 0.7), f"Downloading: {percent}%")
                        if not window_exists:
                            return
                        update_progress(80, "Download complete. Starting installation...")
                        if download_path.endswith('.exe'):
                            if "requires_admin" in ai_suggestion and ai_suggestion["requires_admin"]:
                                powershell_cmd = f'powershell -Command "Start-Process -FilePath \'{download_path}\' -Verb RunAs"'
                            else:
                                powershell_cmd = f'powershell -Command "Start-Process -FilePath \'{download_path}\'"'
                            subprocess.Popen(powershell_cmd, shell=True)
                            update_progress(100, "Installation started. Follow the installer prompts.")
                        elif download_path.endswith('.msi'):
                            powershell_cmd = f'powershell -Command "Start-Process msiexec.exe -ArgumentList \'/i \"{download_path}\"\' -Wait"'
                            subprocess.Popen(powershell_cmd, shell=True)
                            update_progress(100, "Installation started. Follow the installer prompts.")
                        self.speak(f"{app_name} has been downloaded to your Downloads folder and installation has started.")
                    else:
                        update_progress(10, "Searching for application with winget...")
                        temp_script_path = os.path.join(downloads_folder, f"install_{app_name}.ps1")
                        ps_script = f"""
                        $downloadPath = "{downloads_folder.replace('\\', '\\\\')}"
                        Write-Output "Searching for {app_name}..."
                        $searchResults = winget search "{app_name}" --accept-source-agreements
                        $lines = $searchResults -split "`n" | Where-Object {{ $_ -match "{app_name}" }}
                        if ($lines.Count -eq 0) {{
                            $lines = $searchResults -split "`n" | Select-Object -Skip 2
                        }}
                        $appId = $null
                        if ($lines.Count -gt 0) {{
                            $appId = ($lines[0] -split "\\s+")[0]
                            Write-Output "Found app ID: $appId"
                        }}
                        if ($appId) {{
                            Write-Output "Installing $appId to $downloadPath"
                            $process = Start-Process -FilePath "winget" -ArgumentList "install", "$appId", "--location", "$downloadPath", "-e", "--accept-package-agreements", "--accept-source-agreements" -NoNewWindow -PassThru -Wait
                            if ($process.ExitCode -eq 0) {{
                                Write-Output "Installation completed successfully: 100%"
                            }} else {{
                                Write-Output "Installation failed with exit code: $($process.ExitCode)"
                            }}
                        }} else {{
                            Write-Output "App '{app_name}' not found."
                        }}
                        """
                        with open(temp_script_path, 'w') as f:
                            f.write(ps_script)
                        if not window_exists:
                            try:
                                if os.path.exists(temp_script_path):
                                    os.remove(temp_script_path)
                            except:
                                pass
                            return
                        powershell_cmd = f'powershell -ExecutionPolicy Bypass -File "{temp_script_path}"'
                        process = subprocess.Popen(powershell_cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, shell=True)
                        found_app = False
                        installing = False
                        for line in iter(process.stdout.readline, ''):
                            if not window_exists:
                                try:
                                    process.terminate()
                                except:
                                    pass
                                try:
                                    if os.path.exists(temp_script_path):
                                        os.remove(temp_script_path)
                                except:
                                    pass
                                return
                            print(line)
                            if 'Searching' in line:
                                update_progress(15, "Searching for application...")
                            elif 'Found app ID' in line:
                                found_app = True
                                update_progress(20, "Application found, preparing to install...")
                            elif 'Installing' in line and found_app:
                                installing = True
                                update_progress(30, "Starting installation...")
                            elif installing:
                                current = progress_bar["value"]
                                if current < 90:
                                    update_progress(current + 1, "Installing...")
                            elif 'completed successfully' in line:
                                update_progress(100, "Installation complete!")
                                self.speak(f"{app_name} has been successfully installed to your Downloads folder.")
                            elif 'failed' in line or 'not found' in line:
                                update_progress(100, "Installation failed. See console for details.")
                                self.speak(f"Failed to install {app_name}.")
                        try:
                            if os.path.exists(temp_script_path):
                                os.remove(temp_script_path)
                        except:
                            pass
                    time.sleep(3)
                    if window_exists:
                        progress_window.destroy()
                except Exception as e:
                    if not window_exists:
                        return
                    update_progress(100, f"Error: {str(e)}")
                    self.speak(f"Error installing ,try to get download page url")
                    if window_exists:
                        if self.get_download_page_url(app_name):
                            time.sleep(3)
                            progress_window.destroy()
                        else:
                            self.download_app_direct(app_name)
                            time.sleep(3)
                            progress_window.destroy()
            threading.Thread(target=install_thread, daemon=True).start()
        except Exception as e:
            self.handle_errors(e, f"Error setting up installation for {app_name}")
    def get_installation_suggestion(self, app_name):
        """Use AI to determine the best way to install an app"""
        try:
            if hasattr(self, 'gemini_model') and self.gemini_model:
                prompt = f"""
                I need to install the software "{app_name}" on a Windows computer. 
                Provide installation information in JSON format with these fields:
                - best_method: "direct_download" or "winget" or "microsoft_store"
                - direct_url: direct download URL if available
                - file_extension: ".exe" or ".msi" etc.
                - requires_admin: true/false
                - installation_notes: any special instructions
                Only respond with valid JSON. If you're uncertain, only include the fields you're confident about.
                """
                response = self.gemini_model.generate_content(prompt)
                response_text = response.text
                try:
                    import json
                    import re
                    json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
                    if json_match:
                        json_str = json_match.group(0)
                        return json.loads(json_str)
                except:
                    print("Failed to parse AI response as JSON")
                    return None
            return None
        except Exception as e:
            print(f"Error getting installation suggestion: {str(e)}")
            return None
    def update_system_stats(self):
        """Update system statistics"""
        try:
            cpu_percent = psutil.cpu_percent(interval=0.1)
            self.cpu_label.configure(text=f"CPU: {cpu_percent}%")
            memory = psutil.virtual_memory()
            self.memory_label.configure(text=f"RAM: {memory.percent}%")
            self.root.after(1000, self.update_system_stats)
        except Exception as e:
            logging.error(f"Error updating system stats: {str(e)}")
    def change_theme(self, theme_name):
        """Change the application theme"""
        if theme_name in self.themes:
            self.current_theme = theme_name
            self.apply_theme(theme_name)
    def apply_theme(self, theme_name):
        """Apply the selected theme to all widgets"""
        theme = self.themes[theme_name]
        self.root.configure(fg_color=theme["bg"])
        self.main_container.configure(fg_color=theme["bg"])
        self.banner_frame.configure(fg_color=theme["bg"])
        self.gradient_frame.configure(
            fg_color=theme["bg"],
            border_color=theme["highlight"]
        )
        self.wave_frame.configure(fg_color=theme["bg"])
        self.header_label.configure(
            fg_color=theme["bg"],
            text_color=theme["accent"]
        )
        self.subtitle_label.configure(
            fg_color=theme["bg"],
            text_color=theme["secondary"]
        )
        self.status_label.configure(
            fg_color=theme["bg"],
            text_color=theme["error"] if "Not Listening" in self.status_label.cget("text") else theme["success"]
        )
        self.settings_button.configure(
            fg_color=theme["bg"],
            text_color=theme["secondary"]
        )
        self.output_text.configure(
            fg_color=theme["text_bg"],
            text_color=theme["fg"]
        )
        self.cpu_label.configure(
            fg_color=theme["bg"],
            text_color=theme["secondary"]
        )
        self.memory_label.configure(
            fg_color=theme["bg"],
            text_color=theme["secondary"]
        )
        self.time_label.configure(
            fg_color=theme["bg"],
            text_color=theme["secondary"]
        )
    def change_volume(self, value):
        """Change assistant volume"""
        try:
            volume = float(value) / 100
            engine.setProperty('volume', volume)
        except Exception as e:
            logging.error(f"Error changing volume: {str(e)}")
    def handle_shortcuts(self, command):
        """Handle various shortcut commands with minimal delay"""
        try:
            shortcuts = {
                'save as': lambda: pyautogui.hotkey('ctrl', 'shift', 's'),
                'save': lambda: pyautogui.hotkey('ctrl', 's'),
                'new text file': lambda: pyautogui.hotkey('ctrl', 'n'),
                'new text window': lambda: pyautogui.hotkey('ctrl', 'shift', 'n'),
                'open file': lambda: pyautogui.hotkey('ctrl', 'o'),
                'open folder': lambda: pyautogui.hotkey('ctrl', 'shift', 'o'),
                'undo': lambda: pyautogui.hotkey('ctrl', 'z'),
                'select all': lambda: pyautogui.hotkey('ctrl', 'a'),
                'run window': lambda: pyautogui.hotkey('win', 'r'),
                'copy': lambda: pyautogui.hotkey('ctrl', 'c'),
                'copy all': lambda: (pyautogui.hotkey('ctrl', 'a'), pyautogui.hotkey('ctrl', 'c')),
                'paste': lambda: pyautogui.hotkey('ctrl', 'v'),
                'cut': lambda: pyautogui.hotkey('ctrl', 'x'),
                'find': lambda: pyautogui.hotkey('ctrl', 'f'),
                'close window': lambda: pyautogui.hotkey('alt', 'f4'),
                'switch window': lambda: pyautogui.hotkey('alt', 'tab'),
                'minimize all': lambda: pyautogui.hotkey('win', 'm'),
                'task manager': lambda: pyautogui.hotkey('ctrl', 'shift', 'esc'),
                'lock screen': lambda: pyautogui.hotkey('win', 'l')
            }
            pyautogui.PAUSE = 0.1
            pyautogui.MINIMUM_DURATION = 0
            for shortcut, action in shortcuts.items():
                if shortcut in command:
                    action()
                    self.speak(f"Done")
                    return True
            return False
        except Exception as e:
            logging.error(f"Error executing shortcut: {str(e)}")
            self.speak("Error with shortcut")
            return False
    def open_system_app(self, command):
        """Open system applications with proper path handling"""
        office_paths = {
            '64bit': [
                "C:\\Program Files\\Microsoft Office\\root\\Office16",
                "C:\\Program Files\\Microsoft Office\\Office16",
                "C:\\Program Files (x86)\\Microsoft Office\\root\\Office16",
                "C:\\Program Files (x86)\\Microsoft Office\\Office16"
            ]
        }
        app_commands = {
            'word': {
                'name': 'Microsoft Word',
                'cmd': 'WINWORD.EXE',
                'paths': office_paths['64bit']
            },
            'excel': {
                'name': 'Microsoft Excel',
                'cmd': 'EXCEL.EXE',
                'paths': office_paths['64bit']
            },
            'powerpoint': {
                'name': 'Microsoft PowerPoint',
                'cmd': 'POWERPNT.EXE',
                'paths': office_paths['64bit']
            }
        }
        try:
            command = command.lower()
            if command in app_commands:
                app_info = app_commands[command]
                executable_found = False
                for base_path in app_info['paths']:
                    full_path = os.path.join(base_path, app_info['cmd'])
                    if os.path.exists(full_path):
                        subprocess.Popen(full_path)
                        executable_found = True
                        break
                if not executable_found:
                    error_msg = f"Could not find {app_info['name']}. Please verify it is installed correctly."
                    logging.error(error_msg)
                    messagebox.showerror("Application Error", error_msg)
            else:
                pass
        except Exception as e:
            error_msg = f"Error opening application: {str(e)}"
            logging.error(error_msg)
            messagebox.showerror("Error", error_msg)
    def search_and_click_first_result(self, search_term):
        """Ultra-fast Windows search and click"""
        try:
            pyautogui.PAUSE = 0.1
            pyautogui.MINIMUM_DURATION = 0
            pyautogui.hotkey('win', 's')
            time.sleep(0.1)
            pyautogui.write(search_term, interval=0.01)
            time.sleep(0.2)
            pyautogui.press('enter')
            return True
        except Exception as e:
            logging.error(f"Search error: {str(e)}")
            self.speak("Search error")
            return False
    def create_tron_style_frame(self, parent, has_glow=True):
        """Create a TRON Legacy styled frame"""
        frame = ctk.CTkFrame(
            parent,
            fg_color='#0c1221',  # Dark blue background
            border_color='#00f6ff',  # Neon blue border
            border_width=1 if has_glow else 0,
            corner_radius=0
        )
        if has_glow:
            glow = ctk.CTkFrame(
                frame,
                fg_color='#00f6ff',  # Neon blue glow
                height=2,
                corner_radius=0
            )
            glow.pack(fill='x', side='top')
            bottom_glow = ctk.CTkFrame(
                frame,
                fg_color='#00f6ff',  # Neon blue glow
                height=2,
                corner_radius=0
            )
            bottom_glow.pack(fill='x', side='bottom')
        return frame
    def create_tron_button(self, parent, text, command, width=None):
        """Create a TRON Legacy styled button"""
        button_frame = ctk.CTkFrame(
            parent,
            fg_color='#0c1221',
            border_color='#00f6ff',
            border_width=1,
            corner_radius=0
        )
        button = ctk.CTkButton(
            button_frame,
            text=text,
            command=command,
            fg_color='#0c1221',
            text_color='#00f6ff',
            font=('Orbitron', 10, 'bold'),
            corner_radius=0,
            hover_color='#0f1a2b',
            width=width if width else 100,
            height=30,
            border_width=0
        )
        button.pack(padx=1, pady=1)
        def on_enter(e):
            button_frame.configure(border_color='#00ffff')
            button.configure(fg_color='#0f1a2b')
        def on_leave(e):
            button_frame.configure(border_color='#00f6ff')
            button.configure(fg_color='#0c1221')
        button.bind('<Enter>', on_enter)
        button.bind('<Leave>', on_leave)
        return button_frame
    def change_volume(self, value):
        """Change assistant volume"""
        try:
            volume = float(value) / 100
            engine.setProperty('volume', volume)
        except Exception as e:
            logging.error(f"Error changing volume: {str(e)}")
    def handle_shortcuts(self, command):
        """Handle various shortcut commands with minimal delay"""
        try:
            shortcuts = {
                'save as': lambda: pyautogui.hotkey('ctrl', 'shift', 's'),
                'save': lambda: pyautogui.hotkey('ctrl', 's'),
                'new text file': lambda: pyautogui.hotkey('ctrl', 'n'),
                'new text window': lambda: pyautogui.hotkey('ctrl', 'shift', 'n'),
                'open file': lambda: pyautogui.hotkey('ctrl', 'o'),
                'open folder': lambda: pyautogui.hotkey('ctrl', 'shift', 'o'),
                'undo': lambda: pyautogui.hotkey('ctrl', 'z'),
                'select all': lambda: pyautogui.hotkey('ctrl', 'a'),
                'run window': lambda: pyautogui.hotkey('win', 'r'),
                'copy': lambda: pyautogui.hotkey('ctrl', 'c'),
                'copy all': lambda: (pyautogui.hotkey('ctrl', 'a'), pyautogui.hotkey('ctrl', 'c')),
                'paste': lambda: pyautogui.hotkey('ctrl', 'v'),
                'cut': lambda: pyautogui.hotkey('ctrl', 'x'),
                'find': lambda: pyautogui.hotkey('ctrl', 'f'),
                'close window': lambda: pyautogui.hotkey('alt', 'f4'),
                'switch window': lambda: pyautogui.hotkey('alt', 'tab'),
                'minimize all': lambda: pyautogui.hotkey('win', 'm'),
                'task manager': lambda: pyautogui.hotkey('ctrl', 'shift', 'esc'),
                'lock screen': lambda: pyautogui.hotkey('win', 'l')
            }
            pyautogui.PAUSE = 0.1
            pyautogui.MINIMUM_DURATION = 0
            for shortcut, action in shortcuts.items():
                if shortcut in command:
                    action()
                    self.speak(f"Done")
                    return True
            return False
        except Exception as e:
            logging.error(f"Error executing shortcut: {str(e)}")
            self.speak("Error with shortcut")
            return False
    def open_system_app(self, command):
        """Open system applications with proper path handling"""
        office_paths = {
            '64bit': [
                "C:\\Program Files\\Microsoft Office\\root\\Office16",
                "C:\\Program Files\\Microsoft Office\\Office16",
                "C:\\Program Files (x86)\\Microsoft Office\\root\\Office16",
                "C:\\Program Files (x86)\\Microsoft Office\\Office16"
            ]
        }
        app_commands = {
            'word': {
                'name': 'Microsoft Word',
                'cmd': 'WINWORD.EXE',
                'paths': office_paths['64bit']
            },
            'excel': {
                'name': 'Microsoft Excel',
                'cmd': 'EXCEL.EXE',
                'paths': office_paths['64bit']
            },
            'powerpoint': {
                'name': 'Microsoft PowerPoint',
                'cmd': 'POWERPNT.EXE',
                'paths': office_paths['64bit']
            }
        }
        try:
            command = command.lower()
            if command in app_commands:
                app_info = app_commands[command]
                executable_found = False
                for base_path in app_info['paths']:
                    full_path = os.path.join(base_path, app_info['cmd'])
                    if os.path.exists(full_path):
                        subprocess.Popen(full_path)
                        executable_found = True
                        break
                if not executable_found:
                    error_msg = f"Could not find {app_info['name']}. Please verify it is installed correctly."
                    logging.error(error_msg)
                    messagebox.showerror("Application Error", error_msg)
            else:
                pass
        except Exception as e:
            error_msg = f"Error opening application: {str(e)}"
            logging.error(error_msg)
            messagebox.showerror("Error", error_msg)
    def search_and_click_first_result(self, search_term):
        """Ultra-fast Windows search and click"""
        try:
            pyautogui.PAUSE = 0.1
            pyautogui.MINIMUM_DURATION = 0
            pyautogui.hotkey('win', 's')
            time.sleep(0.1)
            pyautogui.write(search_term, interval=0.01)
            time.sleep(0.2)
            pyautogui.press('enter')
            return True
        except Exception as e:
            logging.error(f"Search error: {str(e)}")
            self.speak("Search error")
            return False
    def show_contact_manager(self):
        """Show the contact manager window"""
        ContactManager(self.root)
    def show_programming_window(self):
        """Show the programming window for AI code generation"""
        try:
            try:
                import google.generativeai as genai
                api_key = self.load_setting('gemini_api_key')
                model_name = self.load_setting("gemini_model")
                genai.configure(api_key=api_key)
                model = genai.GenerativeModel(model_name)
            except ImportError:
                self.speak("Please install the google-generativeai package first")
                messagebox.showerror("Missing Dependency", 
                    "Please install the required package:\npip install google-generativeai")
                return
            except Exception as e:
                self.speak("Error initializing Gemini API")
                messagebox.showerror("API Error", f"Failed to initialize Gemini API: {str(e)}")
                return
            if hasattr(self, 'programming_window') and self.programming_window is not None:
                try:
                    self.programming_window.window.lift()
                    self.programming_window.window.focus_force()
                    return
                except tk.TclError:
                    pass
            self.programming_window = ProgrammingWindow(self)
            self.speak("Opening the programming assistant")
        except Exception as e:
            self.handle_errors(e, "Error opening programming window")
            self.speak("Sorry, I encountered an error while opening the programming assistant.")
    def whatsapp_interaction(self, action, contact_name):
        """Handle WhatsApp interactions with auto-resume listening"""
        try:
            was_listening = self.start_button._listening
            if was_listening:
                self.start_button._listening = False
                self.update_status(False)
            conn = sqlite3.connect('samcontact.db')
            cursor = conn.cursor()
            contact_name = contact_name.strip().lower()
            cursor.execute("SELECT name, number FROM contacts WHERE LOWER(name)=?", (contact_name,))
            results = cursor.fetchall()
            if not results:
                cursor.execute("SELECT name, number FROM contacts WHERE LOWER(name) LIKE ?", 
                             ('%' + contact_name + '%',))
                results = cursor.fetchall()
            conn.close()
            if not results:
                self.speak(f'Contact {contact_name} not found. Please add it using the contact manager.')
                if was_listening:
                    self.root.after(1000, self.resume_listening)
                return False
            contact_name = results[0][0]
            mobile_no = results[0][1]
            mobile_no = ''.join(filter(str.isdigit, mobile_no))
            if not mobile_no.startswith('91'):
                mobile_no = '91' + mobile_no
            if action == 'message':
                self.speak("What message do you want to send?")
                message = self.take_command()
                if not message or message == "none":
                    self.speak("No message was provided")
                    if was_listening:
                        self.root.after(1000, self.resume_listening)
                    return False
                url = f"whatsapp://send?phone={mobile_no}&text={quote(message)}"
                webbrowser.open(url)
                time.sleep(3)
                pyautogui.press('enter')
                time.sleep(1)
                self.speak(f"Message sent to {contact_name}")
            elif action == 'call':
                url = f"whatsapp://call?phone={mobile_no}"
                webbrowser.open(url)
                self.speak(f"Starting WhatsApp call with {contact_name}")
            else:
                url = f"whatsapp://videocall?phone={mobile_no}"
                webbrowser.open(url)
                self.speak(f"Starting WhatsApp video call with {contact_name}")
            if was_listening:
                self.root.after(1000, self.resume_listening)
            return True
        except Exception as e:
            logging.error(f"Error in WhatsApp interaction: {str(e)}")
            self.speak(f"Sorry, I encountered an error: {str(e)}")
            if was_listening:
                self.root.after(1000, self.resume_listening)
            return False
    def resume_listening(self):
        """Resume listening mode"""
        try:
            self.start_button._listening = True
            self.start_button.configure(text="◼")
            self.update_status(True)
            if not hasattr(self, 'listening_thread') or not self.listening_thread.is_alive():
                self.listening_thread = threading.Thread(target=self.listen_for_commands)
                self.listening_thread.daemon = True
                self.listening_thread.start()
            if hasattr(self, 'wave_vis') and self.wave_vis:
                self.wave_vis.start_listening_animation()
            self.status_label.configure(text="● Listening", fg='#4CAF50')
        except Exception as e:
            logging.error(f"Error resuming listening: {str(e)}")
    def process_whatsapp_message(self, contact_name):
        """Handle WhatsApp message sending with auto-resume"""
        try:
            was_listening = self.start_button._listening
            print(f"Getting message for contact: {contact_name}")
            self.speak("What message would you like to send?")
            message = self.take_command().lower()
            if message:
                print(f"Message to send: {message}")
                result = self.whatsapp_interaction("message", contact_name)
                if was_listening:
                    self.root.after(1000, self.resume_listening)
                return result
            if was_listening:
                self.root.after(1000, self.resume_listening)
            return False
        except Exception as e:
            print(f"Error in process_whatsapp_message: {e}")
            self.speak("I couldn't process the message. Please try again.")
            if was_listening:
                self.root.after(1000, self.resume_listening)
            return False
    def find_element_with_gemini(self, image_path, element_description):
        """Use Gemini Vision API to find UI elements"""
        try:
            import google.generativeai as genai
            from PIL import Image
            api_key = self.load_setting('gemini_api_key')
            model_name = self.load_setting("gemini_model")
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel(model_name)
            image = Image.open(image_path)
            prompt = f"""
            Analyze this WhatsApp desktop screenshot and find the exact coordinates of the {element_description}.
            Return ONLY the x,y coordinates in format 'x,y' where the element is located.
            If multiple matches are found, return the most likely one.
            If no match is found, return 'not found'.
            """
            response = model.generate_content([prompt, image])
            result = response.text.strip()
            if result and result != 'not found':
                try:
                    x, y = map(int, result.split(','))
                    return (x, y)
                except:
                    return None
            return None
        except Exception as e:
            logging.error(f"Error using Gemini Vision API: {str(e)}")
            return None
    def control_brightness(self, action):
        """Control system brightness"""
        try:
            current = sbc.get_brightness()[0]
            if 'increase' in action or 'up' in action:
                new_brightness = min(current + 20, 100)
                sbc.set_brightness(new_brightness)
                self.speak(f"Brightness increased to {new_brightness} percent")
            elif 'decrease' in action or 'down' in action:
                new_brightness = max(current - 20, 0)
                sbc.set_brightness(new_brightness)
                self.speak(f"Brightness decreased to {new_brightness} percent")
            elif 'set' in action:
                try:
                    level = int(''.join(filter(str.isdigit, action)))
                    level = max(0, min(100, level))
                    sbc.set_brightness(level)
                    self.speak(f"Brightness set to {level} percent")
                except:
                    self.speak("Please specify a brightness level between 0 and 100")
            elif '' in action:
                try:
                    level = int(''.join(filter(str.isdigit, action)))
                    level = max(0, min(100, level))
                    sbc.set_brightness(level)
                    self.speak(f"Brightness set to {level} percent")
                except:
                    self.speak("Please specify a brightness level between 0 and 100")
        except Exception as e:
            self.handle_errors(e, "brightness control")
            self.speak("Sorry, I couldn't adjust the brightness")
    def control_volume(self, action):
        """Control system volume"""
        try:
            from ctypes import cast, POINTER
            from comtypes import CLSCTX_ALL
            from pycaw.pycaw import AudioUtilities, IAudioEndpointVolume
            devices = AudioUtilities.GetSpeakers()
            interface = devices.Activate(IAudioEndpointVolume._iid_, CLSCTX_ALL, None)
            volume = cast(interface, POINTER(IAudioEndpointVolume))
            current_vol = round(volume.GetMasterVolumeLevelScalar() * 100)
            if 'increase' in action or 'up' in action:
                new_vol = min(current_vol + 10, 100)
                volume.SetMasterVolumeLevelScalar(new_vol / 100, None)
                self.speak(f"Volume increased to {new_vol} percent")
            elif 'decrease' in action or 'down' in action:
                new_vol = max(current_vol - 10, 0)
                volume.SetMasterVolumeLevelScalar(new_vol / 100, None)
                self.speak(f"Volume decreased to {new_vol} percent")
            elif 'set' in action:
                try:
                    level = int(''.join(filter(str.isdigit, action)))
                    level = max(0, min(100, level))
                    volume.SetMasterVolumeLevelScalar(level / 100, None)
                    self.speak(f"Volume set to {level} percent")
                except:
                    self.speak("Please specify a volume level between 0 and 100")
            elif '' in action:
                try:
                    level = int(''.join(filter(str.isdigit, action)))
                    level = max(0, min(100, level))
                    volume.SetMasterVolumeLevelScalar(level / 100, None)
                    self.speak(f"Volume set to {level} percent")
                except:
                    self.speak("Please specify a volume level between 0 and 100")
        except Exception as e:
            self.handle_errors(e, "volume control")
            self.speak("Sorry, I couldn't adjust the volume")
    def show_data_analysis_window(self):   
       """Show data analysis window"""
       try:
            import google.generativeai as genai
            api_key = self.load_setting('gemini_api_key')
            model_name = self.load_setting("gemini_model")
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel(model_name)
            data_analysis_api_key = api_key
            genai.configure(api_key=data_analysis_api_key)
            try:
                self.data_analysis_model = genai.GenerativeModel(model_name=model_name)
            except Exception as model_error:
                logging.error(f"Error initializing experimental model: {str(model_error)}")
                self.data_analysis_model = genai.GenerativeModel(model_name="gemini-2.0-pro-02-05")
                messagebox.showinfo("Model Information", 
                                   "Using standard Gemini model as the experimental model is unavailable.")
            analysis_window = tk.Toplevel(self.root)
            analysis_window.title("Data Analysis Dashboard")
            analysis_window.geometry("1600x1000")
            analysis_window.configure(bg=self.themes['Tron Legacy']['bg'])
            analysis_window.minsize(1500, 900)
            analysis_window.transient(self.root)
            analysis_window.grab_set()
            main_container = tk.Frame(
                analysis_window,
                bg=self.themes['Tron Legacy']['bg'],
                padx=15,
                pady=15
            )
            main_container.pack(fill='both', expand=True)
            header_frame = self.create_tron_style_frame(main_container)
            header_frame.pack(fill='x', pady=(0, 15))
            title_label = tk.Label(
                header_frame,
                text="DATA ANALYSIS DASHBOARD",
                font=('Segoe UI', 16, 'bold'),
                bg=header_frame['bg'],
                fg=self.themes['Tron Legacy']['accent']
            )
            title_label.pack(side='left', padx=10, pady=10)
            select_btn = self.create_tron_button(
                header_frame,
                text="Select Dataset",
                command=lambda: self.load_dataset(analysis_window)
            )
            select_btn.pack(side='right', padx=10, pady=10)
            style = ttk.Style()
            style.configure("Custom.TNotebook", background=self.themes['Tron Legacy']['bg'])
            style.configure("Custom.TNotebook.Tab", background=self.themes['Tron Legacy']['bg'], 
                           foreground=self.themes['Tron Legacy']['fg'], padding=[10, 5])
            style.map("Custom.TNotebook.Tab", background=[("selected", self.themes['Tron Legacy']['accent'])],
                     foreground=[("selected", "#000000")])
            notebook = ttk.Notebook(main_container, style="Custom.TNotebook")
            notebook.pack(fill='both', expand=True)
            dashboard_frame = self.create_tron_style_frame(notebook, has_glow=True)
            notebook.add(dashboard_frame, text=" Dashboard ")
            left_panel = self.create_tron_style_frame(dashboard_frame)
            left_panel.pack(side='left', fill='both', expand=True, padx=(0, 5), pady=5)
            right_panel = self.create_tron_style_frame(dashboard_frame)
            right_panel.pack(side='right', fill='both', expand=True, padx=(5, 0), pady=5)
            summary_frame = self.create_tron_style_frame(left_panel)
            summary_frame.pack(fill='x', pady=(0, 10), padx=5)
            summary_label = tk.Label(
                summary_frame,
                text="DATASET SUMMARY",
                font=('Segoe UI', 12, 'bold'),
                bg=summary_frame['bg'],
                fg=self.themes['Tron Legacy']['accent']
            )
            summary_label.pack(anchor='w', padx=10, pady=5)
            self.summary_text = scrolledtext.ScrolledText(
                summary_frame,
                wrap=tk.WORD,
                height=8,
                font=('Consolas', 10),
                bg=self.themes['Tron Legacy']['text_bg'],
                fg=self.themes['Tron Legacy']['fg'],
                insertbackground=self.themes['Tron Legacy']['accent']
            )
            self.summary_text.pack(fill='both', expand=True, padx=10, pady=(0, 10))
            analysis_frame = self.create_tron_style_frame(left_panel)
            analysis_frame.pack(fill='both', expand=True, padx=5)
            analysis_label = tk.Label(
                analysis_frame,
                text="DETAILED ANALYSIS",
                font=('Segoe UI', 12, 'bold'),
                bg=analysis_frame['bg'],
                fg=self.themes['Tron Legacy']['accent']
            )
            analysis_label.pack(anchor='w', padx=10, pady=5)
            self.analysis_text = scrolledtext.ScrolledText(
                analysis_frame,
                wrap=tk.WORD,
                font=('Consolas', 10),
                bg=self.themes['Tron Legacy']['text_bg'],
                fg=self.themes['Tron Legacy']['fg'],
                insertbackground=self.themes['Tron Legacy']['accent']
            )
            self.analysis_text.pack(fill='both', expand=True, padx=10, pady=(0, 10))
            quality_frame = self.create_tron_style_frame(right_panel)
            quality_frame.pack(fill='both', expand=True, padx=5)
            quality_label = tk.Label(
                quality_frame,
                text="DATA QUALITY ASSESSMENT",
                font=('Segoe UI', 12, 'bold'),
                bg=quality_frame['bg'],
                fg=self.themes['Tron Legacy']['accent']
            )
            quality_label.pack(anchor='w', padx=10, pady=5)
            self.quality_text = scrolledtext.ScrolledText(
                quality_frame,
                wrap=tk.WORD,
                font=('Consolas', 10),
                bg=self.themes['Tron Legacy']['text_bg'],
                fg=self.themes['Tron Legacy']['fg'],
                insertbackground=self.themes['Tron Legacy']['accent']
            )
            self.quality_text.pack(fill='both', expand=True, padx=10, pady=(0, 10))
            analysis_tab_frame = self.create_tron_style_frame(notebook)
            notebook.add(analysis_tab_frame, text=" Full Analysis ")
            self.full_analysis_text = scrolledtext.ScrolledText(
                analysis_tab_frame,
                wrap=tk.WORD,
                font=('Consolas', 11),
                bg=self.themes['Tron Legacy']['text_bg'],
                fg=self.themes['Tron Legacy']['fg'],
                insertbackground=self.themes['Tron Legacy']['accent']
            )
            self.full_analysis_text.pack(fill='both', expand=True, padx=10, pady=10)
            data_view_frame = self.create_tron_style_frame(notebook)
            notebook.add(data_view_frame, text=" Data View ")
            data_controls_frame = self.create_tron_style_frame(data_view_frame)
            data_controls_frame.pack(fill='x', padx=10, pady=10)
            limit_label = tk.Label(
                data_controls_frame,
                text="Show rows:",
                font=('Segoe UI', 10),
                bg=data_controls_frame['bg'],
                fg=self.themes['Tron Legacy']['fg']
            )
            limit_label.pack(side='left', padx=(10, 5))
            self.row_limit_var = tk.StringVar(value="100")
            row_limit_options = ["50", "100", "500", "1000", "All"]
            row_limit_menu = ttk.Combobox(
                data_controls_frame, 
                textvariable=self.row_limit_var,
                values=row_limit_options,
                width=5,
                state="readonly"
            )
            row_limit_menu.pack(side='left', padx=5)
            refresh_btn = self.create_tron_button(
                data_controls_frame,
                text="Refresh View",
                command=lambda: self.refresh_data_view()
            )
            refresh_btn.pack(side='left', padx=10)
            search_frame = tk.Frame(
                data_controls_frame,
                bg=data_controls_frame['bg']
            )
            search_frame.pack(side='right', padx=10)
            search_label = tk.Label(
                search_frame,
                text="Search:",
                font=('Segoe UI', 10),
                bg=search_frame['bg'],
                fg=self.themes['Tron Legacy']['fg']
            )
            search_label.pack(side='left', padx=(0, 5))
            self.search_var = tk.StringVar()
            search_entry = tk.Entry(
                search_frame,
                textvariable=self.search_var,
                width=20,
                bg=self.themes['Tron Legacy']['text_bg'],
                fg=self.themes['Tron Legacy']['fg'],
                insertbackground=self.themes['Tron Legacy']['accent']
            )
            search_entry.pack(side='left')
            search_entry.bind('<Return>', lambda e: self.refresh_data_view())
            table_frame = self.create_tron_style_frame(data_view_frame)
            table_frame.pack(fill='both', expand=True, padx=10, pady=(0, 10))
            table_container = tk.Frame(table_frame, bg=table_frame['bg'])
            table_container.pack(fill='both', expand=True, padx=10, pady=10)
            v_scrollbar = tk.Scrollbar(table_container, orient='vertical')
            v_scrollbar.pack(side='right', fill='y')
            h_scrollbar = tk.Scrollbar(table_container, orient='horizontal')
            h_scrollbar.pack(side='bottom', fill='x')
            self.data_view_text = tk.Text(
                table_container,
                wrap='none',
                font=('Consolas', 10),
                bg=self.themes['Tron Legacy']['text_bg'],
                fg=self.themes['Tron Legacy']['fg'],
                insertbackground=self.themes['Tron Legacy']['accent'],
                yscrollcommand=v_scrollbar.set,
                xscrollcommand=h_scrollbar.set
            )
            self.data_view_text.pack(fill='both', expand=True)
            v_scrollbar.configure(command=self.data_view_text.yview)
            h_scrollbar.configure(command=self.data_view_text.xview)
            ai_bot_frame = self.create_tron_style_frame(notebook)
            notebook.add(ai_bot_frame, text=" AI Data Assistant ")
            chat_container = tk.Frame(ai_bot_frame, bg=ai_bot_frame['bg'])
            chat_container.pack(fill='both', expand=True, padx=10, pady=10)
            chat_history_frame = self.create_tron_style_frame(chat_container)
            chat_history_frame.pack(fill='both', expand=True, padx=5, pady=(0, 10))
            chat_history_label = tk.Label(
                chat_history_frame,
                text="ASK ANYTHING ABOUT YOUR DATASET",
                font=('Segoe UI', 12, 'bold'),
                bg=chat_history_frame['bg'],
                fg=self.themes['Tron Legacy']['accent']
            )
            chat_history_label.pack(anchor='w', padx=10, pady=5)
            self.chat_history_text = scrolledtext.ScrolledText(
                chat_history_frame,
                wrap=tk.WORD,
                font=('Consolas', 11),
                bg=self.themes['Tron Legacy']['text_bg'],
                fg=self.themes['Tron Legacy']['fg'],
                insertbackground=self.themes['Tron Legacy']['accent']
            )
            self.chat_history_text.pack(fill='both', expand=True, padx=10, pady=(0, 10))
            self.chat_history_text.insert(tk.END, "AI Data Assistant: ", 'bot_name')
            self.chat_history_text.insert(tk.END, "Hello! I'm your AI Data Assistant. Once you load a dataset, you can ask me anything about it. I can help with data analysis, explanations, and insights.\n\n", 'bot_message')
            input_frame = self.create_tron_style_frame(chat_container)
            input_frame.pack(fill='x', padx=5, pady=0)
            self.chat_input = tk.Entry(
                input_frame,
                font=('Segoe UI', 12),
                bg=self.themes['Tron Legacy']['text_bg'],
                fg=self.themes['Tron Legacy']['fg'],
                insertbackground=self.themes['Tron Legacy']['accent'],
                relief='flat',
                bd=10
            )
            self.chat_input.pack(side='left', fill='x', expand=True, padx=10, pady=10)
            self.chat_input.bind('<Return>', lambda e: self.process_dataset_query())
            send_btn = self.create_tron_button(
                input_frame,
                text="Ask",
                command=self.process_dataset_query,
                width=10
            )
            send_btn.pack(side='right', padx=10, pady=10)
            suggestions_frame = self.create_tron_style_frame(chat_container)
            suggestions_frame.pack(fill='x', padx=5, pady=(10, 0))
            suggestions_label = tk.Label(
                suggestions_frame,
                text="SUGGESTED QUESTIONS",
                font=('Segoe UI', 10, 'bold'),
                bg=suggestions_frame['bg'],
                fg=self.themes['Tron Legacy']['accent']
            )
            suggestions_label.pack(anchor='w', padx=10, pady=5)
            suggestions_buttons_frame = tk.Frame(
                suggestions_frame,
                bg=suggestions_frame['bg']
            )
            suggestions_buttons_frame.pack(fill='x', padx=10, pady=(0, 10))
            self.suggestions_buttons_frame = suggestions_buttons_frame
            self.default_questions = [
                "Summarize this dataset",
                "What are the key insights?",
                "Show me the data distribution",
                "What are the outliers?",
                "How should I clean this data?"
            ]
            for question in self.default_questions:
                suggestion_btn = tk.Button(
                    suggestions_buttons_frame,
                    text=question,
                    font=('Segoe UI', 9),
                    bg=self.themes['Tron Legacy']['bg'],
                    fg=self.themes['Tron Legacy']['fg'],
                    activebackground=self.themes['Tron Legacy']['accent'],
                    activeforeground='black',
                    relief='flat',
                    bd=1,
                    padx=10,
                    pady=5,
                    cursor='hand2',
                    command=lambda q=question: self.use_suggested_question(q)
                )
                suggestion_btn.pack(side='left', padx=5, pady=5)
                suggestion_btn.bind('<Enter>', lambda e, btn=suggestion_btn: btn.configure(
                    bg=self.themes['Tron Legacy']['accent'],
                    fg='black'
                ))
                suggestion_btn.bind('<Leave>', lambda e, btn=suggestion_btn: btn.configure(
                    bg=self.themes['Tron Legacy']['bg'],
                    fg=self.themes['Tron Legacy']['fg']
                ))
            for text_widget in [self.analysis_text, self.full_analysis_text, self.summary_text, 
                               self.quality_text, self.data_view_text, self.chat_history_text]:
                text_widget.tag_configure('header', foreground='#00FF00')  # Bright green for header
                text_widget.tag_configure('insight', foreground='#00F2FF')  # Light blue for insights
                text_widget.tag_configure('warning', foreground='#FFCC00')  # Yellow for warnings
                text_widget.tag_configure('error', foreground='#FF4D4D')  # Red for errors
                text_widget.tag_configure('highlight', foreground='#FF00FF')  # Magenta for highlights
                text_widget.tag_configure('column_header', foreground='#FFFFFF', background='#333333')  # White on dark gray for column headers
                text_widget.tag_configure('row_even', background='#1A1A1A')  # Slightly lighter background for even rows
                text_widget.tag_configure('row_odd', background='#121212')  # Darker background for odd rows
            self.chat_history_text.tag_configure('user_name', foreground='#4CAF50', font=('Consolas', 11, 'bold'))  # Green for user name
            self.chat_history_text.tag_configure('user_message', foreground='#FFFFFF')  # White for user messages
            self.chat_history_text.tag_configure('bot_name', foreground='#00F2FF', font=('Consolas', 11, 'bold'))  # Light blue for bot name
            self.chat_history_text.tag_configure('bot_message', foreground='#E0E0E0')  # Light gray for bot messages
            self.chat_history_text.tag_configure('code', foreground='#FFCC00', background='#1A1A1A')  # Yellow on dark gray for code
            self.current_df = None
            self.chat_input.focus_set()
       except Exception as e:
            self.handle_errors(e, "Error showing analysis window")
    def use_suggested_question(self, question):
        """Use a suggested question in the chat input"""
        self.chat_input.delete(0, tk.END)
        self.chat_input.insert(0, question)
        self.process_dataset_query()
    def process_dataset_query(self):
        """Process a query about the dataset"""
        query = self.chat_input.get().strip()
        if not query:
            return
        self.chat_input.delete(0, tk.END)
        if self.current_df is None:
            self.chat_history_text.insert(tk.END, "You: ", 'user_name')
            self.chat_history_text.insert(tk.END, query + "\n\n", 'user_message')
            self.chat_history_text.insert(tk.END, "AI Data Assistant: ", 'bot_name')
            self.chat_history_text.insert(tk.END, "Please load a dataset first before asking questions.\n\n", 'bot_message')
            self.chat_history_text.see(tk.END)
            return
        self.chat_history_text.insert(tk.END, "You: ", 'user_name')
        self.chat_history_text.insert(tk.END, query + "\n\n", 'user_message')
        self.chat_history_text.insert(tk.END, "AI Data Assistant: ", 'bot_name')
        thinking_index = self.chat_history_text.index(tk.END)
        self.chat_history_text.insert(thinking_index, "Analyzing your dataset...\n", 'bot_message')
        self.chat_history_text.see(tk.END)
        self.chat_history_text.update()
        try:
            df = self.current_df
            dataset_info = {
                'shape': df.shape,
                'columns': df.columns.tolist(),
                'dtypes': {col: str(dtype) for col, dtype in df.dtypes.items()},
                'head': df.head(5).to_dict(),
                'describe': df.describe().to_dict(),
                'missing_values': df.isnull().sum().to_dict()
            }
            prompt = f"""You are an AI Data Assistant analyzing a dataset. 
The user has asked: "{query}"
Here's information about the dataset:
- Shape: {dataset_info['shape']}
- Columns: {dataset_info['columns']}
- Data types: {dataset_info['dtypes']}
- Sample data (first 5 rows): {dataset_info['head']}
- Statistical summary: {dataset_info['describe']}
- Missing values: {dataset_info['missing_values']}
Provide a helpful, accurate, and concise response. If the user asks for calculations or specific insights that require code, include a brief code snippet in Python using pandas that would answer their question.
"""
            response = self.data_analysis_model.generate_content(prompt)
            if not response or not response.text:
                raise Exception("No response from AI model")
            self.chat_history_text.delete(thinking_index, tk.END)
            ai_response = response.text.strip()
            if "```" in ai_response:
                parts = ai_response.split("```")
                for i, part in enumerate(parts):
                    if i % 2 == 0:
                        if part.strip():
                            self.chat_history_text.insert(tk.END, part.strip() + "\n", 'bot_message')
                    else:
                        if part.startswith("python"):
                            part = part[6:].strip()
                        elif part.strip() and part.strip()[0] == "\n":
                            part = part.strip()[1:]
                        if part.strip():
                            self.chat_history_text.insert(tk.END, "\n")
                            self.chat_history_text.insert(tk.END, part.strip() + "\n", 'code')
                            self.chat_history_text.insert(tk.END, "\n")
            else:
                self.chat_history_text.insert(tk.END, ai_response + "\n", 'bot_message')
            self.chat_history_text.insert(tk.END, "\n")
            self.update_suggested_questions(query, dataset_info)
        except Exception as e:
            self.chat_history_text.delete(thinking_index, tk.END)
            error_msg = f"Sorry, I encountered an error: {str(e)}\n\n"
            self.chat_history_text.insert(tk.END, error_msg, 'error')
            self.handle_errors(e, "Error processing dataset query")
        self.chat_history_text.see(tk.END)
    def update_suggested_questions(self, last_query, dataset_info):
        """Update suggested questions based on the dataset and last query"""
        try:
            columns = dataset_info['columns']
            numeric_cols = [col for col, dtype in dataset_info['dtypes'].items() 
                           if 'int' in dtype.lower() or 'float' in dtype.lower()]
            categorical_cols = [col for col, dtype in dataset_info['dtypes'].items() 
                               if 'object' in dtype.lower() or 'category' in dtype.lower()]
            new_questions = []
            if numeric_cols:
                new_questions.append(f"What's the distribution of {numeric_cols[0]}?")
                if len(numeric_cols) > 1:
                    new_questions.append(f"Is there a correlation between {numeric_cols[0]} and {numeric_cols[1]}?")
            if categorical_cols:
                new_questions.append(f"Show me the value counts for {categorical_cols[0]}")
            new_questions.extend([
                "What data cleaning steps do you recommend?",
                "What machine learning models would work well with this data?",
                "Generate a pandas code snippet to analyze this data"
            ])
            for i, (btn, question) in enumerate(zip(self.suggestions_buttons_frame.winfo_children(), new_questions[:5])):
                btn.configure(text=question)
                btn.configure(command=lambda q=question: self.use_suggested_question(q))
        except Exception as e:
            logging.error(f"Error updating suggested questions: {str(e)}")
            pass
    def refresh_data_view(self):
        """Refresh the data view based on current settings"""
        if self.current_df is None:
            self.data_view_text.delete('1.0', tk.END)
            self.data_view_text.insert(tk.END, "No dataset loaded. Please select a dataset first.", 'error')
            return
        try:
            limit = self.row_limit_var.get()
            df = self.current_df
            search_term = self.search_var.get().strip()
            if search_term:
                mask = pd.Series(False, index=df.index)
                for col in df.columns:
                    try:
                        mask = mask | df[col].astype(str).str.contains(search_term, case=False, na=False)
                    except:
                        continue
                df = df[mask]
            if limit != "All":
                df = df.head(int(limit))
            self.data_view_text.delete('1.0', tk.END)
            if len(df) == 0:
                self.data_view_text.insert(tk.END, "No data matches your criteria.", 'warning')
                return
            col_widths = {}
            for col in df.columns:
                col_width = max(
                    len(str(col)),
                    df[col].astype(str).str.len().max() if len(df) > 0 else 0
                )
                col_widths[col] = min(col_width + 2, 50)
            header_row = ""
            for col in df.columns:
                header_row += str(col).ljust(col_widths[col])
            self.data_view_text.insert(tk.END, header_row + "\n", 'column_header')
            separator = ""
            for col in df.columns:
                separator += "-" * col_widths[col]
            self.data_view_text.insert(tk.END, separator + "\n", 'header')
            for i, (_, row) in enumerate(df.iterrows()):
                row_text = ""
                for col in df.columns:
                    cell_value = str(row[col])
                    if len(cell_value) > col_widths[col] - 2:
                        cell_value = cell_value[:col_widths[col] - 5] + "..."
                    row_text += cell_value.ljust(col_widths[col])
                tag = 'row_even' if i % 2 == 0 else 'row_odd'
                self.data_view_text.insert(tk.END, row_text + "\n", tag)
            total_rows = len(self.current_df)
            shown_rows = len(df)
            summary = f"\nShowing {shown_rows} of {total_rows} rows"
            if search_term:
                summary += f" (filtered by '{search_term}')"
            self.data_view_text.insert(tk.END, summary, 'insight')
        except Exception as e:
            self.handle_errors(e, "Error refreshing data view")
            self.data_view_text.delete('1.0', tk.END)
            self.data_view_text.insert(tk.END, f"Error displaying data: {str(e)}", 'error')
    def load_dataset(self, window):
        """Load and analyze dataset"""
        try:
            file_path = filedialog.askopenfilename(
                filetypes=[
                    ("CSV files", "*.csv"),
                    ("Excel files", "*.xlsx;*.xls"),
                    ("All files", "*.*")
                ]
            )
            if not file_path:
                return
            for text_widget in [self.analysis_text, self.full_analysis_text, self.summary_text, 
                               self.quality_text, self.data_view_text]:
                text_widget.delete('1.0', tk.END)
                text_widget.insert(tk.END, "Loading and analyzing data...", 'header')
            if file_path.endswith('.csv'):
                df = pd.read_csv(file_path)
            elif file_path.endswith(('.xlsx', '.xls')):
                df = pd.read_excel(file_path)
            else:
                self.speak("Unsupported file format")
                return
            self.current_df = df
            filename = os.path.basename(file_path)
            window.title(f"Data Analysis Dashboard - {filename}")
            for text_widget in [self.analysis_text, self.full_analysis_text, self.summary_text, 
                               self.quality_text]:
                text_widget.delete('1.0', tk.END)
            self.display_dataset_summary(df)
            self.analyze_headers(df)
            self.perform_analysis(df)
            self.assess_data_quality(df)
            self.refresh_data_view()
        except Exception as e:
            self.handle_errors(e, "Error loading dataset")
            for text_widget in [self.analysis_text, self.full_analysis_text, self.summary_text, 
                               self.quality_text, self.data_view_text]:
                text_widget.insert(tk.END, f"\nError loading dataset: {str(e)}\n", 'error')
    def display_dataset_summary(self, df):
        """Display basic dataset summary information"""
        try:
            num_rows, num_cols = df.shape
            dtypes = df.dtypes.value_counts().to_dict()
            dtype_str = ", ".join([f"{count} {dtype}" for dtype, count in dtypes.items()])
            missing_values = df.isnull().sum().sum()
            missing_percent = (missing_values / (num_rows * num_cols)) * 100
            summary = f"DATASET OVERVIEW\n"
            summary += f"{'=' * 50}\n"
            summary += f"• Rows: {num_rows:,}\n"
            summary += f"• Columns: {num_cols}\n"
            summary += f"• Data types: {dtype_str}\n"
            summary += f"• Missing values: {missing_values:,} ({missing_percent:.2f}%)\n"
            summary += f"{'=' * 50}\n\n"
            summary += "COLUMNS\n"
            summary += f"{'=' * 50}\n"
            for col, dtype in df.dtypes.items():
                missing = df[col].isnull().sum()
                missing_pct = (missing / num_rows) * 100
                summary += f"• {col} ({dtype})"
                if missing > 0:
                    summary += f" - {missing:,} missing ({missing_pct:.2f}%)"
                summary += "\n"
            self.summary_text.insert(tk.END, summary, 'insight')
        except Exception as e:
            self.handle_errors(e, "Error displaying dataset summary")
            self.summary_text.insert(tk.END, f"Error displaying summary: {str(e)}\n", 'error')
    def assess_data_quality(self, df):
        """Assess data quality and display results"""
        try:
            num_rows, num_cols = df.shape
            missing_by_col = df.isnull().sum()
            missing_pct_by_col = (missing_by_col / num_rows) * 100
            high_missing = [col for col, pct in missing_pct_by_col.items() if pct > 5]
            duplicate_rows = df.duplicated().sum()
            duplicate_pct = (duplicate_rows / num_rows) * 100
            outlier_info = []
            for col in df.select_dtypes(include=['number']).columns:
                if df[col].nunique() > 1:
                    q1 = df[col].quantile(0.25)
                    q3 = df[col].quantile(0.75)
                    iqr = q3 - q1
                    lower_bound = q1 - 1.5 * iqr
                    upper_bound = q3 + 1.5 * iqr
                    outliers = df[(df[col] < lower_bound) | (df[col] > upper_bound)][col].count()
                    if outliers > 0:
                        outlier_pct = (outliers / df[col].count()) * 100
                        outlier_info.append(f"• {col}: {outliers:,} potential outliers ({outlier_pct:.2f}%)")
            quality_text = "DATA QUALITY ASSESSMENT\n"
            quality_text += f"{'=' * 50}\n\n"
            missing_overall_pct = (df.isnull().sum().sum() / (num_rows * num_cols)) * 100
            quality_score = 100 - missing_overall_pct - duplicate_pct
            quality_text += f"Overall Quality Score: {quality_score:.1f}%\n\n"
            quality_text += "MISSING VALUES\n"
            quality_text += f"{'=' * 50}\n"
            if missing_by_col.sum() == 0:
                quality_text += "• No missing values detected\n"
            else:
                quality_text += f"• Total missing: {missing_by_col.sum():,} cells ({missing_overall_pct:.2f}%)\n"
                if high_missing:
                    quality_text += "• Columns with high missing rates:\n"
                    for col in high_missing:
                        quality_text += f"  - {col}: {missing_by_col[col]:,} ({missing_pct_by_col[col]:.2f}%)\n"
            quality_text += "\n"
            quality_text += "DUPLICATES\n"
            quality_text += f"{'=' * 50}\n"
            if duplicate_rows == 0:
                quality_text += "• No duplicate rows detected\n"
            else:
                quality_text += f"• {duplicate_rows:,} duplicate rows ({duplicate_pct:.2f}%)\n"
            quality_text += "\n"
            quality_text += "POTENTIAL OUTLIERS\n"
            quality_text += f"{'=' * 50}\n"
            if outlier_info:
                quality_text += "\n".join(outlier_info)
            else:
                quality_text += "• No significant outliers detected\n"
            self.quality_text.insert(tk.END, quality_text, 'insight')
            recommendations = "\nRECOMMENDATIONS\n"
            recommendations += f"{'=' * 50}\n"
            if high_missing:
                recommendations += "• Consider handling missing values in columns with high missing rates\n"
            if duplicate_rows > 0:
                recommendations += "• Review and handle duplicate rows\n"
            if outlier_info:
                recommendations += "• Investigate potential outliers in numeric columns\n"
            if high_missing or duplicate_rows > 0 or outlier_info:
                self.quality_text.insert(tk.END, recommendations, 'warning')
            else:
                self.quality_text.insert(tk.END, "\nRECOMMENDATIONS\n" + "=" * 50 + "\n• Data quality looks good!\n", 'highlight')
        except Exception as e:
            self.handle_errors(e, "Error assessing data quality")
            self.quality_text.insert(tk.END, f"Error assessing data quality: {str(e)}\n", 'error')
    def analyze_headers(self, df):
        """Analyze dataset headers using Gemini"""
        try:
            headers = df.columns.tolist()
            header_prompt = self.load_prompt('data_header_analysis')
            if not header_prompt:
                header_prompt = """Analyze these dataset headers and provide:
    1. Brief explanation of each header
    2. Potential relationships between headers
    3. What insights could be derived
    4. Don't Include this * "' symbols
    5. Should be structure and easy to understand
    6. Should be in bullet points
    Headers: {headers}"""
            formatted_prompt = header_prompt.format(headers=headers)
            response = self.data_analysis_model.generate_content(formatted_prompt)
            if not response or not response.text:
                raise Exception("No response from Gemini API")
            analysis_text = response.text.replace("*", "").strip()
            self.analysis_text.insert(tk.END, "=== HEADER ANALYSIS ===\n\n", 'header')
            self.analysis_text.insert(tk.END, analysis_text + "\n\n", 'insight')
            self.full_analysis_text.insert(tk.END, "=== HEADER ANALYSIS ===\n\n", 'header')
            self.full_analysis_text.insert(tk.END, analysis_text + "\n\n", 'insight')
            self.full_analysis_text.insert(tk.END, "=" * 50 + "\n\n")
        except Exception as e:
            error_msg = f"Error analyzing headers: {str(e)}"
            self.handle_errors(e, error_msg)
            self.analysis_text.insert(tk.END, f"\n{error_msg}\n", 'error')
            self.full_analysis_text.insert(tk.END, f"\n{error_msg}\n", 'error')
    def perform_analysis(self, df):
        """Perform full dataset analysis"""
        try:
            basic_stats = df.describe()
            data_summary = {
                'shape': df.shape,
                'dtypes': df.dtypes.to_dict(),
                'missing_values': df.isnull().sum().to_dict(),
                'basic_stats': basic_stats.to_dict()
            }
            stats_prompt = self.load_prompt('data_statistical_analysis')
            if not stats_prompt:
                stats_prompt = """Analyze this dataset and provide:
    1. Key insights and patterns
    2. Data quality issues
    3. Notable statistics
    4. Recommendations for further analysis
    5. Don't Include this * "' symbols
    6. Should be structure and easy to understand
    7. Should be in bullet points
    Summary: {data_summary}"""
            formatted_prompt = stats_prompt.format(data_summary=data_summary)
            response = self.data_analysis_model.generate_content(formatted_prompt)
            if not response or not response.text:
                raise Exception("No response from Gemini API")
            analysis_text = response.text.replace("*", "").strip()
            self.analysis_text.insert(tk.END, "=== DATASET ANALYSIS ===\n\n", 'header')
            self.analysis_text.insert(tk.END, analysis_text + "\n\n", 'insight')
            self.full_analysis_text.insert(tk.END, "=== DATASET ANALYSIS ===\n\n", 'header')
            self.full_analysis_text.insert(tk.END, analysis_text + "\n\n", 'insight')
            self.full_analysis_text.insert(tk.END, "=== BASIC STATISTICS ===\n\n", 'header')
            self.full_analysis_text.insert(tk.END, str(basic_stats) + "\n\n", 'insight')
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) >= 2:
                self.analyze_correlations(df[numeric_cols].corr())
        except Exception as e:
            error_msg = f"Error performing analysis: {str(e)}"
            self.handle_errors(e, error_msg)
            self.analysis_text.insert(tk.END, f"\n{error_msg}\n", 'error')
            self.full_analysis_text.insert(tk.END, f"\n{error_msg}\n", 'error')
    def analyze_correlations(self, corr):
        """Analyze correlations using Gemini"""
        try:
            strong_correlations = []
            for col1 in corr.columns:
                for col2 in corr.columns:
                    if col1 < col2:
                        correlation = corr.loc[col1, col2]
                        if abs(correlation) > 0.5:
                            strong_correlations.append(
                                f"{col1} - {col2}: {correlation:.2f}"
                            )
            if not strong_correlations:
                correlation_text = "No strong correlations found between variables (threshold: 0.5)"
                self.analysis_text.insert(tk.END, "\n=== CORRELATION ANALYSIS ===\n\n", 'header')
                self.analysis_text.insert(tk.END, correlation_text + "\n\n", 'insight')
                self.full_analysis_text.insert(tk.END, "\n=== CORRELATION ANALYSIS ===\n\n", 'header')
                self.full_analysis_text.insert(tk.END, correlation_text + "\n\n", 'insight')
                return
            corr_prompt = self.load_prompt('data_correlation_analysis')
            if not corr_prompt:
                corr_prompt = """Analyze these correlations and explain:
    1. Strong positive correlations
    2. Strong negative correlations
    3. Key feature relationships
    4. Feature selection recommendations
    5. Structured response with bullet points
    6. Should be structure and easy to understand
    Correlations: {correlations}"""
            formatted_prompt = corr_prompt.format(correlations=strong_correlations)
            response = self.data_analysis_model.generate_content(formatted_prompt)
            if not response or not response.text:
                raise Exception("No response from Gemini API")
            correlation_text = response.text.replace("*", "").strip()
            self.analysis_text.insert(tk.END, "\n=== CORRELATION ANALYSIS ===\n\n", 'header')
            self.analysis_text.insert(tk.END, correlation_text + "\n\n", 'insight')
            self.full_analysis_text.insert(tk.END, "\n=== CORRELATION ANALYSIS ===\n\n", 'header')
            self.full_analysis_text.insert(tk.END, correlation_text + "\n\n", 'insight')
        except Exception as e:
            error_msg = f"Error analyzing correlations: {str(e)}"
            self.handle_errors(e, error_msg)
            self.analysis_text.insert(tk.END, f"\n{error_msg}\n", 'error')
            self.full_analysis_text.insert(tk.END, f"\n{error_msg}\n", 'error')
    def get_battery_info(self):
        """Get battery information"""
        try:
            battery = psutil.sensors_battery()
            if battery:
                percent = battery.percent
                power_plugged = battery.power_plugged
                status = "plugged in" if power_plugged else "not plugged in"
                if not power_plugged:
                    minutes = battery.secsleft // 60
                    hours = minutes // 60
                    mins = minutes % 60
                    time_left = f"{hours} hours and {mins} minutes remaining"
                else:
                    time_left = "charging"
                self.speak(f"Battery is at {percent} percent, {status}, {time_left}")
            else:
                self.speak("No battery detected. This might be a desktop computer.")
        except Exception as e:
            self.handle_errors(e, "battery info")
            self.speak("Sorry, I couldn't get the battery information")
    def analyze_database(self):
        """Open a window to analyze database files with TRON Legacy theme"""
        try:
            analysis_window = tk.Toplevel(self.root)
            analysis_window.title("Database Analysis - TRON Interface")
            analysis_window.geometry("1000x800")
            analysis_window.configure(bg='#0c1221')  # Dark blue background
            analysis_window.transient(self.root)
            analysis_window.grab_set()
            main_frame = self.create_tron_style_frame(analysis_window)
            main_frame.pack(fill='both', expand=True, padx=20, pady=20)
            header_frame = Frame(main_frame, bg='#0c1221')
            header_frame.pack(fill='x', pady=(0, 20))
            Label(
                header_frame,
                text="DATABASE ANALYSIS INTERFACE",
                font=("Orbitron", 24, "bold"),
                bg='#0c1221',
                fg='#00f6ff'
            ).pack(pady=(0, 5))
            Label(
                header_frame,
                text="A.N.I.S GRID SEARCH SYSTEM v1.0 ",
                font=("Orbitron", 12),
                bg='#0c1221',
                fg='#0088ff'
            ).pack()
            file_frame = self.create_tron_style_frame(main_frame)
            file_frame.pack(fill='x', pady=(0, 20))
            self.db_path_var = tk.StringVar()
            entry_frame = Frame(file_frame, bg='#0c1221')
            entry_frame.pack(fill='x', padx=10)
            db_entry = Entry(
                entry_frame,
                textvariable=self.db_path_var,
                font=("Share Tech Mono", 10),
                bg='#0f1a2b',
                fg='#00f6ff',
                insertbackground='#00f6ff',
                relief='flat',
                width=70
            )
            db_entry.pack(side='left', padx=(0, 10))
            browse_btn = self.create_tron_button(entry_frame, "BROWSE", self.browse_database)
            browse_btn.pack(side='left')
            search_frame = self.create_tron_style_frame(main_frame)
            search_frame.pack(fill='x', pady=(0, 20))
            Label(
                search_frame,
                text="ENTER QUERY PARAMETERS:",
                font=("Orbitron", 12),
                bg='#0c1221',
                fg='#00f6ff'
            ).pack(anchor='w', pady=(0, 5))
            self.search_entry = Text(
                search_frame,
                height=3,
                font=("Share Tech Mono", 10),
                bg='#0f1a2b',
                fg='#00f6ff',
                insertbackground='#00f6ff',
                wrap=WORD,
                relief='flat'
            )
            self.search_entry.pack(fill='x', pady=(0, 10))
            button_frame = Frame(search_frame, bg='#0c1221')
            button_frame.pack(pady=10)
            analyze_btn = self.create_tron_button(button_frame, "ANALYZE", self.perform_database_analysis, width=20)
            analyze_btn.pack(side='left', padx=5)
            execute_btn = self.create_tron_button(button_frame, "EXECUTE", lambda: self.perform_database_analysis(modify=True), width=20)
            execute_btn.pack(side='left', padx=5)
            results_frame = self.create_tron_style_frame(main_frame)
            results_frame.pack(fill='both', expand=True)
            Label(
                results_frame,
                text="ANALYSIS OUTPUT:",
                font=("Orbitron", 12),
                bg='#0c1221',
                fg='#00f6ff'
            ).pack(anchor='w', pady=(0, 5))
            results_container = Frame(
                results_frame,
                bg='#0c1221',
                highlightbackground='#00f6ff',
                highlightthickness=1
            )
            results_container.pack(fill='both', expand=True)
            self.results_text = Text(
                results_container,
                wrap=WORD,
                font=("Share Tech Mono", 10),
                bg='#0f1a2b',
                fg='#00f6ff',
                insertbackground='#00f6ff',
                relief='flat',
                padx=10,
                pady=10
            )
            self.results_text.pack(side='left', fill='both', expand=True)
            scrollbar = ttk.Scrollbar(
                results_container,
                orient='vertical',
                command=self.results_text.yview
            )
            scrollbar.pack(side='right', fill='y')
            style = ttk.Style()
            style.configure(
                "Tron.Vertical.TScrollbar",
                troughcolor='#0c1221',
                background='#00f6ff',
                arrowcolor='#00f6ff',
                bordercolor='#00f6ff',
                lightcolor='#00f6ff',
                darkcolor='#00f6ff'
            )
            scrollbar.configure(style="Tron.Vertical.TScrollbar")
            self.results_text['yscrollcommand'] = scrollbar.set
            self.results_text.tag_configure(
                'header',
                foreground='#00ffff',
                font=("Share Tech Mono", 11, 'bold')
            )
            self.results_text.tag_configure(
                'sql',
                foreground='#0088ff',
                font=("Share Tech Mono", 10, 'italic')
            )
            self.results_text.tag_configure(
                'data',
                foreground='#00f6ff'
            )
            self.results_text.tag_configure(
                'error',
                foreground='#ff0055'
            )
        except Exception as e:
            logging.error(f"Error creating database analysis window: {str(e)}")
            messagebox.showerror("Error", "Failed to create analysis window")
    def perform_database_analysis(self, modify=False):
        """Analyze or modify the database based on user query with TRON-style output"""
        try:
            db_path = self.db_path_var.get()
            user_input = self.search_entry.get('1.0', 'end-1c').strip()
            if not db_path:
                messagebox.showwarning("Warning", "Please select a database file")
                return
            if not user_input:
                messagebox.showwarning("Warning", "Please enter your request")
                return
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            self.results_text.delete('1.0', END)
            action = "MODIFYING" if modify else "ANALYZING"
            self.results_text.insert(END, f"INITIATING DATABASE {action}...\n", 'header')
            self.results_text.insert(END, "=" * 50 + "\n\n", 'header')
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
            tables = cursor.fetchall()
            schema_info = []
            sample_data_info = []
            for table in tables:
                table_name = table[0]
                cursor.execute(f"PRAGMA table_info({table_name});")
                columns = cursor.fetchall()
                schema_info.append(f"Table: {table_name}")
                schema_info.append("Columns: " + ", ".join(col[1] for col in columns))
                try:
                    cursor.execute(f"SELECT * FROM {table_name} LIMIT 5")
                    sample_rows = cursor.fetchall()
                    if sample_rows:
                        sample_data_info.append(f"Sample data from {table_name}:")
                        for row in sample_rows:
                            sample_data_info.append(str(row))
                except:
                    pass
            if modify:
                try:
                    import google.generativeai as genai
                    api_key = self.load_setting('gemini_api_key')
                    model_name = self.load_setting("gemini_model")
                    genai.configure(api_key=api_key)
                    model = genai.GenerativeModel(model_name)
                    prompt = f"""Given this database schema and sample data:
                    {chr(10).join(schema_info)}
                    Sample Data:
                    {chr(10).join(sample_data_info)}
                    User wants to modify the database with this request: {user_input}
                    Generate a single SQL query to perform this modification. Important rules:
                    1. Only generate UPDATE, INSERT, DELETE, CREATE, DROP, or ALTER queries
                    2. The query must end with a semicolon
                    3. Use proper SQL syntax and escaping for text values
                    4. Consider data types and constraints
                    5. Return only the SQL query, no explanations
                    6. Make sure the query is safe and won't cause data corruption
                    7. Use appropriate WHERE clauses for UPDATE and DELETE
                    Example formats:
                    UPDATE table SET column = 'value' WHERE condition;
                    INSERT INTO table (column1, column2) VALUES ('value1', 'value2');
                    DELETE FROM table WHERE condition;
                    """
                    response = model.generate_content(prompt)
                    query = response.text.strip()
                    query = query.replace('```sql', '').replace('```', '').strip()
                    query_type = query.strip().upper().split()[0]
                    if query_type not in ('UPDATE', 'INSERT', 'DELETE', 'CREATE', 'DROP', 'ALTER'):
                        raise ValueError("Only UPDATE, INSERT, DELETE, CREATE, DROP, and ALTER queries are allowed in execute mode")
                    self.results_text.insert(END, "GENERATED SQL QUERY:\n", 'header')
                    self.results_text.insert(END, f"{query}\n\n", 'sql')
                    self.results_text.insert(END, "\nCreating backup before modification...\n", 'header')
                    backup_path = self._create_database_backup(db_path)
                    if not backup_path:
                        raise ValueError("Failed to create backup. Aborting modification for safety.")
                    self.results_text.insert(END, f"Backup created successfully at: {backup_path}\n\n", 'data')
                    cursor.execute(query)
                    conn.commit()
                    rows_affected = cursor.rowcount
                    self.results_text.insert(END, f"Query executed successfully\n", 'header')
                    self.results_text.insert(END, f"Rows affected: {rows_affected}\n", 'data')
                    if query_type in ('UPDATE', 'INSERT', 'DELETE'):
                        table_name = self._extract_table_name(query)
                        if table_name:
                            self.results_text.insert(END, f"\nUpdated table contents:\n", 'header')
                            cursor.execute(f"SELECT * FROM {table_name} LIMIT 100")
                            results = cursor.fetchall()
                            if results:
                                col_names = [description[0] for description in cursor.description]
                                self.results_text.insert(END, "COLUMNS: ", 'header')
                                self.results_text.insert(END, " | ".join(col_names) + "\n", 'data')
                                self.results_text.insert(END, "-" * 50 + "\n", 'data')
                                for row in results:
                                    formatted_row = []
                                    for item in row:
                                        if item is None:
                                            formatted_row.append("NULL")
                                        else:
                                            formatted_row.append(str(item))
                                    self.results_text.insert(END, " | ".join(formatted_row) + "\n", 'data')
                except sqlite3.Error as e:
                    error_msg = f"DATABASE ERROR: {str(e)}\n"
                    self.results_text.insert(END, error_msg, 'error')
                    conn.rollback()
                    self.results_text.insert(END, "Changes rolled back. Your database is unchanged.\n", 'header')
                    self.results_text.insert(END, f"You can restore from backup if needed: {backup_path}\n", 'data')
                except ValueError as e:
                    error_msg = f"VALIDATION ERROR: {str(e)}\n"
                    self.results_text.insert(END, error_msg, 'error')
                except Exception as e:
                    error_msg = f"AI ERROR: {str(e)}\n"
                    self.results_text.insert(END, error_msg, 'error')
                    self.results_text.insert(END, "Please try rephrasing your request.\n", 'data')
                    logging.error(f"AI error: {str(e)}")
            else:
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
                tables = cursor.fetchall()
                try:
                    import google.generativeai as genai
                    api_key = self.load_setting('gemini_api_key')
                    model_name = self.load_setting("gemini_model")
                    genai.configure(api_key=api_key)
                    model = genai.GenerativeModel(model_name)
                    schema_info = []
                    sample_data_info = []
                    for table in tables:
                        table_name = table[0]
                        cursor.execute(f"PRAGMA table_info({table_name});")
                        columns = cursor.fetchall()
                        schema_info.append(f"Table: {table_name}")
                        schema_info.append("Columns: " + ", ".join(col[1] for col in columns))
                        try:
                            cursor.execute(f"SELECT * FROM {table_name} LIMIT 5")
                            sample_rows = cursor.fetchall()
                            if sample_rows:
                                sample_data_info.append(f"Sample data from {table_name}:")
                                for row in sample_rows:
                                    sample_data_info.append(str(row))
                        except:
                            pass
                    prompt = f"""Given this database schema and sample data:
                    {chr(10).join(schema_info)}
                    Sample Data:
                    {chr(10).join(sample_data_info)}
                    User query: {user_input}
                    Important instructions for generating SQL queries:
                    1. Generate only pure SQL queries without any formatting
                    2. Each query must end with a semicolon
                    3. Use multiple search patterns for flexible matching:
                       - Use LIKE with wildcards for partial matches: '%word%'
                       - Use word boundary matches: '% word %', 'word %', '% word'
                       - Split search terms and search each individually
                       - Consider both exact and fuzzy matches
                    4. For text columns:
                       - Use COLLATE NOCASE for case-insensitive search
                       - Handle NULL values appropriately
                       - Use OR conditions for multiple possible matches
                    5. If searching numbers:
                       - Consider approximate matches (BETWEEN, >, <)
                       - Handle different number formats
                    6. Generate multiple queries with different matching strategies
                    7. Return only SQL queries, no explanations
                    Example format:
                    SELECT * FROM table WHERE column1 LIKE '%partial%' OR column2 LIKE '%partial%' COLLATE NOCASE;
                    SELECT * FROM table WHERE (column1 LIKE '% word%' OR column1 LIKE '%word %' OR column1 LIKE '%word%') COLLATE NOCASE;
                    """
                    response = model.generate_content(prompt)
                    sql_text = response.text.strip()
                    sql_text = sql_text.replace('```sql', '').replace('```', '').strip()
                    sql_queries = [q.strip() for q in sql_text.split(';') if q.strip()]
                    executed_queries = set()
                    all_results = set()
                    for query in sql_queries:
                        if not query or query in executed_queries:
                            continue
                        executed_queries.add(query)
                        try:
                            if 'LIKE' in query.upper() and 'COLLATE NOCASE' not in query.upper():
                                query = query.replace('LIKE', 'COLLATE NOCASE LIKE')
                            self.results_text.insert(END, "\nEXECUTING QUERY:\n", 'header')
                            self.results_text.insert(END, f"{query}\n\n", 'sql')
                            cursor.execute(query)
                            results = cursor.fetchall()
                            if results:
                                col_names = [description[0] for description in cursor.description]
                                self.results_text.insert(END, "COLUMNS: ", 'header')
                                self.results_text.insert(END, " | ".join(col_names) + "\n", 'data')
                                self.results_text.insert(END, "-" * 50 + "\n", 'data')
                                for row in results:
                                    result_str = " | ".join(str(item) if item is not None else "NULL" for item in row)
                                    if result_str not in all_results:
                                        all_results.add(result_str)
                                        self.results_text.insert(END, result_str + "\n", 'data')
                                self.results_text.insert(END, "\n" + "="*50 + "\n", 'header')
                            else:
                                words = query.split()
                                if len(words) > 1:
                                    for word in words:
                                        alt_query = query.replace(f"%{query}%", f"%{word}%")
                                        if alt_query != query:
                                            cursor.execute(alt_query)
                                            alt_results = cursor.fetchall()
                                            if alt_results:
                                                self.results_text.insert(END, f"\nResults for partial match '{word}':\n", 'header')
                                                for row in alt_results:
                                                    result_str = " | ".join(str(item) if item is not None else "NULL" for item in row)
                                                    if result_str not in all_results:
                                                        all_results.add(result_str)
                                                        self.results_text.insert(END, result_str + "\n", 'data')
                                if not all_results:
                                    self.results_text.insert(END, "NO RESULTS FOUND\n", 'error')
                        except sqlite3.Error as e:
                            error_msg = f"ERROR IN QUERY: {str(e)}\n"
                            self.results_text.insert(END, error_msg, 'error')
                            logging.error(f"SQL error in query '{query}': {str(e)}")
                    if not all_results:
                        self.results_text.insert(END, "\nNO MATCHING RESULTS FOUND\n", 'error')
                        self.results_text.insert(END, "\nSUGGESTIONS:\n", 'header')
                        self.results_text.insert(END, "- Try using fewer words\n", 'data')
                        self.results_text.insert(END, "- Check for typos\n", 'data')
                        self.results_text.insert(END, "- Use partial words\n", 'data')
                        self.results_text.insert(END, "- Try alternative spellings\n", 'data')
                except Exception as e:
                    self.results_text.insert(END, f"ANALYSIS ERROR: {str(e)}\n", 'error')
                    self.results_text.insert(END, "Please try rephrasing your query.\n", 'data')
                    logging.error(f"Analysis error: {str(e)}")
        except Exception as e:
            logging.error(f"Database operation error: {str(e)}")
            messagebox.showerror("Error", f"Operation failed: {str(e)}")
        finally:
            if 'conn' in locals():
                conn.close()
    def _extract_table_name(self, query):
        """Extract table name from SQL query"""
        try:
            import re
            words = query.strip().upper().split()
            if 'UPDATE' in words:
                match = re.search(r'UPDATE\s+([^\s]+)', query, re.IGNORECASE)
            elif 'INSERT' in words:
                match = re.search(r'INSERT\s+INTO\s+([^\s(]+)', query, re.IGNORECASE)
            elif 'DELETE' in words:
                match = re.search(r'DELETE\s+FROM\s+([^\s]+)', query, re.IGNORECASE)
            else:
                return None
            return match.group(1) if match else None
        except:
            return None
    def browse_database(self):
        """Open file dialog to select database file"""
        try:
            file_path = filedialog.askopenfilename(
                title="Select Database File",
                filetypes=[("Database files", "*.db"), ("All files", "*.*")]
            )
            if file_path:
                self.db_path_var.set(file_path)
        except Exception as e:
            logging.error(f"Error browsing database: {str(e)}")
            messagebox.showerror("Error", "Failed to open file dialog")
    def _create_database_backup(self, db_path):
        """Create a backup of the database after modification"""
        try:
            backup_dir = os.path.join(os.path.dirname(db_path), 'backups')
            if not os.path.exists(backup_dir):
                os.makedirs(backup_dir)
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            db_name = os.path.splitext(os.path.basename(db_path))[0]
            backup_path = os.path.join(backup_dir, f"{db_name}_backup_{timestamp}.db")
            source = sqlite3.connect(db_path)
            dest = sqlite3.connect(backup_path)
            with source, dest:
                source.backup(dest)
            logging.info(f"Database backup created successfully at {backup_path}")
            return backup_path
        except Exception as e:
            logging.error(f"Failed to create database backup: {str(e)}")
            return None
    def _create_appearance_settings(self, parent):
        """Create appearance settings tab"""
        appearance_frame = ctk.CTkFrame(parent)
        appearance_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        title_label = ctk.CTkLabel(
            appearance_frame,
            text="Appearance Settings",
            font=ctk.CTkFont(size=18, weight="bold")
        )
        title_label.pack(anchor=tk.W, pady=(0, 15))
        theme_frame = ctk.CTkFrame(appearance_frame)
        theme_frame.pack(fill=tk.X, pady=10)
        theme_label = ctk.CTkLabel(
            theme_frame,
            text="Theme",
            font=ctk.CTkFont(size=16, weight="bold")
        )
        theme_label.pack(anchor=tk.W, padx=15, pady=(10, 5))
        theme_desc = ctk.CTkLabel(
            theme_frame,
            text="Select a color theme for the application interface.",
            wraplength=700
        )
        theme_desc.pack(anchor=tk.W, padx=15, pady=(0, 5))
        self.theme_var = tk.StringVar(value=self.current_theme)
        themes_frame = ctk.CTkFrame(theme_frame, fg_color="transparent")
        themes_frame.pack(fill=tk.X, padx=15, pady=5)
        for idx, theme_name in enumerate(self.themes.keys()):
            theme_radio = ctk.CTkRadioButton(
                themes_frame, 
                text=theme_name.capitalize(), 
                variable=self.theme_var,
                value=theme_name,
                command=lambda t=theme_name: self.apply_theme(t)
            )
            theme_radio.pack(side=tk.LEFT, padx=(0 if idx == 0 else 15, 0))
    def _create_email_settings(self, parent):
        """Create email settings tab"""
        email_frame = ctk.CTkFrame(parent)
        email_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        title_label = ctk.CTkLabel(
            email_frame,
            text="Email Settings",
            font=ctk.CTkFont(size=18, weight="bold")
        )
        title_label.pack(anchor=tk.W, pady=(0, 15))
        config_frame = ctk.CTkFrame(email_frame)
        config_frame.pack(fill=tk.X, pady=10)
        config_label = ctk.CTkLabel(
            config_frame,
            text="Email Configuration",
            font=ctk.CTkFont(size=16, weight="bold")
        )
        config_label.pack(anchor=tk.W, padx=15, pady=(10, 5))
        config_desc = ctk.CTkLabel(
            config_frame,
            text="Configure your email settings for sending emails through voice commands.",
            wraplength=700
        )
        config_desc.pack(anchor=tk.W, padx=15, pady=(0, 5))
        email_label = ctk.CTkLabel(config_frame, text="Email Address:")
        email_label.pack(anchor=tk.W, padx=15, pady=(10, 0))
        email_entry = ctk.CTkEntry(config_frame, width=300)
        email_entry.pack(anchor=tk.W, padx=15, pady=(0, 10))
        self.email_entry = email_entry
        smtp_label = ctk.CTkLabel(config_frame, text="SMTP Server:")
        smtp_label.pack(anchor=tk.W, padx=15, pady=(10, 0))
        smtp_entry = ctk.CTkEntry(config_frame, width=300)
        smtp_entry.pack(anchor=tk.W, padx=15, pady=(0, 10))
        self.smtp_entry = smtp_entry
        port_label = ctk.CTkLabel(config_frame, text="SMTP Port:")
        port_label.pack(anchor=tk.W, padx=15, pady=(10, 0))
        port_entry = ctk.CTkEntry(config_frame, width=100)
        port_entry.pack(anchor=tk.W, padx=15, pady=(0, 10))
        self.port_entry = port_entry
        password_label = ctk.CTkLabel(config_frame, text="Password:")
        password_label.pack(anchor=tk.W, padx=15, pady=(10, 0))
        password_entry = ctk.CTkEntry(config_frame, width=300, show="•")
        password_entry.pack(anchor=tk.W, padx=15, pady=(0, 10))
        self.password_entry = password_entry
        self.use_tls_var = tk.BooleanVar(value=True)
        tls_checkbox = ctk.CTkCheckBox(
            config_frame,
            text="Use TLS",
            variable=self.use_tls_var
        )
        tls_checkbox.pack(anchor=tk.W, padx=15, pady=10)
        test_frame = ctk.CTkFrame(email_frame)
        test_frame.pack(fill=tk.X, pady=10)
        test_label = ctk.CTkLabel(
            test_frame,
            text="Test Email",
            font=ctk.CTkFont(size=16, weight="bold")
        )
        test_label.pack(anchor=tk.W, padx=15, pady=(10, 5))
        test_button = ctk.CTkButton(
            test_frame,
            text="Send Test Email",
            command=self.send_test_email,
            width=200,
            fg_color=self.themes[self.current_theme]["accent"]
        )
        test_button.pack(anchor=tk.W, padx=15, pady=10)
    def _create_prompts_settings(self, parent):
            """Create prompts settings tab"""
            prompts_frame = ctk.CTkFrame(parent)
            prompts_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
            title_label = ctk.CTkLabel(
                prompts_frame,
                text="AI Prompts Settings",
                font=ctk.CTkFont(size=18, weight="bold")
            )
            title_label.pack(anchor=tk.W, pady=(0, 15))
            prompt_types = {
                "system": "System Prompt",
                "commands": "Commands Prompt",
                "screen": "Screen Analysis Prompt"
            }
            for prompt_type, prompt_title in prompt_types.items():
                prompt_frame = ctk.CTkFrame(prompts_frame)
                prompt_frame.pack(fill=tk.X, pady=10)
                prompt_label = ctk.CTkLabel(
                    prompt_frame,
                    text=prompt_title,
                    font=ctk.CTkFont(size=16, weight="bold")
                )
                prompt_label.pack(anchor=tk.W, padx=15, pady=(10, 5))
                prompt_desc = ctk.CTkLabel(
                    prompt_frame,
                    text=f"Edit the {prompt_title.lower()} used by the AI.",
                    wraplength=700
                )
                prompt_desc.pack(anchor=tk.W, padx=15, pady=(0, 5))
                prompt_text = ctk.CTkTextbox(prompt_frame, height=100)
                prompt_text.pack(fill=tk.X, padx=15, pady=5)
                prompt_text.insert("1.0", self.load_prompt(prompt_type))
                setattr(self, f"{prompt_type}_prompt_text", prompt_text)
                button_frame = ctk.CTkFrame(prompt_frame, fg_color="transparent")
                button_frame.pack(fill=tk.X, padx=15, pady=5)
                save_button = ctk.CTkButton(
                    button_frame,
                    text=f"Save {prompt_title}",
                    command=lambda pt=prompt_type, tt=prompt_text: self.save_prompt(pt, tt.get("1.0", "end-1c")),
                    fg_color=self.themes[self.current_theme]["accent"]
                )
                save_button.pack(side=tk.LEFT, padx=(0, 10))
                reset_button = ctk.CTkButton(
                    button_frame,
                    text=f"Reset to Default",
                    command=lambda pt=prompt_type, tt=prompt_text: self.reset_single_prompt(pt, tt),
                    fg_color=self.themes[self.current_theme]["secondary"]
                )
                reset_button.pack(side=tk.LEFT)
    def _create_screen_settings(self, parent):
        """Create screen interaction settings tab"""
        screen_frame = ctk.CTkFrame(parent)
        screen_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        title_label = ctk.CTkLabel(
            screen_frame,
            text="Screen Interaction Settings",
            font=ctk.CTkFont(size=18, weight="bold")
        )
        title_label.pack(anchor=tk.W, pady=(0, 15))
        resolution_frame = ctk.CTkFrame(screen_frame)
        resolution_frame.pack(fill=tk.X, pady=10)
        resolution_label = ctk.CTkLabel(
            resolution_frame,
            text="Screen Resolution Adaptation",
            font=ctk.CTkFont(size=16, weight="bold")
        )
        resolution_label.pack(anchor=tk.W, padx=15, pady=(10, 5))
        if not hasattr(self, '_screen_adapter'):
            self._screen_adapter = ScreenCoordinateAdapter()
        resolution_desc = ctk.CTkLabel(
            resolution_frame,
            text="Coordinates are automatically adapted to your current screen resolution. Update the reference resolution to set your current screen as the new standard.",
            wraplength=700
        )
        resolution_desc.pack(anchor=tk.W, padx=15, pady=(0, 5))
        current_resolution = f"Current Resolution: {self._screen_adapter.current_width}x{self._screen_adapter.current_height}"
        reference_resolution = f"Reference Resolution: {self._screen_adapter.reference_width}x{self._screen_adapter.reference_height}"
        current_res_label = ctk.CTkLabel(
            resolution_frame,
            text=current_resolution
        )
        current_res_label.pack(anchor=tk.W, padx=15, pady=(0, 5))
        ref_res_label = ctk.CTkLabel(
            resolution_frame,
            text=reference_resolution
        )
        ref_res_label.pack(anchor=tk.W, padx=15, pady=(0, 5))
        def update_reference_resolution():
            if self._screen_adapter.update_reference_resolution():
                ref_res_label.configure(text=f"Reference Resolution: {self._screen_adapter.current_width}x{self._screen_adapter.current_height}")
                self.speak("Screen resolution reference updated")
                messagebox.showinfo("Resolution Updated", "Reference resolution has been updated to match your current screen size.")
            else:
                messagebox.showerror("Error", "Failed to update reference resolution.")
        resolution_btn = ctk.CTkButton(
            resolution_frame,
            text="Update Reference Resolution",
            command=update_reference_resolution,
            width=200,
            fg_color=self.themes[self.current_theme]["accent"]
        )
        resolution_btn.pack(anchor=tk.W, padx=15, pady=5)
        training_frame = ctk.CTkFrame(screen_frame)
        training_frame.pack(fill=tk.X, pady=10)
        training_label = ctk.CTkLabel(
            training_frame,
            text="Training Mode",
            font=ctk.CTkFont(size=16, weight="bold")
        )
        training_label.pack(anchor=tk.W, padx=15, pady=(10, 5))
        training_desc = ctk.CTkLabel(
            training_frame,
            text="Training mode lets you identify UI elements by saying 'It's a [element name]' or 'This is a [element name]'",
            wraplength=700
        )
        training_desc.pack(anchor=tk.W, padx=15, pady=(0, 5))
        training_buttons = ctk.CTkFrame(training_frame, fg_color="transparent")
        training_buttons.pack(fill=tk.X, padx=15, pady=5)
        activate_btn = ctk.CTkButton(
            training_buttons,
            text="Activate Training Mode",
            command=self.activate_training_mode,
            width=200,
            fg_color=self.themes[self.current_theme]["accent"]
        )
        activate_btn.pack(side=tk.LEFT, padx=(0, 10))
        deactivate_btn = ctk.CTkButton(
            training_buttons,
            text="Deactivate Training Mode",
            command=self.deactivate_training_mode,
            width=200,
            fg_color=self.themes[self.current_theme]["accent"]
        )
        deactivate_btn.pack(side=tk.LEFT)
        coords_frame = ctk.CTkFrame(screen_frame)
        coords_frame.pack(fill=tk.X, pady=10)
        coords_label = ctk.CTkLabel(
            coords_frame,
            text="Saved Coordinates",
            font=ctk.CTkFont(size=16, weight="bold")
        )
        coords_label.pack(anchor=tk.W, padx=15, pady=(10, 5))
        coords_desc = ctk.CTkLabel(
            coords_frame,
            text="Manage UI element coordinates that have been saved during screen interaction.",
            wraplength=700
        )
        coords_desc.pack(anchor=tk.W, padx=15, pady=(0, 5))
        try:
            conn = sqlite3.connect(DATABASE_FILE)
            cursor = conn.cursor()
            cursor.execute("SELECT COUNT(*) FROM coordinates")
            coord_count = cursor.fetchone()[0]
            conn.close()
        except:
            coord_count = 0
        coords_status = ctk.CTkLabel(
            coords_frame,
            text=f"Currently saved coordinates: {coord_count}",
        )
        coords_status.pack(anchor=tk.W, padx=15, pady=(0, 5))
        coords_btn = ctk.CTkButton(
            coords_frame,
            text="Open Coordinates Manager",
            command=self.show_coordinates_manager,
            width=200,
            fg_color=self.themes[self.current_theme]["accent"]
        )
        coords_btn.pack(anchor=tk.W, padx=15, pady=5)
        history_frame = ctk.CTkFrame(screen_frame)
        history_frame.pack(fill=tk.X, pady=10)
        history_label = ctk.CTkLabel(
            history_frame,
            text="Command History",
            font=ctk.CTkFont(size=16, weight="bold")
        )
        history_label.pack(anchor=tk.W, padx=15, pady=(10, 5))
        history_desc = ctk.CTkLabel(
            history_frame,
            text="View the history of screen interaction commands that have been executed.",
            wraplength=700
        )
        history_desc.pack(anchor=tk.W, padx=15, pady=(0, 5))
        try:
            conn = sqlite3.connect(DATABASE_FILE)
            cursor = conn.cursor()
            cursor.execute("SELECT COUNT(*) FROM conversations WHERE command_type IS NOT NULL")
            cmd_count = cursor.fetchone()[0]
            conn.close()
        except:
            cmd_count = 0
        history_status = ctk.CTkLabel(
            history_frame,
            text=f"Recorded commands: {cmd_count}",
        )
        history_status.pack(anchor=tk.W, padx=15, pady=(0, 5))
        history_btn = ctk.CTkButton(
            history_frame,
            text="View Command History",
            command=self.show_command_history,
            width=200,
            fg_color=self.themes[self.current_theme]["accent"]
        )
        history_btn.pack(anchor=tk.W, padx=15, pady=5)
    def _create_voice_settings(self, parent):
        """Create voice settings tab"""
        voice_frame = ctk.CTkFrame(parent)
        voice_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        title_label = ctk.CTkLabel(
            voice_frame,
            text="Voice Settings",
            font=ctk.CTkFont(size=18, weight="bold")
        )
        title_label.pack(anchor=tk.W, pady=(0, 15))
        volume_frame = ctk.CTkFrame(voice_frame)
        volume_frame.pack(fill=tk.X, pady=10)
        volume_label = ctk.CTkLabel(
            volume_frame,
            text="Voice Volume",
            font=ctk.CTkFont(size=16, weight="bold")
        )
        volume_label.pack(anchor=tk.W, padx=15, pady=(10, 5))
        volume_slider = ctk.CTkSlider(
            volume_frame,
            from_=0,
            to=100,
            number_of_steps=10,
            command=self.change_voice_volume,
            width=300
        )
        volume_slider.pack(anchor=tk.W, padx=15, pady=10)
        volume_slider.set(50)
        self.voice_volume = 0.5
        recognition_frame = ctk.CTkFrame(voice_frame)
        recognition_frame.pack(fill=tk.X, pady=10)
        recognition_label = ctk.CTkLabel(
            recognition_frame,
            text="Voice Recognition",
            font=ctk.CTkFont(size=16, weight="bold")
        )
        recognition_label.pack(anchor=tk.W, padx=15, pady=(10, 5))
        energy_label = ctk.CTkLabel(recognition_frame, text="Energy Threshold:")
        energy_label.pack(anchor=tk.W, padx=15, pady=(10, 0))
        energy_slider = ctk.CTkSlider(
            recognition_frame,
            from_=300,
            to=5000,
            number_of_steps=47,
            width=300
        )
        energy_slider.pack(anchor=tk.W, padx=15, pady=(0, 10))
        energy_slider.set(self.recognizer.energy_threshold)
        self.energy_threshold_slider = energy_slider
        energy_value = ctk.CTkLabel(
            recognition_frame,
            text=f"Current: {int(self.recognizer.energy_threshold)}"
        )
        energy_value.pack(anchor=tk.W, padx=15, pady=(0, 10))
        self.energy_value_label = energy_value
        def update_energy_label(value):
            value = int(value)
            self.energy_value_label.configure(text=f"Current: {value}")
        energy_slider.configure(command=update_energy_label)
    def send_test_email(self):
        """Send a test email using the configured settings"""
        messagebox.showinfo("Test Email", "This feature is not implemented yet.")
class GeminiChatWindow:
    def __init__(self, parent):
        self.parent = parent
        self.window = ctk.CTkToplevel(parent.root)
        self.window.title("Gemini Chat")
        self.window.geometry("800x600")
        self.window.configure(fg_color=parent.themes[parent.current_theme]["bg"])
        self.main_app = parent
        self.theme = parent.themes[parent.current_theme]
        self.is_reading = False
        self.current_highlight_tag = None
        self.loading = False
        self.processing = False
        self.loaded_data = None
        self.loaded_filename = None
        try:
            import google.generativeai as genai
            api_key = self.load_setting('gemini_api_key')
            model_name = self.load_setting("gemini_model")
            genai.configure(api_key=api_key)
            self.model = genai.GenerativeModel(model_name)
            self.chat = self.model.start_chat(history=[])
        except Exception as e:
            messagebox.showerror("Error", f"Failed to initialize Gemini: {str(e)}")
            self.window.destroy()
            return
        self.create_widgets()
        self.engine = pyttsx3.init()
        self.window.after(100, self.center_window)
        self.window.transient(parent.root)
        self.window.grab_set()
        self.bind_shortcuts()
        self.window.protocol("WM_DELETE_WINDOW", self.on_closing)
        self.append_welcome_message()
    def load_setting(self, setting_name):
        """Load setting from database"""
        try:
            conn = sqlite3.connect("settings.db")
            cursor = conn.cursor()
            cursor.execute('''CREATE TABLE IF NOT EXISTS settings
                            (name TEXT PRIMARY KEY, value TEXT)''')
            cursor.execute('SELECT value FROM settings WHERE name = ?', (setting_name,))
            result = cursor.fetchone()
            conn.close()
            return result[0] if result else None
        except Exception as e:
            logging.error(f"Error loading setting: {str(e)}")
            return None
    def append_welcome_message(self):
        """Add welcome message to chat"""
        welcome_text = "Welcome to ANIS Chat! How can I assist you today?"
        self.chat_display._textbox.insert(tk.END, "Response: ", 'assistant')
        self.chat_display._textbox.insert(tk.END, f"{welcome_text}\n")
        self.chat_display.see(tk.END)
    def center_window(self):
        """Center window on screen and ensure proper dimensions"""
        self.window.update_idletasks()
        screen_width = self.window.winfo_screenwidth()
        screen_height = self.window.winfo_screenheight()
        width = min(screen_width - 100, 800)
        height = min(screen_height - 100, 600)
        x = (screen_width // 2) - (width // 2)
        y = (screen_height // 2) - (height // 2)
        self.window.geometry(f'{width}x{height}+{x}+{y}')
        self.window.minsize(width=600, height=400)
    def save_chat(self, event=None):
        try:
            file_path = filedialog.asksaveasfilename(
                defaultextension=".txt",
                filetypes=[("Text files", "*.txt"), ("All files", "*.*")]
            )
            if file_path:
                with open(file_path, 'w', encoding='utf-8') as file:
                    file.write(self.chat_display.get("0.0", "end"))
                dialog = ctk.CTkInputDialog(
                    text="Chat saved successfully!",
                    title="Success",
                    fg_color=self.theme["bg"],
                    button_fg_color=self.theme["accent"]
                )
                dialog.get_input()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save chat: {str(e)}")
    def clear_chat(self, event=None):
        confirm = messagebox.askyesno("Confirm", "Are you sure you want to clear the chat?")
        if confirm:
            self.chat_display.delete("0.0", "end")
            self.chat = self.model.start_chat(history=[])
    def bind_shortcuts(self):
        """Bind keyboard shortcuts to window"""
        self.window.bind('<Return>', self.handle_send)
        self.window.bind('<Control-s>', lambda e: self.save_chat())
        self.window.bind('<Control-l>', lambda e: self.clear_chat())
        self.window.bind('<Control-m>', lambda e: self.toggle_mode())
        self.window.bind('<F1>', lambda e: self.show_shortcuts())
        self.window.bind('<Control-o>', lambda e: self.load_document())
        self.input_text.bind('<Return>', self.handle_send)
    def handle_send(self, event=None):
        """Handle send message event to prevent duplicates"""
        if not self.processing:
            self.send_message()
        return "break"
    def create_widgets(self):
        main_container = ctk.CTkFrame(
            self.window,
            fg_color=self.theme["bg"],
            corner_radius=10,
            border_width=1,
            border_color=self.theme["accent"]
        )
        main_container.pack(fill="both", expand=True, padx=15, pady=15)
        mode_frame = ctk.CTkFrame(
            main_container,
            fg_color=self.theme["bg"],
            corner_radius=0
        )
        mode_frame.pack(fill='x', padx=10, pady=5)
        self.mode_var = ctk.StringVar(value="chat")
        self.mode_selector = ctk.CTkSegmentedButton(
            mode_frame,
            values=["Chat Mode", "Article Mode"],
            command=self.on_mode_change,
            selected_color=self.theme["accent"],
            unselected_color=self.theme["highlight"],
            fg_color=self.theme["text_bg"],
            selected_hover_color=self.theme["secondary"],
            text_color=self.theme["fg"],
            font=ctk.CTkFont(size=12, weight="bold")
        )
        self.mode_selector.pack(side="left", padx=10, pady=5)
        self.mode_selector.set("Chat Mode")
        mode_hint = ctk.CTkLabel(
            mode_frame,
            text="(Ctrl+M to toggle)", 
            font=ctk.CTkFont(size=10),
            text_color=self.theme["secondary"]
        )
        mode_hint.pack(side="left", padx=(0, 10))
        self.load_doc_button = ctk.CTkButton(
            mode_frame,
            text="Load Document",
            command=self.load_document,
            font=ctk.CTkFont(size=12),
            fg_color=self.theme["highlight"],
            hover_color=self.theme["secondary"],
            text_color=self.theme["fg"],
            corner_radius=8,
            width=120,
            height=30
        )
        self.load_doc_button.pack(side="right", padx=10, pady=5)
        self.status_label = ctk.CTkLabel(
            mode_frame,
            text="",
            font=ctk.CTkFont(size=10),
            text_color=self.theme["secondary"]
        )
        self.status_label.pack(side="right", padx=(0, 10))
        chat_frame = ctk.CTkFrame(
            main_container,
            fg_color=self.theme["text_bg"],
            corner_radius=10,
            border_width=1,
            border_color=self.theme["highlight"]
        )
        chat_frame.pack(fill='both', expand=True, padx=10, pady=10)
        self.chat_display = ctk.CTkTextbox(
            chat_frame,
            font=ctk.CTkFont(family="Consolas", size=12),
            fg_color=self.theme["text_bg"],
            text_color=self.theme["fg"],
            corner_radius=8,
            border_width=0,
            wrap="word",
            activate_scrollbars=True
        )
        self.chat_display.pack(fill='both', expand=True, padx=5, pady=5)
        bg_color = self.theme["text_bg"]
        user_bg_color = self.calculate_background_color(bg_color, self.theme["success"])
        assistant_bg_color = self.calculate_background_color(bg_color, self.theme["accent"])
        self.chat_display._textbox.tag_configure('user', foreground=self.theme["success"], font=("Consolas", 12, "bold"))
        self.chat_display._textbox.tag_configure('assistant', foreground=self.theme["accent"], font=("Consolas", 12))
        self.chat_display._textbox.tag_configure('system', foreground=self.theme["secondary"], font=("Consolas", 11, "italic"))
        self.chat_display._textbox.tag_configure('highlight', background=self.theme["accent"], foreground="#ffffff")
        self.chat_display._textbox.tag_configure('error', foreground=self.theme["error"], font=("Consolas", 12, "bold"))
        self.chat_display._textbox.tag_configure('user_bg', background=user_bg_color)
        self.chat_display._textbox.tag_configure('assistant_bg', background=assistant_bg_color)
        self.loading_label = ctk.CTkLabel(
            main_container,
            text="",
            font=ctk.CTkFont(size=12),
            text_color=self.theme["secondary"]
        )
        self.loading_label.pack(fill='x', padx=10)
        self.loading_label.pack_forget()
        input_frame = ctk.CTkFrame(
            main_container,
            fg_color=self.theme["highlight"],
            corner_radius=10,
            border_width=1,
            border_color=self.theme["accent"]
        )
        input_frame.pack(fill='x', padx=10, pady=5)
        input_label = ctk.CTkLabel(
            input_frame,
            text="Your message:",
            font=ctk.CTkFont(size=11, weight="bold"),
            text_color=self.theme["fg"]
        )
        input_label.pack(anchor="w", padx=10, pady=(5, 0))
        input_area = ctk.CTkFrame(
            input_frame,
            fg_color=self.theme["highlight"],
            corner_radius=0
        )
        input_area.pack(fill='x', padx=10, pady=5)
        self.input_text = ctk.CTkTextbox(
            input_area,
            height=60,
            font=ctk.CTkFont(family="Consolas", size=12),
            fg_color=self.theme["bg"],
            text_color=self.theme["fg"],
            corner_radius=8,
            border_width=1,
            border_color=self.theme["accent"],
            wrap="word"
        )
        self.input_text.pack(side='left', fill='x', expand=True, padx=(0, 10))
        self.send_button = ctk.CTkButton(
            input_area,
            text="Send",
            command=self.send_message,
            font=ctk.CTkFont(size=12, weight="bold"),
            fg_color=self.theme["accent"],
            hover_color=self.theme["secondary"],
            text_color="#ffffff",
            corner_radius=8,
            width=100,
            height=40
        )
        self.send_button.pack(side='right', padx=5)
        button_frame = ctk.CTkFrame(
            main_container,
            fg_color=self.theme["bg"],
            corner_radius=0
        )
        button_frame.pack(fill='x', padx=10, pady=5)
        buttons_data = [
            ("Save", self.save_chat, "Ctrl+S"),
            ("Clear", self.clear_chat, "Ctrl+L"),
            ("Help", self.show_shortcuts, "F1")
        ]
        for text, command, shortcut in buttons_data:
            btn = ctk.CTkButton(
                button_frame,
                text=text,
                command=command,
                font=ctk.CTkFont(size=12),
                fg_color=self.theme["highlight"],
                hover_color=self.theme["secondary"],
                text_color=self.theme["fg"],
                corner_radius=8,
                width=100,
                height=30
            )
            btn.pack(side='left', padx=5)
            tooltip = ctk.CTkLabel(
                button_frame,
                text=f"({shortcut})",
                font=ctk.CTkFont(size=9),
                text_color=self.theme["secondary"]
            )
            tooltip.pack(side='left', padx=(0, 10))
    def lighten_color(self, hex_color):
        """Create a lighter version of the given hex color for backgrounds"""
        hex_color = hex_color.lstrip('#')
        r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
        brightness = (r * 299 + g * 587 + b * 114) / 1000
        if brightness < 128:
            r = min(255, r + 20)
            g = min(255, g + 20)
            b = min(255, b + 20)
        else:
            r = min(255, int(r * 0.2 + 200))
            g = min(255, int(g * 0.2 + 200))
            b = min(255, int(b * 0.2 + 200))
        return f'#{r:02x}{g:02x}{b:02x}'
    def calculate_background_color(self, bg_color, accent_color):
        """Calculate a subtle background color by mixing bg_color with a touch of accent_color"""
        bg_color = bg_color.lstrip('#')
        accent_color = accent_color.lstrip('#')
        bg_r = int(bg_color[0:2], 16)
        bg_g = int(bg_color[2:4], 16)
        bg_b = int(bg_color[4:6], 16)
        accent_r = int(accent_color[0:2], 16)
        accent_g = int(accent_color[2:4], 16)
        accent_b = int(accent_color[4:6], 16)
        r = int(bg_r * 0.95 + accent_r * 0.05)
        g = int(bg_g * 0.95 + accent_g * 0.05)
        b = int(bg_b * 0.95 + accent_b * 0.05)
        brightness = (r * 299 + g * 587 + b * 114) / 1000
        if brightness < 128:
            r = min(255, int(r * 1.1))
            g = min(255, int(g * 1.1))
            b = min(255, int(b * 1.1))
        else:
            r = max(0, int(r * 0.95))
            g = max(0, int(g * 0.95))
            b = max(0, int(b * 0.95))
        return f'#{r:02x}{g:02x}{b:02x}'
    def on_mode_change(self, value):
        """Handle mode change from segmented button"""
        new_mode = "chat" if value == "Chat Mode" else "article"
        self.mode_var.set(new_mode)
        self.chat_display.insert(tk.END, f"\nSwitched to {new_mode} mode\n", 'system')
        self.chat_display.see(tk.END)
    def toggle_mode(self, event=None):
        """Toggle between chat and article mode (via keyboard shortcut)"""
        current_mode = self.mode_var.get()
        new_mode = "article" if current_mode == "chat" else "chat"
        self.mode_var.set(new_mode)
        button_value = "Chat Mode" if new_mode == "chat" else "Article Mode"
        self.mode_selector.set(button_value)
        self.chat_display._textbox.insert(tk.END, f"\nSwitched to {new_mode} mode\n", 'system')
        self.chat_display.see(tk.END)
    def show_shortcuts(self, event=None):
        """Show keyboard shortcuts help"""
        shortcuts = """
Keyboard Shortcuts:
─────────────────
Enter        : Send message
Ctrl + S     : Save chat
Ctrl + L     : Clear chat
Ctrl + M     : Toggle mode
Ctrl + O     : Load document
F1           : Show this help
        """
        dialog = ctk.CTkInputDialog(
            title="Keyboard Shortcuts",
            text=shortcuts,
            fg_color=self.theme["bg"],
            button_fg_color=self.theme["accent"],
            button_hover_color=self.theme["secondary"],
            entry_fg_color=self.theme["text_bg"],
            entry_border_color=self.theme["highlight"],
            entry_text_color=self.theme["fg"],
            button_text_color="#ffffff"
        )
        dialog.get_input()
    def show_loading(self, show=True):
        if show:
            self.loading = True
            self.loading_label.configure(text="Generating response...")
            self.loading_label.pack(fill='x', padx=10, pady=(0, 10))
            self.animate_loading()
        else:
            self.loading = False
            self.loading_label.pack_forget()
    def animate_loading(self):
        if not self.loading:
            return
        dots = self.loading_label.cget("text").count(".")
        new_dots = "." * ((dots + 1) % 4)
        self.loading_label.configure(text=f"Generating response{new_dots}")
        self.window.after(500, self.animate_loading)
    def send_message(self):
        if self.processing:
            return
        message = self.input_text.get("0.0", "end").strip()
        if not message:
            return
        self.processing = True
        self.input_text.configure(state='disabled')
        mode = self.mode_var.get()
        self.chat_display._textbox.insert(tk.END, "\n", 'user_bg')
        user_start = self.chat_display._textbox.index(tk.END)
        self.chat_display._textbox.insert(tk.END, "You: ", 'user')
        self.chat_display._textbox.insert(tk.END, f"{message}", 'user_bg')
        user_end = self.chat_display._textbox.index(tk.END)
        self.chat_display._textbox.insert(tk.END, "\n")
        self.chat_display._textbox.tag_add('user_bg', f"{user_start} linestart", f"{user_end} lineend")
        self.chat_display.see(tk.END)
        self.show_loading()
        threading.Thread(target=self._process_message, args=(message, mode), daemon=True).start()
    def _process_message(self, message, mode):
        try:
            if mode == "chat":
                custom_prompt = self.main_app.load_prompt('chat_window')
                document_context = ""
                if self.loaded_data and self.loaded_filename:
                    document_context = f"\n\nUse the following document as reference when answering (from {self.loaded_filename}):\n{self.loaded_data}\n\n"
                if custom_prompt:
                    prompt = f"""{custom_prompt}You should answer like intelligent friend.Don't mention your name gemini when someone ask you name instead of this tell 
 your name is ANIS AI integrated with Operation System With Help Of Gemini and under Developing by 
         Samiulla from 26 August 2024 -- don't  tell this in beginning of your response you should tell only when someone ask you name.Please provide clear, 
                well-structured responses and chat like friend and funny way ,without using asterisks or code blocks unless specifically 
                needed for code examples.Full Form Of Your Name is Artificial Neural Intelligence System .Should be short like 2 lines and sweet.Format your response in a professional, educational manner.Don't read the symbols 
                like dot,comma,exclamation mark,question mark,etc in your response{document_context}\n\nUser Query: {message}"""
                else:
                    prompt = f"""you should answer like intelligent friend.Don't mention your name gemini when someone ask you name instead of this tell 
 your name is ANIS AI integrated with Operation System With Help Of Gemini and under Developing by 
         Samiulla from 26 August 2024 -- don't  tell this in beginning of your response you should tell only when someone ask you name.Please provide clear, 
                well-structured responses and chat like friend and funny way ,without using asterisks or code blocks unless specifically 
                needed for code examples.Full Form Of Your Name is Artificial Neural Intelligence System .Should be short like 2 lines and sweet.Format your response in a professional, educational manner.Don't read the symbols 
                like dot,comma,exclamation mark,question mark,etc in your response{document_context}\n\nUser Query: {message}"""
            else:
                document_context = ""
                if self.loaded_data and self.loaded_filename:
                    document_context = f"\n\nUse the following document as reference when creating your article (from {self.loaded_filename}):\n{self.loaded_data}\n\n"
                custom_prompt = self.main_app.load_prompt('article_mode')
                if custom_prompt:
                    prompt = f"""{custom_prompt} ,You are a professional teacher/expert.Don't mention your name gemini when someone ask you name instead of this tell 
                  your name is ANIS AI integrated with Operation System With Help Of Gemini and under Developing by 
                  Samiulla from 26 August 2024 -- don't  tell this in beginning of your response you should tell only when someone ask you name.Please provide clear, 
                well-structured article responses without using asterisks or code blocks unless specifically 
                needed for code examples.Full Form Of Your Name is Artificial Neural Intelligence System .Should be short and sweet.Format your response in a professional, 
                educational manner.Don't read the symbols like dot,comma,exclamation mark,question mark,etc in your response{document_context}\n\n{message}"""
                else:
                    prompt = f"""You are a professional teacher/expert.Don't mention your name gemini when someone ask you name instead of this tell 
 your name is ANIS AI integrated with Operation System With Help Of Gemini and under Developing by 
         Samiulla from 26 August 2024 -- don't  tell this in beginning of your response you should tell only when someone ask you name.Please provide clear, 
                well-structured responses and chat like friend and funny way ,without using asterisks or code blocks unless specifically 
                needed for code examples.Full Form Of Your Name is Artificial Neural Intelligence System .Should be short like 2 lines and sweet.Format your response in a professional, educational manner.Don't read the symbols 
                like dot,comma,exclamation mark,question mark,etc in your response{document_context}\n\n User Query:{message}"""
            response = self.model.generate_content(prompt)
            response_text = response.text
            response_text = response_text.replace('*', '')
            response_text = response_text.replace('```', '')
            response_text = response_text.strip()
            self.window.after(0, self._update_chat_display, response_text)
        except Exception as e:
            self.window.after(0, self._show_error, str(e))
        finally:
            self.window.after(0, self._reset_processing)
    def _reset_processing(self):
        """Reset processing flag"""
        self.processing = False
    def _update_chat_display(self, response_text):
        self.chat_display._textbox.insert(tk.END, "\n", 'assistant_bg')
        assistant_start = self.chat_display._textbox.index(tk.END)
        self.chat_display._textbox.insert(tk.END, "Response: ", 'assistant')
        self.chat_display._textbox.insert(tk.END, f"{response_text}", 'assistant_bg')
        assistant_end = self.chat_display._textbox.index(tk.END)
        self.chat_display._textbox.insert(tk.END, "\n")
        self.chat_display._textbox.tag_add('assistant_bg', f"{assistant_start} linestart", f"{assistant_end} lineend")
        self.chat_display.see(tk.END)
        self.show_loading(False)
        self.input_text.configure(state='normal')
        self.input_text.delete("0.0", "end")
    def append_welcome_message(self):
        """Add welcome message to chat"""
        welcome_text = "Welcome to ANIS Chat! How can I assist you today?"
        self.chat_display._textbox.insert(tk.END, "\n", 'assistant_bg')
        assistant_start = self.chat_display._textbox.index(tk.END)
        self.chat_display._textbox.insert(tk.END, "Response: ", 'assistant')
        self.chat_display._textbox.insert(tk.END, f"{welcome_text}", 'assistant_bg')
        assistant_end = self.chat_display._textbox.index(tk.END)
        self.chat_display._textbox.insert(tk.END, "\n")
        self.chat_display._textbox.tag_add('assistant_bg', f"{assistant_start} linestart", f"{assistant_end} lineend")
        self.chat_display.see(tk.END)
    def _show_error(self, error_message):
        self.chat_display._textbox.insert(tk.END, f"\nError: {error_message}\n", 'error')
        self.chat_display.see(tk.END)
        self.show_loading(False)
        self.input_text.configure(state='normal')
    def on_closing(self):
        """Handle window closing properly"""
        try:
            if hasattr(self, 'window') and self.window.winfo_exists():
                self.loaded_data = None
                self.loaded_filename = None
                self.window.grab_release()
                self.window.destroy()
        except Exception as e:
            logging.error(f"Error closing window: {e}")
    def load_document(self, event=None):
        """Load document from file (PDF, DOCX, TXT)"""
        try:
            file_path = filedialog.askopenfilename(
                title="Select Document to Load",
                filetypes=[
                    ("All Supported Files", "*.pdf;*.docx;*.txt"),
                    ("PDF Files", "*.pdf"),
                    ("Word Documents", "*.docx"),
                    ("Text Files", "*.txt"),
                    ("All Files", "*.*")
                ]
            )
            if not file_path:
                return
            self.status_label.configure(text="● Loading document...")
            threading.Thread(target=self._load_document_thread, args=(file_path,), daemon=True).start()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load document: {str(e)}")
            self.status_label.configure(text="")
    def _load_document_thread(self, file_path):
        """Background thread for document loading"""
        try:
            filename = os.path.basename(file_path)
            extension = os.path.splitext(file_path)[1].lower()
            if extension == '.pdf':
                self._extract_pdf_text(file_path)
            elif extension == '.docx':
                self._extract_docx_text(file_path)
            elif extension == '.txt':
                self._extract_txt_text(file_path)
            else:
                raise ValueError(f"Unsupported file format: {extension}")
            self.window.after(0, lambda: self._update_document_loaded(filename))
        except Exception as e:
            self.window.after(0, lambda: self._show_document_error(str(e)))
    def _extract_pdf_text(self, file_path):
        """Extract text from PDF file"""
        try:
            import PyPDF2
            text = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                for page_num in range(num_pages):
                    text.append(pdf_reader.pages[page_num].extract_text())
            self.loaded_data = "\n\n".join(text)
            self.loaded_filename = os.path.basename(file_path)
        except ImportError:
            messagebox.showerror("Missing Dependency", "PyPDF2 is required to read PDF files. Please install it with 'pip install PyPDF2'.")
            raise
    def _extract_docx_text(self, file_path):
        """Extract text from DOCX file"""
        try:
            import docx
            doc = docx.Document(file_path)
            self.loaded_data = "\n\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
            self.loaded_filename = os.path.basename(file_path)
        except ImportError:
            messagebox.showerror("Missing Dependency", "python-docx is required to read DOCX files. Please install it with 'pip install python-docx'.")
            raise
    def _extract_txt_text(self, file_path):
        """Extract text from TXT file"""
        with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
            self.loaded_data = file.read()
            self.loaded_filename = os.path.basename(file_path)
    def _update_document_loaded(self, filename):
        """Update UI after document is loaded"""
        short_name = filename
        if len(short_name) > 25:
            short_name = short_name[:22] + "..."
        self.status_label.configure(text=f"● {short_name} loaded")
        self.chat_display._textbox.insert(tk.END, f"\nLoaded document: {filename}\n", 'system')
        self.chat_display.see(tk.END)
        self._analyze_document()
    def _show_document_error(self, error_message):
        """Show error when document loading fails"""
        self.status_label.configure(text="● Error loading document")
        self.chat_display._textbox.insert(tk.END, f"\nError loading document: {error_message}\n", 'error')
        self.chat_display.see(tk.END)
    def _analyze_document(self):
        """Analyze the loaded document with Gemini"""
        try:
            if not self.loaded_data or not self.loaded_filename:
                return
            self.chat_display._textbox.insert(tk.END, "\nAnalyzing document contents...\n", 'system')
            self.chat_display.see(tk.END)
            max_length = 12000
            data_preview = self.loaded_data[:max_length]
            if len(self.loaded_data) > max_length:
                data_preview += f"\n\n[Note: Document was truncated. Total length: {len(self.loaded_data)} characters]"
            analysis_prompt = f"""
            Analyze this document and provide a brief summary of its key points and content.
            Focus on extracting main ideas, concepts, and important details.
            Keep your response concise but comprehensive.
            Document: {self.loaded_filename}
            Content:
            {data_preview}
            """
            threading.Thread(target=self._process_document_analysis, args=(analysis_prompt,), daemon=True).start()
        except Exception as e:
            self.chat_display._textbox.insert(tk.END, f"\nError analyzing document: {str(e)}\n", 'error')
            self.chat_display.see(tk.END)
    def _process_document_analysis(self, analysis_prompt):
        """Process document analysis in background thread"""
        try:
            response = self.model.generate_content(analysis_prompt)
            analysis_text = response.text
            analysis_text = analysis_text.replace('*', '')
            analysis_text = analysis_text.replace('```', '')
            analysis_text = analysis_text.strip()
            self.window.after(0, lambda: self._update_analysis_result(analysis_text))
        except Exception as e:
            self.window.after(0, lambda: self._show_error(f"Error analyzing document: {str(e)}"))
    def _update_analysis_result(self, analysis_text):
        """Update UI with document analysis results"""
        self.chat_display._textbox.insert(tk.END, "\n", 'assistant_bg')
        assistant_start = self.chat_display._textbox.index(tk.END)
        self.chat_display._textbox.insert(tk.END, "Document Analysis: ", 'assistant')
        self.chat_display._textbox.insert(tk.END, f"{analysis_text}", 'assistant_bg')
        assistant_end = self.chat_display._textbox.index(tk.END)
        self.chat_display._textbox.insert(tk.END, "\n")
        self.chat_display._textbox.tag_add('assistant_bg', f"{assistant_start} linestart", f"{assistant_end} lineend")
        self.chat_display.see(tk.END)
        self.chat_display._textbox.insert(tk.END, "\nYou can now ask questions about this document.\n", 'system')
        self.chat_display.see(tk.END)
class SpecializedChatWindow:
    """A dedicated window for specialized chatbots"""
    def __init__(self, parent, specialization):
        self.window = ctk.CTkToplevel(parent)
        self.window.title(f"AI Assistant - {specialization.title()}")
        self.window.geometry("1000x800")
        self.window.minsize(600,500)
        api_key = self.load_setting('gemini_api_key')
        model_name = self.load_setting("gemini_model")
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(model_name)
        self.parent = parent
        self.specialization = specialization
        self.api_key = api_key
        self.processing = False
        self.theme = self.get_theme()
        self.loaded_data = ""
        self.window.configure(fg_color=self.theme["bg"])
        self.create_widgets()
        self.bind_shortcuts()
        self.center_window()
        self.window.protocol("WM_DELETE_WINDOW", self.on_closing)
        self.window.transient(parent)
        self.window.grab_set()
        self.status_label.configure(text="● Initializing AI...")
        self.window.update()
        threading.Thread(target=self._initialize_ai, daemon=True).start()
    def _initialize_ai(self):
        """Initialize AI in background thread to prevent UI freeze"""
        try:
            import google.generativeai as genai
            api_key = self.load_setting('gemini_api_key')
            model_name=self.load_setting('gemini_model')
            genai.configure(api_key=api_key)
            self.gemini_model = genai.GenerativeModel(model_name)
            self.chat = self.gemini_model.start_chat(history=[])
            self.window.after(0, self.create_specialized_system)
        except Exception as e:
            logging.error(f"Error initializing specialized assistant: {str(e)}")
            self.window.after(0, lambda: self._show_error("Failed to initialize specialized assistant. Check your internet connection."))
            self.window.after(1000, self.window.destroy)
    def load_setting(self, setting_name):
        """Load setting from database"""
        try:
            conn = sqlite3.connect("settings.db")
            cursor = conn.cursor()
            cursor.execute('''CREATE TABLE IF NOT EXISTS settings
                            (name TEXT PRIMARY KEY, value TEXT)''')
            cursor.execute('SELECT value FROM settings WHERE name = ?', (setting_name,))
            result = cursor.fetchone()
            conn.close()
            return result[0] if result else None
        except Exception as e:
            logging.error(f"Error loading setting: {str(e)}")
            return None
    def get_theme(self):
        """Get theme colors for consistency with parent app"""
        return {
            "bg": "#1a1a2e",
            "fg": "#e6e6e6",
            "accent": "#2d6cdf",
            "secondary": "#3f72af",
            "highlight": "#0f3460",
            "input_bg": "#0f0f1e",
            "text_bg": "#101024",
            "button": "#1e3a8a",
            "button_hover": "#2d6cdf",
            "error": "#cf6679",
            "highlight_bg": "#d1e0ff"  # Changed from "#2d6cdf20" to a solid light blue color
        }
    def create_specialized_system(self):
        """Create a specialized system prompt for this chatbot"""
        try:
            self.status_label.configure(text="● Generating specialized knowledge...")
            self.window.update()
            meta_prompt = f"""
            I need to create a specialized chatbot for {self.specialization}.
            Please generate a comprehensive system prompt that will:
            1. Define the chatbot's expertise, knowledge boundaries and persona
            2. Specify how it should respond to questions in this domain
            3. Include relevant frameworks, methodologies, or approaches it should use
            4. Define its tone and communication style
            5. List any specific terminology it should be familiar with
            6.Answer should be in simple and easy to understand language and short and concise
            The output should be a well-structured system prompt that I can use directly.
            """
            response = self.gemini_model.generate_content(meta_prompt)
            self.specialized_prompt = response.text
            system_message = f"You are now a specialized assistant for {self.specialization}. Follow these guidelines: {self.specialized_prompt}"
            self.chat.send_message(system_message)
            self.window.after(0, self.append_welcome_message)
        except Exception as e:
            logging.error(f"Error creating specialized system: {str(e)}")
            self.window.after(0, lambda: self._show_error(f"Failed to create specialized system: {str(e)}"))
        finally:
            self.window.after(0, lambda: self.status_label.configure(text="Ready"))
            self.window.update()
    def create_widgets(self):
        """Create UI elements for the chat window"""
        self.main_frame = ctk.CTkFrame(
            self.window, 
            fg_color=self.theme["bg"],
            corner_radius=10,
            border_width=1,
            border_color=self.theme["accent"]
        )
        self.main_frame.pack(fill="both", expand=True, padx=10, pady=10)
        header_frame = ctk.CTkFrame(
            self.main_frame,
            fg_color=self.theme["highlight"],
            corner_radius=6,
            height=40
        )
        header_frame.pack(fill="x", padx=5, pady=5)
        header_frame.pack_propagate(False)
        title_label = ctk.CTkLabel(
            header_frame,
            text=f"{self.specialization.title()} Assistant",
            font=ctk.CTkFont(size=14, weight="bold"),
            text_color=self.theme["accent"]
        )
        title_label.pack(side="left", padx=10)
        self.load_data_button = ctk.CTkButton(
            header_frame,
            text="Load Data",
            font=ctk.CTkFont(size=12),
            fg_color=self.theme["button"],
            hover_color=self.theme["button_hover"],
            text_color="#ffffff",
            corner_radius=6,
            width=80,
            command=self.load_data
        )
        self.load_data_button.pack(side="right", padx=10)
        self.paned_container = ctk.CTkFrame(
            self.main_frame,
            fg_color=self.theme["bg"]
        )
        self.paned_container.pack(fill="both", expand=True, padx=5, pady=5)
        chat_frame = ctk.CTkFrame(
            self.paned_container,
            fg_color=self.theme["text_bg"],
            corner_radius=6,
            border_width=0
        )
        chat_frame.pack(fill="both", expand=True, padx=0, pady=(0, 5))
        chat_header = ctk.CTkFrame(
            chat_frame,
            fg_color=self.theme["highlight"],
            corner_radius=4,
            height=30
        )
        chat_header.pack(fill="x", padx=2, pady=2)
        chat_header.pack_propagate(False)
        chat_header_label = ctk.CTkLabel(
            chat_header,
            text="Conversation",
            font=ctk.CTkFont(size=12, weight="bold"),
            text_color=self.theme["accent"]
        )
        chat_header_label.pack(side="left", padx=10)
        self.chat_display = ctk.CTkTextbox(
            chat_frame,
            wrap="word",
            font=ctk.CTkFont(size=12),
            fg_color=self.theme["text_bg"],
            text_color=self.theme["fg"],
            corner_radius=0
        )
        self.chat_display.pack(fill="both", expand=True, padx=2, pady=2)
        textbox = self.chat_display._textbox
        textbox.tag_configure("timestamp", foreground="#666666")
        textbox.tag_configure("you", foreground="#4ecdc4")
        textbox.tag_configure("assistant", foreground="#2d6cdf")
        textbox.tag_configure("you_message", foreground=self.theme["fg"])
        textbox.tag_configure("assistant_message", foreground=self.theme["fg"])
        textbox.tag_configure("highlight", background=self.theme["highlight_bg"])
        textbox.tag_configure("error", foreground="#cf6679")
        input_container = ctk.CTkFrame(
            self.paned_container,
            fg_color=self.theme["bg"],
            corner_radius=6,
            height=100
        )
        input_container.pack(fill="x", pady=(0, 0))
        input_container.pack_propagate(False)
        input_header = ctk.CTkFrame(
            input_container,
            fg_color=self.theme["highlight"],
            corner_radius=4,
            height=30
        )
        input_header.pack(fill="x", padx=2, pady=2)
        input_header.pack_propagate(False)
        input_header_label = ctk.CTkLabel(
            input_header,
            text="Your Message",
            font=ctk.CTkFont(size=12, weight="bold"),
            text_color=self.theme["accent"]
        )
        input_header_label.pack(side="left", padx=10)
        input_frame = ctk.CTkFrame(
            input_container,
            fg_color=self.theme["bg"],
            corner_radius=0
        )
        input_frame.pack(fill="both", expand=True, padx=2, pady=2)
        self.input_text = ctk.CTkTextbox(
            input_frame,
            height=60,
            wrap="word",
            font=ctk.CTkFont(size=12),
            fg_color=self.theme["input_bg"],
            text_color="#ffffff",  # Brighter text for better visibility
            corner_radius=6,
            border_width=1,
            border_color=self.theme["accent"]
        )
        self.input_text.pack(side="left", fill="both", expand=True, padx=(0, 5))
        self.input_text.insert("1.0", f"Ask about {self.specialization}...")
        self.input_text._textbox.config(foreground="gray")
        self.input_text._textbox.bind("<FocusIn>", self._on_input_focus_in)
        self.input_text._textbox.bind("<FocusOut>", self._on_input_focus_out)
        self.send_button = ctk.CTkButton(
            input_frame,
            text="Send",
            font=ctk.CTkFont(size=12, weight="bold"),
            fg_color=self.theme["accent"],
            hover_color=self.theme["secondary"],
            text_color="#ffffff",
            corner_radius=6,
            width=80,
            height=35,
            command=self.send_message
        )
        self.send_button.pack(side="right", padx=5, pady=5)
        self.status_frame = ctk.CTkFrame(
            self.main_frame,
            fg_color=self.theme["bg"],
            corner_radius=0,
            height=30
        )
        self.status_frame.pack(fill="x", padx=5, pady=(0, 5))
        self.status_label = ctk.CTkLabel(
            self.status_frame,
            text="Ready",
            font=ctk.CTkFont(size=10),
            text_color=self.theme["secondary"]
        )
        self.status_label.pack(side="left", padx=5)
        self.input_text.focus_set()
    def _on_input_focus_in(self, event):
        """Handle input focus in event - clear placeholder"""
        if self.input_text.get("1.0", "end-1c") == f"Ask about {self.specialization}...":
            self.input_text.delete("1.0", "end")
            self.input_text._textbox.config(foreground=self.theme["fg"])
    def _on_input_focus_out(self, event):
        """Handle input focus out event - restore placeholder if empty"""
        if not self.input_text.get("1.0", "end-1c").strip():
            self.input_text.delete("1.0", "end")
            self.input_text.insert("1.0", f"Ask about {self.specialization}...")
            self.input_text._textbox.config(foreground="gray")
    def _process_message(self, message):
        """Process message and get response"""
        try:
            data_context = ""
            if self.loaded_data:
                data_context = f"Remember to use the additional data I've loaded when relevant to this query. "
            prompt = f"""{data_context}As a specialized assistant for {self.specialization}, 
            respond to the following user query with your expertise:
            User query: {message}
            Provide a helpful, accurate, concise, and clear response drawing on your specialized knowledge."""
            import google.generativeai as genai
            api_key = self.load_setting('gemini_api_key')
            model_name = self.load_setting("gemini_model")
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel(model_name)
            genai.configure(api_key=api_key)
            response = self.chat.send_message(prompt)
            response_text = response.text.replace("**", "")
            self.window.after(0, lambda: self._update_chat_display(response_text))
        except Exception as e:
            logging.error(f"Error processing message: {str(e)}")
            self.window.after(0, lambda: self._show_error(f"Error: {str(e)}"))
        finally:
            self.window.after(0, lambda: self.show_loading(False))
            self.processing = False
    def _on_input_focus_in(self, event):
        """Handle input focus in event - clear placeholder"""
        if self.input_text.get("1.0", "end-1c") == f"Ask about {self.specialization}...":
            self.input_text.delete("1.0", "end")
            self.input_text._textbox.config(foreground=self.theme["fg"])
    def _on_input_focus_out(self, event):
        """Handle input focus out event - restore placeholder if empty"""
        if not self.input_text.get("1.0", "end-1c").strip():
            self.input_text.delete("1.0", "end")
            self.input_text.insert("1.0", f"Ask about {self.specialization}...")
            self.input_text._textbox.config(foreground="gray")
    def append_welcome_message(self):
        """Add welcome message to chat"""
        welcome_text = f"Welcome to your specialized {self.specialization.title()} assistant! How can I help you today?"
        self._update_chat_display(welcome_text)
    def _update_chat_display(self, text, is_user=False):
        """Add message to chat display"""
        textbox = self.chat_display._textbox
        text = text.replace("**", "")
        timestamp = datetime.datetime.now().strftime("[%H:%M]")
        textbox.insert(tk.END, f"\n{timestamp} ", "timestamp")
        if is_user:
            textbox.insert(tk.END, "You: ", "you")
            textbox.insert(tk.END, f"{text}\n", "you_message")
        else:
            textbox.insert(tk.END, f"{self.specialization.title()} Assistant: ", "assistant")
            textbox.insert(tk.END, f"{text}\n", "assistant_message")
        self.chat_display.see(tk.END)      
    def _show_error(self, message):
        """Show error message in chat"""
        textbox = self.chat_display._textbox
        textbox.insert(tk.END, f"\nError: {message}\n", "error")
        self.chat_display.see(tk.END)
    def bind_shortcuts(self):
        """Bind keyboard shortcuts"""
        self.window.bind('<Return>', self.handle_send)
        self.window.bind('<Control-s>', lambda e: self.save_chat())
        self.window.bind('<Control-l>', lambda e: self.clear_chat())
        self.window.bind('<Escape>', lambda e: self.on_closing())
        self.input_text.bind('<Return>', self.handle_send)
        self.input_text.bind('<Control-Return>', lambda e: self.input_text.insert(tk.INSERT, '\n'))
    def send_message(self, event=None):
        """Send user message and get response"""
        if self.processing:
            return
        message = self.input_text.get("1.0", tk.END).strip()
        if message == f"Ask about {self.specialization}...":
            return
        if not message:
            return
        self.input_text.delete("1.0", tk.END)
        self._on_input_focus_out(None)
        self._update_chat_display(message, is_user=True)
        self.processing = True
        self.show_loading(True)
        threading.Thread(target=self._process_message, args=(message,), daemon=True).start()
    def handle_send(self, event=None):
        """Handle send button click or Enter key"""
        if event and event.state & 0x4:
            return
        else:
            self.send_message()
            return "break"
    def save_chat(self, event=None):
        """Save chat to file"""
        content = self.chat_display.get("1.0", tk.END)
        if not content.strip():
            return
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        specialization_name = ''.join(c if c.isalnum() else '_' for c in self.specialization.lower())
        filename = f"{specialization_name}_chat_{timestamp}.txt"
        try:
            file_path = filedialog.asksaveasfilename(
                defaultextension=".txt",
                initialfile=filename,
                filetypes=[("Text files", "*.txt"), ("All files", "*.*")]
            )
            if not file_path:
                return
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(f"{self.specialization.title()} Chat - {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                f.write(content)
            self.status_label.configure(text=f"Chat saved to {os.path.basename(file_path)}")
        except Exception as e:
            logging.error(f"Error saving chat: {str(e)}")
            self._show_error(f"Failed to save chat: {str(e)}")
    def clear_chat(self, event=None):
        """Clear chat history"""
        confirm = messagebox.askyesno("Clear Chat", "Are you sure you want to clear the chat history?")
        if confirm:
            self.chat_display.delete("1.0", "end")
            self.chat = self.gemini_model.start_chat(history=[])
            self.create_specialized_system()
            self.append_welcome_message()
    def show_loading(self, show=True):
        """Show/hide loading indicator"""
        if show:
            self.status_label.configure(text="● Thinking...")
            self.send_button.configure(state="disabled")
            self.input_text.configure(state="disabled")
            threading.Thread(target=self.animate_loading, daemon=True).start()
        else:
            self.status_label.configure(text="Ready")
            self.send_button.configure(state="normal")
            self.input_text.configure(state="normal")
            self.input_text.focus_set()
    def animate_loading(self):
        """Animate loading indicator"""
        dots = 1
        while self.processing:
            dots = (dots % 3) + 1
            self.window.after(500, lambda d=dots: self.status_label.configure(
                text=f"● Thinking{'.' * d}"
            ) if self.processing else None)
            time.sleep(0.5)
    def center_window(self):
        """Center window on screen"""
        self.window.update_idletasks()
        width = self.window.winfo_width()
        height = self.window.winfo_height()
        x = (self.window.winfo_screenwidth() // 2) - (width // 2)
        y = (self.window.winfo_screenheight() // 2) - (height // 2)
        self.window.geometry('{}x{}+{}+{}'.format(width, height, x, y))
    def on_closing(self):
        """Clean up and close window"""
        try:
            self.window.grab_release()
            self.window.destroy()
        except Exception as e:
            logging.error(f"Error closing window: {str(e)}")
    def load_data(self):
        """Load data from file to enhance chatbot knowledge"""
        try:
            file_types = [
                ("All Supported Files", "*.txt;*.pdf;*.doc;*.docx"),
                ("Text Files", "*.txt"),
                ("PDF Files", "*.pdf"),
                ("Word Documents", "*.doc;*.docx"),
                ("All Files", "*.*")
            ]
            file_path = filedialog.askopenfilename(
                title="Select Data File",
                filetypes=file_types
            )
            if not file_path:
                return
            file_ext = os.path.splitext(file_path)[1].lower()
            self.status_label.configure(text="● Loading data...")
            self.window.update()
            if file_ext == '.txt':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    self.loaded_data = file.read()
            elif file_ext == '.pdf':
                try:
                    import pypdf
                    pdf_text = []
                    with open(file_path, 'rb') as file:
                        pdf_reader = pypdf.PdfReader(file)
                        for page_num in range(len(pdf_reader.pages)):
                            page = pdf_reader.pages[page_num]
                            pdf_text.append(page.extract_text())
                    self.loaded_data = '\n'.join(pdf_text)
                except ImportError:
                    self._show_error("PDF extraction requires the pypdf library.")
                    self.status_label.configure(text="Ready")
                    return
            elif file_ext in ['.doc', '.docx']:
                try:
                    import docx
                    doc = docx.Document(file_path)
                    doc_text = []
                    for para in doc.paragraphs:
                        doc_text.append(para.text)
                    self.loaded_data = '\n'.join(doc_text)
                except ImportError:
                    self._show_error("Word document extraction requires the python-docx library.")
                    self.status_label.configure(text="Ready")
                    return
            if not self.loaded_data:
                self._show_error(f"No data could be extracted from {os.path.basename(file_path)}")
                self.status_label.configure(text="Ready")
                return
            threading.Thread(target=self._process_loaded_data, args=(file_path,), daemon=True).start()
        except Exception as e:
            logging.error(f"Error loading data: {str(e)}")
            self._show_error(f"Failed to load data: {str(e)}")
            self.status_label.configure(text="Ready")
    def _process_loaded_data(self, file_path):
        """Process loaded data and integrate into the chatbot"""
        try:
            max_length = 20000
            data_preview = self.loaded_data[:max_length]
            if len(self.loaded_data) > max_length:
                data_preview += f"\n\n[Note: Data was truncated. Total length: {len(self.loaded_data)} characters]"
            data_preview = data_preview.replace("**", "")
            self.window.after(0, lambda: self.status_label.configure(text="● Analyzing data..."))
            import google.generativeai as genai
            api_key = self.load_setting('gemini_api_key')
            model_name = self.load_setting("gemini_model")
            genai.configure(api_key=api_key)
            analysis_model = genai.GenerativeModel(model_name)
            process_prompt = f"""
            I need you to deeply analyze this data related to {self.specialization} and extract the most valuable insights.
            Your task:
            1. Identify the key concepts, facts, and relationships in this data
            2. Extract specific terminology and definitions relevant to {self.specialization}
            3. Determine patterns, trends, or important methodologies present
            4. Organize the knowledge into a concise, clear format with distinct sections
            5. Prioritize information most likely to be useful for answering questions about {self.specialization}
            Focus on providing a comprehensive but efficient analysis that can be used as a knowledge base.
            Be concise and clear in your extraction of information.
            Data from {os.path.basename(file_path)}:
            {data_preview}
            """
            response = analysis_model.generate_content(process_prompt)
            analysis_text = response.text.replace("**", "")
            self.window.after(0, lambda: self._update_chat_display(
                f"Data loaded and analyzed from {os.path.basename(file_path)}. I now have additional knowledge to assist with {self.specialization}.", 
                is_user=False
            ))
            data_context = f"""
            I've analyzed the data from {os.path.basename(file_path)} related to {self.specialization}.
            Here's the key information I've extracted:
            {analysis_text}
            Use this specialized knowledge along with your general knowledge when answering questions about {self.specialization}.
            """
            self.chat.send_message(data_context)
        except Exception as e:
            logging.error(f"Error processing loaded data: {str(e)}")
            self.window.after(0, lambda: self._show_error(f"Error processing data: {str(e)}"))
        finally:
            self.window.after(0, lambda: self.status_label.configure(text="Ready"))  
class SummarizeWindow:
    def __init__(self, parent):
        if not parent.winfo_exists():
            raise RuntimeError("Parent window no longer exists")
        self.window = tk.Toplevel(parent)
        self.window.title("Smart Notes Generator")
        self.window.geometry("800x600")
        self.window.configure(bg='#1a237e')
        self.parent = parent
        self.window.transient(parent)
        self.window.grab_set()
        self.create_widgets()
        self.center_window()
        self.window.protocol("WM_DELETE_WINDOW", self.on_closing)
    def on_closing(self):
        """Handle window closing properly"""
        try:
            if hasattr(self, 'window') and self.window.winfo_exists():
                self.window.grab_release()
                self.window.destroy()
        except Exception as e:
            logging.error(f"Error closing window: {e}")
    def cleanup_progress(self):
        """Clean up progress bar and reset button state"""
        try:
            if hasattr(self, 'window') and self.window.winfo_exists():
                if hasattr(self, 'progress_bar'):
                    self.progress_bar.stop()
                    self.progress_bar.pack_forget()
                if hasattr(self, 'convert_btn'):
                    self.convert_btn.configure(state='normal')
        except Exception as e:
            logging.error(f"Error cleaning up progress: {e}")
    def update_notes(self, notes):
        """Update the output text with generated notes"""
        try:
            if hasattr(self, 'window') and self.window.winfo_exists():
                if hasattr(self, 'output_text'):
                    self.output_text.configure(state=tk.NORMAL)
                    self.output_text.delete("1.0", tk.END)
                    self.output_text.insert("1.0", notes)
                    self.output_text.configure(state=tk.DISABLED)
                if hasattr(self, 'status_label'):
                    self.status_label.configure(text="Notes generated successfully")
        except Exception as e:
            logging.error(f"Error updating notes: {e}")
            self.show_error("Failed to update notes display")
    def show_error(self, error_msg):
        """Show error message to user"""
        try:
            if hasattr(self, 'window') and self.window.winfo_exists():
                messagebox.showerror("Error", f"Failed to convert notes: {error_msg}", parent=self.window)
                if hasattr(self, 'status_label'):
                    self.status_label.configure(text="Error generating notes")
        except Exception as e:
            logging.error(f"Error showing error message: {e}")
    def convert_to_notes(self):
        input_text = self.input_text.get('1.0', tk.END).strip()
        if not input_text:
            messagebox.showwarning("Warning", "Please enter some text to convert.", parent=self.window)
            return
        format_type = self.format_var.get()
        def run_conversion():
            try:
                self.progress_bar["value"] = 20
                self.window.update_idletasks()
                notes = self.basic_note_conversion(input_text, format_type)
                self.progress_bar["value"] = 60
                self.window.update_idletasks()
                final_notes = self.post_process_notes(notes, format_type)
                self.progress_bar["value"] = 100
                self.window.update_idletasks()
                NotesOutputWindow(self.window, final_notes)
                self.cleanup_progress()
            except Exception as e:
                self.show_error(f"Error during conversion: {str(e)}")
                self.cleanup_progress()
        threading.Thread(target=run_conversion, daemon=True).start()
    def create_widgets(self):
        main_frame = tk.Frame(self.window, bg='#1a237e')
        main_frame.pack(fill='both', expand=True, padx=20, pady=20)
        title_frame = tk.Frame(main_frame, bg='#1a237e')
        title_frame.pack(fill='x', pady=(0, 20))
        tk.Label(
            title_frame,
            text="Smart Notes Generator",
            font=("Segoe UI", 24, "bold"),
            bg='#1a237e',
            fg='white'
        ).pack(side='left')
        input_frame = tk.Frame(main_frame, bg='#1a237e')
        input_frame.pack(fill='both', expand=True)
        tk.Label(
            input_frame,
            text="Input Text:",
            font=("Segoe UI", 12),
            bg='#1a237e',
            fg='white'
        ).pack(anchor='w', pady=(0, 5))
        input_text_frame = tk.Frame(input_frame, bg='#1a237e')
        input_text_frame.pack(fill='both', expand=True, pady=10)
        input_scroll = ttk.Scrollbar(input_text_frame)
        input_scroll.pack(side='right', fill='y')
        self.input_text = tk.Text(
            input_text_frame,
            height=12,
            bg='#283593',
            fg='white',
            font=("Segoe UI", 11),
                wrap=tk.WORD,
            insertbackground='white',
            yscrollcommand=input_scroll.set
        )
        self.input_text.pack(side='left', fill='both', expand=True)
        input_scroll.configure(command=self.input_text.yview)
        format_frame = tk.Frame(main_frame, bg='#1a237e')
        format_frame.pack(fill='x', pady=10)
        tk.Label(
            format_frame,
            text="Select Note Format:",
            font=("Segoe UI", 12),
            bg='#1a237e',
            fg='white'
        ).pack(anchor='w', pady=(0, 5))
        self.format_var = tk.StringVar(value="bullet")
        formats = {
            "bullet": "• Bullet Points",
            "mindmap": "🔄 Mind Map",
            "numbered": "1. Numbered List",
            "outline": "I. Outline Format",
            "cornell": "📝 Cornell Notes"
        }
        options_frame = tk.Frame(format_frame, bg='#1a237e')
        options_frame.pack(fill='x', pady=5)
        for i, (value, text) in enumerate(formats.items()):
            radio_btn = tk.Radiobutton(
                options_frame,
                text=text,
                variable=self.format_var,
                value=value,
                bg='#1a237e',
                fg='white',
                selectcolor='#283593',
                activebackground='#1a237e',
                activeforeground='white',
                font=("Segoe UI", 11)
            )
            radio_btn.grid(row=i//3, column=i%3, padx=10, pady=5, sticky='w')
        controls_frame = tk.Frame(main_frame, bg='#1a237e')
        controls_frame.pack(fill='x', pady=10)
        self.progress_frame = tk.Frame(controls_frame, bg='#1a237e')
        self.progress_frame.pack(fill='x')
        self.progress_bar = ttk.Progressbar(
            self.progress_frame,
            mode='indeterminate',
            length=300
        )
        buttons_frame = tk.Frame(controls_frame, bg='#1a237e')
        buttons_frame.pack(pady=10)
        self.convert_btn = tk.Button(
            buttons_frame,
            text="🔄 Convert to Notes",
            command=self.convert_to_notes,
            bg='#4361ee',
            fg='white',
            font=("Segoe UI", 11, "bold"),
            relief='flat',
            cursor='hand2',
            width=15,
            height=1
        )
        self.convert_btn.pack(side='left', padx=5)
        self.clear_btn = tk.Button(
            buttons_frame,
            text="🗑️ Clear",
            command=self.clear_text,
            bg='#FF5252',
            fg='white',
            font=("Segoe UI", 11, "bold"),
            relief='flat',
            cursor='hand2',
            width=10,
            height=1
        )
        self.clear_btn.pack(side='left', padx=5)
        output_frame = tk.Frame(main_frame, bg='#1a237e')
        output_frame.pack(fill='both', expand=True)
        output_label_frame = tk.Frame(output_frame, bg='#1a237e')
        output_label_frame.pack(fill='x')
        tk.Label(
            output_label_frame,
            text="Generated Notes:",
            font=("Segoe UI", 12),
            bg='#1a237e',
            fg='white'
        ).pack(side='left')
        button_frame = tk.Frame(output_label_frame, bg='#1a237e')
        button_frame.pack(side='right')
        self.save_btn = tk.Button(
            button_frame,
            text="💾 Save",
            command=self.update_notes,
            bg='#4CAF50',
            fg='white',
            font=("Segoe UI", 11, "bold"),
            relief='flat',
            cursor='hand2'
        )
        self.save_btn.pack(side='right', padx=5)
        self.export_btn = tk.Button(
            button_frame,
            text="📤 Export",
            command=self.export_notes,
            bg='#FF9800',
            fg='white',
            font=("Segoe UI", 11, "bold"),
            relief='flat',
            cursor='hand2'
        )
        self.export_btn.pack(side='right', padx=5)
        output_text_frame = tk.Frame(output_frame, bg='#1a237e')
        output_text_frame.pack(fill='both', expand=True, pady=10)
        output_scroll = ttk.Scrollbar(output_text_frame)
        output_scroll.pack(side='right', fill='y')
        self.output_text = tk.Text(
            output_text_frame,
            height=12,
            bg='#283593',
            fg='white',
            font=("Segoe UI", 11),
            wrap=tk.WORD,
            yscrollcommand=output_scroll.set
        )
        self.output_text.pack(side='left', fill='both', expand=True)
        output_scroll.configure(command=self.output_text.yview)
        self.status_label = tk.Label(
            main_frame,
            text="Ready",
            font=("Segoe UI", 10),
            bg='#1a237e',
            fg='#8f9aff'
        )
        self.status_label.pack(pady=(10, 0))
    def post_process_notes(self, notes, format_type):
        """Post-process generated notes for better formatting"""
        if format_type == "bullet":
            lines = notes.split('\n')
            processed_lines = []
            for line in lines:
                line = line.strip()
                if line:
                    if not line.startswith('•'):
                        line = '• ' + line
                    if line.startswith('• • '):
                        line = '  ○ ' + line.replace('• • ', '')
                    elif line.startswith('• • • '):
                        line = '    ▪ ' + line.replace('• • • ', '')
                processed_lines.append(line)
            return '\n'.join(processed_lines)
        elif format_type == "numbered":
            lines = notes.split('\n')
            current_numbers = [0, 0, 0]
            processed_lines = []
            for line in lines:
                line = line.strip()
                if line:
                    if line[0].isdigit():
                        level = len(line.split('.')[0].split()) - 1
                    else:
                        level = 0
                    current_numbers[level] += 1
                    for i in range(level + 1, len(current_numbers)):
                        current_numbers[i] = 0
                    if level == 0:
                        number = f"{current_numbers[0]}."
                    else:
                        number = '.'.join(str(n) for n in current_numbers[:level+1] if n > 0) + '.'
                    indent = "  " * level
                    processed_lines.append(f"{indent}{number} {line.lstrip('0123456789. ')}")
            return '\n'.join(processed_lines)
        elif format_type == "outline":
            lines = notes.split('\n')
            current_level = 0
            processed_lines = []
            roman_numerals = ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X']
            current_numbers = {'I': 0, 'A': 0, '1': 0}
            for line in lines:
                line = line.strip()
                if line:
                    if line[0].isalpha() and line[0].isupper():
                        current_level = 1
                        current_numbers['A'] += 1
                        current_numbers['1'] = 0
                        prefix = f"  {chr(64 + current_numbers['A'])}."
                    elif line[0].isdigit():
                        current_level = 2
                        current_numbers['1'] += 1
                        prefix = f"    {current_numbers['1']}."
                    else:
                        current_level = 0
                        current_numbers['I'] += 1
                        current_numbers['A'] = 0
                        current_numbers['1'] = 0
                        prefix = f"{roman_numerals[min(current_numbers['I']-1, len(roman_numerals)-1)]}."
                    processed_lines.append(f"{prefix} {line.lstrip('IVX.ABC123 ')}")
            return '\n'.join(processed_lines)
        elif format_type == "mindmap":
            lines = notes.split('\n')
            processed_lines = []
            current_level = 0
            for i, line in enumerate(lines):
                line = line.strip()
                if line:
                    if i == 0:
                        processed_lines.append(f"📌 {line}")
                        processed_lines.append("   │")
                    else:
                        level = 1 if len(line) > 50 else 2 if len(line) > 30 else 3
                        if i == len(lines) - 1:
                            connector = "   └──"
                        else:
                            connector = "   ├──"
                        indent = "   │   " * (level - 1)
                        processed_lines.append(f"{indent}{connector} {line}")
            return '\n'.join(processed_lines)
        elif format_type == "cornell":
            sections = notes.split('\n\n')
            processed_lines = [
                "Cornell Notes",
                "═" * 50,
                "",
                "Main Ideas                    Notes",
                "──────────────               ─────────────────"
            ]
            main_ideas = []
            notes_content = []
            summary = ""
            for section in sections:
                if "Main Idea:" in section or "Key Point:" in section:
                    main_ideas.append(section.replace("Main Idea:", "•").replace("Key Point:", "•"))
                elif "Summary:" in section:
                    summary = section.replace("Summary:", "").strip()
                else:
                    notes_content.append(section)
            max_lines = max(len(main_ideas), len(notes_content))
            for i in range(max_lines):
                main_idea = main_ideas[i] if i < len(main_ideas) else ""
                note = notes_content[i] if i < len(notes_content) else ""
                processed_lines.append(f"{main_idea:<30} {note}")
            processed_lines.extend([
                "",
                "Summary",
                "───────",
                summary
            ])
            return '\n'.join(processed_lines)
        return notes
    def basic_note_conversion(self, text, format_type):
        """Fast and efficient note conversion"""
        sentences = [s.strip() for s in text.replace('\n', ' ').split('.') if s.strip()]
        if format_type == "bullet":
            return '\n'.join(f"• {sentence}" for sentence in sentences)
        elif format_type == "mindmap":
            if not sentences:
                return "Empty text"
            result = [
                "📌 Main Topic: " + sentences[0],
                "   │"
            ]
            current_branch = "   ├──🔸 "
            for i, sentence in enumerate(sentences[1:], 1):
                if i == len(sentences) - 1:
                    result.append("   └──🔸 " + sentence)
                else:
                    result.append(current_branch + sentence)
                    if i % 3 == 0:
                        result.append("   │")
            return '\n'.join(result)
        elif format_type == "numbered":
            return '\n'.join(f"{i+1}. {sentence}" for i, sentence in enumerate(sentences))
        elif format_type == "outline":
            result = []
            for i, sentence in enumerate(sentences):
                if i < 10:
                    result.append(f"{'I' * (i + 1)}. {sentence}")
                else:
                    result.append(f"    {chr(65 + (i % 26))}. {sentence}")
            return '\n'.join(result)
        else:
            if not sentences:
                return "Empty text"
            result = [
                "Cornell Notes",
                "═" * 50,
                "Main Ideas               Notes",
                "──────────               ─────"
            ]
            for i, sentence in enumerate(sentences):
                if i < 5:
                    main_idea = f"• Key Point {i+1}"
                    result.append(f"{main_idea:<20} {sentence}")
                else:
                    result.append(f"{'':20} • {sentence}")
            if sentences:
                result.extend([
                    "",
                    "Summary",
                    "───────",
                    "• " + sentences[0]
                ])
            return '\n'.join(result)
    def clear_text(self):
        """Clear both input and output text areas"""
        self.input_text.delete("1.0", tk.END)
        self.output_text.delete("1.0", tk.END)
        self.status_label.configure(text="Cleared all text")
    def export_notes(self):
        """Export notes to custom location"""
        try:
            from tkinter import filedialog
            notes_text = self.output_text.get("1.0", tk.END).strip()
            if not notes_text:
                messagebox.showwarning("Warning", "No notes to export", parent=self.window)
                return
            file_path = filedialog.asksaveasfilename(
                defaultextension=".txt",
                filetypes=[
                    ("Text files", "*.txt"),
                    ("Markdown files", "*.md"),
                    ("All files", "*.*")
                ],
                title="Export Notes"
            )
            if file_path:
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(notes_text)
                self.status_label.configure(text=f"Exported to {file_path}")
                messagebox.showinfo("Success", "Notes exported successfully!", parent=self.window)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to export notes: {str(e)}", parent=self.window)
    def center_window(self):
        self.window.update_idletasks()
        width = self.window.winfo_width()
        height = self.window.winfo_height()
        x = (self.window.winfo_screenwidth() // 2) - (width // 2)
        y = (self.window.winfo_screenheight() // 2) - (height // 2)
        self.window.geometry(f'{width}x{height}+{x}+{y}')
class EmailSender:
    def __init__(self):
        email = self.load_setting('email_address')
        password = self.load_setting('email_password')
        smtp_server = self.load_setting('smtp_server')
        smtp_port = self.load_setting('smtp_port')
        self.email_config = {}
        try:
            self.email_config = {
            'smtp_server': 'smtp.gmail.com',
            'smtp_port': 587,
            'email': 'samiullas831@gmail.com',
            'password':'hikk vnel ychc play'
        }
        except:
            pass
    def load_setting(self, setting_name):
        """Load setting from database"""
        try:
            conn = sqlite3.connect("settings.db")
            cursor = conn.cursor()
            cursor.execute('''CREATE TABLE IF NOT EXISTS settings
                            (name TEXT PRIMARY KEY, value TEXT)''')
            cursor.execute('SELECT value FROM settings WHERE name = ?', (setting_name,))
            result = cursor.fetchone()
            conn.close()
            return result[0] if result else None
        except Exception as e:
            logging.error(f"Error loading setting: {str(e)}")
            return None
    def send_email(self, to_email, subject, body):
        """Send an email with the provided details"""
        try:
            msg = MIMEMultipart()
            msg['From'] = self.email_config['email']
            msg['To'] = to_email
            msg['Subject'] = subject
            msg.attach(MIMEText(body, 'plain'))
            server = smtplib.SMTP(self.email_config['smtp_server'], self.email_config['smtp_port'])
            server.starttls()
            server.login(self.email_config['email'], self.email_config['password'])
            server.send_message(msg)
            server.quit()
            return True, "Email sent successfully!"
        except Exception as e:
            return False, f"Failed to send email: {str(e)}"
class EmailWindow:
    def __init__(self, parent):
        self.window = tk.Toplevel(parent)
        self.window.title("Send Email")
        self.window.geometry("500x600")
        self.window.configure(bg='#0C141F')  # Dark background
        self.window.transient(parent)
        self.window.grab_set()
        self.email_sender = EmailSender()
        self.create_widgets()
        self.center_window()
    def create_widgets(self):
        main_frame = tk.Frame(self.window, bg='#0C141F')
        main_frame.pack(fill='both', expand=True, padx=20, pady=20)
        content_frame = tk.Frame(main_frame, bg='#0C141F')
        content_frame.pack(fill='both', expand=True)
        tk.Label(content_frame, text="To:", bg='#0C141F', fg='#00F2FF', 
                font=('Orbitron', 12)).pack(anchor='w', pady=(0, 5))
        self.to_entry = tk.Entry(content_frame, width=50,
                               bg='#1D2733',
                               fg='#00F2FF',
                               insertbackground='#00F2FF',
                               font=('Orbitron', 10))
        self.to_entry.pack(fill='x', pady=(0, 15))
        tk.Label(content_frame, text="Subject:", bg='#0C141F', fg='#00F2FF', 
                font=('Orbitron', 12)).pack(anchor='w', pady=(0, 5))
        self.subject_entry = tk.Entry(content_frame, width=50,
                                    bg='#1D2733',
                                    fg='#00F2FF',
                                    insertbackground='#00F2FF',
                                    font=('Orbitron', 10))
        self.subject_entry.pack(fill='x', pady=(0, 15))
        tk.Label(content_frame, text="Message:", bg='#0C141F', fg='#00F2FF', 
                font=('Orbitron', 12)).pack(anchor='w', pady=(0, 5))
        self.message_text = tk.Text(content_frame, height=10, width=50,
                                  bg='#1D2733',
                                  fg='#00F2FF',
                                  insertbackground='#00F2FF',
                                  font=('Orbitron', 10))
        self.message_text.pack(fill='both', expand=True, pady=(0, 20))
        self.message_text.configure(highlightthickness=1,
                                  highlightcolor='#00F2FF',
                                  highlightbackground='#0066CC',
                                  relief='solid',
                                  borderwidth=1)
        button_frame = tk.Frame(main_frame, bg='#0C141F')
        button_frame.pack(side='bottom', fill='x', pady=(20, 0))
        self.send_button = tk.Button(
            button_frame,
            text="Send",
            command=self.send_email,
            bg='#00F2FF',
            fg='#0C141F',
            font=('Orbitron', 10, 'bold'),
            relief='flat',
            padx=20,
            pady=5,
            cursor='hand2'
        )
        self.send_button.pack(side='left', padx=10)
        self.cancel_button = tk.Button(
            button_frame,
            text="Cancel",
            command=self.window.destroy,
            bg='#00F2FF',
            fg='#0C141F',
            font=('Orbitron', 10, 'bold'),
            relief='flat',
            padx=20,
            pady=5,
            cursor='hand2'
        )
        self.cancel_button.pack(side='right', padx=10)
        def on_enter(e, button):
            button.configure(bg='#66F9FF')
        def on_leave(e, button):
            button.configure(bg='#00F2FF')
        self.send_button.bind('<Enter>', lambda e: on_enter(e, self.send_button))
        self.send_button.bind('<Leave>', lambda e: on_leave(e, self.send_button))
        self.cancel_button.bind('<Enter>', lambda e: on_enter(e, self.cancel_button))
        self.cancel_button.bind('<Leave>', lambda e: on_leave(e, self.cancel_button))
    def center_window(self):
        self.window.update_idletasks()
        width = self.window.winfo_width()
        height = self.window.winfo_height()
        x = (self.window.winfo_screenwidth() // 2) - (width // 2)
        y = (self.window.winfo_screenheight() // 2) - (height // 2)
        self.window.geometry(f'{width}x{height}+{x}+{y}')
    def send_email(self):
        to_email = self.to_entry.get()
        subject = self.subject_entry.get()
        body = self.message_text.get("1.0", tk.END)
        if not all([to_email, subject, body.strip()]):
            messagebox.showerror("Error", "Please fill in all fields", parent=self.window)
            return
        success, message = self.email_sender.send_email(to_email, subject, body)
        if success:
            messagebox.showinfo("Success", message, parent=self.window)
            self.window.destroy()
        else:
            messagebox.showerror("Error", message, parent=self.window)
class TranslatorWindow:
    def __init__(self, parent):
        self.window = tk.Toplevel(parent)
        self.window.title("Language Translator")
        self.window.geometry("800x600")
        self.window.configure(bg='#0c141f')  # Dark blue background
        self.PLAY_EMOJI = "▶️"
        self.STOP_EMOJI = "⏹️"
        self.is_translating = False
        style = ttk.Style()
        style.configure('Tron.TFrame', background='#0c141f')
        style.configure('Tron.TLabel', 
                       background='#0c141f', 
                       foreground='#00f2ff',  # Cyan text
                       font=('Orbitron', 10))
        style.configure('Tron.TButton',
                       background='#000000',
                       foreground='#00f2ff',
                       font=('Orbitron', 10))
        style.configure('Tron.TCombobox',
                       background='#1a2634',
                       foreground='#00f2ff',
                       selectbackground='#00f2ff',
                       selectforeground='#0c141f',
                       fieldbackground='#1a2634')
        self.language_dict = {
            "Afrikaans": "af", "Albanian": "sq", "Arabic": "ar", "Bengali": "bn",
            "Bulgarian": "bg", "Chinese (Simplified)": "zh-CN", "Croatian": "hr",
            "Czech": "cs", "Danish": "da", "Dutch": "nl", "English": "en",
            "Finnish": "fi", "French": "fr", "German": "de", "Greek": "el",
            "Gujarati": "gu", "Hindi": "hi", "Hungarian": "hu", "Indonesian": "id",
            "Italian": "it", "Japanese": "ja", "Kannada": "kn", "Korean": "ko",
            "Latin": "la", "Malay": "ms", "Malayalam": "ml", "Marathi": "mr",
            "Nepali": "ne", "Norwegian": "no", "Persian": "fa", "Polish": "pl",
            "Portuguese": "pt", "Punjabi": "pa", "Romanian": "ro", "Russian": "ru",
            "Serbian": "sr", "Slovak": "sk", "Spanish": "es", "Swedish": "sv",
            "Tamil": "ta", "Telugu": "te", "Thai": "th", "Turkish": "tr",
            "Ukrainian": "uk", "Urdu": "ur", "Vietnamese": "vi"
        }
        self.create_widgets()
        self.center_window()
    def create_widgets(self):
        main_frame = ttk.Frame(self.window, padding="20", style='Tron.TFrame')
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        self.window.columnconfigure(0, weight=1)
        self.window.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        ttk.Label(main_frame, text="From Language:", style='Tron.TLabel').grid(row=0, column=0, sticky=tk.W, pady=5)
        self.source_lang = ttk.Combobox(main_frame, values=sorted(self.language_dict.keys()), width=30, state='readonly', style='Tron.TCombobox')
        self.source_lang.grid(row=0, column=1, sticky=(tk.W, tk.E), pady=5, padx=(10, 0))
        self.source_lang.set("English")
        ttk.Label(main_frame, text="To Language:", style='Tron.TLabel').grid(row=1, column=0, sticky=tk.W, pady=5)
        self.target_lang = ttk.Combobox(main_frame, values=sorted(self.language_dict.keys()), width=30, state='readonly', style='Tron.TCombobox')
        self.target_lang.grid(row=1, column=1, sticky=(tk.W, tk.E), pady=5, padx=(10, 0))
        self.target_lang.set("Spanish")
        ttk.Label(main_frame, text="Enter text to translate:", style='Tron.TLabel').grid(row=2, column=0, columnspan=2, sticky=tk.W, pady=(10, 5))
        self.source_text = scrolledtext.ScrolledText(
            main_frame, wrap=tk.WORD, width=60, height=8,
            font=('Consolas', 10),
            bg='#1a2634',  # Darker blue
            fg='#00f2ff',  # Cyan text
            insertbackground='#00f2ff'  # Cyan cursor
        )
        self.source_text.grid(row=3, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=5)
        button_frame = ttk.Frame(main_frame, style='Tron.TFrame')
        button_frame.grid(row=4, column=0, columnspan=2, pady=10)
        self.translate_btn = tk.Button(
            button_frame,
            text=f"{self.PLAY_EMOJI} Translate",
            command=self.toggle_translate,
            font=('Orbitron', 10, 'bold'),
            bg='#1a2634',
            fg='#00f2ff',
            activebackground='#00f2ff',
            activeforeground='#0c141f',
            relief='flat',
            padx=20,
            cursor='hand2'
        )
        self.translate_btn.pack(side=tk.LEFT, padx=5)
        self.is_translating = False
        clear_btn = tk.Button(
            button_frame,
            text="⌧ Clear",
            command=self.clear_text,
            font=('Orbitron', 10),
            bg='#1a2634',
            fg='#00f2ff',
            activebackground='#00f2ff',
            activeforeground='#0c141f',
            relief='flat',
            padx=20,
            cursor='hand2'
        )
        clear_btn.pack(side=tk.LEFT, padx=5)
        ttk.Label(main_frame, text="Translation:", style='Tron.TLabel').grid(row=5, column=0, columnspan=2, sticky=tk.W, pady=(10, 5))
        self.translated_text = scrolledtext.ScrolledText(
            main_frame, wrap=tk.WORD, width=60, height=8,
            font=('Consolas', 10),
            bg='#1a2634',  # Darker blue
            fg='#00f2ff',  # Cyan text
            insertbackground='#00f2ff'  # Cyan cursor
        )
        self.translated_text.grid(row=6, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=5)
        self.translated_text.configure(state='disabled')
    def toggle_translate(self):
        if not self.is_translating:
            self.translate()
            self.translate_btn.configure(text=f"{self.STOP_EMOJI} Stop")
            self.is_translating = True
        else:
            self.stop_translation()
            self.translate_btn.configure(text=f"{self.PLAY_EMOJI} Translate")
            self.is_translating = False
    def stop_translation(self):
        self.translated_text.configure(state='normal')
        self.translated_text.delete("1.0", tk.END)
        self.translated_text.insert("1.0", "Translation stopped")
        self.translated_text.configure(state='disabled')
    def translate(self):
        try:
            source_text = self.source_text.get("1.0", tk.END).strip()
            if not source_text:
                messagebox.showwarning("Warning", "Please enter text to translate.")
                self.translate_btn.configure(text=f"{self.PLAY_EMOJI} Translate")
                self.is_translating = False
                return
            source_lang_code = self.language_dict[self.source_lang.get()]
            target_lang_code = self.language_dict[self.target_lang.get()]
            self.translated_text.configure(state='normal')
            self.translated_text.delete("1.0", tk.END)
            self.translated_text.insert("1.0", "⚡ Translating...")
            self.translated_text.update()
            translator = GoogleTranslator(source=source_lang_code, target=target_lang_code)
            translated = translator.translate(source_text)
            self.translated_text.delete("1.0", tk.END)
            self.translated_text.insert("1.0", translated)
            self.translate_btn.configure(text=f"{self.PLAY_EMOJI} Translate")
            self.is_translating = False
        except Exception as e:
            self.translated_text.delete("1.0", tk.END)
            self.translated_text.insert("1.0", f"Translation failed: {str(e)}")
            self.translate_btn.configure(text=f"{self.PLAY_EMOJI} Translate")
            self.is_translating = False
        finally:
            self.translated_text.configure(state='disabled')
    def clear_text(self):
        self.source_text.delete("1.0", tk.END)
        self.translated_text.configure(state='normal')
        self.translated_text.delete("1.0", tk.END)
        self.translated_text.configure(state='disabled')
    def on_closing(self):
        self.window.destroy()
    def center_window(self):
        self.window.update_idletasks()
        width = self.window.winfo_width()
        height = self.window.winfo_height()
        x = (self.window.winfo_screenwidth() // 2) - (width // 2)
        y = (self.window.winfo_screenheight() // 2) - (height // 2)
        self.window.geometry(f"{width}x{height}+{x}+{y}")
class ContactManager:
    def __init__(self, parent):
        self.window = tk.Toplevel(parent)
        self.window.title("Contact Manager")
        self.window.geometry("600x500")
        self.window.configure(bg='#1a237e')
        self.window.transient(parent)
        self.window.grab_set()
        self.init_database()
        self.create_widgets()
        self.center_window()
    def init_database(self):
        """Initialize the contacts database"""
        try:
            self.conn = sqlite3.connect("samcontact.db")
            self.cursor = self.conn.cursor()
            self.cursor.execute('''
                CREATE TABLE IF NOT EXISTS contacts (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    name TEXT NOT NULL,
                    number TEXT NOT NULL,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            self.conn.commit()
            logging.info("Database initialized successfully")
        except Exception as e:
            logging.error(f"Database initialization error: {str(e)}")
            messagebox.showerror("Database Error", f"Failed to initialize database: {str(e)}", parent=self.window)
    def create_widgets(self):
        main_frame = tk.Frame(self.window, bg='#1a237e')
        main_frame.pack(fill='both', expand=True, padx=20, pady=20)
        tk.Label(
            main_frame,
            text="Contact Manager",
            font=("Segoe UI", 24, "bold"),
            bg='#1a237e',
            fg='white'
        ).pack(pady=(0, 20))
        input_frame = tk.Frame(main_frame, bg='#1a237e')
        input_frame.pack(fill='x', pady=10)
        tk.Label(
            input_frame,
            text="Name:",
            font=("Segoe UI", 12),
            bg='#1a237e',
            fg='white'
        ).pack(anchor='w')
        self.name_entry = tk.Entry(
            input_frame,
            font=("Segoe UI", 11),
            bg='#283593',
            fg='white',
            insertbackground='white'
        )
        self.name_entry.pack(fill='x', pady=(5, 10))
        tk.Label(
            input_frame,
            text="Number:",
            font=("Segoe UI", 12),
            bg='#1a237e',
            fg='white'
        ).pack(anchor='w')
        self.number_entry = tk.Entry(
            input_frame,
            font=("Segoe UI", 11),
            bg='#283593',
            fg='white',
            insertbackground='white'
        )
        self.number_entry.pack(fill='x', pady=(5, 10))
        buttons_frame = tk.Frame(input_frame, bg='#1a237e')
        buttons_frame.pack(fill='x', pady=10)
        tk.Button(
            buttons_frame,
            text="Add Contact",
            command=self.add_contact,
            bg='#4CAF50',
            fg='white',
            font=("Segoe UI", 11, "bold"),
            relief='flat',
            cursor='hand2'
        ).pack(side='left', padx=5)
        tk.Button(
            buttons_frame,
            text="Update Contact",
            command=self.update_contact,
            bg='#FF9800',
            fg='white',
            font=("Segoe UI", 11, "bold"),
            relief='flat',
            cursor='hand2'
        ).pack(side='left', padx=5)
        tk.Button(
            buttons_frame,
            text="Delete Contact",
            command=self.delete_contact,
            bg='#f44336',
            fg='white',
            font=("Segoe UI", 11, "bold"),
            relief='flat',
            cursor='hand2'
        ).pack(side='left', padx=5)
        list_frame = tk.Frame(main_frame, bg='#1a237e')
        list_frame.pack(fill='both', expand=True, pady=10)
        scrollbar = ttk.Scrollbar(list_frame)
        scrollbar.pack(side='right', fill='y')
        self.contacts_list = tk.Listbox(
            list_frame,
            bg='#283593',
            fg='white',
            font=("Segoe UI", 11),
            selectmode='single',
            yscrollcommand=scrollbar.set
        )
        self.contacts_list.pack(side='left', fill='both', expand=True)
        scrollbar.configure(command=self.contacts_list.yview)
        self.contacts_list.bind('<<ListboxSelect>>', self.on_select)
        self.load_contacts()
    def load_contacts(self):
        """Load contacts from database"""
        try:
            self.contacts_list.delete(0, tk.END)
            self.cursor.execute('SELECT name, number FROM contacts ORDER BY name')
            contacts = self.cursor.fetchall()
            for contact in contacts:
                self.contacts_list.insert(tk.END, f"{contact[0]} - {contact[1]}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load contacts: {str(e)}", parent=self.window)
    def add_contact(self):
        """Add a new contact"""
        try:
            name = self.name_entry.get().strip()
            number = self.number_entry.get().strip()
            if not name or not number:
                messagebox.showwarning("Warning", "Please enter both name and number", parent=self.window)
                return
            self.cursor.execute('INSERT INTO contacts (name, number) VALUES (?, ?)', (name, number))
            self.conn.commit()
            self.name_entry.delete(0, tk.END)
            self.number_entry.delete(0, tk.END)
            self.load_contacts()
            messagebox.showinfo("Success", "Contact added successfully!", parent=self.window)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to add contact: {str(e)}", parent=self.window)
    def update_contact(self):
        """Update selected contact"""
        try:
            selection = self.contacts_list.curselection()
            if not selection:
                messagebox.showwarning("Warning", "Please select a contact to update", parent=self.window)
                return
            name = self.name_entry.get().strip()
            number = self.number_entry.get().strip()
            if not name or not number:
                messagebox.showwarning("Warning", "Please enter both name and number", parent=self.window)
                return
            old_contact = self.contacts_list.get(selection[0])
            old_name = old_contact.split(' - ')[0]
            self.cursor.execute('UPDATE contacts SET name=?, number=? WHERE name=?', (name, number, old_name))
            self.conn.commit()
            self.name_entry.delete(0, tk.END)
            self.number_entry.delete(0, tk.END)
            self.load_contacts()
            messagebox.showinfo("Success", "Contact updated successfully!", parent=self.window)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to update contact: {str(e)}", parent=self.window)
    def delete_contact(self):
        """Delete selected contact"""
        try:
            selection = self.contacts_list.curselection()
            if not selection:
                messagebox.showwarning("Warning", "Please select a contact to delete", parent=self.window)
                return
            if not messagebox.askyesno("Confirm", "Are you sure you want to delete this contact?", parent=self.window):
                return
            contact = self.contacts_list.get(selection[0])
            name = contact.split(' - ')[0]
            self.cursor.execute('DELETE FROM contacts WHERE name=?', (name,))
            self.conn.commit()
            self.name_entry.delete(0, tk.END)
            self.number_entry.delete(0, tk.END)
            self.load_contacts()
            messagebox.showinfo("Success", "Contact deleted successfully!", parent=self.window)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to delete contact: {str(e)}", parent=self.window)
    def on_select(self, event):
        """Handle contact selection"""
        try:
            selection = self.contacts_list.curselection()
            if selection:
                contact = self.contacts_list.get(selection[0])
                name, number = contact.split(' - ')
                self.name_entry.delete(0, tk.END)
                self.name_entry.insert(0, name)
                self.number_entry.delete(0, tk.END)
                self.number_entry.insert(0, number)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to select contact: {str(e)}", parent=self.window)
    def center_window(self):
        """Center the window on screen"""
        self.window.update_idletasks()
        width = self.window.winfo_width()
        height = self.window.winfo_height()
        x = (self.window.winfo_screenwidth() // 2) - (width // 2)
        y = (self.window.winfo_screenheight() // 2) - (height // 2)
        self.window.geometry(f'{width}x{height}+{x}+{y}')
class SummaryOutputWindow:
    def __init__(self, parent, summary_text):
        self.window = tk.Toplevel(parent)
        self.window.title("Generated Summary")
        self.window.geometry("600x400")
        self.window.configure(bg='#1a237e')
        self.window.transient(parent)
        self.window.grab_set()
        self.summary_text = summary_text
        self.create_widgets(summary_text)
        self.center_window()
    def create_widgets(self, summary_text):
        main_frame = tk.Frame(self.window, bg='#1a237e')
        main_frame.pack(fill='both', expand=True, padx=20, pady=20)
        title_frame = tk.Frame(main_frame, bg='#1a237e')
        title_frame.pack(fill='x', pady=(0, 20))
        tk.Label(
            title_frame,
            text="Generated Summary",
            font=("Segoe UI", 24, "bold"),
            bg='#1a237e',
            fg='white'
        ).pack(side='left')
        buttons_frame = tk.Frame(title_frame, bg='#1a237e')
        buttons_frame.pack(side='right')
        self.notes_btn = tk.Button(
            buttons_frame,
            text="📝 Convert to Notes",
            command=self.convert_to_notes,
            bg='#4CAF50',
            fg='white',
            font=("Segoe UI", 11, "bold"),
            relief='flat',
            cursor='hand2'
        )
        self.notes_btn.pack(side='right', padx=5)
        self.copy_btn = tk.Button(
            buttons_frame,
            text="📋 Copy",
            command=self.copy_to_clipboard,
            bg='#FF9800',
            fg='white',
            font=("Segoe UI", 11, "bold"),
            relief='flat',
            cursor='hand2'
        )
        self.copy_btn.pack(side='right', padx=5)
        text_frame = tk.Frame(main_frame, bg='#1a237e')
        text_frame.pack(fill='both', expand=True)
        scrollbar = ttk.Scrollbar(text_frame)
        scrollbar.pack(side='right', fill='y')
        self.text = tk.Text(
            text_frame,
            wrap=tk.WORD,
            bg='#283593',
            fg='white',
            font=("Segoe UI", 11),
            insertbackground='white',
            selectbackground='#4361ee',
            selectforeground='white',
            yscrollcommand=scrollbar.set
        )
        self.text.pack(fill='both', expand=True)
        scrollbar.configure(command=self.text.yview)
        self.text.insert('1.0', summary_text)
        self.text.configure(state='disabled')
        self.status_label = tk.Label(
            main_frame,
            text="Ready",
            font=("Segoe UI", 10),
            bg='#1a237e',
            fg='#8f9aff'
        )
        self.status_label.pack(pady=(10, 0))
    def copy_to_clipboard(self):
        """Copy summary to clipboard"""
        try:
            self.window.clipboard_clear()
            self.window.clipboard_append(self.summary_text)
            self.status_label.configure(text="Copied to clipboard")
            messagebox.showinfo("Success", "Summary copied to clipboard!", parent=self.window)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to copy to clipboard: {str(e)}", parent=self.window)
    def center_window(self):
        self.window.update_idletasks()
        width = self.window.winfo_width()
        height = self.window.winfo_height()
        x = (self.window.winfo_screenwidth() // 2) - (width // 2)
        y = (self.window.winfo_screenheight() // 2) - (height // 2)
        self.window.geometry(f'{width}x{height}+{x}+{y}')
class NotesOutputWindow:
    def __init__(self, parent, notes_text):
        self.window = tk.Toplevel(parent)
        self.window.title("Generated Smart Notes")
        self.window.configure(bg='#0C141F')  # Dark blue background
        window_width = 800
        window_height = 600
        self.window.geometry(f"{window_width}x{window_height}")
        self.center_window()
        self.window.resizable(False, False)
        self.create_widgets(notes_text)
    def create_widgets(self, notes_text):
        main_frame = tk.Frame(self.window, bg='#0C141F')
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        title_frame = tk.Frame(main_frame, bg='#0C141F')
        title_frame.pack(fill=tk.X, pady=(0, 15))
        title_label = tk.Label(title_frame, 
                             text="SMART NOTES OUTPUT",
                             font=('Orbitron', 16, 'bold'),
                             fg='#00F0FF',  # Cyan glow
                             bg='#0C141F')
        title_label.pack()
        text_container = tk.Frame(main_frame, 
                                bg='#00F0FF',  # Cyan border
                                padx=2, pady=2)
        text_container.pack(fill=tk.BOTH, expand=True)
        inner_frame = tk.Frame(text_container, bg='#0C141F')
        inner_frame.pack(fill=tk.BOTH, expand=True)
        self.notes_text = tk.Text(inner_frame,
                                wrap=tk.WORD,
                font=('Consolas', 11),
                                bg='#0D1B2A',  # Darker blue for text area
                                fg='#00F0FF',  # Cyan text
                                insertbackground='#00F0FF',  # Cursor color
                                selectbackground='#1B3A5C',  # Selection background
                                selectforeground='#00F0FF',  # Selection text color
                                padx=10, pady=10,
                                relief=tk.FLAT,
                                bd=0)
        self.notes_text.pack(fill=tk.BOTH, expand=True, side=tk.LEFT)
        self.notes_text.insert('1.0', notes_text)
        self.notes_text.configure(state='disabled')
        scrollbar = ttk.Scrollbar(inner_frame, 
                                orient='vertical', 
                                command=self.notes_text.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.notes_text['yscrollcommand'] = scrollbar.set
        button_frame = tk.Frame(main_frame, bg='#0C141F')
        button_frame.pack(fill=tk.X, pady=(15, 0))
        copy_button = tk.Button(button_frame,
                              text="COPY TO CLIPBOARD",
                              command=self.copy_to_clipboard,
                              font=('Orbitron', 10),
                              bg='#1B3A5C',  # Darker blue for button
                              fg='#00F0FF',
                              activebackground='#00F0FF',
                              activeforeground='#0C141F',
                              relief=tk.FLAT,
                              bd=0,
                              padx=20,
                              pady=5)
        copy_button.pack(side=tk.LEFT)
        copy_button.bind('<Enter>', lambda e: copy_button.configure(
            bg='#00F0FF', fg='#0C141F'))
        copy_button.bind('<Leave>', lambda e: copy_button.configure(
            bg='#1B3A5C', fg='#00F0FF'))
        export_button = tk.Button(button_frame,
                                text="EXPORT",
                                command=self.export_notes,
                                font=('Orbitron', 10),
                                bg='#1B3A5C',  # Darker blue for button
                                fg='#00F0FF',
                                activebackground='#00F0FF',
                                activeforeground='#0C141F',
                                relief=tk.FLAT,
                                bd=0,
                                padx=20,
                                pady=5)
        export_button.pack(side=tk.RIGHT)
        export_button.bind('<Enter>', lambda e: export_button.configure(
            bg='#00F0FF', fg='#0C141F'))
        export_button.bind('<Leave>', lambda e: export_button.configure(
            bg='#1B3A5C', fg='#00F0FF'))
    def copy_to_clipboard(self):
        notes_content = self.notes_text.get('1.0', tk.END)
        self.window.clipboard_clear()
        self.window.clipboard_append(notes_content)
        temp_label = tk.Label(self.window,
                            text="Copied to clipboard!",
                            fg='#00F0FF',
                            bg='#0C141F',
                            font=('Orbitron', 10))
        temp_label.place(relx=0.5, rely=0.95, anchor='center')
        self.window.after(2000, temp_label.destroy)
    def export_notes(self):
        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Text files", "*.txt"), ("All files", "*.*")],
            title="Export Notes"
        )
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as file:
                    notes_content = self.notes_text.get('1.0', tk.END)
                    file.write(notes_content)
                temp_label = tk.Label(self.window,
                                    text="Notes exported successfully!",
                                    fg='#00F0FF',
                                    bg='#0C141F',
                                    font=('Orbitron', 10))
                temp_label.place(relx=0.5, rely=0.95, anchor='center')
                self.window.after(2000, temp_label.destroy)
            except Exception as e:
                messagebox.showerror("Export Error", f"Failed to export notes: {str(e)}")
    def center_window(self):
        self.window.update_idletasks()
        width = self.window.winfo_width()
        height = self.window.winfo_height()
        x = (self.window.winfo_screenwidth() // 2) - (width // 2)
        y = (self.window.winfo_screenheight() // 2) - (height // 2)
        self.window.geometry(f'{width}x{height}+{x}+{y}')
class ProgrammingWindow:
    def __init__(self, parent):
        self.window = tk.Toplevel(parent.root)
        self.window.title("AI Programming Assistant")
        self.window.geometry("1200x800")
        self.window.configure(bg="#0C141F")
        self.main_app = parent
        self.buttons = {}
        try:
            import google.generativeai as genai
            api_key = self.load_setting('gemini_api_key')
            model_name = self.load_setting("gemini_model")
            genai.configure(api_key=api_key)
            self.model = genai.GenerativeModel(model_name)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to initialize Gemini: {str(e)}")
            self.window.destroy()
            return
        self.requirements_text = None
        self.code_display = None
        self.status_var = tk.StringVar()
        self.create_widgets()
        self.center_window()
        self.setup_bindings()
        self.window.transient(parent.root)
        self.window.grab_set()
        self.window.protocol("WM_DELETE_WINDOW", self.on_closing)
    def load_setting(self, setting_name):
        """Load setting from database"""
        try:
            conn = sqlite3.connect("settings.db")
            cursor = conn.cursor()
            cursor.execute('''CREATE TABLE IF NOT EXISTS settings
                            (name TEXT PRIMARY KEY, value TEXT)''')
            cursor.execute('SELECT value FROM settings WHERE name = ?', (setting_name,))
            result = cursor.fetchone()
            conn.close()
            return result[0] if result else None
        except Exception as e:
            logging.error(f"Error loading setting: {str(e)}")
            return None
    def create_widgets(self):
        style = ttk.Style()
        style.configure("TRON.TButton",
                       padding=8,
                       font=('Orbitron', 10),
                       background="#00F0FF",
                       foreground="#0C141F")
        style.configure("TRON.TLabel",
                       font=('Orbitron', 12, 'bold'),
                       foreground="#00F0FF",
                       background="#0C141F")
        style.configure("TRON.TFrame",
                       background="#0C141F")
        main_frame = ttk.Frame(self.window, padding="20", style="TRON.TFrame")
        main_frame.pack(fill=tk.BOTH, expand=True)
        title_label = ttk.Label(main_frame, 
                              text="AI PROGRAMMING MATRIX", 
                              font=('Orbitron', 24, 'bold'),
                              foreground="#00F0FF",
                              background="#0C141F")
        title_label.pack(fill=tk.X, pady=(0, 20))
        req_label = ttk.Label(main_frame, 
                            text="PROGRAM REQUIREMENTS",
                            style="TRON.TLabel")
        req_label.pack(fill=tk.X, pady=(0, 5))
        req_frame = ttk.Frame(main_frame, style="TRON.TFrame")
        req_frame.pack(fill=tk.X, pady=(0, 15))
        self.requirements_text = tk.Text(req_frame, 
                              height=5, 
                              wrap=tk.WORD,
                              bg="#1A2030",
                              fg="#00F0FF",
                              insertbackground="#00F0FF",
                              font=('Share Tech Mono', 12),
                bd=2,
                              relief="solid",
                              highlightthickness=1,
                              highlightcolor="#00F0FF",
                              highlightbackground="#004A5D")
        self.requirements_text.pack(fill=tk.X, padx=5, pady=5)
        example_text = "Example: Create a Python program that reads a CSV file, processes the data, and generates a bar chart using matplotlib."
        self.requirements_text.insert('1.0', example_text)
        self.requirements_text.bind('<FocusIn>', lambda e: self.clear_example_text(e))
        btn_frame = ttk.Frame(main_frame, style="TRON.TFrame")
        btn_frame.pack(fill=tk.X, pady=(0, 15))
        button_configs = [
            ("GENERATE PROGRAM", self.generate_program, "Enter"),
            ("CLEAR ALL", self.clear_all, "Ctrl+L"),
            ("COPY CODE", self.copy_to_clipboard, "Ctrl+C"),
            ("SAVE CODE", self.save_code, "Ctrl+S")
        ]
        for text, command, shortcut in button_configs:
            btn = tk.Button(btn_frame,
                          text=f"{text}\n{shortcut}",
                          command=command,
                          font=('Orbitron', 9),
                          bg="#0C141F",
                          fg="#00F0FF",
                          activebackground="#00F0FF",
                          activeforeground="#0C141F",
                          relief="solid",
                          bd=1,
                          padx=15,
                          pady=5)
            btn.pack(side=tk.LEFT, padx=5, expand=True)
            self.buttons[text] = btn
            btn.bind('<Enter>', lambda e, b=btn: self.on_button_hover(e, b))
            btn.bind('<Leave>', lambda e, b=btn: self.on_button_leave(e, b))
        code_label = ttk.Label(main_frame, 
                             text="GENERATED CODE OUTPUT", 
                             style="TRON.TLabel")
        code_label.pack(fill=tk.X, pady=(0, 5))
        code_frame = ttk.Frame(main_frame, style="TRON.TFrame")
        code_frame.pack(fill=tk.BOTH, expand=True)
        self.code_display = tk.Text(code_frame,
                                  wrap=tk.NONE,
                                  bg="#1A2030",
                                  fg="#00F0FF",
                                  insertbackground="#00F0FF",
                                  font=('Share Tech Mono', 12),
                                  bd=2,
                                  relief="solid",
                                  highlightthickness=1,
                                  highlightcolor="#00F0FF",
                                  highlightbackground="#004A5D")
        self.code_display.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        y_scrollbar = ttk.Scrollbar(code_frame, 
                                  orient=tk.VERTICAL,
                                  command=self.code_display.yview)
        y_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        x_scrollbar = ttk.Scrollbar(main_frame,
                                  orient=tk.HORIZONTAL,
                                  command=self.code_display.xview)
        x_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)
        self.code_display.configure(yscrollcommand=y_scrollbar.set,
                                  xscrollcommand=x_scrollbar.set)
        status_bar = ttk.Label(main_frame,
                             textvariable=self.status_var,
                             style="TRON.TLabel",
                             relief=tk.SUNKEN,
                             anchor=tk.W)
        status_bar.pack(side=tk.BOTTOM, fill=tk.X, padx=10, pady=5)
        self.setup_syntax_highlighting()
    def setup_syntax_highlighting(self):
        """Configure syntax highlighting tags"""
        self.code_display.tag_configure(
            "keyword",
            foreground="#00F0FF",
            font=("Consolas", 11, "bold")
        )
        self.code_display.tag_configure(
            "builtin",
            foreground="#FF00FF",
            font=("Consolas", 11, "bold")
        )
        self.code_display.tag_configure(
            "string",
            foreground="#FF9C00"
        )
        self.code_display.tag_configure(
            "comment",
            foreground="#0088CC",
            font=("Consolas", 11, "italic")
        )
        self.code_display.tag_configure(
            "function",
            foreground="#00FF9C"
        )
        self.code_display.tag_configure(
            "class",
            foreground="#FF10F0",
            font=("Consolas", 11, "bold")
        )
        self.code_display.tag_configure(
            "number",
            foreground="#FFB000"
        )
        self.code_display.tag_configure(
            "operator",
            foreground="#FF3366",
            font=("Consolas", 11, "bold")
        )
        self.code_display.tag_configure(
            "decorator",
            foreground="#00FFCC",
            font=("Consolas", 11, "italic")
        )
        self.code_display.tag_configure(
            "type_hint",
            foreground="#BB86FC",
            font=("Consolas", 11, "bold")
        )
        self.code_display.tag_configure(
            "special",
            foreground="#FF5000",
            font=("Consolas", 11, "bold")
        )
    def generate_program(self):
        loading_label = None
        try:
            requirements = self.requirements_text.get("1.0", tk.END).strip()
            if not requirements or requirements == "Example: Create a Python program that reads a CSV file, processes the data, and generates a bar chart using matplotlib.":
                self.show_notification("Please enter program requirements first!")
                return
            self.buttons["GENERATE PROGRAM"].configure(state="disabled")
            loading_label = tk.Label(self.window, text="Generating code...", fg="#00ff00", bg="#1e1e1e")
            loading_label.place(relx=0.5, rely=0.4, anchor="center")
            self.window.update()
            self.code_display.delete("1.0", tk.END)
            for tag in self.code_display.tag_names():
                self.code_display.tag_remove(tag, "1.0", tk.END)
            custom_prompt = self.main_app.load_prompt('programming')
            if custom_prompt:
                prompt = f"""{custom_prompt},
Requirements:
{requirements}
Please ensure:
- Clean, well-documented code
- Proper error handling
- PEP 8 style compliance
- 4 spaces for indentation
- Spaces around operators
- Descriptive variable names
- Type hints where appropriate
- Professional 
Return only the code, no explanations."""
            else:
                prompt = f"""Generate a Python Program based on these requirements:
{requirements}
Please ensure:
- Clean, well-documented code
- Proper error handling
- PEP 8 style compliance
- Professional code structure
Return only the code, no explanations."""
            response = self.model.generate_content(prompt)
            generated_code = response.text.strip()
            self.apply_syntax_highlighting(generated_code)
            self.show_notification("Program generated successfully! Use Ctrl+S to save or Ctrl+C to copy.")
        except Exception as e:
            self.show_notification(f"Error generating program: {str(e)}")
            logging.error(f"Error in generate_program: {str(e)}")
        finally:
            if "GENERATE PROGRAM" in self.buttons:
                self.buttons["GENERATE PROGRAM"].configure(state="normal")
            if loading_label and loading_label.winfo_exists():
                loading_label.destroy()
    def apply_syntax_highlighting(self, code):
        """Apply syntax highlighting to the code"""
        self.code_display.delete("1.0", tk.END)
        import re
        keywords = ["def", "class", "import", "from", "return", "if", "else", "elif",
                   "try", "except", "finally", "for", "while", "in", "and", "or", 
                   "not", "is", "with", "as", "raise", "break", "continue", "pass",
                   "lambda", "assert", "del", "yield", "async", "await"]
        builtins = ["True", "False", "None", "print", "len", "range", "str", "int",
                    "float", "list", "dict", "set", "tuple", "sum", "min", "max",
                    "enumerate", "zip", "map", "filter", "any", "all", "round", "abs"]
        operators = ["+", "-", "*", "/", "//", "%", "**", "=", "==", "!=", "<", ">",
                    "<=", ">=", "+=", "-=", "*=", "/=", "&", "|", "^", "~", "<<", ">>"]
        special_methods = [f"__{name}__" for name in ["init", "str", "repr", "len", 
                         "getitem", "setitem", "delitem", "iter", "next", "enter",
                         "exit", "call", "add", "sub", "mul", "div", "mod", "pow"]]
        lines = code.split("\n")
        for line in lines:
            indent_match = re.match(r'^(\s*)', line)
            indent = indent_match.group(1) if indent_match else ''
            self.code_display.insert(tk.END, indent)
            if line.strip().startswith("#"):
                self.code_display.insert(tk.END, line.strip() + "\n", "comment")
                continue
            if line.strip().startswith("@"):
                self.code_display.insert(tk.END, line.strip() + "\n", "decorator")
                continue
            words = re.findall(r'[\w._]+|"[^"]*"|\'[^\']*\'|#.*|[^\w\s]|\s+', line.strip())
            i = 0
            while i < len(words):
                word = words[i]
                if word.isspace():
                    self.code_display.insert(tk.END, word)
                elif word in keywords:
                    self.code_display.insert(tk.END, word, "keyword")
                elif word in builtins:
                    self.code_display.insert(tk.END, word, "builtin")
                elif word in special_methods:
                    self.code_display.insert(tk.END, word, "special")
                elif word in operators:
                    self.code_display.insert(tk.END, word, "operator")
                elif word.startswith(("'", '"')):
                    self.code_display.insert(tk.END, word, "string")
                elif word.startswith("#"):
                    self.code_display.insert(tk.END, word, "comment")
                elif re.match(r'^[0-9]+(\.[0-9]+)?$', word):
                    self.code_display.insert(tk.END, word, "number")
                elif re.match(r'^[A-Z][a-zA-Z0-9]*$', word):
                    self.code_display.insert(tk.END, word, "class")
                elif re.match(r'^[a-z][a-zA-Z0-9_]*\($', word):
                    self.code_display.insert(tk.END, word, "function")
                elif re.match(r'^[A-Z][a-zA-Z0-9_]*(\[.*\])?$', word):
                    self.code_display.insert(tk.END, word, "type_hint")
                else:
                    self.code_display.insert(tk.END, word)
                i += 1
            self.code_display.insert(tk.END, "\n")
    def setup_bindings(self):
        """Setup keyboard shortcuts"""
        self.window.bind('<Return>', lambda e: self.generate_program())
        self.window.bind('<Control-l>', lambda e: self.clear_all())
        self.window.bind('<Control-c>', lambda e: self.copy_to_clipboard())
        self.window.bind('<Control-s>', lambda e: self.save_code())
    def clear_example_text(self, event):
        """Clear example text when input field is focused"""
        if self.requirements_text.get('1.0', tk.END).strip() == "Example: Create a Python program that reads a CSV file, processes the data, and generates a bar chart using matplotlib.":
            self.requirements_text.delete('1.0', tk.END)
    def on_button_hover(self, event, button):
        """Create glowing effect on hover"""
        button.configure(
            bg="#00F0FF",
            fg="#0C141F",
            bd=2,
            highlightbackground="#00F0FF",
            highlightthickness=2
        )
    def on_button_leave(self, event, button):
        """Remove glowing effect on leave"""
        button.configure(
            bg="#0C141F",
            fg="#00F0FF",
            bd=1,
            highlightbackground="#0C141F",
            highlightthickness=1
        )
    def copy_to_clipboard(self):
        """Copy generated code to clipboard"""
        code = self.code_display.get("1.0", tk.END).strip()
        if code:
            self.window.clipboard_clear()
            self.window.clipboard_append(code)
            self.show_notification("Code copied to clipboard!")
        else:
            self.show_notification("No code to copy!")
    def save_code(self, event=None):
        """Save generated code to file"""
        code = self.code_display.get("1.0", tk.END).strip()
        if not code:
            self.show_notification("No code to save!")
            return
        file_path = filedialog.asksaveasfilename(
            defaultextension=".py",
            filetypes=[("Python files", "*.py"), ("All files", "*.*")],
            title="Save Generated Code"
        )
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as file:
                    file.write(code)
                self.show_notification(f"Code saved successfully to {file_path}")
            except Exception as e:
                self.show_notification(f"Error saving file: {str(e)}")
                messagebox.showerror("Save Error", f"Failed to save file: {str(e)}")
    def clear_all(self):
        """Clear all text fields"""
        self.requirements_text.delete("1.0", tk.END)
        self.code_display.delete("1.0", tk.END)
        self.status_var.set("")
        self.requirements_text.insert('1.0', "Example: Create a Python program that reads a CSV file, processes the data, and generates a bar chart using matplotlib.")
    def show_notification(self, message):
        """Show a notification message in the status bar"""
        if hasattr(self, 'status_var'):
            self.status_var.set(message)
        else:
            messagebox.showinfo("Notification", message)
    def center_window(self):
        """Center the window on the screen"""
        self.window.update_idletasks()
        width = self.window.winfo_width()
        height = self.window.winfo_height()
        x = (self.window.winfo_screenwidth() // 2) - (width // 2)
        y = (self.window.winfo_screenheight() // 2) - (height // 2)
        self.window.geometry(f'{width}x{height}+{x}+{y}')
    def on_closing(self):
        """Handle window closing"""
        self.window.destroy()
class FileHandlingAssistant:
    def __init__(self, parent, assistant):
        self.parent = parent
        self.assistant = assistant
        self.window = None
        self.processed_files = []
        self.assistant_folder = os.path.join(get_app_directory(), "Assistant_Files")
        os.makedirs(self.assistant_folder, exist_ok=True)
        self.gemini_model = self.assistant.gemini_model
    def show_window(self):
        if self.window is not None and self.window.winfo_exists():
            self.window.lift()
            return
        self.window = ctk.CTkToplevel(self.parent)
        self.window.title("File Handling Assistant")
        self.window.geometry("1300x600")
        self.window.transient(self.parent)
        self.window.grab_set()
        self.status_var = tk.StringVar(value="Ready")
        theme = self.assistant.themes[self.assistant.current_theme]
        self.window.configure(fg_color=theme["bg"])
        self.main_frame = ctk.CTkFrame(
            self.window,
            fg_color=theme["bg"],
            corner_radius=10
        )
        self.main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        status_frame = ctk.CTkFrame(
            self.main_frame,
            fg_color=theme["highlight"],
            corner_radius=8,
            height=30
        )
        status_frame.pack(fill=tk.X, padx=10, pady=(0, 10))
        status_label = ctk.CTkLabel(
            status_frame,
            textvariable=self.status_var,
            font=("Segoe UI", 10),
            text_color=theme["fg"]
        )
        status_label.pack(side=tk.LEFT, padx=10, pady=5)
        title_label = ctk.CTkLabel(
            self.main_frame,
            text="File Handling Assistant",
            font=("Segoe UI", 24, "bold"),
            text_color=theme["accent"]
        )
        title_label.pack(pady=(0, 20))
        file_frame = ctk.CTkFrame(
            self.main_frame,
            fg_color=theme["highlight"],
            corner_radius=8
        )
        file_frame.pack(fill=tk.X, padx=10, pady=10)
        self.file_path_var = tk.StringVar()
        file_path_entry = ctk.CTkEntry(
            file_frame,
            textvariable=self.file_path_var,
            width=500,
            font=("Segoe UI", 12),
            placeholder_text="File path...",
            fg_color=theme["text_bg"],
            text_color=theme["fg"]
        )
        file_path_entry.pack(side=tk.LEFT, padx=10, pady=10, fill=tk.X, expand=True)
        browse_button = ctk.CTkButton(
            file_frame,
            text="Browse",
            font=("Segoe UI", 12),
            fg_color=theme["accent"],
            text_color="#ffffff",
            hover_color=theme["secondary"],
            command=self.browse_file
        )
        browse_button.pack(side=tk.RIGHT, padx=10, pady=10)
        commands_frame = ctk.CTkFrame(
            self.main_frame,
            fg_color=theme["highlight"],
            corner_radius=8
        )
        commands_frame.pack(fill=tk.X, padx=10, pady=10)
        command_label = ctk.CTkLabel(
            commands_frame,
            text="File Command:",
            font=("Segoe UI", 12, "bold"),
            text_color=theme["fg"]
        )
        command_label.pack(side=tk.LEFT, padx=10, pady=10)
        self.command_var = tk.StringVar()
        command_entry = ctk.CTkEntry(
            commands_frame,
            textvariable=self.command_var,
            width=400,
            font=("Segoe UI", 12),
            placeholder_text="Enter command (e.g., 'convert to pdf', 'summarize')",
            fg_color=theme["text_bg"],
            text_color=theme["fg"]
        )
        command_entry.pack(side=tk.LEFT, padx=10, pady=10, fill=tk.X, expand=True)
        voice_loop_button = ctk.CTkButton(
            commands_frame,
            text="Voice Mode",
            font=("Segoe UI", 12),
            width=100,
            fg_color=theme["success"],
            text_color="#ffffff",
            hover_color=theme["secondary"],
            command=self.voice_command_loop
        )
        voice_loop_button.pack(side=tk.RIGHT, padx=5, pady=10)
        voice_cmd_button = ctk.CTkButton(
            commands_frame,
            text="🎤",
            font=("Segoe UI", 16),
            width=40,
            fg_color=theme["accent"],
            text_color="#ffffff",
            hover_color=theme["secondary"],
            command=self.voice_command
        )
        voice_cmd_button.pack(side=tk.RIGHT, padx=5, pady=10)
        execute_button = ctk.CTkButton(
            commands_frame,
            text="Execute",
            font=("Segoe UI", 12),
            fg_color=theme["accent"],
            text_color="#ffffff",
            hover_color=theme["secondary"],
            command=self.process_file_command
        )
        execute_button.pack(side=tk.RIGHT, padx=10, pady=10)
        recent_files_button = ctk.CTkButton(
            commands_frame,
            text="Recent Files",
            font=("Segoe UI", 12),
            width=100,
            fg_color=theme["accent"],
            text_color="#ffffff",
            hover_color=theme["secondary"],
            command=self.show_recent_files_window
        )
        recent_files_button.pack(side=tk.LEFT, padx=10, pady=10)
        output_frame = ctk.CTkFrame(
            self.main_frame,
            fg_color=theme["highlight"],
            corner_radius=8
        )
        output_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        output_label = ctk.CTkLabel(
            output_frame,
            text="Output:",
            font=("Segoe UI", 12, "bold"),
            text_color=theme["fg"]
        )
        output_label.pack(anchor=tk.W, padx=10, pady=(10, 5))
        self.output_text = ctk.CTkTextbox(
            output_frame,
            font=("Segoe UI", 12),
            fg_color=theme["text_bg"],
            text_color=theme["fg"],
            corner_radius=6
        )
        self.output_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=(0, 10))
        recent_frame = ctk.CTkFrame(
            self.main_frame,
            fg_color=theme["highlight"],
            corner_radius=8
        )
        recent_frame.pack(fill=tk.X, padx=10, pady=10)
        recent_label = ctk.CTkLabel(
            recent_frame,
            text="Recent Files:",
            font=("Segoe UI", 12, "bold"),
            text_color=theme["fg"]
        )
        recent_label.pack(anchor=tk.W, padx=10, pady=(10, 5))
        self.recent_files_listbox = tk.Listbox(
            recent_frame,
            bg=theme["text_bg"],
            fg=theme["fg"],
            font=("Segoe UI", 11),
            height=4,
            selectbackground=theme["accent"]
        )
        def on_enter_key(event):
            self.process_file_command()
            return "break"
        self.window.bind("<Return>", on_enter_key)
        for widget in self.window.winfo_children():
            if isinstance(widget, ctk.CTkFrame):
                for child in widget.winfo_children():
                    if isinstance(child, ctk.CTkFrame):
                        for entry in child.winfo_children():
                            if isinstance(entry, ctk.CTkEntry) and entry.cget("placeholder_text") and "command" in entry.cget("placeholder_text").lower():
                                entry.bind("<Return>", on_enter_key)
        self.recent_files_listbox.pack(fill=tk.X, padx=10, pady=(0, 10))
        self.recent_files_listbox.bind("<Double-Button-1>", self.select_recent_file)
        self.load_recent_files()
        self.status_var.set("Select a file and enter a command, or try voice commands")
        self.update_output("Welcome to File Handling Assistant!\nYou can:\n- Browse for a file\n- Use voice commands\n- Try commands like 'convert to pdf', 'summarize this file', etc.\n- View recently processed files")
    def voice_command(self):
        """Use a separate thread for voice command input"""
        if hasattr(self.assistant, 'start_button'):
            self.was_main_listening = False
            if hasattr(self.assistant.start_button, '_listening'):
                self.was_main_listening = self.assistant.start_button._listening
            if hasattr(self.assistant.start_button, '_listening'):
                self.assistant.start_button._listening = False
            self.assistant.start_button.configure(text="▶")
            self.assistant.status_label.configure(text="● Not Listening", text_color='#ff4d4d')
            if hasattr(self.assistant, 'wave_vis'):
                self.assistant.wave_vis.stop_listening_animation()
            time.sleep(0.1)
        self._voice_command_active = True
        for widget in self.window.winfo_children():
            if isinstance(widget, ctk.CTkFrame):
                for child in widget.winfo_children():
                    if isinstance(child, ctk.CTkFrame):
                        for button in child.winfo_children():
                            if isinstance(button, ctk.CTkButton) and button.cget("text") == "🎤":
                                button.configure(state=DISABLED)
                                self.voice_button_ref = button
                                break
        self.update_output("Listening for voice command...")
        self.status_var.set("Listening...")
        def voice_listening_thread():
            try:
                r = sr.Recognizer()
                with sr.Microphone() as source:
                    r.adjust_for_ambient_noise(source, duration=0.5)
                    self.window.after(0, lambda: self.status_var.set("Listening..."))
                    try:
                        audio = r.listen(source, timeout=5, phrase_time_limit=10)
                        self.window.after(0, lambda: self.status_var.set("Processing voice..."))
                        try:
                            command = r.recognize_google(audio)
                            if command:
                                self.window.after(0, lambda cmd=command: 
                                    self.update_output(f"Voice command received: {cmd}"))
                                self.window.after(0, lambda cmd=command: self.command_var.set(cmd))
                                self.window.after(0, lambda: self.process_voice_command(
                                    self.command_var.get(), self.file_path_var.get()))
                        except sr.UnknownValueError:
                            self.window.after(0, lambda: self.update_output("Voice not recognized"))
                        except sr.RequestError as e:
                            self.window.after(0, lambda err=e: 
                                self.update_output(f"Speech service error: {err}"))
                    except sr.WaitTimeoutError:
                        self.window.after(0, lambda: self.update_output("No speech detected within timeout"))
            except Exception as e:
                self.window.after(0, lambda err=e: self.update_output(f"Voice recognition error: {err}"))
            finally:
                if hasattr(self, 'voice_button_ref') and self.voice_button_ref:
                    try:
                        if self.voice_button_ref.winfo_exists():
                            self.voice_button_ref.configure(state=NORMAL)
                    except Exception:
                        pass
                self._voice_command_active = False
                self.window.after(0, self.restore_main_listening)
        voice_thread = threading.Thread(target=voice_listening_thread)
        voice_thread.daemon = True
        voice_thread.start()
    def restore_main_listening(self):
        """Restore the main assistant's listening state if it was active before"""
        self.status_var.set("Ready")
        self._voice_command_active = False
        self._voice_loop_active = False
        if hasattr(self, 'was_main_listening') and self.was_main_listening:
            try:
                if hasattr(self.assistant, 'start_button'):
                    time.sleep(0.2)
                    self.assistant.start_button._listening = True
                    self.assistant.start_button.configure(text="■")
                    self.assistant.status_label.configure(text="● Listening", text_color='#4CAF50')
                    if hasattr(self.assistant, 'wave_vis'):
                        self.assistant.wave_vis.start_listening_animation()
                    logging.info("Restored main assistant listening state")
            except Exception as e:
                logging.error(f"Error restoring main listening: {e}")
        self.was_main_listening = False
    def voice_command_loop(self):
        """Use a separate thread for continuous voice command input"""
        if hasattr(self.assistant, 'start_button'):
            self.was_main_listening = False
            if hasattr(self.assistant.start_button, '_listening'):
                self.was_main_listening = self.assistant.start_button._listening
            if hasattr(self.assistant.start_button, '_listening'):
                self.assistant.start_button._listening = False
            self.assistant.start_button.configure(text="▶")
            self.assistant.status_label.configure(text="● Not Listening", text_color='#ff4d4d')
            if hasattr(self.assistant, 'wave_vis'):
                self.assistant.wave_vis.stop_listening_animation()
            time.sleep(0.2)
        self._voice_loop_active = True
        self.update_output("Voice command mode activated - speak commands")
        self.status_var.set("Voice mode active")
        for widget in self.window.winfo_children():
            if isinstance(widget, ctk.CTkFrame):
                for child in widget.winfo_children():
                    if isinstance(child, ctk.CTkFrame):
                        for button in child.winfo_children():
                            if isinstance(button, ctk.CTkButton) and button.cget("text") == "Voice Mode":
                                button.configure(text="Stop Voice Mode", fg_color="#FF5252")
                                self.voice_loop_button_ref = button
                                break
        def voice_listening_thread():
            try:
                r = sr.Recognizer()
                r.pause_threshold = 0.8
                r.dynamic_energy_threshold = True
                while getattr(self, '_voice_loop_active', False):
                    try:
                        with sr.Microphone() as source:
                            r.adjust_for_ambient_noise(source, duration=0.5)
                            self.window.after(0, lambda: self.status_var.set("Listening for command..."))
                            try:
                                audio = r.listen(source, timeout=3, phrase_time_limit=5)
                                self.window.after(0, lambda: self.status_var.set("Processing voice..."))
                                try:
                                    command = r.recognize_google(audio, language='en-in')
                                    if command.lower() in ["stop", "stop voice mode", "exit voice mode", "exit"]:
                                        self.window.after(0, self.stop_voice_command_loop)
                                        break
                                    self.window.after(0, lambda cmd=command: self.command_var.set(cmd))
                                    self.window.after(0, lambda cmd=command: self.update_output(f"Received voice command: {cmd}"))
                                    file_path = self.file_path_var.get()
                                    if file_path and os.path.exists(file_path):
                                        self.window.after(0, lambda cmd=command, path=file_path: 
                                                        self.process_voice_command(cmd, path))
                                    else:
                                        self.window.after(0, lambda: self.update_output("Please select a file first"))
                                    time.sleep(1)
                                except sr.UnknownValueError:
                                    pass
                                except sr.RequestError:
                                    self.window.after(0, lambda: self.update_output(
                                        "Speech recognition service unavailable. Please try again."))
                                    time.sleep(2)
                            except sr.WaitTimeoutError:
                                pass
                    except Exception as loop_e:
                        self.window.after(0, lambda err=loop_e: 
                            self.update_output(f"Voice loop error: {err}"))
                        time.sleep(1)
                    time.sleep(0.2)
            except Exception as e:
                self.window.after(0, lambda err=e: 
                    self.update_output(f"Voice recognition error: {err}"))
            finally:
                self.window.after(0, self.reset_voice_ui)
                self.window.after(100, self.restore_main_listening)
        self.voice_loop_thread = threading.Thread(target=voice_listening_thread)
        self.voice_loop_thread.daemon = True
        self.voice_loop_thread.start()
    def stop_voice_command_loop(self):
        """Stop the continuous voice command loop"""
        self._voice_loop_active = False
        self.update_output("Voice command mode deactivated")
        self.status_var.set("Ready")
        self.reset_voice_ui()
    def update_output(self, text):
        self.output_text.configure(state=NORMAL)
        self.output_text.insert(END, f"{text}\n")
        self.output_text.configure(state=NORMAL)
        self.output_text.see(END)
        self.status_var.set(text.split("\n")[0] if "\n" in text else text)
    def convert_file(self, file_path, command):
        """Convert a file from one format to another"""
        target_format = None
        valid_formats = {
            'pdf': 'pdf', 'adobe': 'pdf',
            'txt': 'txt', 'text': 'txt', 'note': 'txt',
            'html': 'html', 'web': 'html', 'webpage': 'html',
            'md': 'md', 'markdown': 'md',
            'csv': 'csv', 'excel': 'csv', 'spreadsheet': 'csv',
            'json': 'json', 'data': 'json',
            'png': 'png', 'jpg': 'jpg', 'jpeg': 'jpg', 'webp': 'webp',
            'bmp': 'bmp', 'gif': 'gif', 'tiff': 'tiff',
            'docx': 'docx', 'document': 'docx', 'doc': 'docx', 'word': 'docx',
            'rtf': 'rtf'
        }
        image_extensions = ['.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff', '.webp']
        format_words = []
        cmd_parts = command.lower().split()
        for i, word in enumerate(cmd_parts):
            if word == "to" and i < len(cmd_parts) - 1:
                if cmd_parts[i+1] in valid_formats:
                    format_words.append(cmd_parts[i+1])
                elif i < len(cmd_parts) - 2 and f"{cmd_parts[i+1]} {cmd_parts[i+2]}" in valid_formats:
                    format_words.append(f"{cmd_parts[i+1]} {cmd_parts[i+2]}")
            elif word == "as" and i < len(cmd_parts) - 1:
                if cmd_parts[i+1] in valid_formats:
                    format_words.append(cmd_parts[i+1])
            elif "format" in word and i > 0:
                if cmd_parts[i-1] in valid_formats:
                    format_words.append(cmd_parts[i-1])
            elif word in valid_formats:
                format_words.append(word)
        for format_word in format_words:
            if format_word in valid_formats:
                target_format = valid_formats[format_word]
                break
        if not target_format:
            common_words = ["convert", "to", "into", "as", "a", "the", "file", "document", "format", "image", "picture"]
            potential_formats = [word for word in cmd_parts if word not in common_words and len(word) > 1]
            for word in potential_formats:
                if word in valid_formats:
                    target_format = valid_formats[word]
                    break
                elif len(potential_formats) == 1:
                    target_format = word
                    self.window.after(0, lambda fmt=word: self.update_output(
                        f"Warning: '{fmt}' might not be a standard format. Trying anyway."))
                    break
        if not target_format:
            self.window.after(0, lambda: self.update_output(
                "Error: Could not determine target format from command. Please specify clearly (e.g., 'convert to png', 'save as pdf')."))
            return
        file_name = os.path.basename(file_path)
        file_base, file_ext = os.path.splitext(file_name)
        output_path = os.path.join(self.assistant_folder, f"{file_base}.{target_format}")
        is_image = file_ext.lower() in image_extensions
        if target_format == 'docx':
            self.window.after(0, lambda: self.update_output(f"Converting to DOCX (Word) format..."))
            try:
                if file_ext.lower() == '.txt' or file_ext.lower() == '.md' or file_ext.lower() == '.html':
                    try:
                        from docx import Document
                        document = Document()
                        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                            content = f.read()
                        for para in content.split('\n'):
                            if para.strip():
                                document.add_paragraph(para)
                        document.save(output_path)
                        self.window.after(0, lambda: self.update_output(f"File successfully converted to DOCX and saved to {output_path}"))
                        self.add_to_assistant_file(output_path, "converted")
                        return
                    except ImportError:
                        self.window.after(0, lambda: self.update_output("Error: python-docx library not found. Cannot convert to DOCX."))
                        return
                elif is_image:
                    try:
                        from docx import Document
                        import pytesseract
                        from PIL import Image
                        try:
                            img = Image.open(file_path)
                            if img.mode != 'L':
                                img = img.convert('L')
                            from PIL import ImageEnhance
                            enhancer = ImageEnhance.Contrast(img)
                            img = enhancer.enhance(2.0)
                            text = pytesseract.image_to_string(img)
                        except Exception as ocr_e:
                            self.window.after(0, lambda: self.update_output(f"Error during OCR: {str(ocr_e)}"))
                            text = "OCR failed. Image content could not be extracted."
                        document = Document()
                        document.add_heading(f"OCR Result: {file_name}", 0)
                        document.add_paragraph(text)
                        document.add_page_break()
                        document.add_heading("Original Image", 1)
                        try:
                            document.add_picture(file_path, width=4000000)
                        except:
                            document.add_paragraph("Could not embed original image")
                        document.save(output_path)
                        self.window.after(0, lambda: self.update_output(f"Image converted to DOCX with OCR at {output_path}"))
                        self.add_to_assistant_file(output_path, "converted")
                        return
                    except ImportError:
                        self.window.after(0, lambda: self.update_output("Error: Required libraries (python-docx/pytesseract) not found."))
                        return
                elif file_ext.lower() == '.pdf':
                    try:
                        from docx import Document
                        import PyPDF2
                        pdf_text = ""
                        with open(file_path, 'rb') as pdf_file:
                            pdf_reader = PyPDF2.PdfReader(pdf_file)
                            for page_num in range(len(pdf_reader.pages)):
                                page = pdf_reader.pages[page_num]
                                pdf_text += page.extract_text() + "\n\n"
                        document = Document()
                        document.add_heading(f"Converted from PDF: {file_name}", 0)
                        for para in pdf_text.split('\n'):
                            if para.strip():
                                document.add_paragraph(para)
                        document.save(output_path)
                        self.window.after(0, lambda: self.update_output(f"PDF converted to DOCX and saved to {output_path}"))
                        self.add_to_assistant_file(output_path, "converted")
                        return
                    except ImportError:
                        self.window.after(0, lambda: self.update_output("Error: Required libraries not found for PDF to DOCX conversion."))
                        return
                    except Exception as pdf_e:
                        self.window.after(0, lambda: self.update_output(f"Error converting PDF to DOCX: {str(pdf_e)}"))
                        return
                else:
                    self.window.after(0, lambda: self.update_output(f"Unsupported source format {file_ext} for DOCX conversion."))
                    return
            except Exception as docx_e:
                self.window.after(0, lambda: self.update_output(f"Error during DOCX conversion: {str(docx_e)}"))
                return
        try:
            if is_image:
                self.window.after(0, lambda: self.update_output(f"Attempting image conversion to {target_format}..."))
                from PIL import Image
                img = Image.open(file_path)
                quality = 90
                for i, word in enumerate(cmd_parts):
                    if word == "quality" and i < len(cmd_parts) - 1:
                        try:
                            q = int(cmd_parts[i+1])
                            if 1 <= q <= 100:
                                quality = q
                                self.window.after(0, lambda q=q: self.update_output(f"Using quality setting: {q}"))
                        except ValueError:
                            pass
                if target_format.lower() in ['jpg', 'jpeg']:
                    if img.mode in ('RGBA', 'LA', 'P'):
                        bg = Image.new("RGB", img.size, (255, 255, 255))
                        bg.paste(img, (0, 0), img if img.mode == 'RGBA' else None)
                        img = bg
                    img.save(output_path, quality=quality, optimize=True)
                elif target_format.lower() == 'png':
                    if img.mode == 'RGBA':
                        img.save(output_path, optimize=True)
                    else:
                        converted = img.convert('RGBA')
                        converted.save(output_path, optimize=True)
                elif target_format.lower() == 'webp':
                    img.save(output_path, quality=quality, method=4)
                elif target_format.lower() == 'bmp':
                    if img.mode in ('RGBA', 'LA'):
                        bg = Image.new("RGB", img.size, (255, 255, 255))
                        bg.paste(img, (0, 0), img if img.mode == 'RGBA' else None)
                        img = bg
                    img.save(output_path)
                elif target_format.lower() == 'gif':
                    if img.mode != 'P':
                        img = img.convert('RGBA')
                        img = img.quantize(method=2)
                    img.save(output_path, optimize=True)
                elif target_format.lower() == 'tiff':
                    img.save(output_path, compression='tiff_lzw')
                elif target_format.lower() == 'pdf':
                    try:
                        from reportlab.lib.pagesizes import letter
                        from reportlab.pdfgen import canvas
                        from io import BytesIO
                        if img.mode != 'RGB' and img.mode != 'RGBA':
                            img = img.convert('RGB')
                        width, height = img.size
                        aspect = height / width
                        page_width, page_height = letter
                        pdf_width = page_width * 0.9
                        pdf_height = pdf_width * aspect
                        if pdf_height > page_height * 0.9:
                            pdf_height = page_height * 0.9
                            pdf_width = pdf_height / aspect
                        img_bytes = BytesIO()
                        img.save(img_bytes, format='PNG')
                        img_bytes.seek(0)
                        c = canvas.Canvas(output_path, pagesize=letter)
                        x = (page_width - pdf_width) / 2
                        y = (page_height - pdf_height) / 2
                        c.drawImage(img_bytes, x, y, width=pdf_width, height=pdf_height)
                        c.save()
                    except ImportError:
                        if img.mode != 'RGB':
                            img = img.convert('RGB')
                        img.save(output_path, "PDF", resolution=300.0)
                else:
                    img.save(output_path)
                success_msg = f"Image successfully converted and saved to {output_path}"
                self.window.after(0, lambda msg=success_msg: self.update_output(msg))
                self.add_to_assistant_file(output_path, "converted")
            elif file_ext.lower() in ['.txt', '.md', '.html', '.csv', '.json'] or target_format in ['txt', 'md', 'html', 'csv', 'json']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                if target_format == 'pdf':
                    try:
                        try:
                            from reportlab.lib.pagesizes import letter
                            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Preformatted
                            from reportlab.lib.styles import getSampleStyleSheet
                            from reportlab.lib.enums import TA_LEFT
                            from reportlab.pdfbase import pdfmetrics
                            from reportlab.pdfbase.ttfonts import TTFont
                            doc = SimpleDocTemplate(output_path, pagesize=letter)
                            styles = getSampleStyleSheet()
                            try:
                                font_path = "DejaVuSans.ttf"
                                if os.path.exists(font_path):
                                    pdfmetrics.registerFont(TTFont('DejaVuSans', font_path))
                                    body_font = 'DejaVuSans'
                                else:
                                    try:
                                        pdfmetrics.registerFont(TTFont('Arial', 'arialuni.ttf'))
                                        body_font = 'Arial'
                                    except:
                                        body_font = 'Helvetica'
                            except Exception as font_e:
                                logging.warning(f"Font registration failed: {font_e}, using Helvetica.")
                                body_font = 'Helvetica'
                            normal_style = styles['Normal']
                            normal_style.fontName = body_font
                            normal_style.fontSize = 10
                            normal_style.leading = 12
                            normal_style.alignment = TA_LEFT
                            code_style = styles['Code']
                            code_style.fontName = 'Courier'
                            code_style.fontSize = 9
                            code_style.leading = 11
                            code_style.alignment = TA_LEFT
                            story = []
                            lines = content.splitlines()
                            preformatted_block = []
                            for line in lines:
                                is_code_like = line.strip().startswith((' ', '\t', '#', '//', '/*', 'def ', 'class ', 'import '))
                                try:
                                    line_processed = line.encode('utf-8', 'replace').decode('utf-8', 'replace')
                                    line_processed = line_processed.replace('\t', '    ')
                                except Exception:
                                    line_processed = "".join(c if ord(c) < 128 else '?' for c in line)
                                if is_code_like:
                                    if preformatted_block:
                                        story.append(Paragraph("<br/>".join(preformatted_block), normal_style))
                                        story.append(Spacer(1, 6))
                                        preformatted_block = []
                                    story.append(Preformatted(line_processed, code_style))
                                    story.append(Spacer(1, 0))
                                else:
                                    preformatted_block.append(line_processed)
                            if preformatted_block:
                                story.append(Paragraph("<br/>".join(preformatted_block), normal_style))
                            doc.build(story)
                            success_msg = f"File converted to PDF (ReportLab) and saved to {output_path}"
                            self.window.after(0, lambda msg=success_msg: self.update_output(msg))
                        except ImportError:
                            from fpdf import FPDF
                            pdf = FPDF()
                            pdf.add_page()
                            try:
                                pdf.add_font('DejaVu', '', 'DejaVuSans.ttf', uni=True)
                                pdf.set_font('DejaVu', '', 12)
                            except RuntimeError:
                                pdf.set_font("Arial", size=12)
                            filtered_content = "".join(char if ord(char) < 256 else '?' for char in content)
                            pdf.multi_cell(0, 10, filtered_content)
                            pdf.output(output_path)
                            success_msg = f"File converted to PDF (FPDF) and saved to {output_path}"
                            self.window.after(0, lambda msg=success_msg: self.update_output(msg))
                    except ImportError:
                        error_msg = "Error: PDF libraries (ReportLab/FPDF) not installed. Please install with 'pip install fpdf reportlab'"
                        self.window.after(0, lambda msg=error_msg: self.update_output(msg))
                        return
                    except Exception as pdf_e:
                        error_msg = f"Error during PDF generation: {str(pdf_e)}"
                        self.window.after(0, lambda msg=error_msg: self.update_output(msg))
                        return
                else:
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(content)
                    success_msg = f"File converted and saved to {output_path}"
                    self.window.after(0, lambda msg=success_msg: self.update_output(msg))
                self.add_to_assistant_file(output_path, "converted")
            elif target_format.lower() == 'pdf' and file_ext.lower() == '.docx':
                self.window.after(0, lambda: self.update_output("Converting DOCX to PDF..."))
                def docx_to_pdf_worker():
                    conversion_success = False
                    error_msg = ""
                    try:
                        from docx2pdf import convert
                        convert(file_path, output_path)
                        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                            self.window.after(0, lambda: self.update_output("DOCX converted to PDF using docx2pdf"))
                            return True
                    except ImportError:
                        error_msg = "docx2pdf not available"
                    except Exception as e:
                        error_msg = f"docx2pdf error: {str(e)}"
                    try:
                        from docx import Document
                        from reportlab.lib.pagesizes import letter
                        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                        from reportlab.lib.styles import getSampleStyleSheet
                        doc = Document(file_path)
                        pdf = SimpleDocTemplate(output_path, pagesize=letter)
                        styles = getSampleStyleSheet()
                        story = []
                        for para in doc.paragraphs:
                            if para.text.strip():
                                style_name = 'Normal'
                                if para.style.name.startswith('Heading'):
                                    if '1' in para.style.name:
                                        style_name = 'Heading1'
                                    elif '2' in para.style.name:
                                        style_name = 'Heading2'
                                    else:
                                        style_name = 'Heading3'
                                clean_text = para.text.encode('ascii', 'replace').decode('ascii')
                                p = Paragraph(clean_text, styles[style_name])
                                story.append(p)
                                story.append(Spacer(1, 6))
                        if story:
                            pdf.build(story)
                            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                                self.window.after(0, lambda: self.update_output("DOCX converted to PDF using ReportLab"))
                                return True
                        else:
                            error_msg = "Document has no text content"
                    except ImportError:
                        error_msg = f"{error_msg}, ReportLab not available"
                    except Exception as e:
                        error_msg = f"{error_msg}, ReportLab error: {str(e)}"
                    try:
                        from fpdf import FPDF
                        from docx import Document
                        doc = Document(file_path)
                        text_content = ""
                        for para in doc.paragraphs:
                            text_content += para.text + "\n\n"
                        pdf = FPDF()
                        pdf.add_page()
                        pdf.set_font("Arial", size=12)
                        safe_text = "".join(c if ord(c) < 128 else ' ' for c in text_content)
                        chunks = [safe_text[i:i+80] for i in range(0, len(safe_text), 80)]
                        for chunk in chunks:
                            if chunk.strip():
                                pdf.multi_cell(0, 10, chunk)
                        pdf.output(output_path)
                        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                            self.window.after(0, lambda: self.update_output("DOCX converted to PDF using FPDF"))
                            return True
                    except ImportError:
                        error_msg = f"{error_msg}, FPDF not available"
                    except Exception as e:
                        error_msg = f"{error_msg}, FPDF error: {str(e)}"
                    self.window.after(0, lambda: self.update_output(f"All PDF conversion methods failed: {error_msg}"))
                    return False
                conversion_thread = threading.Thread(target=docx_to_pdf_worker)
                conversion_thread.daemon = True
                conversion_thread.start()
                conversion_thread.join(timeout=30)
                if conversion_thread.is_alive():
                    self.window.after(0, lambda: self.update_output("Conversion process taking too long, terminating"))
                    return
                if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                    self.window.after(0, lambda: self.update_output(f"File successfully converted and saved to {output_path}"))
                    self.add_to_assistant_file(output_path, "converted")
                else:
                    self.window.after(0, lambda: self.update_output("Conversion failed or produced empty file"))
        except Exception as e:
            error_msg = f"Error converting file: {str(e)}"
            logging.error(f"Conversion Error: {error_msg} for file {file_path} to {target_format}")
            self.window.after(0, lambda msg=error_msg: self.update_output(msg))
    def add_to_assistant_file(self, file_path, operation_type):
        """Add file to the assistant's database with metadata"""
        try:
            db_path = os.path.join(self.assistant_folder, "files_database.db")
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS processed_files (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    file_path TEXT,
                    file_name TEXT,
                    operation_type TEXT,
                    timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            file_name = os.path.basename(file_path)
            cursor.execute(
                "INSERT INTO processed_files (file_path, file_name, operation_type) VALUES (?, ?, ?)",
                (file_path, file_name, operation_type)
            )
            conn.commit()
            conn.close()
        except Exception as e:
            logging.error(f"Error adding file to database: {str(e)}")
    def summarize_file(self, file_path, command=None):
        file_name = os.path.basename(file_path)
        file_base, file_ext = os.path.splitext(file_name)
        export_format = None
        if command:
             match = re.search(r'(?:export|save)\s+(?:it|result|summary)\s+as\s+(\w+)', command.lower())
             if match:
                 export_format = match.group(1).lower()
                 self.window.after(0, lambda: self.update_output(f"Detected export request for summary: {export_format}"))
        output_base_path = os.path.join(self.assistant_folder, f"{file_base}_summary")
        content = ""
        try:
            self.window.after(0, lambda: self.update_output(f"Reading file for summarization: {file_name}"))
            file_ext_lower = file_ext.lower()
            if file_ext_lower in ['.txt', '.md', '.csv', '.json', '.html', '.xml', '.log']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                self.window.after(0, lambda: self.update_output("Text file read successfully."))
            elif file_ext_lower == '.docx':
                content = self._extract_text_from_docx(file_path)
                if not content or content.startswith("[Error"):
                    self.window.after(0, lambda: self.update_output("Failed to extract text from DOCX file for summary."))
                    return
                self.window.after(0, lambda: self.update_output("DOCX file processed for summary."))
            elif file_ext_lower == '.pdf':
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as pdf_file:
                        pdf_reader = PyPDF2.PdfReader(pdf_file)
                        for page_num in range(len(pdf_reader.pages)):
                            page = pdf_reader.pages[page_num]
                            content += page.extract_text() + "\n"
                    if not content.strip():
                         self.window.after(0, lambda: self.update_output("PDF text extraction yielded no content, attempting OCR..."))
                         content = self._extract_text_from_pdf_with_ocr(file_path)
                    self.window.after(0, lambda: self.update_output("PDF processed for summary."))
                except ImportError:
                     self.window.after(0, lambda: self.update_output("PyPDF2 not found for PDF summary. Install it."))
                     return
                except Exception as pdf_err:
                     self.window.after(0, lambda: self.update_output(f"Error reading PDF for summary: {pdf_err}"))
                     return
            elif file_ext_lower in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.gif', '.webp']:
                 content = self._extract_text_from_image(file_path)
                 self.window.after(0, lambda: self.update_output("Image OCR completed for summary."))
            else:
                 self.window.after(0, lambda: self.update_output(f"Attempting to read unsupported format {file_ext_lower} as text for summary."))
                 try:
                     with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                         content = f.read()
                 except Exception as read_err:
                     self.window.after(0, lambda: self.update_output(f"Could not read file as text: {read_err}"))
                     return
            if not content or not content.strip():
                 self.window.after(0, lambda: self.update_output("No content found in the file to summarize."))
                 return
            self.window.after(0, lambda: self.update_output("Summarizing file using AI..."))
            prompt = f"Summarize the following content concisely:\n\n{content[:15000]}"
            try:
                response = self.gemini_model.generate_content(prompt)
                summary = response.text.strip()
                if not summary:
                     self.window.after(0, lambda: self.update_output("AI returned an empty summary."))
                     return
                if export_format:
                    output_path = self._export_result(summary, output_base_path, export_format)
                    if output_path:
                        success_msg = f"Summary Preview:\n{summary[:500]}..."
                        self.window.after(0, lambda msg=success_msg: self.update_output(msg))
                else:
                    output_path_txt = output_base_path + ".txt"
                    with open(output_path_txt, 'w', encoding='utf-8') as f:
                        f.write(summary)
                    success_msg1 = f"File summarized and saved to {output_path_txt}"
                    success_msg2 = f"Summary Preview:\n{summary[:500]}..."
                    self.window.after(0, lambda msg=success_msg1: self.update_output(msg))
                    self.window.after(0, lambda msg=success_msg2: self.update_output(msg))
                    self.add_to_assistant_file(output_path_txt, "summary")
            except Exception as e:
                if "Content Policy Violation" in str(e):
                     error_msg = f"AI Error: Could not summarize due to content policy restrictions."
                elif "quota" in str(e).lower():
                     error_msg = f"AI Error: API quota likely exceeded."
                else:
                    error_msg = f"Error using AI to summarize: {str(e)}"
                self.window.after(0, lambda msg=error_msg: self.update_output(msg))
        except FileNotFoundError:
             self.window.after(0, lambda: self.update_output(f"Error: File not found at {file_path}"))
        except Exception as e:
            error_msg = f"Error reading file for summary: {str(e)}"
            self.window.after(0, lambda msg=error_msg: self.update_output(msg))
    def move_file(self, file_path, command):
        target_folder = None
        words = command.split()
        for i, word in enumerate(words):
            if word == "to" and i < len(words) - 1:
                target_folder = words[i+1]
                break
        if not target_folder:
            self.window.after(0, lambda: self.update_output("Error: Could not determine target folder"))
            return
        if target_folder.lower() == "assistant":
            target_folder = self.assistant_folder
        else:
            home_dir = os.path.expanduser("~")
            common_locations = [
                os.path.join(home_dir, "Desktop"),
                os.path.join(home_dir, "Documents"),
                os.path.join(home_dir, "Downloads"),
                os.path.join(home_dir, target_folder)
            ]
            for location in common_locations:
                if os.path.exists(location) and os.path.isdir(location):
                    target_folder = location
                    break
            if not os.path.exists(target_folder):
                self.window.after(0, lambda: self.update_output(f"Error: Target folder '{target_folder}' not found"))
                return
        file_name = os.path.basename(file_path)
        target_path = os.path.join(target_folder, file_name)
        try:
            shutil.move(file_path, target_path)
            self.window.after(0, lambda: self.update_output(f"File moved to {target_path}"))
        except Exception as e:
            self.window.after(0, lambda: self.update_output(f"Error moving file: {str(e)}"))
    def copy_file(self, file_path, command):
        target_folder = self.assistant_folder
        words = command.split()
        for i, word in enumerate(words):
            if word == "to" and i < len(words) - 1:
                folder_name = words[i+1]
                if folder_name.lower() != "assistant":
                    home_dir = os.path.expanduser("~")
                    common_locations = [
                        os.path.join(home_dir, "Desktop"),
                        os.path.join(home_dir, "Documents"),
                        os.path.join(home_dir, "Downloads"),
                        os.path.join(home_dir, folder_name)
                    ]
                    for location in common_locations:
                        if os.path.exists(location) and os.path.isdir(location):
                            target_folder = location
                            break
                    if not os.path.exists(target_folder):
                        self.window.after(0, lambda: self.update_output(f"Folder '{folder_name}' not found. Using Assistant folder instead."))
        file_name = os.path.basename(file_path)
        target_path = os.path.join(target_folder, file_name)
        try:
            shutil.copy2(file_path, target_path)
            self.window.after(0, lambda: self.update_output(f"File copied to {target_path}"))
        except Exception as e:
            self.window.after(0, lambda: self.update_output(f"Error copying file: {str(e)}"))
    def extract_from_file(self, file_path, command):
        """Extract specific information from a file based on the command"""
        extract_what = None
        export_format = None
        try:
            if "extract" in command.lower():
                extract_index = command.lower().find("extract")
                remaining_text = command.lower()[extract_index + 7:].strip()
                for phrase in ["export as", "save as", "export it as", "save it as"]:
                    if phrase in remaining_text:
                        remaining_text = remaining_text.split(phrase)[0].strip()
                extract_what = remaining_text
            export_match = re.search(r'(?:export|save)\s+(?:it|result|extraction)\s+as\s+(\w+)', command.lower())
            if export_match:
                export_format = export_match.group(1).lower()
        except Exception as e:
            self.window.after(0, lambda: self.update_output(f"Error parsing command: {str(e)}"))
            extract_what = self._determine_extraction_target_locally(command)
        if not extract_what:
            self.window.after(0, lambda: self.update_output("Determining extraction target..."))
            extract_what = self._determine_extraction_target_locally(command)
            if not extract_what:
                try:
                    self.window.after(0, lambda: self.update_output("Using Gemini to interpret your request..."))
                    extract_what = self._use_gemini_to_interpret_extraction(command)
                except Exception as e:
                    self.window.after(0, lambda: self.update_output(f"Error connecting to AI service: {str(e)}"))
                    extract_what = command.lower().replace("extract", "").strip()
                    self.window.after(0, lambda: self.update_output(f"Fallback extraction target: {extract_what}"))
        file_ext = os.path.splitext(file_path)[1].lower()
        content = ""
        extraction_source = "text"
        try:
            self.window.after(0, lambda: self.update_output(f"Reading file: {os.path.basename(file_path)}"))
            if not os.path.exists(file_path):
                self.window.after(0, lambda: self.update_output(f"Error: File not found at {file_path}"))
                return
            if file_ext == '.pdf':
                try:
                    self.window.after(0, lambda: self.update_output("Extracting text from PDF..."))
                    content = self._extract_from_pdf_enhanced(file_path)
                    content_preview = content[:100].replace('\n', ' ')
                    self.window.after(0, lambda: self.update_output(f"Initial PDF extraction result: {len(content)} characters"))
                    if not content.strip() or len(content.strip()) < 100:
                        self.window.after(0, lambda: self.update_output("PDF may be scanned or has little text, using OCR..."))
                        content = self._extract_text_from_pdf_with_ocr(file_path)
                        self.window.after(0, lambda: self.update_output(f"OCR extraction result: {len(content)} characters"))
                except Exception as e:
                    self.window.after(0, lambda: self.update_output(f"PDF extraction error: {str(e)}. Trying OCR..."))
                    try:
                        content = self._extract_text_from_pdf_with_ocr(file_path)
                        self.window.after(0, lambda: self.update_output(f"OCR fallback extraction: {len(content)} characters"))
                    except Exception as ocr_e:
                        self.window.after(0, lambda: self.update_output(f"OCR extraction failed: {str(ocr_e)}"))
                        content = ""
            if file_ext == '.pdf':
                self.window.after(0, lambda: self.update_output("Extracting text from PDF..."))
                content = self._extract_from_pdf_enhanced(file_path)
                self.window.after(0, lambda: self.update_output(f"PDF extraction result: {len(content)} characters"))
            elif file_ext == '.docx':
                self.window.after(0, lambda: self.update_output("Extracting text from DOCX document..."))
                content = self._extract_text_from_docx(file_path)
            elif file_ext in ['.txt', '.md', '.csv', '.json', '.html', '.xml', '.log']:
                content = self._extract_text_from_txt(file_path)
            else:
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                    self.window.after(0, lambda: self.update_output(f"Generic file extraction: {len(content)} characters"))
                except:
                    self.window.after(0, lambda: self.update_output("File can't be read as text, trying OCR as fallback..."))
                    try:
                        content = self._enhanced_image_ocr(file_path)
                        extraction_source = "image"
                        self.window.after(0, lambda: self.update_output(f"OCR fallback extraction: {len(content)} characters"))
                    except Exception as ocr_e:
                        self.window.after(0, lambda: self.update_output(f"OCR fallback failed: {str(ocr_e)}"))
                        content = ""
            if not content or not content.strip():
                self.window.after(0, lambda: self.update_output("Error: Could not extract any content from the file"))
                return
            self.window.after(0, lambda: self.update_output(f"Extracting '{extract_what}' using AI..."))
            extracted = self._extract_with_ai_and_timeout(content, extract_what, extraction_source, timeout=60)
            if not extracted or extracted == "EXTRACTION_FAILED":
                self.window.after(0, lambda: self.update_output("AI extraction failed, using local methods..."))
                extracted = self._extract_locally(content, extract_what)
            if not extracted or "No matching information found" in extracted:
                self.window.after(0, lambda: self.update_output(
                    f"Could not find information matching '{extract_what}'. Try different search terms."))
                return
            file_name = os.path.basename(file_path)
            file_base, _ = os.path.splitext(file_name)
            clean_extract_what = re.sub(r'[^a-zA-Z0-9_]', '_', extract_what).strip('_')
            output_base_path = os.path.join(self.assistant_folder, f"{file_base}_extracted_{clean_extract_what[:20]}")
            if export_format:
                output_path = self._export_result(extracted, output_base_path, export_format)
                if output_path:
                    self.window.after(0, lambda: self.update_output(f"Extracted content exported as {export_format}"))
                    self.window.after(0, lambda: self.update_output(f"Preview:\n{extracted[:500]}..."))
            else:
                output_path_txt = output_base_path + ".txt"
                with open(output_path_txt, 'w', encoding='utf-8') as f:
                    f.write(extracted)
                self.window.after(0, lambda: self.update_output(f"Extraction saved to {output_path_txt}"))
                self.window.after(0, lambda: self.update_output(f"Preview:\n{extracted[:500]}..."))
                self.add_to_assistant_file(output_path_txt, "extract")
        except Exception as e:
            self.window.after(0, lambda: self.update_output(f"Extraction error: {str(e)}"))
    def _extract_with_ai_and_timeout(self, content, extract_what, extraction_source="text", timeout=60):
        """Extract information using AI with a timeout to prevent hanging"""
        result = [""]
        extraction_complete = threading.Event()
        def extraction_worker():
            try:
                prompt = f"""Task: Find and extract information related to: '{extract_what}'
                From this {extraction_source} content:
                ---
                {content[:1000000]}
                ---
                FORMAT YOUR RESPONSE LIKE THIS:
                [EXTRACTED_CONTENT_START]
                (only the extracted information goes here, nothing else)
                [EXTRACTED_CONTENT_END]
                IMPORTANT INSTRUCTIONS:
                1. Extract information specifically about: '{extract_what}'
                2. If the exact term isn't found, look for SIMILAR or RELATED information
                3. Include context around the extracted information for better understanding
                4. Be thorough - extract ALL relevant information, even if it appears in different parts
                5. Format appropriately (tables as tables, lists as lists, etc.)
                6. If no matching or related information exists, write ONLY: "No matching information found"
                7.You should give complete extraction text .
                """
                response = self.gemini_model.generate_content(
                    prompt
                )
                extracted_text = response.text.strip()
                match = re.search(r'\[EXTRACTED_CONTENT_START\](.*?)\[EXTRACTED_CONTENT_END\]', 
                                extracted_text, re.DOTALL)
                if match:
                    result[0] = match.group(1).strip()
                else:
                    result[0] = extracted_text
                extraction_complete.set()
            except Exception as e:
                self.window.after(0, lambda: self.update_output(f"AI extraction error: {str(e)}"))
                result[0] = "EXTRACTION_FAILED"
                extraction_complete.set()
        extraction_thread = threading.Thread(target=extraction_worker)
        extraction_thread.daemon = True
        extraction_thread.start()
        wait_start = time.time()
        for i in range(timeout):
            if extraction_complete.wait(1):
                break
            if i % 5 == 0:
                elapsed = time.time() - wait_start
                self.window.after(0, lambda e=elapsed: self.update_output(f"Extraction in progress... ({int(e)}s elapsed)"))
        if not extraction_complete.is_set():
            self.window.after(0, lambda: self.update_output(f"AI extraction timed out after {timeout} seconds"))
            return "EXTRACTION_FAILED"
        return result[0]
    def _determine_extraction_target_locally(self, command):
        """Parse the command locally to determine what to extract without using AI"""
        command_lower = command.lower()
        patterns = [
            r"extract\s+(all\s+)?(the\s+)?([a-z0-9_\s]+?)(?:\s+from|\s+in|\s+on|$)",
            r"get\s+(all\s+)?(the\s+)?([a-z0-9_\s]+?)(?:\s+from|\s+in|\s+on|$)",
            r"pull\s+(all\s+)?(the\s+)?([a-z0-9_\s]+?)(?:\s+from|\s+in|\s+on|$)",
            r"find\s+(all\s+)?(the\s+)?([a-z0-9_\s]+?)(?:\s+from|\s+in|\s+on|$)"
        ]
        for pattern in patterns:
            match = re.search(pattern, command_lower)
            if match:
                return match.group(3).strip()
        common_words = ["extract", "the", "all", "from", "in", "this", "file", "on", "get", "find", "pull"]
        result = command_lower
        for word in common_words:
            result = result.replace(word, "")
        return result.strip()
    def _extract_locally(self, content, extract_what):
        """Extract information locally without using AI"""
        self.window.after(0, lambda: self.update_output(f"Performing local extraction for '{extract_what}'"))
        search_terms = extract_what.lower().split()
        lines = content.split('\n')
        relevant_sections = []
        current_section = []
        in_relevant_section = False
        section_end_count = 0
        for i, line in enumerate(lines):
            line_lower = line.lower()
            if any(term in line_lower for term in search_terms):
                if not in_relevant_section:
                    in_relevant_section = True
                    current_section = []
                    for j in range(max(0, i-2), i):
                        current_section.append(lines[j])
                current_section.append(line)
                section_end_count = 0
            elif in_relevant_section:
                current_section.append(line)
                section_end_count += 1
                if section_end_count >= 5:
                    relevant_sections.append('\n'.join(current_section))
                    current_section = []
                    in_relevant_section = False
        if current_section:
            relevant_sections.append('\n'.join(current_section))
        if not relevant_sections:
            if any(term in extract_what.lower() for term in ['path', 'os.path', 'directory', 'folder', 'file']):
                path_patterns = [
                    r'os\.path\.join\([^)]+\)',
                    r'os\.path\.[a-z]+\([^)]+\)', 
                    r'Path\([^)]+\)',
                    r'["\']/[^"\']+["\']', 
                    r'["\']\\[^"\']+["\']',
                    r'[a-zA-Z]:\\[^"\']+',
                ]
                results = []
                for pattern in path_patterns:
                    matches = re.findall(pattern, content)
                    results.extend(matches)
                if results:
                    return "Found path-related code:\n\n" + "\n".join(results)
        if relevant_sections:
            return "\n\n---\n\n".join(relevant_sections)
        else:
            return "No matching information found for: " + extract_what
    def _use_gemini_to_interpret_extraction(self, command):
        """Use Gemini to interpret what needs to be extracted, with timeout"""
        result = [""]
        interpretation_complete = threading.Event()
        def interpretation_worker():
            try:
                prompt = f"""
                Analyze this file processing command: "{command}"
                What specific information is the user trying to extract from the file?
                Return only the information to be extracted, nothing else.
                For example, if the user says "extract phone numbers from this file",
                return only "phone numbers".
                """
                response = self.gemini_model.generate_content(prompt)
                result[0] = response.text.strip()
                interpretation_complete.set()
            except Exception as e:
                result[0] = ""
                interpretation_complete.set()
        interpretation_thread = threading.Thread(target=interpretation_worker)
        interpretation_thread.daemon = True
        interpretation_thread.start()
        if not interpretation_complete.wait(10):
            self.window.after(0, lambda: self.update_output("Interpretation timed out, using fallback"))
            return self._determine_extraction_target_locally(command)
        if result[0]:
            self.window.after(0, lambda: self.update_output(f"Interpreted extraction target: {result[0]}"))
            return result[0]
        else:
            return self._determine_extraction_target_locally(command)
    def _extract_text_from_image(self, image_path):
        """Extract text from an image using OCR"""
        try:
            import pytesseract
            from PIL import Image
            try:
                pytesseract.get_tesseract_version()
            except Exception as e:
                self.window.after(0, lambda: self.update_output(
                    "Tesseract OCR not found. Please install Tesseract OCR and ensure it's in your PATH."))
                raise Exception("Tesseract OCR not available")
            img = Image.open(image_path)
            max_dimension = 3000
            if max(img.size) > max_dimension:
                ratio = max_dimension / max(img.size)
                new_size = (int(img.size[0] * ratio), int(img.size[1] * ratio))
                img = img.resize(new_size, Image.LANCZOS)
            if img.mode != 'L':
                img = img.convert('L')
            from PIL import ImageEnhance
            enhancer = ImageEnhance.Contrast(img)
            img = enhancer.enhance(2.0)
            text = pytesseract.image_to_string(img)
            return text
        except ImportError:
            self.window.after(0, lambda: self.update_output(
                "Error: Missing libraries for OCR. Install pytesseract."))
            raise Exception("OCR libraries not available")
    def handle_custom_command(self, file_path, command):
        """Handle custom commands using Gemini AI, with export option"""
        export_format = None
        export_match = re.search(r'(?:export|save)\s+(?:it|result)\s+as\s+(\w+)', command.lower())
        if export_match:
            export_format = export_match.group(1).lower()
            self.window.after(0, lambda: self.update_output(f"Detected export request for custom command result: {export_format}"))
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            self.window.after(0, lambda: self.update_output(f"Processing custom command '{command}' using AI..."))
            prompt = f"""
            Act as a file processing assistant. Execute the following command on the provided file content:
            Command: "{command}"
            File Content:
            ---
            {content[:15000]}
            ---
            Provide only the result of the command execution. If the command is unclear or cannot be performed on the text, state that clearly.
            Result:"""
            response = self.gemini_model.generate_content(prompt)
            result = response.text.strip()
            if not result:
                 self.window.after(0, lambda: self.update_output("AI returned an empty result for the custom command."))
                 return
            file_name = os.path.basename(file_path)
            file_base, file_ext = os.path.splitext(file_name)
            clean_command = re.sub(r'[^a-zA-Z0-9_]', '_', command).strip('_')
            output_base_path = os.path.join(self.assistant_folder, f"{file_base}_custom_{clean_command[:20]}")
            if export_format:
                output_path = self._export_result(result, output_base_path, export_format)
                if output_path:
                    success_msg = f"Result Preview:\n{result[:500]}..."
                    self.window.after(0, lambda msg=success_msg: self.update_output(msg))
            else:
                output_path_txt = output_base_path + ".txt"
                with open(output_path_txt, 'w', encoding='utf-8') as f:
                    f.write(result)
                success_msg1 = f"Custom command processed and saved to {output_path_txt}"
                success_msg2 = f"Result Preview:\n{result[:500]}..."
                self.window.after(0, lambda msg=success_msg1: self.update_output(msg))
                self.window.after(0, lambda msg=success_msg2: self.update_output(msg))
                self.add_to_assistant_file(output_path_txt, "custom_command")
        except FileNotFoundError:
             self.window.after(0, lambda: self.update_output(f"Error: File not found at {file_path}"))
        except Exception as e:
            if "Content Policy Violation" in str(e):
                error_msg = f"AI Error: Could not process custom command due to content policy."
            elif "quota" in str(e).lower():
                 error_msg = f"AI Error: API quota likely exceeded."
            else:
                 error_msg = f"Error processing custom command: {str(e)}"
            self.window.after(0, lambda msg=error_msg: self.update_output(msg))
    def _export_result(self, content, base_path, format_type):
        """Export the extracted content in the specified format"""
        try:
            format_type = format_type.lower()
            self.window.after(0, lambda: self.update_output(f"Exporting as {format_type}..."))
            if format_type == "txt" or format_type == "text":
                output_path = base_path + ".txt"
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                self.window.after(0, lambda: self.update_output(f"Exported to text file: {output_path}"))
                self.add_to_assistant_file(output_path, "extract")
                return output_path
            elif format_type == "json":
                try:
                    if content.strip().startswith('{') and content.strip().endswith('}'):
                        import json
                        parsed = json.loads(content)
                        formatted_content = json.dumps(parsed, indent=2)
                    else:
                        formatted_content = json.dumps({"extracted_content": content}, indent=2)
                except:
                    import json
                    formatted_content = json.dumps({"extracted_content": content}, indent=2)
                output_path = base_path + ".json"
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(formatted_content)
                self.window.after(0, lambda: self.update_output(f"Exported to JSON file: {output_path}"))
                self.add_to_assistant_file(output_path, "extract")
                return output_path
            elif format_type == "csv":
                output_path = base_path + ".csv"
                lines = content.strip().split('\n')
                if len(lines) <= 1 or ',' not in content:
                    with open(output_path, 'w', encoding='utf-8', newline='') as f:
                        import csv
                        writer = csv.writer(f)
                        writer.writerow(["extracted_content"])
                        writer.writerow([content])
                else:
                    with open(output_path, 'w', encoding='utf-8', newline='') as f:
                        f.write(content)
                self.window.after(0, lambda: self.update_output(f"Exported to CSV file: {output_path}"))
                self.add_to_assistant_file(output_path, "extract")
                return output_path
            elif format_type == "docx" or format_type == "word":
                try:
                    from docx import Document
                    document = Document()
                    document.add_paragraph(content)
                    output_path = base_path + ".docx"
                    document.save(output_path)
                    self.window.after(0, lambda: self.update_output(f"Exported to Word document: {output_path}"))
                    self.add_to_assistant_file(output_path, "extract")
                    return output_path
                except ImportError:
                    self.window.after(0, lambda: self.update_output("Error: python-docx library not found. Saving as text instead."))
                    output_path = base_path + ".txt"
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(content)
                    self.window.after(0, lambda: self.update_output(f"Exported to text file (fallback): {output_path}"))
                    self.add_to_assistant_file(output_path, "extract")
                    return output_path
            elif format_type == "md" or format_type == "markdown":
                output_path = base_path + ".md"
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                self.window.after(0, lambda: self.update_output(f"Exported to Markdown file: {output_path}"))
                self.add_to_assistant_file(output_path, "extract")
                return output_path
            elif format_type == "html":
                output_path = base_path + ".html"
                html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Extracted Content</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
        .container {{ max-width: 800px; margin: 0 auto; }}
        pre {{ background-color:
    </style>
</head>
<body>
    <div class="container">
        <h1>Extracted Content</h1>
        <pre>{content.replace("<", "&lt;").replace(">", "&gt;")}</pre>
    </div>
</body>
</html>"""
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
                self.window.after(0, lambda: self.update_output(f"Exported to HTML file: {output_path}"))
                self.add_to_assistant_file(output_path, "extract")
                return output_path
            elif format_type == "pdf":
                try:
                    from fpdf import FPDF
                    pdf = FPDF()
                    pdf.add_page()
                    pdf.set_font("Arial", size=12)
                    chunks = [content[i:i+80] for i in range(0, len(content), 80)]
                    for chunk in chunks:
                        pdf.multi_cell(0, 10, chunk)
                    output_path = base_path + ".pdf"
                    pdf.output(output_path)
                    self.window.after(0, lambda: self.update_output(f"Exported to PDF file: {output_path}"))
                    self.add_to_assistant_file(output_path, "extract")
                    return output_path
                except ImportError:
                    self.window.after(0, lambda: self.update_output("Error: FPDF library not found. Saving as text instead."))
                    output_path = base_path + ".txt"
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(content)
                    self.window.after(0, lambda: self.update_output(f"Exported to text file (fallback): {output_path}"))
                    self.add_to_assistant_file(output_path, "extract")
                    return output_path
            else:
                self.window.after(0, lambda: self.update_output(f"Unsupported format: {format_type}. Using text format instead."))
                output_path = base_path + ".txt"
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                self.window.after(0, lambda: self.update_output(f"Exported to text file (fallback): {output_path}"))
                self.add_to_assistant_file(output_path, "extract")
                return output_path
        except Exception as e:
            error_msg = f"Error during export: {str(e)}"
            self.window.after(0, lambda msg=error_msg: self.update_output(msg))
            try:
                output_path = base_path + ".txt"
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                self.window.after(0, lambda: self.update_output(f"Exported to text file (fallback): {output_path}"))
                self.add_to_assistant_file(output_path, "extract")
                return output_path
            except:
                return None
    def add_to_recent_files(self, file_path):
        """Add a file to recent files list"""
        if file_path not in self.processed_files:
            self.processed_files.append(file_path)
            if len(self.processed_files) > 10:
                self.processed_files.pop(0)
        self.update_recent_files_list()
        self.save_recent_files()
        self.add_to_assistant_file(file_path, "opened")
    def update_recent_files_list(self):
        """Update the recent files listbox"""
        self.recent_files_listbox.delete(0, tk.END)
        for file_path in reversed(self.processed_files):
            self.recent_files_listbox.insert(0, os.path.basename(file_path))
    def select_recent_file(self, event):
        """Handle selection of a recent file"""
        selection = self.recent_files_listbox.curselection()
        if selection:
            index = selection[0]
            reversed_index = len(self.processed_files) - 1 - index
            if 0 <= reversed_index < len(self.processed_files):
                file_path = self.processed_files[reversed_index]
                self.file_path_var.set(file_path)
                self.update_output(f"Selected recent file: {os.path.basename(file_path)}")
    def save_recent_files(self):
        """Save recent files to a file"""
        recent_files_path = os.path.join(self.assistant_folder, "recent_files.txt")
        try:
            with open(recent_files_path, 'w', encoding='utf-8') as f:
                for file_path in self.processed_files:
                    f.write(f"{file_path}\n")
        except Exception as e:
            logging.error(f"Error saving recent files: {str(e)}")
    def load_recent_files(self):
        """Load recent files from the file"""
        recent_files_path = os.path.join(self.assistant_folder, "recent_files.txt")
        if os.path.exists(recent_files_path):
            try:
                with open(recent_files_path, 'r', encoding='utf-8') as f:
                    self.processed_files = [line.strip() for line in f.readlines() if line.strip()]
                self.update_recent_files_list()
            except Exception as e:
                logging.error(f"Error loading recent files: {str(e)}")
    def interpret_command(self, command, file_path):
        """Use AI to understand natural language commands"""
        self.update_output("Interpreting your command...")
        self.status_var.set("Interpreting...")
        def interpretation_thread():
            try:
                prompt = f"""
                I need to understand the following file processing command: "{command}"
                Please interpret this command and tell me which of these operations it's asking for:
                1. Convert file format - Examples: 
                   - "convert to pdf"
                   - "make this a text file"
                   - "save as HTML"
                   - "export to PDF format"
                   - "change to doc format"
                2. Summarize content - Examples:
                   - "summarize this"
                   - "give me a summary"
                   - "create a summary of this file"
                3. Move file - Examples:
                   - "move to desktop"
                   - "put this in documents folder"
                4. Copy file - Examples:
                   - "copy to downloads"
                   - "duplicate this file to documents"
                5. Extract information - Examples:
                   - "extract phone numbers"
                   - "find all emails"
                   - "pull out the dates"
                Return ONLY a JSON response with these fields:
                {{
                  "operation": "convert|summarize|move|copy|extract|unknown",
                  "target_format": "pdf|txt|etc" (for convert),
                  "target_location": "folder_name" (for move/copy),
                  "extract_what": "text to search for" (for extract)
                }}
                IMPORTANT NOTES:
                - Look for export/save/make as indicators of convert operations
                - The phrase "export to X format" or "save as X" means convert to X format
                - When "document" is mentioned as a format, assume PDF
                - Be flexible with format names (txt, text, pdf, document, etc.)
                """
                response = self.gemini_model.generate_content(prompt)
                interpretation = response.text
                try:
                    import json
                    import re
                    json_match = re.search(r'\{.*\}', interpretation, re.DOTALL)
                    if json_match:
                        json_str = json_match.group(0)
                        interpretation_data = json.loads(json_str)
                        operation = interpretation_data.get("operation", "unknown")
                        if operation == "convert":
                            target_format = interpretation_data.get("target_format", "")
                            if target_format:
                                new_command = f"convert to {target_format}"
                                self.window.after(0, lambda cmd=new_command: self.command_var.set(cmd))
                                self.window.after(0, lambda cmd=new_command: 
                                    self.update_output(f"Interpreted as: {cmd}"))
                                self.window.after(0, lambda: self.convert_file(file_path, new_command))
                            else:
                                self.window.after(0, lambda: 
                                    self.update_output("Couldn't determine the target format"))
                        elif operation == "summarize":
                            self.window.after(0, lambda: 
                                self.update_output("Interpreted as: summarize"))
                            self.window.after(0, lambda: self.summarize_file(file_path))
                        elif operation == "move":
                            target = interpretation_data.get("target_location", "")
                            if target:
                                new_command = f"move to {target}"
                                self.window.after(0, lambda cmd=new_command: self.command_var.set(cmd))
                                self.window.after(0, lambda cmd=new_command: 
                                    self.update_output(f"Interpreted as: {cmd}"))
                                self.window.after(0, lambda: self.move_file(file_path, new_command))
                            else:
                                self.window.after(0, lambda: 
                                    self.update_output("Couldn't determine where to move the file"))
                        elif operation == "copy":
                            target = interpretation_data.get("target_location", "")
                            if target:
                                new_command = f"copy to {target}"
                                self.window.after(0, lambda cmd=new_command: self.command_var.set(cmd))
                                self.window.after(0, lambda cmd=new_command: 
                                    self.update_output(f"Interpreted as: {cmd}"))
                                self.window.after(0, lambda: self.copy_file(file_path, new_command))
                            else:
                                self.window.after(0, lambda: 
                                    self.update_output("Couldn't determine where to copy the file"))
                        elif operation == "extract":
                            extract_what = interpretation_data.get("extract_what", "")
                            if extract_what:
                                new_command = f"extract {extract_what}"
                                self.window.after(0, lambda cmd=new_command: self.command_var.set(cmd))
                                self.window.after(0, lambda cmd=new_command: 
                                    self.update_output(f"Interpreted as: {cmd}"))
                                self.window.after(0, lambda: self.extract_from_file(file_path, new_command))
                            else:
                                self.window.after(0, lambda: 
                                    self.update_output("Couldn't determine what to extract"))
                        else:
                            self.window.after(0, lambda: 
                                self.update_output("I'm not sure how to process that command. Please try a different phrasing."))
                            self.window.after(0, lambda: self.handle_custom_command(file_path, command))
                    else:
                        self.window.after(0, lambda: 
                            self.update_output("Couldn't interpret the command. Trying general processing..."))
                        self.window.after(0, lambda: self.handle_custom_command(file_path, command))
                except Exception as e:
                    error_msg = f"Error interpreting command: {str(e)}. Trying general processing..."
                    self.window.after(0, lambda msg=error_msg: self.update_output(msg))
                    self.window.after(0, lambda: self.handle_custom_command(file_path, command))
                self.window.after(0, lambda: self.add_to_recent_files(file_path))
                self.window.after(0, lambda: self.status_var.set("Completed"))
            except Exception as e:
                error_msg = f"Error processing command: {str(e)}"
                self.window.after(0, lambda msg=error_msg: self.update_output(f"Error: {str(e)}"))
                self.window.after(0, lambda: self.status_var.set("Error"))
        interpretation_thread = threading.Thread(target=interpretation_thread)
        interpretation_thread.daemon = True
        interpretation_thread.start()
    def show_recent_files_window(self):
        """Show a window with recently processed files by type"""
        if hasattr(self, 'recent_window') and self.recent_window is not None and self.recent_window.winfo_exists():
            self.recent_window.lift()
            return
        self.recent_window = ctk.CTkToplevel(self.window)
        self.recent_window.title("Recent Files")
        self.recent_window.geometry("700x500")
        self.recent_window.transient(self.window)
        theme = self.assistant.themes[self.assistant.current_theme]
        self.recent_window.configure(fg_color=theme["bg"])
        tabview = ctk.CTkTabview(
            self.recent_window,
            fg_color=theme["highlight"],
            segmented_button_fg_color=theme["bg"],
            segmented_button_selected_color=theme["accent"],
            segmented_button_unselected_color=theme["bg"]
        )
        tabview.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        all_tab = tabview.add("All Files")
        summary_tab = tabview.add("Summaries")
        convert_tab = tabview.add("Converted")
        extract_tab = tabview.add("Extracted")
        try:
            db_path = os.path.join(self.assistant_folder, "files_database.db")
            if not os.path.exists(db_path):
                conn = sqlite3.connect(db_path)
                cursor = conn.cursor()
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS processed_files (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        file_path TEXT,
                        file_name TEXT,
                        operation_type TEXT,
                        timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
                    )
                ''')
                conn.commit()
                conn.close()
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            cursor.execute(
                "SELECT id, file_path, file_name, operation_type, timestamp FROM processed_files ORDER BY timestamp DESC"
            )
            all_files = cursor.fetchall()
            cursor.execute(
                "SELECT id, file_path, file_name, timestamp FROM processed_files WHERE operation_type='summary' ORDER BY timestamp DESC"
            )
            summaries = cursor.fetchall()
            cursor.execute(
                "SELECT id, file_path, file_name, timestamp FROM processed_files WHERE operation_type='converted' ORDER BY timestamp DESC"
            )
            converted = cursor.fetchall()
            cursor.execute(
                "SELECT id, file_path, file_name, timestamp FROM processed_files WHERE operation_type='extract' ORDER BY timestamp DESC"
            )
            extracted = cursor.fetchall()
            conn.close()
            self.create_files_list(all_tab, all_files, show_type=True)
            self.create_files_list(summary_tab, summaries)
            self.create_files_list(convert_tab, converted)
            self.create_files_list(extract_tab, extracted)
        except Exception as e:
            error_frame = ctk.CTkFrame(all_tab, fg_color=theme["bg"])
            error_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
            ctk.CTkLabel(
                error_frame,
                text=f"Error loading files: {str(e)}",
                font=("Segoe UI", 14),
                text_color=theme["error"]
            ).pack(pady=20)
    def create_files_list(self, parent, files_data, show_type=False):
        """Create a scrollable list of files with actions"""
        theme = self.assistant.themes[self.assistant.current_theme]
        scroll_frame = ctk.CTkScrollableFrame(
            parent,
            fg_color=theme["bg"],
            corner_radius=0
        )
        scroll_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        if not files_data:
            ctk.CTkLabel(
                scroll_frame,
                text="No files found",
                font=("Segoe UI", 14),
                text_color=theme["fg"]
            ).pack(pady=20)
            return
        for file_data in files_data:
            if show_type:
                file_id, file_path, file_name, operation_type, timestamp = file_data
            else:
                file_id, file_path, file_name, timestamp = file_data
                operation_type = ""
            card = ctk.CTkFrame(
                scroll_frame,
                fg_color=theme["highlight"],
                corner_radius=8
            )
            card.pack(fill=tk.X, padx=5, pady=5)
            info_frame = ctk.CTkFrame(
                card,
                fg_color=theme["highlight"],
                corner_radius=0
            )
            info_frame.pack(fill=tk.X, padx=10, pady=5)
            name_label = ctk.CTkLabel(
                info_frame,
                text=file_name,
                font=("Segoe UI", 12, "bold"),
                text_color=theme["fg"]
            )
            name_label.pack(anchor=tk.W)
            if show_type:
                type_date = ctk.CTkLabel(
                    info_frame,
                    text=f"Type: {operation_type.capitalize()} | Date: {timestamp}",
                    font=("Segoe UI", 10),
                    text_color=theme["secondary"]
                )
                type_date.pack(anchor=tk.W)
            else:
                date_label = ctk.CTkLabel(
                    info_frame,
                    text=f"Date: {timestamp}",
                    font=("Segoe UI", 10),
                    text_color=theme["secondary"]
                )
                date_label.pack(anchor=tk.W)
            action_frame = ctk.CTkFrame(
                card,
                fg_color=theme["highlight"],
                corner_radius=0
            )
            action_frame.pack(fill=tk.X, padx=10, pady=5)
            open_button = ctk.CTkButton(
                action_frame,
                text="Open",
                font=("Segoe UI", 11),
                width=80,
                height=28,
                fg_color=theme["accent"],
                hover_color=theme["secondary"],
                command=lambda path=file_path: self.open_file(path)
            )
            open_button.pack(side=tk.LEFT, padx=5, pady=5)
            select_button = ctk.CTkButton(
                action_frame,
                text="Select",
                font=("Segoe UI", 11),
                width=80,
                height=28,
                fg_color=theme["success"],
                hover_color=theme["secondary"],
                command=lambda path=file_path: self.select_recent_file_from_path(path)
            )
            select_button.pack(side=tk.LEFT, padx=5, pady=5)
            delete_button = ctk.CTkButton(
                action_frame,
                text="Delete",
                font=("Segoe UI", 11),
                width=80,
                height=28,
                fg_color=theme["error"],
                hover_color=theme["secondary"],
                command=lambda id=file_id, card=card: self.delete_file_from_db(id, card)
            )
            delete_button.pack(side=tk.RIGHT, padx=5, pady=5)
    def open_file(self, file_path):
        """Open a file with the default application"""
        try:
            if os.path.exists(file_path):
                if platform.system() == 'Windows':
                    os.startfile(file_path)
                elif platform.system() == 'Darwin':
                    subprocess.run(['open', file_path])
                else:
                    subprocess.run(['xdg-open', file_path])
            else:
                messagebox.showerror("Error", f"File not found: {file_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Error opening file: {str(e)}")
    def select_recent_file_from_path(self, file_path):
        """Select a file from recent files window and load it into the main interface"""
        try:
            if os.path.exists(file_path):
                self.file_path_var.set(file_path)
                self.update_output(f"Selected file: {os.path.basename(file_path)}")
                if hasattr(self, 'recent_window') and self.recent_window is not None:
                    self.recent_window.destroy()
            else:
                messagebox.showerror("Error", f"File not found: {file_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Error selecting file: {str(e)}")
    def delete_file_from_db(self, file_id, card_widget):
        """Delete a file from the database"""
        try:
            if not messagebox.askyesno("Confirm", "Are you sure you want to delete this file from the database?"):
                return
            db_path = os.path.join(self.assistant_folder, "files_database.db")
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            cursor.execute("SELECT file_path FROM processed_files WHERE id=?", (file_id,))
            result = cursor.fetchone()
            if result:
                file_path = result[0]
                cursor.execute("DELETE FROM processed_files WHERE id=?", (file_id,))
                conn.commit()
                if messagebox.askyesno("Delete File", "Do you also want to delete the actual file?"):
                    if os.path.exists(file_path):
                        os.remove(file_path)
                card_widget.destroy()
            conn.close()
        except Exception as e:
            messagebox.showerror("Error", f"Error deleting file: {str(e)}")
    def process_voice_command(self, command, file_path=None):
        """Process a voice command from the continuous mode"""
        command_lower = command.lower()
        if 'browse files' in command_lower or 'browse file' in command_lower or 'load file' in command_lower:
            self.update_output("Opening file browser...")
            self.browse_file()
            return
        # Handle "show file" or "open file" to open the current file
        if (command_lower == 'show file' or command_lower == 'open file') and file_path:
            self.update_output(f"Opening current file: {os.path.basename(file_path)}")
            self.open_file(file_path)
            return
        recent_file_patterns = [
            'show the recently', 'show recent', 'open recent', 'show the recent', 
            'show me the recent', 'show the latest', 'open the latest'
        ]
        for pattern in recent_file_patterns:
            if pattern in command_lower:
                recent_info = self._parse_recent_file_command(command)
                if recent_info["is_open_request"]:
                    self.window.after(0, lambda: self.update_output(f"Opening recent file (Type: {recent_info['filter_op_type']})"))
                    self.window.after(0, lambda: self.open_recent_file_by_filter(
                        filter_op_type=recent_info["filter_op_type"],
                        initial_tab_name=recent_info["initial_tab"],
                        filter_date=recent_info["filter_date"],
                        filter_weekday=recent_info["filter_weekday"]
                    ))
                    return True
                break
        if 'browse files' in command_lower or 'browse file' in command_lower or 'load file' in command_lower:
            self.speak("Opening file browser")
            self.root.after(0, lambda: self.browse_and_process_file(command))
        if "show recent" in command_lower or "show files" in command_lower or "recent files" in command_lower or "recently converted" in command_lower or "show converted" in command_lower:
            try:
                operation_type = None
                if "converted" in command_lower:
                    operation_type = "converted"
                elif "summarized" in command_lower or "summarised" in command_lower:
                    operation_type = "summary"
                elif "extracted" in command_lower:
                    operation_type = "extract"
                elif "opened" in command_lower:
                    operation_type = "opened"
                if (any(open_term in command_lower for open_term in 
                        ['show the', 'open the', 'display the', 'show me the']) 
                        and operation_type):
                    self.window.after(0, lambda: self.update_output(f"Opening most recent {operation_type} file"))
                    self.window.after(0, lambda op_type=operation_type: self.open_recent_file_by_filter(filter_op_type=op_type))
                    return True
                self.show_recent_files_window()
                if operation_type and hasattr(self, 'recent_window') and self.recent_window.winfo_exists():
                    tab_map = {
                        "converted": "Converted",
                        "summary": "Summaries",
                        "extract": "Extracted",
                        "opened": "All Files"
                    }
                    if operation_type in tab_map:
                        for widget in self.recent_window.winfo_children():
                            if isinstance(widget, ctk.CTkTabview):
                                try:
                                    widget.set(tab_map[operation_type])
                                    self.update_output(f"Showing {tab_map[operation_type]} tab")
                                    break
                                except Exception as tab_e:
                                    self.update_output(f"Could not select tab: {str(tab_e)}")
            except Exception as e:
                self.update_output(f"Error showing recent files: {str(e)}")
            return True
        self.command_var.set(command)
        self.process_file_command()
        return False
    def reset_voice_ui(self):
        """Reset UI elements after voice command mode is stopped"""
        if hasattr(self, 'voice_loop_button_ref') and self.voice_loop_button_ref:
            try:
                if self.voice_loop_button_ref.winfo_exists():
                    theme = self.assistant.themes[self.assistant.current_theme]
                    self.voice_loop_button_ref.configure(
                        text="Voice Mode", 
                        fg_color=theme["success"]
                    )
            except Exception as e:
                logging.error(f"Error resetting voice UI: {e}")
        for widget in self.window.winfo_children():
            if isinstance(widget, ctk.CTkFrame):
                for child in widget.winfo_children():
                    if isinstance(child, ctk.CTkFrame):
                        for button in child.winfo_children():
                            if isinstance(button, ctk.CTkButton) and button.cget("text") == "🎤":
                                button.configure(state=NORMAL)
        self.restore_main_listening()
    def browse_file(self):
        """Browse for a file without closing the window"""
        try:
            if self.window and self.window.winfo_exists():
                self.window.grab_release()
            file_path = filedialog.askopenfilename(
                parent=self.window,
                title="Select File", 
                filetypes=[
                    ("All Files", "*.*"),
                    ("PDF Files", "*.pdf"),
                    ("Word Documents", "*.docx"),
                    ("Text Files", "*.txt")
                ]
            )
            if file_path:
                self.file_path_var.set(file_path)
                self.update_output(f"Selected file: {os.path.basename(file_path)}")
                self.add_to_recent_files(file_path)
            if self.window and self.window.winfo_exists():
                self.window.focus_force()
                self.window.grab_set()
        except Exception as e:
            self.update_output(f"Error browsing file: {str(e)}")
    def _open_file_dialog(self):
        """Open file dialog in a way that doesn't close the parent window"""
        try:
            if self.window is None or not self.window.winfo_exists():
                return
            self.window.attributes('-topmost', True)
            file_path = filedialog.askopenfilename(
                parent=self.window,
                title="Select File",
                filetypes=(("All Files", "*.*"), ("Text Files", "*.txt"), ("PDF Files", "*.pdf"))
            )
            self.window.attributes('-topmost', False)
            if file_path:
                self.file_path_var.set(file_path)
                self.update_output(f"Selected file: {os.path.basename(file_path)}")
                self.add_to_recent_files(file_path)
            if self.window.winfo_exists():
                self.window.grab_set()
                self.window.focus_force()
        except Exception as e:
            error_msg = f"Error browsing file: {str(e)}"
            if hasattr(self, 'update_output'):
                self.update_output(error_msg)
            logging.error(error_msg)
    def _parse_recent_file_command(self, command):
        """Parses command to find operation type, timeframe, and intent (show vs open) for recent files."""
        command_lower = command.lower()
        operation_type = None
        filter_date = None
        filter_weekday = None
        initial_tab = "All Files"
        is_open_request = False
        type_keywords = {
            'summary': 'Summaries', 'summarise':'Summaries', 'summaries': 'Summaries', 'summarized': 'Summaries',
            'converted': 'Converted', 'conversion': 'Converted', 'conversions': 'Converted',
            'extracted': 'Extracted', 'extraction': 'Extracted', 'extractions': 'Extracted', 'extract': 'Extracted',
            'exported': 'Exported', 'exports': 'Exported', 'export': 'Exported'
        }
        timeframe_keywords = {
            'today': 'today', 'todays': 'today', 'this day': 'today',
            'yesterday': 'yesterday', 'yesterdays': 'yesterday',
            'monday': 0, 'mondays': 0, 'on monday': 0,
            'tuesday': 1, 'tuesdays': 1, 'on tuesday': 1,
            'wednesday': 2, 'wednesdays': 2, 'on wednesday': 2,
            'thursday': 3, 'thursdays': 3, 'on thursday': 3,
            'friday': 4, 'fridays': 4, 'on friday': 4,
            'saturday': 5, 'saturdays': 5, 'on saturday': 5,
            'sunday': 6, 'sundays': 6, 'on sunday': 6,
            'this week': 'this_week',
            'last week': 'last_week',
            'this month': 'this_month',
            'last month': 'last_month'
        }
        weekday_map_to_sqlite = {
             0: '1', 1: '2', 2: '3', 3: '4', 4: '5', 5: '6', 6: '0'
        }
        open_patterns = [
            'show the recently', 'show the recent', 'show recent', 'open recent', 'open the recent',
            'show the last', 'open the last', 'open last', 'display recent', 'show me the recent',
            'show me the last', 'show the latest', 'open the latest'
        ]
        for pattern in open_patterns:
            if pattern in command_lower:
                is_open_request = True
                break
        words = command_lower.split()
        if not is_open_request and words[0] in ['open', 'load', 'display', 'view'] and ('recent' in words or 'file' in words or any(t in words for t in type_keywords) or any(t in words for t in timeframe_keywords)):
             is_open_request = True
        for i, word in enumerate(words):
            if word in type_keywords:
                initial_tab = type_keywords[word]
                db_op_type_map = {
                    'Summaries': 'summary',
                    'Converted': 'converted',
                    'Extracted': 'extract',
                    'Exported': 'exported'
                }
                if initial_tab == 'Exported':
                     operation_type = 'exported%'
                else:
                     operation_type = next((db_type for tab, db_type in db_op_type_map.items() if tab == initial_tab), None)
                break
            if i < len(words) - 1 and f"{word} {words[i+1]}" in type_keywords:
                compound = f"{word} {words[i+1]}"
                initial_tab = type_keywords[compound]
                db_op_type_map = {
                    'Summaries': 'summary',
                    'Converted': 'converted',
                    'Extracted': 'extract',
                    'Exported': 'exported'
                }
                if initial_tab == 'Exported':
                     operation_type = 'exported%'
                else:
                     operation_type = next((db_type for tab, db_type in db_op_type_map.items() if tab == initial_tab), None)
                break
        for i, word in enumerate(words):
            if word in timeframe_keywords:
                timeframe = timeframe_keywords[word]
                self._process_timeframe(timeframe, filter_date, filter_weekday, weekday_map_to_sqlite)
                break
            if i < len(words) - 1 and f"{word} {words[i+1]}" in timeframe_keywords:
                compound = f"{word} {words[i+1]}"
                timeframe = timeframe_keywords[compound]
                filter_date, filter_weekday = self._process_timeframe(timeframe, weekday_map_to_sqlite)
                break
            if word in ['on', 'from', 'since', 'after'] and i < len(words) - 1:
                date_str = words[i+1]
                if re.match(r'\d{4}-\d{2}-\d{2}', date_str):
                    filter_date = date_str
                    break
                if i < len(words) - 2:
                    potential_date = f"{words[i+1]} {words[i+2]}"
                    try:
                        parsed_date = datetime.datetime.strptime(potential_date, "%B %d").replace(
                            year=datetime.datetime.now().year).date()
                        filter_date = parsed_date.isoformat()
                        break
                    except ValueError:
                        try:
                            parsed_date = datetime.datetime.strptime(potential_date, "%d %B").replace(
                                year=datetime.datetime.now().year).date()
                            filter_date = parsed_date.isoformat()
                            break
                        except ValueError:
                            pass
        if not is_open_request and ('recently' in command_lower or 'latest' in command_lower) and operation_type:
            is_open_request = True
        is_recent_command = filter_date is not None or filter_weekday is not None or any(w in command_lower for w in ['recent', 'show file', 'list file', 'open file', 'open summary', 'open converted', 'open extracted', 'latest', 'recently'])
        filter_op_type_for_query = None
        if operation_type:
             if operation_type.endswith('%'):
                  filter_op_type_for_query = operation_type
             elif any(k in words for k in type_keywords):
                 filter_op_type_for_query = operation_type
        if is_open_request and not filter_op_type_for_query and any(term in command_lower for term in ['recent file', 'recently', 'latest file']):
            filter_op_type_for_query = 'opened'
        return {
            "is_recent_command": is_recent_command,
            "is_open_request": is_open_request,
            "initial_tab": initial_tab,
            "filter_op_type": filter_op_type_for_query,
            "filter_date": filter_date,
            "filter_weekday": filter_weekday
        }
    def _process_timeframe(self, timeframe, weekday_map_to_sqlite):
        """Process timeframe value and return appropriate filter_date and filter_weekday values"""
        filter_date = None
        filter_weekday = None
        today = datetime.date.today()
        if timeframe == 'today':
            filter_date = today.isoformat()
        elif timeframe == 'yesterday':
            filter_date = (today - datetime.timedelta(days=1)).isoformat()
        elif timeframe == 'this_week':
            start_of_week = today - datetime.timedelta(days=today.weekday())
            filter_date = f"date(timestamp) >= '{start_of_week.isoformat()}'"
        elif timeframe == 'last_week':
            end_of_last_week = today - datetime.timedelta(days=today.weekday() + 1)
            start_of_last_week = end_of_last_week - datetime.timedelta(days=6)
            filter_date = f"date(timestamp) BETWEEN '{start_of_last_week.isoformat()}' AND '{end_of_last_week.isoformat()}'"
        elif timeframe == 'this_month':
            first_of_month = today.replace(day=1)
            filter_date = f"date(timestamp) >= '{first_of_month.isoformat()}'"
        elif timeframe == 'last_month':
            first_of_this_month = today.replace(day=1)
            last_of_last_month = first_of_this_month - datetime.timedelta(days=1)
            first_of_last_month = last_of_last_month.replace(day=1)
            filter_date = f"date(timestamp) BETWEEN '{first_of_last_month.isoformat()}' AND '{last_of_last_month.isoformat()}'"
        elif isinstance(timeframe, int):
            sqlite_weekday = weekday_map_to_sqlite.get(timeframe)
            if sqlite_weekday:
                filter_weekday = sqlite_weekday
        return filter_date, filter_weekday
    def process_file_command(self):
        file_path = self.file_path_var.get().strip()
        command = self.command_var.get().strip()
        command_lower = command.lower()
        if 'browse files' in command_lower or 'browse file' in command_lower or 'load file' in command_lower:
            self.update_output("Opening file browser...")
            self.browse_file()
            return
        # Handle "show file" or "open file" to open the current file
        if (command_lower == 'show file' or command_lower == 'open file') and file_path:
            self.update_output(f"Opening current file: {os.path.basename(file_path)}")
            self.open_file(file_path)
            return
        recent_file_patterns = [
            'show recently', 'show recent', 'open recent', 'display recent',
            'recently converted', 'recently opened', 'recently extracted',
            'show the recently', 'show the recent', 'open the recent',
            'recent files', 'recent file', 'last converted', 'last opened'
        ]
        if any(pattern in command_lower for pattern in recent_file_patterns):
            self.update_output(f"Detected recent files command: {command}")
            recent_info = self._parse_recent_file_command(command)
            if recent_info["is_recent_command"]:
                if recent_info["is_open_request"]:
                    self.window.after(0, lambda: self.open_recent_file_by_filter(
                        filter_op_type=recent_info["filter_op_type"],
                        filter_date=recent_info["filter_date"],
                        filter_weekday=recent_info["filter_weekday"]
                    ))
                    self.update_output(f"Opening most recent {recent_info['filter_op_type'] or 'file'}")
                else:
                    self.window.after(0, lambda: self.show_recent_files_window(
                        initial_tab_name=recent_info["initial_tab"],
                        filter_op_type=recent_info["filter_op_type"],
                        filter_date=recent_info["filter_date"],
                        filter_weekday=recent_info["filter_weekday"]
                    ))
                    self.update_output(f"Showing recent files with filter: {recent_info['filter_op_type'] or 'All'}")
                return
        if not file_path and any(pattern in command_lower for pattern in 
                                ['show recent', 'open recent', 'recent files']):
            self.update_output("Showing recent files...")
            self.show_recent_files_window()
            return
        if not file_path:
            self.update_output("Error: No file selected")
            return
        if not os.path.exists(file_path):
            self.update_output("Error: File not found")
            return
        if not command:
            self.update_output("Error: No command specified")
            return
        self.update_output(f"Processing command: {command} on file: {os.path.basename(file_path)}")
        self.status_var.set("Processing...")
        def process_thread():
            try:
                if "convert" in command:
                    self.convert_file(file_path, command)
                elif "summarize" in command or "summarise" in command:
                    self.summarize_file(file_path)
                elif "move" in command:
                    self.move_file(file_path, command)
                elif 'browse files' in command_lower or 'browse file' in command_lower or 'load file' in command_lower:
                    self.speak("Opening file browser")
                    self.root.after(0, lambda: self.browse_and_process_file(command))
                elif "copy" in command:
                    self.copy_file(file_path, command)
                elif "extract" in command:
                    self.extract_from_file(file_path, command)
                else:
                    self.handle_custom_command(file_path, command)
                self.window.after(0, lambda: self.add_to_recent_files(file_path))
                self.window.after(0, lambda: self.status_var.set("Completed"))
            except Exception as e:
                error_msg = f"Error processing command: {str(e)}"
                self.window.after(0, lambda msg=error_msg: self.update_output(msg))
                self.window.after(0, lambda: self.status_var.set("Error"))
        processing_thread = threading.Thread(target=process_thread)
        processing_thread.daemon = True
        processing_thread.start()
    def open_recent_file_by_filter(self, filter_op_type=None, filter_date=None, filter_weekday=None):
         """Finds the single most recent file matching filters and opens it."""
         try:
             db_path = os.path.join(self.assistant_folder, "files_database.db")
             if not os.path.exists(db_path):
                  self.update_output("Recent files database not found.")
                  return
             conn = sqlite3.connect(db_path)
             cursor = conn.cursor()
             query = "SELECT file_path, file_name, operation_type, timestamp FROM processed_files"
             conditions = []
             params = []
             type_desc = ""
             if filter_op_type:
                 if filter_op_type.endswith('%'):
                     type_desc = "exported"
                     conditions.append("operation_type LIKE ?")
                     params.append(filter_op_type)
                 else:
                     type_desc = filter_op_type
                     conditions.append("operation_type = ?")
                     params.append(filter_op_type)
             if filter_date:
                 if filter_date.startswith("date(timestamp)"):
                     conditions.append(filter_date)
                 elif filter_date.startswith("date(timestamp) BETWEEN"):
                     conditions.append(filter_date)
                 else:
                     conditions.append("date(timestamp) = date(?)")
                     params.append(filter_date)
             if filter_weekday is not None:
                 conditions.append("strftime('%w', timestamp) = ?")
                 params.append(filter_weekday)
                 weekday_names = ['Sunday', 'Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']
                 weekday_map = {'0': 6, '1': 0, '2': 1, '3': 2, '4': 3, '5': 4, '6': 5}
                 if filter_weekday in weekday_map:
                     weekday_name = weekday_names[weekday_map[filter_weekday]]
                     if type_desc:
                         type_desc = f"{type_desc} on {weekday_name}"
                     else:
                         type_desc = f"processed on {weekday_name}"
             search_desc = type_desc or "processed recently"
             self.update_output(f"Searching for the most recent file {search_desc}...")
             if conditions:
                 query += " WHERE " + " AND ".join(conditions)
             query += " ORDER BY timestamp DESC LIMIT 1"
             cursor.execute(query, params)
             result = cursor.fetchone()
             if result:
                 file_path, file_name, operation_type, timestamp = result
                 if os.path.exists(file_path):
                     operation_desc = {
                         'summary': 'summarized',
                         'converted': 'converted',
                         'extract': 'extracted from',
                         'opened': 'opened'
                     }.get(operation_type, operation_type)
                     try:
                         ts_dt = datetime.datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                         formatted_time = ts_dt.strftime("%Y-%m-%d %H:%M")
                     except:
                         formatted_time = timestamp
                     self.update_output(f"Opening most recent {operation_desc} file: {file_name}")
                     self.update_output(f"This file was {operation_desc} on {formatted_time}")
                     self.open_file(file_path)
                     self.select_recent_file_from_path(file_path)
                 else:
                     self.update_output(f"File found in database but no longer exists: {file_path}")
                     cursor.execute("DELETE FROM processed_files WHERE file_path = ?", (file_path,))
                     conn.commit()
                     self.update_output("The file entry has been removed from the database.")
             else:
                 if type_desc:
                     self.update_output(f"No {type_desc} files found in the database.")
                 else:
                     self.update_output("No matching files found with those filters.")
             conn.close()
         except Exception as e:
             self.update_output(f"Error opening recent file: {str(e)}")
             logging.error(f"Error in open_recent_file_by_filter: {str(e)}")
    def _extract_from_pdf_enhanced(self, file_path):
        """Extract text from PDF file using only PyPDF2"""
        try:
            self.window.after(0, lambda: self.update_output("Extracting PDF text..."))
            import PyPDF2
            text = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                self.window.after(0, lambda: self.update_output(f"PDF has {num_pages} pages"))
                for page_num in range(num_pages):
                    try:
                        page_text = pdf_reader.pages[page_num].extract_text()
                        text.append(page_text)
                        if page_num % 5 == 0 or page_num == num_pages - 1:
                            self.window.after(0, lambda p=page_num, t=num_pages: 
                                self.update_output(f"Reading PDF: page {p+1}/{t}"))
                    except Exception as page_error:
                        self.window.after(0, lambda p=page_num: 
                            self.update_output(f"Could not extract text from page {p+1}"))
            content = "\n\n".join(text)
            self.window.after(0, lambda c=len(content): self.update_output(f"PDF extraction complete: {c} characters"))
            return content
        except Exception as e:
            self.window.after(0, lambda err=str(e): self.update_output(f"PDF extraction error: {err}"))
            return ""
    def _extract_text_from_docx(self, file_path):
        """Extract text from DOCX file"""
        try:
            self.window.after(0, lambda: self.update_output("Extracting DOCX text..."))
            try:
                import docx
                doc = docx.Document(file_path)
                paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
                content = "\n\n".join(paragraphs)
                self.window.after(0, lambda: self.update_output(f"DOCX extraction complete: {len(content)} characters"))
                return content
            except ImportError:
                self.window.after(0, lambda: self.update_output("python-docx not found, trying docx2txt..."))
                import docx2txt
                content = docx2txt.process(file_path)
                self.window.after(0, lambda: self.update_output(f"DOCX extraction complete: {len(content)} characters"))
                return content
        except Exception as e:
            self.window.after(0, lambda err=str(e): self.update_output(f"DOCX extraction error: {err}"))
            return ""
    def _extract_text_from_txt(self, file_path):
        """Extract text from TXT file"""
        try:
            self.window.after(0, lambda: self.update_output("Reading text file..."))
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                content = file.read()
            self.window.after(0, lambda: self.update_output(f"Text file read complete: {len(content)} characters"))
            return content
        except Exception as e:
            self.window.after(0, lambda err=str(e): self.update_output(f"Text file reading error: {err}"))
            return ""
def check_dependencies():
    """Check if all required dependencies are available"""
    try:
        required_modules = [
            'speech_recognition', 'pyttsx3', 'pyautogui', 'pillow', 'numpy',
            'requests', 'psutil', 'googletrans', 'transformers', 'torch'
        ]
        missing_modules = []
        for module in required_modules:
            try:
                importlib.import_module(module)
            except ImportError:
                missing_modules.append(module)
        if missing_modules:
            error_msg = f"Missing required dependencies: {', '.join(missing_modules)}"
            logging.error(error_msg)
            messagebox.showerror("Dependency Error", error_msg)
            return False
        return True
    except Exception as e:
        logging.critical(f"Fatal error checking dependencies: {str(e)}")
        messagebox.showerror("Critical Error", f"Failed to check dependencies: {str(e)}")
        return False
def get_system_info():
    """Get basic system information with error handling"""
    info = {}
    try:
        info['system'] = platform.system()
    except Exception as e:
        info['system'] = f"Unknown (Error: {str(e)})"
    try:
        info['version'] = platform.version()
    except Exception as e:
        info['version'] = f"Unknown (Error: {str(e)})"
    try:
        info['machine'] = platform.machine()
    except Exception as e:
        info['machine'] = f"Unknown (Error: {str(e)})"
    try:
        info['processor'] = platform.processor()
    except Exception as e:
        info['processor'] = f"Unknown (Error: {str(e)})"
    try:
        info['python_version'] = platform.python_version()
    except Exception as e:
        info['python_version'] = f"Unknown (Error: {str(e)})"
    return info
def create_backup():
    """Create backup of important files"""
    try:
        app_dir = get_app_directory()
        backup_dir = os.path.join(app_dir, "backups")
        if not os.path.exists(backup_dir):
            os.makedirs(backup_dir)
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_file = os.path.join(backup_dir, f"backup_{timestamp}.zip")
        files_to_backup = [
            os.path.join(get_log_directory(), "assistant_log.txt"),
            os.path.join(app_dir, "data")
        ]
        import zipfile
        with zipfile.ZipFile(backup_file, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for file in files_to_backup:
                if os.path.exists(file):
                    if os.path.isfile(file):
                        try:
                            zipf.write(file, os.path.basename(file))
                            logging.info(f"Backed up file: {file}")
                        except Exception as e:
                            logging.warning(f"Failed to backup file {file}: {str(e)}")
                    elif os.path.isdir(file):
                        for root, dirs, files in os.walk(file):
                            for f in files:
                                try:
                                    file_path = os.path.join(root, f)
                                    arcname = os.path.relpath(file_path, os.path.dirname(file))
                                    zipf.write(file_path, arcname)
                                    logging.info(f"Backed up file: {file_path}")
                                except Exception as e:
                                    logging.warning(f"Failed to backup file {file_path}: {str(e)}")
        logging.info(f"Backup created successfully: {backup_file}")
        return backup_file
    except Exception as e:
        logging.error(f"Backup creation failed: {str(e)}")
        return None
def initialize_task_db():
        """Initialize the task database with proper schema"""
        try:
            db_path = "task.db"
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            cursor.execute('''CREATE TABLE IF NOT EXISTS scheduled_tasks
                             (id INTEGER PRIMARY KEY AUTOINCREMENT,
                              task TEXT NOT NULL,
                              seconds INTEGER NOT NULL,
                              scheduled_time TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                              execution_time TIMESTAMP,
                              completed INTEGER DEFAULT 0,
                              status TEXT DEFAULT 'scheduled')''')
            conn.commit()
            conn.close()
            logging.info(f"Database initialized at {db_path}")
        except Exception as e:
            logging.error(f"Error initializing database: {str(e)}")
            raise
def initialize_auto_login_db():
    """Initialize the auto login database."""
    try:
        conn = sqlite3.connect('auto_login.db')
        cursor = conn.cursor()
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS credentials (
                id INTEGER PRIMARY KEY,
                website TEXT NOT NULL,
                email TEXT NOT NULL,
                password TEXT NOT NULL
            )
        ''')
        conn.commit()
        conn.close()
    except Exception as e:
        logging.error(f"Error initializing auto login database: {str(e)}")
def initialize_url_db():
    """Initialize the URL database."""
    try:
        conn = sqlite3.connect('url.db')
        cursor = conn.cursor()
        cursor.execute('''CREATE TABLE IF NOT EXISTS urls (
            id INTEGER PRIMARY KEY,
            website TEXT NOT NULL UNIQUE,
            url TEXT NOT NULL
        )''')
        conn.commit()
        conn.close()
    except Exception as e:
        logging.error(f"Error initializing URL database: {str(e)}")
def initialize_contact_database():
        """Initialize the contacts database"""
        try:
            conn = sqlite3.connect("samcontact.db")
            cursor = conn.cursor()
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS contacts (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    name TEXT NOT NULL,
                    number TEXT NOT NULL,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            conn.commit()
            logging.info("Database initialized successfully")
        except Exception as e:
            logging.error(f"Database initialization error: {str(e)}")
def add_url(website, url):
    """Add a new URL to the database."""
    try:
        conn = sqlite3.connect('url.db')
        cursor = conn.cursor()
        cursor.execute('INSERT OR REPLACE INTO urls (website, url) VALUES (?, ?)', (website.lower(), url))
        conn.commit()
        conn.close()
    except Exception as e:
        logging.error(f"Error adding URL: {str(e)}")
def get_url(website):
    """Retrieve a URL from the database by website name."""
    try:
        conn = sqlite3.connect('url.db')
        cursor = conn.cursor()
        cursor.execute('SELECT url FROM urls WHERE website = ?', (website.lower(),))
        result = cursor.fetchone()
        conn.close()
        return result[0] if result else None
    except Exception as e:
        logging.error(f"Error retrieving URL: {str(e)}")
        return None
def click_at_coordinates(x, y):
    """Moves mouse to (x, y) and performs a left click."""
    try:
        target_x, target_y = int(x), int(y)
        print(f"Attempting to click at ({target_x}, {target_y})...")
        pyautogui.moveTo(target_x, target_y, duration=0.2)
        time.sleep(0.1)
        pyautogui.click(target_x, target_y)
        print(f"Clicked at ({target_x}, {target_y}).")
        return True
    except Exception as e:
        print(f"Error performing click at ({x}, {y}): {e}")
        return False
def type_text(text_to_type):
    """Uses pyautogui to type the given text."""
    try:
        print(f"Attempting to type: '{text_to_type}'")
        pyautogui.write(text_to_type, interval=0.05)
        print("Typing complete.")
        return True
    except Exception as e:
        print(f"Error performing typing: {e}")
        return False
def get_active_app_info():
    """Get information about the currently active application window."""
    result = {"app_name": "Unknown", "window_title": "Unknown"}
    try:
        active_window = gw.getActiveWindow()
        if active_window:
            result["window_title"] = active_window.title
            app_parts = result["window_title"].split(" - ")
            result["app_name"] = app_parts[-1].strip() if len(app_parts) > 1 else result["window_title"]
            result["app_name"] = result["app_name"][:50]
        else:
            logging.warning("Could not get active window (None returned).")
    except Exception as e:
        logging.error(f"Error getting active window info: {e}")
    if not result["app_name"] or result["app_name"] == "Unknown":
        result["app_name"] = "unknown_app"
    return result
def capture_screen_to_pil():
    """Captures the primary screen and returns it as a PIL Image object."""
    try:
        time.sleep(0.2)
        screenshot = ImageGrab.grab()
        print(f"Screenshot captured (size: {screenshot.size})")
        return screenshot
    except Exception as e:
        print(f"Error capturing screen: {e}")
        return None
class ScreenCoordinateAdapter:
    """Adapts coordinates based on screen resolution differences between saved and current resolutions."""
    def __init__(self):
        self.current_width, self.current_height = pyautogui.size()
        self.reference_width, self.reference_height = self._get_reference_resolution()
        logging.info(f"Screen adapter initialized: Current: {self.current_width}x{self.current_height}, Reference: {self.reference_width}x{self.reference_height}")
    def _get_reference_resolution(self):
        """Get reference resolution from settings file, or create if not exists"""
        reference_file = "screen_resolution.json"
        try:
            if os.path.exists(reference_file):
                with open(reference_file, 'r') as f:
                    data = json.load(f)
                    return data.get('width', self.current_width), data.get('height', self.current_height)
            else:
                self._save_reference_resolution(self.current_width, self.current_height)
                return self.current_width, self.current_height
        except Exception as e:
            logging.error(f"Error getting reference resolution: {e}")
            return self.current_width, self.current_height
    def _save_reference_resolution(self, width, height):
        """Save reference resolution to settings file"""
        reference_file = "screen_resolution.json"
        try:
            with open(reference_file, 'w') as f:
                json.dump({'width': width, 'height': height}, f)
            logging.info(f"Saved reference resolution: {width}x{height}")
            return True
        except Exception as e:
            logging.error(f"Error saving reference resolution: {e}")
            return False
    def update_reference_resolution(self):
        """Update the reference resolution to current screen size"""
        return self._save_reference_resolution(self.current_width, self.current_height)
    def adapt_coordinates(self, x, y):
        """
        Adapt coordinates based on the difference between reference and current screen size
        Returns: (adapted_x, adapted_y)
        """
        if self.reference_width == self.current_width and self.reference_height == self.current_height:
            return x, y
        x_scale = self.current_width / self.reference_width
        y_scale = self.current_height / self.reference_height
        adapted_x = int(x * x_scale)
        adapted_y = int(y * y_scale)
        logging.info(f"Adapted coordinates from ({x}, {y}) to ({adapted_x}, {adapted_y})")
        return adapted_x, adapted_y
if __name__ == "__main__":
    try:
        if not create_required_directories():
            print("Warning: Some directories could not be created")
        if not setup_logging():
            print("Warning: Using fallback logging configuration")
        initialize_url_db()
        initialize_auto_login_db()
        initialize_contact_database()
        initialize_task_db()
        root = ctk.CTk()
        app = VoiceAssistant(root)
        root.mainloop()
    except Exception as e:
        print(f"Fatal error during startup: {str(e)}")
        sys.exit(1)
def format_time(seconds):
    """Convert seconds to formatted time string"""
    minutes, seconds = divmod(seconds, 60)
    hours, minutes = divmod(minutes, 60)
    return f"{int(hours):02d}:{int(minutes):02d}"