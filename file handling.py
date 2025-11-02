import tkinter as tk
from tkinter import ttk, scrolledtext, Label, Frame, Button,Entry, Text, END, NORMAL, DISABLED, WORD,Canvas, Tk, filedialog
import customtkinter as ctk
import pyttsx3
import datetime
import speech_recognition as sr
from PIL import Image, ImageGrab, ImageDraw, ImageFont
import webbrowser
import numpy as np
import requests
import io
import base64
from openai import OpenAI
from urllib.parse import urlparse
import os
import pyperclip
import pandas as pd
import time
import math
import threading
import logging
from urllib3.util.retry import Retry
from requests.adapters import HTTPAdapter
import pygame
import screen_brightness_control as sbc
import keyboard
import pyautogui
from PIL import Image, ImageGrab
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import smtplib
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import psutil
from tkinter import messagebox
from deep_translator import GoogleTranslator
import urllib.parse
import sys
import psutil
import random
import platform
import subprocess
import importlib
from pathlib import Path
import sqlite3
import google.generativeai as genai
from urllib.parse import quote
import shutil
import re
def get_app_directory():
    try:
        home = str(Path.home())
        app_dir = os.path.join(home, "VoiceAssistant")
        return app_dir
    except Exception as e:
        print(f"Error getting app directory: {str(e)}")
        return os.path.join(os.path.dirname(os.path.abspath(__file__)), "VoiceAssistant")
def ensure_directory_exists(directory):
    try:
        os.makedirs(directory, exist_ok=True)
        return True
    except Exception as e:
        print(f"Error creating directory {directory}: {str(e)}")
        return False
def get_log_directory():
    try:
        app_dir = get_app_directory()
        log_dir = os.path.join(app_dir, "logs")
        return log_dir
    except Exception as e:
        print(f"Error getting log directory: {str(e)}")
        return os.path.join(os.path.dirname(os.path.abspath(__file__)), "logs")
def setup_logging():
    try:
        log_dir = get_log_directory()
        if not ensure_directory_exists(log_dir):
            print("Warning: Failed to create log directory. Using basic console logging.")
            logging.basicConfig(
                level=logging.INFO,
                format='%(asctime)s - %(levelname)s - %(message)s',
                handlers=[logging.StreamHandler()]
            )
            return False
        log_file = os.path.join(log_dir, "voice_assistant.log")
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler(log_file),
                logging.StreamHandler()
            ]
        )
        logging.info("Logging initialized successfully")
        return True
    except Exception as e:
        print(f"Warning: Error in logging setup: {str(e)}. Using basic console logging.")
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[logging.StreamHandler()]
        )
        return False
client = OpenAI(
    base_url="https://openrouter.ai/api/v1",
    api_key="Your_OpenRouter_API_Key"
)
def create_required_directories():
    try:
        app_dir = get_app_directory()
        required_dirs = {
            'app': app_dir,
            'logs': os.path.join(app_dir, "logs"),
            'data': os.path.join(app_dir, "data"),
            'config': os.path.join(app_dir, "configure"),
            'temp': os.path.join(app_dir, "temp")
        }
        success = True
        for dir_name, dir_path in required_dirs.items():
            if not ensure_directory_exists(dir_path):
                print(f"Warning: Failed to create {dir_name} directory at {dir_path}")
                success = False
        return success
    except Exception as e:
        print(f"Error creating required directories: {str(e)}")
        return False
ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("blue")
engine = pyttsx3.init('sapi5')
voices = engine.getProperty('voices')
engine.setProperty('voice', voices[0].id)
def play_sound():
    try:
        pygame.mixer.init()
        pygame.mixer.music.load(r'C:\Users\HP\Desktop\Ai\start_sounds\notification sound 2_15-03-25_15-27-59-839.mp3')
        pygame.mixer.music.play()
        while pygame.mixer.music.get_busy():
            pygame.time.Clock().tick(10)
    except Exception as e:
        print(f"Failed to play sound: {e}")
def play_sound2():
    try:
        pygame.mixer.init()
        pygame.mixer.music.load(r'C:\Users\HP\Desktop\Ai\start_sounds\mixkit-sci-fi-click-900.mp3')
        pygame.mixer.music.play()
        while pygame.mixer.music.get_busy():
            pygame.time.Clock().tick(10)
    except Exception as e:
        print(f"Failed to play sound: {e}")
def play_sound3():
    try:
        pygame.mixer.init()
        pygame.mixer.music.load(r'C:\Users\HP\Desktop\Ai\start_sounds\mixkit-fast-sci-fi-bleep-903.mp3')
        pygame.mixer.music.play()
        while pygame.mixer.music.get_busy():
            pygame.time.Clock().tick(10)
    except Exception as e:
        print(f"Failed to play sound: {e}")
def play_sound4():
    try:
        pygame.mixer.init()
        pygame.mixer.music.load(r'C:\Users\HP\Desktop\Ai\start_sounds\mixkit-fast-sci-fi-bleep-903.mp3')
        pygame.mixer.music.play()
        while pygame.mixer.music.get_busy():
            pygame.time.Clock().tick(10)
    except Exception as e:
        print(f"Failed to play sound: {e}")
class WaveVisualizer(tk.Canvas):
    def _convert_tk_properties(self, **kwargs):
        """Convert customtkinter properties to tkinter equivalents"""
        property_map = {
            'fg_color': 'bg',
            'border_color': 'highlightbackground',
            'border_width': 'highlightthickness'
        }
        converted_kwargs = {}
        for key, value in kwargs.items():
            if key in property_map:
                converted_kwargs[property_map[key]] = value
            else:
                converted_kwargs[key] = value
        return converted_kwargs
    def __init__(self, parent, **kwargs):
        kwargs = self._convert_tk_properties(**kwargs)
        super().__init__(parent, **kwargs)
        self.configure(bg='#000000')  # Black background
        self.is_active = True
        self.animation_running = False
        self.after_id = None
        self._lock = threading.Lock()
        self.particles = []
        self.colors = ['#FF0000', '#FF7F00', '#FFFF00', '#00FF00', '#0000FF', '#4B0082', '#8B00FF']  # Rainbow colors
        self.init_particles()
    def init_particles(self):
        """Initialize particles for the visualization"""
        self.particles = []
        num_particles = 50
        for i in range(num_particles):
            angle = (2 * math.pi * i) / num_particles
            self.particles.append({
                'angle': angle,
                'radius': 20,
                'speed': 0.1 + random.random() * 0.2,
                'color': random.choice(self.colors),
                'size': random.randint(2, 5)
            })
    def speaking_animation(self):
        if not self.is_active or self.animation_running:
            return
        def _animate():
            if not self.is_active or not self.animation_running:
                return False
            try:
                with self._lock:
                    self.delete("wave")
                    width = self.winfo_width()
                    height = self.winfo_height()
                    center_x = width / 2
                    center_y = height / 2
                    current_time = time.time()
                    for particle in self.particles:
                        particle['radius'] += math.sin(current_time * 2) * 0.5
                        particle['angle'] += particle['speed']
                        x = center_x + particle['radius'] * math.cos(particle['angle']) * 5
                        y = center_y + particle['radius'] * math.sin(particle['angle']) * 5
                        size = particle['size'] * (1 + math.sin(current_time * 4) * 0.3)
                        glow_size = size * 2
                        self.create_oval(
                            x - glow_size, y - glow_size,
                            x + glow_size, y + glow_size,
                            fill='', outline=particle['color'],
                            width=1, tags="wave",
                            stipple='gray50'
                        )
                        self.create_oval(
                            x - size, y - size,
                            x + size, y + size,
                            fill=particle['color'],
                            outline='',
                            tags="wave"
                        )
                        if random.random() < 0.3:
                            next_particle = random.choice(self.particles)
                            next_x = center_x + next_particle['radius'] * math.cos(next_particle['angle']) * 5
                            next_y = center_y + next_particle['radius'] * math.sin(next_particle['angle']) * 5
                            self.create_line(
                                x, y, next_x, next_y,
                                fill=particle['color'],
                                width=1,
                                tags="wave",
                                stipple='gray50'
                            )
                    return True
            except tk.TclError:
                self.is_active = False
                return False
            except Exception as e:
                logging.error(f"Animation error: {str(e)}")
                return False
        def animate_loop():
            if _animate():
                self.after_id = self.after(30, animate_loop)
            else:
                self.animation_running = False
                self.after_id = None
        with self._lock:
            if not self.animation_running and self.is_active:
                self.animation_running = True
                self.after_id = self.after(30, animate_loop)
    def start_listening_animation(self):
        if not self.is_active or self.animation_running:
            return
        def _animate():
            if not self.is_active or not self.animation_running:
                return False
            try:
                with self._lock:
                    self.delete("wave")
                    width = self.winfo_width()
                    height = self.winfo_height()
                    center_x = width / 2
                    center_y = height / 2
                    current_time = time.time()
                    max_radius = min(width, height) / 3
                    num_circles = 8
                    for i in range(num_circles):
                        radius = (max_radius * (i + 1) / num_circles) + (math.sin(current_time * 2) * 10)
                        color = self.colors[i % len(self.colors)]
                        self.create_oval(
                            center_x - radius, center_y - radius,
                            center_x + radius, center_y + radius,
                            outline=color,
                            width=2,
                            tags="wave"
                        )
                        for j in range(8):
                            angle = (2 * math.pi * j / 8) + current_time * 2
                            dot_x = center_x + radius * math.cos(angle)
                            dot_y = center_y + radius * math.sin(angle)
                            dot_size = 3 + math.sin(current_time * 4 + j) * 2
                            self.create_oval(
                                dot_x - dot_size, dot_y - dot_size,
                                dot_x + dot_size, dot_y + dot_size,
                                fill=color,
                                outline='',
                                tags="wave"
                            )
                    return True
            except tk.TclError:
                self.is_active = False
                return False
            except Exception as e:
                logging.error(f"Animation error: {str(e)}")
                return False
        def animate_loop():
            if _animate():
                self.after_id = self.after(30, animate_loop)
            else:
                self.animation_running = False
                self.after_id = None
        with self._lock:
            if not self.animation_running and self.is_active:
                self.animation_running = True
                self.after_id = self.after(30, animate_loop)
    def stop_listening_animation(self):
        with self._lock:
            self.animation_running = False
            if self.after_id:
                self.after_cancel(self.after_id)
                self.after_id = None
            if self.is_active:
                try:
                    self.delete("wave")
                except tk.TclError:
                    self.is_active = False
    def cleanup(self):
        """Clean up the visualizer before window destruction"""
        with self._lock:
            self.is_active = False
            self.animation_running = False
            if self.after_id:
                self.after_cancel(self.after_id)
                self.after_id = None
            try:
                self.delete("all")
            except tk.TclError:
                pass
class VoiceAssistant:
    def __init__(self, root):
        self.root = root
        self.programming_window = None
        self.root.title("A.N.I.S AI Assistant")
        self.root.geometry("650x700")
        self.root.resizable(False, False)
        self.default_prompts = {
            'study_notes': """You are a helpful study assistant focusing on {topic}. 
            Please provide a clear and informative response to this question: {question},should be short and professional
            Include key points, examples, and explanations where relevant.swer should be short and concise in 1 minutes or less."""
        }
        self.programming_window = None
        self.study_mode = False
        self.current_topic = None
        self.study_progress = {}
        self.api_key = "Your_Gemini_API_Key"
        self.quiz_window = None
        self.quiz_questions = []
        self.quiz_answers = []
        self.quiz_correct_answers = []
        self.quiz_score = 0
        self.quiz_total_questions = 0
        self.main_container = ctk.CTkFrame(
            self.root,
            fg_color='#0a0a1f',
            corner_radius=0
        )
        self.main_container.pack(fill=tk.BOTH, expand=True, padx=30, pady=20)
        self.banner_frame = ctk.CTkFrame(
            self.main_container,
            fg_color='#0a0a1f',
            height=100,
            corner_radius=0
        )
        self.banner_frame.pack(fill=tk.X, pady=(0, 20))
        self.gradient_frame = ctk.CTkFrame(
            self.main_container,
            fg_color='#0a0a1f',
            border_color='#2a2a4a',
            border_width=2,
            corner_radius=10
        )
        self.gradient_frame.pack(fill=tk.X, pady=20)
        self.wave_frame = ctk.CTkFrame(
            self.gradient_frame,
            fg_color='#0a0a1f',
            corner_radius=8
        )
        self.wave_frame.pack(fill=tk.X, padx=15, pady=15)
        self.themes = {
            "Dark Blue": {
                 "bg": "#0D0221",
                "fg": "#00FF9F",
                "accent": "#FF2975",
                "secondary": "#2DE2E6",
                "highlight": "#261447",
                "text_bg": "#0F0326",
                "success": "#00FF9F",
                "error": "#FF2975"
            },
            "Cyberpunk": {
                  "bg": "#0a0a1f",
                "fg": "#e2e2e2",
                "accent": "#4361ee",
                "secondary": "#8f9aff",
                "highlight": "#2a2a4a",
                "text_bg": "#12122a",
                "success": "#4CAF50",
                "error": "#ff4d4d"
            },
            "Forest": {
                "bg": "#1A2F1D",
                "fg": "#C7F9CC",
                "accent": "#57BA5E",
                "secondary": "#80ED99",
                "highlight": "#2D4A31",
                "text_bg": "#1F3823",
                "success": "#57BA5E",
                "error": "#FF6B6B"
            },
            "Ocean": {
                "bg": "#0A192F",
                "fg": "#E6F1FF",
                "accent": "#64FFDA",
                "secondary": "#8892B0",
                "highlight": "#172A45",
                "text_bg": "#0C1B2B",
                "success": "#64FFDA",
                "error": "#FF647F"
            },
            "Tron Legacy": {
                "bg": "#0C141F",
                "fg": "#00F2FF",
                "accent": "#15f4ee",
                "secondary": "#2a2a4a",
                "highlight": "#2a2a4a",
                "text_bg": "#12122a",
                "success": "#4CAF50",
                "error": "#ff4d4d"
            }
        }
        self.current_theme = "Dark Blue"
        self.wave_vis = WaveVisualizer(
            self.wave_frame, 
            width=820,
            height=150,
            bg='#000000',
            highlightthickness=1,
            highlightbackground='#2a2a4a'
        )
        self.wave_vis.pack(pady=5)
        self.create_gui()
        self.apply_theme(self.current_theme)
        self.root.protocol("WM_DELETE_WINDOW", self.exit_assistant)
        try:
            import google.generativeai as genai
            genai.configure(api_key="Your_Gemini_API_Key")
            self.gemini_model = genai.GenerativeModel('gemini-2.0-flash')
            self.chat = self.gemini_model.start_chat(history=[])
        except Exception as e:
            self.handle_errors(e, "Error initializing Gemini API")
        self.file_assistant = FileHandlingAssistant(self.root, self)
    def stop_typing(self):
        """Stop typing mode"""
        self.typing_active = False
        self.status_label.configure(text="Listening...")
    def create_gui(self):
        top_control_frame = ctk.CTkFrame(
            self.main_container,
            fg_color='#0a0a1f',
            height=60,
            corner_radius=0
        )
        top_control_frame.pack(fill=tk.X, pady=(0, 10))
        top_control_frame.pack_propagate(False)
        stats_frame = ctk.CTkFrame(
            top_control_frame,
            fg_color='#0a0a1f',
            corner_radius=0
        )
        stats_frame.pack(side=tk.LEFT, padx=20)
        self.cpu_label = ctk.CTkLabel(
            stats_frame,
            text="CPU: 0%",
            font=("Segoe UI", 10),
            text_color='#8f9aff'
        )
        self.cpu_label.pack(side=tk.LEFT, padx=(0, 15))
        self.memory_label = ctk.CTkLabel(
            stats_frame,
            text="RAM: 0%",
            font=("Segoe UI", 10),
            text_color='#8f9aff'
        )
        self.memory_label.pack(side=tk.LEFT, padx=(0, 15))
        self.time_label = ctk.CTkLabel(
            stats_frame,
            text="",
            font=("Segoe UI", 10),
            text_color='#8f9aff'
        )
        self.time_label.pack(side=tk.LEFT)
        self.update_time()
        control_buttons_frame = ctk.CTkFrame(
            top_control_frame,
            fg_color='#0a0a1f',
        )
        control_buttons_frame.pack(side=tk.RIGHT, padx=20)
        self.settings_button = ctk.CTkButton(
            control_buttons_frame,
            text="⚙",
            font=('Segoe UI', 16),
            fg_color='#0a0a1f',
            text_color='#8f9aff',
            border_width=0,
            corner_radius=8,
            hover_color='#2a2a4a',
            command=self.show_settings
        )
        self.settings_button.pack(side=tk.RIGHT, padx=(10))
        play_container = ctk.CTkFrame(
            control_buttons_frame,
            fg_color='#0a0a1f',
            border_color='#00ff00',
        )
        play_container.pack(side=tk.RIGHT, padx=10)
        self.start_button = ctk.CTkButton(
            play_container,
            text="▶",
            command=self.toggle_listening,
            font=('Segoe UI', 16, 'bold'),
            fg_color='#00ff00',
            text_color='#000000',
            hover_color='#00cc00',
            width=40,
            height=40,
            corner_radius=20,
            border_width=0
        )
        self.start_button.pack(padx=2, pady=2)
        self.mic_frame = ctk.CTkFrame(
            self.main_container,
            fg_color='#0a0a1f',
            height=30,
            corner_radius=0
        )
        self.mic_frame.pack(fill=tk.X, pady=(0, 10))
        self.mic_dots = []
        for i in range(10):
            dot = ctk.CTkLabel(
                self.mic_frame,
                text="●",
                font=("Segoe UI", 8),
                text_color='#1f1f3f'  # Start with dark color
            )
            dot.pack(side=tk.LEFT, padx=2)
            self.mic_dots.append(dot)
        self.start_button._clicked = False
        self.start_button._listening = False
        self.listening_thread = None
        self.update_system_stats()
        self.header_label = ctk.CTkLabel(
            self.banner_frame, 
            text="A.N.I.S",
            font=("Segoe UI", 50, "bold"),
            text_color='#4361ee'
        )
        self.header_label.pack(pady=(10, 0))
        self.subtitle_label = ctk.CTkLabel(
            self.banner_frame,
            text="Artificial Neural Intelligence System",
            font=("Segoe UI Light", 14),
            text_color='#8f9aff'
        )
        self.subtitle_label.pack()
        self.status_frame = ctk.CTkFrame(self.main_container, fg_color='#0a0a1f', corner_radius=0)
        self.status_frame.pack(fill=tk.X, pady=10)
        self.status_label = ctk.CTkLabel(
            self.status_frame,
            text="● Not Listening",
            font=("Segoe UI", 12),
            text_color='#ff4d4d'  # Red when not listening
        )
        self.status_label.pack(side=tk.LEFT, padx=20)
        text_container = ctk.CTkFrame(
            self.main_container,
            fg_color='#0a0a1f',
            border_color='#2a2a4a',
            border_width=1,
            corner_radius=10
        )
        text_container.pack(fill=tk.BOTH, expand=True, pady=20)
        text_frame = ctk.CTkFrame(
            text_container,
            fg_color='#0a0a1f',
            corner_radius=8
        )
        text_frame.pack(fill=tk.BOTH, expand=True, padx=2, pady=2)
        style = ttk.Style()
        style.configure(
            "Custom.Vertical.TScrollbar",
            background='#4361ee',
            troughcolor='#1a1a35',
            width=10,
            arrowsize=0
        )
        self.output_text = ctk.CTkTextbox(
            text_frame,
            height=350,
            width=750,
            fg_color='#12122a',
            text_color='#e2e2e2',
            font=("JetBrains Mono", 10),
            corner_radius=8,
            border_width=0,
            wrap="word"
        )
        self.output_text.pack(fill='both', expand=True, padx=3, pady=3)
        self.output_text._textbox.tag_configure(
            "user_msg", 
            foreground="#BF00FF",
            font=("JetBrains Mono", 9, "bold")
        )
        self.output_text._textbox.tag_configure(
            "assistant_msg", 
            foreground="#4CAF50",
            font=("JetBrains Mono", 10)
        )
        self.output_text._textbox.tag_configure(
            "time_stamp", 
            foreground="#666666",
            font=("JetBrains Mono", 8)
        )
        def insert_with_timestamp(self, text, tag):
            timestamp = datetime.datetime.now().strftime("%H:%M:%S")
            self.output_text.configure(state=NORMAL)
            self.output_text._textbox.insert(END, f"[{timestamp}] ", "time_stamp")
            self.output_text._textbox.insert(END, f"{text}\n", tag)
            self.output_text.configure(state=DISABLED)
            self.output_text.see(END)
        self.insert_with_timestamp = insert_with_timestamp.__get__(self)
    def update_time(self):
        """Update the time display"""
        try:
            if hasattr(self, 'time_label') and self.time_label.winfo_exists():
                current_time = datetime.datetime.now().strftime("%I:%M:%S %p")
                self.time_label.configure(text=current_time)
                self._after_id = self.root.after(1000, self.update_time)
        except Exception as e:
            logging.error(f"Error updating time: {str(e)}")
            return
    def update_status(self, is_listening=False):
        """Update the status label safely"""
        try:
            if hasattr(self, 'status_label') and self.status_label.winfo_exists():
                if is_listening:
                    self.status_label.configure(
                        text="● Listening",
                        text_color='#4CAF50'  # Green when listening
                    )
                else:
                    self.status_label.configure(
                        text="● Not Listening",
                        text_color='#ff4d4d'  # Red when not listening
                    )
        except tk.TclError:
            pass
    def type_and_enter(self, command):
        """Type the text at cursor position with minimal delay"""
        try:
            if command.startswith('type '):
                text = command[5:].strip()
                if text:
                    self.typing_active = True
                    self.status_label.configure(text="Typing Mode Active")
                    pyautogui.typewrite(text, interval=0.01)
                    self.speak('Done')
                else:
                    self.speak('Please provide text to type')
        except Exception as e:
            logging.error(f"Error in typing: {str(e)}")
            self.speak("Error typing")
    def process_command(self, command):
        """Process voice commands"""
        try:
            play_sound2()
            command = command.lower().strip()
            if any(phrase in command for phrase in ['file assistant', 'file handling', 'handle file', 'file prompt']):
                self.speak("Opening file handling assistant")
                self.root.after(0, self.file_assistant.show_window)
                return False
            elif 'browse file' in command:
                self.speak("Opening file browser")
                self.root.after(0, lambda: self.browse_and_process_file(command))
                return False
            elif any(phrase in command for phrase in ['show yesterday files', 'show yesterday file', 'yesterday file', 'recent files']):
                self.speak("Showing recent files")
                self.root.after(0, self.show_recent_files)
                return False
            elif any(phrase in command for phrase in ['show the recently converted', 'show recent converted', 'open recent converted', 
                                                     'show the recently summarized', 'show recent summarized', 'open recent summarized',
                                                     'show the recently extracted', 'show recent extracted', 'open recent extracted',
                                                     'show the latest file', 'show recently processed', 'open latest file']):
                self.speak("Opening recently processed file")
                self.root.after(0, lambda cmd=command: self.open_recent_file(cmd))
                return False
            elif 'comments' in command or 'in python' in command or 'coding' in command:
                command = command.replace('code', '').strip()
                command = command.replace('coding','').strip()
                command = command.replace('comments', '').strip()
                command = command.replace('in python', '').strip()
                self.show_program_output_code(command)
                return False
            elif "auto install" in command or 'install' in command:
                app_name = command.replace("auto install", "").replace("install","").replace("how to","").strip()
                if app_name:
                    self.speak(f"Attempting to install {app_name}")
                    self.root.after(0, lambda: self.auto_install_app(app_name))
                else:
                    self.speak("Please specify the app name to install.")
                return False 
            elif any(cmd in command for cmd in ['analyse news','analyse text','analyse the news','analyse the highlighted','analyse highlights','analyse highlight','analyse the highlight','check news','check text','check link','analyse link']):
                highlighted_text = pyautogui.hotkey('ctrl', 'c')
                time.sleep(0.1)
                text_to_analyze = pyperclip.paste()
                if text_to_analyze:
                    self.speak("Analyzing the selected content")
                    self.analyze_news(text_to_analyze)
                else:
                    self.speak("Please highlight some text or a link to analyze")
                return False
            elif 'study mode on' in command:
                self.study_mode = True
                self.speak("Study mode activated. What would you like to learn?")
                topic = self.take_command().lower()
                if topic != 'none':
                    self.current_topic = topic
                    self.speak(f"Great! Let's start learning about {topic}")
                    search_query = f"{topic} course tutorial"
                    google_url = f'https://www.google.com/search?q={urllib.parse.quote(search_query)}'
                    search_url = f"https://www.youtube.com/results?search_query={urllib.parse.quote(search_query)}"
                    html = urllib.request.urlopen(search_url)
                    video_ids = re.findall(r"watch\?v=(\S{11})", html.read().decode())
                    if video_ids:
                            first_video_url = f"https://www.youtube.com/watch?v={video_ids[0]}"
                    else:
                            self.speak("No videos found for your search")
                    webbrowser.open(google_url)
                    webbrowser.open(first_video_url)
                    pyautogui.hotkey('enter')
                    self.speak(f"Playing video for {search_query} on YouTube")
                    self.root.after(0, self.show_chats_window)
                    if topic not in self.study_progress:
                        self.study_progress[topic] = {
                            'level': 1,
                            'completed_tasks': 0,
                            'quiz_scores': []
                        }
                return False
            elif 'play music' in command or 'play random music' in command:
                    try:
                        music_dir = r'C:\Users\HP\Music'
                        songs = os.listdir(music_dir)
                        if songs:
                            random_song = np.random.choice(songs)
                            os.startfile(os.path.join(music_dir, random_song))
                            self.speak(f'Playing {random_song}')
                        else:
                            self.speak("No music files found in the directory.")
                    except Exception as e:
                        self.speak(f"Failed to play random music: {e}")    
            elif 'study mode off' in command:
                if self.study_mode:
                    self.study_mode = False
                    self.current_topic = None
                    self.speak("Study mode deactivated. Great job with your studies!")
                else:
                    self.speak("Study mode is not currently active.")
                return False
            elif 'quiz' in command:
                self.show_quiz_window()
            elif 'next' in command:
                current_level = self.study_progress[self.current_topic]['level']
                self.study_progress[self.current_topic]['level'] = current_level + 1
                self.study_progress[self.current_topic]['completed_tasks'] += 1
                self.speak(f"Moving to level {current_level + 1} of {self.current_topic}")
                search_query = f"{self.current_topic} advanced tutorial level {current_level + 1}"
                google_url = f'https://www.google.com/search?q={urllib.parse.quote(search_query)}'
                youtube_url = f"https://www.youtube.com/results?search_query={urllib.parse.quote(search_query)}"
                webbrowser.open(google_url)
                webbrowser.open(youtube_url)
            elif 'complete task' in command and self.study_mode:
                if self.current_topic:
                    try:
                        self.speak("Would you like to take a quiz, get certification, or move to the next level?")
                        choice = self.take_command()
                        if choice is None:
                            self.speak("I didn't catch that. Please try again.")
                            return False
                        choice = choice.lower()
                        if 'quiz' in choice:
                            self.show_quiz_window()
                        elif 'certification' in choice or 'certificate' in choice:
                            self.speak("Starting certification exam for " + self.current_topic)
                            score = self.start_certification_exam()
                            if score >= 60:
                                self.speak("Congratulations! You passed the certification exam.")
                                self.generate_certificate()
                            else:
                                self.speak(f"You scored {score}%. You need 60% to pass. Please try again after more practice.")
                        elif 'next' in choice:
                            current_level = self.study_progress[self.current_topic]['level']
                            self.study_progress[self.current_topic]['level'] = current_level + 1
                            self.study_progress[self.current_topic]['completed_tasks'] += 1
                            self.speak(f"Moving to level {current_level + 1} of {self.current_topic}")
                            search_query = f"{self.current_topic} advanced tutorial level {current_level + 1}"
                            google_url = f'https://www.google.com/search?q={urllib.parse.quote(search_query)}'
                            youtube_url = f"https://www.youtube.com/results?search_query={urllib.parse.quote(search_query)}"
                            webbrowser.open(google_url)
                            webbrowser.open(youtube_url)
                        else:
                            self.speak("Please say either 'quiz', 'certification', or 'next' to proceed.")
                    except Exception as e:
                        self.handle_errors(e, "Error processing complete task command")
                else:
                    self.speak("No topic is currently selected in study mode.")
                return False
            elif any(phrase in command for phrase in ['what do you see', 'analyze screen', 'describe screen', 'what is on screen', 'what is on the screen']):
                self.speak("Ok Sir")
                description = self.analyze_screen()
                self.speak(description)
                return False
            elif 'reset conversation' in command:
                if hasattr(self, 'chat_window'):
                    self.chat_window.destroy()
                self.root.after(0, self.show_chats_window)
                self.speak("Conversation has been reset")
                return False
            elif 'analyse database' in command or 'analyze database' in command or 'database analysis' in command:
                self.speak("Opening database analysis window")
                self.root.after(0, self.analyze_database)
                return False
            elif any(cmd in command for cmd in ['open google']):
                            webbrowser.open('https://www.google.com')
                            self.speak("Opening Google")
                            return False
            elif 'search for' in command:
                search_query = command.replace('search for', '').strip()
                if search_query:
                    search_url = f'https://www.google.com/search?q={urllib.parse.quote(search_query)}'
                    webbrowser.open(search_url)
                    self.speak(f"Searching Google for {search_query}")
                    return False
            elif 'open youtube' in command:
                webbrowser.open('https://www.youtube.com')
                self.speak("Opening YouTube")
                return False
            elif 'open wikipedia' in command:
                webbrowser.open('https://www.wikipedia.org')
                self.speak("Opening Wikipedia")
                return False
            elif any(cmd in command for cmd in ['programming window','programming assistant','programming']):
                self.speak("Opening the programming assistant")
                try:
                    import google.generativeai as genai
                    self.root.after(0, self.show_programming_window)
                except ImportError:
                    self.speak("Please install the required package google-generativeai first")
                    logging.error("google-generativeai package not installed")
                return False
            elif any(exit_cmd in command for exit_cmd in ['exit','you can leave','quit', 'goodbye', 'bye','see you later','you can','see you again']):
                self.speak("Thank you sir,see you soon")
                self.exit_assistant()
                return True
            elif 'brightness' in command:
                self.control_brightness(command)
                return False
            elif 'volume' in command:
                self.control_volume(command)
                return False
            elif 'battery' in command:
                self.get_battery_info()
                return False
            elif 'storage' in command or 'disk space' in command:
                self.get_storage_info()
                return False
            elif any(word in command for word in ['time', 'date', 'day']):
                self.get_time_date()
                return False
            elif command.startswith('type '):
                self.type_and_enter(command)
                return False
            elif command in ['stop typing', 'end typing', 'exit typing']:
                self.stop_typing()
                return False
            elif 'contact manager' in command or 'show contacts' in command:
                self.speak("Opening contact manager")
                self.root.after(0, self.show_contact_manager)
                return False
            elif 'tab' in command:
                tab_number = None
                if 'first' in command or '1st' in command or 'one' in command:
                    tab_number = 1
                elif 'second' in command or '2nd' in command or 'two' in command:
                    tab_number = 2
                elif 'third' in command or '3rd' in command or 'three' in command:
                    tab_number = 3
                elif 'fourth' in command or '4th' in command or 'four' in command:
                    tab_number = 4
                else:
                    numbers = [int(s) for s in command.split() if s.isdigit()]
                    if numbers:
                        tab_number = numbers[0]
                if tab_number is not None:
                    self.switch_tab(tab_number - 1)
                    self.speak(f"Switched to tab {tab_number}")
                    return False
            elif 'send a message to' in command:
                contact_name = command.replace('send a message to', '').strip()
                if contact_name:
                    self.whatsapp_interaction('message', contact_name)
                    return True
            elif 'video call' in command:
                contact_name = command.replace('video call', '').strip()
                if contact_name:
                    self.whatsapp_interaction('video', contact_name)
                    return True
            elif 'whatsapp call' in command:
                contact_name = command.replace('whatsapp call', '').strip()
                if contact_name:
                    self.whatsapp_interaction('call', contact_name)
                    return True
            elif "auto login" in command:
                website = command.replace("auto login", "").strip()
                if website:
                    if self.is_valid_website(website):
                        self.auto_login(website)
                    else:
                        self.speak(f"I cannot log into {website}. Please specify a valid website.")
                else:
                    self.speak("Please specify the website for auto login.")
                return False
            elif "password manager" in command:
                self.show_password_manager()
                return False
            elif command in ['delete text','delete the text']:
                pyautogui.hotkey('delete')
                return False
            is_scheduled_task = False
            if command.startswith("__scheduled__:"):
                is_scheduled_task = True
                command = command.replace("__scheduled__:", "", 1).strip()
                logging.info(f"Processing scheduled task command: '{command}'")
            if is_scheduled_task:
                return self._execute_direct_command(command)
            elif not is_scheduled_task and 'schedule' in command and ('task' in command or 'reminder' in command or '' in command):
                max_retries = 3
                retry_count = 0
                while retry_count < max_retries:
                    self.speak("How long until the tasks should start?")
                    duration = self.take_command()
                    if duration and duration.lower() != 'none':
                        seconds = self.parse_duration(duration)
                        if seconds > 0:
                            self.speak("What tasks would you like to schedule?.")
                            tasks = self.take_command()
                            if tasks and tasks.lower() != 'none':
                                task_list = [t.strip() for t in tasks.split(',') if t.strip()]
                                if task_list:
                                    self.show_task_scheduler(preset_time=seconds, preset_tasks=','.join(task_list))
                                    return False
                            self.show_task_scheduler(preset_time=seconds)
                            return False
                        else:
                            retry_count += 1
                            if retry_count < max_retries:
                                self.speak("I couldn't understand the time duration. Please try saying it like '5 minutes' or '2 hours' or '30 seconds'.")
                            else:
                                self.speak("I'm still having trouble understanding the time. Would you like to enter it manually in the scheduler window?")
                                response = self.take_command()
                                if any(word in response.lower() for word in ['yes', 'yeah', 'sure', 'okay']):
                                    self.show_task_scheduler()
                                else:
                                    self.speak("Task scheduling cancelled.")
                                return False
                    else:
                        retry_count += 1
                        if retry_count < max_retries:
                            self.speak("I didn't catch that. Please try again with the duration.")
                        else:
                            self.speak("I'm having trouble understanding. Would you like to open the scheduler window to enter the time manually?")
                            response = self.take_command()
                            if any(word in response.lower() for word in ['yes', 'yeah', 'sure', 'okay']):
                                self.show_task_scheduler()
                            else:
                                self.speak("Task scheduling cancelled.")
                            return False
                return False
            elif 'new text file' in command:
                pyautogui.hotkey('ctrl', 'n')
                self.speak("Done")
                return False
            elif 'new text window' in command:
                pyautogui.hotkey('ctrl', 'shift', 'n')
                self.speak("Done")
                return False
            elif 'open file' in command:
                pyautogui.hotkey('ctrl', 'o')
                self.speak("Done")
                return False
            elif 'data analysis' in command:
                self.speak("Here You can analyse the Data")
                self.root.after(0, self.show_data_analysis_window)
                return False
            elif 'open folder' in command:
                pyautogui.hotkey('ctrl', 'shift', 'o')
                self.speak("Done")
                return False
            elif 'translate' in command:
                target_lang = None
                for lang_name in self.language_dict.keys():
                    if f"to {lang_name.lower()}" in command:
                        target_lang = lang_name
                        break
                self.show_translation_comparison(target_lang)
                return False
            elif 'undo' in command:
                pyautogui.hotkey('ctrl', 'z')
                self.speak("Done")
                return False
            elif 'select all' in command:
                pyautogui.hotkey('ctrl', 'a')
                self.speak("Done")
                return False
            elif any(phrase in command for phrase in ['program for highlighted', 'create a program']):
                self.show_program_output()
                return False 
            elif 'run window' in command:
                pyautogui.hotkey('win', 'r')
                self.speak("Done")
                return False
            elif 'copy' in command:
                pyautogui.hotkey('ctrl', 'c')
                self.speak("Done")
                return False
            elif 'copy all' in command:
                pyautogui.hotkey('ctrl', 'a')
                pyautogui.hotkey('ctrl', 'c')
                self.speak("Done")
                return False
            elif 'paste' in command:
                pyautogui.hotkey('ctrl', 'v')
                self.speak("Done")
                return False
            elif 'cut' in command:
                pyautogui.hotkey('ctrl', 'x')
                self.speak("Done")
                return False
            elif 'find' in command:
                pyautogui.hotkey('ctrl', 'f')
                self.speak("Done")
                return False
            elif 'close window' in command:
                pyautogui.hotkey('alt', 'f4')
                self.speak("Done")
                return False
            elif 'enter' in command:
                pyautogui.hotkey('enter')
                self.speak("Done")
                return False
            elif command in ['check properties','check property','check property of']:
                pyautogui.hotkey('alt', 'enter')
                self.speak("Done")
                return False
            elif any(phrase in command for phrase in ['maximize all', 'maximize the all window', 'maximize all windows', 'maximize window', 'maximize the window', 'restore all window', 'restore window', 'restore all']):
                pyautogui.hotkey('win', 'shift', 'm')
                self.speak("All windows maximized")
                return False
            elif any(phrase in command for phrase in ['save notes in file','save note', 'save this note', 'save the note', 'save notes']):
                try:
                    notes_dir = r"C:\Users\HP\\Desktop\Notes"
                    now = datetime.datetime.now()
                    week_num = now.isocalendar()[1]
                    weekday = now.strftime('%A')
                    week_dir = os.path.join(notes_dir, f"Week_{week_num}")
                    day_dir = os.path.join(week_dir, weekday)
                    os.makedirs(day_dir, exist_ok=True)
                    active_window = pyautogui.getActiveWindow()
                    if "Notepad" in str(active_window.title):
                        pyautogui.hotkey('ctrl', 'a')
                        time.sleep(0.2)
                        pyautogui.hotkey('ctrl', 'c')
                        time.sleep(0.2)
                        import win32clipboard
                        win32clipboard.OpenClipboard()
                        try:
                            text = win32clipboard.GetClipboardData()
                        except:
                            text = ""
                        finally:
                            win32clipboard.CloseClipboard()
                        if text.strip():
                            filename = self.generate_filename_from_content(text)
                            filename = self.sanitize_filename(filename)
                            filepath = os.path.join(day_dir, filename)
                            with open(filepath, 'w', encoding='utf-8') as f:
                                f.write(text)
                            self.speak(f"Notes saved as {filename} in {weekday} folder for Week {week_num}")
                        else:
                            self.speak("No content to save")
                    else:
                        self.speak("Please make sure Notepad is active")
                except Exception as e:
                    logging.error(f"Error saving notes: {str(e)}")
                    self.speak("Sorry, I couldn't save the notes")
                return False
            elif 'switch window' in command:
                pyautogui.hotkey('alt', 'tab')
                self.speak("Done")
                return False
            elif 'minimise all' in command or 'minimise the all window' in command or 'minimise the window' in command or 'minimise window' in command:
                pyautogui.hotkey('win', 'm')
                self.speak("Done")
                return False
            elif 'task manager' in command:
                pyautogui.hotkey('ctrl', 'shift', 'esc')
                self.speak("Done")
                return False
            elif 'lock screen' in command:
                pyautogui.hotkey('win', 'l')
                self.speak("Done")
                return False
            elif 'search youtube' in command or 'youtube search' in command or 'youtube' in command:
                search_query = command.replace('search youtube', '').replace('youtube search', '').replace('youtube', '').strip()
                if search_query:
                    try:
                        search_url = f"https://www.youtube.com/results?search_query={urllib.parse.quote(search_query)}"
                        html = urllib.request.urlopen(search_url)
                        video_ids = re.findall(r"watch\?v=(\S{11})", html.read().decode())
                        if video_ids:
                            first_video_url = f"https://www.youtube.com/watch?v={video_ids[0]}"
                            webbrowser.open(first_video_url)
                            self.speak(f"Playing video for {search_query} on YouTube")
                            pyautogui.hotkey('enter')
                        else:
                            self.speak("No videos found for your search")
                    except Exception as e:
                        logging.error(f"Error searching YouTube: {str(e)}")
                        self.speak("Sorry, I encountered an error while searching YouTube")
                else:
                    self.speak("Please specify what to search on YouTube")
                return False
            elif any(cmd in command for cmd in ['chatbot', 'gemini chat','gemini','chat box' ]):
                self.speak("Opening ANIS chat")
                self.root.after(0, self.show_chat_window)
                return False
            elif 'translate window' in command or 'translator' in command:
                self.speak("Opening translator")
                self.root.after(0, self.show_translator_window)
                return False
            elif 'make notes' in command or 'summary' in command:
                self.speak("Opening summarizer window")
                self.root.after(0, self.show_summarizer_window)
                return False
            elif 'show email' in command or 'show mail' in command or 'send email' in command or 'send mail' in command or 'send a mail' in command or'send a email' in command or 'send an email' in command:
                self.root.after(0, self.show_email_window)
                return False
            elif 'open notes' in command:
                self.root.after(0, self.show_notes_window)
                return False
            elif 'check system health' in command or 'check system' in command or 'check system status' in command or 'check system information' in command or 'check system details' in command or 'check system status' in command or 'check system information' in command or 'check system details' in command:
                self.speak("Checking system health")
                self.check_system_health()
                return False
            elif any(phrase in command for phrase in ['search and open', 'find and open', 'search for and open']):
                search_term = command
                for phrase in ['search and open', 'find and open', 'search for and open']:
                    search_term = search_term.replace(phrase, '').strip()
                if search_term:
                    self.search_and_click_first_result(search_term)
                else:
                    self.speak("Please specify what to search for")
                return False
            elif any(phrase in command for phrase in ['save notes in file', 'save note', 'save this note', 'save the note', 'save notes']):
                try:
                    notes_dir = r"C:\Users\HP\Desktop\Ai\Notes"
                    now = datetime.datetime.now()
                    week_num = now.isocalendar()[1]
                    weekday = now.strftime('%A')
                    week_dir = os.path.join(notes_dir, f"Week_{week_num}")
                    day_dir = os.path.join(week_dir, weekday)
                    os.makedirs(day_dir, exist_ok=True)
                    timestamp = now.strftime('%Y%m%d_%H%M%S')
                    filename = f"note_{timestamp}.txt"
                    filepath = os.path.join(day_dir, filename)
                    active_window = pyautogui.getActiveWindow()
                    if "Notepad" in active_window.title:
                        pyautogui.hotkey('ctrl', 'a')
                        time.sleep(0.1)
                        pyautogui.hotkey('ctrl', 'c')
                        time.sleep(0.1)
                        text = self.window.clipboard_get()
                        with open(filepath, 'w', encoding='utf-8') as f:
                            f.write(text)
                        self.speak(f"Notes saved in {weekday} folder for Week {week_num}")
                    else:
                        self.speak("Please make sure Notepad is active")
                except Exception as e:
                    logging.error(f"Error saving notes: {str(e)}")
                    self.speak("Sorry, I couldn't save the notes")
                return False
            elif 'open' in command:
                if self.open_system_app(command):
                    return False
                search_term = command.replace('open', '').strip()
                if search_term:
                    self.search_and_click_first_result(search_term)
                return False
            elif 'take screenshot' in command:
                    self.speak("Taking screenshot")
                    self.take_screenshot()
                    return False
            elif any(phrase in command for phrase in ['save highlighted text','save the highlights','save the hightlight','save selection', 'save selected text', 'save highlight', 'save the highlighted words']):
                try:
                    notes_dir = r"C:\Users\HP\Desktop\Ai\Important Notes"
                    if not os.path.exists(notes_dir):
                        os.makedirs(notes_dir)
                    now = datetime.datetime.now()
                    week_num = now.isocalendar()[1]
                    weekday = now.strftime('%A')
                    week_dir = os.path.join(notes_dir, f"Week_{week_num}")
                    day_dir = os.path.join(week_dir, weekday)
                    os.makedirs(day_dir, exist_ok=True)
                    pyautogui.hotkey('ctrl', 'c')
                    time.sleep(0.5)
                    try:
                        selected_text = self.root.clipboard_get()
                    except:
                        self.speak("No text was selected")
                        return False
                    if not selected_text or selected_text.strip() == "":
                        self.speak("Please select some text first")
                        return False
                    try:
                        import google.generativeai as genai
                        genai.configure(api_key="AIzaSyAqlSk_zL7ID0a_tiBP_E6sIurmXB43F4k")
                        model = genai.GenerativeModel('gemini-2.0-flash')
                        prompt = f"Generate a short, descriptive filename (2-3 words) for this text: {selected_text[:200]}"
                        response = model.generate_content(prompt)
                        suggested_name = response.text.strip()
                        timestamp = now.strftime('%Y%m%d_%H%M%S')
                        filename = f"{suggested_name}_{timestamp}.txt"
                    except Exception as e:
                        logging.error(f"Error generating filename: {str(e)}")
                        timestamp = now.strftime('%Y%m%d_%H%M%S')
                        filename = f"selected_text_{timestamp}.txt"
                    filename = self.sanitize_filename(filename)
                    filepath = os.path.join(day_dir, filename)
                    with open(filepath, 'w', encoding='utf-8') as f:
                        f.write(f"Date: {now.strftime('%Y-%m-%d %H:%M:%S')}\n")
                        f.write("-" * 50 + "\n\n")
                        f.write(selected_text)
                    if os.path.exists(filepath):
                        self.speak(f"Selected text saved successfully in {weekday} folder")
                        logging.info(f"Saved selected text to: {filepath}")
                    else:
                        self.speak("Failed to save the text")
                except Exception as e:
                    logging.error(f"Error saving selected text: {str(e)}")
                    self.speak("Sorry, I couldn't save the selected text")
                return False
            elif 'google search' in command or 'search google' in command:
                    search_query = command.replace('google search', '').replace('search google', '').replace('google', '').strip()
                    if search_query:
                        self.speak(f"Searching Google for {search_query}")
                        url = f"https://www.google.com/search?q={urllib.parse.quote(search_query)}"
                        webbrowser.open(url)
                    else:
                        self.speak("Opening Google")
                        webbrowser.open("https://www.google.com")
                    return False
            try:
                self.type_and_enter("Processing your request...")
                custom_prompt = self.load_prompt('voice_commands')
                if custom_prompt:
                    prompt = f"""GuideLines: {custom_prompt}Guidelines:
                - answer should be in two lines for any query
                - Use simple language
                - Avoid special characters or formatting
                - Use proper punctuation
                - Keep responses brief and direct
                -your name is A N I S AI integrated with Operation System With Help Of Gemini and under Developing by 
                Samiulla from 26 August 2024 don't  tell this in beginning of your response you should tell only when someone ask you name
                - No emojis or symbols
                - answer in a friendly manner
                - answer like a human
                - answer in a way that is easy to understand
                - answer in a way that is not too formal
                - answer in a way that is not too technical
                - if need techinical answer in techincal terms
                - answer should be short like 2 lines and concise
                - avoid geographical deep terms
                - try to answer in 2 lines simple and short
                - some of the time you can answer in a funny way
                \n\nUser Query: {command}"""
                else:
                    prompt = f"""Please provide a clear and concise response to: {command}
                    Guidelines:
                - Use simple language
                - Avoid special characters or formatting
                - Use proper punctuation
                - Keep responses brief and direct
                -your name is A N I S AI integrated with Operation System With Help Of Gemini and under Developing by 
                Samiulla from 26 August 2024 don't  tell this in beginning of your response you should tell only when someone ask you name
                - No emojis or symbols
                - answer in a friendly manner
                - answer like a human
                - answer in a way that is easy to understand
                - answer in a way that is not too formal
                - answer in a way that is not too technical
                - if need techinical answer in techincal terms
                - answer should be short like 2 lines and concise
                - avoid geographical deep terms
                - try to answer in 2 lines simple and short
                - some of the time you can answer in a funny way
                """
                response = self.chat.send_message(prompt)
                response_text = response.text.strip()
                response_text = response_text.replace('*', '').replace('#', '').replace('`', '')
                self.stop_typing()
                self.speak(response_text)
                self.log_interaction(command, response_text)
                return False
            except Exception as e:
                self.handle_errors(e, "Error getting Gemini response")
                self.speak("I'm having trouble understanding that right now. Could you please try again?")
                return False
        except Exception as e:
            self.handle_errors(e, "Error processing command")
            self.speak("Sorry, I encountered an error processing that command")
            return False
    def toggle_listening(self):
        """Toggle between play and pause states"""
        try:
            if not self.start_button._listening:
                self.start_button._listening = True
                self.start_button.configure(text="◼")
                self.listening_thread = threading.Thread(target=self.listen_for_commands)
                self.listening_thread.daemon = True
                self.listening_thread.start()
                self.status_label.configure(text="● Listening", text_color='#4CAF50')
            else:
                self.start_button._listening = False
                self.start_button.configure(text="▶")
                self.status_label.configure(text="Not Listening", text_color='#f44336')
        except Exception as e:
            logging.error(f"Error toggling listening state: {str(e)}")
            self.speak("Sorry, I encountered an error with the voice recognition")
    def start_listening(self):
        """Start the listening process"""
        try:
            self.start_button._listening = True
            self.start_button.configure(text="◼")
            self.start_color_transition(self.start_button, '#27ae60')
            self.start_button.master.configure(highlightbackground='#27ae60')
            self.listening_thread = threading.Thread(target=self.listen_for_commands, daemon=True)
            self.listening_thread.start()
            self.status_label.configure(text="● Listening", text_color='#4CAF50')
        except Exception as e:
            logging.error(f"Error in start_listening: {str(e)}")
            self.status_label.configure(text="● Not Listening", text_color='#ff4d4d')
    def listen_for_commands(self):
        """Continuously listen for commands until paused or exit command is given"""
        try:
            while self.start_button._listening:
                command = self.take_command()
                if command != 'none':
                    result = self.process_command(command)
                    if result is True:
                        self.start_button.configure(text="▶")
                        self.start_button._listening = False
                        self.status_label.configure(text="● Not Listening", text_color='#ff4d4d')
                        break
                    if self.start_button._listening:
                        self.start_button.configure(text="◼")
                        self.status_label.configure(text="● Listening", text_color='#4CAF50')
        except Exception as e:
            logging.error(f"Error in listen_for_commands: {str(e)}")
            self.speak("I encountered an error. Please try again.")
            self.start_button.configure(text="▶")
            self.start_button._listening = False
            self.status_label.configure(text="● Not Listening", text_color='#ff4d4d')
    def exit_assistant(self):
        """Safely exit the voice assistant"""
        try:
            self.start_button._listening = False
            if hasattr(self, '_after_id'):
                self.root.after_cancel(self._after_id)
            if hasattr(self, 'engine'):
                self.engine.stop()
            if hasattr(self, 'wave_vis'):
                self.wave_vis.cleanup()
            if hasattr(self, 'root') and self.root.winfo_exists():
                self.root.quit()
                self.root.destroy()
        except Exception as e:
            logging.error(f"Error during exit: {str(e)}")
            sys.exit(0)
    def log_interaction(self, user_input, assistant_response):
        try:
            timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            log_entry = (
                f"Time: {timestamp}\n"
                f"User: {user_input}\n"
                f"Assistant: {assistant_response}\n"
                f"{'-'*50}\n"
            )
            with open("interaction_log.txt", "a", encoding="utf-8") as log_file:
                log_file.write(log_entry)
        except Exception as e:
            logging.error(f"Logging error: {str(e)}")
    def handle_errors(self, error, context=""):
        """Generic error handler for the assistant"""
        error_message = f"Error in {context}: {str(error)}"
        logging.error(error_message)
        user_message = "I encountered an error. Please try again."
        self.speak(user_message)
        self.update_status(False)
        self.log_interaction(context, user_message)
    def process_text_command(self, text_input):
        """Process text commands without voice input"""
        try:
            if text_input:
                self.output_text.configure(state=NORMAL)
                self.output_text._textbox.insert(END, f"Text Input: {text_input}\n")
                self.output_text.configure(state=DISABLED)
                self.output_text.see(END)
                return self.process_command(text_input.lower())
            return False
        except Exception as e:
            self.handle_errors(e, "text command processing")
            return False
    def setup_keyboard_shortcuts(self):
        """Setup keyboard shortcuts for the assistant"""
        try:
            keyboard_thread = threading.Thread(
                target=self.handle_keyboard_shortcuts,
                daemon=True
            )
            keyboard_thread.start()
        except Exception as e:
            self.handle_errors(e, "keyboard shortcut setup")
    def speak(self, text, emotion=None):
        try:
            if hasattr(self, 'wave_vis') and self.wave_vis.is_active:
                self.wave_vis.speaking_animation()
            self.insert_with_timestamp(f"A.N.I.S: {text}", "assistant_msg")
            engine.say(text)
            engine.runAndWait()
            if hasattr(self, 'wave_vis') and self.wave_vis.is_active:
                self.wave_vis.stop_listening_animation()
        except Exception as e:
            logging.error(f"Error in speak method: {str(e)}")
    def take_command(self):
        r = sr.Recognizer()
        with sr.Microphone() as source:
            try:
                if hasattr(self, 'wave_vis') and self.wave_vis.is_active:
                    self.wave_vis.start_listening_animation()
                self.status_label.configure(text="Listening...")
                self.root.update()
                r.adjust_for_ambient_noise(source, duration=1)
                r.dynamic_energy_threshold = True
                r.pause_threshold = 0.8
                r.non_speaking_duration = 0.8
                r.phrase_threshold = 0.3
                audio = r.listen(source, timeout=10 ,phrase_time_limit=8)
                self.status_label.configure(text="Recognizing...")
                self.root.update()
                query = r.recognize_google(audio, language='en-in')
                self.insert_with_timestamp(f"You: {query}", "user_msg")
                return query.lower()
            except sr.WaitTimeoutError:
                return "none"
            except sr.UnknownValueError:
                return "none"
            except sr.RequestError:
                self.speak("Sorry, there was an error with the speech recognition service")
                return "none"
            except Exception as e:
                logging.error(f"Error in speech recognition: {str(e)}")
                return "none"
            finally:
                if hasattr(self, 'wave_vis') and self.wave_vis.is_active:
                    self.wave_vis.stop_listening_animation()
                self.status_label.configure(text="Not Listening")
    def update_system_stats(self):
        """Update system statistics"""
        try:
            cpu_percent = psutil.cpu_percent(interval=0.1)
            self.cpu_label.configure(text=f"CPU: {cpu_percent}%")
            memory = psutil.virtual_memory()
            self.memory_label.configure(text=f"RAM: {memory.percent}%")
            self.root.after(1000, self.update_system_stats)
        except Exception as e:
            logging.error(f"Error updating system stats: {str(e)}")
    def show_settings(self):
        """Show settings window"""
        settings_window = ctk.CTkToplevel(self.root)
        settings_window.title("Settings")
        settings_window.geometry("600x700")
        settings_window.configure(fg_color=self.themes[self.current_theme]["bg"])
        settings_window.transient(self.root)
        settings_window.grab_set()
        tabview = ctk.CTkTabview(
            settings_window,
            fg_color=self.themes[self.current_theme]["bg"],
            segmented_button_fg_color=self.themes[self.current_theme]["bg"],
            segmented_button_selected_color=self.themes[self.current_theme]["accent"],
            segmented_button_selected_hover_color=self.themes[self.current_theme]["secondary"],
            segmented_button_unselected_color=self.themes[self.current_theme]["bg"],
            segmented_button_unselected_hover_color=self.themes[self.current_theme]["highlight"]
        )
        tabview.pack(fill='both', expand=True, padx=10, pady=10)
        theme_tab = tabview.add("Themes")
        volume_tab = tabview.add("Volume")
        help_tab = tabview.add("Help")
        prompts_tab = tabview.add("Prompts")
        theme_frame = ctk.CTkFrame(
            theme_tab,
            fg_color=self.themes[self.current_theme]["bg"],
            corner_radius=0
        )
        theme_frame.pack(fill='both', expand=True, padx=20, pady=20)
        ctk.CTkLabel(
            theme_frame,
            text="Theme Selection",
            font=("Segoe UI", 16, "bold"),
            text_color=self.themes[self.current_theme]["accent"]
        ).pack(anchor='w', pady=(0, 20))
        theme_var = tk.StringVar(value=self.current_theme)
        for theme_name in self.themes.keys():
            button_container = ctk.CTkFrame(
                theme_frame,
                fg_color=self.themes[self.current_theme]["bg"],
                corner_radius=8
            )
            button_container.pack(fill='x', pady=5, padx=10)
            theme_button = ctk.CTkRadioButton(
                button_container,
                text=theme_name,
                variable=theme_var,
                value=theme_name,
                command=lambda t=theme_name: self.change_theme(t),
                fg_color=self.themes[self.current_theme]["accent"],
                border_color=self.themes[self.current_theme]["fg"],
                text_color=self.themes[self.current_theme]["fg"],
                hover_color=self.themes[self.current_theme]["secondary"]
            )
            theme_button.pack(fill='x', padx=15, pady=8)
            if theme_name == self.current_theme:
                theme_button.select()
        volume_frame = ctk.CTkFrame(
            volume_tab,
            fg_color=self.themes[self.current_theme]["bg"],
            corner_radius=0
        )
        volume_frame.pack(fill='both', expand=True, padx=20, pady=20)
        ctk.CTkLabel(
            volume_frame,
            text="Volume Settings",
            font=("Segoe UI", 16, "bold"),
            text_color=self.themes[self.current_theme]["accent"]
        ).pack(anchor='w', pady=(0, 20))
        ctk.CTkLabel(
            volume_frame,
            text="Assistant Voice Volume",
            font=("Segoe UI", 12),
            text_color=self.themes[self.current_theme]["fg"]
        ).pack(anchor='w', pady=(0, 10))
        volume_control_frame = ctk.CTkFrame(
            volume_frame,
            fg_color=self.themes[self.current_theme]["bg"],
            corner_radius=0
        )
        volume_control_frame.pack(fill='x', pady=10)
        self.volume_scale = ctk.CTkSlider(
            volume_control_frame,
            from_=0,
            to=100,
            orientation='horizontal',
            number_of_steps=100,
            command=self.change_volume
        )
        self.volume_scale.set(engine.getProperty('volume') * 100)
        self.volume_scale.pack(fill='x', pady=10)
        volume_labels_frame = ctk.CTkFrame(
            volume_control_frame,
            fg_color=self.themes[self.current_theme]["bg"],
            corner_radius=0
        )
        volume_labels_frame.pack(fill='x')
        ctk.CTkLabel(
            volume_labels_frame,
            text="0%",
            font=("Segoe UI", 11),
            text_color=self.themes[self.current_theme]["fg"]
        ).pack(side='left')
        ctk.CTkLabel(
            volume_labels_frame,
            text="100%",
            font=("Segoe UI", 11),
            text_color=self.themes[self.current_theme]["fg"]
        ).pack(side='right')
        help_frame = ctk.CTkFrame(
            help_tab,
            fg_color=self.themes[self.current_theme]["bg"],
            corner_radius=0
        )
        help_frame.pack(fill='both', expand=True, padx=20, pady=20)
        ctk.CTkLabel(
            help_frame,
            text="Available Commands",
            font=("Segoe UI", 16, "bold"),
            text_color=self.themes[self.current_theme]["accent"]
        ).pack(anchor='w', pady=(0, 20))
        help_frame_container = ctk.CTkFrame(
            help_frame,
            fg_color=self.themes[self.current_theme]["bg"],
            corner_radius=0
        )
        help_frame_container.pack(fill='both', expand=True)
        help_text = ctk.CTkTextbox(
            help_frame_container,
            height=300,
            fg_color=self.themes[self.current_theme]["text_bg"],
            text_color=self.themes[self.current_theme]["fg"],
            font=("Segoe UI", 11),
            corner_radius=8,
            border_width=0
        )
        help_text.pack(fill='both', expand=True)
        help_content = """
    🎯 Voice Commands:
    🔍 Search & Browse:
    • "wikipedia [topic]" - Search Wikipedia
    • "search youtube [query]" - Search on YouTube
    • "google search [query]" OR "search google [query]" - Search on Google
    • "open [website]" - Open websites (youtube/google/whatsapp/facebook/twitter/instagram)
    🖱️ Screen Interaction:
    • "click on [element]" - Click on screen elements
    • "type [text]" - Type at cursor position
    • "stop typing" - End typing mode
    🎮 System Controls:
    • "check system health" - View system status
    • "take screenshot" - Capture screen
    • "create backup" - Backup system data
    🎵 Media Controls:
    • "play music" - Play random music
    • "play [song name]" - Play specific song
    🌍 Translator:
    • "translate" - Open translator
    • "translate [text]" - Translate text
    y
    ⌨️ Keyboard Shortcuts:
    • Ctrl+Shift+A - Start listening
    • Ctrl+Shift+X - Exit assistant
    • Ctrl+Shift+S - Take screenshot
    • Ctrl+Shift+H - Check system health
    • Ctrl+Shift+B - Create backup
    ❌ Exit Commands:
    • "exit" or "quit" or "goodbye"
    • "you can leave"
    💡 Tips:
    • Click the play button (▶) to start
    • Click stop button (◼) to pause
    • Watch the wave animation for voice activity
    • Monitor system stats in top bar
    """
        help_text.insert('1.0', help_content)
        help_text.configure(state='disabled')
        prompts_main_frame = ctk.CTkScrollableFrame(
            prompts_tab,
            fg_color=self.themes[self.current_theme]["bg"],
            corner_radius=0
        )
        prompts_main_frame.pack(fill='both', expand=True)
        ctk.CTkLabel(
            prompts_main_frame,
            text="Custom Agent Prompt Guidelines",
            font=("Segoe UI", 16, "bold"),
            text_color=self.themes[self.current_theme]["accent"]
        ).pack(anchor='w', pady=(0, 20))
        prompt_sections = {
            'voice_commands': "Voice Commands Prompt",
            'chat_window': "Chat Window Prompt",
            'article_mode': "Article Mode Prompt",
            'programming': "Programming Assistant Prompt",
            'header_analysis': "Data Headers Analysis Prompt",
            'full_analysis': "Full Dataset Analysis Prompt",
            'commanding_prompt': "Commanding Prompt",
        }
        for key, title in prompt_sections.items():
            section_frame = ctk.CTkFrame(
                prompts_main_frame,
                fg_color=self.themes[self.current_theme]["bg"],
                corner_radius=0
            )
            section_frame.pack(fill='x', pady=10)
            ctk.CTkLabel(
                section_frame,
                text=title,
                font=("Segoe UI", 12),
                text_color=self.themes[self.current_theme]["fg"]
            ).pack(anchor='w')
            placeholder_text = {
                'voice_commands': "{command} - Voice command, {context} - Command context",
                'chat_window': "{message} - User message, {history} - Chat history",
                'article_mode': "{text} - Article content, {mode} - Analysis mode",
                'programming': "{code} - Code snippet, {language} - Programming language",
                'header_analysis': "{headers} - Dataset column names",
                'full_analysis': "{data_summary} - Dataset statistics",
                'commanding_prompt': "{commands} - Command, {contexts} - Command context",
            }
            ctk.CTkLabel(
                section_frame,
                text=placeholder_text[key],
                font=("Segoe UI", 9, "italic"),
                text_color=self.themes[self.current_theme]["secondary"]
            ).pack(anchor='w', pady=(0, 5))
            text_widget = ctk.CTkTextbox(
                section_frame,
                height=100,
                fg_color=self.themes[self.current_theme]["text_bg"],
                text_color=self.themes[self.current_theme]["fg"],
                font=("Segoe UI", 10),
                corner_radius=6
            )
            text_widget.pack(fill='x', pady=5)
            current_prompt = self.load_prompt(key)
            if current_prompt:
                text_widget.insert('1.0', current_prompt)
            else:
                default_prompts = {
                    'voice_commands': """Process this voice command and determine the appropriate action.
    Consider the context and user preferences.
    Command: {command}
    Context: {context}""",
                    'chat_window': """You are a helpful AI assistant. Respond to the user's message
    while maintaining context of the conversation.
    Message: {message}
    History: {history}""",
                    'article_mode': """Analyze this article and provide key insights.
    Include main points and relevant details.
    Text: {text}
    Mode: {mode}""",
                    'programming': """Review this code and provide:
    1. Code analysis
    2. Suggestions for improvement
    3. Best practices
    4. Potential issues
    Code: {code}
    Language: {language}""",
                    'header_analysis': """Analyze these dataset headers and provide:
    1. Brief explanation of each header
    2. Potential relationships between headers
    3. What insights could be derived
    Headers: {headers}""",
                    'full_analysis': """Analyze this dataset and provide:
    1. Key insights and patterns
    2. Data quality issues
    3. Notable correlations
    4. Machine learning potential
    Data Summary: {data_summary}""",
                    'correlation_analysis': """Analyze the correlation matrix and explain:
    1. Strong positive correlations
    2. Strong negative correlations
    3. Important feature relationships
    """
                }
                text_widget.insert('1.0', default_prompts.get(key, ""))
            button_frame = ctk.CTkFrame(
                section_frame,
                fg_color=self.themes[self.current_theme]["bg"],
                corner_radius=0
            )
            button_frame.pack(fill='x', pady=5)
            ctk.CTkButton(
                button_frame,
                text="Save Prompt",
                command=lambda k=key, t=text_widget: self.save_prompt(k, t.get('1.0', 'end-1c')),
                fg_color=self.themes[self.current_theme]["accent"],
                text_color="#ffffff",
                font=('Segoe UI', 10),
                corner_radius=6,
                hover_color=self.themes[self.current_theme]["highlight"]
            ).pack(side='right', padx=5)
            ctk.CTkButton(
                button_frame,
                text="Reset",
                command=lambda k=key, t=text_widget: self.reset_single_prompt(k, t),
                fg_color=self.themes[self.current_theme]["secondary"],
                text_color="#ffffff",
                font=('Segoe UI', 10),
                corner_radius=6,
                hover_color=self.themes[self.current_theme]["highlight"]
            ).pack(side='right', padx=5)
        def on_close():
            settings_window.destroy()
        ctk.CTkButton(
            settings_window,
            text="Close",
            command=on_close,
            fg_color=self.themes[self.current_theme]["accent"],
            text_color="#ffffff",
            font=('Segoe UI', 12, 'bold'),
            width=100,
            height=40,
            corner_radius=8,
            hover_color=self.themes[self.current_theme]["secondary"]
        ).pack(pady=20)
    def load_prompt(self, prompt_type):
        """Load prompt from database"""
        try:
            db_path = os.path.join(get_app_directory(), "configure", "prompts.db")
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            cursor.execute('''CREATE TABLE IF NOT EXISTS prompts
                            (type TEXT PRIMARY KEY, content TEXT)''')
            cursor.execute('SELECT content FROM prompts WHERE type = ?', (prompt_type,))
            result = cursor.fetchone()
            conn.close()
            return result[0] if result else ""
        except Exception as e:
            logging.error(f"Error loading prompt: {str(e)}")
            return ""
    def save_prompt(self, prompt_type, content):
        """Save prompt to database"""
        try:
            db_path = os.path.join(get_app_directory(), "configure", "prompts.db")
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            cursor.execute('''CREATE TABLE IF NOT EXISTS prompts
                            (type TEXT PRIMARY KEY, content TEXT)''')
            cursor.execute('''INSERT OR REPLACE INTO prompts (type, content)
                            VALUES (?, ?)''', (prompt_type, content))
            conn.commit()
            conn.close()
            messagebox.showinfo("Success", "Prompt saved successfully!")
        except Exception as e:
            logging.error(f"Error saving prompt: {str(e)}")
            messagebox.showerror("Error", f"Failed to save prompt: {str(e)}")
    def change_theme(self, theme_name):
        """Change the application theme"""
        if theme_name in self.themes:
            self.current_theme = theme_name
            self.apply_theme(theme_name)
    def apply_theme(self, theme_name):
        """Apply the selected theme to all widgets"""
        theme = self.themes[theme_name]
        self.root.configure(fg_color=theme["bg"])
        self.main_container.configure(fg_color=theme["bg"])
        self.banner_frame.configure(fg_color=theme["bg"])
        self.gradient_frame.configure(
            fg_color=theme["bg"],
            border_color=theme["highlight"]
        )
        self.wave_frame.configure(fg_color=theme["bg"])
        self.header_label.configure(
            fg_color=theme["bg"],
            text_color=theme["accent"]
        )
        self.subtitle_label.configure(
            fg_color=theme["bg"],
            text_color=theme["secondary"]
        )
        self.status_label.configure(
            fg_color=theme["bg"],
            text_color=theme["error"] if "Not Listening" in self.status_label.cget("text") else theme["success"]
        )
        self.settings_button.configure(
            fg_color=theme["bg"],
            text_color=theme["secondary"]
        )
        self.output_text.configure(
            fg_color=theme["text_bg"],
            text_color=theme["fg"]
        )
        self.cpu_label.configure(
            fg_color=theme["bg"],
            text_color=theme["secondary"]
        )
        self.memory_label.configure(
            fg_color=theme["bg"],
            text_color=theme["secondary"]
        )
        self.time_label.configure(
            fg_color=theme["bg"],
            text_color=theme["secondary"]
        )
    def open_system_app(self, command):
        """Open system applications with proper path handling"""
        office_paths = {
            '64bit': [
                "C:\\Program Files\\Microsoft Office\\root\\Office16",
                "C:\\Program Files\\Microsoft Office\\Office16",
                "C:\\Program Files (x86)\\Microsoft Office\\root\\Office16",
                "C:\\Program Files (x86)\\Microsoft Office\\Office16"
            ]
        }
        app_commands = {
            'word': {
                'name': 'Microsoft Word',
                'cmd': 'WINWORD.EXE',
                'paths': office_paths['64bit']
            },
            'excel': {
                'name': 'Microsoft Excel',
                'cmd': 'EXCEL.EXE',
                'paths': office_paths['64bit']
            },
            'powerpoint': {
                'name': 'Microsoft PowerPoint',
                'cmd': 'POWERPNT.EXE',
                'paths': office_paths['64bit']
            }
        }
        try:
            command = command.lower()
            if command in app_commands:
                app_info = app_commands[command]
                executable_found = False
                for base_path in app_info['paths']:
                    full_path = os.path.join(base_path, app_info['cmd'])
                    if os.path.exists(full_path):
                        subprocess.Popen(full_path)
                        executable_found = True
                        break
                if not executable_found:
                    error_msg = f"Could not find {app_info['name']}. Please verify it is installed correctly."
                    logging.error(error_msg)
                    messagebox.showerror("Application Error", error_msg)
            else:
                pass
        except Exception as e:
            error_msg = f"Error opening application: {str(e)}"
            logging.error(error_msg)
            messagebox.showerror("Error", error_msg)
    def search_and_click_first_result(self, search_term):
        """Ultra-fast Windows search and click"""
        try:
            pyautogui.PAUSE = 0.1
            pyautogui.MINIMUM_DURATION = 0
            pyautogui.hotkey('win', 's')
            time.sleep(0.1)
            pyautogui.write(search_term, interval=0.01)
            time.sleep(0.2)
            pyautogui.press('enter')
            self.speak(f"Opening {search_term}")
            return True
        except Exception as e:
            logging.error(f"Search error: {str(e)}")
            self.speak("Search error")
            return False
    def create_tron_style_frame(self, parent, has_glow=True):
        """Create a TRON Legacy styled frame"""
        frame = ctk.CTkFrame(
            parent,
            fg_color='#0c1221',  # Dark blue background
            border_color='#00f6ff',  # Neon blue border
            border_width=1 if has_glow else 0,
            corner_radius=0
        )
        if has_glow:
            glow = ctk.CTkFrame(
                frame,
                fg_color='#00f6ff',  # Neon blue glow
                height=2,
                corner_radius=0
            )
            glow.pack(fill='x', side='top')
            bottom_glow = ctk.CTkFrame(
                frame,
                fg_color='#00f6ff',  # Neon blue glow
                height=2,
                corner_radius=0
            )
            bottom_glow.pack(fill='x', side='bottom')
        return frame
    def create_tron_button(self, parent, text, command, width=None):
        """Create a TRON Legacy styled button"""
        button_frame = ctk.CTkFrame(
            parent,
            fg_color='#0c1221',
            border_color='#00f6ff',
            border_width=1,
            corner_radius=0
        )
        button = ctk.CTkButton(
            button_frame,
            text=text,
            command=command,
            fg_color='#0c1221',
            text_color='#00f6ff',
            font=('Orbitron', 10, 'bold'),
            corner_radius=0,
            hover_color='#0f1a2b',
            width=width if width else 100,
            height=30,
            border_width=0
        )
        button.pack(padx=1, pady=1)
        def on_enter(e):
            button_frame.configure(border_color='#00ffff')
            button.configure(fg_color='#0f1a2b')
        def on_leave(e):
            button_frame.configure(border_color='#00f6ff')
            button.configure(fg_color='#0c1221')
        button.bind('<Enter>', on_enter)
        button.bind('<Leave>', on_leave)
        return button_frame
    def change_volume(self, value):
        """Change assistant volume"""
        try:
            volume = float(value) / 100
            engine.setProperty('volume', volume)
        except Exception as e:
            logging.error(f"Error changing volume: {str(e)}")
    def get_storage_info(self):
        try:
            total, used, free = shutil.disk_usage("/")
            used_percentage = (used / total) * 100
            message = f"{used_percentage:.1f}% of storage is used"
            self.speak(message)
        except Exception as e:
            self.handle_errors(e, "Error getting storage info")
            self.speak("Sorry, I couldn't get the storage information")
    def get_time_date(self):
        """Get current time and date information"""
        try:
            now = datetime.datetime.now()
            date_str = now.strftime("%B %d, %Y")
            time_str = now.strftime("%I:%M %p")
            day_str = now.strftime("%A")
            self.speak(f"Today is {day_str}, {date_str}. The current time is {time_str}")
        except Exception as e:
            self.handle_errors(e, "time and date")
            self.speak("Sorry, I couldn't get the time and date information")
    def reset_single_prompt(self, key, text_widget):
        """Reset a single prompt to default value"""
        try:
            default_prompts = {
                'voice_commands': """Process this voice command and determine the appropriate action.
Consider the context and user preferences.
Command: {command}
Context: {context}""",
                'chat_window': """You are a helpful AI assistant. Respond to the user's message
while maintaining context of the conversation.
Message: {message}
History: {history}""",
                'article_mode': """Analyze this article and provide key insights.
Include main points and relevant details.
Text: {text}
Mode: {mode}""",
                'programming': """Review this code and provide:
1. Code analysis
2. Suggestions for improvement
3. Best practices
4. Potential issues
Code: {code}
Language: {language}""",
                'header_analysis': """Analyze these dataset headers and provide:
1. Brief explanation of each header
2. Potential relationships between headers
3. What insights could be derived
Headers: {headers}""",
                'full_analysis': """Analyze this dataset and provide:
1. Key insights and patterns
2. Data quality issues
3. Notable correlations
4. Machine learning potential
Data Summary: {data_summary}""",
                'correlation_analysis': """Analyze the correlation matrix and explain:
1. Strong positive correlations
2. Strong negative correlations
3. Important feature relationships
"""
            }
            text_widget.delete('1.0', 'end')
            text_widget.insert('1.0', default_prompts.get(key, ""))
            db_path = os.path.join(get_app_directory(), "configrue", "prompts.db")
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            cursor.execute('DELETE FROM prompts WHERE type = ?', (key,))
            conn.commit()
            conn.close()
            messagebox.showinfo("Success", f"Reset {key} prompt to default")
        except Exception as e:
            logging.error(f"Error resetting prompt: {str(e)}")
            messagebox.showerror("Error", f"Failed to reset prompt: {str(e)}")
    def _convert_tk_properties(self, **kwargs):
        """Convert tkinter properties to customtkinter equivalents"""
        property_map = {
            'bg': 'fg_color',
            'highlightbackground': 'border_color',
            'highlightthickness': 'border_width',
            'bd': 'border_width'
        }
        converted_kwargs = {}
        for key, value in kwargs.items():
            if key in property_map:
                converted_kwargs[property_map[key]] = value
            else:
                converted_kwargs[key] = value
        return converted_kwargs
    def cleanup(self):
        """Clean up the visualizer before window destruction"""
        with self._lock:
            self.is_active = False
            self.animation_running = False
            if self.after_id:
                self.after_cancel(self.after_id)
                self.after_id = None
            try:
                self.delete("all")
            except tk.TclError:
                pass
    def browse_and_process_file(self, command):
        """Browse for a file and process it based on the command"""
        file_path = filedialog.askopenfilename()
        if file_path:
            self.speak(f"Selected file: {os.path.basename(file_path)}")
            self.file_assistant.file_path_var.set(file_path)
            self.file_assistant.show_window()
    def show_recent_files(self):
        """Show recent files window"""
        self.file_assistant.show_recent_files_window()
        self.speak("Showing your recently processed files")
    def show_recent_files_window(self, initial_tab_name="All Files", filter_op_type=None, filter_date=None, filter_weekday=None):
        """Show a window with recently processed files by type, with optional filtering."""
        if hasattr(self, 'recent_window') and self.recent_window is not None and self.recent_window.winfo_exists():
            try:
                 self.recent_window.destroy()
            except tk.TclError:
                 pass
        self.recent_window = ctk.CTkToplevel(self.window or self.parent)
        self.recent_window.title("Recent Files")
        self.recent_window.geometry("700x550")
        if self.window and self.window.winfo_exists():
            self.recent_window.transient(self.window)
        else:
             self.recent_window.transient(self.parent)
        theme = self.assistant.themes[self.assistant.current_theme]
        self.recent_window.configure(fg_color=theme["bg"])
        self.recent_window.grab_set()
        filter_info_text = "Filters: None"
        filters_applied = []
        if filter_op_type:
            filters_applied.append(f"Type={filter_op_type}")
        if filter_date:
            filters_applied.append(f"Date={filter_date}")
        if filter_weekday:
            weekday_names = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
            weekday_int_map = {'1': 0, '2': 1, '3': 2, '4': 3, '5': 4, '6': 5, '0': 6}
            try:
                 filters_applied.append(f"Day={weekday_names[weekday_int_map[filter_weekday]]}")
            except KeyError:
                 filters_applied.append(f"Day=Invalid({filter_weekday})")
        if filters_applied:
            filter_info_text = "Filters: " + ", ".join(filters_applied)
        filter_label = ctk.CTkLabel(
             self.recent_window,
             text=filter_info_text,
             font=("Segoe UI", 10),
             text_color=theme["secondary"]
        )
        filter_label.pack(pady=(5, 0), padx=20, anchor='w')
        tabview = ctk.CTkTabview(
            self.recent_window,
            fg_color=theme["highlight"],
            segmented_button_fg_color=theme["bg"],
            segmented_button_selected_color=theme["accent"],
            segmented_button_unselected_color=theme["bg"]
        )
        tabview.pack(fill=tk.BOTH, expand=True, padx=20, pady=(5, 20))
        tab_names = ["All Files", "Summaries", "Converted", "Extracted", "Exported", "Opened", "Other"]
        tabs = {}
        for name in tab_names:
             tabs[name] = tabview.add(name)
        try:
            db_path = os.path.join(self.assistant_folder, "files_database.db")
            if not os.path.exists(db_path):
                conn = sqlite3.connect(db_path)
                cursor = conn.cursor()
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS processed_files (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        file_path TEXT,
                        file_name TEXT,
                        operation_type TEXT,
                        timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
                    )
                ''')
                conn.commit()
                conn.close()
                all_files_data = []
            else:
                conn = sqlite3.connect(db_path)
                conn.execute("PRAGMA journal_mode=WAL;")
                cursor = conn.cursor()
                query = "SELECT id, file_path, file_name, operation_type, timestamp FROM processed_files"
                conditions = []
                params = []
                if filter_op_type:
                     if filter_op_type.startswith('exported'):
                          conditions.append("operation_type LIKE ?")
                          params.append('exported_%')
                     elif filter_op_type == 'other':
                          known_types = ('summary', 'converted', 'extract', 'opened')
                          conditions.append(f"operation_type NOT LIKE 'exported_%' AND operation_type NOT IN ({','.join('?'*len(known_types))})")
                          params.extend(known_types)
                     else:
                         conditions.append("operation_type = ?")
                         params.append(filter_op_type)
                if filter_date:
                    conditions.append("DATE(timestamp) = DATE(?)")
                    params.append(filter_date)
                elif filter_weekday:
                    conditions.append("STRFTIME('%w', timestamp) = ?")
                    params.append(filter_weekday)
                if conditions:
                    query += " WHERE " + " AND ".join(conditions)
                query += " ORDER BY timestamp DESC LIMIT 100"
                cursor.execute(query, params)
                all_files_data = cursor.fetchall()
                conn.close()
            tab_data = {
                "All Files": [],
                "Summaries": [],
                "Converted": [],
                "Extracted": [],
                "Exported": [],
                "Opened": [],
                "Other": []
            }
            for file_row in all_files_data:
                file_id, file_path, file_name, op_type, timestamp = file_row
                op_type_lower = op_type.lower()
                tab_data["All Files"].append(file_row)
                if op_type_lower == 'summary':
                    tab_data["Summaries"].append(file_row)
                elif op_type_lower == 'converted':
                    tab_data["Converted"].append(file_row)
                elif op_type_lower == 'extract':
                    tab_data["Extracted"].append(file_row)
                elif op_type_lower.startswith('exported'):
                     tab_data["Exported"].append(file_row)
                elif op_type_lower == 'opened':
                     tab_data["Opened"].append(file_row)
                else:
                     tab_data["Other"].append(file_row)
            for tab_name in tab_names:
                 parent_tab = tabs.get(tab_name)
                 if parent_tab:
                     self.create_files_list(parent_tab, tab_data[tab_name], show_type=(tab_name == "All Files"))
            try:
                 if initial_tab_name in tab_names:
                     tabview.set(initial_tab_name)
                 else:
                     tabview.set("All Files")
            except Exception as tab_e:
                 logging.error(f"Error setting tab: {tab_e}")
                 tabview.set("All Files")
        except sqlite3.Error as db_e:
             logging.error(f"Database error in show_recent_files: {db_e}")
             error_frame = ctk.CTkFrame(tabs["All Files"], fg_color=theme["bg"])
             error_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
             ctk.CTkLabel(
                 error_frame,
                 text=f"Error loading files from database: {db_e}",
                 font=("Segoe UI", 14),
                 text_color=theme["error"]
             ).pack(pady=20)
        except Exception as e:
            logging.error(f"Error showing recent files window: {str(e)}")
            error_frame = ctk.CTkFrame(tabs["All Files"], fg_color=theme["bg"])
            error_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
            ctk.CTkLabel(
                error_frame,
                text=f"An unexpected error occurred: {str(e)}",
                font=("Segoe UI", 14),
                text_color=theme["error"]
            ).pack(pady=20)
class FileHandlingAssistant:
    def __init__(self, parent, assistant):
        self.parent = parent
        self.assistant = assistant
        self.window = None
        self.processed_files = []
        self.assistant_folder = os.path.join(get_app_directory(), "Assistant_Files")
        os.makedirs(self.assistant_folder, exist_ok=True)
        self.gemini_model = self.assistant.gemini_model
    def show_window(self):
        if self.window is not None and self.window.winfo_exists():
            self.window.lift()
            return
        self.window = ctk.CTkToplevel(self.parent)
        self.window.title("File Handling Assistant")
        self.window.geometry("1300x600")
        self.window.transient(self.parent)
        self.window.grab_set()
        self.status_var = tk.StringVar(value="Ready")
        theme = self.assistant.themes[self.assistant.current_theme]
        self.window.configure(fg_color=theme["bg"])
        self.main_frame = ctk.CTkFrame(
            self.window,
            fg_color=theme["bg"],
            corner_radius=10
        )
        self.main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        status_frame = ctk.CTkFrame(
            self.main_frame,
            fg_color=theme["highlight"],
            corner_radius=8,
            height=30
        )
        status_frame.pack(fill=tk.X, padx=10, pady=(0, 10))
        status_label = ctk.CTkLabel(
            status_frame,
            textvariable=self.status_var,
            font=("Segoe UI", 10),
            text_color=theme["fg"]
        )
        status_label.pack(side=tk.LEFT, padx=10, pady=5)
        title_label = ctk.CTkLabel(
            self.main_frame,
            text="File Handling Assistant",
            font=("Segoe UI", 24, "bold"),
            text_color=theme["accent"]
        )
        title_label.pack(pady=(0, 20))
        file_frame = ctk.CTkFrame(
            self.main_frame,
            fg_color=theme["highlight"],
            corner_radius=8
        )
        file_frame.pack(fill=tk.X, padx=10, pady=10)
        self.file_path_var = tk.StringVar()
        file_path_entry = ctk.CTkEntry(
            file_frame,
            textvariable=self.file_path_var,
            width=500,
            font=("Segoe UI", 12),
            placeholder_text="File path...",
            fg_color=theme["text_bg"],
            text_color=theme["fg"]
        )
        file_path_entry.pack(side=tk.LEFT, padx=10, pady=10, fill=tk.X, expand=True)
        browse_button = ctk.CTkButton(
            file_frame,
            text="Browse",
            font=("Segoe UI", 12),
            fg_color=theme["accent"],
            text_color="#ffffff",
            hover_color=theme["secondary"],
            command=self.browse_file
        )
        browse_button.pack(side=tk.RIGHT, padx=10, pady=10)
        commands_frame = ctk.CTkFrame(
            self.main_frame,
            fg_color=theme["highlight"],
            corner_radius=8
        )
        commands_frame.pack(fill=tk.X, padx=10, pady=10)
        command_label = ctk.CTkLabel(
            commands_frame,
            text="File Command:",
            font=("Segoe UI", 12, "bold"),
            text_color=theme["fg"]
        )
        command_label.pack(side=tk.LEFT, padx=10, pady=10)
        self.command_var = tk.StringVar()
        command_entry = ctk.CTkEntry(
            commands_frame,
            textvariable=self.command_var,
            width=400,
            font=("Segoe UI", 12),
            placeholder_text="Enter command (e.g., 'convert to pdf', 'summarize')",
            fg_color=theme["text_bg"],
            text_color=theme["fg"]
        )
        command_entry.pack(side=tk.LEFT, padx=10, pady=10, fill=tk.X, expand=True)
        voice_loop_button = ctk.CTkButton(
            commands_frame,
            text="Voice Mode",
            font=("Segoe UI", 12),
            width=100,
            fg_color=theme["success"],
            text_color="#ffffff",
            hover_color=theme["secondary"],
            command=self.voice_command_loop
        )
        voice_loop_button.pack(side=tk.RIGHT, padx=5, pady=10)
        voice_cmd_button = ctk.CTkButton(
            commands_frame,
            text="🎤",
            font=("Segoe UI", 16),
            width=40,
            fg_color=theme["accent"],
            text_color="#ffffff",
            hover_color=theme["secondary"],
            command=self.voice_command
        )
        voice_cmd_button.pack(side=tk.RIGHT, padx=5, pady=10)
        execute_button = ctk.CTkButton(
            commands_frame,
            text="Execute",
            font=("Segoe UI", 12),
            fg_color=theme["accent"],
            text_color="#ffffff",
            hover_color=theme["secondary"],
            command=self.process_file_command
        )
        execute_button.pack(side=tk.RIGHT, padx=10, pady=10)
        recent_files_button = ctk.CTkButton(
            commands_frame,
            text="Recent Files",
            font=("Segoe UI", 12),
            width=100,
            fg_color=theme["accent"],
            text_color="#ffffff",
            hover_color=theme["secondary"],
            command=self.show_recent_files_window
        )
        recent_files_button.pack(side=tk.LEFT, padx=10, pady=10)
        output_frame = ctk.CTkFrame(
            self.main_frame,
            fg_color=theme["highlight"],
            corner_radius=8
        )
        output_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        output_label = ctk.CTkLabel(
            output_frame,
            text="Output:",
            font=("Segoe UI", 12, "bold"),
            text_color=theme["fg"]
        )
        output_label.pack(anchor=tk.W, padx=10, pady=(10, 5))
        self.output_text = ctk.CTkTextbox(
            output_frame,
            font=("Segoe UI", 12),
            fg_color=theme["text_bg"],
            text_color=theme["fg"],
            corner_radius=6
        )
        self.output_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=(0, 10))
        recent_frame = ctk.CTkFrame(
            self.main_frame,
            fg_color=theme["highlight"],
            corner_radius=8
        )
        recent_frame.pack(fill=tk.X, padx=10, pady=10)
        recent_label = ctk.CTkLabel(
            recent_frame,
            text="Recent Files:",
            font=("Segoe UI", 12, "bold"),
            text_color=theme["fg"]
        )
        recent_label.pack(anchor=tk.W, padx=10, pady=(10, 5))
        self.recent_files_listbox = tk.Listbox(
            recent_frame,
            bg=theme["text_bg"],
            fg=theme["fg"],
            font=("Segoe UI", 11),
            height=4,
            selectbackground=theme["accent"]
        )
        self.recent_files_listbox.pack(fill=tk.X, padx=10, pady=(0, 10))
        self.recent_files_listbox.bind("<Double-Button-1>", self.select_recent_file)
        self.load_recent_files()
        self.status_var.set("Select a file and enter a command, or try voice commands")
        self.update_output("Welcome to File Handling Assistant!\nYou can:\n- Browse for a file\n- Use voice commands\n- Try commands like 'convert to pdf', 'summarize this file', etc.\n- View recently processed files")
    def voice_command(self):
        """Use a separate thread for voice command input"""
        if hasattr(self.assistant, 'start_button'):
            self.was_main_listening = False
            if hasattr(self.assistant.start_button, '_listening'):
                self.was_main_listening = self.assistant.start_button._listening
            if hasattr(self.assistant.start_button, '_listening'):
                self.assistant.start_button._listening = False
            self.assistant.start_button.configure(text="▶")
            self.assistant.status_label.configure(text="● Not Listening", text_color='#ff4d4d')
            if hasattr(self.assistant, 'wave_vis'):
                self.assistant.wave_vis.stop_listening_animation()
            time.sleep(0.1)
        self._voice_command_active = True
        for widget in self.window.winfo_children():
            if isinstance(widget, ctk.CTkFrame):
                for child in widget.winfo_children():
                    if isinstance(child, ctk.CTkFrame):
                        for button in child.winfo_children():
                            if isinstance(button, ctk.CTkButton) and button.cget("text") == "🎤":
                                button.configure(state=DISABLED)
                                self.voice_button_ref = button
                                break
        self.update_output("Listening for voice command...")
        self.status_var.set("Listening...")
        def voice_listening_thread():
            try:
                r = sr.Recognizer()
                with sr.Microphone() as source:
                    r.adjust_for_ambient_noise(source, duration=0.5)
                    self.window.after(0, lambda: self.status_var.set("Listening..."))
                    try:
                        audio = r.listen(source, timeout=5, phrase_time_limit=10)
                        self.window.after(0, lambda: self.status_var.set("Processing voice..."))
                        try:
                            command = r.recognize_google(audio)
                            if command:
                                self.window.after(0, lambda cmd=command: 
                                    self.update_output(f"Voice command received: {cmd}"))
                                self.window.after(0, lambda cmd=command: self.command_var.set(cmd))
                                self.window.after(0, lambda: self.process_voice_command(
                                    self.command_var.get(), self.file_path_var.get()))
                        except sr.UnknownValueError:
                            self.window.after(0, lambda: self.update_output("Voice not recognized"))
                        except sr.RequestError as e:
                            self.window.after(0, lambda err=e: 
                                self.update_output(f"Speech service error: {err}"))
                    except sr.WaitTimeoutError:
                        self.window.after(0, lambda: self.update_output("No speech detected within timeout"))
            except Exception as e:
                self.window.after(0, lambda err=e: self.update_output(f"Voice recognition error: {err}"))
            finally:
                if hasattr(self, 'voice_button_ref') and self.voice_button_ref:
                    try:
                        if self.voice_button_ref.winfo_exists():
                            self.voice_button_ref.configure(state=NORMAL)
                    except Exception:
                        pass
                self._voice_command_active = False
                self.window.after(0, self.restore_main_listening)
        voice_thread = threading.Thread(target=voice_listening_thread)
        voice_thread.daemon = True
        voice_thread.start()
    def restore_main_listening(self):
        """Restore the main assistant's listening state if it was active before"""
        self.status_var.set("Ready")
        self._voice_command_active = False
        self._voice_loop_active = False
        if hasattr(self, 'was_main_listening') and self.was_main_listening:
            try:
                if hasattr(self.assistant, 'start_button'):
                    time.sleep(0.2)
                    self.assistant.start_button._listening = True
                    self.assistant.start_button.configure(text="■")
                    self.assistant.status_label.configure(text="● Listening", text_color='#4CAF50')
                    if hasattr(self.assistant, 'wave_vis'):
                        self.assistant.wave_vis.start_listening_animation()
                    logging.info("Restored main assistant listening state")
            except Exception as e:
                logging.error(f"Error restoring main listening: {e}")
        self.was_main_listening = False
    def voice_command_loop(self):
        """Use a separate thread for continuous voice command input"""
        if hasattr(self.assistant, 'start_button'):
            self.was_main_listening = False
            if hasattr(self.assistant.start_button, '_listening'):
                self.was_main_listening = self.assistant.start_button._listening
            if hasattr(self.assistant.start_button, '_listening'):
                self.assistant.start_button._listening = False
            self.assistant.start_button.configure(text="▶")
            self.assistant.status_label.configure(text="● Not Listening", text_color='#ff4d4d')
            if hasattr(self.assistant, 'wave_vis'):
                self.assistant.wave_vis.stop_listening_animation()
            time.sleep(0.2)
        self._voice_loop_active = True
        self.update_output("Voice command mode activated - speak commands")
        self.status_var.set("Voice mode active")
        for widget in self.window.winfo_children():
            if isinstance(widget, ctk.CTkFrame):
                for child in widget.winfo_children():
                    if isinstance(child, ctk.CTkFrame):
                        for button in child.winfo_children():
                            if isinstance(button, ctk.CTkButton) and button.cget("text") == "Voice Mode":
                                button.configure(text="Stop Voice Mode", fg_color="#FF5252")
                                self.voice_loop_button_ref = button
                                break
        def voice_listening_thread():
            try:
                r = sr.Recognizer()
                r.pause_threshold = 0.8
                r.dynamic_energy_threshold = True
                while getattr(self, '_voice_loop_active', False):
                    try:
                        with sr.Microphone() as source:
                            r.adjust_for_ambient_noise(source, duration=0.5)
                            self.window.after(0, lambda: self.status_var.set("Listening for command..."))
                            try:
                                audio = r.listen(source, timeout=3, phrase_time_limit=5)
                                self.window.after(0, lambda: self.status_var.set("Processing voice..."))
                                try:
                                    command = r.recognize_google(audio, language='en-in')
                                    if command.lower() in ["stop", "stop voice mode", "exit voice mode", "exit"]:
                                        self.window.after(0, self.stop_voice_command_loop)
                                        break
                                    self.window.after(0, lambda cmd=command: self.command_var.set(cmd))
                                    self.window.after(0, lambda cmd=command: self.update_output(f"Received voice command: {cmd}"))
                                    file_path = self.file_path_var.get()
                                    if file_path and os.path.exists(file_path):
                                        self.window.after(0, lambda cmd=command, path=file_path: 
                                                        self.process_voice_command(cmd, path))
                                    else:
                                        self.window.after(0, lambda: self.update_output("Please select a file first"))
                                    time.sleep(1)
                                except sr.UnknownValueError:
                                    pass
                                except sr.RequestError:
                                    self.window.after(0, lambda: self.update_output(
                                        "Speech recognition service unavailable. Please try again."))
                                    time.sleep(2)
                            except sr.WaitTimeoutError:
                                pass
                    except Exception as loop_e:
                        self.window.after(0, lambda err=loop_e: 
                            self.update_output(f"Voice loop error: {err}"))
                        time.sleep(1)
                    time.sleep(0.2)
            except Exception as e:
                self.window.after(0, lambda err=e: 
                    self.update_output(f"Voice recognition error: {err}"))
            finally:
                self.window.after(0, self.reset_voice_ui)
                self.window.after(100, self.restore_main_listening)
        self.voice_loop_thread = threading.Thread(target=voice_listening_thread)
        self.voice_loop_thread.daemon = True
        self.voice_loop_thread.start()
    def stop_voice_command_loop(self):
        """Stop the continuous voice command loop"""
        self._voice_loop_active = False
        self.update_output("Voice command mode deactivated")
        self.status_var.set("Ready")
        self.reset_voice_ui()
    def process_file_command(self):
        file_path = self.file_path_var.get().strip()
        command = self.command_var.get().strip()
        recent_info = self._parse_recent_file_command(command)
        if recent_info["is_recent_command"]:
            if recent_info["is_open_request"]:
                self.window.after(0, lambda: self.open_recent_file_by_filter(
                    filter_op_type=recent_info["filter_op_type"],
                    filter_date=recent_info["filter_date"],
                    filter_weekday=recent_info["filter_weekday"]
                ))
                self.update_output(f"Opening most recent {recent_info['filter_op_type'] or 'file'}")
            else:
                self.window.after(0, lambda: self.show_recent_files_window(
                    initial_tab_name=recent_info["initial_tab"],
                    filter_op_type=recent_info["filter_op_type"],
                    filter_date=recent_info["filter_date"],
                    filter_weekday=recent_info["filter_weekday"]
                ))
                self.update_output(f"Showing recent files (Filter: Type={recent_info['filter_op_type']}, Date={recent_info['filter_date']}, Weekday={recent_info['filter_weekday']})")
            return
        command_lower = command.lower()
        if not file_path:
            self.update_output("Error: No file selected")
            return
        if not os.path.exists(file_path):
            self.update_output("Error: File not found")
            return
        if not command:
            self.update_output("Error: No command specified")
            return
        self.update_output(f"Processing command: {command} on file: {os.path.basename(file_path)}")
        self.status_var.set("Processing...")
        if any(export_word in command_lower for export_word in ["export", "save as", "output as"]):
            export_format = None
            cmd_parts = command_lower.split()
            for i, word in enumerate(cmd_parts):
                if word in ["to", "as", "in"] and i < len(cmd_parts) - 1:
                    export_format = cmd_parts[i+1]
                    break
            if export_format:
                new_command = f"convert to {export_format}"
                self.update_output(f"Interpreting export command as: {new_command}")
                command = new_command
                command_lower = command.lower()
        core_action = None
        if "convert" in command_lower:
            core_action = "convert"
        elif "summarize" in command_lower or "summarise" in command_lower:
            core_action = "summarize"
        elif "move" in command_lower:
            core_action = "move"
        elif "copy" in command_lower:
            core_action = "copy"
        elif "extract" in command_lower:
            core_action = "extract"
        if not any(keyword in command_lower for keyword in ["convert", "summarize", "summarise", "move", "copy", "extract"]):
            self.interpret_command(command, file_path)
            return
        def process_thread():
            try:
                action_handled = False
                if core_action == "convert":
                    self.convert_file(file_path, command)
                    action_handled = True
                elif core_action == "summarize":
                    self.summarize_file(file_path, command)
                    action_handled = True
                elif core_action == "move":
                    self.move_file(file_path, command)
                    action_handled = True
                elif core_action == "copy":
                    self.copy_file(file_path, command)
                    action_handled = True
                elif core_action == "extract":
                    self.extract_from_file(file_path, command)
                    action_handled = True
                if not action_handled:
                    if "export" in command_lower or "save as" in command_lower:
                        self.window.after(0, lambda: self.update_output("Command contains export/save but no clear action (summarize/extract). Attempting custom processing..."))
                        self.handle_custom_command(file_path, command)
                    else:
                        self.interpret_command(command, file_path)
                if core_action != "move":
                    if os.path.exists(file_path):
                        self.window.after(0, lambda: self.add_to_recent_files(file_path))
                self.window.after(0, lambda: self.status_var.set("Completed"))
            except Exception as e:
                error_msg = f"Error processing command '{command}': {str(e)}"
                logging.error(error_msg)
                self.window.after(0, lambda msg=error_msg: self.update_output(f"Error: {str(e)}"))
                self.window.after(0, lambda: self.status_var.set("Error"))
        processing_thread = threading.Thread(target=process_thread)
        processing_thread.daemon = True
        processing_thread.start()
    def update_output(self, text):
        self.output_text.configure(state=NORMAL)
        self.output_text.insert(END, f"{text}\n")
        self.output_text.configure(state=NORMAL)
        self.output_text.see(END)
        self.status_var.set(text.split("\n")[0] if "\n" in text else text)
    def convert_file(self, file_path, command):
        """Convert a file from one format to another"""
        target_format = None
        valid_formats = {
            'pdf': 'pdf', 'adobe': 'pdf',
            'txt': 'txt', 'text': 'txt', 'note': 'txt',
            'html': 'html', 'web': 'html', 'webpage': 'html',
            'md': 'md', 'markdown': 'md',
            'csv': 'csv', 'excel': 'csv', 'spreadsheet': 'csv',
            'json': 'json', 'data': 'json',
            'png': 'png', 'jpg': 'jpg', 'jpeg': 'jpg', 'webp': 'webp',
            'bmp': 'bmp', 'gif': 'gif', 'tiff': 'tiff',
            'docx': 'docx', 'document': 'docx', 'doc': 'docx', 'word': 'docx',
            'rtf': 'rtf'
        }
        image_extensions = ['.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff', '.webp']
        format_words = []
        cmd_parts = command.lower().split()
        for i, word in enumerate(cmd_parts):
            if word == "to" and i < len(cmd_parts) - 1:
                if cmd_parts[i+1] in valid_formats:
                    format_words.append(cmd_parts[i+1])
                elif i < len(cmd_parts) - 2 and f"{cmd_parts[i+1]} {cmd_parts[i+2]}" in valid_formats:
                    format_words.append(f"{cmd_parts[i+1]} {cmd_parts[i+2]}")
            elif word == "as" and i < len(cmd_parts) - 1:
                if cmd_parts[i+1] in valid_formats:
                    format_words.append(cmd_parts[i+1])
            elif "format" in word and i > 0:
                if cmd_parts[i-1] in valid_formats:
                    format_words.append(cmd_parts[i-1])
            elif word in valid_formats:
                format_words.append(word)
        for format_word in format_words:
            if format_word in valid_formats:
                target_format = valid_formats[format_word]
                break
        if not target_format:
            common_words = ["convert", "to", "into", "as", "a", "the", "file", "document", "format", "image", "picture"]
            potential_formats = [word for word in cmd_parts if word not in common_words and len(word) > 1]
            for word in potential_formats:
                if word in valid_formats:
                    target_format = valid_formats[word]
                    break
                elif len(potential_formats) == 1:
                    target_format = word
                    self.window.after(0, lambda fmt=word: self.update_output(
                        f"Warning: '{fmt}' might not be a standard format. Trying anyway."))
                    break
        if not target_format:
            self.window.after(0, lambda: self.update_output(
                "Error: Could not determine target format from command. Please specify clearly (e.g., 'convert to png', 'save as pdf')."))
            return
        file_name = os.path.basename(file_path)
        file_base, file_ext = os.path.splitext(file_name)
        output_path = os.path.join(self.assistant_folder, f"{file_base}.{target_format}")
        is_image = file_ext.lower() in image_extensions
        
        # DOCX specific handling
        if target_format == 'docx':
            self.window.after(0, lambda: self.update_output(f"Converting to DOCX (Word) format..."))
            try:
                if file_ext.lower() == '.txt' or file_ext.lower() == '.md' or file_ext.lower() == '.html':
                    try:
                        from docx import Document
                        document = Document()
                        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                            content = f.read()
                        for para in content.split('\n'):
                            if para.strip():
                                document.add_paragraph(para)
                        document.save(output_path)
                        self.window.after(0, lambda: self.update_output(f"File successfully converted to DOCX and saved to {output_path}"))
                        self.add_to_assistant_file(output_path, "converted")
                        return
                    except ImportError:
                        self.window.after(0, lambda: self.update_output("Error: python-docx library not found. Cannot convert to DOCX."))
                        return
                elif is_image:
                    try:
                        from docx import Document
                        import pytesseract
                        from PIL import Image
                        try:
                            img = Image.open(file_path)
                            if img.mode != 'L':
                                img = img.convert('L')
                            from PIL import ImageEnhance
                            enhancer = ImageEnhance.Contrast(img)
                            img = enhancer.enhance(2.0)
                            text = pytesseract.image_to_string(img)
                        except Exception as ocr_e:
                            self.window.after(0, lambda: self.update_output(f"Error during OCR: {str(ocr_e)}"))
                            text = "OCR failed. Image content could not be extracted."
                        document = Document()
                        document.add_heading(f"OCR Result: {file_name}", 0)
                        document.add_paragraph(text)
                        document.add_page_break()
                        document.add_heading("Original Image", 1)
                        try:
                            document.add_picture(file_path, width=4000000)
                        except:
                            document.add_paragraph("Could not embed original image")
                        document.save(output_path)
                        self.window.after(0, lambda: self.update_output(f"Image converted to DOCX with OCR at {output_path}"))
                        self.add_to_assistant_file(output_path, "converted")
                        return
                    except ImportError:
                        self.window.after(0, lambda: self.update_output("Error: Required libraries (python-docx/pytesseract) not found."))
                        return
                elif file_ext.lower() == '.pdf':
                    try:
                        from docx import Document
                        import PyPDF2
                        pdf_text = ""
                        with open(file_path, 'rb') as pdf_file:
                            pdf_reader = PyPDF2.PdfReader(pdf_file)
                            for page_num in range(len(pdf_reader.pages)):
                                page = pdf_reader.pages[page_num]
                                pdf_text += page.extract_text() + "\n\n"
                        document = Document()
                        document.add_heading(f"Converted from PDF: {file_name}", 0)
                        for para in pdf_text.split('\n'):
                            if para.strip():
                                document.add_paragraph(para)
                        document.save(output_path)
                        self.window.after(0, lambda: self.update_output(f"PDF converted to DOCX and saved to {output_path}"))
                        self.add_to_assistant_file(output_path, "converted")
                        return
                    except ImportError:
                        self.window.after(0, lambda: self.update_output("Error: Required libraries not found for PDF to DOCX conversion."))
                        return
                    except Exception as pdf_e:
                        self.window.after(0, lambda: self.update_output(f"Error converting PDF to DOCX: {str(pdf_e)}"))
                        return
                else:
                    self.window.after(0, lambda: self.update_output(f"Unsupported source format {file_ext} for DOCX conversion."))
                    return
            except Exception as docx_e:
                self.window.after(0, lambda: self.update_output(f"Error during DOCX conversion: {str(docx_e)}"))
                return
        
        # For images and PDF conversions    
        try:
            if is_image:
                self.window.after(0, lambda: self.update_output(f"Attempting image conversion to {target_format}..."))
                from PIL import Image
                img = Image.open(file_path)
                quality = 90
                for i, word in enumerate(cmd_parts):
                    if word == "quality" and i < len(cmd_parts) - 1:
                        try:
                            q = int(cmd_parts[i+1])
                            if 1 <= q <= 100:
                                quality = q
                                self.window.after(0, lambda q=q: self.update_output(f"Using quality setting: {q}"))
                        except ValueError:
                            pass
                if target_format.lower() in ['jpg', 'jpeg']:
                    if img.mode in ('RGBA', 'LA', 'P'):
                        bg = Image.new("RGB", img.size, (255, 255, 255))
                        bg.paste(img, (0, 0), img if img.mode == 'RGBA' else None)
                        img = bg
                    img.save(output_path, quality=quality, optimize=True)
                elif target_format.lower() == 'png':
                    if img.mode == 'RGBA':
                        img.save(output_path, optimize=True)
                    else:
                        converted = img.convert('RGBA')
                        converted.save(output_path, optimize=True)
                elif target_format.lower() == 'webp':
                    img.save(output_path, quality=quality, method=4)
                elif target_format.lower() == 'bmp':
                    if img.mode in ('RGBA', 'LA'):
                        bg = Image.new("RGB", img.size, (255, 255, 255))
                        bg.paste(img, (0, 0), img if img.mode == 'RGBA' else None)
                        img = bg
                    img.save(output_path)
                elif target_format.lower() == 'gif':
                    if img.mode != 'P':
                        img = img.convert('RGBA')
                        img = img.quantize(method=2)
                    img.save(output_path, optimize=True)
                elif target_format.lower() == 'tiff':
                    img.save(output_path, compression='tiff_lzw')
                elif target_format.lower() == 'pdf':
                    try:
                        from reportlab.lib.pagesizes import letter
                        from reportlab.pdfgen import canvas
                        from io import BytesIO
                        if img.mode != 'RGB' and img.mode != 'RGBA':
                            img = img.convert('RGB')
                        width, height = img.size
                        aspect = height / width
                        page_width, page_height = letter
                        pdf_width = page_width * 0.9
                        pdf_height = pdf_width * aspect
                        if pdf_height > page_height * 0.9:
                            pdf_height = page_height * 0.9
                            pdf_width = pdf_height / aspect
                        img_bytes = BytesIO()
                        img.save(img_bytes, format='PNG')
                        img_bytes.seek(0)
                        c = canvas.Canvas(output_path, pagesize=letter)
                        x = (page_width - pdf_width) / 2
                        y = (page_height - pdf_height) / 2
                        c.drawImage(img_bytes, x, y, width=pdf_width, height=pdf_height)
                        c.save()
                    except ImportError:
                        if img.mode != 'RGB':
                            img = img.convert('RGB')
                        img.save(output_path, "PDF", resolution=300.0)
                else:
                    img.save(output_path)
                success_msg = f"Image successfully converted and saved to {output_path}"
                self.window.after(0, lambda msg=success_msg: self.update_output(msg))
                self.add_to_assistant_file(output_path, "converted")
            elif file_ext.lower() in ['.txt', '.md', '.html', '.csv', '.json'] or target_format in ['txt', 'md', 'html', 'csv', 'json']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                if target_format == 'pdf':
                    try:
                        try:
                            from reportlab.lib.pagesizes import letter
                            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Preformatted
                            from reportlab.lib.styles import getSampleStyleSheet
                            from reportlab.lib.enums import TA_LEFT
                            from reportlab.pdfbase import pdfmetrics
                            from reportlab.pdfbase.ttfonts import TTFont
                            doc = SimpleDocTemplate(output_path, pagesize=letter)
                            styles = getSampleStyleSheet()
                            try:
                                font_path = "DejaVuSans.ttf"
                                if os.path.exists(font_path):
                                    pdfmetrics.registerFont(TTFont('DejaVuSans', font_path))
                                    body_font = 'DejaVuSans'
                                else:
                                    try:
                                        pdfmetrics.registerFont(TTFont('Arial', 'arialuni.ttf'))
                                        body_font = 'Arial'
                                    except:
                                        body_font = 'Helvetica'
                            except Exception as font_e:
                                logging.warning(f"Font registration failed: {font_e}, using Helvetica.")
                                body_font = 'Helvetica'
                            normal_style = styles['Normal']
                            normal_style.fontName = body_font
                            normal_style.fontSize = 10
                            normal_style.leading = 12
                            normal_style.alignment = TA_LEFT
                            code_style = styles['Code']
                            code_style.fontName = 'Courier'
                            code_style.fontSize = 9
                            code_style.leading = 11
                            code_style.alignment = TA_LEFT
                            story = []
                            lines = content.splitlines()
                            preformatted_block = []
                            for line in lines:
                                is_code_like = line.strip().startswith((' ', '\t', '#', '//', '/*', 'def ', 'class ', 'import '))
                                try:
                                    line_processed = line.encode('utf-8', 'replace').decode('utf-8', 'replace')
                                    line_processed = line_processed.replace('\t', '    ')
                                except Exception:
                                    line_processed = "".join(c if ord(c) < 128 else '?' for c in line)
                                if is_code_like:
                                    if preformatted_block:
                                        story.append(Paragraph("<br/>".join(preformatted_block), normal_style))
                                        story.append(Spacer(1, 6))
                                        preformatted_block = []
                                    story.append(Preformatted(line_processed, code_style))
                                    story.append(Spacer(1, 0))
                                else:
                                    preformatted_block.append(line_processed)
                            if preformatted_block:
                                story.append(Paragraph("<br/>".join(preformatted_block), normal_style))
                            doc.build(story)
                            success_msg = f"File converted to PDF (ReportLab) and saved to {output_path}"
                            self.window.after(0, lambda msg=success_msg: self.update_output(msg))
                        except ImportError:
                            from fpdf import FPDF
                            pdf = FPDF()
                            pdf.add_page()
                            try:
                                pdf.add_font('DejaVu', '', 'DejaVuSans.ttf', uni=True)
                                pdf.set_font('DejaVu', '', 12)
                            except RuntimeError:
                                pdf.set_font("Arial", size=12)
                            filtered_content = "".join(char if ord(char) < 256 else '?' for char in content)
                            pdf.multi_cell(0, 10, filtered_content)
                            pdf.output(output_path)
                            success_msg = f"File converted to PDF (FPDF) and saved to {output_path}"
                            self.window.after(0, lambda msg=success_msg: self.update_output(msg))
                    except ImportError:
                        error_msg = "Error: PDF libraries (ReportLab/FPDF) not installed. Please install with 'pip install fpdf reportlab'"
                        self.window.after(0, lambda msg=error_msg: self.update_output(msg))
                        return
                    except Exception as pdf_e:
                        error_msg = f"Error during PDF generation: {str(pdf_e)}"
                        self.window.after(0, lambda msg=error_msg: self.update_output(msg))
                        return
                else:
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(content)
                    success_msg = f"File converted and saved to {output_path}"
                    self.window.after(0, lambda msg=success_msg: self.update_output(msg))
                self.add_to_assistant_file(output_path, "converted")
                
            # === IMPROVED DOCX TO PDF CONVERSION ===
            elif target_format.lower() == 'pdf' and file_ext.lower() == '.docx':
                self.window.after(0, lambda: self.update_output("Converting DOCX to PDF..."))
                
                def docx_to_pdf_worker():
                    conversion_success = False
                    error_msg = ""
                    
                    # ATTEMPT 1: Try docx2pdf (most reliable but requires MS Word)
                    try:
                        from docx2pdf import convert
                        convert(file_path, output_path)
                        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                            self.window.after(0, lambda: self.update_output("DOCX converted to PDF using docx2pdf"))
                            return True
                    except ImportError:
                        error_msg = "docx2pdf not available"
                    except Exception as e:
                        error_msg = f"docx2pdf error: {str(e)}"
                    
                    # ATTEMPT 2: ReportLab conversion (high quality, no dependencies)
                    try:
                        from docx import Document
                        from reportlab.lib.pagesizes import letter
                        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                        from reportlab.lib.styles import getSampleStyleSheet
                        
                        # Extract docx text
                        doc = Document(file_path)
                        pdf = SimpleDocTemplate(output_path, pagesize=letter)
                        styles = getSampleStyleSheet()
                        story = []
                        
                        # Process paragraphs with proper styling
                        for para in doc.paragraphs:
                            if para.text.strip():
                                style_name = 'Normal'
                                if para.style.name.startswith('Heading'):
                                    if '1' in para.style.name:
                                        style_name = 'Heading1'
                                    elif '2' in para.style.name:
                                        style_name = 'Heading2'
                                    else:
                                        style_name = 'Heading3'
                                
                                # Clean text of any problematic characters
                                clean_text = para.text.encode('ascii', 'replace').decode('ascii')
                                p = Paragraph(clean_text, styles[style_name])
                                story.append(p)
                                story.append(Spacer(1, 6))
                        
                        # Build the PDF
                        if story:
                            pdf.build(story)
                            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                                self.window.after(0, lambda: self.update_output("DOCX converted to PDF using ReportLab"))
                                return True
                        else:
                            error_msg = "Document has no text content"
                    except ImportError:
                        error_msg = f"{error_msg}, ReportLab not available"
                    except Exception as e:
                        error_msg = f"{error_msg}, ReportLab error: {str(e)}"
                    
                    # ATTEMPT 3: FPDF (simpler, more compatible)
                    try:
                        from fpdf import FPDF
                        from docx import Document
                        
                        # Extract text content
                        doc = Document(file_path)
                        text_content = ""
                        for para in doc.paragraphs:
                            text_content += para.text + "\n\n"
                        
                        # Create PDF
                        pdf = FPDF()
                        pdf.add_page()
                        pdf.set_font("Arial", size=12)
                        
                        # Handle encoding issues
                        safe_text = "".join(c if ord(c) < 128 else ' ' for c in text_content)
                        
                        # Split text into manageable chunks
                        chunks = [safe_text[i:i+80] for i in range(0, len(safe_text), 80)]
                        for chunk in chunks:
                            if chunk.strip():
                                pdf.multi_cell(0, 10, chunk)
                        
                        pdf.output(output_path)
                        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                            self.window.after(0, lambda: self.update_output("DOCX converted to PDF using FPDF"))
                            return True
                    except ImportError:
                        error_msg = f"{error_msg}, FPDF not available"
                    except Exception as e:
                        error_msg = f"{error_msg}, FPDF error: {str(e)}"
                    
                    # If all methods failed
                    self.window.after(0, lambda: self.update_output(f"All PDF conversion methods failed: {error_msg}"))
                    return False
                
                # Run conversion in a separate thread with timeout
                conversion_thread = threading.Thread(target=docx_to_pdf_worker)
                conversion_thread.daemon = True
                conversion_thread.start()
                
                # Wait with timeout (30 seconds max)
                conversion_thread.join(timeout=30)
                
                if conversion_thread.is_alive():
                    self.window.after(0, lambda: self.update_output("Conversion process taking too long, terminating"))
                    return
                
                if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                    self.window.after(0, lambda: self.update_output(f"File successfully converted and saved to {output_path}"))
                    self.add_to_assistant_file(output_path, "converted")
                else:
                    self.window.after(0, lambda: self.update_output("Conversion failed or produced empty file"))
                    
        except Exception as e:
            error_msg = f"Error converting file: {str(e)}"
            logging.error(f"Conversion Error: {error_msg} for file {file_path} to {target_format}")
            self.window.after(0, lambda msg=error_msg: self.update_output(msg))
    def add_to_assistant_file(self, file_path, operation_type):
        """Add file to the assistant's database with metadata"""
        try:
            db_path = os.path.join(self.assistant_folder, "files_database.db")
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS processed_files (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    file_path TEXT,
                    file_name TEXT,
                    operation_type TEXT,
                    timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            file_name = os.path.basename(file_path)
            cursor.execute(
                "INSERT INTO processed_files (file_path, file_name, operation_type) VALUES (?, ?, ?)",
                (file_path, file_name, operation_type)
            )
            conn.commit()
            conn.close()
        except Exception as e:
            logging.error(f"Error adding file to database: {str(e)}")
    def summarize_file(self, file_path, command=None):
        file_name = os.path.basename(file_path)
        file_base, file_ext = os.path.splitext(file_name)
        export_format = None
        if command:
             match = re.search(r'(?:export|save)\s+(?:it|result|summary)\s+as\s+(\w+)', command.lower())
             if match:
                 export_format = match.group(1).lower()
                 self.window.after(0, lambda: self.update_output(f"Detected export request for summary: {export_format}"))
        output_base_path = os.path.join(self.assistant_folder, f"{file_base}_summary")
        content = ""
        try:
            self.window.after(0, lambda: self.update_output(f"Reading file for summarization: {file_name}"))
            file_ext_lower = file_ext.lower()
            if file_ext_lower in ['.txt', '.md', '.csv', '.json', '.html', '.xml', '.log']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                self.window.after(0, lambda: self.update_output("Text file read successfully."))
            elif file_ext_lower == '.docx':
                content = self._extract_text_from_docx(file_path)
                if not content or content.startswith("[Error"):
                    self.window.after(0, lambda: self.update_output("Failed to extract text from DOCX file for summary."))
                    return
                self.window.after(0, lambda: self.update_output("DOCX file processed for summary."))
            elif file_ext_lower == '.pdf':
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as pdf_file:
                        pdf_reader = PyPDF2.PdfReader(pdf_file)
                        for page_num in range(len(pdf_reader.pages)):
                            page = pdf_reader.pages[page_num]
                            content += page.extract_text() + "\n"
                    if not content.strip():
                         self.window.after(0, lambda: self.update_output("PDF text extraction yielded no content, attempting OCR..."))
                         content = self._extract_text_from_pdf_with_ocr(file_path)
                    self.window.after(0, lambda: self.update_output("PDF processed for summary."))
                except ImportError:
                     self.window.after(0, lambda: self.update_output("PyPDF2 not found for PDF summary. Install it."))
                     return
                except Exception as pdf_err:
                     self.window.after(0, lambda: self.update_output(f"Error reading PDF for summary: {pdf_err}"))
                     return
            elif file_ext_lower in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.gif', '.webp']:
                 content = self._extract_text_from_image(file_path)
                 self.window.after(0, lambda: self.update_output("Image OCR completed for summary."))
            else:
                 self.window.after(0, lambda: self.update_output(f"Attempting to read unsupported format {file_ext_lower} as text for summary."))
                 try:
                     with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                         content = f.read()
                 except Exception as read_err:
                     self.window.after(0, lambda: self.update_output(f"Could not read file as text: {read_err}"))
                     return
            if not content or not content.strip():
                 self.window.after(0, lambda: self.update_output("No content found in the file to summarize."))
                 return
            self.window.after(0, lambda: self.update_output("Summarizing file using AI..."))
            prompt = f"Summarize the following content concisely:\n\n{content[:15000]}"
            try:
                response = self.gemini_model.generate_content(prompt)
                summary = response.text.strip()
                if not summary:
                     self.window.after(0, lambda: self.update_output("AI returned an empty summary."))
                     return
                if export_format:
                    output_path = self._export_result(summary, output_base_path, export_format)
                    if output_path:
                        success_msg = f"Summary Preview:\n{summary[:500]}..."
                        self.window.after(0, lambda msg=success_msg: self.update_output(msg))
                else:
                    output_path_txt = output_base_path + ".txt"
                    with open(output_path_txt, 'w', encoding='utf-8') as f:
                        f.write(summary)
                    success_msg1 = f"File summarized and saved to {output_path_txt}"
                    success_msg2 = f"Summary Preview:\n{summary[:500]}..."
                    self.window.after(0, lambda msg=success_msg1: self.update_output(msg))
                    self.window.after(0, lambda msg=success_msg2: self.update_output(msg))
                    self.add_to_assistant_file(output_path_txt, "summary")
            except Exception as e:
                if "Content Policy Violation" in str(e):
                     error_msg = f"AI Error: Could not summarize due to content policy restrictions."
                elif "quota" in str(e).lower():
                     error_msg = f"AI Error: API quota likely exceeded."
                else:
                    error_msg = f"Error using AI to summarize: {str(e)}"
                self.window.after(0, lambda msg=error_msg: self.update_output(msg))
        except FileNotFoundError:
             self.window.after(0, lambda: self.update_output(f"Error: File not found at {file_path}"))
        except Exception as e:
            error_msg = f"Error reading file for summary: {str(e)}"
            self.window.after(0, lambda msg=error_msg: self.update_output(msg))
    def move_file(self, file_path, command):
        target_folder = None
        words = command.split()
        for i, word in enumerate(words):
            if word == "to" and i < len(words) - 1:
                target_folder = words[i+1]
                break
        if not target_folder:
            self.window.after(0, lambda: self.update_output("Error: Could not determine target folder"))
            return
        if target_folder.lower() == "assistant":
            target_folder = self.assistant_folder
        else:
            home_dir = os.path.expanduser("~")
            common_locations = [
                os.path.join(home_dir, "Desktop"),
                os.path.join(home_dir, "Documents"),
                os.path.join(home_dir, "Downloads"),
                os.path.join(home_dir, target_folder)
            ]
            for location in common_locations:
                if os.path.exists(location) and os.path.isdir(location):
                    target_folder = location
                    break
            if not os.path.exists(target_folder):
                self.window.after(0, lambda: self.update_output(f"Error: Target folder '{target_folder}' not found"))
                return
        file_name = os.path.basename(file_path)
        target_path = os.path.join(target_folder, file_name)
        try:
            shutil.move(file_path, target_path)
            self.window.after(0, lambda: self.update_output(f"File moved to {target_path}"))
        except Exception as e:
            self.window.after(0, lambda: self.update_output(f"Error moving file: {str(e)}"))
    def copy_file(self, file_path, command):
        target_folder = self.assistant_folder
        words = command.split()
        for i, word in enumerate(words):
            if word == "to" and i < len(words) - 1:
                folder_name = words[i+1]
                if folder_name.lower() != "assistant":
                    home_dir = os.path.expanduser("~")
                    common_locations = [
                        os.path.join(home_dir, "Desktop"),
                        os.path.join(home_dir, "Documents"),
                        os.path.join(home_dir, "Downloads"),
                        os.path.join(home_dir, folder_name)
                    ]
                    for location in common_locations:
                        if os.path.exists(location) and os.path.isdir(location):
                            target_folder = location
                            break
                    if not os.path.exists(target_folder):
                        self.window.after(0, lambda: self.update_output(f"Folder '{folder_name}' not found. Using Assistant folder instead."))
        file_name = os.path.basename(file_path)
        target_path = os.path.join(target_folder, file_name)
        try:
            shutil.copy2(file_path, target_path)
            self.window.after(0, lambda: self.update_output(f"File copied to {target_path}"))
        except Exception as e:
            self.window.after(0, lambda: self.update_output(f"Error copying file: {str(e)}"))
    def extract_from_file(self, file_path, command):
        extract_what = None
        export_format = None
        words = command.lower().split()
        try:
            extract_index = words.index("extract")
            potential_what = []
            for i in range(extract_index + 1, len(words)):
                if words[i] in ["export", "save"]:
                    break
                potential_what.append(words[i])
            extract_what = " ".join(potential_what).strip()
            export_match = re.search(r'(?:export|save)\s+(?:it|result|extraction)\s+as\s+(\w+)', command.lower())
            if export_match:
                export_format = export_match.group(1).lower()
                self.window.after(0, lambda: self.update_output(f"Detected export request for extraction: {export_format}"))
        except ValueError:
             self.window.after(0, lambda: self.update_output("Error: 'extract' keyword not found in command."))
             return
        except Exception as parse_e:
             self.window.after(0, lambda: self.update_output(f"Error parsing extract command: {parse_e}"))
             return
        if not extract_what:
            self.window.after(0, lambda: self.update_output("Error: Could not determine what to extract from the command."))
            return
        file_ext = os.path.splitext(file_path)[1].lower()
        content = ""
        extraction_source = "text"
        try:
            self.window.after(0, lambda: self.update_output(f"Reading file: {os.path.basename(file_path)}"))
            if file_ext in ['.txt', '.md', '.csv', '.json', '.html', '.xml', '.log']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                self.window.after(0, lambda: self.update_output("Text file processed successfully"))
            elif file_ext == '.docx':
                try:
                    content = self._extract_text_from_docx(file_path)
                    if not content or content.startswith("[Error"):
                        self.window.after(0, lambda: self.update_output("Failed to extract text from DOCX file"))
                        return
                    self.window.after(0, lambda: self.update_output("Word document processed successfully"))
                except Exception as docx_e:
                    self.window.after(0, lambda: self.update_output(f"Error processing DOCX file: {str(docx_e)}"))
                    return
            elif file_ext == '.pdf':
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as pdf_file:
                        try:
                            pdf_reader = PyPDF2.PdfReader(pdf_file)
                            content = ""
                            for page_num in range(len(pdf_reader.pages)):
                                page = pdf_reader.pages[page_num]
                                content += page.extract_text() + "\n"
                        except Exception as pdf_read_error:
                            self.window.after(0, lambda: self.update_output(f"Error reading PDF: {str(pdf_read_error)}"))
                            return
                    if not content.strip():
                        self.window.after(0, lambda: self.update_output("PDF may be scanned - attempting OCR..."))
                        try:
                            content = self._extract_text_from_pdf_with_ocr(file_path)
                        except Exception as ocr_error:
                            self.window.after(0, lambda: self.update_output(f"OCR failed: {str(ocr_error)}"))
                            return
                    self.window.after(0, lambda: self.update_output("PDF processed successfully"))
                except ImportError:
                    try:
                        import pdfplumber
                        with pdfplumber.open(file_path) as pdf:
                            content = ""
                            for page in pdf.pages:
                                content += page.extract_text() + "\n"
                        self.window.after(0, lambda: self.update_output("PDF processed using pdfplumber"))
                    except ImportError:
                        self.window.after(0, lambda: self.update_output("Error: Missing libraries for PDF processing. Install PyPDF2 or pdfplumber"))
                        return
                    except Exception as plumb_error:
                        self.window.after(0, lambda: self.update_output(f"Error with pdfplumber: {str(plumb_error)}"))
                        return
            elif file_ext == '.doc':
                try:
                    self.window.after(0, lambda: self.update_output("Attempting to process .doc file..."))
                    import subprocess
                    try:
                        result = subprocess.run(['antiword', '-h'], stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
                        result = subprocess.run(['antiword', file_path], stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
                        if result.returncode == 0:
                            content = result.stdout
                            self.window.after(0, lambda: self.update_output("DOC file processed with antiword"))
                        else:
                            self.window.after(0, lambda: self.update_output(f"Antiword error: {result.stderr}"))
                            return
                    except FileNotFoundError:
                        try:
                            import textract
                            content = textract.process(file_path).decode('utf-8')
                            self.window.after(0, lambda: self.update_output("DOC file processed with textract"))
                        except ImportError:
                            self.window.after(0, lambda: self.update_output("Error: Cannot process .doc files. Please install antiword or textract, or convert to .docx"))
                            return
                except Exception as doc_error:
                    self.window.after(0, lambda: self.update_output(f"Failed to process .doc file: {str(doc_error)}"))
                    return
            elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.gif', '.webp']:
                try:
                    content = self._extract_text_from_image(file_path)
                    extraction_source = "image"
                    self.window.after(0, lambda: self.update_output("Image processed with OCR"))
                except Exception as ocr_e:
                    self.window.after(0, lambda: self.update_output(f"Error processing image with OCR: {str(ocr_e)}"))
                    return
            else:
                self.window.after(0, lambda: self.update_output(f"Unsupported file format: {file_ext}. Trying general text extraction..."))
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                except Exception:
                    self.window.after(0, lambda: self.update_output("Could not read file as text. Please convert to a supported format."))
                    return
            if not content or not content.strip():
                self.window.after(0, lambda: self.update_output("Error: No content could be extracted from the file"))
                return
            status_msg = f"Extracting '{extract_what}' using AI from {extraction_source}..."
            self.window.after(0, lambda msg=status_msg: self.update_output(msg))
            prompt = f"""Task: Extract information based on the request: '{extract_what}'.
            Input Content Type: {extraction_source}
            Input Content:
            ---
            {content[:20000]}
            ---
            Output ONLY the extracted information that strictly matches the request '{extract_what}'.
            Be thorough and precise in extracting all relevant information.
            Do not include explanations, greetings, or apologies. If no matching information is found, output the exact phrase: NO_INFO_FOUND
            Extracted Data:"""
            try:
                response = self.gemini_model.generate_content(prompt)
                extracted = response.text.strip()
                if not extracted or extracted == "NO_INFO_FOUND":
                     self.window.after(0, lambda: self.update_output(f"AI could not find information matching '{extract_what}'."))
                     return
                file_name = os.path.basename(file_path)
                file_base, file_ext = os.path.splitext(file_name)
                clean_extract_what = re.sub(r'[^a-zA-Z0-9_]', '_', extract_what).strip('_')
                output_base_path = os.path.join(self.assistant_folder, f"{file_base}_extracted_{clean_extract_what[:20]}")
                if export_format:
                    output_path = self._export_result(extracted, output_base_path, export_format)
                    if output_path:
                        success_msg = f"Extracted Preview:\n{extracted[:500]}..."
                        self.window.after(0, lambda msg=success_msg: self.update_output(msg))
                else:
                    output_path_txt = output_base_path + ".txt"
                    with open(output_path_txt, 'w', encoding='utf-8') as f:
                        f.write(extracted)
                    success_msg1 = f"Extraction saved to {output_path_txt}"
                    success_msg2 = f"Extracted Preview:\n{extracted[:500]}..."
                    self.window.after(0, lambda msg=success_msg1: self.update_output(msg))
                    self.window.after(0, lambda msg=success_msg2: self.update_output(msg))
                    self.add_to_assistant_file(output_path_txt, "extract")
            except Exception as e:
                if "Content Policy Violation" in str(e):
                     error_msg = f"AI Error: Could not extract due to content policy."
                elif "quota" in str(e).lower():
                    error_msg = f"AI Error: API quota likely exceeded."
                else:
                    error_msg = f"Error using AI to extract: {str(e)}"
                self.window.after(0, lambda msg=error_msg: self.update_output(msg))
        except FileNotFoundError:
             self.window.after(0, lambda: self.update_output(f"Error: File not found at {file_path}"))
        except Exception as e:
            error_msg = f"Error processing file for extraction: {str(e)}"
            self.window.after(0, lambda msg=error_msg: self.update_output(msg))
    def _extract_text_from_image(self, image_path):
        """Extract text from an image using OCR"""
        try:
            import pytesseract
            from PIL import Image
            try:
                pytesseract.get_tesseract_version()
            except Exception as e:
                self.window.after(0, lambda: self.update_output(
                    "Tesseract OCR not found. Please install Tesseract OCR and ensure it's in your PATH."))
                raise Exception("Tesseract OCR not available")
            img = Image.open(image_path)
            max_dimension = 3000
            if max(img.size) > max_dimension:
                ratio = max_dimension / max(img.size)
                new_size = (int(img.size[0] * ratio), int(img.size[1] * ratio))
                img = img.resize(new_size, Image.LANCZOS)
            if img.mode != 'L':
                img = img.convert('L')
            from PIL import ImageEnhance
            enhancer = ImageEnhance.Contrast(img)
            img = enhancer.enhance(2.0)
            text = pytesseract.image_to_string(img)
            return text
        except ImportError:
            self.window.after(0, lambda: self.update_output(
                "Error: Missing libraries for OCR. Install pytesseract."))
            raise Exception("OCR libraries not available")
    def _extract_text_from_pdf_with_ocr(self, pdf_path):
        """Extract text from a scanned PDF using OCR"""
        try:
            import pytesseract
            from PIL import Image
            import fitz
            try:
                pytesseract.get_tesseract_version()
            except Exception:
                self.window.after(0, lambda: self.update_output(
                    "Tesseract OCR not found. Please install Tesseract OCR and ensure it's in your PATH."))
                raise Exception("Tesseract OCR not available")
            text_content = ""
            pdf_document = fitz.open(pdf_path)
            self.window.after(0, lambda: self.update_output(f"Processing PDF with {len(pdf_document)} pages using OCR..."))
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                img = img.convert('L')
                from PIL import ImageEnhance
                enhancer = ImageEnhance.Contrast(img)
                img = enhancer.enhance(2.0)
                page_text = pytesseract.image_to_string(img)
                text_content += page_text + "\n\n"
                if page_num % 5 == 0 or page_num == len(pdf_document) - 1:
                    self.window.after(0, lambda p=page_num, t=len(pdf_document): 
                        self.update_output(f"OCR progress: {p+1}/{t} pages"))
            return text_content
        except ImportError:
            self.window.after(0, lambda: self.update_output(
                "Error: Missing libraries for PDF OCR. Install pytesseract and PyMuPDF (fitz)."))
            raise Exception("PDF OCR libraries not available")
    def handle_custom_command(self, file_path, command):
        """Handle custom commands using Gemini AI, with export option"""
        export_format = None
        export_match = re.search(r'(?:export|save)\s+(?:it|result)\s+as\s+(\w+)', command.lower())
        if export_match:
            export_format = export_match.group(1).lower()
            self.window.after(0, lambda: self.update_output(f"Detected export request for custom command result: {export_format}"))
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            self.window.after(0, lambda: self.update_output(f"Processing custom command '{command}' using AI..."))
            prompt = f"""
            Act as a file processing assistant. Execute the following command on the provided file content:
            Command: "{command}"
            File Content:
            ---
            {content[:15000]}
            ---
            Provide only the result of the command execution. If the command is unclear or cannot be performed on the text, state that clearly.
            Result:"""
            response = self.gemini_model.generate_content(prompt)
            result = response.text.strip()
            if not result:
                 self.window.after(0, lambda: self.update_output("AI returned an empty result for the custom command."))
                 return
            file_name = os.path.basename(file_path)
            file_base, file_ext = os.path.splitext(file_name)
            clean_command = re.sub(r'[^a-zA-Z0-9_]', '_', command).strip('_')
            output_base_path = os.path.join(self.assistant_folder, f"{file_base}_custom_{clean_command[:20]}")
            if export_format:
                output_path = self._export_result(result, output_base_path, export_format)
                if output_path:
                    success_msg = f"Result Preview:\n{result[:500]}..."
                    self.window.after(0, lambda msg=success_msg: self.update_output(msg))
            else:
                output_path_txt = output_base_path + ".txt"
                with open(output_path_txt, 'w', encoding='utf-8') as f:
                    f.write(result)
                success_msg1 = f"Custom command processed and saved to {output_path_txt}"
                success_msg2 = f"Result Preview:\n{result[:500]}..."
                self.window.after(0, lambda msg=success_msg1: self.update_output(msg))
                self.window.after(0, lambda msg=success_msg2: self.update_output(msg))
                self.add_to_assistant_file(output_path_txt, "custom_command")
        except FileNotFoundError:
             self.window.after(0, lambda: self.update_output(f"Error: File not found at {file_path}"))
        except Exception as e:
            if "Content Policy Violation" in str(e):
                error_msg = f"AI Error: Could not process custom command due to content policy."
            elif "quota" in str(e).lower():
                 error_msg = f"AI Error: API quota likely exceeded."
            else:
                 error_msg = f"Error processing custom command: {str(e)}"
            self.window.after(0, lambda msg=error_msg: self.update_output(msg))
    def _export_result(self, content, base_path, format_type):
        """Export the extracted content in the specified format"""
        try:
            format_type = format_type.lower()
            self.window.after(0, lambda: self.update_output(f"Exporting as {format_type}..."))
            if format_type == "txt" or format_type == "text":
                output_path = base_path + ".txt"
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                self.window.after(0, lambda: self.update_output(f"Exported to text file: {output_path}"))
                self.add_to_assistant_file(output_path, "extract")
                return output_path
            elif format_type == "json":
                try:
                    if content.strip().startswith('{') and content.strip().endswith('}'):
                        import json
                        parsed = json.loads(content)
                        formatted_content = json.dumps(parsed, indent=2)
                    else:
                        formatted_content = json.dumps({"extracted_content": content}, indent=2)
                except:
                    import json
                    formatted_content = json.dumps({"extracted_content": content}, indent=2)
                output_path = base_path + ".json"
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(formatted_content)
                self.window.after(0, lambda: self.update_output(f"Exported to JSON file: {output_path}"))
                self.add_to_assistant_file(output_path, "extract")
                return output_path
            elif format_type == "csv":
                output_path = base_path + ".csv"
                lines = content.strip().split('\n')
                if len(lines) <= 1 or ',' not in content:
                    with open(output_path, 'w', encoding='utf-8', newline='') as f:
                        import csv
                        writer = csv.writer(f)
                        writer.writerow(["extracted_content"])
                        writer.writerow([content])
                else:
                    with open(output_path, 'w', encoding='utf-8', newline='') as f:
                        f.write(content)
                self.window.after(0, lambda: self.update_output(f"Exported to CSV file: {output_path}"))
                self.add_to_assistant_file(output_path, "extract")
                return output_path
            elif format_type == "docx" or format_type == "word":
                try:
                    from docx import Document
                    document = Document()
                    document.add_paragraph(content)
                    output_path = base_path + ".docx"
                    document.save(output_path)
                    self.window.after(0, lambda: self.update_output(f"Exported to Word document: {output_path}"))
                    self.add_to_assistant_file(output_path, "extract")
                    return output_path
                except ImportError:
                    self.window.after(0, lambda: self.update_output("Error: python-docx library not found. Saving as text instead."))
                    output_path = base_path + ".txt"
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(content)
                    self.window.after(0, lambda: self.update_output(f"Exported to text file (fallback): {output_path}"))
                    self.add_to_assistant_file(output_path, "extract")
                    return output_path
            elif format_type == "md" or format_type == "markdown":
                output_path = base_path + ".md"
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                self.window.after(0, lambda: self.update_output(f"Exported to Markdown file: {output_path}"))
                self.add_to_assistant_file(output_path, "extract")
                return output_path
            elif format_type == "html":
                output_path = base_path + ".html"
                html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Extracted Content</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
        .container {{ max-width: 800px; margin: 0 auto; }}
        pre {{ background-color:
    </style>
</head>
<body>
    <div class="container">
        <h1>Extracted Content</h1>
        <pre>{content.replace("<", "&lt;").replace(">", "&gt;")}</pre>
    </div>
</body>
</html>"""
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
                self.window.after(0, lambda: self.update_output(f"Exported to HTML file: {output_path}"))
                self.add_to_assistant_file(output_path, "extract")
                return output_path
            elif format_type == "pdf":
                try:
                    from fpdf import FPDF
                    pdf = FPDF()
                    pdf.add_page()
                    pdf.set_font("Arial", size=12)
                    chunks = [content[i:i+80] for i in range(0, len(content), 80)]
                    for chunk in chunks:
                        pdf.multi_cell(0, 10, chunk)
                    output_path = base_path + ".pdf"
                    pdf.output(output_path)
                    self.window.after(0, lambda: self.update_output(f"Exported to PDF file: {output_path}"))
                    self.add_to_assistant_file(output_path, "extract")
                    return output_path
                except ImportError:
                    self.window.after(0, lambda: self.update_output("Error: FPDF library not found. Saving as text instead."))
                    output_path = base_path + ".txt"
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(content)
                    self.window.after(0, lambda: self.update_output(f"Exported to text file (fallback): {output_path}"))
                    self.add_to_assistant_file(output_path, "extract")
                    return output_path
            else:
                self.window.after(0, lambda: self.update_output(f"Unsupported format: {format_type}. Using text format instead."))
                output_path = base_path + ".txt"
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                self.window.after(0, lambda: self.update_output(f"Exported to text file (fallback): {output_path}"))
                self.add_to_assistant_file(output_path, "extract")
                return output_path
        except Exception as e:
            error_msg = f"Error during export: {str(e)}"
            self.window.after(0, lambda msg=error_msg: self.update_output(msg))
            try:
                output_path = base_path + ".txt"
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                self.window.after(0, lambda: self.update_output(f"Exported to text file (fallback): {output_path}"))
                self.add_to_assistant_file(output_path, "extract")
                return output_path
            except:
                return None
    def add_to_recent_files(self, file_path):
        """Add a file to recent files list"""
        if file_path not in self.processed_files:
            self.processed_files.append(file_path)
            if len(self.processed_files) > 10:
                self.processed_files.pop(0)
        self.update_recent_files_list()
        self.save_recent_files()
        self.add_to_assistant_file(file_path, "opened")
    def update_recent_files_list(self):
        """Update the recent files listbox"""
        self.recent_files_listbox.delete(0, tk.END)
        for file_path in reversed(self.processed_files):
            self.recent_files_listbox.insert(0, os.path.basename(file_path))
    def select_recent_file(self, event):
        """Handle selection of a recent file"""
        selection = self.recent_files_listbox.curselection()
        if selection:
            index = selection[0]
            reversed_index = len(self.processed_files) - 1 - index
            if 0 <= reversed_index < len(self.processed_files):
                file_path = self.processed_files[reversed_index]
                self.file_path_var.set(file_path)
                self.update_output(f"Selected recent file: {os.path.basename(file_path)}")
    def save_recent_files(self):
        """Save recent files to a file"""
        recent_files_path = os.path.join(self.assistant_folder, "recent_files.txt")
        try:
            with open(recent_files_path, 'w', encoding='utf-8') as f:
                for file_path in self.processed_files:
                    f.write(f"{file_path}\n")
        except Exception as e:
            logging.error(f"Error saving recent files: {str(e)}")
    def load_recent_files(self):
        """Load recent files from the file"""
        recent_files_path = os.path.join(self.assistant_folder, "recent_files.txt")
        if os.path.exists(recent_files_path):
            try:
                with open(recent_files_path, 'r', encoding='utf-8') as f:
                    self.processed_files = [line.strip() for line in f.readlines() if line.strip()]
                self.update_recent_files_list()
            except Exception as e:
                logging.error(f"Error loading recent files: {str(e)}")
    def interpret_command(self, command, file_path):
        """Use AI to understand natural language commands"""
        self.update_output("Interpreting your command...")
        self.status_var.set("Interpreting...")
        def interpretation_thread():
            try:
                prompt = f"""
                I need to understand the following file processing command: "{command}"
                Please interpret this command and tell me which of these operations it's asking for:
                1. Convert file format - Examples: 
                   - "convert to pdf"
                   - "make this a text file"
                   - "save as HTML"
                   - "export to PDF format"
                   - "change to doc format"
                2. Summarize content - Examples:
                   - "summarize this"
                   - "give me a summary"
                   - "create a summary of this file"
                3. Move file - Examples:
                   - "move to desktop"
                   - "put this in documents folder"
                4. Copy file - Examples:
                   - "copy to downloads"
                   - "duplicate this file to documents"
                5. Extract information - Examples:
                   - "extract phone numbers"
                   - "find all emails"
                   - "pull out the dates"
                Return ONLY a JSON response with these fields:
                {{
                  "operation": "convert|summarize|move|copy|extract|unknown",
                  "target_format": "pdf|txt|etc" (for convert),
                  "target_location": "folder_name" (for move/copy),
                  "extract_what": "text to search for" (for extract)
                }}
                IMPORTANT NOTES:
                - Look for export/save/make as indicators of convert operations
                - The phrase "export to X format" or "save as X" means convert to X format
                - When "document" is mentioned as a format, assume PDF
                - Be flexible with format names (txt, text, pdf, document, etc.)
                """
                response = self.gemini_model.generate_content(prompt)
                interpretation = response.text
                try:
                    import json
                    import re
                    json_match = re.search(r'\{.*\}', interpretation, re.DOTALL)
                    if json_match:
                        json_str = json_match.group(0)
                        interpretation_data = json.loads(json_str)
                        operation = interpretation_data.get("operation", "unknown")
                        if operation == "convert":
                            target_format = interpretation_data.get("target_format", "")
                            if target_format:
                                new_command = f"convert to {target_format}"
                                self.window.after(0, lambda cmd=new_command: self.command_var.set(cmd))
                                self.window.after(0, lambda cmd=new_command: 
                                    self.update_output(f"Interpreted as: {cmd}"))
                                self.window.after(0, lambda: self.convert_file(file_path, new_command))
                            else:
                                self.window.after(0, lambda: 
                                    self.update_output("Couldn't determine the target format"))
                        elif operation == "summarize":
                            self.window.after(0, lambda: 
                                self.update_output("Interpreted as: summarize"))
                            self.window.after(0, lambda: self.summarize_file(file_path))
                        elif operation == "move":
                            target = interpretation_data.get("target_location", "")
                            if target:
                                new_command = f"move to {target}"
                                self.window.after(0, lambda cmd=new_command: self.command_var.set(cmd))
                                self.window.after(0, lambda cmd=new_command: 
                                    self.update_output(f"Interpreted as: {cmd}"))
                                self.window.after(0, lambda: self.move_file(file_path, new_command))
                            else:
                                self.window.after(0, lambda: 
                                    self.update_output("Couldn't determine where to move the file"))
                        elif operation == "copy":
                            target = interpretation_data.get("target_location", "")
                            if target:
                                new_command = f"copy to {target}"
                                self.window.after(0, lambda cmd=new_command: self.command_var.set(cmd))
                                self.window.after(0, lambda cmd=new_command: 
                                    self.update_output(f"Interpreted as: {cmd}"))
                                self.window.after(0, lambda: self.copy_file(file_path, new_command))
                            else:
                                self.window.after(0, lambda: 
                                    self.update_output("Couldn't determine where to copy the file"))
                        elif operation == "extract":
                            extract_what = interpretation_data.get("extract_what", "")
                            if extract_what:
                                new_command = f"extract {extract_what}"
                                self.window.after(0, lambda cmd=new_command: self.command_var.set(cmd))
                                self.window.after(0, lambda cmd=new_command: 
                                    self.update_output(f"Interpreted as: {cmd}"))
                                self.window.after(0, lambda: self.extract_from_file(file_path, new_command))
                            else:
                                self.window.after(0, lambda: 
                                    self.update_output("Couldn't determine what to extract"))
                        else:
                            self.window.after(0, lambda: 
                                self.update_output("I'm not sure how to process that command. Please try a different phrasing."))
                            self.window.after(0, lambda: self.handle_custom_command(file_path, command))
                    else:
                        self.window.after(0, lambda: 
                            self.update_output("Couldn't interpret the command. Trying general processing..."))
                        self.window.after(0, lambda: self.handle_custom_command(file_path, command))
                except Exception as e:
                    error_msg = f"Error interpreting command: {str(e)}. Trying general processing..."
                    self.window.after(0, lambda msg=error_msg: self.update_output(msg))
                    self.window.after(0, lambda: self.handle_custom_command(file_path, command))
                self.window.after(0, lambda: self.add_to_recent_files(file_path))
                self.window.after(0, lambda: self.status_var.set("Completed"))
            except Exception as e:
                error_msg = f"Error processing command: {str(e)}"
                self.window.after(0, lambda msg=error_msg: self.update_output(f"Error: {str(e)}"))
                self.window.after(0, lambda: self.status_var.set("Error"))
        interpretation_thread = threading.Thread(target=interpretation_thread)
        interpretation_thread.daemon = True
        interpretation_thread.start()
    def show_recent_files_window(self):
        """Show a window with recently processed files by type"""
        if hasattr(self, 'recent_window') and self.recent_window is not None and self.recent_window.winfo_exists():
            self.recent_window.lift()
            return
        self.recent_window = ctk.CTkToplevel(self.window)
        self.recent_window.title("Recent Files")
        self.recent_window.geometry("700x500")
        self.recent_window.transient(self.window)
        theme = self.assistant.themes[self.assistant.current_theme]
        self.recent_window.configure(fg_color=theme["bg"])
        tabview = ctk.CTkTabview(
            self.recent_window,
            fg_color=theme["highlight"],
            segmented_button_fg_color=theme["bg"],
            segmented_button_selected_color=theme["accent"],
            segmented_button_unselected_color=theme["bg"]
        )
        tabview.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        all_tab = tabview.add("All Files")
        summary_tab = tabview.add("Summaries")
        convert_tab = tabview.add("Converted")
        extract_tab = tabview.add("Extracted")
        try:
            db_path = os.path.join(self.assistant_folder, "files_database.db")
            if not os.path.exists(db_path):
                conn = sqlite3.connect(db_path)
                cursor = conn.cursor()
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS processed_files (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        file_path TEXT,
                        file_name TEXT,
                        operation_type TEXT,
                        timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
                    )
                ''')
                conn.commit()
                conn.close()
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            cursor.execute(
                "SELECT id, file_path, file_name, operation_type, timestamp FROM processed_files ORDER BY timestamp DESC"
            )
            all_files = cursor.fetchall()
            cursor.execute(
                "SELECT id, file_path, file_name, timestamp FROM processed_files WHERE operation_type='summary' ORDER BY timestamp DESC"
            )
            summaries = cursor.fetchall()
            cursor.execute(
                "SELECT id, file_path, file_name, timestamp FROM processed_files WHERE operation_type='converted' ORDER BY timestamp DESC"
            )
            converted = cursor.fetchall()
            cursor.execute(
                "SELECT id, file_path, file_name, timestamp FROM processed_files WHERE operation_type='extract' ORDER BY timestamp DESC"
            )
            extracted = cursor.fetchall()
            conn.close()
            self.create_files_list(all_tab, all_files, show_type=True)
            self.create_files_list(summary_tab, summaries)
            self.create_files_list(convert_tab, converted)
            self.create_files_list(extract_tab, extracted)
        except Exception as e:
            error_frame = ctk.CTkFrame(all_tab, fg_color=theme["bg"])
            error_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
            ctk.CTkLabel(
                error_frame,
                text=f"Error loading files: {str(e)}",
                font=("Segoe UI", 14),
                text_color=theme["error"]
            ).pack(pady=20)
    def create_files_list(self, parent, files_data, show_type=False):
        """Create a scrollable list of files with actions"""
        theme = self.assistant.themes[self.assistant.current_theme]
        scroll_frame = ctk.CTkScrollableFrame(
            parent,
            fg_color=theme["bg"],
            corner_radius=0
        )
        scroll_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        if not files_data:
            ctk.CTkLabel(
                scroll_frame,
                text="No files found",
                font=("Segoe UI", 14),
                text_color=theme["fg"]
            ).pack(pady=20)
            return
        for file_data in files_data:
            if show_type:
                file_id, file_path, file_name, operation_type, timestamp = file_data
            else:
                file_id, file_path, file_name, timestamp = file_data
                operation_type = ""
            card = ctk.CTkFrame(
                scroll_frame,
                fg_color=theme["highlight"],
                corner_radius=8
            )
            card.pack(fill=tk.X, padx=5, pady=5)
            info_frame = ctk.CTkFrame(
                card,
                fg_color=theme["highlight"],
                corner_radius=0
            )
            info_frame.pack(fill=tk.X, padx=10, pady=5)
            name_label = ctk.CTkLabel(
                info_frame,
                text=file_name,
                font=("Segoe UI", 12, "bold"),
                text_color=theme["fg"]
            )
            name_label.pack(anchor=tk.W)
            if show_type:
                type_date = ctk.CTkLabel(
                    info_frame,
                    text=f"Type: {operation_type.capitalize()} | Date: {timestamp}",
                    font=("Segoe UI", 10),
                    text_color=theme["secondary"]
                )
                type_date.pack(anchor=tk.W)
            else:
                date_label = ctk.CTkLabel(
                    info_frame,
                    text=f"Date: {timestamp}",
                    font=("Segoe UI", 10),
                    text_color=theme["secondary"]
                )
                date_label.pack(anchor=tk.W)
            action_frame = ctk.CTkFrame(
                card,
                fg_color=theme["highlight"],
                corner_radius=0
            )
            action_frame.pack(fill=tk.X, padx=10, pady=5)
            open_button = ctk.CTkButton(
                action_frame,
                text="Open",
                font=("Segoe UI", 11),
                width=80,
                height=28,
                fg_color=theme["accent"],
                hover_color=theme["secondary"],
                command=lambda path=file_path: self.open_file(path)
            )
            open_button.pack(side=tk.LEFT, padx=5, pady=5)
            select_button = ctk.CTkButton(
                action_frame,
                text="Select",
                font=("Segoe UI", 11),
                width=80,
                height=28,
                fg_color=theme["success"],
                hover_color=theme["secondary"],
                command=lambda path=file_path: self.select_recent_file_from_path(path)
            )
            select_button.pack(side=tk.LEFT, padx=5, pady=5)
            delete_button = ctk.CTkButton(
                action_frame,
                text="Delete",
                font=("Segoe UI", 11),
                width=80,
                height=28,
                fg_color=theme["error"],
                hover_color=theme["secondary"],
                command=lambda id=file_id, card=card: self.delete_file_from_db(id, card)
            )
            delete_button.pack(side=tk.RIGHT, padx=5, pady=5)
    def open_file(self, file_path):
        """Open a file with the default application"""
        try:
            if os.path.exists(file_path):
                if platform.system() == 'Windows':
                    os.startfile(file_path)
                elif platform.system() == 'Darwin':
                    subprocess.run(['open', file_path])
                else:
                    subprocess.run(['xdg-open', file_path])
            else:
                messagebox.showerror("Error", f"File not found: {file_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Error opening file: {str(e)}")
    def select_recent_file_from_path(self, file_path):
        """Select a file from recent files window and load it into the main interface"""
        try:
            if os.path.exists(file_path):
                self.file_path_var.set(file_path)
                self.update_output(f"Selected file: {os.path.basename(file_path)}")
                if hasattr(self, 'recent_window') and self.recent_window is not None:
                    self.recent_window.destroy()
            else:
                messagebox.showerror("Error", f"File not found: {file_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Error selecting file: {str(e)}")
    def delete_file_from_db(self, file_id, card_widget):
        """Delete a file from the database"""
        try:
            if not messagebox.askyesno("Confirm", "Are you sure you want to delete this file from the database?"):
                return
            db_path = os.path.join(self.assistant_folder, "files_database.db")
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            cursor.execute("SELECT file_path FROM processed_files WHERE id=?", (file_id,))
            result = cursor.fetchone()
            if result:
                file_path = result[0]
                cursor.execute("DELETE FROM processed_files WHERE id=?", (file_id,))
                conn.commit()
                if messagebox.askyesno("Delete File", "Do you also want to delete the actual file?"):
                    if os.path.exists(file_path):
                        os.remove(file_path)
                card_widget.destroy()
            conn.close()
        except Exception as e:
            messagebox.showerror("Error", f"Error deleting file: {str(e)}")
    def process_voice_command(self, command, file_path=None):
        """Process a voice command from the continuous mode"""
        command_lower = command.lower()
        recent_file_patterns = [
            'show the recently', 'show recent', 'open recent', 'show the recent', 
            'show me the recent', 'show the latest', 'open the latest'
        ]
        for pattern in recent_file_patterns:
            if pattern in command_lower:
                recent_info = self._parse_recent_file_command(command)
                if recent_info["is_open_request"]:
                    self.window.after(0, lambda: self.update_output(f"Opening recent file (Type: {recent_info['filter_op_type']})"))
                    self.window.after(0, lambda: self.open_recent_file_by_filter(
                        filter_op_type=recent_info["filter_op_type"],
                        filter_date=recent_info["filter_date"],
                        filter_weekday=recent_info["filter_weekday"]
                    ))
                    return True
                break
        if "show recent" in command_lower or "show files" in command_lower or "recent files" in command_lower or "recently converted" in command_lower or "show converted" in command_lower:
            try:
                operation_type = None
                if "converted" in command_lower:
                    operation_type = "converted"
                elif "summarized" in command_lower or "summarised" in command_lower:
                    operation_type = "summary"
                elif "extracted" in command_lower:
                    operation_type = "extract"
                elif "opened" in command_lower:
                    operation_type = "opened"
                if (any(open_term in command_lower for open_term in 
                        ['show the', 'open the', 'display the', 'show me the']) 
                        and operation_type):
                    self.window.after(0, lambda: self.update_output(f"Opening most recent {operation_type} file"))
                    self.window.after(0, lambda op_type=operation_type: self.open_recent_file_by_filter(filter_op_type=op_type))
                    return True
                self.show_recent_files_window()
                if operation_type and hasattr(self, 'recent_window') and self.recent_window.winfo_exists():
                    tab_map = {
                        "converted": "Converted",
                        "summary": "Summaries",
                        "extract": "Extracted",
                        "opened": "All Files"
                    }
                    if operation_type in tab_map:
                        for widget in self.recent_window.winfo_children():
                            if isinstance(widget, ctk.CTkTabview):
                                try:
                                    widget.set(tab_map[operation_type])
                                    self.update_output(f"Showing {tab_map[operation_type]} tab")
                                    break
                                except Exception as tab_e:
                                    self.update_output(f"Could not select tab: {str(tab_e)}")
            except Exception as e:
                self.update_output(f"Error showing recent files: {str(e)}")
            return True
        self.command_var.set(command)
        self.process_file_command()
        return False
    def reset_voice_ui(self):
        """Reset UI elements after voice command mode is stopped"""
        if hasattr(self, 'voice_loop_button_ref') and self.voice_loop_button_ref:
            try:
                if self.voice_loop_button_ref.winfo_exists():
                    theme = self.assistant.themes[self.assistant.current_theme]
                    self.voice_loop_button_ref.configure(
                        text="Voice Mode", 
                        fg_color=theme["success"]
                    )
            except Exception as e:
                logging.error(f"Error resetting voice UI: {e}")
        for widget in self.window.winfo_children():
            if isinstance(widget, ctk.CTkFrame):
                for child in widget.winfo_children():
                    if isinstance(child, ctk.CTkFrame):
                        for button in child.winfo_children():
                            if isinstance(button, ctk.CTkButton) and button.cget("text") == "🎤":
                                button.configure(state=NORMAL)
        self.restore_main_listening()
    def browse_file(self):
        """Browse for a file without closing the window"""
        if self.window.winfo_exists():
            self.window.grab_release()
        self.window.after(100, self._open_file_dialog)
    def _open_file_dialog(self):
        """Open file dialog in a way that doesn't close the parent window"""
        try:
            if self.window is None or not self.window.winfo_exists():
                return
            self.window.attributes('-topmost', True)
            file_path = filedialog.askopenfilename(
                parent=self.window,
                title="Select File",
                filetypes=(("All Files", "*.*"), ("Text Files", "*.txt"), ("PDF Files", "*.pdf"))
            )
            self.window.attributes('-topmost', False)
            if file_path:
                self.file_path_var.set(file_path)
                self.update_output(f"Selected file: {os.path.basename(file_path)}")
                self.add_to_recent_files(file_path)
            if self.window.winfo_exists():
                self.window.grab_set()
                self.window.focus_force()
        except Exception as e:
            error_msg = f"Error browsing file: {str(e)}"
            if hasattr(self, 'update_output'):
                self.update_output(error_msg)
            logging.error(error_msg)
    def _parse_recent_file_command(self, command):
        """Parses command to find operation type, timeframe, and intent (show vs open) for recent files."""
        command_lower = command.lower()
        operation_type = None
        filter_date = None
        filter_weekday = None
        initial_tab = "All Files"
        is_open_request = False
        type_keywords = {
            'summary': 'Summaries', 'summarise':'Summaries', 'summaries': 'Summaries', 'summarized': 'Summaries',
            'converted': 'Converted', 'conversion': 'Converted', 'conversions': 'Converted',
            'extracted': 'Extracted', 'extraction': 'Extracted', 'extractions': 'Extracted', 'extract': 'Extracted',
            'exported': 'Exported', 'exports': 'Exported', 'export': 'Exported'
        }
        timeframe_keywords = {
            'today': 'today', 'todays': 'today', 'this day': 'today',
            'yesterday': 'yesterday', 'yesterdays': 'yesterday',
            'monday': 0, 'mondays': 0, 'on monday': 0,
            'tuesday': 1, 'tuesdays': 1, 'on tuesday': 1,
            'wednesday': 2, 'wednesdays': 2, 'on wednesday': 2,
            'thursday': 3, 'thursdays': 3, 'on thursday': 3,
            'friday': 4, 'fridays': 4, 'on friday': 4,
            'saturday': 5, 'saturdays': 5, 'on saturday': 5,
            'sunday': 6, 'sundays': 6, 'on sunday': 6,
            'this week': 'this_week',
            'last week': 'last_week',
            'this month': 'this_month',
            'last month': 'last_month'
        }
        weekday_map_to_sqlite = {
             0: '1', 1: '2', 2: '3', 3: '4', 4: '5', 5: '6', 6: '0'
        }
        open_patterns = [
            'show the recently', 'show the recent', 'show recent', 'open recent', 'open the recent',
            'show the last', 'open the last', 'open last', 'display recent', 'show me the recent',
            'show me the last', 'show the latest', 'open the latest'
        ]
        for pattern in open_patterns:
            if pattern in command_lower:
                is_open_request = True
                break
        words = command_lower.split()
        if not is_open_request and words[0] in ['open', 'load', 'display', 'view'] and ('recent' in words or 'file' in words or any(t in words for t in type_keywords) or any(t in words for t in timeframe_keywords)):
             is_open_request = True
        for i, word in enumerate(words):
            if word in type_keywords:
                initial_tab = type_keywords[word]
                db_op_type_map = {
                    'Summaries': 'summary',
                    'Converted': 'converted',
                    'Extracted': 'extract',
                    'Exported': 'exported'
                }
                if initial_tab == 'Exported':
                     operation_type = 'exported%'
                else:
                     operation_type = next((db_type for tab, db_type in db_op_type_map.items() if tab == initial_tab), None)
                break
            if i < len(words) - 1 and f"{word} {words[i+1]}" in type_keywords:
                compound = f"{word} {words[i+1]}"
                initial_tab = type_keywords[compound]
                db_op_type_map = {
                    'Summaries': 'summary',
                    'Converted': 'converted',
                    'Extracted': 'extract',
                    'Exported': 'exported'
                }
                if initial_tab == 'Exported':
                     operation_type = 'exported%'
                else:
                     operation_type = next((db_type for tab, db_type in db_op_type_map.items() if tab == initial_tab), None)
                break
        for i, word in enumerate(words):
            if word in timeframe_keywords:
                timeframe = timeframe_keywords[word]
                self._process_timeframe(timeframe, filter_date, filter_weekday, weekday_map_to_sqlite)
                break
            if i < len(words) - 1 and f"{word} {words[i+1]}" in timeframe_keywords:
                compound = f"{word} {words[i+1]}"
                timeframe = timeframe_keywords[compound]
                filter_date, filter_weekday = self._process_timeframe(timeframe, weekday_map_to_sqlite)
                break
            if word in ['on', 'from', 'since', 'after'] and i < len(words) - 1:
                date_str = words[i+1]
                if re.match(r'\d{4}-\d{2}-\d{2}', date_str):
                    filter_date = date_str
                    break
                if i < len(words) - 2:
                    potential_date = f"{words[i+1]} {words[i+2]}"
                    try:
                        parsed_date = datetime.datetime.strptime(potential_date, "%B %d").replace(
                            year=datetime.datetime.now().year).date()
                        filter_date = parsed_date.isoformat()
                        break
                    except ValueError:
                        try:
                            parsed_date = datetime.datetime.strptime(potential_date, "%d %B").replace(
                                year=datetime.datetime.now().year).date()
                            filter_date = parsed_date.isoformat()
                            break
                        except ValueError:
                            pass
        if not is_open_request and ('recently' in command_lower or 'latest' in command_lower) and operation_type:
            is_open_request = True
        is_recent_command = filter_date is not None or filter_weekday is not None or any(w in command_lower for w in ['recent', 'show file', 'list file', 'open file', 'open summary', 'open converted', 'open extracted', 'latest', 'recently'])
        filter_op_type_for_query = None
        if operation_type:
             if operation_type.endswith('%'):
                  filter_op_type_for_query = operation_type
             elif any(k in words for k in type_keywords):
                 filter_op_type_for_query = operation_type
        if is_open_request and not filter_op_type_for_query and any(term in command_lower for term in ['recent file', 'recently', 'latest file']):
            filter_op_type_for_query = 'opened'
        return {
            "is_recent_command": is_recent_command,
            "is_open_request": is_open_request,
            "initial_tab": initial_tab,
            "filter_op_type": filter_op_type_for_query,
            "filter_date": filter_date,
            "filter_weekday": filter_weekday
        }
    def _process_timeframe(self, timeframe, weekday_map_to_sqlite):
        """Process timeframe value and return appropriate filter_date and filter_weekday values"""
        filter_date = None
        filter_weekday = None
        today = datetime.date.today()
        if timeframe == 'today':
            filter_date = today.isoformat()
        elif timeframe == 'yesterday':
            filter_date = (today - datetime.timedelta(days=1)).isoformat()
        elif timeframe == 'this_week':
            start_of_week = today - datetime.timedelta(days=today.weekday())
            filter_date = f"date(timestamp) >= '{start_of_week.isoformat()}'"
        elif timeframe == 'last_week':
            end_of_last_week = today - datetime.timedelta(days=today.weekday() + 1)
            start_of_last_week = end_of_last_week - datetime.timedelta(days=6)
            filter_date = f"date(timestamp) BETWEEN '{start_of_last_week.isoformat()}' AND '{end_of_last_week.isoformat()}'"
        elif timeframe == 'this_month':
            first_of_month = today.replace(day=1)
            filter_date = f"date(timestamp) >= '{first_of_month.isoformat()}'"
        elif timeframe == 'last_month':
            first_of_this_month = today.replace(day=1)
            last_of_last_month = first_of_this_month - datetime.timedelta(days=1)
            first_of_last_month = last_of_last_month.replace(day=1)
            filter_date = f"date(timestamp) BETWEEN '{first_of_last_month.isoformat()}' AND '{last_of_last_month.isoformat()}'"
        elif isinstance(timeframe, int):
            sqlite_weekday = weekday_map_to_sqlite.get(timeframe)
            if sqlite_weekday:
                filter_weekday = sqlite_weekday
        return filter_date, filter_weekday
    def process_file_command(self):
        file_path = self.file_path_var.get().strip()
        command = self.command_var.get().strip().lower()
        if not file_path:
            self.update_output("Error: No file selected")
            return
        if not os.path.exists(file_path):
            self.update_output("Error: File not found")
            return
        if not command:
            self.update_output("Error: No command specified")
            return
        self.update_output(f"Processing command: {command} on file: {os.path.basename(file_path)}")
        self.status_var.set("Processing...")
        if any(export_word in command for export_word in ["export", "save as", "output as"]):
            export_format = None
            cmd_parts = command.split()
            for i, word in enumerate(cmd_parts):
                if word in ["to", "as", "in"] and i < len(cmd_parts) - 1:
                    export_format = cmd_parts[i+1]
                    break
            if export_format:
                new_command = f"convert to {export_format}"
                self.update_output(f"Interpreting export command as: {new_command}")
                command = new_command
        if not any(keyword in command for keyword in ["convert", "summarize", "summarise", "move", "copy", "extract"]):
            self.interpret_command(command, file_path)
            return
        def process_thread():
            try:
                if "convert" in command:
                    self.convert_file(file_path, command)
                elif "summarize" in command or "summarise" in command:
                    self.summarize_file(file_path)
                elif "move" in command:
                    self.move_file(file_path, command)
                elif "copy" in command:
                    self.copy_file(file_path, command)
                elif "extract" in command:
                    self.extract_from_file(file_path, command)
                else:
                    self.handle_custom_command(file_path, command)
                self.window.after(0, lambda: self.add_to_recent_files(file_path))
                self.window.after(0, lambda: self.status_var.set("Completed"))
            except Exception as e:
                error_msg = f"Error processing command: {str(e)}"
                self.window.after(0, lambda msg=error_msg: self.update_output(msg))
                self.window.after(0, lambda: self.status_var.set("Error"))
        processing_thread = threading.Thread(target=process_thread)
        processing_thread.daemon = True
        processing_thread.start()
    def open_recent_file_by_filter(self, filter_op_type=None, filter_date=None, filter_weekday=None):
         """Finds the single most recent file matching filters and opens it."""
         try:
             db_path = os.path.join(self.assistant_folder, "files_database.db")
             if not os.path.exists(db_path):
                  self.update_output("Recent files database not found.")
                  return
             conn = sqlite3.connect(db_path)
             cursor = conn.cursor()
             query = "SELECT file_path, file_name, operation_type, timestamp FROM processed_files"
             conditions = []
             params = []
             type_desc = ""
             if filter_op_type:
                 if filter_op_type.endswith('%'):
                     type_desc = "exported"
                     conditions.append("operation_type LIKE ?")
                     params.append(filter_op_type)
                 else:
                     type_desc = filter_op_type
                     conditions.append("operation_type = ?")
                     params.append(filter_op_type)
             if filter_date:
                 if filter_date.startswith("date(timestamp)"):
                     conditions.append(filter_date)
                 elif filter_date.startswith("date(timestamp) BETWEEN"):
                     conditions.append(filter_date)
                 else:
                     conditions.append("date(timestamp) = date(?)")
                     params.append(filter_date)
             if filter_weekday is not None:
                 conditions.append("strftime('%w', timestamp) = ?")
                 params.append(filter_weekday)
                 weekday_names = ['Sunday', 'Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']
                 weekday_map = {'0': 6, '1': 0, '2': 1, '3': 2, '4': 3, '5': 4, '6': 5}
                 if filter_weekday in weekday_map:
                     weekday_name = weekday_names[weekday_map[filter_weekday]]
                     if type_desc:
                         type_desc = f"{type_desc} on {weekday_name}"
                     else:
                         type_desc = f"processed on {weekday_name}"
             search_desc = type_desc or "processed recently"
             self.update_output(f"Searching for the most recent file {search_desc}...")
             if conditions:
                 query += " WHERE " + " AND ".join(conditions)
             query += " ORDER BY timestamp DESC LIMIT 1"
             cursor.execute(query, params)
             result = cursor.fetchone()
             if result:
                 file_path, file_name, operation_type, timestamp = result
                 if os.path.exists(file_path):
                     operation_desc = {
                         'summary': 'summarized',
                         'converted': 'converted',
                         'extract': 'extracted from',
                         'opened': 'opened'
                     }.get(operation_type, operation_type)
                     try:
                         ts_dt = datetime.datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                         formatted_time = ts_dt.strftime("%Y-%m-%d %H:%M")
                     except:
                         formatted_time = timestamp
                     self.update_output(f"Opening most recent {operation_desc} file: {file_name}")
                     self.update_output(f"This file was {operation_desc} on {formatted_time}")
                     self.open_file(file_path)
                     self.select_recent_file_from_path(file_path)
                 else:
                     self.update_output(f"File found in database but no longer exists: {file_path}")
                     cursor.execute("DELETE FROM processed_files WHERE file_path = ?", (file_path,))
                     conn.commit()
                     self.update_output("The file entry has been removed from the database.")
             else:
                 if type_desc:
                     self.update_output(f"No {type_desc} files found in the database.")
                 else:
                     self.update_output("No matching files found with those filters.")
             conn.close()
         except Exception as e:
             self.update_output(f"Error opening recent file: {str(e)}")
             logging.error(f"Error in open_recent_file_by_filter: {str(e)}")
    def _extract_text_from_docx(self, file_path):
        """Extract text from a DOCX file with error handling for corrupted files"""
        try:
            try:
                import docx2txt
                text = docx2txt.process(file_path)
                if text and len(text.strip()) > 0:
                    return text
            except ImportError:
                pass
            except Exception as e:
                self.window.after(0, lambda: self.update_output(f"Warning: docx2txt failed: {str(e)}. Trying python-docx..."))
            try:
                from docx import Document
                doc = Document(file_path)
                text = []
                for para in doc.paragraphs:
                    text.append(para.text)
                return '\n'.join(text)
            except ImportError:
                raise Exception("No DOCX processing libraries available. Please install docx2txt or python-docx.")
            except Exception as e:
                raise Exception(f"Failed to process DOCX file: {str(e)}. The file may be corrupted.")
        except Exception as e:
            self.window.after(0, lambda: self.update_output(f"Error extracting text from DOCX: {str(e)}"))
            return f"[Error extracting text: {str(e)}]"
    def extract_from_file(self, file_path, command):
        extract_what = None
        export_format = None
        words = command.lower().split()
        try:
            if "extract" in command.lower():
                extract_index = command.lower().find("extract")
                remaining_text = command.lower()[extract_index + 7:].strip()
                for phrase in ["export as", "save as", "export it as", "save it as"]:
                    if phrase in remaining_text:
                        remaining_text = remaining_text.split(phrase)[0].strip()
                extract_what = remaining_text
            export_match = re.search(r'(?:export|save)\s+(?:it|result|extraction)\s+as\s+(\w+)', command.lower())
            if export_match:
                export_format = export_match.group(1).lower()
        except Exception as e:
            self.window.after(0, lambda: self.update_output(f"Error parsing command: {str(e)}"))
            extract_what = self._use_gemini_to_interpret_extraction(command)
        if not extract_what:
            self.window.after(0, lambda: self.update_output("Using Gemini to interpret your extraction request..."))
            extract_what = self._use_gemini_to_interpret_extraction(command)
        file_ext = os.path.splitext(file_path)[1].lower()
        content = ""
        extraction_source = "text"
        try:
            self.window.after(0, lambda: self.update_output(f"Reading file: {os.path.basename(file_path)}"))
            if file_ext == '.pdf':
                try:
                    content = self._extract_from_pdf_enhanced(file_path)
                    if not content.strip():
                        self.window.after(0, lambda: self.update_output("PDF seems to be scanned, using OCR..."))
                        content = self._extract_text_from_pdf_with_ocr(file_path)
                except Exception as e:
                    self.window.after(0, lambda: self.update_output(f"PDF extraction error: {str(e)}. Trying OCR..."))
                    content = self._extract_text_from_pdf_with_ocr(file_path)
            elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.gif', '.webp']:
                content = self._enhanced_image_ocr(file_path)
                extraction_source = "image"
            elif file_ext == '.docx':
                content = self._extract_text_from_docx_enhanced(file_path)
            elif file_ext in ['.txt', '.md', '.csv', '.json', '.html', '.xml', '.log']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
            else:
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                except:
                    self.window.after(0, lambda: self.update_output("Trying OCR as fallback..."))
                    content = self._enhanced_image_ocr(file_path)
                    extraction_source = "image"
            prompt = f"""Task: Extract ONLY the following information: '{extract_what}'
            From this {extraction_source} content:
            ---
            {content[:20000]}
            ---
            FORMAT YOUR RESPONSE LIKE THIS:
            [EXTRACTED_CONTENT_START]
            (only the extracted information goes here, nothing else)
            [EXTRACTED_CONTENT_END]
            IMPORTANT INSTRUCTIONS:
            1. Extract EXACTLY what was requested: '{extract_what}'
            2. Include ALL instances of the requested information
            3. Format appropriately (tables as tables, lists as lists, etc.)
            4. If no matching information exists, write ONLY: "No matching information found"
            5. No explanations, introductions, or analysis - ONLY the extracted content
            """
            response = self.gemini_model.generate_content(prompt)
            extracted_text = response.text.strip()
            match = re.search(r'\[EXTRACTED_CONTENT_START\](.*?)\[EXTRACTED_CONTENT_END\]', 
                            extracted_text, re.DOTALL)
            if match:
                extracted = match.group(1).strip()
            else:
                extracted = extracted_text
            if not extracted or "No matching information found" in extracted:
                self.window.after(0, lambda: self.update_output(
                    f"Could not find information matching '{extract_what}'. Try a different query."))
                return
            file_name = os.path.basename(file_path)
            file_base, _ = os.path.splitext(file_name)
            clean_extract_what = re.sub(r'[^a-zA-Z0-9_]', '_', extract_what).strip('_')
            output_base_path = os.path.join(self.assistant_folder, f"{file_base}_extracted_{clean_extract_what[:20]}")
            if export_format:
                output_path = self._export_result(extracted, output_base_path, export_format)
            else:
                output_path_txt = output_base_path + ".txt"
                with open(output_path_txt, 'w', encoding='utf-8') as f:
                    f.write(extracted)
                self.window.after(0, lambda: self.update_output(f"Extraction saved to {output_path_txt}"))
                self.window.after(0, lambda: self.update_output(f"Preview:\n{extracted[:500]}..."))
                self.add_to_assistant_file(output_path_txt, "extract")
        except Exception as e:
            self.window.after(0, lambda: self.update_output(f"Extraction error: {str(e)}"))
    def _extract_from_pdf_enhanced(self, file_path):
        """Enhanced PDF text extraction using multiple libraries"""
        content = ""
        try:
            import PyPDF2
            with open(file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    content += page.extract_text() + "\n"
        except:
            pass
        if not content.strip():
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        content += page.extract_text() + "\n"
            except:
                pass
        return content
    def _enhanced_image_ocr(self, image_path):
        """Better OCR for images with preprocessing"""
        try:
            import pytesseract
            from PIL import Image, ImageEnhance, ImageFilter
            img = Image.open(image_path)
            max_dimension = 3000
            if max(img.size) > max_dimension:
                ratio = max_dimension / max(img.size)
                new_size = (int(img.size[0] * ratio), int(img.size[1] * ratio))
                img = img.resize(new_size, Image.LANCZOS)
            if img.mode != 'L':
                img = img.convert('L')
            img = img.filter(ImageFilter.SHARPEN)
            enhancer = ImageEnhance.Contrast(img)
            img = enhancer.enhance(2.0)
            text = pytesseract.image_to_string(img, lang='eng')
            return text
        except Exception as e:
            self.window.after(0, lambda: self.update_output(f"OCR error: {str(e)}"))
            return ""
    def _extract_text_from_docx_enhanced(self, file_path):
        """More robust DOCX text extraction"""
        content = ""
        try:
            import docx2txt
            content = docx2txt.process(file_path)
        except:
            pass
        if not content.strip():
            try:
                from docx import Document
                doc = Document(file_path)
                for para in doc.paragraphs:
                    content += para.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            row_text.append(cell.text)
                        content += " | ".join(row_text) + "\n"
                    content += "\n"
            except:
                pass
        return content
    def _use_gemini_to_interpret_extraction(self, command):
        """Use Gemini to interpret what needs to be extracted"""
        try:
            prompt = f"""
            Analyze this file processing command: "{command}"
            What specific information is the user trying to extract from the file?
            Return only the information to be extracted, nothing else.
            For example, if the user says "extract phone numbers from this file",
            return only "phone numbers".
            """
            response = self.gemini_model.generate_content(prompt)
            extract_what = response.text.strip()
            self.window.after(0, lambda: self.update_output(f"Interpreted extraction target: {extract_what}"))
            return extract_what
        except Exception as e:
            self.window.after(0, lambda: self.update_output(f"Error interpreting command: {str(e)}"))
            return None
if __name__ == "__main__":
    try:
        if not create_required_directories():
            print("Warning: Some directories could not be created")
        if not setup_logging():
            print("Warning: Using fallback logging configuration")
        root = ctk.CTk()
        app = VoiceAssistant(root)
        root.mainloop()
    except Exception as e:
        print(f"Fatal error during startup: {str(e)}")
        sys.exit(1)
def format_time(seconds):
    """Convert seconds to formatted time string"""
    minutes, seconds = divmod(seconds, 60)
    hours, minutes = divmod(minutes, 60)

    return f"{int(hours):02d}:{int(minutes):02d}"
